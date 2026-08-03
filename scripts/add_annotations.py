# -*- coding: utf-8 -*-
"""
Add inline annotation notes to the polished DOCX at specific paragraph locations.
Annotations appear as small-font (8pt), colored (blue) text appended to the paragraph.
"""
import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

SRC = r"c:\Users\郝\Desktop\claude\向晶\页岩气藏地应力演化与裂缝扩展耦合机制研究进展与展望-润色版.docx"
DST = r"c:\Users\郝\Desktop\claude\向晶\页岩气藏地应力演化与裂缝扩展耦合机制研究进展与展望-润色版-带批注.docx"
ANN_PATH = r"c:\Users\郝\Desktop\claude\向晶\annotations.json"

with open(ANN_PATH, 'r', encoding='utf-8') as f:
    annotations = json.load(f)

doc = Document(SRC)
count = 0

for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue

    # Find matching annotations
    for ann in annotations:
        if ann['keyword'] in text:
            # Add annotation as a new run with small blue font
            # Use a line break first
            spacer = para.add_run('\n')
            spacer.font.size = Pt(4)
            
            note = para.add_run(ann['text'])
            note.font.size = Pt(8)
            note.font.color.rgb = RGBColor(0, 70, 180)  # Dark blue
            note.font.name = '宋体'
            # Also make it slightly lighter weight
            
            count += 1
            # Remove this annotation so we don't add it to multiple paragraphs
            annotations.remove(ann)
            break

doc.save(DST)
print(f"Added {count} annotations to {DST}")
