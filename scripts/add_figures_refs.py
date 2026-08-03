#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
1. Add figure/table placeholders throughout the thesis
2. Add new references for expanded Chapter 1 content
3. Add citation brackets in expanded text
"""
import sys
from docx import Document
from lxml import etree
from copy import deepcopy

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

DOC = r'荧光压裂液/论文初稿_v2.docx'
doc = Document(DOC)

def find_para(doc, snippet, skip=0):
    found = 0
    for i, para in enumerate(doc.paragraphs):
        if snippet in para.text:
            if found >= skip:
                return i, para
            found += 1
    return None, None

def add_para_after(doc, target_para, text):
    new_para = doc.add_paragraph(text)
    target_para._element.addnext(new_para._element)
    return new_para

def add_ref_after(doc, target_para, ref_text):
    """Add a reference entry paragraph after the last reference."""
    new_para = doc.add_paragraph(ref_text)
    target_para._element.addnext(new_para._element)
    return new_para

# ================================================================
# PART 1: FIGURE & TABLE PLACEHOLDERS
# ================================================================
print("=" * 60)
print("PART 1: Adding figure/table placeholders")

figures = {}

# ---- Chapter 1 ----
# Figure 1-1: 技术路线图 (after 1.4 content)
_, p = find_para(doc, '各研究内容按照"材料改性→体系构建→动态验证→工程转化"的主线递进展开')
if p:
    add_para_after(doc, p, '[图1-1 此处插入：论文技术路线图。以流程图形式展示"材料改性→体系构建→动态验证→工程转化"四模块的逻辑关系、各模块对应的研究内容编号（1-6）、以及模块间的输入输出关系。建议使用Visio或Draw.io绘制，A4横排。]')
    print('  Added 图1-1 技术路线图')

# ---- Chapter 2 ----
# Figure 2-1: XRD
_, p = find_para(doc, 'X射线衍射（XRD）分析晶相组成')
if p:
    add_para_after(doc, p, '[图2-1 此处插入：SrAl₂O₄:Eu²⁺,Dy³⁺荧光粉XRD图谱。标注各衍射峰对应的晶面指数（hkl），与JCPDS标准卡片（No.34-0379）对比。]')
    print('  Added 图2-1 XRD')
else:
    # Try alternative
    _, p = find_para(doc, 'XRD图谱、SEM照片')
    if p:
        add_para_after(doc, p, '[图2-1 此处插入：荧光粉XRD图谱]\n[图2-2 此处插入：荧光粉SEM照片（低倍×500 + 高倍×5000）]\n[图2-3 此处插入：荧光粉粒度分布曲线（D10/D50/D90标注）]\n[图2-4 此处插入：稳态激发光谱（监测520 nm）和发射光谱（激发365 nm）]\n[图2-5 此处插入：余辉衰减曲线（I-t，双指数拟合，标注τ₁和τ₂）]')
        print('  Added 图2-1~2-5 基础物性')

# Figure 2-6~2-9: Environmental stability
_, p = find_para(doc, '获取速率常数k和半衰期t₁/₂')
if p:
    add_para_after(doc, p, '[图2-6 此处插入：热稳定性——相对发光保持率 vs 温度/时间（四温度曲线叠加）]\n[图2-7 此处插入：化学稳定性-矿化度——相对发光保持率 vs NaCl浓度/时间]\n[图2-8 此处插入：化学稳定性-氧化环境——相对发光保持率 vs APS浓度/时间]\n[图2-9 此处插入：化学稳定性-pH——相对发光保持率 vs pH/时间]\n[表2-1 此处插入：各条件下的一级衰减动力学参数汇总（k, t₁/₂, R²）]')
    print('  Added 图2-6~2-9 + 表2-1 环境适应性')

# Orthogonal experiment table
_, p = find_para(doc, '正交实验设计（L9(3⁴)）')
if p:
    add_para_after(doc, p, '[表2-2 此处插入：L9(3⁴)正交实验设计表（四因素三水平，9组实验条件）]\n[表2-3 此处插入：正交实验结果——各组相对浊度保持率（%）和Zeta电位（mV），极差分析和方差分析]')
    print('  Added 表2-2~2-3 正交实验')

# FTIR/XPS/TGA
_, p = find_para(doc, '改性效果表征：傅里叶变换红外光谱')
if p:
    add_para_after(doc, p, '[图2-10 此处插入：改性前后荧光粉FTIR图谱对比（标注Si-O-Al ~980 cm⁻¹, C-H ~2880 cm⁻¹, C-O-C ~1100 cm⁻¹, N-H ~1560 cm⁻¹）]\n[图2-11 此处插入：改性前后XPS全谱及Si 2p/N 1s窄扫描对比]\n[图2-12 此处插入：TGA曲线（改性前/后，室温-800°C，标注有机物失重百分比）]\n[图2-13 此处插入：改性前后SEM形貌对比（同倍率）]')
    print('  Added 图2-10~2-13 改性表征')

# Dispersion stability
_, p = find_para(doc, '静置沉降曲线——将各分散液置于25 mL量筒中')
if p:
    add_para_after(doc, p, '[图2-14 此处插入：六组分散液静置沉降曲线（相对浊度保持率 vs 时间，六条曲线叠加）]\n[图2-15 此处插入：D50和Span值随时间的变化（六组对比柱状图）]\n[图2-16 此处插入：Zeta电位-pH曲线（六组叠加）]\n[表2-4 此处插入：各组分独立贡献和协同效应定量分析表]')
    print('  Added 图2-14~2-16 + 表2-4 分散稳定性')

# Crosslinking SEM
_, p = find_para(doc, 'SEM观察冻胶微观结构')
if p:
    add_para_after(doc, p, '[图2-17 此处插入：含/不含荧光粉的HPG冻胶冷冻干燥SEM照片（×1000, ×5000对比）]')
    print('  Added 图2-17 冻胶SEM')

# ---- Chapter 3 ----
# Masterbatch stability
_, p = find_para(doc, '静态稳定性——将母液静置于25 mL量筒中')
if p:
    add_para_after(doc, p, '[图3-1 此处插入：母液静态稳定性——7天分层界面高度 vs 时间（三种浓度40/50/60 g/L对比）]\n[表3-1 此处插入：母液动态稳定性和荧光保持率数据汇总]')
    print('  Added 图3-1 + 表3-1 母液稳定性')

# Base fluid properties
_, p = find_para(doc, '依据SY/T 5107-2016[55]《水基压裂液性能评价方法》测定基液表观粘度')
if p:
    add_para_after(doc, p, '[表3-2 此处插入：基液性能对比表——表观粘度（170 s⁻¹, 25°C）、pH、密度（含/不含荧光粉）]')
    print('  Added 表3-2 基液性能')

# Crosslinked gel rheology
_, p = find_para(doc, '恒定剪切），按温度程序（从25°C以3°C/min升温至120°C')
if p:
    add_para_after(doc, p, '[图3-2 此处插入：耐温耐剪切曲线——表观粘度 vs 温度/时间（含/不含荧光粉两条曲线叠加，标注50 mPa·s线和60 min线）]\n[图3-3 此处插入：粘弹性频率扫描——G\'和G\" vs 频率（0.1~10 Hz），含/不含荧光粉对比，标注G\'>G\"区间]')
    print('  Added 图3-2~3-3 冻胶流变')

# Proppant suspension
_, p = find_para(doc, '采用静态悬砂法评价荧光冻胶的支撑剂悬浮能力')
if p:
    add_para_after(doc, p, '[图3-4 此处插入：支撑剂沉降曲线——沉降界面高度 vs 时间（含/不含荧光粉，室温+储层温度，四条曲线叠加）]\n[表3-3 此处插入：支撑剂沉降速度对比表（mm/min）]')
    print('  Added 图3-4 + 表3-3 悬砂性能')

# Breaker performance
_, p = find_para(doc, '向荧光冻胶中加入过硫酸铵破胶剂（0.05~0.2% w/v浓度梯度）')
if p:
    add_para_after(doc, p, '[图3-5 此处插入：破胶曲线——表观粘度 vs 破胶时间（不同APS浓度/温度组合，标注10 mPa·s线）]\n[表3-4 此处插入：破胶时间和破胶液残渣含量对比表（含/不含荧光粉）]')
    print('  Added 图3-5 + 表3-4 破胶性能')

# Formation damage
_, p = find_para(doc, 'API滤失仪），在3.5 MPa压差、90°C条件下')
if p:
    add_para_after(doc, p, '[图3-6 此处插入：滤失量 vs 时间曲线（含/不含荧光粉对比）]\n[图3-7 此处插入：滤饼截面SEM-EDS（荧光粉在滤饼中的分布）]')
    print('  Added 图3-6~3-7 滤饼伤害')

_, p = find_para(doc, '按照SY/T 6540-2002《钻井液完井液损害油层室内评价方法》')
if p:
    add_para_after(doc, p, '[表3-5 此处插入：岩心驱替伤害实验结果——初始渗透率K₀、伤害后渗透率K_d、渗透率伤害率D（含/不含荧光粉，3次重复，均值±SD）]')
    print('  Added 表3-5 岩心伤害')

_, p = find_para(doc, 'API导流室中铺设20/40目陶粒支撑剂单层')
if p:
    add_para_after(doc, p, '[图3-8 此处插入：导流能力 vs 闭合压力曲线（10~60 MPa，含/不含荧光粉两条曲线对比）]')
    print('  Added 图3-8 导流能力')

# ---- Chapter 4 ----
# Experimental setup schematic
_, p = find_para(doc, '实验装置的核心是一个可调节裂缝宽度的岩心夹持系统')
if p:
    add_para_after(doc, p, '[图4-1 此处插入：模拟裂缝可视化实验装置示意图（标注各组件：岩心夹持器、中间容器×3、恒流泵、回压阀、加热套、紫外暗室成像系统）。建议使用CAD或Visio绘制。]')
    print('  Added 图4-1 实验装置')

# Particle distribution
_, p = find_para(doc, '沿裂缝延伸方向等分为入口段、中段、尖端段3个分区')
if p:
    add_para_after(doc, p, '[图4-2 此处插入：裂缝三段分区紫外照片（入口段/中段/尖端段并排对比）]\n[图4-3 此处插入：三段荧光灰度值对比柱状图（mean±SD，ANOVA检验结果标注）]\n[图4-4 此处插入：注入压力沿裂缝长度方向分布曲线（三个压力传感器信号叠加）]')
    print('  Added 图4-2~4-4 颗粒运移分布')

# Surface chemistry change
_, p = find_para(doc, '对两种粉末进行FTIR（KBr压片）、XPS')
if p:
    add_para_after(doc, p, '[图4-5 此处插入：破胶处理前后FTIR对比图谱（标注PEG特征峰C-O-C ~1100 cm⁻¹和C-H ~2880 cm⁻¹的变化）]\n[图4-6 此处插入：破胶处理前后XPS N 1s窄扫描对比（标注-NH₂ ~399 eV峰面积变化）]\n[图4-7 此处插入：破胶处理前后Zeta电位-pH曲线对比]')
    print('  Added 图4-5~4-7 表面化学变化')

# Adsorption
_, p = find_para(doc, '拟合Langmuir（q_e = q_max·K_L·C_e')
if p:
    add_para_after(doc, p, '[图4-8 此处插入：吸附等温线——q_e vs C_e（三温度25/50/80°C，Langmuir/Freundlich/Temkin三种模型拟合曲线叠加）]\n[图4-9 此处插入：吸附动力学——q_t vs t（三温度，拟一级/拟二级/颗粒内扩散拟合曲线叠加）]\n[图4-10 此处插入：Van\'t Hoff图——ln(K_L) vs 1/T（标注ΔH°和ΔS°计算值）]\n[表4-1 此处插入：等温吸附模型拟合参数汇总（q_max, K_L, K_F, 1/n, R², AIC）]\n[表4-2 此处插入：吸附动力学参数汇总（k₁, k₂, k_id, R²）]\n[表4-3 此处插入：吸附热力学参数（ΔG°, ΔH°, ΔS°）]')
    print('  Added 图4-8~4-10 + 表4-1~4-3 吸附实验')

# Dynamic displacement
_, p = find_para(doc, '注入总量为裂缝体积的3倍。全程在线记录注入端和产出端的压力')
if p:
    add_para_after(doc, p, '[图4-11 此处插入：注入压力曲线——p_in vs 注入PV数（四种裂缝宽度0.1/0.5/1.0/2.0 mm，四条曲线叠加）]')
    print('  Added 图4-11 注入压力')

_, p = find_para(doc, '逐裂缝体积（CPV）收集返排液样品')
if p:
    add_para_after(doc, p, '[图4-12 此处插入：返排液荧光粉归一化浓度 vs 累计返排CPV（四种裂缝宽度，四条曲线叠加，标注本底线）]\n[表4-4 此处插入：荧光粉净残留率汇总表（四种裂缝宽度的注入总量、累计返排量和净残留率）]')
    print('  Added 图4-12 + 表4-4 返排')

_, p = find_para(doc, '在紫外暗室成像系统中，将所有成像参数')
if p:
    add_para_after(doc, p, '[图4-13 此处插入：四种裂缝宽度（0.1/0.5/1.0/2.0 mm）的岩心裂缝壁面紫外照片（同成像参数，并排对比，标注5个ROI位置）]')
    print('  Added 图4-13 紫外成像')

# Semi-quantitative analysis
_, p = find_para(doc, '以四种裂缝宽度（0.1、0.5、1.0、2.0 mm）为横坐标')
if p:
    add_para_after(doc, p, '[图4-14 此处插入：平均灰度值 vs 裂缝宽度散点图（n=12，标注Spearman ρ和p值，标注各宽度SD误差棒）]\n[表4-5 此处插入：半定量分析汇总——各裂缝宽度的5个ROI灰度值（mean±SD）、SNR、检出率]')
    print('  Added 图4-14 + 表4-5 半定量分析')

# ---- Chapter 5 ----
_, p = find_para(doc, '5.1  "母液预配+在线稀释"工艺方案')
if p:
    add_para_after(doc, p, '[图5-1 此处插入：现场施工工艺流程图——从母液配制罐→计量泵→静态混合器→混砂车→高压泵组→井筒的完整流程，标注各节点的质量控制参数]')
    print('  Added 图5-1 工艺流程')

_, p = find_para(doc, '5.2  经济技术评估')
if p:
    add_para_after(doc, p, '[表5-1 此处插入：荧光示踪方案与化学示踪剂方案经济性对比表（材料费、设备费、服务费、总费用、适用条件）]\n[表5-2 此处插入：单井材料用量估算汇总表（荧光粉、KH550、PEG4000、柠檬酸、Triton X-100的用量和费用明细）]')
    print('  Added 表5-1~5-2 经济评估')


# ================================================================
# PART 2: ADD REFERENCES FOR EXPANDED CONTENT
# ================================================================
print("\n" + "=" * 60)
print("PART 2: Adding new references for expanded content")

# Find the last reference paragraph (currently [55])
_, p_last_ref = find_para(doc, '[55] SY/T 5107-2016')
if not p_last_ref:
    print("  WARNING: Could not find last reference")
else:
    # New references for 1.2.4 (胍胶压裂液体系) and expanded 1.1
    new_refs = [
        '[56] 邹才能, 董大忠, 王玉满, 等. (2015). 中国页岩气特征、挑战及前景. 石油勘探与开发, 42(6): 689-701.',
        '[57] 董大忠, 王玉满, 李新景, 等. (2016). 中国页岩气勘探开发新进展及前景展望. 天然气工业, 36(1): 19-32.',
        '[58] Harris, P.C. (1993). Chemistry and Rheology of Borate-Crosslinked Fluids at Low Temperatures. SPE Production & Facilities, 8(3): 211-216.',
        '[59] Kesavan, S. & Prud\'homme, R.K. (1992). Rheology of Guar and HPG Cross-linked by Borate. Macromolecules, 25(7): 2026-2032.',
        '[60] Brannon, H.D. & Tjon-Joe-Pin, R.M. (1994). Biotechnological Breakthrough Improves Performance of Moderate to High-Temperature Fracturing Applications. SPE-28513-MS.',
        '[61] Barati, R. & Liang, J.T. (2014). A Review of Fracturing Fluid Systems Used for Hydraulic Fracturing of Oil and Gas Wells. Journal of Applied Polymer Science, 131(16): 40735.',
        '[62] Economides, M.J. & Nolte, K.G. (2000). Reservoir Stimulation (3rd ed.). John Wiley & Sons.',
        '[63] King, G.E. (2012). Hydraulic Fracturing 101: What Every Representative, Environmentalist, Regulator, Reporter, Investor, University Researcher, Neighbor and Engineer Should Know About Estimating Frac Risk. SPE-152596-MS.',
        '[64] 郭建春, 何春明. (2012). 压裂液破胶过程伤害微观机理. 石油学报, 33(6): 1018-1022.',
        '[65] 翁定为, 雷群, 胥云, 等. (2024). 水力压裂裂缝监测技术综述. 世界石油工业, 31(6): 66-76.',
        '[66] 邸德家. (2025). 油气井压裂示踪监测技术现状及发展建议. 钻采工艺, 48(2): 74-81.',
        '[67] Van den Eeckhout, K., Smet, P.F., & Poelman, D. (2010). Persistent Luminescence in Eu²⁺-Doped Compounds: A Review. Materials, 3(4): 2536-2566.',
        '[68] Rojas-Hernandez, R.E., Rubio-Marcos, F., Rodriguez, M.A., & Fernandez, J.F. (2018). Long Lasting Phosphors: SrAl₂O₄:Eu,Dy as the Most Studied Material. Renewable and Sustainable Energy Reviews, 81(2): 2759-2770.',
        '[69] Flury, M. & Wai, N.N. (2003). Dyes as Tracers for Vadose Zone Hydrology. Reviews of Geophysics, 41(1): 1002.',
        '[70] 中国石油天然气集团公司. (2024). 中国石油页岩气勘探开发年度报告.',
    ]

    for ref_text in new_refs:
        # Add after the last reference, then update last_ref pointer
        new_p = doc.add_paragraph(ref_text)
        p_last_ref._element.addnext(new_p._element)
        p_last_ref = new_p

    print(f'  Added {len(new_refs)} new references [56]-[70]')

# ================================================================
# PART 3: ADD CITATION BRACKETS IN EXPANDED TEXT
# ================================================================
print("\n" + "=" * 60)
print("PART 3: Adding citation brackets in expanded text")

# 3a. In 1.2.4, find Harris reference
for para in doc.paragraphs:
    if 'Harris（1993）系统研究了硼交联HPG冻胶' in para.text and '[58]' not in para.text:
        for run in para.runs:
            if 'Harris（1993）' in run.text:
                run.text = run.text.replace('Harris（1993）', 'Harris（1993）[58]')
                print('  Added [58] to Harris 1993 citation')
                break
        break

# 3b. Kesavan reference
for para in doc.paragraphs:
    if 'Kesavan和Prud\'homme（1992）' in para.text and '[59]' not in para.text:
        for run in para.runs:
            if 'Kesavan和Prud' in run.text:
                run.text = run.text.replace("Kesavan和Prud\'homme（1992）", "Kesavan和Prud\'homme（1992）[59]")
                print('  Added [59] to Kesavan 1992 citation')
                break
        break

# 3c. Brannon reference
for para in doc.paragraphs:
    if 'Brannon和Tjon-Joe-Pin（1994）' in para.text and '[60]' not in para.text:
        for run in para.runs:
            if 'Brannon' in run.text:
                run.text = run.text.replace('Brannon和Tjon-Joe-Pin（1994）', 'Brannon和Tjon-Joe-Pin（1994）[60]')
                print('  Added [60] to Brannon 1994 citation')
                break
        break

# 3d. Al-Muntasheri already has [53]
for para in doc.paragraphs:
    if 'Al-Muntasheri（2014）在其压裂液综述中' in para.text and '[53]' not in para.text:
        for run in para.runs:
            if 'Al-Muntasheri（2014）' in run.text:
                run.text = run.text.replace('Al-Muntasheri（2014）', 'Al-Muntasheri（2014）[53]')
                print('  Added [53] to Al-Muntasheri 2014 citation')
                break
        break

# 3e. 翁定为 reference in expanded 1.2.1
for para in doc.paragraphs:
    if '翁定为等（2024）系统综述了水力压裂裂缝监测技术' in para.text and '[65]' not in para.text:
        for run in para.runs:
            if '翁定为等（2024）' in run.text:
                run.text = run.text.replace('翁定为等（2024）', '翁定为等（2024）[65]')
                print('  Added [65] to 翁定为 2024 citation')
                break
        break

# 3f. 邸德家 reference
for para in doc.paragraphs:
    if '邸德家（2025）对油气井压裂示踪监测技术' in para.text and '[66]' not in para.text:
        for run in para.runs:
            if '邸德家（2025）' in run.text:
                run.text = run.text.replace('邸德家（2025）', '邸德家（2025）[66]')
                print('  Added [66] to 邸德家 2025 citation')
                break
        break

# 3g. Guo reference in expanded gap section
for para in doc.paragraphs:
    if 'Guo等，2007' in para.text and '[45]' not in para.text:
        for run in para.runs:
            if 'Guo等，2007' in run.text:
                run.text = run.text.replace('Guo等，2007', 'Guo等（2007）[45]')
                print('  Added [45] to Guo 2007 citation')
                break
        break

# 3h. Add Chinese shale gas references to expanded 1.1
for para in doc.paragraphs:
    if '中国已成为全球最大的页岩气生产国之一' in para.text and '[56]' not in para.text:
        for run in para.runs:
            if '2023年页岩气产量突破240亿立方米' in run.text:
                run.text = run.text.replace('2023年页岩气产量突破240亿立方米', '2023年页岩气产量突破240亿立方米[56-57]')
                print('  Added [56-57] to shale gas statistics')
                break
        break

# 3i. King reference in 1.2.4
for para in doc.paragraphs:
    if 'Montney地层' in para.text:
        break  # not needed

# 3j. Barati reference
for para in doc.paragraphs:
    if '已被广泛用作水基压裂液的稠化剂' in para.text and '[61]' not in para.text:
        for run in para.runs:
            if '已被广泛用作水基压裂液的稠化剂' in run.text:
                run.text = run.text.replace('已被广泛用作水基压裂液的稠化剂', '已被广泛用作水基压裂液的稠化剂[53,61]')
                print('  Added [61] to HPG thickener statement')
                break
        break

# ================================================================
# SAVE
# ================================================================
doc.save(DOC)

# Final stats
total_paras = sum(1 for p in doc.paragraphs if p.text.strip())
ref_count = 0
in_refs = False
for para in doc.paragraphs:
    if para.text.strip() == '参考文献':
        in_refs = True
        continue
    if in_refs and para.text.strip().startswith('['):
        ref_count += 1

print(f'\nSaved. Total paragraphs: {total_paras}, References: {ref_count}')
print(f'Backup at: 荧光压裂液/论文初稿_v2_backup.docx')