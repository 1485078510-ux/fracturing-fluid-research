# -*- coding: utf-8 -*-
"""Add remaining revision suggestions as small blue annotations to the final DOCX."""
import json
from docx import Document
from docx.shared import Pt, RGBColor

SRC = r"c:\Users\郝\Desktop\claude\向晶\页岩气藏地应力演化与裂缝扩展耦合机制研究进展与展望-润色版-最终版.docx"
DST = r"c:\Users\郝\Desktop\claude\向晶\页岩气藏地应力演化与裂缝扩展耦合机制研究进展与展望-润色版-最终版-带批注.docx"
ANN_PATH = r"c:\Users\郝\Desktop\claude\向晶\annotations_v2.json"

with open(ANN_PATH, 'r', encoding='utf-8') as f:
    annotations = json.load(f)

doc = Document(SRC)
count = 0
remaining = list(annotations)  # Make a mutable copy

for para in doc.paragraphs:
    text = para.text
    if not text.strip():
        continue

    matched = None
    for ann in remaining:
        if ann['keyword'] in text:
            matched = ann
            break

    if matched:
        # Add annotation in small blue font
        spacer = para.add_run('\n')
        spacer.font.size = Pt(4)

        note_run = para.add_run(matched['text'])
        note_run.font.size = Pt(8)
        note_run.font.color.rgb = RGBColor(0, 70, 180)
        note_run.font.name = '宋体'

        count += 1
        remaining.remove(matched)

doc.save(DST)
print(f"Added {count} annotations (out of {len(annotations)} total)")
print(f"Missed: {len(remaining)}")
for a in remaining:
    print(f"  - keyword: {a['keyword'][:60]}...")
print(f"Saved to: {DST}")
