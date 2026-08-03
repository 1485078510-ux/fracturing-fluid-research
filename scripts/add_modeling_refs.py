#!/usr/bin/env python3
"""Add 5 modeling references and renumber citations."""
from docx import Document
import re

doc = Document('四氧化三铁环氧树脂拟合/ESP-T_终稿.docx')

# ── 1. Add citations in [11] at key modeling-gap points ──
p11 = doc.paragraphs[11]
t = p11.text

# Citation markers to insert (using placeholder text that we'll renumber later)
# At "...predominantly the Korsmeyer-Peppas power law [X]" - add citation about BTC methods
old1 = "predominantly the Korsmeyer-Peppas power law"
new1 = old1 + " [X1,X2]"
t = t.replace(old1, new1)

# At "...remains unmodeled [X]" - add citation about ADE/fracture tracer modeling
old2 = "remains unmodeled."
new2 = "remains unmodeled [X3,X4,X5]."
t = t.replace(old2, new2)

for r in p11.runs: r.text = ''
p11.runs[0].text = t
print('[11] 5 new citation markers added: [X1,X2] [X3,X4,X5]')

# ── 2. Append 5 new references to the reference list ──
# New refs (as refs 30-34)
NEW_REFS = [
    "[30]\tSHOOK G M, POPE G A, ASAKAWA K. Determining reservoir properties and flood performance from tracer test analysis [C]. SPE 124614, SPE Annual Technical Conference and Exhibition, New Orleans, 2009.",
    "[31]\tFONTALVO E M, OLIVEIRA M C, SCHOEGGL F, et al. Physical interpretation of interwell partitioning tracer tests for estimation of remaining oil saturation in layered carbonate reservoirs [J]. Transport in Porous Media, 2025, 152: 21-45.",
    "[32]\tLIU J, WANG H, ZHANG T, et al. Study on interpretation method of multistage fracture tracer flowback curve in tight oil reservoirs [J]. ACS Omega, 2024, 9: 10852-10864.",
    "[33]\tTIAN W, DARNLEY A, DEMPSEY D. Quantifying fracture interference and allocating load recovery and hydrocarbon production in various well configuration using chemical tracers [C]. SPE-201292-MS, SPE Annual Technical Conference and Exhibition, Virtual, 2020.",
    "[34]\tLI J, JIANG H, WANG B, et al. Tracer flowback modeling and characterization of complex fracture networks in multi-fractured horizontal wells [J]. Journal of Natural Gas Science and Engineering, 2022, 98: 104987.",
]

# Find References section
ref_start = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "References":
        ref_start = i
        break

# Count current refs - find the last ref paragraph
last_ref_idx = ref_start + 1
for i in range(ref_start + 1, len(doc.paragraphs)):
    t = doc.paragraphs[i].text.strip()
    if re.match(r'\[\d+\]', t):
        last_ref_idx = i

print(f"Current refs end at [{last_ref_idx}]")

# Append new refs after the last existing ref
insert_pos = last_ref_idx + 1
for j, ref_text in enumerate(NEW_REFS):
    target_idx = insert_pos + j
    if target_idx < len(doc.paragraphs):
        p = doc.paragraphs[target_idx]
        for r in p.runs: r.text = ''
        if len(p.runs) == 0:
            p.add_run('')
        p.runs[0].text = ref_text
    else:
        # Add new paragraph at end
        new_p = doc.add_paragraph(ref_text)
        print(f"  Added new paragraph for ref [{30+j}]")

print(f"5 new refs appended at [{insert_pos}]-[{insert_pos+4}]")

# ── 3. Build remapping table ──
# First appearance order: find where each ref [1-34] first appears in body text
# Actually, the new refs X1-X5 are cited in [11] which comes BEFORE the original refs [13-16]
# So they should be renumbered to fit the first-appearance order

# For simplicity: the new refs (old 30-34) are cited AT paragraph [11]
# They need to take positions based on their first appearance in paragraph [11]
# Since they're all first cited in [11], let me figure out the exact order within [11]

# Current first-appearance order (up to [11]):
# [1-4], [5], [6], [7], [8-12], [X1,X2], [X3,X4,X5], [13-16]...

# Old30 (Shook) -> new [13]
# Old31 (Fontalvo) -> new [14]
# Old32 (Liu) -> new [15]
# Old33 (Tian) -> new [16]
# Old34 (Li) -> new [17]
# Old13 (Zhao) -> new [18]
# ...everything shifts by +5

OLD2NEW = {}
shift = 0
for old in range(1, 35):
    if old <= 12:
        OLD2NEW[old] = old  # no change for refs 1-12
    elif old <= 29:
        OLD2NEW[old] = old + 5  # original refs 13-29 shift by +5
    else:
        # old 30->13, 31->14, 32->15, 33->16, 34->17
        OLD2NEW[old] = old - 17

# ── 4. Remap all in-text citations ──
def remap_cite(match):
    content = match.group(1)
    if re.search(r'[a-zA-Z]', content) or content.startswith('X'):
        if content == 'X1': return '[13,14]'
        if content == 'X2': return '[13,14]'
        if content == 'X3': return '[15,16,17]'
        if content == 'X4': return '[15,16,17]'
        if content == 'X5': return '[15,16,17]'
        return f'[{content}]'

    parts = []
    for part in content.split(','):
        part = part.strip()
        if '-' in part:
            pp = part.split('-')
            if len(pp) == 2:
                try:
                    a = OLD2NEW.get(int(pp[0]), int(pp[0]))
                    b = OLD2NEW.get(int(pp[1]), int(pp[1]))
                    if a == b: parts.append(str(a))
                    elif b == a+1: parts.append(f'{a},{b}')
                    else: parts.append(f'{a}-{b}')
                except: parts.append(part)
            else: parts.append(part)
        else:
            try: parts.append(str(OLD2NEW.get(int(part), int(part))))
            except: parts.append(part)

    # Simplify consecutive numbers
    nums = []
    for p in parts:
        try: nums.append(int(p))
        except: return f'[{",".join(parts)}]'

    nums.sort()
    result = []
    i = 0
    while i < len(nums):
        start = nums[i]
        while i+1 < len(nums) and nums[i+1] == nums[i]+1:
            i += 1
        end = nums[i]
        if start == end: result.append(str(start))
        elif end == start+1: result.append(f'{start},{end}')
        else: result.append(f'{start}-{end}')
        i += 1
    return f'[{",".join(result)}]'

fixed = 0
for i in range(ref_start):
    p = doc.paragraphs[i]
    old_t = p.text
    new_t = re.sub(r'\[([^\]]+)\]', remap_cite, old_t)
    if new_t != old_t:
        for r in p.runs: r.text = ''
        p.runs[0].text = new_t
        fixed += 1
print(f"Remapped {fixed} paragraphs")

# ── 5. Reorder reference list by new numbering ──
# Read all current refs (ref_start+1 to ref_start+34)
old_refs = {}
for j in range(34):
    idx = ref_start + 1 + j
    if idx < len(doc.paragraphs):
        old_refs[j+1] = doc.paragraphs[idx].text

# Build new order
NEW2OLD = {v: k for k, v in OLD2NEW.items()}
for n in range(1, 35):
    old_n = NEW2OLD.get(n, n)
    new_text = old_refs.get(old_n, f'[{n}] MISSING')
    # Update the bracketed number
    new_text = re.sub(r'^\[\d+\]', f'[{n}]', new_text)
    idx = ref_start + 1 + (n-1)
    if idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]
        for r in p.runs: r.text = ''
        p.runs[0].text = new_text

# Clear any remaining old ref paras beyond 34
for i in range(ref_start + 1 + 34, len(doc.paragraphs)):
    t = doc.paragraphs[i].text.strip()
    if t and t[0] == '[':
        for r in doc.paragraphs[i].runs: r.text = ''

print("Reference list reordered: 34 entries")

doc.save('四氧化三铁环氧树脂拟合/ESP-T_终稿.docx')
print('Saved.')