import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"c:\Users\郝\Desktop\claude\荧光压裂液\荧光压裂液专利申请书_优化版.docx"
OUT = r"c:\Users\郝\Desktop\claude\荧光压裂液\荧光压裂液专利申请书_多体系版.docx"

doc = Document(SRC)

def replace_in_para(p, old, new):
    for run in p.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)

def find_and_replace(doc, old, new):
    count = 0
    for p in doc.paragraphs:
        if old in p.text:
            replace_in_para(p, old, new)
            count += 1
    return count

mods = []
mods.append(("压裂液注入阶段要求荧光粉在胍胶基液中保持均匀稳定悬浮",
             "压裂液注入阶段要求荧光粉在各类稠化剂基液（包括植物胶类、合成聚合物类及粘弹性表面活性剂类）中保持均匀稳定悬浮"))

mods.append(("与常规水力压裂工艺（胍胶基液、硼交联剂、过硫酸铵破胶）全流程兼容",
             "与常规水力压裂工艺（各类水基压裂液体系，包括植物胶类、合成聚合物类及清洁压裂液，以及硼交联剂、过硫酸铵破胶等）全流程兼容"))

mods.append(("阻断其对胍胶交联反应的干扰", "阻断其对稠化剂交联反应的干扰"))

old_t = "稠化剂，选自羟丙基胍胶（HPG）、瓜尔胶、羧甲基羟丙基胍胶（CMHPG）中的一种或多种，优选HPG，在压裂液终液中的浓度为0.3~1.0 wt%"
new_t = ("稠化剂，选自植物胶类（羟丙基胍胶（HPG）、瓜尔胶、羧甲基羟丙基胍胶（CMHPG）、田菁胶）、"
         "合成聚合物类（聚丙烯酰胺（PAM）、部分水解聚丙烯酰胺（HPAM）、疏水缔合聚合物）、"
         "粘弹性表面活性剂类（季铵盐型、甜菜碱型、氧化胺型，用于构建清洁压裂液体系）"
         "或滑溜水用减阻剂类（阴离子型PAM基减阻剂）中的一类或多类，优选羟丙基胍胶（HPG）；"
         "当选用植物胶类稠化剂时，其在压裂液终液中的浓度为0.3~1.0 wt%；"
         "当选用合成聚合物类稠化剂时，其在压裂液终液中的浓度为0.1~0.8 wt%；"
         "当选用粘弹性表面活性剂类时，其在压裂液终液中的浓度为1.0~5.0 wt%")
mods.append((old_t, new_t))

old_xl = "交联剂，选自有机硼交联剂、有机锆交联剂或有机钛交联剂中的一种或多种，优选有机硼延缓交联剂，在压裂液终液中的浓度为0.1~0.5 vol%"
new_xl = ("交联剂，当稠化剂为植物胶类或合成聚合物类时，选自有机硼交联剂、有机锆交联剂或有机钛交联剂中的一种或多种，优选有机硼延缓交联剂，在压裂液终液中的浓度为0.1~0.5 vol%；"
          "当稠化剂为粘弹性表面活性剂类或滑溜水用减阻剂类时，可选择不添加交联剂或添加微量交联促进剂")
for p in doc.paragraphs:
    t = p.text.strip()
    if old_xl in p.text and not (t.startswith("12.") or t.startswith("13.") or t.startswith("14.") or t.startswith("15.") or t.startswith("根据权利")):
        replace_in_para(p, old_xl, new_xl)
        print(f"  Crosslinker: {p.text[:60]}...")

mods.append(("注入HPG基液主流中", "注入稠化剂基液主流中"))
mods.append(("保障荧光粉在胍胶基液中的均匀分散", "保障荧光粉在各类稠化剂基液中的均匀分散"))
mods.append(("保障荧光粉在胍胶冻胶中的稳定悬浮和随携砂液的高效运移", 
             "保障荧光粉在各类水基压裂液冻胶/溶液中的稳定悬浮和随携砂液的高效运移"))

mods.append(("以HPG胍胶冻胶为常规压裂液载体（区别于可凝固树脂体系），采用有机硼延缓交联剂和过硫酸铵破胶剂，可完全沿用现有压裂泵注设备和施工程序",
             "以各类水基压裂液（包括HPG胍胶冻胶、聚丙烯酰胺合成聚合物压裂液、粘弹性表面活性剂清洁压裂液及滑溜水等体系）为常规载体（区别于可凝固树脂体系），可兼容有机硼/有机锆/有机钛等金属交联剂或非交联自增稠机制以及过硫酸铵等氧化破胶剂，完全沿用现有压裂泵注设备和施工程序"))

mods.append(("阻断其对胍胶分子链邻位顺式羟基的竞争性交联干扰（化学螯合层）",
             "阻断其对植物胶类稠化剂分子链邻位顺式羟基的竞争性交联干扰，或消除其对合成聚合物类稠化剂金属交联位点的占用（化学螯合层）"))

mods.append(("荧光压裂液体系以HPG为稠化剂、有机硼为交联剂、过硫酸铵为破胶剂",
             "荧光压裂液体系以羟丙基胍胶（HPG）等植物胶类或聚丙烯酰胺等合成聚合物类为稠化剂、以有机硼等金属交联剂（或采用VES非交联自增稠机制）构建冻胶/溶液体系、过硫酸铵为破胶剂"))

print("Applying modifications...")
total = 0
for old, new in mods:
    c = find_and_replace(doc, old, new)
    if c > 0:
        total += c
        print(f"  [{c}x] {old[:70]}")
    else:
        print(f"  [MISS] {old[:70]}")
print(f"Total: {total} replacements")

# Update claim 13 thickener list
claim13_old = "所述稠化剂选自羟丙基胍胶（HPG）、瓜尔胶、羧甲基羟丙基胍胶（CMHPG）中的一种或多种，在压裂液终液中的浓度为0.3~1.0 wt%"
claim13_new = ("所述稠化剂选自以下类别中的一种或多种：植物胶类（羟丙基胍胶（HPG）、瓜尔胶、羧甲基羟丙基胍胶（CMHPG）），"
               "合成聚合物类（聚丙烯酰胺（PAM）、部分水解聚丙烯酰胺（HPAM）），"
               "粘弹性表面活性剂类（季铵盐型、甜菜碱型），或滑溜水用减阻剂类；"
               "当选用植物胶类或合成聚合物类稠化剂时，其在压裂液终液中的浓度为0.1~1.0 wt%")
for p in doc.paragraphs:
    t = p.text.strip()
    if claim13_old in p.text and t.startswith("13."):
        replace_in_para(p, claim13_old, claim13_new)
        print("  Claim 13 updated")

# Add multi-system note after industrial note
for i, p in enumerate(doc.paragraphs):
    if "工业应用前景" in p.text and "页岩气" in p.text:
        new_p = OxmlElement('w:p')
        new_pPr = OxmlElement('w:pPr')
        sp_el = OxmlElement('w:spacing')
        sp_el.set(qn('w:line'), '360'); sp_el.set(qn('w:lineRule'), 'auto')
        new_pPr.append(sp_el)
        ind_el = OxmlElement('w:ind')
        ind_el.set(qn('w:firstLine'), '480')
        new_pPr.append(ind_el)
        new_p.append(new_pPr)
        
        new_r = OxmlElement('w:r')
        new_rPr = OxmlElement('w:rPr')
        rf = OxmlElement('w:rFonts')
        rf.set(qn('w:ascii'), 'Times New Roman')
        rf.set(qn('w:hAnsi'), 'Times New Roman')
        rf.set(qn('w:eastAsia'), '宋体')
        new_rPr.append(rf)
        sz_el = OxmlElement('w:sz'); sz_el.set(qn('w:val'), '24')
        new_rPr.append(sz_el)
        new_r.append(new_rPr)
        
        new_t = OxmlElement('w:t')
        new_t.set(qn('xml:space'), 'preserve')
        new_t.text = (
            "需要说明的是，以上实施例以羟丙基胍胶（HPG）为代表性稠化剂进行实验验证，"
            "但本发明所述改性稀土铝酸盐荧光粉的适用体系不限于此。基于相同的双层改性加协同分散技术原理，"
            "该荧光粉同样适用于以下压裂液体系：（1）植物胶类体系，包括瓜尔胶、羧甲基羟丙基胍胶（CMHPG）、"
            "田菁胶等，其交联机理与HPG体系一致，改性荧光粉中螯合剂对多价金属阳离子的络合作用可普遍阻断"
            "其对各类植物胶邻位顺式羟基交联位点的竞争干扰；（2）合成聚合物类体系，包括聚丙烯酰胺（PAM）、"
            "部分水解聚丙烯酰胺（HPAM）、疏水缔合聚合物等，PEG外层的空间位阻效应同样适用于非植物胶类高分子"
            "溶液环境，保障颗粒分散稳定性，且活性氨基与砂岩壁面的锚定机制独立于稠化剂类型；"
            "（3）粘弹性表面活性剂（VES）清洁压裂液体系，包括季铵盐型、甜菜碱型、氧化胺型等胶束自组装体系，"
            "由于VES体系不含高分子聚合物，不存在多价金属阳离子对交联的竞争干扰问题，改性荧光粉的分散"
            "与锚定性能预期更为优越；（4）滑溜水体系，以低浓度PAM基减阻剂为主要添加剂，荧光粉在低粘度"
            "环境中的分散稳定性由PEG外层空间位阻和协同分散助剂体系共同保障。本领域技术人员可根据具体"
            "储层条件和施工要求，在上述体系中选择合适的稠化剂类型，并按相应行业标准调整配方参数。"
        )
        new_r.append(new_t)
        new_p.append(new_r)
        
        parent = p._element.getparent()
        idx = list(parent).index(p._element)
        parent.insert(idx + 1, new_p)
        print("  Multi-system note added")
        break

doc.save(OUT)
print(f"\nSaved to: {OUT}")