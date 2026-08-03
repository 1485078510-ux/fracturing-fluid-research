#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Add organic transitions between sections within each chapter."""
import sys
from docx import Document

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

DOC = r'荧光压裂液/论文初稿_v2.docx'
doc = Document(DOC)

def find_para(doc, snippet, skip=0):
    found = 0
    for i, para in enumerate(doc.paragraphs):
        if snippet in para.text:
            if found >= skip:
                return i, para
            found += 1
    return None, None

def add_para_after(doc, target_para, text):
    new_para = doc.add_paragraph(text)
    target_para._element.addnext(new_para._element)
    return new_para

# ================================================================
# CHAPTER 2 transitions
# ================================================================
print("=== Chapter 2 section transitions ===")

# 2.2 → 2.3: After 2.2 data placeholder or last content, before 2.3 heading
_, p = find_para(doc, '2.3  双层表面改性工艺优化')
if p:
    # Find the paragraph immediately before 2.3 heading
    for i, para in enumerate(doc.paragraphs):
        if '2.3  双层表面改性工艺优化' in para.text and i > 0:
            prev = doc.paragraphs[i-1]
            trans = (
                '2.2节的储层环境适应性评价回答了一个前置性问题：SrAl₂O₄:Eu²⁺,Dy³⁺荧光粉'
                '在井下工况中"能不能用"。实验将给出其适用边界——在什么温度、矿化度、pH和氧化'
                '条件下，荧光信号能够保持足够的强度和时间窗口。然而，即便荧光粉在这些条件下表现'
                '良好，其原始表面是亲水性的无机氧化物，与有机聚合物基液之间缺乏亲和性，且在水中'
                '会缓慢水解。因此，必须对荧光粉进行表面化学改造——这正是2.3节的核心工作。'
            )
            add_para_after(doc, prev, trans)
            print('  Added 2.2→2.3')
            break

# 2.4 → 2.5 (本章小结): After 2.4 data placeholder
_, p = find_para(doc, '2.5  本章小结')
if p:
    for i, para in enumerate(doc.paragraphs):
        if '2.5  本章小结' in para.text and i > 0:
            prev = doc.paragraphs[i-1]
            trans = (
                '2.4节的六组对照实验完成了从"单组分改性"到"多组分协同"的递进，确立了荧光粉'
                '在HPG基液中长期稳定悬浮的配方基础。从原始荧光粉到改性荧光粉再到稳定的悬浮分散液，'
                '本章完成了材料层面的三步递进。以下本章小结汇总第二章的主要发现。'
            )
            add_para_after(doc, prev, trans)
            print('  Added 2.4→2.5')
            break


# ================================================================
# CHAPTER 3 transitions
# ================================================================
print("\n=== Chapter 3 section transitions ===")

# 3.1 → 3.2: After 3.1 data placeholder, before 3.2 heading
_, p = find_para(doc, '3.2  基液性能评价')
if p:
    for i, para in enumerate(doc.paragraphs):
        if '3.2  基液性能评价' in para.text and i > 0:
            prev = doc.paragraphs[i-1]
            trans = (
                '3.1节制备的荧光母液是后续所有压裂液性能评价的"原料"。将母液引入HPG基液后，'
                '第一个需要确认的问题是：基液的基础物理性质——粘度、pH、密度——是否因荧光粉的'
                '加入而偏离了常规HPG压裂液的正常范围？如果基液性质已经显著改变，后续的交联和'
                '破胶设计将失去参照基准。3.2节对此进行定量比较。'
            )
            add_para_after(doc, prev, trans)
            print('  Added 3.1→3.2')
            break

# 3.2 → 3.3: After 3.2 data placeholder
_, p = find_para(doc, '3.3  交联冻胶性能评价')
if p:
    for i, para in enumerate(doc.paragraphs):
        if '3.3  交联冻胶性能评价' in para.text and i > 0:
            prev = doc.paragraphs[i-1]
            trans = (
                '3.2节确认了基液性质在荧光粉引入后未发生显著偏移。接下来的核心问题是交联——'
                '这是胍胶压裂液从"稠水"变为"冻胶"的关键化学步骤，直接决定了造缝宽度和支撑剂'
                '输送能力。3.3节系统评价荧光粉对交联过程和冻胶流变学特性的影响。'
            )
            add_para_after(doc, prev, trans)
            print('  Added 3.2→3.3')
            break

# 3.3 → 3.4: After 3.3 data placeholder
_, p = find_para(doc, '3.4  悬砂性能评价')
if p:
    for i, para in enumerate(doc.paragraphs):
        if '3.4  悬砂性能评价' in para.text and i > 0:
            prev = doc.paragraphs[i-1]
            trans = (
                '3.3节的流变学表征提供了冻胶粘弹性的基础数据，但工程师更关心的是一个直观的问题：'
                '支撑剂在荧光冻胶中沉降有多快？G\'>G\"的弹性主导准则只是间接判断，静态悬砂实验'
                '可以直接测量支撑剂的沉降速度。3.4节将这一工程关切转化为定量实验。'
            )
            add_para_after(doc, prev, trans)
            print('  Added 3.3→3.4')
            break

# 3.4 → 3.5: After 3.4 data placeholder
_, p = find_para(doc, '3.5  破胶性能评价')
if p:
    for i, para in enumerate(doc.paragraphs):
        if '3.5  破胶性能评价' in para.text and i > 0:
            prev = doc.paragraphs[i-1]
            trans = (
                '3.3和3.4节评价了冻胶在注入阶段的携砂能力。然而，压裂液在完成造缝和支撑剂'
                '输送后，必须彻底破胶——将高粘度冻胶降解为低粘度液体以便返排，否则残留的冻胶'
                '将永久堵塞裂缝。荧光粉的存在是否会干扰破胶过程？3.5节对此进行考察。'
            )
            add_para_after(doc, prev, trans)
            print('  Added 3.4→3.5')
            break

# 3.6 → 3.7: After 3.6 data placeholder
_, p = find_para(doc, '3.7  本章小结')
if p:
    for i, para in enumerate(doc.paragraphs):
        if '3.7  本章小结' in para.text and i > 0:
            prev = doc.paragraphs[i-1]
            trans = (
                '3.6节的地层伤害评价是压裂液-储层相容性检验的最后一道关口。至此，第三章从基液'
                '到冻胶、从悬砂到破胶、从工程性能到地层伤害，完成了荧光压裂液体系的全方位体检。'
                '以下本章小结汇总各项评价结果。'
            )
            add_para_after(doc, prev, trans)
            print('  Added 3.6→3.7')
            break


# ================================================================
# CHAPTER 4 transitions
# ================================================================
print("\n=== Chapter 4 section transitions ===")

# 4.2 → 4.3: After 4.2 data placeholder
_, p = find_para(doc, '4.3  破胶前后荧光粉表面化学变化')
if p:
    for i, para in enumerate(doc.paragraphs):
        if '4.3  破胶前后荧光粉表面化学变化' in para.text and i > 0:
            prev = doc.paragraphs[i-1]
            trans = (
                '4.2节的颗粒运移实验关注的是"荧光粉到了哪里"——即空间分布均匀性的问题。'
                '但本研究的核心工作假说——破胶诱导的表面功能切换——需要在分子层面获得直接化学'
                '证据才能成立。如果PEG包覆层在破胶后并未脱附，那么后续的壁面锚定将缺乏驱动机制。'
                '4.3节转向这个分子层面的关键检验。'
            )
            add_para_after(doc, prev, trans)
            print('  Added 4.2→4.3')
            break

# 4.4 → 4.5: After 4.4 data placeholder
_, p = find_para(doc, '4.5  动态驱替四阶段实验')
if p:
    for i, para in enumerate(doc.paragraphs):
        if '4.5  动态驱替四阶段实验' in para.text and i > 0:
            prev = doc.paragraphs[i-1]
            trans = (
                '4.4节的静态吸附实验在热力学和动力学层面验证了改性荧光粉对砂岩壁面的锚定能力。'
                '但静态条件是简化的——它忽略了注入阶段的流动剪切历史、裂缝壁面的粗糙形貌和返排'
                '阶段的液流冲刷。真实的锚定过程是在动态流动中完成的，4.5节将此前的静态认识投入'
                '到动态、多阶段的压裂模拟中进行检验。'
            )
            add_para_after(doc, prev, trans)
            print('  Added 4.4→4.5')
            break

# 4.5 → 4.6: After 4.5 data placeholder
_, p = find_para(doc, '4.6  荧光信号与裂缝几何的半定量关系')
if p:
    for i, para in enumerate(doc.paragraphs):
        if '4.6  荧光信号与裂缝几何的半定量关系' in para.text and i > 0:
            prev = doc.paragraphs[i-1]
            trans = (
                '4.5节的四阶段实验从定性层面验证了荧光示踪技术的工程可行性——可注入、可锚定、'
                '可成像。但一项诊断技术的价值取决于它能否提供超越"有/无"的信息。如果裂缝宽度'
                '越大、荧光灰度越高，且这一关系具有统计显著性，那么紫外图像就不只是"看到了裂缝"，'
                '而是可以半定量地"读出裂缝宽度"。4.6节探索这一可能性的边界。'
            )
            add_para_after(doc, prev, trans)
            print('  Added 4.5→4.6')
            break

# 4.6 → 4.7: After 4.6 data placeholder
_, p = find_para(doc, '4.7  本章小结')
if p:
    for i, para in enumerate(doc.paragraphs):
        if '4.7  本章小结' in para.text and i > 0:
            prev = doc.paragraphs[i-1]
            trans = (
                '4.6节的分析为荧光信号的解读建立了从定性到半定量的递进框架。至此，第四章从'
                '颗粒运移、界面化学、吸附参数和动态驱替四个互补维度，构成了荧光示踪技术工程'
                '可行性的完整证据链。以下本章小结汇总各项发现。'
            )
            add_para_after(doc, prev, trans)
            print('  Added 4.6→4.7')
            break


# ================================================================
# CHAPTER 5 transitions
# ================================================================
print("\n=== Chapter 5 section transitions ===")

# 5.1 → 5.2
_, p = find_para(doc, '5.2  经济技术评估')
if p:
    for i, para in enumerate(doc.paragraphs):
        if '5.2  经济技术评估' in para.text and i > 0:
            prev = doc.paragraphs[i-1]
            trans = (
                '5.1节设计了现场施工的工艺流程和质量控制方案。一项新技术的现场接受度不仅取决'
                '于技术可行性，还取决于经济合理性——如果示踪方案的成本远超其带来的诊断收益，'
                '工程推广将失去商业驱动力。5.2节从材料用量、成本构成和场景化经济分析三个角度'
                '进行初步评估。'
            )
            add_para_after(doc, prev, trans)
            print('  Added 5.1→5.2')
            break

# 5.2 → 5.3
_, p = find_para(doc, '5.3  技术定位与适用条件')
if p:
    for i, para in enumerate(doc.paragraphs):
        if '5.3  技术定位与适用条件' in para.text and i > 0:
            prev = doc.paragraphs[i-1]
            trans = (
                '5.2节的经济性分析表明，荧光示踪方案在特定场景下具有经济可行性。然而，任何'
                '技术都有其适用边界——超出边界的使用不仅无效，还可能产生误导性的诊断结论。'
                '5.3节在本论文全部研究工作的基础上，系统界定荧光示踪压裂液技术的应用场景和边界条件。'
            )
            add_para_after(doc, prev, trans)
            print('  Added 5.2→5.3')
            break


# ================================================================
# SAVE
# ================================================================
doc.save(DOC)
print(f'\nAll section transitions saved to {DOC}')