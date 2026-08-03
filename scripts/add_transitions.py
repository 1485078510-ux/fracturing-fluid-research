#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Add organic inter-chapter bridges throughout the thesis."""
import sys
from docx import Document

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

DOC = r'荧光压裂液/论文初稿_v2.docx'
doc = Document(DOC)

def find_para(doc, snippet, skip=0):
    found = 0
    for i, para in enumerate(doc.paragraphs):
        if snippet in para.text:
            if found >= skip:
                return i, para
            found += 1
    return None, None

def add_para_after(doc, target_para, text):
    new_para = doc.add_paragraph(text)
    target_para._element.addnext(new_para._element)
    return new_para


# ================================================================
# 1. Chapter 1 closing bridge → Chapter 2
#    Insert after 1.5 (论文结构安排), before 第二章 heading
# ================================================================
print("1. Adding Chapter 1→2 bridge...")
_, p_15_end = find_para(doc, '第六章：结论与展望。总结全文主要结论和创新点')
if p_15_end:
    bridge_1_2 = (
        '以上六章沿着从宏观问题到微观机制、从材料设计到工程应用的逻辑链展开。'
        '第一章提出的三个层次技术缺口——材料悬浮稳定性、工艺功能切换、动态评价方法——'
        '分别在后续的第二、三、四章中得到回应。从下一章开始，论文进入具体的实验研究工作：'
        '首先从示踪材料的"原材料"——SrAl₂O₄:Eu²⁺,Dy³⁺荧光粉——出发，建立其性能基线，'
        '并解决其在压裂液环境中的表面化学问题。'
    )
    add_para_after(doc, p_15_end, bridge_1_2)
    print('  Added Ch1→Ch2 bridge')


# ================================================================
# 2. Chapter 2 closing bridge → Chapter 3
#    Insert after 2.5 本章小结 (the placeholder paragraph)
# ================================================================
print("2. Adding Chapter 2→3 bridge...")
_, p_25_end = find_para(doc, '2.5  本章小结')
# Find the placeholder after 2.5 heading
in_25 = False
for i, para in enumerate(doc.paragraphs):
    if '2.5  本章小结' in para.text:
        in_25 = True
        continue
    if in_25 and para.text.strip():
        # This is the placeholder or content after 2.5 heading
        bridge_2_3 = (
            '第二章的工作为荧光粉从"工业原料"到"压裂液功能组分"的转化奠定了材料学基础。'
            '然而，改性荧光粉在HPG基液中的良好分散仅是第一步——它能否在完整的压裂液体系中'
            '（含交联剂、破胶剂、支撑剂）保持性能稳定，且不损害压裂液既有的工程功能，'
            '是决定其能否从实验室走向压裂现场的第二个关键问题。第三章将改性荧光粉引入HPG压裂液'
            '全配方体系，依照行业标准进行全面性能评价。'
        )
        add_para_after(doc, para, bridge_2_3)
        print(f'  Added Ch2→Ch3 bridge after P{i}')
        break


# ================================================================
# 3. Chapter 3 closing bridge → Chapter 4
#    Insert after 3.7 本章小结
# ================================================================
print("3. Adding Chapter 3→4 bridge...")
in_37 = False
for i, para in enumerate(doc.paragraphs):
    if '3.7  本章小结' in para.text:
        in_37 = True
        continue
    if in_37 and para.text.strip():
        bridge_3_4 = (
            '通过第三章的系统评价，荧光压裂液在基液、冻胶、悬砂和破胶等各项工程性能上'
            '与常规HPG压裂液保持了可接受的偏差范围，验证了"荧光示踪功能不损害压裂液工程性能"'
            '这一基本前提。然而，上述评价均属于压裂液自身的"体检"——它们回答的是"这罐液体合不合格"，'
            '而非"荧光粉在裂缝里是否真的附着在壁面上"。后者需要将流体注入真实的裂缝模型，'
            '经历注入、关井、破胶、返排的完整工程序列，然后取出岩心在紫外光下观察——'
            '这正是第四章的核心任务。'
        )
        add_para_after(doc, para, bridge_3_4)
        print(f'  Added Ch3→Ch4 bridge after P{i}')
        break


# ================================================================
# 4. Chapter 4 closing bridge → Chapter 5
#    Insert after 4.7 本章小结
# ================================================================
print("4. Adding Chapter 4→5 bridge...")
in_47 = False
for i, para in enumerate(doc.paragraphs):
    if '4.7  本章小结' in para.text:
        in_47 = True
        continue
    if in_47 and para.text.strip():
        bridge_4_5 = (
            '第四章通过静态吸附实验和动态驱替四阶段模拟，在实验室尺度上验证了荧光示踪方案'
            '的工程可行性：荧光粉可被注入裂缝、在破胶后锚定于壁面、并在返排冲刷后保留足够的'
            '信号强度以供紫外成像。然而，从实验室的50 mm岩心夹持器到1500 m水平井的施工现场，'
            '中间存在一个尺度跨越和工程转化的问题。实验室确定的最优配方能否在规模化配制中保持'
            '稳定性？施工工艺如何嵌入现有的压裂泵注流程？经济上是否具有竞争力？'
            '第五章将从工程转化的视角逐一审视这些问题。'
        )
        add_para_after(doc, para, bridge_4_5)
        print(f'  Added Ch4→Ch5 bridge after P{i}')
        break


# ================================================================
# 5. Chapter 5 closing bridge → Chapter 6
#    Insert after the last content paragraph of 5.3
# ================================================================
print("5. Adding Chapter 5→6 bridge...")
_, p_53_last = find_para(doc, '不适用于页岩基质微裂缝')
if p_53_last:
    bridge_5_6 = (
        '第五章从工程现场的角度审视了荧光示踪方案的可行性。需要强调的是，本章的所有估算'
        '和分析均建立在实验室确定的最优参数基础之上，在实际工程应用前需通过现场先导性试验'
        '进行校准和修正。综合前五章从材料到工程的完整研究链条，第六章对全文的主要发现、'
        '创新贡献和不足之处进行总结，并提出后续研究方向的展望。'
    )
    add_para_after(doc, p_53_last, bridge_5_6)
    print('  Added Ch5→Ch6 bridge')


# ================================================================
# 6. Improve chapter openings to echo previous chapter's ending
# ================================================================
print("6. Enhancing chapter openings...")

# Ch2 opening — already decent, add small echo
_, p_ch2_open = find_para(doc, '本章围绕SrAl₂O₄:Eu²⁺,Dy³⁺荧光粉的材料学基础展开工作，是全文研究的起点')
if p_ch2_open:
    new_open = (
        '如第一章所述，荧光示踪方案面临的首要技术障碍是材料层面的：高密度微米粉体在低密度'
        '压裂液基液中的悬浮稳定性，以及在储层环境中的化学稳定性。本章围绕这两个核心问题，'
        '从SrAl₂O₄:Eu²⁺,Dy³⁺荧光粉的基础物性出发，系统评价其环境适应性，并通过表面化学'
        '改性赋予其工程所需的多重功能。'
    )
    for run in p_ch2_open.runs:
        run.text = ''
    if p_ch2_open.runs:
        p_ch2_open.runs[0].text = new_open
    print('  Enhanced Ch2 opening')

# Ch3 opening — echo Ch2's dispersion work
_, p_ch3_open = find_para(doc, '在第二章完成荧光粉表面改性和分散稳定性优化的基础上')
if p_ch3_open:
    new_open = (
        '第二章解决了荧光粉"如何在HPG基液中均匀悬浮"的材料学问题。然而，一种功能添加剂'
        '的终极检验标准不是它自身性能有多好，而是它"不破坏原有体系"。对荧光示踪材料而言，'
        '这意味着引入荧光粉后的HPG压裂液必须在基液粘度、交联效率、悬砂能力、破胶彻底性和'
        '地层伤害等核心工程指标上，与不含荧光粉的常规HPG压裂液保持在可接受的偏差范围内。'
        '本章将第二章确定的最优改性方案引入HPG压裂液全配方体系，依据石油天然气行业标准'
        '进行系统性对标评价。'
    )
    for run in p_ch3_open.runs:
        run.text = ''
    if p_ch3_open.runs:
        p_ch3_open.runs[0].text = new_open
    print('  Enhanced Ch3 opening')

# Ch4 opening
_, p_ch4_open = find_para(doc, '第二、三章分别从材料和流体层面验证了荧光示踪方案的要素可行性')
if p_ch4_open:
    new_open = (
        '第二章和第三章分别从材料改性和流体性能两个层面回答了"荧光粉能否用"和"压裂液是否'
        '仍然合格"的问题。但这两个层面的回答都是静态的、离体的——它们没有触及一个关键的动态'
        '问题：在真实的裂缝流动环境中，经历注入剪切、高温关井破胶和返排冲刷之后，荧光粉是否'
        '仍然牢固附着在裂缝壁面上？换言之，前三章验证了荧光压裂液的"自身品质"，但尚未验证其'
        '"工作效果"。本章构建模拟裂缝可视化实验装置，通过颗粒运移、界面化学、静态吸附和动态'
        '驱替四个递进层次的实验，将验证从"离体"推进到"在体"。'
    )
    for run in p_ch4_open.runs:
        run.text = ''
    if p_ch4_open.runs:
        p_ch4_open.runs[0].text = new_open
    print('  Enhanced Ch4 opening')

# Ch5 opening
_, p_ch5_open = find_para(doc, '前述章节从实验室尺度验证了荧光压裂液体系的技术可行性')
if p_ch5_open:
    new_open = (
        '第四章的动态驱替实验在实验室尺度（50 mm岩心）上验证了荧光示踪的基本工作原理。'
        '然而，一项井下技术的完整开发链条需要回答三个递进问题：实验室里行不行？现场施工'
        '可不可行？经济上划不划算？前四章回答了第一个问题，本章转向后两个问题——'
        '设计"母液预配+在线稀释"的现场施工工艺方案，以四川盆地典型页岩气水平井为算例进行'
        '材料用量和经济性估算，并评估荧光示踪方案的环保安全性。需要指出，本章的分析以实验室'
        '确定的最优参数为前提，工程实施前需通过现场先导性试验验证。'
    )
    for run in p_ch5_open.runs:
        run.text = ''
    if p_ch5_open.runs:
        p_ch5_open.runs[0].text = new_open
    print('  Enhanced Ch5 opening')


# ================================================================
# SAVE
# ================================================================
doc.save(DOC)
print(f'\nAll inter-chapter bridges saved to {DOC}')