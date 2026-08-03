"""Align body text with the revised Introduction narrative."""
from docx import Document
from docx.shared import Pt

INPUT = r"c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP-T_投稿文件\ESP-T_Final_4-revised.docx"
doc = Document(INPUT)

# === PARA 49: Fix Section 3.2 ending ===
# Old: "These kinetic parameters provide the release-side boundary condition..."
# Problem: K-P parameters are NOT boundary conditions for the ADE model.
# Fix: They characterize release independently; later corroborated by Pe.

p49 = doc.paragraphs[49]
old_text = p49.text
new_text = old_text.replace(
    "These kinetic parameters provide the release-side boundary condition for the transport model developed in Section 3.3.",
    "These kinetic parameters characterize the release mechanism independently of any flow configuration. In Section 3.3, the non-Fickian transport regime identified here will be corroborated by the Peclet number independently obtained from breakthrough curve analysis."
)
if old_text != new_text:
    for r in list(p49.runs):
        r._element.getparent().remove(r._element)
    for child in list(p49._element):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('r', 'drawing', 'pict'):
            p49._element.remove(child)
    p49.add_run(new_text)
    print("[49] Section 3.2 ending: FIXED")
else:
    print("[49] No change needed")

# === PARA 104: Rewrite "Taken together..." summary ===
# Old: Grandiose language ("reverses this logic", "shift from descriptive to predictive")
# New: Factual summary matching Introduction's four validation lines

NEW_SUMMARY = (
    "Taken together, the results of this section demonstrate that the two-component model "
    "recovers physically meaningful parameters from a single breakthrough curve. The "
    "dual-component structure is a statistical necessity (ΔAICc = 32.7 relative to "
    "the best single-component alternative); the fitted flow rate agrees with the independent "
    "pump setting without being constrained in the objective function; the Peclet number "
    "independently corroborates the non-Fickian transport regime identified via K-P kinetics "
    "(Section 3.2); and the ~47% erfc tail contribution is stable under variation of the "
    "transition-width parameter. The sustained-release fraction has a practical implication: "
    "nearly half of the tracer detected at the wellhead originates from ongoing "
    "matrix-diffusion-controlled release rather than from the initial shut-in accumulation "
    "slug, suggesting that a single shut-in can support monitoring over an extended "
    "production period."
)

p104 = doc.paragraphs[104]
for r in list(p104.runs):
    r._element.getparent().remove(r._element)
for child in list(p104._element):
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag in ('r', 'drawing', 'pict'):
        p104._element.remove(child)
p104.add_run(NEW_SUMMARY)
print("[104] 'Taken together...' summary: REPLACED")

# === CHECK FOR OTHER INCONSISTENCIES ===
# Search for phrases that contradict the joint-estimation framing
problematic = [
    'inverse problem',
    'bridging the gap',
    'reverses this logic',
    'shift from descriptive to predictive',
    'coupled inverse',
]
print("\n=== Remaining consistency check ===")
for i, p in enumerate(doc.paragraphs):
    text = p.text.lower()
    for phrase in problematic:
        if phrase in text:
            print(f"  WARNING: [{i}] contains '{phrase}': {p.text[:120]}...")
            break

doc.save(INPUT)
print("\nSaved.")
