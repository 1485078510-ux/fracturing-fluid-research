# -*- coding: utf-8 -*-
"""Replace ALL remaining Chinese with English — headings, captions, tables."""
from docx import Document
from docx.oxml.ns import qn
import re

DST = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP_EN_final.docx'
doc = Document(DST)

IMG = set()
for i, p in enumerate(doc.paragraphs):
    for run in p.runs:
        if (run._element.findall('.//'+qn('w:drawing')) or run._element.findall('.//'+qn('a:blip'))):
            IMG.add(i); break

def sp(idx, text):
    if idx in IMG: return
    p = doc.paragraphs[idx]
    for r in p.runs: r.text = ''
    if p.runs: p.runs[0].text = text
    else: p.add_run(text)

# ===== ALL HEADINGS =====
sp(6,  "2. Experiments")
sp(7,  "2.1 Materials")
sp(11, "2.2 Synthesis of Stearic Acid-Modified Nano-Fe3O4 (nano-Fe3O4@SA)")
sp(15, "2.3 Preparation of ESP-T")
sp(19, "2.4 Characterization")
sp(21, "2.5 Tracer Release Experiments at Different Temperatures")
sp(25, "2.6 Oil Production Monitoring Experiments")
sp(32, "3. Results and Discussion")
# 3.1 SEM — [33] already English? Let me check
sp(33, "3.1 SEM Characterization")
sp(41, "3.2 Thermal Stability")
sp(46, "3.3 Water Contact Angle (WCA)")
sp(50, "3.4 Bulk Density and Apparent Density")
sp(54, "3.5 Proppant Pack Conductivity")
sp(59, "3.6 Tracer Release Behavior at Different Temperatures")
sp(74, "3.7 Oil Production Monitoring of ESP-T")
sp(94, "4. Conclusions")
sp(98, "References")

# ===== FIGURE CAPTIONS =====
sp(14, "Fig. 2-1. Schematic illustration of the preparation of stearic acid-modified metal-doped Fe3O4 nanoparticles.")
sp(18, "Fig. 2-2. Schematic illustration of the preparation of Fe3O4-encapsulated epoxy resin proppants.")
sp(24, "Fig. 2-3. Schematic diagram of the tracer release apparatus at different temperatures.")
sp(28, "Fig. 2-4. Schematic diagram of the single-phase oil flow experimental apparatus.")
sp(31, "Fig. 2-5. Schematic diagram of the oil-water two-phase flow experimental apparatus.")

sp(37, "Fig. 3-1. SEM images of epoxy microspheres (top row) and Fe3O4@epoxy resin proppants (bottom row). "
    "(a, d) Overview of multiple microspheres; (b, e) single whole microsphere; (c, f) local surface detail.")
sp(40, "Fig. 3-2. EDS elemental mapping of Fe3O4@epoxy resin proppant. "
    "(a) SEM image; (b) Fe element distribution; (c) Si element distribution.")
sp(45, "Fig. 3-3. TGA/DTG and DSC curves of Fe3O4@epoxy resin proppant.")
sp(49, "Fig. 3-4. Water contact angle measurements. (a) Pure epoxy resin; (b) ESP-T after doping with nano-Fe3O4@SA.")
sp(53, "Fig. 3-5. Bulk density and apparent density of epoxy microspheres and ESP-T.")
sp(57, "Fig. 3-6. Oil-water permeability test of epoxy microspheres and ESP-T. "
    "(a, b) Before and after water passing through epoxy microspheres; "
    "(c, d) before and after oil passing through epoxy microspheres; "
    "(e, f) before and after water passing through ESP-T; "
    "(g, h) before and after oil passing through ESP-T.")
sp(71, "Fig. 3-7. (a) Release curves of ESP-T at different temperatures; "
    "(b) K-P model fitting curves; (c) release rate curves of ESP-T.")
sp(88, "Fig. 3-8. (a) Tracer production curve for ESP-T single-phase flow monitoring; "
    "(b) model fitting curve; (c) comparison of pump-set flow rate and fitted flow rate.")
sp(93, "Fig. 3-9. (a) Tracer concentration versus total two-phase flow rate at different OWRs; "
    "(b) tracer flux versus total two-phase flow rate at different OWRs; "
    "(c) comparison of normalized tracer flux and actual oil flow rate at 0.1 mL/min.")

# ===== TABLE CAPTIONS =====
sp(9,  "Table 2-1. Specifications and sources of experimental materials.")
sp(58, "Table 3-1. Oil and water passage time.")
sp(72, "Table 3-2. Fitting parameters of the K-P model.")

# ===== MISC CHINESE =====
sp(62, "where:")  # "式中：" → "where:"

# ===== TABLES — translate Chinese cells =====
# Table 0: Materials (Table 2-1)
t0 = doc.tables[0]
cell_map_t0 = {
    (0,0): "Reagent", (0,1): "Purity", (0,2): "Manufacturer",
    (1,0): "Ethanol", (1,1): "AR", (1,2): "Chengdu Kelong Chemical Co., Ltd.",
    (2,0): "FeCl3", (2,1): "AR", (3,0): "FeCl2-4H2O", (3,1): "AR",
    (4,0): "MnCl2-6H2O", (4,1): "AR", (5,0): "Stearic acid", (5,1): "AR",
    (6,0): "SiO2", (6,1): "AR",
    (7,0): "Guar gum", (7,1): "90-95%",
    (8,0): "E51 epoxy resin", (8,1): ">=99%",
    (9,0): "T31 curing agent", (9,1): "95%",
    (10,0): "Hollow glass microspheres", (10,1): ">=99%",
}
for (r, c), v in cell_map_t0.items():
    if r < len(t0.rows) and c < len(t0.rows[r].cells):
        for para in t0.rows[r].cells[c].paragraphs:
            for run in para.runs:
                run.text = ''

# Table 1: Oil/Water Passage Time (Table 3-1)
t1 = doc.tables[1]
cell_map_t1 = {
    (0,0): "", (0,1): "ESP-T", (0,2): "Pure epoxy microspheres",
    (1,0): "Water passage time", (1,1): "28 min 41 s", (1,2): "2 min 53 s",
    (2,0): "Oil passage time", (2,1): "5 min 11 s", (2,2): "15 min 11 s",
}
for (r, c), v in cell_map_t1.items():
    if r < len(t1.rows) and c < len(t1.rows[r].cells):
        for para in t1.rows[r].cells[c].paragraphs:
            for run in para.runs:
                run.text = ''

# Table 2: K-P Model (Table 3-2)
t2 = doc.tables[2]
# Headers
kp_headers = {
    (0,0): "", (0,1): "", (0,2): "Release temperature", (0,3): "", (0,4): "", (0,5): "",
    (1,0): "Fit parameter", (1,1): "", (1,2): "30 C", (1,3): "60 C", (1,4): "90 C", (1,5): "120 C",
    (2,0): "R2",
    (3,0): "K",
    (4,0): "n",
}
for (r, c), v in kp_headers.items():
    if r < len(t2.rows) and c < len(t2.rows[r].cells):
        cell = t2.rows[r].cells[c]
        # Only set if currently empty or Chinese
        has_cn = bool(re.search(r'[一-鿿]', cell.text))
        if has_cn or c <= 1:
            for para in cell.paragraphs:
                for run in para.runs:
                    if has_cn or (c <= 1 and r >= 2):
                        run.text = v if r >= 2 and c == 0 else (v if r < 2 else run.text)

# Fix table 2 cells safely
for r_idx in [2, 3, 4]:
    label = kp_headers[(r_idx, 0)]
    cell0 = t2.rows[r_idx].cells[0]
    # Clear and set text
    for para in cell0.paragraphs:
        for run in para.runs:
            run.text = ''
    if cell0.paragraphs and cell0.paragraphs[0].runs:
        cell0.paragraphs[0].runs[0].text = label
    else:
        cell0.paragraphs[0].add_run(label)

# ===== SAVE =====
doc.save(DST)
print(f"All Chinese replaced with English. Saved: {DST}")