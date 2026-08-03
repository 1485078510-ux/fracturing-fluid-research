#!/usr/bin/env python3
"""Apply all changes: intro rewrite + ADE deepening + new references."""
from docx import Document

try:
    doc = Document('四氧化三铁环氧树脂拟合/ESP-T_final.docx')
except:
    doc = Document('四氧化三铁环氧树脂拟合/ESP-T_final_backup.docx')

# Intro paras
intro_map = {
    11: (
        "Conventional oil-soluble chemical tracers exhibit limited compatibility "
        "with water-based fracturing fluids and cannot sustain long-term monitoring "
        "[5-9]. Tracer proppants—composite particles that immobilize tracer agents "
        "within a solid carrier and are co-injected with proppant-laden fracturing "
        "fluids—have emerged as a promising alternative, eliminating the need for "
        "separate tracer-injection operations. Zhao et al. [10] coated rhodamine 6G "
        "onto polyvinyl alcohol-modified ceramic particles to retard aqueous release; "
        "Zhou et al. [11] encapsulated carbon quantum dots in polymer-coated ceramic "
        "particles; Li et al. [12] employed rare-earth tracers with ammonium "
        "polymethacrylate coatings; and Gong et al. [13] embedded oleophilic Fe3O4 "
        "tracers into polystyrene (PS) microspheres for oil-phase monitoring. Despite "
        "this progress, all of these studies share a critical methodological "
        "limitation: they characterize tracer release using empirical kinetic "
        "models—predominantly the Korsmeyer-Peppas power law—but none establishes a "
        "quantitative, physically grounded linkage between the release kinetics "
        "measured in batch experiments and the resulting tracer breakthrough signal "
        "observed at the wellhead. The advection-dispersion transport of released "
        "tracers from the proppant pack through the production tubing to the sampling "
        "point remains unmodeled. Consequently, per-interval oil production rates "
        "cannot be quantitatively inferred from tracer concentration histories, "
        "regardless of how accurately the release kinetics are characterized. This "
        "transport-to-signal gap represents the central unresolved challenge in "
        "tracer-proppant-based production monitoring."
    ),
    13: (
        "Beyond the transport-modeling gap, material-level constraints further limit "
        "existing tracer proppants. Coated proppants suffer from high particle density "
        "that impedes transport in fracturing fluids, abrupt monitoring cessation upon "
        "coating dissolution, and—in the case of pure PS microspheres—inadequate "
        "thermal stability and mechanical strength under downhole conditions [14-17]. "
        "Epoxy resin, a thermosetting polymer with a highly cross-linked network, "
        "offers outstanding thermal stability (decomposition temperature exceeding "
        "350 degC), exceptional chemical resistance, and tunable mechanical "
        "properties—making it a compelling alternative matrix material [18,19]. "
        "Direct emulsion-polymerization synthesis of epoxy resin microspheres further "
        "overcomes the limitations of coating-based fabrication—enabling low-density "
        "tailoring, nanoparticle modification, and in-situ tracer encapsulation within "
        "a single synthetic step. Li et al. [20] and Wei et al. [21] demonstrated "
        "epoxy-encapsulated water-soluble tracers for aqueous-phase inflow profiling; "
        "however, the use of an epoxy matrix for oleophilic, oil-phase tracer "
        "release—and, critically, its integration with a predictive transport model "
        "capable of translating tracer signals into quantitative per-interval "
        "production allocation—has not been reported."
    ),
    15: (
        "Here we address this dual gap. First, we develop a piecewise "
        "advection-dispersion model with smooth hyperbolic-tangent (tanh) transition "
        "that decomposes the tracer breakthrough curve into a Gaussian pulse "
        "component—representing the advective-dispersive transport of the tracer slug "
        "accumulated during the well shut-in period—and an erfc tailing "
        "component—representing the sustained, matrix-diffusion-controlled release of "
        "residual tracer from the polymer carrier. The model establishes a direct "
        "quantitative relationship between the measured tracer concentration history "
        "and the per-interval oil production rate, enabling physical interpretation "
        "of parameters that would otherwise remain as empirical fitting coefficients. "
        "Second, we validate this model using a purpose-designed oleophilic tracer "
        "proppant (ESP-T) fabricated by encapsulating stearic acid-modified "
        "nano-Fe3O4 (nano-Fe3O4@SA) within an epoxy resin matrix via emulsion "
        "polymerization. We systematically characterize ESP-Ts microstructure, "
        "thermal stability, wettability, oil-water transport selectivity, and "
        "mechanical integrity; quantify its temperature-dependent release kinetics "
        "via the Korsmeyer-Peppas model; and demonstrate the integrated "
        "material-model framework through single-phase and steady-state two-phase "
        "core displacement experiments. This work establishes a coupled "
        "experimental-modeling methodology that translates sustained tracer release "
        "signals into quantitative per-interval production allocation—a broadly "
        "applicable framework for unconventional reservoir management."
    ),
}

for idx, new_text in intro_map.items():
    p = doc.paragraphs[idx]
    for run in p.runs:
        run.text = ""
    p.runs[0].text = new_text
    print(f"Intro [{idx}] rewritten")

# ADE deepening
deep = (
    " The fitted parameters carry several physically significant implications "
    "that extend beyond empirical curve-fitting. First, the Peclet number "
    "Pe = x/alpha = 0.934 places the transport regime precisely at the "
    "transition between advection-dominated and dispersion-dominated behavior. "
    "This regime is characteristic of a gradually releasing source: pure "
    "advection would produce a sharp, temporally compressed Gaussian pulse, "
    "while the comparable dispersion broadens the peak and generates the "
    "extended tail that carries 47% of the integrated signal. The near-unity Pe "
    "thus constitutes independent physical corroboration of the matrix-diffusion-"
    "controlled release mechanism identified in Section 3.6—the two independent "
    "analyses (K-P kinetics and ADE transport) converge on the same physical "
    "picture. Second, the tanh blending function provides a physically motivated, "
    "continuously differentiable transition between the Gaussian rise and erfc "
    "tail regimes. A piecewise step-function transition would introduce a "
    "non-physical discontinuity in the first derivative of the concentration "
    "profile at the crossover—an artifact that the tanh formulation eliminates "
    "while preserving the two-component decomposition central to the model's "
    "interpretability. Third, the close agreement between the pump-set flow rate "
    "(0.50 mL/min) and the independently fitted flow rate (0.46 mL/min, 8% "
    "relative error) demonstrates that the model parameters carry well-defined "
    "physical meaning and can be recovered from the tracer signal alone—an "
    "essential property for field applications where independent flow-rate "
    "measurements may not be available for every interval. Fourth, the 47% "
    "erfc tail contribution has direct practical significance for long-term "
    "monitoring: it quantitatively confirms that nearly half of the total "
    "detectable tracer signal originates from sustained matrix-diffusion-"
    "controlled release rather than from the initial concentration pulse "
    "generated during shut-in. This establishes that ESP-T-based production "
    "monitoring is fundamentally a diffusion-governed, extended-duration "
    "process—a conclusion inaccessible from empirical release fitting alone, "
    "and one that directly informs the design of sampling protocols and "
    "shut-in durations in field deployments."
)

for i, p in enumerate(doc.paragraphs):
    if "establishes the physical basis for using ESP-T as a long-term production monitor" in p.text:
        enhanced = p.text.rstrip() + deep
        for run in p.runs:
            run.text = ""
        p.runs[0].text = enhanced
        print(f"ADE deepened at [{i}]")
        break

# References
NEW_REFS = [
    "[1]\tIEA. World Energy Outlook 2024 [R]. Paris: International Energy Agency, 2024.",
    "[2]\tMONTGOMERY S L, JARVIE D M, BOWKER K A, et al. Mississippian Barnett Shale, Fort Worth basin, north-central Texas: Gas-shale play with multi-trillion cubic foot potential [J]. AAPG Bulletin, 2005, 89(2): 155-175.",
    "[3]\tBARATI R, LIANG J T. A review of fracturing fluid systems used for hydraulic fracturing of oil and gas wells [J]. Journal of Applied Polymer Science, 2014, 131(16): 40735.",
    "[4]\tMEDEIROS F, KURTOGLU B, OZKAN E, et al. Analysis of production data from hydraulically fractured horizontal wells in shale reservoirs [J]. SPE Reservoir Evaluation & Engineering, 2010, 13(3): 559-568.",
    "[5]\tPATIDAR A K, JOSHI D, DRISTANT U, et al. A review of tracer testing techniques in porous media specially attributed to the oil and gas industry [J]. Journal of Petroleum Exploration and Production Technology, 2022, 12(12): 3339-3356.",
    "[6]\tSANNI M, AL-ABBAD M, KOKAL S, et al. Pushing the envelope of residual oil measurement: A field case study of a new class of inter-well chemical tracers [J]. Journal of Petroleum Science and Engineering, 2018, 163: 538-545.",
    "[7]\tSILVA M, STRAY H, BJORNSTAD T. Stability assessment of PITT tracer candidate compounds - The case of pyrazines [J]. Journal of Petroleum Science and Engineering, 2019, 182: 106269.",
    "[8]\tWATKINS J W, MARDOCK E S. Use of radioactive iodine as a tracer in water-flooding operations [J]. Journal of Petroleum Technology, 1954, 6(9): 117-124.",
    "[9]\tYANG H, GUO K, LIN L, et al. Application of micro-substance tracer test in fractured horizontal wells [J]. Journal of Petroleum Exploration and Production Technology, 2024, 14(5): 1235-1246.",
    "[10]\tZHAO B, PANTHI K, MOHANTY K K. Tracer eluting proppants for hydraulic fracture characterization [J]. Journal of Petroleum Science and Engineering, 2020, 190: 107048.",
    "[11]\tZHOU Y, LIU H, GAO J, et al. Coated proppants with self-suspension and tracer slow-release functions [J]. Journal of Petroleum Science and Engineering, 2022, 208: 109645.",
    "[12]\tLI N, CHENG Q, GONG Z, et al. Release kinetics of rare earth tracer from polymer-coated proppants for hydraulic fracture analysis [J]. Geoenergy Science and Engineering, 2023, 227: 211782.",
    "[13]\tGONG Z, LI N, KANG W, et al. Novel oleophilic tracer-slow-released proppant for monitoring the oil production contribution [J]. Fuel, 2024, 364: 130945.",
    "[14]\tWANG G, MA Q, REN L, et al. A comprehensive review of multifunctional proppants [J]. ACS Omega, 2024, 9(44): 44120-44133.",
    "[15]\tKRISHNAN M R, LI W, ALHARBI B, et al. In-situ high-strength poly(styrene-methyl methacrylate)-2D nanofiller composite microbeads as potential proppants in hydraulic fracturing [J]. Geoenergy Science and Engineering, 2025, 257: 214195.",
    "[16]\tGUO X, WEI K, NI T, et al. Preparation and performance analysis of polyethylene glycol/epoxy resin composite phase change material [J]. Journal of Energy Storage, 2024, 88: 111525.",
    "[17]\tLIANG C, LUO W, YAN C, et al. Ultra-lightweight proppant synthesized from PMMA/pine bark composite: Low-cost material and outstanding properties [J]. Chemistry Letters, 2016, 45(8): 994-996.",
    "[18]\tZOVEIDAVIANPOOR M, GHARIBI A, BIN JAAFAR M Z. Experimental characterization of a new high-strength ultra-lightweight composite proppant derived from renewable resources [J]. Journal of Petroleum Science and Engineering, 2018, 170: 1038-1047.",
    "[19]\tSABINS F, APBLETT A, SHAFER R, et al. Epoxy resin exhibits long-term durability and chemical stability as a well sealant [C]. SPE-204374-MS, SPE International Conference on Oilfield Chemistry, The Woodlands, Texas, 2021.",
    "[20]\tLI H, LIU Z, LI Y, et al. Evaluation of the release mechanism of sustained-release tracers and its application in horizontal well inflow profile monitoring [J]. ACS Omega, 2021, 6(29): 19269-19280.",
    "[21]\tWEI M, WANG Y, DUAN Y, et al. Screening and performance evaluation of epoxy resin long-term sustained-release solid tracer [J]. International Journal of Oil, Gas and Coal Technology, 2024, 36(2): 170-196.",
    "[22]\tGONG Z, LI N, QIN M, et al. Magnetic nano-Fe3O4-based oleophilic tracer for stability studies of nano-tracer in oilfields condition [J]. Colloids and Surfaces A: Physicochemical and Engineering Aspects, 2024, 683: 133085.",
    "[23]\tWANG L, TIAN Y, YU X, et al. Advances in improved/enhanced oil recovery technologies for tight and shale reservoirs [J]. Fuel, 2017, 210: 425-445.",
    "[24]\tMOHR S H, WANG J, ELLEM G, et al. Projection of world fossil fuels by country [J]. Fuel, 2015, 141: 120-135.",
    "[25]\tCHONG Z R, YANG S H B, BABU P, et al. Review of natural gas hydrates as an energy resource: Prospects and challenges [J]. Applied Energy, 2016, 162: 1633-1652.",
    "[26]\tRITGER P L, PEPPAS N A. A simple equation for description of solute release I. Fickian and non-Fickian release from non-swellable devices in the form of slabs, spheres, cylinders or discs [J]. Journal of Controlled Release, 1987, 5(1): 23-36.",
    "[27]\tPEPPAS N A, SAHLIN J J. A simple equation for the description of solute release. III. Coupling of diffusion and relaxation [J]. International Journal of Pharmaceutics, 1989, 57(2): 169-172.",
    "[28]\tVAN GENUCHTEN M T, ALVES W J. Analytical solutions of the one-dimensional convective-dispersive solute transport equation [R]. USDA Technical Bulletin No. 1661, U.S. Department of Agriculture, 1982.",
    "[29]\tKRUMBEIN W C, SLOSS L L. Stratigraphy and Sedimentation [M]. 2nd ed. San Francisco: W.H. Freeman and Company, 1963.",
]

ref_start = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "References":
        ref_start = i
        break

if ref_start:
    for i in range(ref_start + 1, len(doc.paragraphs)):
        for run in doc.paragraphs[i].runs:
            run.text = ""
    for j, ref_text in enumerate(NEW_REFS):
        ti = ref_start + 1 + j
        if ti < len(doc.paragraphs):
            p = doc.paragraphs[ti]
            if len(p.runs) == 0:
                p.add_run("")
            for run in p.runs:
                run.text = ""
            p.runs[0].text = ref_text
    # Clean trailing trash
    for i in range(ref_start + 1 + len(NEW_REFS), len(doc.paragraphs)):
        t = doc.paragraphs[i].text.strip()
        if t and (t[0] == "[" or "Reagents" in t or "ESP-T" in t or "Water passage" in t):
            for run in doc.paragraphs[i].runs:
                run.text = ""
    print(f"References: {len(NEW_REFS)} entries")
else:
    print("WARNING: References heading not found!")

# Save
out = "四氧化三铁环氧树脂拟合/ESP-T_FINAL.docx"
doc.save(out)
print(f"Saved: {out}")