# -*- coding: utf-8 -*-
"""Apply content expansions to the polished DOCX with yellow highlighting for new text."""
import json
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"c:\Users\郝\Desktop\claude\向晶\页岩气藏地应力演化与裂缝扩展耦合机制研究进展与展望-润色版.docx"
DST = r"c:\Users\郝\Desktop\claude\向晶\页岩气藏地应力演化与裂缝扩展耦合机制研究进展与展望-润色版-最终版.docx"
EXP_PATH = r"c:\Users\郝\Desktop\claude\向晶\expansions.json"

def add_highlight(run, color="yellow"):
    rPr = run._element.get_or_add_rPr()
    h = OxmlElement('w:highlight')
    h.set(qn('w:val'), color)
    rPr.append(h)

def get_font_info(para):
    if para.runs:
        r = para.runs[0]
        return {'size': r.font.size, 'bold': r.font.bold, 'name': r.font.name}
    return {'size': Pt(12), 'bold': None, 'name': '宋体'}

def apply_font(run, fi):
    if fi['size']: run.font.size = fi['size']
    if fi['bold'] is not None: run.font.bold = fi['bold']
    if fi['name']: run.font.name = fi['name']

with open(EXP_PATH, 'r', encoding='utf-8') as f:
    expansions = json.load(f)

doc = Document(SRC)
count_replace = 0
count_append = 0

# Build a set of keywords to skip already-processed paragraphs
used_keywords = set()

for para in doc.paragraphs:
    text = para.text
    if not text.strip():
        continue

    # Find matching expansion
    for exp in expansions:
        kw = exp['keyword']
        if kw in used_keywords:
            continue
        if kw in text:
            used_keywords.add(kw)
            fi = get_font_info(para)
            action = exp['action']
            new_text = exp['new_text']

            if action == 'replace':
                # Replace entire paragraph with new text, highlight all
                for run in list(para.runs):
                    run._element.getparent().remove(run._element)
                new_run = para.add_run(new_text)
                apply_font(new_run, fi)
                add_highlight(new_run, "yellow")
                count_replace += 1
                print(f'  REPLACE: {exp["note"][:60]}')

            elif action == 'append':
                # Append new text as highlighted run
                spacer = para.add_run('\n')
                spacer.font.size = Pt(4)
                new_run = para.add_run(new_text)
                apply_font(new_run, fi)
                add_highlight(new_run, "yellow")
                count_append += 1
                print(f'  APPEND: {exp["note"][:60]}')

            break  # Move to next paragraph

doc.save(DST)
print(f'\nDone! {count_replace} replacements, {count_append} appends.')
print(f'Saved to: {DST}')
