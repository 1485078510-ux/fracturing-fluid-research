# -*- coding: utf-8 -*-
"""Final CET-4 prep doc: 1 primary topic + 3 backups."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor

doc = Document()
style = doc.styles['Normal']
style.font.size = Pt(11.5)
style.paragraph_format.line_spacing = 1.45

BLUE = (0x1A, 0x56, 0xDB)
RED = (0xCC, 0x33, 0x33)
GRAY = (0x6B, 0x72, 0x80)
GREEN = (0x1F, 0x88, 0x54)

def AP(text, bold=False, size=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)

def CODE(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(1)

def SEP():
    p = doc.add_paragraph()
    p.add_run("─" * 55).font.size = Pt(6)

# ================================================================
AP("2026.06 CET-4 作文最后一背", bold=True, size=18, color=BLUE)
AP("主打方向：数字生活与心理健康  |  备胎：AI、乡村振兴、传统文化", size=10, color=GRAY)
SEP()

# ================================================================
AP("先背框架（三段骨架，任何题目照抄）", bold=True, size=14, color=RED)
SEP()

AP("第一段（3句）", bold=True, size=12, color=BLUE)
CODE("In recent years, the issue of [话题] has sparked considerable")
CODE("discussion, particularly regarding [具体方面]. While this trend")
CODE("brings notable benefits, it also raises concerns that deserve")
CODE("serious attention.")

AP("第二段（5-6句）← 塞三样东西：例子 + 让步反驳 + 长句", bold=True, size=12, color=BLUE)
CODE("Several factors explain this phenomenon. At its root, [根本原因]")
CODE("has created fertile ground for it to take hold. The effects are")
CODE("already visible: [具体影响A], and more importantly, [更深影响B].")
CODE("What deserves equal attention is [被忽视的问题] -- a point often")
CODE("neglected in casual discussion.")
CODE("")
CODE("  [此处塞：具体例子, 2-3句]")
CODE("  [此处塞：让步反驳, 2句]")
CODE("  [此处塞：复合长句, 1句]")

AP("第三段（2-3句）", bold=True, size=12, color=BLUE)
CODE("In light of the above, I am convinced that [重申观点] is not only")
CODE("desirable but achievable -- provided that [关键条件] is taken")
CODE("seriously. The path forward calls for a balance between [A] and [B],")
CODE("a challenge today's students are ready to meet.")

doc.add_page_break()

# ================================================================
# HOW TO WRITE EXAMPLES
# ================================================================
AP("实例怎么写（20秒速成，考场现编也像真的）", bold=True, size=15, color=RED)
AP("阅卷老师看例子只看三点：有人名/身份 + 有细节 + 有结果。满足就加分。", size=10, color=GRAY)
SEP()

AP("万能公式：", bold=True, size=13, color=BLUE)
CODE("A [身份, 如classmate/roommate/friend/senior] of mine, who once")
CODE("[过去的困难或状态], now [做法/改变]. [1个具体细节 -- 数字/")
CODE("场景/引用原话]. [结果]. [一句感悟收尾].")

AP("不是真事不要紧，编得像真的就行。像真话的秘诀：加一个具体细节。", size=10, color=GRAY)
SEP()

AP("对比一下你就懂了 -- ", bold=True, size=12, color=BLUE)

AP("假大空写法（阅卷老师直接跳过）：", bold=True, color=RED)
CODE("Many students spend too much time on social media. This is bad for")
CODE("their health. We should reduce our screen time.              ← 0分例子")

AP("有血有肉写法（阅卷老师眼睛一亮）：", bold=True, color=GREEN)
CODE("A roommate of mine used to scroll through short videos for hours")
CODE("before sleep, telling himself it was just a way to unwind. Within")
CODE("a semester, his grades slipped and he confided that he felt")
CODE("strangely empty -- constantly stimulated yet never truly rested.")
CODE("                                               ← 14分例子")

AP("差在哪：第一个只有观点没有画面，第二个有场景有时间线有原话。", size=10, color=GRAY)
SEP()

AP("考场30秒速编公式：套一个你认识的人 + 一个谁都有过的经历", bold=True, size=12, color=BLUE)
AP("步骤：", bold=True)
CODE("1. 定身份：classmate / roommate / friend / senior / cousin")
CODE("2. 定过去困境：used to [沉迷某事] / once struggled with [困难]")
CODE("3. 定转折点：within [时间] / after [事件] / he realized that...")
CODE("4. 定改变：has since [新做法] / replaced A with B / set a limit on...")
CODE("5. 加原话（最能骗到分的一步）：he told me / he confided that /")
CODE("   in his words / as he puts it / he says ...")

AP("直接引语加进去，整个例子立刻像真的。阅卷老师默认你不会编原话，所以看到引号就信。", size=10, color=GRAY)
SEP()

AP("四个话题，每个给你：方框的多个填法（选最顺手的）→ 填完长什么样。", bold=True, size=12, color=BLUE)
SEP()

# ================================================================
# TOPIC 1: DIGITAL
# ================================================================
AP("【数字生活】身份=室友/自己, 困境=刷手机上瘾", bold=True, size=13, color=RED)
SEP()

AP("每个方框的可选填法（挑最顺手的用，别全背）：", bold=True, size=11, color=BLUE)
CODE("")
CODE("[某平台]       选一个:")
CODE("  Douyin / short video apps / social media feeds / Weibo / Bilibili")
CODE("")
CODE("[时间长度]     选一个:")
CODE("  a single semester / just a few months / within weeks / by midterm")
CODE("")
CODE("[具体后果]     选一个:")
CODE("  his grades began to slip / he couldn't focus during lectures")
CODE("  his sleep quality collapsed / he felt anxious when offline")
CODE("  he realized he hadn't finished a single book all semester")
CODE("")
CODE("[改变行动]     选一个:")
CODE("  set a daily one-hour screen limit / deleted the app from his phone")
CODE("  replaced pre-sleep scrolling with reading / turned off all notifications")
CODE("  started leaving his phone outside the bedroom at night")
CODE("")
CODE("[一句原话]     选一个:")
CODE("  'I didn't even realize I was addicted.'")
CODE("  'I was constantly stimulated yet never truly rested.'")
CODE("  'The algorithm knew what I wanted before I did.'")
CODE("  'I had forgotten what it felt like to be bored -- and that was the problem.'")
CODE("")
CODE("[你的感悟]     选一个:")
CODE("  the most draining habits are the ones we never stop to notice")
CODE("  what we consume online ends up consuming us in return")
CODE("  real rest doesn't come from a screen")

AP("填完长这样（6句 = 身份 + 过去 + 后果 + 改变 + 原话 + 感悟）：", bold=True, size=11, color=GREEN)
CODE("A roommate of mine used to spend hours scrolling through Douyin")
CODE("before sleep, telling himself it was just a way to unwind. Within a")
CODE("single semester, his grades began to slip and he found himself")
CODE("unable to focus during lectures. He has since set a daily screen")
CODE("limit and replaced late-night scrolling with reading. 'I was")
CODE("constantly stimulated yet never truly rested,' he confided. His")
CODE("experience taught me that the most draining habits are often the")
CODE("ones we never stop to notice.")

doc.add_page_break()

# ================================================================
# TOPIC 2: AI
# ================================================================
AP("【AI学习】身份=同学, 困境=写作/学习困难", bold=True, size=13, color=RED)
SEP()

AP("每个方框的可选填法：", bold=True, size=11, color=BLUE)
CODE("")
CODE("[困难领域]     选一个:")
CODE("  academic writing / English composition / math problem sets")
CODE("  preparing for exams / organizing lecture notes / debugging code")
CODE("")
CODE("[AI工具]       选一个:")
CODE("  ChatGPT / an AI grammar checker / an AI study assistant")
CODE("  an intelligent tutoring app / a large language model")
CODE("")
CODE("[用来做什么]   选一个:")
CODE("  polish her essay drafts / generate study outlines")
CODE("  explain difficult concepts in plain language")
CODE("  check her grammar and suggest improvements")
CODE("")
CODE("[克制做法]     选一个:")
CODE("  always rewrites every suggestion in her own words")
CODE("  double-checks every factual claim the AI makes")
CODE("  treats the AI as a starting point, never the final answer")
CODE("  sets a rule: AI can suggest, but she must decide")
CODE("")
CODE("[一句原话]     选一个:")
CODE("  'AI gives me a starting point, not the final answer.'")
CODE("  'The tool is fast, but the thinking still has to be mine.'")
CODE("  'I use it like a tutor -- it explains, but I still have to learn.'")
CODE("  'Without the rewrites, I wouldn't actually be improving.'")
CODE("")
CODE("[你的感悟]     选一个:")
CODE("  a tool is only as wise as the person using it")
CODE("  technology amplifies effort; it doesn't replace it")
CODE("  the line between assistance and dependence is thinner than we think")

AP("填完长这样：", bold=True, size=11, color=GREEN)
CODE("A classmate of mine once struggled with academic writing. She started")
CODE("using an AI grammar checker to polish her essay drafts, yet she")
CODE("always rewrites every suggestion in her own words -- treating the")
CODE("tool as a brainstorming partner rather than a shortcut. 'The tool")
CODE("is fast, but the thinking still has to be mine,' she says. Her")
CODE("writing has actually improved, proving that technology amplifies")
CODE("effort rather than replacing it.")

doc.add_page_break()

# ================================================================
# TOPIC 3: RURAL
# ================================================================
AP("【乡村振兴】身份=学长/亲戚, 困境=留城vs返乡", bold=True, size=13, color=RED)
SEP()

AP("每个方框的可选填法：", bold=True, size=11, color=BLUE)
CODE("")
CODE("[大城市工作]   选一个:")
CODE("  a corporate job in Beijing / a marketing position in Shanghai")
CODE("  a well-paying tech job in Shenzhen / a bank internship in Guangzhou")
CODE("")
CODE("[回乡地点]     选一个:")
CODE("  his hometown in Yunnan / his village in Guizhou")
CODE("  his family's farm in Sichuan / a small town in Jiangxi")
CODE("")
CODE("[具体做法]     选一个:")
CODE("  helps local farmers sell produce through livestreaming")
CODE("  started an e-commerce business for local specialty products")
CODE("  opened a rural guesthouse showcasing local culture")
CODE("  teaches villagers how to use Douyin to market their crafts")
CODE("")
CODE("[具体结果]     选一个:")
CODE("  his monthly income now exceeds that of his urban classmates")
CODE("  his products now ship to customers in over 20 provinces")
CODE("  the village's collective income has tripled since he returned")
CODE("")
CODE("[一句原话]     选一个:")
CODE("  'I'm building something that actually belongs to me.'")
CODE("  'I wake up every day knowing my work matters to real people.'")
CODE("  'The city had a thousand people like me. The village needed me.'")
CODE("  'I came back to help my hometown; it ended up giving me a career.'")
CODE("")
CODE("[你的感悟]     选一个:")
CODE("  success doesn't have a single postal code")
CODE("  the countryside is not a fallback -- it's a new frontier")
CODE("  meaningful work is not about where you are, but what you build")

AP("填完长这样：", bold=True, size=11, color=GREEN)
CODE("A senior I know faced a choice between a corporate job in Beijing")
CODE("and returning to his hometown in Yunnan. He chose the latter and")
CODE("now helps local tea farmers sell their products through")
CODE("livestreaming. His monthly income now exceeds that of many of his")
CODE("urban classmates. 'The city had a thousand people like me. The")
CODE("village actually needed me,' he told me last month. His story made")
CODE("me realize that meaningful work is not about where you are, but")
CODE("what you build.")

doc.add_page_break()

# ================================================================
# TOPIC 4: CULTURE
# ================================================================
AP("【传统文化】身份=社团/自己, 困境=年轻人不了解传统", bold=True, size=13, color=RED)
SEP()

AP("每个方框的可选填法：", bold=True, size=11, color=BLUE)
CODE("")
CODE("[做的事]       选一个:")
CODE("  makes short videos about traditional festivals")
CODE("  teaches calligraphy to elementary school children")
CODE("  organizes campus hanfu exhibitions and photo shoots")
CODE("  runs a WeChat account sharing folk stories from different regions")
CODE("")
CODE("[某节日/某活动] 选一个:")
CODE("  the Dragon Boat Festival / the Mid-Autumn Festival")
CODE("  the Spring Festival / the Lantern Festival")
CODE("  a campus Cultural Heritage Week event")
CODE("")
CODE("[具体做了什么] 选一个:")
CODE("  filmed the entire process of making zongzi from scratch")
CODE("  taught a group of kids to write Spring Festival couplets")
CODE("  recreated historical costumes and explained their cultural meaning")
CODE("  organized a paper-cutting workshop open to all students")
CODE("")
CODE("[好结果]       选一个:")
CODE("  the video received over 50,000 views overnight")
CODE("  over 200 students showed up to our one-day workshop")
CODE("  the kids asked when we were coming back next")
CODE("")
CODE("[一句评论]     选一个:")
CODE("  'I finally understand what this festival actually means.'")
CODE("  'I had no idea this was so interesting -- why did no one show us before?'")
CODE("  'This made me want to learn more about my own culture.'")
CODE("")
CODE("[你的感悟]     选一个:")
CODE("  tradition only feels distant when no one bothers to tell its story")
CODE("  culture doesn't survive by being preserved -- it survives by being lived")
CODE("  the best way to protect heritage is to make it feel alive today")

AP("填完长这样：", bold=True, size=11, color=GREEN)
CODE("Last semester, I joined a campus club that makes short videos about")
CODE("traditional festivals. For the Dragon Boat Festival, we filmed the")
CODE("entire process of making zongzi -- from soaking bamboo leaves to")
CODE("the final unwrapping. To our surprise, the video received over")
CODE("50,000 views. 'I finally understand what this festival actually")
CODE("means,' one viewer commented. That moment made me realize that")
CODE("tradition only feels distant when no one bothers to tell its story.")
SEP()

AP("一句话总结：每个方框挑1个填进去。填完就是6句话 = 一个完整的例子。原话那句必须留。", bold=True, size=12, color=RED)

doc.add_page_break()

# ================================================================
# PRIMARY TOPIC
# ================================================================
AP("★★★ 主打方向：数字生活与心理健康 ★★★", bold=True, size=15, color=RED)
AP("可能命题：社交媒体对大学生的影响 / 如何在数字时代保持专注 / 信息过载与心理压力", size=10, color=GRAY)
SEP()

AP("话题关键词（轮换用，展示词汇量）", bold=True, size=12, color=GREEN)
CODE("mental well-being  |  information overload  |  screen time")
CODE("digital detox  |  attention fragmentation  |  shallow engagement")
CODE("a constant state of distraction  |  genuine human connection")

AP("万能例子（必背！）", bold=True, size=12, color=GREEN)
CODE("A roommate of mine used to scroll through short videos for hours")
CODE("before sleep, telling himself it was just a way to unwind. Within")
CODE("a semester, his grades slipped and he confided that he felt")
CODE("strangely empty -- constantly stimulated yet never truly rested.")
CODE("He has since set a daily screen-time limit and replaced his pre-")
CODE("sleep scrolling with reading. The change, he says, has been")
CODE("nothing short of transformative. His experience taught me that the")
CODE("most addictive forms of digital consumption are often the ones that")
CODE("leave us the most drained.")

AP("让步反驳句", bold=True, size=12, color=GREEN)
CODE("Admittedly, social media and short videos do offer genuine value --")
CODE("they inform, entertain, and connect us. Yet this concern, legitimate")
CODE("as it is, overlooks a crucial point: the problem is not these")
CODE("platforms themselves, but the algorithms designed to maximize")
CODE("engagement at the expense of our attention and well-being.")

AP("复合长句", bold=True, size=12, color=GREEN)
CODE("What makes this issue particularly insidious is that the erosion of")
CODE("attention happens so gradually that most students do not notice it")
CODE("until they find themselves unable to read a full chapter or sit")
CODE("through a conversation without reaching for their phone -- a quiet")
CODE("loss that no exam score can measure but that shapes the very quality")
CODE("of a life.")

SEP()

AP("完整成文 (约150词) — 三件套位置已标注", bold=True, size=12, color=GREEN)

CODE("In recent years, the impact of social media on students' mental")
CODE("well-being has sparked considerable discussion. While these")
CODE("platforms bring notable benefits, they also raise concerns.")
CODE("")
CODE("At its root, attention-driven algorithms have allowed shallow")
CODE("scrolling to replace deep focus.")
CODE("")
CODE(">>> [例子塞这里, 3-4句] <<<")
CODE("A roommate of mine used to scroll Douyin for hours before sleep.")
CODE("Within a semester, his grades slipped and he confided that he felt")
CODE("strangely empty -- constantly stimulated yet never truly rested. He")
CODE("has since set a daily screen limit; the change, he says, has been")
CODE("transformative.")
CODE("")
CODE(">>> [让步反驳塞这里, 2句] <<<")
CODE("Admittedly, social media offers value in keeping us connected. Yet")
CODE("the real issue is not the platforms but the algorithms that hijack")
CODE("our attention.")
CODE("")
CODE("In light of the above, reclaiming our focus is both urgent and")
CODE("achievable -- provided we treat screen time as a health issue, not")
CODE("merely a habit. The path forward calls for balance between staying")
CODE("connected and staying focused, a challenge today's students are")
CODE("ready to meet.")

doc.add_page_break()

# ================================================================
# BACKUP 1: AI
# ================================================================
AP("★★☆ 备胎1：AI与学习", bold=True, size=14, color=BLUE)
AP("背：例子 + 话题词。完整成文略，用框架自己填。", size=10, color=GRAY)
SEP()

CODE("话题词：artificial intelligence / AI-powered tools / digital literacy")
CODE("          independent thinking / academic integrity / double-edged sword")

AP("例子", bold=True, size=12, color=GREEN)
CODE("A classmate of mine, who once struggled with academic writing, now")
CODE("uses AI grammar checkers to polish her drafts. Yet she makes a point")
CODE("of rewriting every AI-generated suggestion in her own words --")
CODE("treating the tool as a brainstorming partner rather than a shortcut.")

AP("让步反驳", bold=True, size=12, color=GREEN)
CODE("Admittedly, over-reliance on AI could erode independent thinking.")
CODE("Yet this overlooks a crucial point: the real threat is not AI itself,")
CODE("but the absence of guidance on how to use it wisely.")

doc.add_page_break()

# ================================================================
# BACKUP 2: Rural
# ================================================================
AP("★★☆ 备胎2：乡村振兴", bold=True, size=14, color=BLUE)
AP("背：例子 + 话题词。", size=10, color=GRAY)
SEP()

CODE("话题词：rural revitalization / entrepreneurship / digital skills")
CODE("          bridge the urban-rural gap / a two-way journey / livestreaming")

AP("例子", bold=True, size=12, color=GREEN)
CODE("A senior I know turned down a corporate job in Beijing to return to")
CODE("his hometown, where he now helps local farmers sell tea through")
CODE("livestreaming. His income exceeds that of many urban classmates --")
CODE("but more importantly, he feels he is building something that belongs")
CODE("to him.")

AP("让步反驳", bold=True, size=12, color=GREEN)
CODE("Admittedly, returning to the countryside means giving up urban")
CODE("conveniences. Yet this overlooks a deeper shift: with digital")
CODE("infrastructure reaching remote villages, the opportunity gap is")
CODE("narrowing faster than most people realize.")

doc.add_page_break()

# ================================================================
# BACKUP 3: Culture
# ================================================================
AP("★☆☆ 备胎3：传统文化传承", bold=True, size=14, color=BLUE)
AP("背：例子 + 话题词。", size=10, color=GRAY)
SEP()

CODE("话题词：cultural heritage / cultural confidence / intangible heritage")
CODE("          breathe new life into / bridge tradition and modernity")

AP("例子", bold=True, size=12, color=GREEN)
CODE("Last semester, I joined a campus club that makes short videos about")
CODE("traditional festivals. For the Dragon Boat Festival, we filmed the")
CODE("making of zongzi from start to finish. The video received over 50,000")
CODE("views, with young viewers commenting that they finally understood")
CODE("what the festival was really about beyond a day off.")

AP("让步反驳", bold=True, size=12, color=GREEN)
CODE("Admittedly, traditional customs can feel distant from modern life.")
CODE("Yet this overlooks the real issue: tradition has rarely been presented")
CODE("in a language young people actually speak -- and that is exactly")
CODE("where digital tools can make all the difference.")

doc.add_page_break()

# ================================================================
# TRANSLATION VOCAB
# ================================================================
AP("翻译必备词汇短语（文化类为主，大概率考）", bold=True, size=15, color=RED)
AP("作文考社会/科技 → 翻译就考文化。按互斥规律，这次翻译大概率是传统文化方向。", size=10, color=GRAY)
SEP()

AP("翻译万能句式（汉译英三段骨架，任何文化题都能套）", bold=True, size=13, color=BLUE)
CODE("句式1  [A] is a [category] that [定义], dating back to [朝代].")
CODE("       例: Paper-cutting is a traditional folk art that originated")
CODE("           in China, dating back to the Han Dynasty.")
CODE("")
CODE("句式2  In recent years, [A] has gained growing recognition, not only")
CODE("        in China but also across the world.")
CODE("       例: In recent years, Tai Chi has gained growing recognition,")
CODE("           not only in China but also across the world.")
CODE("")
CODE("句式3  [A] not only represents [文化意义], but also serves as a")
CODE("        bridge between [传统] and [现代].")
CODE("       例: Calligraphy not only represents the beauty of Chinese")
CODE("           characters, but also serves as a bridge between ancient")
CODE("           wisdom and modern aesthetics.")
SEP()

AP("翻译核心词（按必考频率排序）", bold=True, size=13, color=RED)

AP("朝代时间（必考，翻译里一定出现）", bold=True, size=12, color=BLUE)
CODE("Han Dynasty 汉朝     Tang Dynasty 唐朝     Song Dynasty 宋朝")
CODE("Ming Dynasty 明朝     Qing Dynasty 清朝")
CODE("date back to [时间] 追溯到     over 2,000 years 两千多年")
CODE("ancient times 古代     have a history of... 有...的历史")

AP("文化符号（最可能考的一类）", bold=True, size=12, color=BLUE)
CODE("paper-cutting 剪纸     embroidery 刺绣     shadow puppetry 皮影戏")
CODE("calligraphy 书法     porcelain / ceramics 陶瓷")
CODE("traditional Chinese painting 国画     seal carving 篆刻")
CODE("folk art 民间艺术     intangible cultural heritage 非物质文化遗产")
CODE("cultural relic 文物     time-honored 历史悠久的")

AP("传统建筑/场所", bold=True, size=12, color=BLUE)
CODE("temple 寺庙     palace 宫殿     courtyard 庭院")
CODE("the Forbidden City 故宫     the Great Wall 长城")
CODE("the Silk Road 丝绸之路     ancient capital 古都")
CODE("UNESCO World Heritage Site 联合国教科文组织世界遗产")

AP("传统节日/习俗", bold=True, size=12, color=BLUE)
CODE("the Spring Festival 春节     the Mid-Autumn Festival 中秋节")
CODE("the Dragon Boat Festival 端午节     the Lantern Festival 元宵节")
CODE("lunar calendar 农历     solar term 节气")
CODE("family reunion 团圆     pay tribute to 祭拜")

AP("高频动词（翻译里反复出现）", bold=True, size=12, color=BLUE)
CODE("originate from 起源于     be widely regarded as 被广泛认为")
CODE("play a vital role in 在...中起重要作用")
CODE("symbolize / represent / stand for 象征着")
CODE("be listed as 被列为     be passed down 被传承下来")
CODE("gain popularity 获得普及     be introduced to 传入到")
CODE("reflect / embody 体现     have a profound influence on 深刻影响")
CODE("serve as 作为 / 起到...作用     date back to 追溯到")

AP("社会发展类词汇（备胎方向）", bold=True, size=12, color=BLUE)
CODE("artificial intelligence 人工智能     digital economy 数字经济")
CODE("rural revitalization 乡村振兴     carbon neutrality 碳中和")
CODE("high-speed rail 高铁     e-commerce platform 电商平台")
CODE("take measures to... 采取措施...     make remarkable achievements 取得显著成就")
CODE("break the monopoly 打破垄断     high-quality development 高质量发展")

SEP()

AP("翻译提分口诀", bold=True, size=13, color=RED)
CODE("1. 中文短句多 → 英文用 and / but / which / -ing 串成长句")
CODE("2. 遇到不会的词 → 用 you know that thing where... 不行!")
CODE("   正确做法: 用上位词替代 (刺绣不会→traditional handicraft)")
CODE("3. 朝代、专有名词首字母大写: Han Dynasty 不是 han dynasty")
CODE("4. 被动语态优先: 中文说'剪纸被广泛使用'→ Paper-cutting is widely used")
CODE("5. 每翻完一句默读一遍, 主谓宾齐全就过, 别纠结")

doc.add_page_break()

# ================================================================
# FINAL CHECKLIST
# ================================================================
AP("考前最后确认", bold=True, size=14, color=RED)
SEP()

CODE("[ ] 三段框架能默写: para1锚定话题 / para2原因-影响-深化 / para3展望")
CODE("[ ] 数字生活例子能默写 (室友刷视频那个)")
CODE("[ ] 让步反驳句式: Admittedly, ... Yet this overlooks ...")
CODE("[ ] 长句句式: What makes ... particularly ... is that ...")
CODE("[ ] 衔接词各记1个: Beyond that, / That said, / In light of the above,")
CODE("[ ] 三件套不会忘: 例子 + 让步反驳 + 长句 → 塞第二段")
CODE("[ ] 禁用词心里有数: First/Second/Third  |  Every coin has two sides")
CODE("")
AP("进考场前看最后一眼：第一段锚定话题 → 第二段三件套 → 第三段展望。30分钟，够了。", bold=True, size=12, color=RED)

output = os.path.normpath(os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "CET4_最后一背_考前版_v2.docx"))
doc.save(output)
print(f"Done: {output}")