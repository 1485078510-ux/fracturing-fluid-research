# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor

doc = Document()
style = doc.styles['Normal']
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

BLUE = (0x1A, 0x56, 0xDB)
RED = (0xCC, 0x33, 0x33)

def AP(text, bold=False, size=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)

def CODE(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(12)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)

def BLANK():
    """Empty line as spacing between paragraphs"""
    p = doc.add_paragraph()
    p.add_run(" ").font.size = Pt(6)

# ================================================================
AP("四级作文万能成文框架", bold=True, size=18, color=BLUE)
AP("[ ] 内填你的内容, 其余照抄, 直接成文", size=11, color=RED)
AP("目标11-14分 · 通吃90%题目", size=10)

BLANK()

# ========== 第一段 ==========
AP("第一段 (2-3句)", bold=True, size=14, color=BLUE)
BLANK()

AP("▎选一个:", size=10, color=RED)
BLANK()

AP("A. 现象利弊类 (AI/快递/绿色校园……)", bold=True, size=11)
CODE("In recent years, the issue of [话题关键词] has sparked considerable")
CODE("discussion. While this trend brings notable benefits, it also raises")
CODE("concerns that deserve serious attention.")

BLANK()

AP("B. 观点论证类 (是否该做某事)", bold=True, size=11)
CODE("Opinions vary greatly when it comes to [话题]. Some argue that [正方],")
CODE("while others believe that [反方]. From where I stand, the [former/latter]")
CODE("view holds more weight.")

BLANK()

AP("C. 建议/意义类 (如何培养某能力/某事的意义)", bold=True, size=11)
CODE("How can [人群] best [目标]? This question has become increasingly")
CODE("pressing. In my view, the answer lies in [方案A] and [方案B].")

BLANK()

AP("D. 描述/应用文类 (投稿/推荐/印象最深的事)", bold=True, size=11)
CODE("Among the many [experiences/courses/activities] I have had,")
CODE("[具体对象] stands out as the most memorable, not only because of")
CODE("[原因], but also for its profound impact on my [成长方面].")

BLANK()
BLANK()

# ========== 第二段 ==========
AP("第二段 (4-5句)  ← 拉分关键！", bold=True, size=14, color=BLUE)
AP("!!! 严禁 First... Second... Finally !!!", bold=True, color=RED)
BLANK()

AP("▎选一个:", size=10, color=RED)
BLANK()

AP("1. 原因→影响→深化 (利弊分析/现象类)", bold=True, size=11)
CODE("Several factors explain this [phenomenon/trend]. At its root,")
CODE("[根本原因] has created fertile ground for it to take hold.")
CODE("The effects are already visible: [具体影响A], and more importantly,")
CODE("[更深影响B]. What deserves equal attention is [被忽视的问题]")
CODE("-- a point often neglected in casual discussion.")

BLANK()

AP("2. 论点→论据→让步反驳 (观点论证/思辨类)", bold=True, size=11)
CODE("My support for [你的立场] rests on both practical and principled")
CODE("grounds. From a practical standpoint, [具体理由]. Take [简短例子]")
CODE("as an illustration: [1-2句例证]. Admittedly, those who favor [对方观点]")
CODE("have a point. Yet this concern, legitimate as it is, overlooks a")
CODE("crucial fact: [你的反驳]. The deeper issue is not merely [表面争议],")
CODE("but rather how we can [本质目标].")

BLANK()

AP("3. 问题→方案→可行性 (提建议/措施类)", bold=True, size=11)
CODE("To meet this challenge, action is needed on multiple fronts. At the")
CODE("university level, [校方] could introduce [措施A], which has proven")
CODE("effective in [类似案例]. On the student side, we should take the")
CODE("initiative to [行动B] -- treating [困难] as an opportunity for")
CODE("growth. While these steps demand effort, the payoff -- [预期成果]")
CODE("-- makes them a worthy investment.")

BLANK()
BLANK()

# ========== 第三段 ==========
AP("第三段 (2-3句)", bold=True, size=14, color=BLUE)
BLANK()

AP("▎选一个:", size=10, color=RED)
BLANK()

AP("X. 展望建议 (万能, 90%题目可用)", bold=True, size=11)
CODE("In light of the above, I am convinced that [重申观点] is not only")
CODE("desirable but achievable -- provided that [关键条件] is taken")
CODE("seriously. The path forward calls for a balance between [维度A]")
CODE("and [维度B], a challenge today's students are ready to meet.")

BLANK()

AP("Y. 呼吁行动 (建议/环保/应用文)", bold=True, size=11)
CODE("The measures outlined above are by no means exhaustive, but they")
CODE("offer a practical starting point. What matters most now is not")
CODE("further debate, but concrete action. The time to start is now.")

BLANK()

AP("Z. 反思升华 (思辨/哲理/科技类)", bold=True, size=11)
CODE("Perhaps the real question is not whether [表面议题], but what kind")
CODE("of [更大价值] we wish to embrace. In a world of relentless change,")
CODE("[核心能力] may well be the compass that guides us -- not toward")
CODE("easy answers, but toward better questions.")

BLANK()
BLANK()

# ========== 必杀三招 ==========
AP("--- 第二段必塞三样东西 (直接从11分拉到14分) ---", bold=True, size=12, color=RED)
BLANK()

AP("1. 塞一个具体例子:", bold=True)
CODE("A classmate of mine, who once [困难], now [做法].")
AP("2. 塞一句让步反驳:", bold=True)
CODE("Admittedly, [对方观点] carries weight. Yet this overlooks [你的反驳].")
AP("3. 塞一个长句:", bold=True)
CODE("What makes [话题] particularly [adj.] is that it [动词], influencing")
CODE("not only [A] but also [B] -- a reality we can no longer ignore.")

BLANK()
BLANK()

AP("--- 绝对禁用 (用了直接降档) ---", bold=True, size=12, color=RED)
CODE("First... Second... Third...    Every coin has two sides.")
CODE("With the development of society...    That's all. Thank you.")

BLANK()

AP("心法: 框架照抄, [ ] 内写自己的内容。每篇都不同, 但骨架一样稳。", bold=True, size=13, color=BLUE)

# Save
output = os.path.normpath(os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "CET4_成文框架_极简版.docx"))
doc.save(output)
print(f"Done: {output}")