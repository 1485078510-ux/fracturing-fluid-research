# -*- coding: utf-8 -*-
"""Generate a lean, memorizable CET-4 template."""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

style = doc.styles['Normal']
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.4
style.paragraph_format.space_after = Pt(6)

BLUE = (0x1A, 0x56, 0xDB)
RED = (0xCC, 0x33, 0x33)
GRAY = (0x6B, 0x72, 0x80)

def AP(text, bold=False, size=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)
    return p

def CODE(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(0.5)

def SEP():
    p = doc.add_paragraph()
    p.add_run("─" * 50).font.size = Pt(6)

# ================================================================
AP("四级作文万能模板 · 纯净背诵版", bold=True, size=16, color=BLUE)
AP("30分钟速成 · 直接套用 · 目标11-14分", size=10, color=GRAY)
SEP()

# ================================================================
AP("用法：拿到题目 → 判断属于哪一类 → 从对应句式库中选句子 → 把 [ ] 替换成题目关键词 → 成文", bold=True, size=11, color=RED)
SEP()

# ========== PARA 1 ==========
AP("第一段：引入 (2-3句, 30-40词)", bold=True, size=13, color=BLUE)

AP("【A. 现象/利弊类】如：AI学习、快递发展、绿色校园、设施开放", bold=True)
CODE("In recent years, the issue of [话题] has sparked considerable discussion,")
CODE("particularly regarding [具体切入点]. While this trend brings notable")
CODE("benefits, it also raises concerns that deserve serious attention.")

AP("【B. 观点论证类】如：是否该必修某课、是否该做某事", bold=True)
CODE("Opinions vary greatly when it comes to [话题]. Some argue that [正方],")
CODE("while others firmly believe that [反方]. From where I stand, the former/")
CODE("latter view holds more weight -- for reasons I will elaborate below.")

AP("【C. 建议/意义类】如：如何培养某能力、某事的意义", bold=True)
CODE("How can [人群] best [目标] in an era of [时代背景]? This question has")
CODE("become increasingly pressing. In my view, the answer involves a")
CODE("combination of [方案A] and [方案B], supported by effort at multiple levels.")

AP("【D. 描述/应用文类】如：投稿、推荐、印象最深的事", bold=True)
CODE("Among the many [experiences / courses / activities] I have had as a")
CODE("university student, [具体对象] stands out as the most memorable, not")
CODE("only because of [直接原因], but also for the profound impact it has")
CODE("left on my [成长维度].")

SEP()

# ========== PARA 2 ==========
AP("第二段：论证 (4-5句, 60-70词)  ← 拉分关键！", bold=True, size=13, color=BLUE)
AP("!!! 严禁用 First... Second... Finally !!!  改用逻辑递进链", bold=True, color=RED)

AP("【路径1. 原因→影响→深化】适合：利弊分析、现象评述", bold=True)
CODE("Several factors help explain this [phenomenon/trend]. At its root,")
CODE("[根本原因] has created fertile ground for it to take hold. The ripple")
CODE("effects are already visible: [具体影响A] has begun to reshape [方面],")
CODE("and more importantly, [更深影响B] is gradually altering how [人群]")
CODE("approach [领域]. What deserves equal attention, however, is [被忽视")
CODE("的连锁反应] -- a point often neglected in casual discussion.")

AP("【路径2. 论点→论据→让步反驳】适合：观点论证、是否该做某事", bold=True)
CODE("My support for [你的立场] rests on both practical and principled")
CODE("grounds. From a practical standpoint, [具体理由]. Take [简短例子] as")
CODE("an illustration: [1-2句例证细节]. Admittedly, those who favor [对方")
CODE("观点] have a point when they note that [对方理由]. Yet this concern,")
CODE("legitimate as it is, overlooks a crucial fact: [你的反驳]. The deeper")
CODE("issue, then, is not merely [表面争议], but rather how we can [本质目标].")

AP("【路径3. 问题→方案→可行性】适合：提建议、措施类、应用文", bold=True)
CODE("To meet this challenge, action is needed on multiple fronts. At the")
CODE("university level, [校方] could introduce [措施A], which has already")
CODE("yielded encouraging results in [同类案例]. On the student side, we")
CODE("ourselves should take the initiative to [行动B] -- treating [困难]")
CODE("not as a burden but as an opportunity for growth. While these steps")
CODE("demand genuine effort, the payoff -- [预期成果] -- makes them a")
CODE("worthy investment of time and energy.")

SEP()

# ========== PARA 3 ==========
AP("第三段：结尾 (2-3句, 25-35词)  ← 冲刺14分！", bold=True, size=13, color=BLUE)

AP("【X. 展望建议型】万能收束，90%题目可用", bold=True)
CODE("In light of the discussion above, I am convinced that [重申观点] is")
CODE("not only desirable but achievable -- provided that [关键条件] is given")
CODE("the attention it deserves. The path forward calls for a thoughtful")
CODE("balance between [维度A] and [维度B], a challenge today's university")
CODE("students are well equipped to meet.")

AP("【Y. 呼吁行动型】适合：建议措施、环保、应用文", bold=True)
CODE("The measures outlined above -- from [措施A] to [措施B] -- are by no")
CODE("means exhaustive, but they offer a practical starting point. What")
CODE("matters most at this juncture is not further debate, but concrete")
CODE("steps. The time to start is now.")

AP("【Z. 反思升华型】适合：思辨类、科技伦理、哲理话题", bold=True)
CODE("Perhaps the real question is not whether [表面议题], but what kind of")
CODE("[更大价值] we wish to embrace as a generation. In a world of relentless")
CODE("change, [核心能力] may well prove to be the compass that guides us --")
CODE("not toward easy answers, but toward better questions.")

doc.add_page_break()

# ========== 加分必杀技 ==========
AP("六大加分必杀技 (背下来, 每个值1-2分)", bold=True, size=13, color=BLUE)

AP("1. 首段同义替换", bold=True, color=RED)
AP("拿到题目花30秒想3个话题同义词, 全文至少用2个不同说法。如 AI: artificial intelligence / AI-powered tools / intelligent assistants", size=10)

AP("2. 段落间用代词回指, 不用序号", bold=True, color=RED)
CODE("[X] First... Second... Finally...     ← 找死")
CODE("[OK] At its root, ... Such efficiency, however, ... What deserves")
CODE("     equal attention is...            ← 自然衔接, 高分信号")
AP("备用衔接词: Beyond that, | That said, | This, in turn, leads to... | More significantly,", size=10)

AP("3. 插入一个具体例证 (必须!)", bold=True, color=RED)
CODE("A classmate of mine, who once struggled with [困难], now [做法] --")
CODE("yet she is always careful to [克制], treating [工具/方法] as a [积极")
CODE("比喻] rather than a [消极比喻].")

AP("4. 写一个复合长句 (展示语法)", bold=True, color=RED)
CODE("What makes [话题] particularly [adj.] is that it [动词短语],")
CODE("influencing not only [A] but also [B] -- a reality that [人群]")
CODE("can no longer afford to ignore.")

AP("5. 一句让步反驳 (展示思辨)", bold=True, color=RED)
CODE("Admittedly, [对方观点] carries some weight. Yet this concern,")
CODE("legitimate as it is, overlooks a crucial point: [你的反驳].")

AP("6. 结尾人文温度 (最后3秒好感)", bold=True, color=RED)
CODE("Ultimately, [主题升华]. If we [积极行动], the [场景] of the future")
CODE("may well be a place where [人文价值], not [表面指标], takes center stage.")

doc.add_page_break()

# ========== 救命词汇 ==========
AP("六大话题救命词汇 (每个方向背3个)", bold=True, size=13, color=BLUE)

words = [
    ("AI/科技", "digital literacy | algorithm-driven | harness the power of | double-edged sword"),
    ("校园学习", "academic integrity | strike a balance | foster an environment | self-directed learning"),
    ("环保绿色", "carbon neutrality | sustainable development | take concrete steps | every small effort counts"),
    ("文化传承", "cultural confidence | breathe new life into | bridge past and present | time-honored"),
    ("个人成长", "resilience | critical thinking | step out of one's comfort zone | lifelong learning"),
    ("社会议题", "digital divide | work-life balance | bridge the gap | public concern"),
]
for topic, w in words:
    AP(f"  {topic}: {w}", size=10)

SEP()

# ========== 禁令 + 公式 ==========
AP("绝对禁止用语 (用了直接降档)", bold=True, size=13, color=RED)
AP("First... Second... Third...  |  Every coin has two sides.  |  With the development of society...  |  That's all. Thank you.", size=10)
SEP()

AP("分数公式", bold=True, size=13, color=BLUE)
AP("结构清晰 + 切题 + 无明显语法错 = 11分", size=11)
AP("以上 + 让步反驳 + 复合长句 + 具体例证 = 13-14分", size=11)
AP("以上 + 人文温度结尾 = 14-15分", size=11)
SEP()

AP("心法：框架在心里，表达在笔头。模板是骨架，你的思考才是血肉。", bold=True, size=12, color=BLUE)

# Save
output = os.path.normpath(os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "CET4_万能模板_纯净背诵版.docx"))
doc.save(output)
print(f"Done: {output}")