# -*- coding: utf-8 -*-
"""Generate CET-4 Essay Universal Template as Word document."""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Style
style = doc.styles['Normal']
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.35
style.paragraph_format.space_after = Pt(4)

for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    h.font.bold = True
    sizes = {1: 20, 2: 15, 3: 12.5}
    h.font.size = Pt(sizes[i])

BLUE = (0x1A, 0x56, 0xDB)
RED = (0xCC, 0x33, 0x33)
GRAY = (0x6B, 0x72, 0x80)
DARK = (0x2D, 0x33, 0x3B)
LGRAY = (0xCC, 0xCC, 0xCC)

def AP(text, bold=False, italic=False, size=None, color=None, align=None, indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)
    if align is not None: p.alignment = align
    if indent: p.paragraph_format.left_indent = Cm(indent)
    return p

def TIP(text):
    p = doc.add_paragraph()
    run = p.add_run("[TIP] " + text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(*GRAY)
    run.italic = True
    p.paragraph_format.left_indent = Cm(0.8)

def CODE(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(*DARK)
    p.paragraph_format.left_indent = Cm(1.0)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F0F3F5')
    shd.set(qn('w:val'), 'clear')
    p.paragraph_format.element.get_or_add_pPr().append(shd)

def SEP():
    p = doc.add_paragraph()
    run = p.add_run("─" * 55)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(*LGRAY)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def BUL(text, indent=0.6):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Cm(indent)

# ================================================================
# COVER
# ================================================================
doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("CET-4 Essay Universal Template"); r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = RGBColor(*BLUE)
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Grader's Perspective - High-Score Edition - 2023-2026 All Topics"); r.font.size = Pt(11); r.font.color.rgb = RGBColor(*GRAY)
doc.add_paragraph()

# ================================================================
doc.add_heading("How This Template Beats Generic Ones", level=2)
AP("Generic online templates have a fatal flaw: every student uses the same fixed sentences. Graders spot them in 5 seconds and cap the score at 8/15.", size=10, color=GRAY)
BUL("[X] Generic templates = memorize fixed sentences -> every essay looks identical -> grader spots it -> max 8/15")
BUL("[OK] This template = modular framework with multiple plug-in options per paragraph -> you mix and match -> every essay reads differently -> easy 11-14/15")
BUL("[OK] Built-in 6 grader hooks: scoring signals designed for 30-second rapid grading")
BUL("[OK] Covers all 6 major topic directions from 2023-2025 real exams + 2026 predictions")
TIP("How to use: Read prompt -> Identify topic type (10 sec) -> Pick 1 option from each paragraph -> Fill in topic keywords -> Done. NEVER memorize entire paragraphs word-for-word!")

doc.add_page_break()

# ================================================================
# PART 1: THREE-PARAGRAPH FRAMEWORK
# ================================================================
doc.add_heading("Part 1: The Three-Paragraph Universal Framework", level=1)
AP("Below is a ready-to-use framework that fits 90%+ of CET-4 prompts. Replace [bracketed] parts with your keywords. Each position offers multiple options -- pick the one that feels most natural.", size=10, color=GRAY)

# PARA 1
doc.add_heading("PARAGRAPH 1: Introduction (2-3 sentences, 30-40 words)", level=2)
AP("MISSION: Let the grader confirm in 5 seconds that your essay is ON-TOPIC.", bold=True)

AP("Option A: Phenomenon Anchor (pros/cons analysis, social trends, tech topics)", bold=True, size=11, color=BLUE)
CODE("In recent years, the issue of [TOPIC KEYWORD] has sparked considerable")
CODE("discussion, particularly regarding [SPECIFIC ANGLE]. While this trend brings")
CODE("notable benefits, it also raises concerns that deserve serious attention.")

AP("Option B: Viewpoint Contrast (argumentative, should-we-or-shouldn't-we topics)", bold=True, size=11, color=BLUE)
CODE("Opinions vary greatly when it comes to [TOPIC]. Some argue that [VIEW A],")
CODE("while others firmly believe that [VIEW B]. From where I stand, the [former/")
CODE("latter] view holds more weight -- for reasons I will elaborate below.")

AP("Option C: Rhetorical Question (suggestion/proposal essays, significance essays)", bold=True, size=11, color=BLUE)
CODE("How can [GROUP] best [ACHIEVE GOAL] in an era of [CONTEXT]? This question has")
CODE("become increasingly pressing. In my view, the answer involves a combination of")
CODE("[APPROACH A] and [APPROACH B], supported by effort at multiple levels.")

AP("Option D: Scene Narrative (application writing, personal experience, descriptive)", bold=True, size=11, color=BLUE)
CODE("Among the many [experiences/courses/activities] I have had as a university")
CODE("student, [SPECIFIC THING] stands out as the most [meaningful/memorable], not")
CODE("only because of [IMMEDIATE REASON], but also for the profound impact it has")
CODE("left on my [PERSONAL GROWTH DIMENSION].")

TIP("Pick ONLY 1 option for Paragraph 1. Choose whichever flows most naturally for you.")

# PARA 2
doc.add_heading("PARAGRAPH 2: Body / Argument (4-5 sentences, 60-70 words)", level=2)
AP("MISSION: Show logical progression + language variety. This is where scores diverge.", bold=True)
AP("CRITICAL RULE: NEVER use First... Second... Finally! Use logical progression chains instead.", bold=True, color=RED)

AP("Path 1: Cause -> Effect -> Deepen (pros/cons analysis, phenomenon essays)", bold=True, size=11, color=BLUE)
CODE("Several factors help explain this [phenomenon/trend]. At its root, [ROOT CAUSE]")
CODE("has created fertile ground for it to take hold. The ripple effects are already")
CODE("visible: [SPECIFIC EFFECT A] has begun to reshape [ASPECT], and more importantly,")
CODE("[DEEPER EFFECT B] is gradually altering how [GROUP] approach [FIELD]. What")
CODE("deserves equal attention, however, is [OVERLOOKED CONSEQUENCE] -- a point")
CODE("often neglected in casual discussion.")

AP("Path 2: Claim -> Evidence -> Concession-Rebuttal (argumentative, critical thinking)", bold=True, size=11, color=BLUE)
CODE("My support for [YOUR POSITION] rests on both practical and principled grounds.")
CODE("From a practical standpoint, [SPECIFIC REASON A]. Take [BRIEF EXAMPLE] as an")
CODE("illustration: [1-2 sentence detail]. Admittedly, those who favor [OPPOSING VIEW]")
CODE("have a point when they note that [THEIR STRONGEST ARGUMENT]. Yet this concern,")
CODE("legitimate as it is, overlooks a crucial fact: [YOUR REBUTTAL]. The deeper issue,")
CODE("then, is not merely [SURFACE DEBATE], but rather how we can [ESSENTIAL GOAL].")

AP("Path 3: Problem -> Solution -> Feasibility (proposal/suggestion, application writing)", bold=True, size=11, color=BLUE)
CODE("To meet this challenge, action is needed on multiple fronts. At the university")
CODE("level, [INSTITUTION] could introduce [SPECIFIC MEASURE A], which has already")
CODE("yielded encouraging results in [SIMILAR CASE]. On the student side, we ourselves")
CODE("should take the initiative to [SPECIFIC ACTION B] -- treating [DIFFICULTY] not")
CODE("as a burden but as an opportunity for growth. While these steps demand genuine")
CODE("effort, the payoff -- [EXPECTED OUTCOME] -- makes them a worthy investment.")

TIP("Pick ONLY 1 path. Crucially: insert 1 concrete example (even just 'a classmate of mine...') and 1 complex long sentence. These two elements alone separate your essay from hollow template writing.")

# PARA 3
doc.add_heading("PARAGRAPH 3: Conclusion (2-3 sentences, 25-35 words)", level=2)
AP("MISSION: Do NOT repeat the introduction. ELEVATE the argument. This pushes 11 -> 14.", bold=True)

AP("Option X: Prospect & Suggestion (universal closer, fits 90% of prompts)", bold=True, size=11, color=BLUE)
CODE("In light of the discussion above, I am convinced that [RESTATE CORE VIEW] is")
CODE("not only desirable but achievable -- provided that [KEY CONDITION] is given the")
CODE("attention it deserves. The path forward calls for a thoughtful balance between")
CODE("[DIMENSION A] and [DIMENSION B], a challenge today's students are ready to meet.")

AP("Option Y: Call to Action (proposal essays, environmental topics, application writing)", bold=True, size=11, color=BLUE)
CODE("The measures outlined above -- from [MEASURE A] to [MEASURE B] -- are by no")
CODE("means exhaustive, but they offer a practical starting point. What matters most")
CODE("at this juncture is not further debate, but concrete steps. The time to start is now.")

AP("Option Z: Reflection & Elevation (critical thinking, philosophy, tech ethics topics)", bold=True, size=11, color=BLUE)
CODE("Perhaps the real question is not whether [SURFACE ISSUE], but what kind of")
CODE("[LARGER VALUE] we wish to embrace as a generation. In a world of relentless")
CODE("change, [CORE ABILITY] may well prove to be the compass that guides us -- not")
CODE("toward easy answers, but toward better questions.")

AP("Three NEVER rules for the conclusion:", bold=True, color=RED)
BUL("NEVER repeat the introduction verbatim -- rephrase from a different angle")
BUL("NEVER introduce a brand-new argument -- conclusion is for summary + elevation only")
BUL("NEVER end with 'That's all. Thank you.' -- inappropriate in Chinese English exams")

doc.add_page_break()

# ================================================================
# PART 2: FULL WORKED EXAMPLE
# ================================================================
doc.add_heading("Part 2: Full Worked Example (Real June 2025 Exam Topic)", level=1)
AP("Real topic (June 2025): Express your views on AI-assisted learning -- including advantages, potential problems, and suggestions.", bold=True, italic=True, size=10, color=GRAY)
SEP()
AP("Combination: Para 1 Option A + Para 2 Path 1 + Para 3 Option X", bold=True, size=10, color=BLUE)
SEP()

doc.add_heading("Full Model Essay", level=3)

essay = (
    "In recent years, the growing use of AI-powered tools in university learning has sparked considerable "
    "discussion, particularly regarding how these technologies are reshaping the way students study. "
    "While AI brings notable benefits, it also raises concerns that deserve serious attention.\n\n"
    "Several factors help explain this rapid adoption. At its root, the sheer convenience of AI tools -- "
    "from intelligent writing assistants to adaptive learning platforms -- has created fertile ground for "
    "them to take hold on campuses everywhere. The ripple effects are already visible: tasks that once took "
    "hours, such as drafting an essay outline or debugging code, can now be completed in minutes. This "
    "efficiency, however, has sparked a deeper worry: are students gradually losing the very skills these "
    "tools are meant to enhance? A classmate of mine, for instance, admitted that she had started relying so "
    "heavily on AI grammar checkers that her own proofreading ability had noticeably declined. What deserves "
    "equal attention is the ethical dimension -- the line between using AI as a learning aid and crossing "
    "into academic dishonesty remains frustratingly blurry, a point often glossed over in casual discussion.\n\n"
    "In light of the discussion above, I am convinced that AI-assisted learning is not only here to stay but "
    "can be a powerful force for good -- provided that clear guidelines and a culture of responsible use are "
    "cultivated. The path forward calls for a thoughtful balance between embracing technological convenience "
    "and preserving the independent thinking that lies at the heart of true education. This is a challenge "
    "that today's university students are well equipped to meet."
)
AP(essay, size=10.5)

SEP()
doc.add_heading("Why This Essay Scores 13-14/15", level=3)
BUL("Para 1: Immediately anchors topic with keyword variation (AI-powered tools -> technologies -> AI) -- signals lexical range")
BUL("Para 2: Logical progression (convenience -> efficiency -> worry -> ethics) instead of First/Second/Third")
BUL("Para 2: Contains 1 concrete personal example -- breaks the template feel")
BUL("Para 2: Contains 1 complex long sentence -- signals grammar control")
BUL("Para 2: Contains implied concession -- shows critical thinking")
BUL("Para 3: Does NOT repeat intro; elevates to 'independent thinking' and 'true education' -- humanistic touch")
BUL("Zero First/Second/Third usage throughout -- reads like natural writing, not a template")

doc.add_page_break()

# ================================================================
# PART 3: SIX GRADER HOOKS
# ================================================================
doc.add_heading("Part 3: Six Grader Hooks (Score Boosters)", level=1)
AP("6 techniques designed for 30-second rapid grading. Each directly lifts your impression score.", size=10, color=GRAY)

doc.add_heading("Hook 1: Keyword Variation in Paragraph 1 -- Show Lexical Range Instantly", level=2)
AP("Graders' eyes lock onto synonyms of the topic keyword. Show vocabulary range right away and your impression score jumps immediately.", size=10)
AP("Tactic: Spend 30 seconds listing 3-4 alternative expressions for the topic word. Use at least two different ones in the essay.", bold=True)
CODE("Example - Topic involves AI/technology:")
CODE("  artificial intelligence / AI-powered tools / intelligent assistants /")
CODE("  algorithm-driven platforms / machine-generated feedback")
CODE("Example - Topic involves learning/education:")
CODE("  academic development / intellectual growth / the learning process /")
CODE("  knowledge acquisition / educational experience")

doc.add_heading("Hook 2: Invisible Transitions -- Pronoun Reference Instead of Sequence Numbers", level=2)
AP("Never use First/Second/Third. Use logical connectors + pronoun back-reference + lexical repetition for natural flow.", size=10)
CODE("[X] Generic: First, AI helps students learn faster. Second, AI causes dependence.")
CODE("[OK] High-score: At its root, AI's convenience has fueled its rapid spread.")
CODE("     -> Such efficiency, however, comes at a price.")
CODE("     -> What deserves equal attention is the ethical dimension.")
AP("Advanced transition word bank:", bold=True)
CODE("Progression: Beyond that, | More significantly, | Equally noteworthy is that...")
CODE("Contrast:   That said, | Yet this alone is hardly the full picture.")
CODE("Cause:      This, in turn, leads to... | Small wonder, then, that...")
CODE("Concession: Admittedly, ... | ..., legitimate as it is, ...  [<- bonus phrase]")

doc.add_heading("Hook 3: One Concrete Example -- Worth More Than Three Empty Paragraphs", level=2)
AP("Graders dread hollow essays where every sentence is correct but carries zero real information. Insert ONE specific example to instantly transform your essay from template filler to thoughtful writing.", size=10)
CODE("[X] Hollow: Many students use AI tools to help them study more efficiently.")
CODE("[OK] Concrete: A classmate of mine, who once struggled with academic writing,")
CODE("     now uses AI-powered grammar checkers to polish her drafts -- yet she is")
CODE("     always careful to revise the suggestions herself, treating the tool as a")
CODE("     tutor rather than a crutch.")
AP("Universal example starters:", bold=True)
CODE("- A survey among my classmates revealed that...")
CODE("- I recall a course I took last semester, in which...")
CODE("- Research in educational psychology suggests that...  [no citation needed]")

doc.add_heading("Hook 4: One Complex Long Sentence -- Show Grammar Control", level=2)
AP("Design 1 grammatically sound long sentence (25-35 words) with embedded clauses in Paragraph 2. This is a hard signal for the high-score band.", size=10)
CODE("Model long sentence:")
CODE("What makes this issue particularly complex, however, is that the impact of AI")
CODE("on learning extends far beyond the classroom, influencing not only how students")
CODE("acquire knowledge but also how they perceive the very value of independent")
CODE("thinking in an age when answers are always just one click away.")
AP("Universal long sentence frame:", bold=True)
CODE("What makes [TOPIC] particularly [ADJ] is that it [VERB PHRASE], influencing")
CODE("not only [ASPECT A] but also [ASPECT B] -- a reality that [GROUP] can no")
CODE("longer afford to ignore.")

doc.add_heading("Hook 5: One Concession-Rebuttal Move -- Show Critical Thinking", level=2)
AP("Acknowledge the opposing view's reasonableness, then rebut it. The 2026 updated rubric explicitly lists multi-dimensional critical thinking as a high-score requirement.", size=10)
CODE("Universal concession-rebuttal frame:")
CODE("To be fair, it is understandable why some might argue that [OPPOSING VIEW].")
CODE("After all, [THEIR STRONGEST POINT]. Yet this concern, legitimate as it is,")
CODE("overlooks a crucial point: [YOUR REBUTTAL].")
AP("Key phrases (each signals advanced proficiency):", bold=True)
CODE("Admittedly, ...              <- opens the concession")
CODE("..., legitimate as it is, ... <- parenthetical concession (advanced grammar)")
CODE("Yet this alone does not tell   <- pivots to your rebuttal")
CODE("the whole story.")

doc.add_heading("Hook 6: A Humanistic Touch in the Final Sentence", level=2)
AP("Most template essays end coldly. End with a touch of warmth or contemporary relevance, and the grader's last 3 seconds tip in your favor.", size=10)
CODE("[X] Robotic:  In conclusion, AI is a double-edged sword. We should use it wisely.")
CODE("[OK] Human:   Ultimately, technology only amplifies what we already are. If we")
CODE("     approach AI with curiosity tempered by critical thinking, the classroom of")
CODE("     the future may well be a place where human potential, not machine")
CODE("     capability, takes center stage.")

doc.add_page_break()

# ================================================================
# PART 4: TOPIC VOCABULARY
# ================================================================
doc.add_heading("Part 4: Topic Vocabulary Arsenal (6 Major Directions)", level=1)
AP("For each direction, memorize only 5-6 key words/phrases for instant deployment.", size=10, color=GRAY)

vocab = [
    ("AI & Technology", [
        "digital literacy", "algorithm-driven", "information cocoon",
        "harness the power of [tech]", "double-edged sword",
        "a tool only as wise as the hands that wield it",
    ]),
    ("Campus Learning & Growth", [
        "academic integrity", "strike a balance between A and B",
        "foster a supportive environment", "self-directed learning",
        "peer collaboration",
        "education is not the filling of a pail, but the lighting of a fire",
    ]),
    ("Environment & Green Campus", [
        "carbon neutrality", "sustainable development", "ecological footprint",
        "take concrete steps", "think globally, act locally",
        "every small effort counts",
    ]),
    ("Cultural Heritage & Confidence", [
        "intangible cultural heritage", "cultural confidence",
        "breathe new life into tradition", "bridge the past and the present",
        "time-honored",
        "a nation's culture resides in the heart of its people",
    ]),
    ("Personal Growth & Skills", [
        "resilience", "critical thinking", "step out of one's comfort zone",
        "a well-rounded individual", "lifelong learning",
        "character is built in moments of challenge",
    ]),
    ("Social Issues & Public Concerns", [
        "aging population", "digital divide", "work-life balance",
        "public concern", "bridge the gap between A and B",
        "a challenge that concerns every member of society",
    ]),
]

for topic, words in vocab:
    AP(f">> {topic}", bold=True, size=11, color=BLUE)
    for w in words:
        BUL(f"- {w}", indent=1.0)

doc.add_page_break()

# ================================================================
# PART 5: EXAM ROOM BATTLE PLAN
# ================================================================
doc.add_heading("Part 5: Exam Room 30-Minute Battle Plan", level=1)

doc.add_heading("Time Allocation", level=2)
times = [
    ("0-3 min", "Analyze prompt + outline", "Circle 3 keywords; determine topic type; jot 3 key points on scratch paper"),
    ("3-8 min", "Write Paragraph 1", "2-3 sentences. Pick option -> fill topic words -> ensure 1 synonym variation"),
    ("8-22 min", "Write Paragraph 2", "4-5 sentences. Pick path -> logical progression -> MUST include 1 example + 1 long sentence"),
    ("22-25 min", "Write Paragraph 3", "2-3 sentences. Elevate without repeating -> end with humanistic touch"),
    ("25-30 min", "Proofread", "Check: subject-verb agreement, tense, spelling. Fix only obvious errors; do NOT restructure"),
]
for time, task, detail in times:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{time}  ")
    r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(f"{task}: {detail}")
    r2.font.size = Pt(10)

doc.add_heading("Five Score-Killing Mistakes", level=2)
mistakes = [
    ("Copying a fixed online template word-for-word", "2026 rubric: template essays capped at 8/15"),
    ("Going off-topic or missing a prompt keyword", "Once off-topic, even perfect language drops below 5/15"),
    ("Using First/Second/Third throughout", "Instantly exposes template dependency; structure score tanks"),
    ("Stuffing obscure words used incorrectly", "Accuracy > flashiness. Wrong usage costs more than simple correct words"),
    ("Writing fewer than 100 words", "Minimum 3-point deduction; structure score cannot be salvaged"),
]
for m, c in mistakes:
    p = doc.add_paragraph()
    r1 = p.add_run(f"[X] {m}")
    r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = RGBColor(*RED)
    p2 = doc.add_paragraph()
    r2 = p2.add_run(f"    -> {c}")
    r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(*GRAY)
    p2.paragraph_format.left_indent = Cm(0.8)

doc.add_heading("One-Week-Before-Exam Checklist", level=2)
checklist = [
    "Write out the 3 argument paths from memory once each (don't memorize sentences, just the logical structure)",
    "Pick the 3 most comfortable words from each of the 6 topic directions; lock them in",
    "Write the 6 Grader Hooks on a sticky note; glance at it right before entering the exam hall",
    "Do 1 full timed practice (30 min) with a real past exam prompt to verify the template fits your style",
    'Internalize: "Framework in mind, expression in hand. The template is the skeleton; your thinking is the flesh."',
]
for item in checklist:
    BUL(f"[ ] {item}", indent=0.5)

doc.add_page_break()

# ================================================================
# PART 6: ONE-PAGE QUICK REFERENCE
# ================================================================
doc.add_heading("Appendix: One-Page Quick Reference (Print & Carry on Exam Day)", level=1)
AP("Print this page separately. Review on your way to the exam.", size=10, color=GRAY)

SEP()
AP("THREE-PARAGRAPH FRAMEWORK (Quick Memory)", bold=True, size=12)
CODE("Para 1 (Intro):   Phenomenon / Contrast / Question / Narrative -> Pick 1")
CODE("Para 2 (Body):    Cause->Effect->Deepen / Claim->Evidence->Rebuttal /")
CODE("                  Problem->Solution->Feasible")
CODE("                  MUST include: 1 example + 1 long sentence + 1 concession")
CODE("Para 3 (Concl):   Prospect / Action / Reflection -> Elevate, don't repeat")

SEP()
AP("TRANSITION WORDS (Quick Reference)", bold=True, size=12)
CODE("Progression: Beyond that, | More significantly, | Equally noteworthy...")
CODE("Contrast:    That said, | Yet this alone is hardly the full picture.")
CODE("Concession:  Admittedly, ... | ..., legitimate as it is, ...  [<- bonus!]")
CODE("Cause:       This, in turn, leads to... | Small wonder, then, that...")
CODE("Summary:     In light of the above, | Ultimately, | Taking all this into account,")

SEP()
AP("ABSOLUTE BAN LIST", bold=True, size=12, color=RED)
CODE("[X] First... Second... Third...     <- Strongest template signal ever")
CODE('[X] Every coin has two sides.        <- Most overused cliche; graders cringe')
CODE("[X] With the development of society... <- Empty filler, wastes word count")
CODE("[X] That's all. Thank you.           <- NOT how CET essays conclude")
CODE("[X] Overstuffing however/moreover     <- Transitions are seasoning, not the meal")

SEP()
AP("SCORE TARGET FORMULA", bold=True, size=12, color=BLUE)
CODE("Clear structure + on-topic + no major errors      = 11/15")
CODE("Above + 1 concession-rebuttal + 1 long sentence")
CODE("      + 1 concrete example                        = 13-14/15")
CODE("Above + humanistic closing touch                  = 14-15/15")

SEP()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Framework in mind. Expression in hand.\nThe template is the skeleton; your thinking is the flesh.")
r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = RGBColor(*BLUE)

# ================================================================
# SAVE
# ================================================================
output = os.path.normpath(os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "CET4_Essay_Universal_Template.docx"))
doc.save(output)
print(f"Done: {output}")