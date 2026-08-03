# -*- coding: utf-8 -*-
"""Generate topic-specific filled templates for June 2026."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor

doc = Document()
style = doc.styles['Normal']
style.font.size = Pt(11.5)
style.paragraph_format.line_spacing = 1.45

BLUE = (0x1A, 0x56, 0xDB)
RED = (0xCC, 0x33, 0x33)
GRAY = (0x6B, 0x72, 0x80)
GREEN = (0x1F, 0x88, 0x54)

def AP(text, bold=False, size=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)

def CODE(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(2)

def SEP():
    p = doc.add_paragraph()
    p.add_run("─" * 55).font.size = Pt(6)

# ================================================================
AP("2026年6月四级 三大方向 · 成文弹药", bold=True, size=18, color=BLUE)
AP("直接背, 考场填关键词就能用", size=11, color=RED)
SEP()

# ================================================================
# TOPIC 1: AI
# ================================================================
AP("Topic 1  AI与学习 (概率最高)", bold=True, size=14, color=RED)

AP("万能例子 (必背)", bold=True, size=12, color=GREEN)
AP("A classmate of mine, who once spent hours drafting essays from scratch, now uses AI writing assistants to generate outlines and polish her grammar. Yet she makes a point of rewriting every AI-generated sentence in her own words -- treating the tool as a brainstorming partner rather than a shortcut. Her writing, she told me, has actually improved, because AI frees her to focus on ideas instead of mechanics.", size=10.5)

AP("让步反驳句", bold=True, size=12, color=GREEN)
CODE("Admittedly, over-reliance on AI could erode students' independent")
CODE("thinking. Yet this concern, legitimate as it is, overlooks a crucial")
CODE("point: the real threat is not AI itself, but the absence of guidance")
CODE("on how to use it responsibly.")

AP("长句", bold=True, size=12, color=GREEN)
CODE("What makes this issue particularly urgent is that AI's influence on")
CODE("education extends far beyond the classroom, reshaping not only how")
CODE("students acquire knowledge but also how they define the very meaning")
CODE("of learning in an age when answers are always one click away.")

AP("必背词汇", bold=True, size=12, color=GREEN)
AP("artificial intelligence / AI-powered tools / digital literacy / independent thinking / double-edged sword / harness the power of / algorithm-driven / academic integrity", size=10)

AP("完整成文", bold=True, size=12, color=GREEN)
CODE("In recent years, the issue of AI-assisted learning has sparked")
CODE("considerable discussion on university campuses. While this trend")
CODE("brings notable benefits, it also raises concerns that deserve")
CODE("serious attention.")
CODE("")
CODE("Several factors explain this phenomenon. At its root, the sheer")
CODE("convenience of AI-powered tools has created fertile ground for them")
CODE("to take hold. The effects are already visible: tasks that once took")
CODE("hours can now be completed in minutes, allowing students to focus on")
CODE("higher-level thinking. What deserves equal attention, however, is the")
CODE("risk of over-dependence -- a point often neglected in casual")
CODE("discussion. A classmate of mine, who once spent hours drafting essays")
CODE("from scratch, now uses AI assistants to generate outlines -- yet she")
CODE("always rewrites every suggestion in her own words, treating the tool")
CODE("as a partner rather than a crutch. Admittedly, the concern about")
CODE("eroding independent thinking carries weight. Yet this overlooks a")
CODE("crucial fact: the real issue is not AI itself, but how we choose to")
CODE("use it. What makes this particularly urgent is that AI's influence")
CODE("extends beyond the classroom, reshaping not only how students learn")
CODE("but also how they define the very meaning of learning.")
CODE("")
CODE("In light of the above, I am convinced that AI-assisted learning is")
CODE("not only here to stay but can be a powerful force for good --")
CODE("provided that clear guidelines and a culture of responsible use are")
CODE("cultivated. The path forward calls for a balance between embracing")
CODE("technological convenience and preserving independent thinking, a")
CODE("challenge today's students are ready to meet.")

doc.add_page_break()

# ================================================================
# TOPIC 2: Rural Revitalization
# ================================================================
AP("Topic 2  乡村振兴 (新热点)", bold=True, size=14, color=RED)

AP("万能例子 (必背)", bold=True, size=12, color=GREEN)
AP("A senior student I know turned down a corporate job offer in Beijing to return to his hometown in Yunnan. Using the digital marketing skills he learned at university, he now helps local tea farmers sell their products through livestreaming. His monthly income exceeds that of many of his urban classmates -- but more importantly, he told me, he feels he is building something that truly belongs to him.", size=10.5)

AP("让步反驳句", bold=True, size=12, color=GREEN)
CODE("Admittedly, returning to the countryside means giving up the")
CODE("conveniences and opportunities of big cities. Yet this concern,")
CODE("legitimate as it is, overlooks a deeper shift: with digital")
CODE("infrastructure reaching even remote villages, the urban-rural gap")
CODE("in career prospects is narrowing faster than most people realize.")

AP("长句", bold=True, size=12, color=GREEN)
CODE("What makes rural revitalization particularly meaningful is that it")
CODE("represents not a one-sided sacrifice by the young, but a genuine")
CODE("two-way journey -- one in which graduates bring skills and fresh")
CODE("ideas to the countryside while discovering career paths that no")
CODE("crowded city could offer them.")

AP("必背词汇", bold=True, size=12, color=GREEN)
AP("rural revitalization / entrepreneurship / digital skills / bridge the urban-rural gap / livestreaming commerce / a two-way journey / tap into local resources / remote villages", size=10)

AP("完整成文", bold=True, size=12, color=GREEN)
CODE("In recent years, the issue of rural revitalization has sparked")
CODE("considerable discussion across Chinese society. While this national")
CODE("drive brings notable benefits, it also raises a pressing question:")
CODE("how can young graduates play a meaningful role in this process?")
CODE("")
CODE("Several factors explain why this matters now more than ever. At its")
CODE("root, the rapid expansion of digital infrastructure into remote")
CODE("villages has created fertile ground for young talent to make a real")
CODE("impact. The effects are already visible: e-commerce and livestreaming")
CODE("have transformed how local products reach national markets, and more")
CODE("importantly, they have begun to rewrite the old narrative that")
CODE("success can only be found in big cities. A senior student I know")
CODE("turned down a corporate job in Beijing to help his hometown tea")
CODE("farmers sell through livestreaming -- his income now exceeds that of")
CODE("many urban classmates, but he says the real reward is building")
CODE("something that truly belongs to him. Admittedly, returning to the")
CODE("countryside means giving up urban conveniences. Yet this overlooks a")
CODE("deeper shift: the urban-rural gap in opportunity is narrowing faster")
CODE("than most realize. What makes this particularly meaningful is that")
CODE("rural revitalization is not a one-sided sacrifice but a genuine")
CODE("two-way journey -- graduates bring fresh ideas to the countryside")
CODE("while discovering careers no crowded city could offer.")
CODE("")
CODE("In light of the above, I am convinced that engaging with rural")
CODE("development is not only meaningful but a genuinely smart career move")
CODE("-- provided that universities equip students with the digital and")
CODE("entrepreneurial skills such work demands. The path forward calls for")
CODE("a vision in which the countryside is seen not as a fallback, but as")
CODE("a new frontier, a challenge today's graduates are ready to embrace.")

doc.add_page_break()

# ================================================================
# TOPIC 3: Traditional Culture
# ================================================================
AP("Topic 3  传统文化传承与创新 (常青树)", bold=True, size=14, color=RED)

AP("万能例子 (必背)", bold=True, size=12, color=GREEN)
AP("Last semester, I joined a campus club that makes short videos about traditional Chinese festivals. For the Dragon Boat Festival, we filmed the entire process of making zongzi -- from soaking bamboo leaves to the final unwrapping -- and posted it on Douyin. To our surprise, the video received over 50,000 views, with comments from young viewers saying they finally understood what the festival was actually about beyond a day off.", size=10.5)

AP("让步反驳句", bold=True, size=12, color=GREEN)
CODE("Admittedly, traditional practices can feel distant from the fast-paced")
CODE("lives of today's young people. Yet this concern, legitimate as it is,")
CODE("overlooks a crucial point: the problem is not that tradition is")
CODE("irrelevant, but that it has rarely been presented in a language")
CODE("young people actually speak -- and that is exactly where digital")
CODE("tools can make all the difference.")

AP("长句", bold=True, size=12, color=GREEN)
CODE("What makes cultural heritage particularly worth protecting is that")
CODE("it represents far more than old artifacts or customs -- it is the")
CODE("collective memory of a people, a thread connecting past and present")
CODE("that, once severed, can never be fully rewoven.")

AP("必背词汇", bold=True, size=12, color=GREEN)
AP("cultural heritage / cultural confidence / intangible cultural heritage / breathe new life into / bridge tradition and modernity / time-honored / collective memory / pass down", size=10)

AP("完整成文", bold=True, size=12, color=GREEN)
CODE("In recent years, the issue of preserving traditional culture has")
CODE("sparked considerable discussion, particularly regarding how young")
CODE("people can play a role in keeping heritage alive in the digital age.")
CODE("While this effort faces real challenges, it also opens up creative")
CODE("possibilities that previous generations never had.")
CODE("")
CODE("Several factors explain why this matters. At its root, a growing")
CODE("sense of cultural confidence among young Chinese has created fertile")
CODE("ground for tradition to be rediscovered and reinvented. The effects")
CODE("are already visible: hanfu clubs on campus, traditional crafts")
CODE("showcased on Douyin, and museum exhibitions that draw massive young")
CODE("crowds. What deserves equal attention, however, is that much of our")
CODE("intangible heritage still risks fading away -- a point often")
CODE("neglected in casual discussion. Last semester, I joined a campus club")
CODE("that makes short videos about traditional festivals. For the Dragon")
CODE("Boat Festival, we filmed the making of zongzi from start to finish.")
CODE("To our surprise, the video received over 50,000 views, with young")
CODE("viewers commenting that they finally understood what the festival was")
CODE("really about. Admittedly, traditional customs can feel distant from")
CODE("modern life. Yet this overlooks a crucial point: the issue is not")
CODE("that tradition is irrelevant, but that it has rarely been presented")
CODE("in a language young people actually speak. What makes cultural")
CODE("heritage worth protecting is that it represents far more than old")
CODE("customs -- it is a thread connecting past and present that, once")
CODE("severed, can never be fully rewoven.")
CODE("")
CODE("In light of the above, I am convinced that preserving traditional")
CODE("culture is not only a duty but a genuinely creative opportunity --")
CODE("provided that young people take the lead in making heritage relevant")
CODE("to the digital age. The path forward calls for a blend of respect")
CODE("for the past and courage to innovate, a challenge today's students")
CODE("are more than ready to embrace.")

doc.add_page_break()

# ================================================================
# 通用弹药
# ================================================================
AP("通用弹药 (三个方向共享)", bold=True, size=14, color=BLUE)

AP("任何题目都能塞的万能例子框架:", bold=True, size=12, color=GREEN)
CODE("A [classmate/friend/senior] of mine, who once [过去的困难/状态],")
CODE("now [改变后的做法]. [具体成果/数据]. [一句感悟], [名字] told me,")
CODE("[直接引语 -- 这句话让例子有真实感].")

AP("任何题目都能塞的让步反驳:", bold=True, size=12, color=GREEN)
CODE("Admittedly, [对方合理之处] carries some weight. Yet this concern,")
CODE("legitimate as it is, overlooks [你的核心反驳].")

AP("任何题目都能塞的长句:", bold=True, size=12, color=GREEN)
CODE("What makes [话题] particularly [adj.] is that it [动词短语],")
CODE("[现在分词短语补充说明] -- [破折号后点睛].")

AP("衔接词库 (轮换用, 别重复):", bold=True, size=12, color=GREEN)
CODE("递进: Beyond that, | More significantly, | What deserves equal attention is...")
CODE("转折: Yet this alone is hardly the full picture. | That said,")
CODE("因果: This, in turn, leads to... | Small wonder, then, that...")
CODE("总结: In light of the above, | Ultimately, | Taking all this into account,")

AP("绝对禁用:", bold=True, size=12, color=RED)
CODE("First... Second... Third...    Every coin has two sides.")
CODE("With the development of society...    That's all. Thank you.")

SEP()
AP("考前一晚：三个方向的例子各默写1遍 + 框架过1遍。够了。", bold=True, size=13, color=BLUE)

output = os.path.normpath(os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "CET4_三大方向_成文弹药.docx"))
doc.save(output)
print(f"Done: {output}")