#!/usr/bin/env python3
"""Build clean DOCX from scratch. No fragmentation, consistent formatting."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree
import copy

# Create new document
doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

def H1(text):
    """Section heading"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    return p

def H2(text):
    """Subsection heading"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    return p

def P(text):
    """Body paragraph"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    return p

def empty():
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Sustained-Release Tracer Proppant with Transport-Based Breakthrough Curve Interpretation for Production Monitoring in Fractured Wells")
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'
empty()

# ═══════════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════════
P("Abstract: Per-interval production monitoring in fractured unconventional wells relies on tracer-based methods, yet interpreting the breakthrough curve of a sustained-release tracer proppant to extract quantitative production information remains an unresolved challenge. Here we address this by developing a two-component transport model that decomposes the breakthrough curve into a Gaussian pulse, representing the shut-in accumulation slug, and an erfc tail, representing sustained matrix-diffusion-controlled release, joined by a smooth tanh weighting function. The model was validated using an oleophilic epoxy/Fe3O4 sustained-release tracer proppant (ESP-T) synthesized by emulsion polymerization of stearic acid-modified nano-Fe3O4@SA. Fitted to single-phase displacement data, the model achieves R2 = 0.9939 and recovers the effective production flow rate Q within 8% of the independently set pump value (0.46 vs. 0.50 mL/min). The erfc tail accounts for 47% of the integrated tracer signal. ESP-T exhibits thermal stability to 357 degC, a water contact angle of 104.6 deg (vs. 72.3 deg for neat epoxy), and a 66% reduction in oil filtration time, consistent with a water-resistant, oil-permeable transport character. Under steady-state two-phase flow, tracer flux tracks the oil production rate across oil-water ratios from 4:1 to 1:4. The combined experimental-modeling approach provides a framework for per-interval production monitoring without downhole instrumentation.")
empty()
P("Key words: Unconventional oil and gas reservoirs; Hydraulic fracturing; Tracer proppant; Epoxy resin; Fe3O4 nanoparticles; Advection-dispersion model; Piecewise modeling; Production allocation; Release kinetics")
empty()

# ═══════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════
H1("1 Introduction")

P("Unconventional oil and gas resources—shale oil, tight oil, and coalbed methane—account for more than half of global hydrocarbon reserves [1-4]. Their development depends on hydraulic fracturing, in which high-pressure fluid creates artificial fractures in low-permeability formations, providing conductive pathways for hydrocarbons to reach the wellbore [5,6]. A typical horizontal well may be fractured in twenty or more stages along its lateral section.")

P("After stimulation, the contribution of each fractured interval to total production varies considerably, and individual intervals decline at different rates [7]. Operators currently lack a practical method for measuring how much oil each interval produces over time. Production logs require well intervention; distributed fiber-optic sensing is expensive and restricted to new completions. A method that quantifies per-interval production from surface samples alone, without downhole tools, would directly inform refracturing candidate selection, infill drilling decisions, and stimulation design for subsequent wells.")

P("Tracer-based methods offer this capability. A chemically distinct tracer is added to each fracturing stage; the tracer appears in the produced fluid, and its concentration history at the wellhead carries information about that stage's flow contribution. Sustained-release tracer proppants [13-16] extend the monitoring window from hours to weeks by immobilizing the tracer within a solid carrier that is co-injected with the proppant and releases it gradually. However, a sustained-release proppant creates an interpretation problem not encountered with instantaneous tracer injection. The breakthrough curve recorded at the wellhead is the sum of two overlapping contributions: a concentration pulse from tracer that accumulated in the proppant pack during the shut-in period, and a long, slowly decaying tail from tracer that continues to release from the carrier matrix throughout production. These contributions cannot be separated by inspecting the curve. Current practice characterizes release with the empirical Korsmeyer-Peppas power law [17,18], which describes how much tracer leaves the carrier but provides no route to interpreting the wellhead signal. ADE-based fitting has been used to extract fracture geometry [19]; tracer mass-balance methods allocate production without exploiting breakthrough curve shape [20]. A transport-based framework that resolves the curve into its component contributions and recovers a production rate from the fitted parameters has received limited attention [21-26].")

P("Building such a framework requires two things: a transport model capable of decomposing the composite breakthrough curve into physically meaningful components, and a tracer proppant whose release behavior is dominated by matrix diffusion—so that the sustained tail is well developed—and whose surface is oleophilic, so that the released tracer partitions selectively into the oil phase. Existing tracer proppants do not meet both requirements. Coated proppants lose their tracer function when the coating dissolves; polystyrene microspheres lack thermal stability above approximately 200 degC [27-30]. Epoxy resin offers a compelling alternative: its highly cross-linked network provides thermal stability exceeding 350 degC, chemical resistance, and tunable mechanical properties [31,32]. Emulsion polymerization consolidates nanoparticle modification and tracer encapsulation into a single synthetic step. Epoxy-encapsulated water-soluble tracers have been demonstrated [33,34]; an oleophilic variant for oil-phase monitoring has not been reported.")

P("In this work, we develop a transport model that interprets the breakthrough curve of a sustained-release tracer proppant by decomposing it into a Gaussian pulse (the shut-in accumulation slug) and an erfc tail (the sustained matrix-diffusion-controlled release), joined by a smooth tanh transition. The decomposition recovers the effective production flow rate Q directly from the tracer signal. We synthesize an oleophilic epoxy/Fe3O4 sustained-release tracer proppant (ESP-T) by emulsion polymerization, using stearic acid-modified nano-Fe3O4@SA as the oleophilic tracer. We characterize its microstructure, thermal stability, wettability, oil-water transport selectivity, and mechanical integrity (Sections 3.1-3.5); quantify its temperature-dependent release kinetics via the Korsmeyer-Peppas model (Section 3.6); apply the transport model to recover production rates from single-phase breakthrough data (Section 3.7); and extend the analysis to steady-state oil-water two-phase flow using a tracer-flux methodology (Section 3.7). The study proceeds from material design through release characterization to transport modeling and production-rate quantification.")
empty()

# ═══════════════════════════════════════════════════════════════════
# 2. EXPERIMENTS
# ═══════════════════════════════════════════════════════════════════
H1("2 Experiments")
H2("2.1 Materials")
P("All chemical reagents used in the experiments were of analytical grade or industrial grade and used without further purification. Analytical Grade (AR) reagents, including ethanol (purity >= 99.7%), ferric chloride (FeCl3, purity >= 99.0%), ferrous chloride tetrahydrate (FeCl2.4H2O, purity >= 99.0%), manganese chloride hexahydrate (MnCl2.6H2O, purity >= 99.0%), stearic acid (purity >= 99.5%) and silicon dioxide (SiO2, purity >= 99.99%), were purchased from Chengdu Kelong Chemical Co., Ltd. Industrial grade reagents included guar gum (purity 90%-95%), E51 epoxy resin (epoxy value meeting industrial standards, purity >= 99%), T31 curing agent (effective component content 95%), and hollow glass microspheres (main component purity >= 99%). The detailed specifications and sources of all materials are summarized in Table 2-1.")
P("[Table 2-1] Specifications and sources of experimental materials")
empty()

H2("2.2 Fabrication of stearic acid-modified nano-Fe3O4 (nano-Fe3O4@SA)")
P("A clean three-necked flask was mounted in a thermostatic water bath and connected to nitrogen purge and mechanical stirring apparatuses. The flask was charged with 100 mL deionized water and heated to 80 degC. Upon reaching the target temperature, 2.703 g (0.01 mol) FeCl3 was added, and the system was purged with nitrogen for 15 min, followed by the introduction of 1.15 g (0.058 mol) FeCl2.4H2O and 2x10-5 mol MnCl2.6H2O. After stirring for 10 min, 5.5 mL ammonia solution was rapidly added, and the reaction proceeded at pH 10 for 2 h. Upon completion, the resulting black suspension was transferred to an Erlenmeyer flask and magnetically separated, with the supernatant discarded. The precipitate was washed with anhydrous ethanol under sonication repeatedly until neutral pH, then collected. Finally, the precipitate was blended with an ethanolic stearic acid solution and subjected to sonication for oleophilic modification. After washing, the product was diluted to 100 mL and stored for subsequent use [21].")
P("[Figure 2-1] Schematic illustration for the preparation of stearic acid-modified metal-doped Fe3O4 nanoparticles")
empty()

H2("2.3 Preparation of Epoxy Resin Microspheres and ESP-T")
P("A SiO2 aqueous dispersion was prepared by adding 0.3 g of 20 nm SiO2 to 150 mL deionized water under stirring at 380 RPM. Separately, 0.9 g guar gum and 0.5 g of 45 um SiO2 were weighed. An epoxy pre-mixture consisting of 20 mL E51 epoxy resin, 6 g of Fe3O4-ethanol mixture, 1 g hollow glass microspheres, and 7 g T31 curing agent was homogenized by manual stirring for 3 min. The pre-mixture was transferred into the SiO2 dispersion and agitated for 1 min to complete emulsification, followed by the addition of 45 um SiO2 and guar gum to stabilize the system. After microsphere formation, the product was cured at 50 degC for 1 h, rinsed with deionized water, and dried at 80 degC for 10 h. For pure epoxy microspheres, 6 g of pure ethanol replaced the Fe3O4-ethanol mixture and the guar gum dosage was 0.9 g.")
P("[Figure 2-2] Schematic illustration for the preparation of Fe3O4-encapsulated epoxy resin proppants")
empty()

H2("2.4 Characterization Methods")
P("The microstructure, elemental composition and distribution of neat epoxy resin microspheres and ESP-T were analyzed by scanning electron microscopy (SEM, ZEISS-Sigma 500). Samples were sputter-coated with a thin conductive metal layer under vacuum, mounted onto conductive adhesive tape, and observed at magnifications ranging from low power up to 50,000x. Characteristic X-rays were excited by a field-emission electron gun, and two-dimensional elemental mapping was performed over a selected region to visualize the distribution of Fe element in the proppants.")
P("Surface morphology and particle dispersion of neat epoxy resin microspheres and ESP-T were examined using an optical microscope (Leica DM2700P). Thermogravimetric analysis (TGA, TA Instruments Q500, USA) was conducted to evaluate the thermal stability of the two proppants. Tests were performed in an air atmosphere at a heating rate of 10 degC/min over a temperature range of 50-800 degC.")
P("Water contact angle (WCA) measurements were carried out using a video optical contact angle analyzer (OCA20, Germany). Pressed proppant flakes were placed on the test stage, and a 5 uL droplet of deionized water was deposited onto the sample surface. Images were captured after the droplet stabilized, and five replicate measurements were performed for each sample.")
P("The physical and mechanical properties as well as field applicability of the proppants were comprehensively evaluated using roundness, sphericity, bulk density, apparent density, acid solubility, and crush rate. The oil-water conductivity of the proppants was indirectly assessed via filtration time. The bottom of a funnel was blocked with a 200-mesh screen to prevent proppant leakage while allowing water or oil to pass through. A total of 2.0 g of proppant was loaded into the funnel, followed by the addition of 20 mL deionized water or dodecane, and the time required for complete permeation was recorded.")
empty()

H2("2.5 Tracer Release Behavior at Different Temperatures")
P("Glass vials containing 100 mL dodecane were placed in thermostatic oil baths set at 30, 60, 90, and 120 degC, respectively, followed by the addition of 5 g of 40-70 mesh ESP-T. The vials were then hermetically sealed to prevent dodecane volatilization. Sampling was conducted at 12 h intervals, and the concentration of tracer-doped metal ions in the dodecane medium was quantified via inductively coupled plasma mass spectrometry (ICP-MS, PerkinElmer NexION 300X), from which the time-dependent tracer release curves were plotted.")
P("[Figure 2-3] Schematic diagram of the tracer release apparatus at different temperatures")
empty()

H2("2.6 Core Displacement Experiments for Production Monitoring")
P("For the single-phase oil production monitoring experiment, dodecane was employed to simulate crude oil, with the focus on the steady oil production regime of a single fractured interval. A dynamic displacement apparatus was utilized to mimic the dynamic oil production behavior of horizontal wells, as illustrated in Figure 2-4. The pore spaces of a steel core (diameter 2.5 cm, length 5.0 cm) were packed with ESP-T, and both ends were sealed with 200-mesh metal screens to prevent proppant migration. The proppant-packed core was placed into a core holder, and a confining pressure of 5 MPa was applied. Dodecane was pumped at 5 mL/min until the tubing string was saturated, after which the outlet valve was closed and the system left undisturbed for 96 h. Thereafter, dodecane was injected at 0.5 mL/min. Sampling was conducted at 4-min intervals, with 2 mL of effluent collected per sample for 20 sets. The concentrations of the labeled elements were quantified by ICP-MS.")
P("[Figure 2-4] Schematic diagram of the single-phase oil flow experimental apparatus")
P("For the oil-water two-phase flow experiment, the tracer was maintained in a stable release regime throughout the steady-flow period without shut-in. The total two-phase flow rate and oil-water ratio were modulated by adjusting the injection rates of dodecane and water. Three oil-water volume ratios (4:1, 1:1, 1:4) and four total two-phase flow rates (0.1, 0.2, 0.3, 0.4 mL/min) were employed (Figure 2-5). The proppant-packed core was placed under 5 MPa confining pressure, and dodecane and water were co-injected. Sampling was carried out at 5-min intervals, with 2 mL of effluent collected per sample for 20 sets. The volume ratio of the dodecane phase to the aqueous phase was recorded, and the concentrations of the labeled metal elements were quantified by ICP-MS.")
P("[Figure 2-5] Schematic diagram of the oil-water two-phase flow experimental apparatus")
empty()

# ═══════════════════════════════════════════════════════════════════
# 3. RESULTS AND DISCUSSION
# ═══════════════════════════════════════════════════════════════════
H1("3 Results and Discussion")

# 3.1 SEM
H2("3.1 SEM Characterization")
P("ESP-T microspheres are well-formed and highly uniform. SEM imaging at low magnification (Figure 3-1, panels a and d) reveals excellent sphericity and no inter-particle agglomeration in both pure epoxy and ESP-T samples, confirming that the emulsion polymerization process accommodates the nano-Fe3O4@SA filler without disruption. The most visible difference is in surface texture: pure epoxy surfaces are smooth with occasional particulate debris, whereas ESP-T surfaces are uniformly rough, providing the first indication that the nanofiller has been incorporated throughout the particle.")
P("At higher magnification (panels b and e), the differences sharpen. Pure epoxy microspheres display subtle wrinkles and shallow depressions, artefacts of volume shrinkage during thermoset curing. ESP-T shows a markedly different morphology: the surface is densely covered with micron- and submicron-scale protrusions distributed as discrete islands rather than as coalesced masses. This morphology indicates that nano-Fe3O4@SA disperses as nanoclusters within the epoxy matrix, and that stearic acid modification has successfully prevented the uncontrolled macroscopic phase separation that would otherwise produce large agglomerates.")
P("The highest-magnification images (panels c and f) reveal how the nanofiller interacts with the matrix. Pure epoxy exhibits the dense, undulating lava-like texture expected of a cross-linked polymer network. ESP-T shows a debris-like surface morphology in which individual nanoclusters, wrapped in stearic acid, can be distinguished. The key observation concerns the interface: no cracks or gaps separate these nanoclusters from the underlying epoxy. They appear embedded rather than merely attached. This tight integration is direct structural evidence that the long alkyl chains of stearic acid form physical entanglements and hydrophobic interactions with the epoxy chains, producing the robust interfacial adhesion needed for a durable composite proppant.")
P("Taken together, the images suggest a formation pathway in which hydrophobically modified nano-Fe3O4@SA nanoclusters disperse uniformly within the emulsified epoxy droplets and become immobilized as the polymer network cross-links. During curing shrinkage, a fraction of these clusters is extruded toward the particle surface, producing the surface-enriched rough nanostructures visible in panels c and f. We note that this mechanism is inferred from post-cure microscopy; real-time in-situ characterization would be required to confirm the kinetic pathway directly.")
P("[Figure 3-1] SEM images of epoxy microspheres and Fe3O4@epoxy proppants. The epoxy microspheres are at the top and the Fe3O4@epoxy proppants at the bottom. (a) and (d): multiple whole microspheres; (b) and (e): single whole microsphere; (c) and (f): local high-magnification views.")
P("Elemental mapping by EDS (Figure 3-2) confirms what the morphology suggests. The Fe signal is distributed across the entire particle cross-section, demonstrating that nano-Fe3O4@SA resides within the epoxy matrix rather than merely decorating the surface. Minor inhomogeneities in signal intensity point to distinct phase domains within the composite, consistent with the nanocluster dispersion observed in SEM. The imaging and elemental data together establish that the designed ESP-T structure has been successfully realized.")
P("[Figure 3-2] Mapping images of Fe3O4@epoxy resin proppant. (a) SEM image, (b) Fe distribution, (c) Si distribution.")
empty()

# 3.2 Thermal
H2("3.2 Thermal Stability")
P("Thermal analysis (Figure 3-3) identifies three decomposition stages. The first, 50-350 degC, involves a minor weight loss of 5.70%, attributable to adsorbed water and residual ethanol. The second stage, 350-400 degC, constitutes the primary decomposition: cleavage of C-O-C and C-C bonds in the epoxy network releases CO2 and small hydrocarbons, with the DTG curve peaking at 357.27 degC and accounting for 72.5% of the total mass loss. Above 400 degC the residual mass stabilizes; further weight loss derives from oxidative combustion of the carbonaceous residue. The final residue comprises hollow glass microspheres, thermally stable nano-Fe3O4@SA, and a minor carbonaceous fraction. A brief DSC endotherm accompanies the decomposition, consistent with the DTG profile.")
P("The initial decomposition temperature of 357.27 degC is far above the temperature range encountered in oil and gas wells (typically 80-150 degC, reaching 200 degC in deep wells). Below 200 degC the material loses only trace moisture, with no detectable degradation of the epoxy matrix. This thermal margin indicates that the epoxy matrix remains structurally intact across the full operational temperature range of oil and gas wells.")
P("[Figure 3-3] TGA/DTG and DSC curves of Fe3O4@epoxy resin proppant")
empty()

# 3.3 WCA
H2("3.3 Water Contact Angle (WCA)")
P("The morphological change documented by SEM corresponds to a pronounced shift in surface wettability. Water contact angle measurements (Figure 3-4) give an average of 72.3 deg for pure epoxy microspheres, a value consistent with the hydroxyl groups present on epoxy chains. ESP-T yields 104.6 deg, a 32.3 deg increase that crosses the hydrophobic threshold. The mechanism behind this transition is the stearic acid modification of the nano-Fe3O4@SA: carboxyl groups coordinate with surface hydroxyls on the Fe3O4, while the C17 alkyl chains orient outward, constructing a low-energy hydrophobic film. The surface enrichment of these chains, driven by the nanocluster extrusion documented in SEM, amplifies the effect by concentrating the hydrophobic moieties at the proppant surface.")
P("[Figure 3-4] (a) WCA of epoxy resin patch, (b) WCA of patch after doping with nano-Fe3O4@SA")
empty()

# 3.4 Density
H2("3.4 Physical and Mechanical Properties")
P("Both proppants are lighter than water. Pure epoxy microspheres have a bulk density of 0.6179 g/cm3 and an apparent density of 1.02 g/cm3 (Figure 3-5); ESP-T yields 0.646 and 1.072 g/cm3, respectively. The increase reflects the higher density of nano-Fe3O4@SA (approximately 5.18 g/cm3) relative to epoxy (approximately 1.1 g/cm3). Both values remain below the density of water, meaning the proppants can be suspended in water-based fracturing fluids without additional agents. The 0.05 g/cm3 increase in apparent density is consistent with nanoclusters filling internal pores, producing a more compact particle structure. ESP-T meets key industry specifications: sphericity and roundness exceed 0.9 (Krumbien-Sloss chart [22]), acid solubility is 3.3% (well below the 5% threshold), and the crush rate at 52 MPa is 2.9%, comparable to pure epoxy microspheres (2.6%). The slight increase in crush rate is attributable to the hollow glass microspheres; the epoxy matrix itself retains its structural integrity after nano-Fe3O4@SA incorporation. Both values remain well below the typical benchmark for ultra-lightweight proppants (<10% at 52 MPa per SY/T 5107-2016).")
P("[Figure 3-5] Bulk density and apparent density of epoxy microspheres and ESP-T")
empty()

# 3.5 Conductivity
H2("3.5 Proppant Pack Conductivity")
P("The wettability reversal produces a corresponding reversal in transport behavior. Filtration tests (Figure 3-6, Table 3-1) quantify the contrast: water passes through a pure epoxy pack in under three minutes, but requires nearly half an hour to percolate through ESP-T under identical conditions. For dodecane the situation inverts: pure epoxy passes oil in just over 15 min; ESP-T does so in 5 min, a 66% reduction. The hydrophobic surface repels water from inter-particle pores while allowing oil to spread freely. In a formation producing both phases, this selectivity favors oil flow toward the wellbore while suppressing water breakthrough, a characteristic with direct implications for production management in maturing wells.")
P("[Figure 3-6] Oil-water permeability test of epoxy resin microspheres and ESP-T. (a,b) water before and after pure epoxy; (c,d) oil before and after pure epoxy; (e,f) water before and after ESP-T; (g,h) oil before and after ESP-T.")
P("[Table 3-1] Oil and Water Passage Time")
empty()

# 3.6 Release
H2("3.6 Tracer Release Behavior at Different Temperatures")
P("Reservoir temperature increases with depth, and because tracer release from the epoxy matrix is diffusion-controlled, temperature directly affects the release rate. Figure 3-7(a) shows the release profiles of ESP-T in dodecane at 30, 60, 90, and 120 degC, plotted as normalized concentration C/C0 versus time. Release accelerates with temperature at all time points. The release rate is highest during the first 24 hours and decays gradually (Figure 3-7c). At 120 degC, cumulative release over 14 days remains well above the ICP-MS detection threshold, confirming that ESP-T sustains a measurable tracer signal over timescales relevant to long-term monitoring.")
P("Trace release kinetics were analyzed using the Korsmeyer-Peppas (K-P) model, C/C0 = K tn, where C is the concentration at time t, C0 the maximum release concentration, K the kinetic rate constant, and n the diffusion exponent [23]. For spherical carriers, the exponent n distinguishes the dominant mechanism: n <= 0.43 indicates Fickian-diffusion-controlled release, 0.43 < n < 0.85 indicates anomalous transport co-governed by diffusion and polymer relaxation, and n >= 0.85 indicates Case-II-relaxation-controlled release [24].")
P("Fitted to the release data at each temperature (Figure 3-7b, Table 3-2), the K-P model gives rate constants K that increase systematically from 0.055 at 30 degC to 0.196 at 120 degC, consistent with the thermally activated nature of both diffusion and polymer relaxation. All fitted n values fall within 0.45-0.85, the range corresponding to anomalous transport. R2 exceeds 0.94 at every temperature. We note that the K-P model, as a power-law formulation, lacks an upper asymptote and is strictly valid only for the early-to-intermediate release regime (Mt/Minf < 0.6). The extrapolated C/C0 values at 14 days exceed unity, reflecting this inherent limitation; the reported R2 values therefore pertain to the 0-14 day fitting range and should not be extrapolated further.")
P("The release mechanism inferred from these data involves solvent-driven swelling. Dodecane molecules permeate into the cross-linked epoxy network, generating a dual-state structure with an inner glassy core and an outer gel layer. Swelling reduces polymer chain entanglement, creating expanded transport channels through which the tracer diffuses into the external medium. Higher temperature enhances solvent permeability, accelerates swelling, and weakens intermolecular forces within the network, enlarging the pore dimensions available for tracer diffusion. This mechanism accounts for both the temperature dependence of K and the mixed Fickian/Case-II character reflected in the n values.")
P("[Figure 3-7] (a) Release curves of ESP-T at different temperatures, (b) K-P model fitting curves, (c) Release rate change curves of ESP-T")
P("[Table 3-2] Fitting Parameters of K-P Model")
empty()

# 3.7 Core
H2("3.7 Tracer Breakthrough Modeling and Production Rate Quantification")
P("The breakthrough curve of a sustained-release tracer proppant contains two overlapping signals: a concentration pulse from tracer accumulated in the proppant pack during shut-in, and a persistent tail from tracer that continues to release from the carrier matrix throughout production. Because these contributions overlap in time, neither can be isolated by inspection.")
P("We separate them by modeling transport explicitly using the one-dimensional advection-dispersion equation, dC/dt + v dC/dx = D d2C/dx2, where v = 4Q/(pi d2) is the flow velocity and D = alpha v the dispersion coefficient. The velocity v connects the model directly to the volumetric production rate Q. For the shut-in slug we apply the classical instantaneous-injection solution [25], which yields a Gaussian pulse: C_rise = (M/A)/sqrt(4pi D t) exp[-(x-vt)2/(4Dt)]. For the sustained release we apply the continuous-source solution, which yields a complementary error function: C_fall = (C0/2) erfc[(x-vt)/sqrt(4Dt)]. The two solutions are joined through a hyperbolic-tangent weighting function, w(t) = 1/2 [1 + tanh((t0-t)/sigma)], producing a smooth transition from the pulse-dominated early period to the diffusion-dominated tail. The full composite model, C(t) = cb + w(t) C_rise + [1-w(t)] C_fall, is Eq. (2). The parameter sigma is fixed at the sampling interval (4 min); the six free parameters are cb (tail-plateau concentration), A (slug amplitude), a (sustained-release amplitude), alpha (longitudinal dispersivity), Q (effective flow rate), and t0 (crossover center time).")
P("The recovery of Q from the breakthrough curve follows from the model structure. The flow velocity v appears in both components: it determines the arrival time and width of the Gaussian pulse, and the decay rate of the erfc tail. Because Q is the only parameter shared by both components, it is constrained by two independent portions of the data—the pulse shape and the tail decay—making its fitted value well-determined. The model was not given the pump rate as an input; it recovered Q from the shape of the breakthrough curve.")
P("Fitted to the single-phase breakthrough curve obtained with ESP-T (Figure 3-8b, Table 3-3), the model achieves R2 = 0.9939 with RMSE = 0.0210. The recovered flow rate Q = 0.46 mL/min agrees with the independently set pump rate of 0.50 mL/min within 8%. The fitted mean residence time MRT = 37.4 min matches the convective travel time x/v = 38.6 min. The Peclet number Pe = x/alpha = 0.934 places the transport at the advection-dispersion transition, consistent with a gradual-release source. Integrating the two components separately shows that 47% of the total integrated signal originates from the erfc tail. Nearly half the tracer detected at the wellhead during the sampling period comes from sustained release rather than from the shut-in slug.")
P("Q, MRT, and Pe, obtained from a single fit, are each consistent with independently known values: the pump setting, the tubing geometry, and the release kinetics from Section 3.6. This internal consistency indicates that the model recovers physical quantities rather than accommodating noise. The 47% tail contribution captures a key feature of sustained-release proppants: the signal does not return to baseline after the slug passes because release continues. The tail-plateau concentration cb reflects the steady-state release flux; because ESP-T release is thermally activated (K rising from 0.055 at 30 degC to 0.196 at 120 degC), cb is expected to scale with reservoir temperature. This interpretation approach differs from prior ADE-based analyses of tracer data, which have extracted fracture geometry [19], performed mass-balance allocation [20], or characterized fracture networks [26]. In those studies, transport modeling served reservoir description. Here, the transport decomposition itself yields the production metric, by resolving the breakthrough curve into contributions that a sustained-release proppant generates.")
P("[Table 3-3] Fitted Parameters of the Piecewise ADE Model and Derived Transport Properties")
P("[Figure 3-8] (a) Tracer release production curve of ESP-T single-phase flow monitoring, (b) Fitting curve of ESP-T tracer release model, (c) Comparison of pump set flow rate and fitted flow rate")
P("The single-phase analysis demonstrates the model under idealized conditions. To test whether it remains valid when both oil and water flow through the proppant pack, we conducted steady-state displacement experiments at three oil-water ratios (OWR = 4:1, 1:1, 1:4) and four total flow rates (0.1-0.4 mL/min). At OWR = 1:1, the early effluent showed a higher oil fraction, consistent with ESP-T's oil-permeable character established in Section 3.5.")
P("Figure 3-9(a) shows that tracer concentration decreases with increasing total flow rate due to dilution and is largely independent of OWR, consistent with competitive wetting of the proppant surface by the two phases. Concentration alone is therefore a poor proxy for the oil production rate, because it confounds release with dilution. We resolve this by defining the tracer flux FO as the mass of tracer passing the sampling point per unit time, the product of the oil-phase concentration and the oil volumetric flow rate. FO equals the release rate from the ESP-T pack [8] and is independent of dilution. Figure 3-9(b) confirms that FO increases with OWR, as expected when more oil contacts the proppant, but remains constant across total flow rates at fixed OWR. The release rate is governed by the oil-wetted area (Section 3.3), not by the flow velocity through the pack.")
P("Figure 3-9(c) normalizes FO by the steady-state value from single-phase oil displacement (3.187 ug/min) and compares the result with the known oil flow rate at each OWR. The agreement confirms that the oil production rate of a labeled interval can be recovered from FO when the total flow rate is known, even with a water phase present. Several boundaries apply. The model assumes one interval with uniform packing; real wells have multiple intervals whose signals overlap. The FO calibration is steady-state; transient conditions during start-up, shut-in, or rapid drawdown may cause deviations. Dodecane is a model oil; crude oil components, particularly asphaltenes, may alter wetting and diffusion behavior. The long-term chemical stability of the epoxy matrix and stearic acid coating above 120 degC, or in the presence of H2S, CO2, and high-salinity brines, remains to be evaluated.")
P("[Figure 3-9] (a) Tracer concentration changes with oil-water ratio and total two-phase flow, (b) Tracer flux changes with oil-water ratio and total two-phase flow, (c) Comparison of tracer flux and actual flow under different oil-water ratios at 0.1 mL/min")
empty()

# ═══════════════════════════════════════════════════════════════════
# 4. CONCLUSIONS
# ═══════════════════════════════════════════════════════════════════
H1("4. Conclusions")
P("This work addresses the interpretation of breakthrough curves from sustained-release tracer proppants. A two-component transport model was developed that decomposes the curve into a Gaussian pulse (the shut-in accumulation slug) and an erfc tail (the sustained matrix-diffusion-controlled release), joined by a smooth tanh transition. An oleophilic epoxy/Fe3O4 sustained-release tracer proppant (ESP-T) was synthesized by emulsion polymerization and used to validate the model in single-phase post-shut-in flow and steady-state two-phase flow.")
P("The model recovers the effective production flow rate Q from the tracer signal (fitted 0.46 vs. pump-set 0.50 mL/min, 8% error). The erfc tail accounts for 47% of the integrated signal, confirming that a single shut-in period suffices for sustained monitoring. The fitted parameters are internally consistent with independently known experimental values.")
P("ESP-T provides a practical sustained-release platform: thermal stability to 357 degC, oleophilic surface (WCA 104.6 deg, 66% reduction in oil filtration time), and self-suspension in fracturing fluids (bulk density 0.646 g/cm3). Under steady-state two-phase flow, tracer flux tracks the oil production rate across oil-water ratios from 4:1 to 1:4.")
P("The validation is limited to laboratory-scale, single-interval experiments with dodecane as the model oil. Field application requires testing under multi-interval, multiphase conditions with crude oil. The present work establishes the interpretation framework and the experimental techniques needed for such testing.")
empty()

# ═══════════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════════
H1("References")
refs = [
    "[1] MONTGOMERY S L, JARVIE D M, BOWKER K A, et al. Mississippian Barnett Shale, Fort Worth basin, north-central Texas: Gas-shale play with multi-trillion cubic foot potential [J]. AAPG Bulletin, 2005, 89(2): 155-175.",
    "[2] WANG L, TIAN Y, YU X, et al. Advances in improved/enhanced oil recovery technologies for tight and shale reservoirs [J]. Fuel, 2017, 210: 425-445.",
    "[3] MOHR S H, WANG J, ELLEM G, et al. Projection of world fossil fuels by country [J]. Fuel, 2015, 141: 120-135.",
    "[4] CHONG Z R, YANG S H B, BABU P, et al. Review of natural gas hydrates as an energy resource: Prospects and challenges [J]. Applied Energy, 2016, 162: 1633-1652.",
    "[5] IEA. World Energy Outlook 2024 [R]. Paris: International Energy Agency, 2024.",
    "[6] BARATI R, LIANG J T. A review of fracturing fluid systems used for hydraulic fracturing of oil and gas wells [J]. Journal of Applied Polymer Science, 2014, 131(16): 40735.",
    "[7] MEDEIROS F, KURTOGLU B, OZKAN E, et al. Analysis of production data from hydraulically fractured horizontal wells in shale reservoirs [J]. SPE Reservoir Evaluation & Engineering, 2010, 13(3): 559-568.",
    "[8] PATIDAR A K, JOSHI D, DRISTANT U, et al. A review of tracer testing techniques in porous media specially attributed to the oil and gas industry [J]. Journal of Petroleum Exploration and Production Technology, 2022, 12(12): 3339-3356.",
    "[9] SANNI M, AL-ABBAD M, KOKAL S, et al. Pushing the envelope of residual oil measurement: A field case study of a new class of inter-well chemical tracers [J]. Journal of Petroleum Science and Engineering, 2018, 163: 538-545.",
    "[10] SILVA M, STRAY H, BJORNSTAD T. Stability assessment of PITT tracer candidate compounds - The case of pyrazines [J]. Journal of Petroleum Science and Engineering, 2019, 182: 106269.",
    "[11] WATKINS J W, MARDOCK E S. Use of radioactive iodine as a tracer in water-flooding operations [J]. Journal of Petroleum Technology, 1954, 6(9): 117-124.",
    "[12] YANG H, GUO K, LIN L, et al. Application of micro-substance tracer test in fractured horizontal wells [J]. Journal of Petroleum Exploration and Production Technology, 2024, 14(5): 1235-1246.",
    "[13] ZHAO B, PANTHI K, MOHANTY K K. Tracer eluting proppants for hydraulic fracture characterization [J]. Journal of Petroleum Science and Engineering, 2020, 190: 107048.",
    "[14] ZHOU Y, LIU H, GAO J, et al. Coated proppants with self-suspension and tracer slow-release functions [J]. Journal of Petroleum Science and Engineering, 2022, 208: 109645.",
    "[15] LI N, CHENG Q, GONG Z, et al. Release kinetics of rare earth tracer from polymer-coated proppants for hydraulic fracture analysis [J]. Geoenergy Science and Engineering, 2023, 227: 211782.",
    "[16] GONG Z, LI N, KANG W, et al. Novel oleophilic tracer-slow-released proppant for monitoring the oil production contribution [J]. Fuel, 2024, 364: 130945.",
    "[17] SHOOK G M, POPE G A, ASAKAWA K. Determining reservoir properties and flood performance from tracer test analysis [C]. SPE 124614, SPE Annual Technical Conference and Exhibition, New Orleans, 2009.",
    "[18] FONTALVO E M, OLIVEIRA M C, SCHOEGGL F, et al. Physical interpretation of interwell partitioning tracer tests for estimation of remaining oil saturation in layered carbonate reservoirs [J]. Transport in Porous Media, 2025, 152: 21-45.",
    "[19] LIU J, WANG H, ZHANG T, et al. Study on interpretation method of multistage fracture tracer flowback curve in tight oil reservoirs [J]. ACS Omega, 2024, 9: 10852-10864.",
    "[20] TIAN W, DARNLEY A, DEMPSEY D. Quantifying fracture interference and allocating load recovery and hydrocarbon production in various well configuration using chemical tracers [C]. SPE-201292-MS, SPE Annual Technical Conference and Exhibition, Virtual, 2020.",
    "[21] GONG Z, LI N, QIN M, et al. Magnetic nano-Fe3O4-based oleophilic tracer for stability studies of nano-tracer in oilfields condition [J]. Colloids and Surfaces A: Physicochemical and Engineering Aspects, 2024, 683: 133085.",
    "[22] KRUMBEIN W C, SLOSS L L. Stratigraphy and Sedimentation [M]. 2nd ed. San Francisco: W.H. Freeman and Company, 1963.",
    "[23] RITGER P L, PEPPAS N A. A simple equation for description of solute release I. Fickian and non-Fickian release from non-swellable devices in the form of slabs, spheres, cylinders or discs [J]. Journal of Controlled Release, 1987, 5(1): 23-36.",
    "[24] PEPPAS N A, SAHLIN J J. A simple equation for the description of solute release. III. Coupling of diffusion and relaxation [J]. International Journal of Pharmaceutics, 1989, 57(2): 169-172.",
    "[25] VAN GENUCHTEN M T, ALVES W J. Analytical solutions of the one-dimensional convective-dispersive solute transport equation [R]. USDA Technical Bulletin No. 1661, U.S. Department of Agriculture, 1982.",
    "[26] LI J, JIANG H, WANG B, et al. Tracer flowback modeling and characterization of complex fracture networks in multi-fractured horizontal wells [J]. Journal of Natural Gas Science and Engineering, 2022, 98: 104987.",
    "[27] WANG G, MA Q, REN L, et al. A comprehensive review of multifunctional proppants [J]. ACS Omega, 2024, 9(44): 44120-44133.",
    "[28] KRISHNAN M R, LI W, ALHARBI B, et al. In-situ high-strength poly(styrene-methyl methacrylate)-2D nanofiller composite microbeads as potential proppants in hydraulic fracturing [J]. Geoenergy Science and Engineering, 2025, 257: 214195.",
    "[29] GUO X, WEI K, NI T, et al. Preparation and performance analysis of polyethylene glycol/epoxy resin composite phase change material [J]. Journal of Energy Storage, 2024, 88: 111525.",
    "[30] LIANG C, LUO W, YAN C, et al. Ultra-lightweight proppant synthesized from PMMA/pine bark composite: Low-cost material and outstanding properties [J]. Chemistry Letters, 2016, 45(8): 994-996.",
    "[31] ZOVEIDAVIANPOOR M, GHARIBI A, BIN JAAFAR M Z. Experimental characterization of a new high-strength ultra-lightweight composite proppant derived from renewable resources [J]. Journal of Petroleum Science and Engineering, 2018, 170: 1038-1047.",
    "[32] SABINS F, APBLETT A, SHAFER R, et al. Epoxy resin exhibits long-term durability and chemical stability as a well sealant [C]. SPE-204374-MS, SPE International Conference on Oilfield Chemistry, The Woodlands, Texas, 2021.",
    "[33] LI H, LIU Z, LI Y, et al. Evaluation of the release mechanism of sustained-release tracers and its application in horizontal well inflow profile monitoring [J]. ACS Omega, 2021, 6(29): 19269-19280.",
    "[34] WEI M, WANG Y, DUAN Y, et al. Screening and performance evaluation of epoxy resin long-term sustained-release solid tracer [J]. International Journal of Oil, Gas and Coal Technology, 2024, 36(2): 170-196.",
]
for r in refs:
    P(r)

# Save
import time
out = f'四氧化三铁环氧树脂拟合/ESP-T_clean_{int(time.time())}.docx'
doc.save(out)
print(f'Saved: {out}')
print('Clean document built from scratch.')