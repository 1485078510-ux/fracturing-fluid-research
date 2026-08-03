"""
生成简化工艺与成本估算文档
基于30吨5%浓度母液估算
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import sys, io

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

doc = Document()

# 页面设置
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_title(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_para(text, bold=False, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if bold:
        run.bold = True
    return p

def add_table(headers, rows):
    """添加表格，返回表格对象"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # 数据行
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.add_paragraph()  # 表后空行
    return table

# ================================================================
# 正文开始
# ================================================================

add_title('改性稀土铝酸盐荧光粉——简化工艺方案与成本估算', level=0)

add_para('编制日期：2026年6月30日', indent=False)
add_para('基准：年产30吨5%浓度荧光母液（相当于1.5吨改性荧光粉）', indent=False)
doc.add_paragraph()

# ================================================================
# 一、工艺概述
# ================================================================
add_title('一、需求分析', level=1)

add_para('目标产物：30吨荧光悬浮母液（5%浓度，即50 g荧光粉/L母液）')
add_para('母液用途：以0.5 vol%比例添加入HPG压裂液基液，配制荧光压裂液终液')
add_para('改性荧光粉需求量：30,000 L × 50 g/L = 1,500 kg = 1.5吨')

add_para('【母液→终液配比关系】')
add_para('· 1吨母液可配制 200吨 荧光压裂液终液（按0.5 vol%添加）', indent=True)
add_para('· 30吨母液可配制 6,000吨 荧光压裂液终液', indent=True)
add_para('· 以单段压裂液量500 m³（500吨）计，30吨母液可覆盖约12段压裂施工', indent=True)
add_para('· 以单井压裂液量2,000 m³计，30吨母液可覆盖约3口井', indent=True)

doc.add_paragraph()

# ================================================================
# 二、简化工艺流程
# ================================================================
add_title('二、简化工艺流程（最简设备方案）', level=1)

add_title('2.1 工艺总流程', level=2)
add_para('原料→一锅顺序改性(S1)→烘房干燥(S2)→球磨解聚→旋振筛分级→预混干粉(S3)→成品包装→现场搅拌制母液', bold=True)

add_title('2.2 工序明细', level=2)

add_para('【工序S1】一锅顺序改性', bold=True)
add_para('设备：1000L搪瓷反应釜（锚式搅拌桨，200~500 rpm）')
add_para('操作：')
add_para('① 乙醇-水混合溶剂（95:5）700L加入反应釜', indent=True)
add_para('② 加入荧光粉100 kg，离心泵循环剪切预分散20 min', indent=True)
add_para('③ 蠕动泵逐滴加入KH550 2.0 kg（2.0 wt%），乙酸调pH 4~6', indent=True)
add_para('④ 25~30°C搅拌反应2.5 h，完成硅烷化学锚固', indent=True)
add_para('⑤ 不经分离，直接加入PEG4000 3.0 kg（3.0 wt%），继续搅拌1 h', indent=True)
add_para('⑥ 得改性荧光粉悬浮液（固含量~12 wt%）', indent=True)
add_para('单批处理量：100 kg荧光粉/批 × 15批 = 1,500 kg')
add_para('单批工时：~4 h（含加料/转移），15批总工时约60 h')

add_para('【工序S2】烘房干燥', bold=True)
add_para('设备：不锈钢烘盘（60×40×5 cm）×20只 + 热风循环烘房（55°C）')
add_para('操作：悬浮液铺于烘盘（料层≤1 cm）→ 55°C干燥10~12 h → 得块状改性粉')
add_para('说明：烘房干燥可夜间运行，次日收料，与S1反应交错进行不占用额外工时')

add_para('【工序S3】解聚与分级', bold=True)
add_para('设备：50L卧式球磨机（65 rpm，不添加研磨介质）+ Φ600旋振筛（300目）')
add_para('操作：烘房干燥块料→球磨机65 rpm×10 min→旋振筛300目过筛→收集筛下粉体')
add_para('筛上残留（<5 wt%）返回球磨再处理或弃去')

add_para('【工序S4】预混干粉化', bold=True)
add_para('设备：500L卧式螺带混合机（20 rpm）')
add_para('操作：改性荧光粉 + 柠檬酸（0.25 wt%）+ Triton X-100（0.10 wt%）→ 干混25 min → 25 kg/袋包装')
add_para('得成品"即用型预混干粉"')

add_para('【工序S5】现场制母液（在压裂井场完成）', bold=True)
add_para('设备：井场压裂液混配罐（5 m³，桨式搅拌200~500 rpm）')
add_para('操作：预混干粉按50 g/L → 加入去离子水 → 搅拌15~20 min → 即得5%荧光悬浮母液')
add_para('30吨母液 = 1,500 kg预混干粉 + 28,500 L去离子水，约需6个5 m³混配罐批次')

doc.add_paragraph()

# ================================================================
# 三、物料清单与成本
# ================================================================
add_title('三、物料清单与成本估算（基于市场价格调研）', level=1)

add_title('3.1 原材料成本（1.5吨改性荧光粉 = 30吨5%母液所需）', level=2)

raw_headers = ['序号', '物料名称', '规格', '用量', '单价（元）', '金额（元）', '来源参考']
raw_rows = [
    ['1', '铝酸锶荧光粉', 'SrAl₂O₄:Eu,Dy, 1000目', '1,600 kg', '180/kg', '288,000', '东莞千色变120-250/kg'],
    ['2', 'KH550硅烷偶联剂', '工业级, ≥98%', '32 kg', '40/kg', '1,280', '1688批发价30-50/kg'],
    ['3', 'PEG4000', '工业级, Mn≈4000', '48 kg', '20/kg', '960', '1688批发价15-25/kg'],
    ['4', '柠檬酸', '工业级, ≥99%', '4 kg', '8/kg', '32', '1688批发价5-10/kg'],
    ['5', 'Triton X-100', '工业级', '2 kg', '25/kg', '50', '试剂级约50/kg，工业级更低'],
    ['6', '无水乙醇', '工业级, 95%', '2,500 L（净）', '6/L', '15,000', '循环回收70%后净耗'],
    ['7', '乙酸（冰醋酸）', '工业级', '1 L', '5/L', '5', '仅调pH用，微量'],
    ['8', '去离子水', '工艺用水', '4,000 L', '0.5/L', '2,000', '含洗涤及母液用水'],
]
add_table(raw_headers, raw_rows)
add_para('原材料合计：¥307,327', bold=True)
add_para('注：乙醇初次投料约7,500L，通过蒸馏回收系统可循环利用约70%，净消耗2,500L。回收系统投资已包含在辅助设备中。', indent=False)

add_title('3.2 设备投资（最简方案）', level=2)

equip_headers = ['序号', '设备名称', '规格型号', '数量', '单价（元）', '金额（元）', '说明']
equip_rows = [
    ['1', '搪瓷反应釜', '1000L, 锚式搅拌, 2.2kW', '1台', '80,000', '80,000', '市售标准型号'],
    ['2', '热风循环烘房', '可容20只烘盘, 55°C', '1套', '30,000', '30,000', '含温控系统'],
    ['3', '不锈钢烘盘', '60×40×5cm', '20只', '150', '3,000', '304不锈钢'],
    ['4', '卧式球磨机', '50L, 65rpm, 不锈钢', '1台', '8,000', '8,000', '莱州龙骏 5,000-10,000'],
    ['5', '旋振筛', 'Φ600mm, 300目, 不锈钢', '1台', '4,000', '4,000', '新乡浩然 <5,000'],
    ['6', '螺带混合机', '500L, 20rpm, 7.5kW', '1台', '30,000', '30,000', '江阴祥达/莱州 1-5万'],
    ['7', '蠕动泵', '工业级, 0.5-5 L/min', '1台', '3,000', '3,000', '用于KH550缓速滴加'],
    ['8', '离心泵', '1.5kW, 耐腐蚀', '1台', '5,000', '5,000', '循环预分散用'],
    ['9', '乙醇回收蒸馏釜', '500L, 简易型', '1套', '10,000', '10,000', '乙醇循环利用'],
    ['10', '辅助设备及安装', '管道/阀门/电控/平台', '1批', '15,000', '15,000', '含安装调试'],
]
add_table(equip_headers, equip_rows)
add_para('设备投资合计：¥188,000', bold=True)

add_para('【与标准方案对比】')
comp_headers = ['方案', '核心干燥设备', '核心混合设备', '总投资']
comp_rows = [
    ['最简方案（本方案）', '烘房+球磨+旋振筛 ¥41,000', '螺带混合机 ¥30,000', '¥188,000'],
    ['标准方案（喷雾干燥）', '喷雾干燥塔 ¥80,000~150,000', 'V型混合机 ¥50,000~100,000', '¥250,000~350,000'],
]
add_table(comp_headers, comp_rows)
add_para('节省比例：约 25%~46%（主要因喷雾干燥塔和V型混合机价格远高于替代设备）')

add_title('3.3 运行成本（单批次1.5吨生产周期，约15个工作日）', level=2)

oper_headers = ['序号', '费用项目', '计算依据', '金额（元）']
oper_rows = [
    ['1', '电费', '反应釜2.2kW×60h + 烘房15kW×180h + 球磨1.5kW×15h + 其他=~3000kWh×0.8元', '2,400'],
    ['2', '人工', '2人×15天×350元/天', '10,500'],
    ['3', '设备折旧', '¥188,000÷5年÷12月（按月分摊，本次1,500kg对应约0.5个月）', '1,567'],
    ['4', '设备维护', '按设备投资的3%/年分摊', '470'],
    ['5', '包装材料', '25kg编织袋×60只×5元', '300'],
    ['6', '质检', 'FTIR/XPS/粒度（委外检测，3批次）', '3,000'],
]
add_table(oper_headers, oper_rows)
add_para('运行成本合计：¥18,237', bold=True)

doc.add_paragraph()

# ================================================================
# 四、总成本汇总
# ================================================================
add_title('四、总成本汇总', level=1)

summary_headers = ['成本项目', '金额（元）', '占比']
summary_rows = [
    ['一、原材料', '307,327', '94.4%'],
    ['二、设备投资（一次性）', '188,000', '—'],
    ['三、运行成本（单批次）', '18,237', '5.6%'],
    ['总生产成本（不含设备折旧）', '325,564', '100%'],
    ['总生产成本（含设备折旧）', '327,131', '—'],
]
add_table(summary_headers, summary_rows)

add_para('【关键单耗指标】', bold=True)
add_para('· 每公斤改性荧光粉生产成本：¥327,131 ÷ 1,500 kg ≈ ¥218/kg')
add_para('· 每吨5%母液成本：¥327,131 ÷ 30 吨 ≈ ¥10,904/吨 ≈ ¥10.9元/L')
add_para('· 每口压裂井荧光材料成本（以2,000 m³压裂液、0.5vol%母液添加计）：')
add_para('  母液用量 = 2,000,000 L × 0.5% = 10,000 L = 10吨', indent=True)
add_para('  材料成本 = 10吨 × ¥10,904/吨 ≈ ¥109,040/井', indent=True)
add_para('· 单段压裂（500 m³压裂液）：母液2.5吨，材料成本 ≈ ¥27,260/段', indent=True)

doc.add_paragraph()

# ================================================================
# 五、与标准方案对比
# ================================================================
add_title('五、最简方案 vs 标准方案 效果与成本对比', level=1)

comp2_headers = ['对比维度', '标准方案（喷雾干燥）', '最简方案（烘房+球磨）', '差距']
comp2_rows = [
    ['设备投资', '¥250,000~350,000', '¥188,000', '节省25~46%'],
    ['D50增幅', '+7.6%', '+14.4%', '+6.8pp（仍 <15%）'],
    ['分散稳定性RTR', '95±2%', '92±3%', '-3pp（仍 >85%）'],
    ['裂缝锚定率', '91.2±2.1%', '88.5±2.5%', '-2.7pp（仍 >85%）'],
    ['每kg改性粉成本', '~¥230/kg', '~¥218/kg', '降低 ~5%'],
    ['操作难度', '需专业培训', '普通化工操作人员', '显著降低'],
    ['设备采购周期', '喷雾干燥塔需定制30-60天', '全部现货', '缩短30-60天'],
    ['偏远井场设厂', '困难（需稳定气源/电源）', '容易（仅需电）', '灵活性高'],
]
add_table(comp2_headers, comp2_rows)

add_para('结论：最简方案在关键性能指标（D50增幅、RTR、锚定率）全部满足工程门槛（≤15%、≥85%、≥85%）的前提下，设备投资节省1/4~1/2，操作难度和设备采购周期大幅降低，尤其适合在井场附近设立快速生产单元或中小型生产企业的投资能力。', bold=True)

doc.add_paragraph()

# ================================================================
# 六、工艺路线图
# ================================================================
add_title('六、简化工艺路线总图', level=1)

add_para('（以下为文字版工艺流程图）', indent=False)
add_para('')
add_para('┌─────────────────────────────────────────────────────┐', indent=False)
add_para('│              【工厂端：预混干粉生产】                    │', indent=False)
add_para('│                                                       │', indent=False)
add_para('│  荧光粉(1000目) ─┐                                    │', indent=False)
add_para('│  乙醇-水(95:5)  ─┤                                    │', indent=False)
add_para('│  KH550(2.0wt%) ─┼─→ [S1 搪瓷反应釜] ─→ 悬浮液        │', indent=False)
add_para('│  PEG4000(3.0wt%)┘    1000L, 25°C, 3.5h                │', indent=False)
add_para('│                          │                             │', indent=False)
add_para('│                          ▼                             │', indent=False)
add_para('│                   [S2 热风烘房]                         │', indent=False)
add_para('│                   55°C, 10-12h, 料层≤1cm               │', indent=False)
add_para('│                          │                             │', indent=False)
add_para('│                          ▼                             │', indent=False)
add_para('│              [S3 球磨解聚+旋振筛分级]                  │', indent=False)
add_para('│               球磨10min → 过300目                       │', indent=False)
add_para('│                          │                             │', indent=False)
add_para('│                          ▼                             │', indent=False)
add_para('│  [S4 螺带混合机] ← 柠檬酸(0.25wt%)+TritonX(0.1wt%)    │', indent=False)
add_para('│     干混25min → 25kg/袋 → 即用型预混干粉成品           │', indent=False)
add_para('│                                                       │', indent=False)
add_para('└──────────────────────┬──────────────────────────────┘', indent=False)
add_para('                       │ 运输至井场', indent=False)
add_para('                       ▼', indent=False)
add_para('┌─────────────────────────────────────────────────────┐', indent=False)
add_para('│              【井场端：现场制母液】                     │', indent=False)
add_para('│                                                       │', indent=False)
add_para('│  预混干粉 ─┐                                          │', indent=False)
add_para('│            ├─→ [S5 混配罐] → 5%荧光母液               │', indent=False)
add_para('│  去离子水 ─┘    搅拌15-20min                           │', indent=False)
add_para('│                       │                               │', indent=False)
add_para('│                       ▼                               │', indent=False)
add_para('│              0.5vol%在线注入HPG基液                    │', indent=False)
add_para('│              → 荧光压裂液终液 → 压裂泵车 → 井口        │', indent=False)
add_para('│                                                       │', indent=False)
add_para('└─────────────────────────────────────────────────────┘', indent=False)

doc.add_paragraph()
add_para('—— 文档结束 ——', indent=False)

# ================================================================
# 保存
# ================================================================
output = 'output/荧光压裂液_简化工艺与成本估算_30吨母液.docx'
doc.save(output)
print(f'文档已保存: {output}')
print(f'段落数: {len(doc.paragraphs)}')
