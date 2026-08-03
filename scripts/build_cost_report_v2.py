"""
更新成本估算文档 —— 荧光粉改为600目，修正各项单价
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import sys, io

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_title(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def add_para(text, bold=False, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if bold:
        run.bold = True
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers), style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.add_paragraph()
    return table

# ================================================================
add_title('改性稀土铝酸盐荧光粉——简化工艺方案与成本估算', level=0)
add_para('编制日期：2026年6月30日', indent=False)
add_para('基准：年产30吨5%浓度荧光母液（相当于1.5吨改性荧光粉）', indent=False)
add_para('荧光粉规格：SrAl₂O₄:Eu²⁺,Dy³⁺，600目（~23μm），黄绿色长余辉型', indent=False)
doc.add_paragraph()

# ================================================================
add_title('一、需求分析', level=1)

add_para('目标产物：30吨荧光悬浮母液（5%浓度，即50 g荧光粉/L母液）')
add_para('改性荧光粉需求量：30,000 L × 50 g/L = 1,500 kg = 1.5吨')

add_para('【600目选型依据】', bold=True)
add_para('① 经济性：600目为市场通用规格，厂家多、供货快，单价约为1000目的60~70%', indent=True)
add_para('② 分散性：粒径23μm比表面积更小，改性过程中团聚倾向更低，可适当降低对超声预分散和喷雾干燥的设备依赖', indent=True)
add_para('③ 裂缝进入能力：水力压裂裂缝宽度通常0.1~5mm，23μm颗粒完全可随携砂液进入并附着', indent=True)
add_para('④ 荧光效果：由于单位质量颗粒数更少（约为1000目的1/4），单颗粒亮度更高（晶格更完整），附着后荧光信号可辨识性不受影响', indent=True)

add_para('【母液→终液配比】', bold=True)
add_para('· 1吨母液 → 200吨荧光压裂液终液（0.5 vol%添加）', indent=True)
add_para('· 30吨母液 → 6,000吨压裂液 → 覆盖约12段或3口井', indent=True)

doc.add_paragraph()

# ================================================================
add_title('二、简化工艺流程（5个工序，全部通用设备）', level=1)

add_title('2.1 工艺总流程', level=2)
add_para('原料 → [S1]一锅顺序改性 → [S2]烘房干燥 → [S3]球磨解聚+旋振筛分级 → [S4]螺带混合预混干粉 → 成品包装 → [S5]井场混配罐制母液', bold=True)

add_title('2.2 各工序详细参数', level=2)

add_para('【S1】一锅顺序改性 —— 设备：1000L搪瓷反应釜', bold=True)
add_para('投料（单批100kg粉）：乙醇-水(95:5)700L + 荧光粉100kg → 离心泵循环预分散20min → 蠕动泵滴加KH550 2.0kg(2.0wt%)，乙酸调pH4~6 → 25~30°C搅拌2.5h → 直接加入PEG4000 3.0kg(3.0wt%) → 继续搅拌1h')
add_para('单批工时~4h，共15批，总工时~60h（可两班倒7~8天完成）')

add_para('【S2】烘房干燥 —— 设备：热风循环烘房 + 不锈钢烘盘×20只', bold=True)
add_para('悬浮液铺盘（料层≤1cm）→ 55°C×10~12h → 得块状改性粉。夜间运行，与S1反应日班交错。')

add_para('【S3】解聚与分级 —— 设备：50L卧式球磨机 + Φ600旋振筛(300目)', bold=True)
add_para('烘房块料 → 球磨65rpm×10min（无介质）→ 过300目旋振筛 → 筛下粉收集，筛上<5%返回或弃去。')

add_para('【S4】预混干粉化 —— 设备：500L螺带混合机', bold=True)
add_para('改性粉 + 柠檬酸0.25wt% + Triton X-100 0.10wt% → 20rpm干混25min → 25kg/袋包装 → 即用型成品')

add_para('【S5】现场制母液 —— 设备：井场混配罐（已有）', bold=True)
add_para('预混干粉50g/L浓度 + 去离子水 → 200~500rpm搅拌15~20min → 5%荧光母液。约6个5m³罐批次。')

doc.add_paragraph()

# ================================================================
add_title('三、成本估算', level=1)

add_title('3.1 荧光粉原料价格调研', level=2)

add_para('以下为2024-2025年中国市场SrAl₂O₄:Eu²⁺,Dy³⁺铝酸锶夜光粉真实报价汇总：')

price_headers = ['采购量级', '600目价格', '1000目价格', '数据来源']
price_rows = [
    ['1~9 kg（零售）', '¥120~180/kg', '¥230~435/kg', '千色变/华利化学/1688'],
    ['20~99 kg（小批发）', '¥100~150/kg', '¥180~250/kg', '马可波罗/1688厂家'],
    ['100~999 kg（中批量）', '¥90~130/kg', '¥120~200/kg', '厂家直供议价'],
    ['1,000 kg+（大批量）', '¥80~110/kg', '¥107~150/kg', '年度框架协议'],
    ['5,000 kg+（吨级）', '¥70~90/kg', '¥80~120/kg', 'OEM大单'],
]
add_table(price_headers, price_rows)

add_para('本项目年采购量1,600kg，取600目中批量保守价 ¥110/kg 作预算基准（实际议价空间至 ¥90~100/kg）。', bold=True)

add_title('3.2 原材料成本（1.5吨改性粉 = 30吨5%母液）', level=2)

raw_headers = ['序号', '物料', '规格', '用量', '单价', '金额', '价格来源']
raw_rows = [
    ['1', '铝酸锶荧光粉', '600目, SrAl₂O₄:Eu,Dy', '1,600 kg', '¥110/kg', '¥176,000', '1688/马可波罗 600目 ¥120-145/kg零售，批量¥90-110'],
    ['2', 'KH550硅烷偶联剂', '工业级 ≥98%', '32 kg', '¥40/kg', '¥1,280', '1688批发 ¥30-50/kg'],
    ['3', 'PEG4000', '工业级 Mn≈4000', '48 kg', '¥20/kg', '¥960', '1688批发 ¥15-25/kg'],
    ['4', '柠檬酸', '工业级 ≥99%', '4 kg', '¥8/kg', '¥32', '1688批发 ¥5-10/kg'],
    ['5', 'Triton X-100', '工业级', '2 kg', '¥25/kg', '¥50', '试剂级¥50/kg，工业级更低'],
    ['6', '无水乙醇(95%)', '工业级, 循环回收70%', '2,500 L(净耗)', '¥6/L', '¥15,000', '初次7,500L，回收后净耗2,500L'],
    ['7', '冰醋酸', '工业级', '1 L', '¥5/L', '¥5', '仅调pH'],
    ['8', '去离子水', '工艺用水', '4,000 L', '¥0.5/L', '¥2,000', '含洗涤+母液配制'],
]
add_table(raw_headers, raw_rows)
add_para('原材料合计：¥195,327', bold=True)

add_title('3.3 设备投资', level=2)

equip_headers = ['序号', '设备', '规格', '数量', '单价', '金额', '说明']
equip_rows = [
    ['1', '搪瓷反应釜', '1000L, 锚式搅拌, 2.2kW', '1台', '¥80,000', '¥80,000', '市售标准型号'],
    ['2', '热风循环烘房', '20只烘盘容量, 55°C', '1套', '¥30,000', '¥30,000', '含温控'],
    ['3', '不锈钢烘盘', '60×40×5cm, 304', '20只', '¥150', '¥3,000', ''],
    ['4', '卧式球磨机', '50L, 65rpm, 不锈钢', '1台', '¥8,000', '¥8,000', '莱州龙骏 ¥5,000-10,000'],
    ['5', '旋振筛', 'Φ600, 300目, 不锈钢', '1台', '¥4,000', '¥4,000', '新乡浩然 <¥5,000'],
    ['6', '螺带混合机', '500L, 20rpm, 7.5kW', '1台', '¥30,000', '¥30,000', '江阴祥达 ¥1-5万'],
    ['7', '蠕动泵', '工业级 0.5-5L/min', '1台', '¥3,000', '¥3,000', 'KH550缓速滴加'],
    ['8', '离心泵', '1.5kW, 耐腐蚀', '1台', '¥5,000', '¥5,000', '循环预分散'],
    ['9', '乙醇回收蒸馏釜', '500L, 简易型', '1套', '¥10,000', '¥10,000', '溶剂循环'],
    ['10', '辅助及安装', '管道/阀门/电控/平台', '1批', '¥15,000', '¥15,000', ''],
]
add_table(equip_headers, equip_rows)
add_para('设备投资合计：¥188,000（一次性）', bold=True)

add_title('3.4 运行成本（单批次1.5吨，约15个工作日）', level=2)

oper_headers = ['序号', '项目', '计算依据', '金额']
oper_rows = [
    ['1', '电费', '釜2.2kW×60h + 烘房15kW×180h + 球磨1.5kW×15h + 其他 ≈3,000kWh×¥0.8', '¥2,400'],
    ['2', '人工', '2人×15天×¥350/天', '¥10,500'],
    ['3', '设备折旧', '¥188,000÷5年÷12月×0.5月（本次分摊）', '¥1,567'],
    ['4', '维护', '设备投资×3%/年分摊', '¥470'],
    ['5', '包装', '25kg编织袋×60只×¥5', '¥300'],
    ['6', '质检', 'FTIR/XPS/粒度 委外×3批', '¥3,000'],
]
add_table(oper_headers, oper_rows)
add_para('运行成本合计：¥18,237', bold=True)

doc.add_paragraph()

# ================================================================
add_title('四、总成本汇总', level=1)

summary_headers = ['成本项目', '金额', '占比']
summary_rows = [
    ['原材料', '¥195,327', '91.5%'],
    ['运行成本', '¥18,237', '8.5%'],
    ['设备投资（一次性，不计入单批）', '¥188,000', '—'],
    ['总生产成本（不含折旧）', '¥213,564', '100%'],
    ['总生产成本（含折旧）', '¥215,131', '—'],
]
add_table(summary_headers, summary_rows)

add_para('【关键单耗指标】', bold=True)
add_para('· 每kg改性荧光粉生产成本：¥215,131 ÷ 1,500kg ≈ ¥143/kg')
add_para('· 每吨5%母液成本：¥215,131 ÷ 30吨 ≈ ¥7,171/吨')
add_para('· 单段压裂（500m³压裂液）：母液2.5吨 ≈ ¥17,928')
add_para('· 单井压裂（2,000m³压裂液）：母液10吨 ≈ ¥71,710')

doc.add_paragraph()

# ================================================================
add_title('五、1000目→600目 降本效果对比', level=1)

comp_headers = ['对比项', '1000目方案（原）', '600目方案（现）', '降幅']
comp_rows = [
    ['荧光粉单价', '¥160/kg', '¥110/kg', '↓31%'],
    ['荧光粉总成本', '¥256,000', '¥176,000', '↓¥80,000'],
    ['原材料总成本', '¥275,327', '¥195,327', '↓29%'],
    ['每kg改性粉成本', '¥197/kg', '¥143/kg', '↓27%'],
    ['每吨母液成本', '¥9,844', '¥7,171', '↓27%'],
    ['单段压裂材料成本', '¥24,610', '¥17,928', '↓27%'],
    ['单井压裂材料成本', '¥98,440', '¥71,710', '↓27%'],
]
add_table(comp_headers, comp_rows)

doc.add_paragraph()

# ================================================================
add_title('六、完整工艺路线', level=1)

add_para('┌──────────────────────────────────────────────────┐', indent=False)
add_para('│            【工厂端：预混干粉生产】                │', indent=False)
add_para('│                                                    │', indent=False)
add_para('│  荧光粉(600目) ─┐                                 │', indent=False)
add_para('│  乙醇-水(95:5) ─┤                                 │', indent=False)
add_para('│  KH550 2.0wt%  ─┼→ [S1] 1000L搪瓷釜 → 悬浮液     │', indent=False)
add_para('│  PEG4000 3.0wt% ┘   25°C, 3.5h, 一锅顺序加入       │', indent=False)
add_para('│                       │                            │', indent=False)
add_para('│                       ▼                            │', indent=False)
add_para('│                 [S2] 烘房 55°C×12h                  │', indent=False)
add_para('│                       │                            │', indent=False)
add_para('│                       ▼                            │', indent=False)
add_para('│            [S3] 球磨10min → 旋振筛300目            │', indent=False)
add_para('│                       │                            │', indent=False)
add_para('│                       ▼                            │', indent=False)
add_para('│    [S4] 螺带混合机 ← 柠檬酸+Triton X-100          │', indent=False)
add_para('│          干混25min → 25kg/袋 → 成品                │', indent=False)
add_para('│                                                    │', indent=False)
add_para('└────────────────────┬───────────────────────────────┘', indent=False)
add_para('                    │ 运输至井场', indent=False)
add_para('                    ▼', indent=False)
add_para('┌──────────────────────────────────────────────────┐', indent=False)
add_para('│            【井场端：现场制母液】                  │', indent=False)
add_para('│                                                    │', indent=False)
add_para('│  预混干粉50g/L + 去离子水                          │', indent=False)
add_para('│        │                                           │', indent=False)
add_para('│        ▼                                           │', indent=False)
add_para('│  [S5] 井场混配罐 搅拌15-20min → 5%荧光母液        │', indent=False)
add_para('│        │                                           │', indent=False)
add_para('│        ▼                                           │', indent=False)
add_para('│  0.5vol%在线注入HPG基液 → 压裂泵车 → 井口         │', indent=False)
add_para('│                                                    │', indent=False)
add_para('└──────────────────────────────────────────────────┘', indent=False)

doc.add_paragraph()
add_para('—— 文档结束 ——', indent=False)

output = 'output/荧光压裂液_简化工艺与成本估算_30吨母液_600目.docx'
doc.save(output)
print(f'已保存: {output}')
