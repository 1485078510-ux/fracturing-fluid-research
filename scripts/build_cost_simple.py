"""
精简版：5%荧光母液生产成本
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2); s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = '宋体'; style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def h(text, level=1):
    hd = doc.add_heading(text, level=level)
    for r in hd.runs: r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def p(text, bold=False):
    pg = doc.add_paragraph()
    pg.paragraph_format.first_line_indent = Pt(24)
    run = pg.add_run(text)
    run.font.size = Pt(11); run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if bold: run.bold = True

def tbl(headers, rows):
    from docx.enum.table import WD_TABLE_ALIGNMENT
    t = doc.add_table(rows=1+len(rows), cols=len(headers), style='Table Grid')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, hd in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = ''
        rn = c.paragraphs[0].add_run(hd); rn.bold = True; rn.font.size = Pt(10)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.rows[i+1].cells[j]; c.text = ''
            rn = c.paragraphs[0].add_run(str(val)); rn.font.size = Pt(10)
    doc.add_paragraph()
    return t

# ================================================================
h('5%荧光悬浮母液 —— 简化工艺与成本', level=0)
doc.add_paragraph('规格：蓝光型铝酸盐荧光粉 600目 | 规模：30吨母液/年（=1.5吨改性干粉）')
doc.add_paragraph()

# --- 工艺流程 ---
h('一、工艺流程')

h('1.1 工厂端：改性荧光粉生产', level=2)
p('每批处理100kg荧光粉，共16批。')

tbl(
    ['顺序', '操作', '加入物料', '条件', '时间'],
    [
        ['①', '投料分散', '乙醇475L + 水25L + 荧光粉100kg', '离心泵循环剪切', '20 min'],
        ['②', '调pH', '冰醋酸', '调至pH 4~6', '—'],
        ['③', '滴加KH550', 'KH550 2.0kg（2.0wt%）', '蠕动泵缓滴，25~30°C搅拌', '2.5 h'],
        ['④', '加入PEG', 'PEG4000 3.0kg（3.0wt%）', '不经分离，直接加入同一釜', '1 h'],
        ['⑤', '放料', '悬浮液→烘盘', '料层≤1cm', '—'],
        ['⑥', '干燥', '—', '烘房55°C', '10~12 h（夜间）'],
        ['⑦', '解聚', '—', '球磨65rpm，无介质', '10 min'],
        ['⑧', '过筛', '—', '旋振筛300目', '—'],
        ['⑨', '包装', '—', '25kg/袋', '—'],
    ]
)

h('1.2 井场端：母液配制', level=2)
p('配5%浓度母液（50g粉/L），单罐5m³，共6批次。')

tbl(
    ['顺序', '操作', '加入物料', '条件', '时间'],
    [
        ['①', '加水', '去离子水 ~4,750L/罐', '混配罐', '—'],
        ['②', '溶助剂', '柠檬酸 0.625kg/罐\nTriton X-100 0.25kg/罐', '搅拌溶解', '2~3 min'],
        ['③', '加粉', '改性荧光粉 250kg/罐', '边搅边加', '—'],
        ['④', '搅拌分散', '—', '桨式200~500rpm', '15~20 min'],
        ['⑤', '出液', '→ 5%荧光母液', '在线计量泵注入HPG基液', '—'],
    ]
)
p('柠檬酸用量：粉重的0.25wt% → 1,500kg × 0.25% = 3.75kg')
p('Triton用量：粉重的0.10wt% → 1,500kg × 0.10% = 1.50kg')

doc.add_paragraph()

# --- 原料明细 ---
h('二、工厂端：原料与加工（产1.5吨改性荧光粉）')

doc.add_paragraph('说明：柠檬酸和Triton X-100按专利步骤(3)在井场配制母液时加入，不属工厂改性原料。溶剂用量依据专利实施例1（100g粉/500mL，固液比1:5）。荧光粉取蓝光型（Sr₄Al₁₄O₂₅:Eu,Dy或CaAl₂O₄:Eu,Nd），¥140/kg（600目大批量，蓝光比黄绿光贵约30~50%，参考秀彩¥175~180零售、MINHUI ¥137起）。')

h('2.1 工厂原料（5种）', level=2)
tbl(
    ['序号', '物料', '化学式/规格', '用量', '单价', '金额'],
    [
        ['1', '铝酸锶蓝光荧光粉', 'Sr₄Al₁₄O₂₅:Eu²⁺,Dy³⁺\n或CaAl₂O₄:Eu²⁺,Nd³⁺\n600目（~23μm），蓝光~490nm', '1,600 kg', '¥140/kg', '¥224,000'],
        ['2', 'KH550硅烷偶联剂', 'NH₂(CH₂)₃Si(OC₂H₅)₃\n工业级≥98%，用量2.0wt%', '32 kg', '¥40/kg', '¥1,280'],
        ['3', 'PEG4000', 'HO(CH₂CH₂O)nH，Mn≈4000\n工业级，用量3.0wt%', '48 kg', '¥20/kg', '¥960'],
        ['4', '无水乙醇(95%)', 'C₂H₅OH，工业级\n总投7,600L，蒸馏回收85%\n净耗1,140L', '1,140 L', '¥6/L', '¥6,840'],
        ['5', '冰醋酸', 'CH₃COOH，工业级\n调pH 4~6，微量', '1 L', '¥5/L', '¥5'],
    ]
)
p('工厂原料合计：¥233,085', bold=True)

h('2.2 工厂加工步骤（4步）', level=2)
tbl(
    ['步骤', '操作', '设备', '费用'],
    [
        ['S1 一锅改性', '16批×3.5h：粉+溶剂→预分散→滴加KH550反应2.5h→加PEG4000反应1h', '1000L搪瓷釜', '¥2,906'],
        ['S2 烘房干燥', '悬浮液铺盘→55°C×12h（夜间运行）', '烘房+烘盘×20', '¥4,960'],
        ['S3 球磨过筛', '块料→球磨10min→过300目', '球磨机+旋振筛', '¥724'],
        ['S4 质检包装', '委外检测3批 + 25kg/袋包装', '—', '¥3,300'],
    ]
)
p('工厂加工合计：¥11,890', bold=True)
p('工厂总成本（原料+加工）：¥233,085 + ¥11,890 = ¥244,975 → 改性荧光粉 ¥163.32/kg', bold=True)

doc.add_paragraph()

# --- 井场端 ---
h('三、井场端：母液配制')

p('按专利步骤(3)，柠檬酸和Triton X-100在配制母液时加入水中溶解，再分散改性荧光粉。')

tbl(
    ['物料/操作', '用量/说明', '金额'],
    [
        ['改性荧光粉', '1,500 kg（工厂产出）', '¥244,975'],
        ['去离子水', '28,500 L × ¥0.5/L', '¥14,250'],
        ['柠檬酸', 'C₆H₈O₇，工业级\n用量0.25wt%（对粉重）= 3.75 kg × ¥8/kg', '¥30'],
        ['Triton X-100', '辛基酚聚氧乙烯醚，工业级\n用量0.10wt%（对粉重）= 1.5 kg × ¥25/kg', '¥38'],
        ['搅拌', '混配罐6批次 × 15min，电+人工', '¥400'],
    ]
)

doc.add_paragraph()

# --- 汇总 ---
h('四、汇总')

tbl(
    ['项目', '金额', '占比'],
    [
        ['工厂端（原料+加工）', '¥244,975', '94.3%'],
        ['井场端（水+助剂+搅拌）', '¥14,718', '5.7%'],
        ['30吨母液总成本', '¥259,693', '100%'],
    ]
)

p('')
p('每吨母液：¥8,656 | 每升母液：¥8.66', bold=True)
p('单段压裂（母液2.5吨）：¥21,641 | 单井压裂（母液10吨）：¥86,564', bold=True)

# --- 设备 ---
h('五、所需设备（一次性投资 ¥128,000）')

tbl(
    ['设备', '规格', '估价'],
    [
        ['搪瓷反应釜', '1000L', '¥80,000'],
        ['热风烘房+烘盘', '20盘', '¥33,000'],
        ['卧式球磨机', '50L', '¥8,000'],
        ['旋振筛', 'Φ600, 300目', '¥4,000'],
        ['蠕动泵+离心泵', '各1台', '¥8,000'],
        ['乙醇回收釜+辅材', '500L', '¥25,000'],
    ]
)
doc.add_paragraph('注：较之前方案省去螺带混合机（¥30,000），因柠檬酸和Triton在井场加，工厂无需预混。')

doc.add_paragraph()
doc.add_paragraph('—— 完 ——')

doc.save('output/荧光压裂液_母液成本_简化版.docx')
print('已保存')
