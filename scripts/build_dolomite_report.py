#!/usr/bin/env python3
"""生成色斑白云岩文献调研报告 Word 文档"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ---- 全局样式设置 ----
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(4)

# 设置标题样式
for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    hfont = heading_style.font
    hfont.name = '黑体'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    hfont.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        hfont.size = Pt(16)
        heading_style.paragraph_format.space_before = Pt(18)
        heading_style.paragraph_format.space_after = Pt(10)
    elif level == 2:
        hfont.size = Pt(14)
        heading_style.paragraph_format.space_before = Pt(14)
        heading_style.paragraph_format.space_after = Pt(8)
    else:
        hfont.size = Pt(12)
        heading_style.paragraph_format.space_before = Pt(10)
        heading_style.paragraph_format.space_after = Pt(6)

def add_para(text, bold=False, italic=False, size=None, indent_cm=0):
    """添加段落"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(indent_cm if indent_cm else 0.74)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_bullet(text, level=0):
    """添加项目符号"""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * (level + 1))
    return p

def set_cell_shading(cell, color):
    """设置单元格底色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(headers, rows, col_widths=None):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'D9E2F3')

    # 数据行
    for r, row in enumerate(rows):
        for c, cell_text in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(cell_text))
            run.font.size = Pt(9.5)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # 表后空行
    return table

def add_reference(text):
    """添加参考文献条目"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

# ============================================================
# 封面 / 标题
# ============================================================
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_para.paragraph_format.space_before = Pt(60)
run = title_para.add_run('色斑白云岩成因与致色原因\n文献调研报告')
run.bold = True
run.font.size = Pt(22)
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.color.rgb = RGBColor(0, 51, 102)

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_para.paragraph_format.space_before = Pt(20)
run = sub_para.add_run(f'——受白云石内部成分、组构及特殊离子含量差异控制的\n不发育溶蚀特征的不同颜色白云岩斑块研究')
run.font.size = Pt(12)
run.font.name = '楷体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
run.font.color.rgb = RGBColor(80, 80, 80)

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_para.paragraph_format.space_before = Pt(30)
run = date_para.add_run(f'生成日期：{datetime.date.today().strftime("%Y年%m月%d日")}')
run.font.size = Pt(11)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.color.rgb = RGBColor(120, 120, 120)

doc.add_page_break()

# ============================================================
# 目录占位
# ============================================================
doc.add_heading('目  录', level=1)
add_para('（请在 Word 中插入自动目录：引用 → 目录 → 自动目录）', italic=True, size=10)
doc.add_page_break()

# ============================================================
# 一、概念与分类
# ============================================================
doc.add_heading('一、色斑白云岩的概念与分类', level=1)

doc.add_heading('1.1 定义', level=2)
add_para('色斑白云岩（又称斑化白云岩、斑状白云岩、豹斑白云岩/豹斑云质灰岩；英文：mottled dolomite, color-spotted dolomite, patchy dolomite）是指白云石呈斑块状、条带状、网状或不规则状不均匀分布于灰岩基质中，形成不同颜色斑块相互拼接或不均匀分布的岩石类型。其核心特征为：受白云石内部成分、组构及特殊离子含量的差异控制，不发育溶蚀特征，常见浅红色、米黄色、绿灰色、浅灰色及深灰色等白云岩斑块。')

doc.add_heading('1.2 按白云石化程度分类', level=2)
add_table(
    ['类型', '白云石含量', '生物扰动指数', '特征描述'],
    [
        ['不均匀斑化白云岩', '> 90%', 'BI = 6（完全扰动）', '白云石占绝对主导，原岩组构基本消失'],
        ['明显斑化含灰云岩\n（豹斑云质灰岩）', '50% ~ 90%', 'BI = 4~5（中-强扰动）', '潜穴充填白云石，基质为泥晶方解石，斑块明显'],
        ['略微斑化灰质云岩\n（云斑灰岩/豹皮灰岩）', '< 50%', 'BI = 2~3（弱-中扰动）', '基质灰岩为主，白云石斑块零星分布'],
    ],
    col_widths=[4.0, 2.5, 3.5, 6.0]
)

doc.add_heading('1.3 颜色特征', level=2)
add_para('色斑白云岩的常见颜色包括：浅红色、米黄色、绿灰色、浅灰色、深灰色、黄褐色、紫红色等。颜色分异主要受 Fe 价态与含量、Mn 含量、有机质丰度、晶粒大小及白云石化程度共同控制（详见第三章）。')

# ============================================================
# 二、成因机制
# ============================================================
doc.add_heading('二、成因机制——主要学说', level=1)
add_para('色斑白云岩的成因研究经历了从描述性岩石学到多因素耦合机制认识的演进过程。目前国际主流观点认为，生物扰动造成的原始渗透率差异是色斑结构形成的最普遍初始条件，而白云石化流体性质、成岩阶段与构造背景的差异则决定了斑块的具体形态与地球化学特征。')

doc.add_heading('2.1 生物扰动 + 渗透回流白云石化模式', level=2)
add_para('【中国鄂尔多斯盆地主流模式】', bold=True, size=11)
add_para('这是目前国内研究最为深入、证据最为充分的色斑白云岩成因模式，以鄂尔多斯盆地奥陶系马家沟组为代表。')
add_para('核心逻辑链：生物扰动 → 原始组构不均一 → 渗透率差异 → 选择性白云石化 → 斑状构造')
add_para('具体过程如下：')

steps = [
    '固底底质条件下，Thalassinoides（海生迹）等造迹生物（三叶虫类、甲壳类、蠕虫类）在碳酸盐沉积物中挖掘复杂潜穴系统，形成高孔渗通道网络；',
    '早期同沉积胶结（镁方解石）使潜穴间基质硬化、渗透率大幅降低，而潜穴充填物保持未胶结状态、维持高渗透率；',
    '海平面高频振荡过程中，蒸发浓缩形成的富Mg²⁺卤水沿潜穴/岩溶通道向下渗透回流，优先交代高渗透潜穴充填物中的方解石/文石为白云石；',
    '埋藏阶段温压升高，白云石发生重结晶，进一步调整晶体结构及有序度，晶粒增大。'
]
for s in steps:
    add_bullet(s)

add_para('遗迹化石组合：鄂尔多斯盆地马家沟组四段共识别出遗迹化石7属8种，以 Thalassinoides（海生迹）的两种形态——Thalassinoides network（水平迷宫形）和 Thalassinoides boxwork（三维箱状结构）占主导，属于固底底质条件的 Glossifungites 遗迹相。')

add_para('垂向序列演变（向上变浅序列）：')
add_bullet('下部：水平状豹斑灰岩（Thalassinoides network）')
add_bullet('中部：斜交管状豹斑')
add_bullet('上部：不规则花斑状豹斑云岩（Thalassinoides boxwork）')
add_para('该演变序列反映了从正常浪基面以下到潮间带的水动力逐渐降低过程。')

doc.add_heading('2.2 生物潜穴优先白云石化经典模式（加拿大 Tyndall Stone）', level=2)
add_para('【国际奠基性研究】', bold=True, size=11)
add_para('Kendall (1977) 对加拿大 Saskatchewan 和 Manitoba 地区奥陶系 Yeoman 组（Red River 群下部）Tyndall Stone 的研究是色斑白云岩成因研究的奠基性工作。Tyndall Stone 以米黄色白云石斑块嵌布于深灰色未白云石化泥晶灰岩基质中而闻名，被广泛用作建筑装饰石材。')

add_para('Kendall 的关键发现：')
findings = [
    '斑块是"被选择性白云石化的潜穴充填物"，而非围绕潜穴的白云石化晕——这一判断基于斑块内部偏心位置的小型潜穴、半月形充填构造等微观证据；',
    '关键证据链——硬石膏假晶（代替石盐晶体）仅出现在白云石斑块内部，说明潜穴充填物直至深埋成岩晚期仍保持未胶结、可位移状态（石盐晶体能在其中生长），而基质早已硬化；',
    'Mg来源的两种可能：(a) 上覆 Herald 组蒸发岩的回流卤水；(b) 原始镁方解石沉积物稳定化过程中释放的Mg²⁺被"自噬"再利用；',
    '后期存在第二期基质白云石化，可形成第二代颜色斑块，造成多期叠加的复杂斑块格局。'
]
for f in findings:
    add_bullet(f)

add_para('Gingras, Pemberton & Henk (2001) 进一步补充了微生物介导机制：潜穴微环境富集有机质 → 微生物硫酸盐还原/发酵作用 → 改变局部pH和碱度 → 促进白云石在潜穴内部优先沉淀。富集δ¹³C的同位素证据支持这一解释。')

doc.add_heading('2.3 多期埋藏白云石化模式（川西北栖霞组）', level=2)
add_para('曾鑫耀在川西北双鱼石地区下二叠统栖霞组识别出三期白云岩化过程：')

add_table(
    ['成岩阶段', '产物', '特征'],
    [
        ['早期埋藏白云岩化', '零星白云石（Rd1）+ 斑状白云石（Rd2）', '形成豹斑云质灰岩，白云石呈斑块状不均匀分布'],
        ['晚期埋藏白云岩化', '中粗晶白云石（Rd3）+ 基质细晶白云石（Rd4）', '大规模白云石化，斑块边界模糊化'],
        ['过度白云岩化', '裂缝/孔隙充填鞍状白云石（Cd1, Cd2）', '热液叠加，鞍状白云石具弯曲晶面和波状消光'],
    ],
    col_widths=[3.5, 5.5, 6.5]
)
add_para('地球化学特征表明，云化流体为保存在地层中的封存海水，无显著外源流体加入（δ¹³C与同期海水一致，⁸⁷Sr/⁸⁶Sr略高于同期海水）。')

doc.add_heading('2.4 构造-热液控制模式', level=2)
add_para('Navarro-Ciurana et al. (2016) 在西班牙东南部 Prebetic 带 Riópar 地区描述了沿断层分布的不规则斑块状热液白云岩体（patchy geobodies）。流体温度 150–250°C，由低盐度与高盐度卤水混合触发白云石化，伴有鞍形白云石和 MVT 型硫化物矿化。该模式适用于构造活动区、与断裂系统密切相关的斑块状白云岩。')

doc.add_heading('2.5 鲕粒差异白云石化模式', level=2)
add_para('韩宇等（2023）在重庆云阳飞仙关组发现，鲕粒内部差异白云石化受渐进式胶结 + 准同生大气淡水淋滤共同控制：完全白云石化的鲕粒经历了早期文石质均一化胶结 → 无差异白云石化；非完全白云石化的鲕粒经历了文石质/方解石差异胶结 → 仅沿文石质胶结物白云石化，形成斑状结构。')

doc.add_heading('2.6 成因机制总结', level=2)
add_table(
    ['成因模式', '核心驱动力', '关键证据', '典型实例'],
    [
        ['生物扰动+渗透回流', '潜穴渗透率差异', '潜穴形态、潜穴内石盐假晶', '鄂尔多斯马家沟组'],
        ['埋藏白云石化（多期）', '封存海水+温压升高', 'δ¹³C/δ¹⁸O、⁸⁷Sr/⁸⁶Sr', '川西北栖霞组'],
        ['构造-热液', '断裂流体+高温卤水', '鞍形白云石、硫化物、δ¹⁸O亏损', '西班牙 Riópar'],
        ['差异胶结+大气淡水', '鲕粒胶结物矿物差异', '文石/方解石差异胶结', '重庆飞仙关组'],
        ['微生物介导', '有机质+微生物代谢', 'δ¹³C富集、微组构', '加拿大 Yeoman Fm'],
    ],
    col_widths=[3.5, 3.5, 4.5, 4.0]
)

# ============================================================
# 三、致色原因
# ============================================================
doc.add_heading('三、致色原因——Fe、Mn等元素对颜色的控制', level=1)
add_para('色斑白云岩的颜色差异是多种因素在不同尺度上叠加的结果。宏观颜色主要受铁（Fe）的价态和赋存形式控制，微观阴极发光特征主要受 Fe/Mn 比值控制，有机质含量和晶体粒度则进一步叠加影响。')

doc.add_heading('3.1 颜色的基本分类', level=2)
add_table(
    ['颜色类型', '成因阶段', '典型颜色', '特征'],
    [
        ['自生色', '沉积期及早期成岩', '灰色、深灰色、黑灰色', '反映原始沉积环境，颜色越深→水体越深、有机质越高、还原性越强'],
        ['次生色', '后生作用/风化阶段', '白色、灰绿色、黄褐色、紫红色', '由次生矿物（如赤铁矿）造成，常切穿层面，分布不均匀，多呈块状、斑点状'],
    ],
    col_widths=[2.5, 3.0, 4.0, 6.0]
)

doc.add_heading('3.2 铁（Fe）——最重要的致色元素', level=2)
add_para('铁是色斑白云岩颜色变化的决定性致色元素，其价态决定了颜色的基色调：')

add_table(
    ['Fe 价态/矿物', '颜色表现', '成岩环境', '形成机制'],
    [
        ['Fe³⁺（赤铁矿 Fe₂O₃）', '紫红、肉红、黄褐', '氧化环境\n（表生风化/大气淡水）', '含铁流体氧化沉淀，赤铁矿以独立晶体形式出现在白云石菱面体边缘或核心'],
        ['Fe²⁺（亚铁化合物）', '蓝灰、暗色、深灰', '还原环境\n（埋藏成岩）', 'Fe²⁺进入白云石晶格替代Mg²⁺，或形成亚铁化合物沿晶界分布'],
    ],
    col_widths=[4.0, 3.0, 3.0, 5.5]
)

add_para('川西北泥盆系观雾山组实例：X射线衍射（XRD）分析表明，紫红色白云岩中检测出约3%的赤铁矿（Fe₂O₃），而正常灰色白云岩不含赤铁矿。铁元素来源于下伏地层的含铁砂岩（赤铁矿砂岩），经大气淡水风化淋滤 → 含铁水体汇聚 → 浸染、氧化碳酸盐颗粒，导致岩石呈现紫红、肉红等色调。镜下可见白云石核心被铁质染色呈深褐色，而核心外部胶结物几乎未被染色——这说明铁质浸染主要发生在早成岩期。')

add_para('Kendall (1977) 的经典观察进一步证实：白云石菱面体本身是无色透明的，铁氧化物以独立晶体形式出现在白云石菱面体边缘，而非晶格内均匀替代——这意味着 Fe 是由白云石化流体同期引入或活化，而非后期风化淋滤带入。渗流大气水可在成岩后期氧化 Fe²⁺→Fe³⁺，形成条带状或环带状的铁质染色纹理。')

doc.add_heading('3.3 锰（Mn）与阴极发光（CL）特征', level=2)
add_para('Mn 和 Fe 在白云石晶格中的含量与比值是控制阴极发光颜色和强度的核心参数：')

add_table(
    ['元素', '在CL中的角色', '临界浓度', '发光效应'],
    [
        ['Mn²⁺', '激活剂（Activator）', '≥ 10–40 ppm 即可激活', '产生橙红色、橙黄色、玫瑰红色发光'],
        ['Fe²⁺', '猝灭剂（Quencher）', '≥ 35 ppm 开始抑制\n> 9160 μg/g 显著减弱', '抑制发光强度，Fe²⁺含量越高发光越弱直至不发光'],
        ['Fe²⁺/Mn²⁺比值', '综合控制参数', '比值高→发光暗/不发光\n比值低→发光亮', '反映成岩流体氧化还原条件的演化'],
    ],
    col_widths=[3.0, 3.0, 4.0, 5.5]
)

add_para('Machel (1985) 及 Machel & Burton (1991) 指出，至少有26种因素影响碳酸盐矿物的阴极发光行为，单独依赖 CL 颜色推断古氧化还原条件存在严重风险。必须结合稳定同位素（δ¹³C、δ¹⁸O）、流体包裹体、微量元素（ICP-MS/SIMS/EPMA）和岩石学鉴定进行综合判断。')

add_para('白云石 CL 环带的环境指示意义：')
add_bullet('晶核→边缘（浅埋成岩环境）：Fe²⁺、Mn²⁺逐渐富集 → CL由亮变暗 → 单偏光下表现为"雾心亮边"')
add_bullet('热液鞍形白云石：多期Fe条带（由内到外Fe含量依次降低）→ 多环带发光 → 反映多期次深部热液侵入')

doc.add_heading('3.4 有机质含量', level=2)
add_para('生物扰动区（潜穴内部）残留的有机质可使白云石颜色显著加深（灰-深灰色），而未受扰动的基质区域因有机质含量低而颜色较浅。潜穴微环境中有机质的富集还可通过微生物硫酸盐还原/发酵作用改变局部pH和碱度，间接影响Fe、Mn的溶解度，从而对颜色产生次级控制。')

doc.add_heading('3.5 晶粒大小与光的物理散射', level=2)
add_para('晶体粒度差异本身即可形成肉眼可辨的颜色对比：')
add_bullet('细晶白云石（< 50 μm）：比表面积大，光散射强 → 颜色偏浅（灰白、米黄色）')
add_bullet('粗晶白云石（> 200 μm）：透明度高，显示基质/杂质颜色 → 颜色偏深')
add_bullet('不同斑块的重结晶程度差异 → 晶粒大小差异 → 相同矿物成分下呈现不同色调')

doc.add_heading('3.6 致色因素综合汇总', level=2)
add_table(
    ['致色机制', '控制因素', '颜色表现'],
    [
        ['Fe³⁺（赤铁矿）浸染', '氧化还原条件、铁源供给', '紫红、肉红、黄褐色斑块'],
        ['Fe²⁺/Mn²⁺ 比值', '成岩流体氧化还原演化', '阴极发光颜色与强度（橙红↔暗红↔不发光）'],
        ['有机质含量差异', '生物扰动强度、有机质保存条件', '深灰-黑色斑块（有机质富集区）'],
        ['晶粒大小差异', '重结晶程度', '浅色（细晶）vs 深色（粗晶）'],
        ['白云石化程度差异', '选择性交代', '白云石（米黄/灰）vs 方解石基质（深灰）'],
    ],
    col_widths=[4.0, 5.5, 6.0]
)

# ============================================================
# 四、典型实例
# ============================================================
doc.add_heading('四、国内外典型实例', level=1)

doc.add_heading('4.1 国内典型实例', level=2)
add_table(
    ['地区/层位', '类型', '成因模式', '主要研究者'],
    [
        ['鄂尔多斯盆地\n马家沟组 O₂m', '斑化白云岩\n豹斑灰岩', '生物扰动+渗透回流白云石化', '廖慧鸿、苏中堂(2022)\n许杰等(2022)'],
        ['鄂尔多斯西北部\n桌子山组 O₂', '斑状白云岩', '差异白云石化+表生岩溶叠加', '何小会、董兆雄等(2012)'],
        ['塔里木盆地塔中\n奥陶系 O', '斑状白云岩\n（6类之一）', '埋藏白云石化+热液叠加', '胡明毅等(2011)'],
        ['川西北双鱼石\n栖霞组 P₁q', '豹斑云质灰岩', '多期埋藏白云石化\n（三期）', '曾鑫耀（硕士论文）'],
        ['川中磨溪-高石梯\n栖霞组 P₁q', '斑状白云岩', '多期白云石化+热液', '何溥为、胥旺等'],
        ['豫西登封\n朱砂洞组 ∈₂', '豹斑构造\n（条带→花斑→网格）', '生物扰动优先白云石化', '地质科技通报(2014)'],
        ['川西北\n观雾山组 D₂g', '紫红色白云岩', '赤铁矿（Fe³⁺）浸染致色', '方少仙、董兆雄等'],
        ['重庆云阳\n飞仙关组 T₁f', '鲕粒不均匀白云石化', '差异胶结+大气淡水淋滤', '韩宇、张云峰等(2023)'],
    ],
    col_widths=[3.5, 3.5, 4.5, 4.0]
)

doc.add_heading('4.2 国外典型实例', level=2)
add_table(
    ['地区/层位', '类型', '成因模式', '主要文献'],
    [
        ['加拿大 Saskatchewan\nManitoba·Yeoman Fm\n（奥陶系）', 'Tyndall Stone\n白云石斑块', '潜穴选择性白云石化\n（奠基性研究）', 'Kendall (1977) ⭐'],
        ['西班牙 Riópar\nPrebetic Zone', '断层控制斑块状\n热液白云岩', '构造热液（150–250°C）\nMVT型矿化共生', 'Navarro-Ciurana et al.\n(2016)'],
        ['巴基斯坦\nSamana Suk Fm\n（侏罗系）', '斑块状白云岩', '断裂/缝合线控制\n白云石化', 'Springer (2020)'],
        ['美国 California\nBeck Spring Dolomite\n（前寒武系）', '纤维状白云石\n胶结物颜色条带', '混合水带白云石化\n斑块状发光', 'GeoScienceWorld\n文献'],
        ['加拿大 Alberta\nRainbow Buildups\n（中泥盆统）', '漂浮白云石菱面体\n/斑块', '早期成岩选择性交代\n基质优先白云石化', 'Qing & Mountjoy\n(1989), JSP'],
    ],
    col_widths=[3.5, 3.5, 4.5, 4.0]
)

# ============================================================
# 五、研究方法
# ============================================================
doc.add_heading('五、主要研究方法', level=1)
add_para('色斑白云岩的成因与致色研究需要多学科、多尺度手段的综合运用：')

methods = [
    '岩石学方法：偏光显微镜薄片鉴定（晶粒大小、组构、孔隙类型）、阴极发光显微镜（CL，Fe/Mn分布与成岩流体示踪）、扫描电镜（SEM，微米-纳米尺度组构观察）、X射线衍射（XRD，鉴定致色矿物相如赤铁矿）；',
    '地球化学方法：碳氧稳定同位素（δ¹³C、δ¹⁸O，流体来源与成岩温度）、锶同位素（⁸⁷Sr/⁸⁶Sr，流体来源与海水对比）、主微量元素分析（Fe、Mn、Sr、Na、Ba，致色元素定量）、稀土元素配分模式（REE，氧化还原条件与流体来源）；',
    '微区分析技术：电子探针（EPMA，微米级Fe/Mn定量与环带分析）、二次离子质谱（SIMS，微量元素与同位素微区原位分析）、激光剥蚀ICP-MS（LA-ICP-MS，高灵敏度微量元素与U-Pb定年）；',
    '流体包裹体分析：均一温度（成岩温度约束）、盐度（流体性质约束）、激光拉曼（包裹体成分鉴定）；',
    '遗迹化石学方法：潜穴形态分类与三维重建、生物扰动指数（BI）定量评价、遗迹相分析与沉积环境判别。'
]
for m in methods:
    add_bullet(m)

# ============================================================
# 六、关键文献索引
# ============================================================
doc.add_heading('六、关键参考文献索引', level=1)

doc.add_heading('6.1 中文核心文献', level=2)

zh_refs = [
    '[1] 廖慧鸿, 苏中堂, 黄文明, 等. 鄂尔多斯盆地奥陶系马家沟组五段5亚段斑化白云岩成因[J]. 石油实验地质, 2022, 44(5): 835-844.',
    '[2] 许杰, 苏中堂, 等. 鄂尔多斯盆地东缘奥陶系马家沟组四段豹斑状云质灰岩特征及成因：以关家崖剖面为例[J]. 古地理学报, 2022, 24(2).',
    '[3] 何小会, 董兆雄, 张文涛, 等. 鄂尔多斯盆地西北部中奥陶统斑状白云岩成因研究[J]. 四川地质学报, 2012(4).',
    '[4] 胡明毅, 胡忠贵, 李思田, 等. 塔中地区奥陶系白云岩岩石地球化学特征及成因机理分析[J]. 地质学报, 2011.',
    '[5] 韩宇, 张云峰, 等. 鲕粒的不均匀白云石化及其成因——以重庆云阳上坝剖面飞仙关组为例[J]. 天然气勘探与开发, 2023, 46(1): 19-31.',
    '[6] 何溥为, 胥旺, 等. 川中磨溪—高石梯地区栖霞组白云岩特征及成因机制[J]. 天然气地球科学.',
    '[7] 曾鑫耀. 川西北部双鱼石地区下二叠统栖霞组白云岩成因研究[D]. 中国石油大学（北京）, 硕士论文.',
    '[8] 方少仙, 董兆雄, 侯方浩, 等. 层状白云岩储层特征与成因——以黔桂地区泥盆系、石炭系及湘鄂交界地区三叠系为例[M]. 北京: 地质出版社.',
    '[9] 彭阳, 章雨旭, 等. 豹斑灰岩相关研究[J]. 地球学报, 2000, 21(1).',
    '[10] 豫西登封地区寒武系第二统朱砂洞组生物成因的豹斑构造[J]. 地质科技通报, 2014, 33(5).',
    '[11] 王窅廷, 等. 鄂尔多斯盆地奥陶系豹斑灰岩成因研究[J]. 古地理学报, 2024(4).',
    '[12] 李祖兵, 肖尧, 刘均. 白云岩成因模式及分析[M]. （专著——涵盖塞卜哈、渗透回流、埋藏、热液、混合水、玄武岩淋滤及生物成因等多种模式）',
]
for ref in zh_refs:
    add_reference(ref)

doc.add_heading('6.2 英文核心文献', level=2)

en_refs = [
    '[13] Kendall, A.C. (1977). Origin of dolomite mottling in Ordovician limestones from Saskatchewan and Manitoba. Bulletin of Canadian Petroleum Geology, 25(3): 480–504. ⭐奠基性文献，被引约68次.',
    '[14] Gingras, M.K., Pemberton, S.G. & Henk, F. (2001). Conceptual Models for Burrow-related, Selective Dolomitization. AAPG Annual Meeting, Abstract #0278.',
    '[15] Navarro-Ciurana, D., et al. (2016). Petrography and geochemistry of fault-controlled hydrothermal dolomites in the Riópar area (Prebetic Zone, SE Spain). Marine and Petroleum Geology, 71: 1–25.',
    '[16] Machel, H.G. (1985). Cathodoluminescence in calcite and dolomite and its chemical interpretation. Geoscience Canada, 12(4): 139–147.',
    '[17] Machel, H.G. & Burton, E.A. (1991). Factors Governing Cathodoluminescence in Calcite and Dolomite, and their Implications for Studies of Carbonate Diagenesis. SEPM Short Course 25, 37–57.',
    '[18] Qing, H. & Mountjoy, E.W. (1989). Multistage dolomitization in Rainbow buildups, Middle Devonian Keg River Formation, Alberta, Canada. Journal of Sedimentary Petrology, 59(1).',
    '[19] Diagenetic evolution and associated dolomitization events in the middle Jurassic Samana Suk Formation, Lesser Himalayan Hill Ranges, NW Pakistan. Carbonates and Evaporites (Springer), 2020.',
    '[20] Hiatt, E.E. & Pufahl, P.K. (2014). Cathodoluminescence applications in carbonate diagenesis studies. SEPM Special Publication.',
]
for ref in en_refs:
    add_reference(ref)

# ============================================================
# 七、总结
# ============================================================
doc.add_heading('七、总结与展望', level=1)
add_para('综合以上文献调研，色斑白云岩的成因与致色是一个多因素耦合的复杂地质过程，可以概括为以下核心认识：')

doc.add_heading('7.1 成因认识', level=2)
conclusions = [
    '根本驱动：原始沉积组构的非均质性是色斑形成的初始条件，其中以生物扰动（Thalassinoides 等造迹生物潜穴系统）造成的渗透率差异最为普遍和重要；',
    '关键过程：富Mg流体（蒸发浓缩海水、封存地层水或热液）沿高渗透率路径（潜穴、裂缝、缝合线、粒间孔）的选择性白云石化是斑块形成的核心机制；',
    '多期叠加：多数色斑白云岩经历了多期白云石化（准同生期→浅埋藏→深埋藏→热液叠加）的复合改造，不同期次白云石的空间叠置形成了复杂的斑块格局；',
    '构造控制：断裂和裂缝系统既是白云石化流体的运移通道，也能直接控制斑块状白云岩体的空间分布（热液白云岩模式）。'
]
for c in conclusions:
    add_bullet(c)

doc.add_heading('7.2 致色认识', level=2)
coloring = [
    'Fe 是宏观颜色的决定性元素：Fe³⁺（赤铁矿）赋予红褐色调，Fe²⁺赋予蓝灰-暗色调，氧化还原条件的时空变化直接控制颜色分异；',
    'Fe/Mn 比值是微观发光特征的核心参数：Mn²⁺激活阴极发光（橙红-玫瑰红），Fe²⁺猝灭发光，二者比值记录了成岩流体的氧化还原演化历史；',
    '有机质含量差异和晶粒大小差异构成次级致色因素，与Fe/Mn共同叠加形成肉眼可见的多色斑块；',
    '致色元素的赋存位置（晶格内 vs. 晶界 vs. 独立矿物）和赋存时间（早成岩期 vs. 晚成岩期 vs. 表生期）是理解颜色成因的重要维度。'
]
for c in coloring:
    add_bullet(c)

doc.add_heading('7.3 研究展望', level=2)
outlook = [
    '高精度微区分析（LA-ICP-MS mapping、SIMS、nanoSIMS）的应用将实现对致色元素在微米-纳米尺度上分布的精确定量；',
    '非传统稳定同位素（如δ²⁶Mg、δ⁴⁴Ca）可为白云石化流体的来源与演化提供新的约束；',
    '数值模拟（反应-运移模型）可定量评估流体性质、流动速率与白云石化斑块形态之间的关系；',
    '微生物分子化石（生物标志化合物）的研究可进一步揭示微生物活动在早期白云石化与颜色分异中的作用。'
]
for o in outlook:
    add_bullet(o)

# ============================================================
# 末尾说明
# ============================================================
doc.add_paragraph()
add_para('— 报告完 —', italic=True, size=10)
doc.add_paragraph()
add_para('说明：本报告基于公开学术数据库（知网CNKI、Semantic Scholar、GeoScienceWorld、ScienceDirect等）检索结果整理生成。建议通过学校图书馆数据库获取文献全文，并使用 DOI 在 doi.org 验证文献真实性。', italic=True, size=9)
add_para(f'报告生成日期：{datetime.date.today().strftime("%Y年%m月%d日")}', italic=True, size=9)

# ============================================================
# 保存
# ============================================================
output_path = r'c:\Users\郝\Desktop\claude\output\色斑白云岩成因与致色原因_文献调研报告.docx'
doc.save(output_path)
print(f'报告已保存至: {output_path}')