# -*- coding: utf-8 -*-
"""Build thesis draft with 6-chapter structure and 100 real references."""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

style = doc.styles['Normal']
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rPr.insert(0, rFonts)
rFonts.set(qn('w:ascii'), 'Times New Roman')
rFonts.set(qn('w:hAnsi'), 'Times New Roman')
rFonts.set(qn('w:eastAsia'), '宋体')
rFonts.set(qn('w:cs'), 'Times New Roman')

def set_font(run, western='Times New Roman', east_asian='宋体', size=None, bold=None):
    rPr = run._r.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rPr.insert(0, rf)
    rf.set(qn('w:ascii'), western)
    rf.set(qn('w:hAnsi'), western)
    rf.set(qn('w:eastAsia'), east_asian)
    rf.set(qn('w:cs'), western)
    if size: run.font.size = Pt(size)
    if bold is not None: run.bold = bold

def h1(text):
    h = doc.add_heading(text, level=1)
    for r in h.runs: set_font(r, 'Arial', '黑体')
def h2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs: set_font(r, 'Arial', '黑体')
def h3(text):
    h = doc.add_heading(text, level=3)
    for r in h.runs: set_font(r, 'Arial', '黑体')
def p(text, indent=True, size=12):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.5
    if indent: para.paragraph_format.first_line_indent = Pt(24)
    r = para.add_run(text)
    set_font(r, size=size)
def pref(text):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.15
    para.paragraph_format.first_line_indent = Pt(0)
    r = para.add_run(text)
    set_font(r, size=10)

# ======== TITLE PAGE ========
for _ in range(4): doc.add_paragraph()
pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = pp.add_run('硕士学位论文'); set_font(r, 'Arial', '黑体', 28, True)
doc.add_paragraph()
pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = pp.add_run('用于压裂裂缝监测的荧光压裂液体系构建与性能研究'); set_font(r, 'Arial', '黑体', 18, True)
doc.add_paragraph()
for t in ['培养单位：成都理工大学能源学院','专    业：石油与天然气工程','研究方向：油气田开发','导    师：李娜','研 究 生：郝乐乐']:
    pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pp.add_run(t); set_font(r, size=14)
doc.add_page_break()

# ======== 摘要 ========
h1('摘  要')
p('水力压裂是非常规油气资源商业化开发的核心技术，压裂裂缝的延伸范围直接决定储层改造效果。然而，现有裂缝监测技术——微地震监测、分布式光纤传感（DAS/DTS）、示踪剂监测等——均属间接测量范畴：微地震依赖速度模型反演声发射事件，DAS依赖光纤-地层耦合的应变响应，示踪剂依赖返排浓度曲线反演。这些方法回答的是"裂缝可能在哪儿"，而非"压裂液实际到了哪儿"。尤其对于压后取心这一裂缝诊断最直接的验证途径，现有技术难以在岩心尺度上留下可辨识的压裂液波及标记，致使裂缝模型的可靠性始终未经实物校准。')
p('针对上述问题，本文以无机稀土铝酸盐长余辉荧光粉SrAl₂O₄:Eu²⁺,Dy³⁺为示踪介质，以羟丙基胍胶（HPG）压裂液为载体，通过表面化学改性与体系工程构建，研制了一种与现有压裂工艺兼容、以压后取心紫外直接观察为检测手段的荧光压裂液体系。选择该无机荧光粉的技术逻辑在于：其发光中心Eu²⁺受刚性晶格保护，对过硫酸铵氧化破胶剂具有本征化学惰性——该特性保证了示踪材料在破胶返排全过程中的荧光信号完整性，是本方案区别于已有荧光示踪技术（有机荧光树脂体系及半导体量子点微球体系）的核心优势。本文围绕"材料改性—体系构建—动态验证—工程转化"主线开展工作。在分析目标上，将荧光信号的判读定位于"压裂液波及范围实物验证"——发亮即表明压裂液到过该区域——而非裂缝几何参数的绝对定量测量，这一工程定位使方法的可靠性建立在坚实的实证基础上。')
p('关键词：水力压裂；裂缝监测；荧光示踪；稀土铝酸盐荧光粉；胍胶压裂液；表面改性', indent=False)
doc.add_page_break()

# ======== ABSTRACT ========
h1('Abstract')
p('Hydraulic fracturing is the enabling technology for commercial development of unconventional oil and gas resources. The spatial extent of hydraulic fractures directly governs the stimulated reservoir volume and, consequently, well productivity. Yet, all existing fracture diagnostic techniques—microseismic monitoring, distributed fiber optic sensing (DAS/DTS), and chemical or radioactive tracers—operate on the principle of indirect inference: microseismic events are inverted through velocity models, DAS strain signals depend on fiber-formation coupling, and tracer return concentrations require flowback deconvolution. These methods answer "where fractures might be," not "where the fracturing fluid actually reached." This gap between geophysical inference and physical evidence is particularly acute for post-frac coring—the most direct validation method available—where no existing technique can leave an identifiable mark on the core surface to indicate fluid sweep extent.')
p('This thesis addresses this gap by developing a fluorescent hydraulic fracturing fluid system that is fully compatible with conventional HPG (hydroxypropyl guar) fracturing operations and enables direct visual verification of fluid sweep through post-frac core examination under ultraviolet (UV) light. The tracer material is SrAl₂O₄:Eu²⁺,Dy³⁺, an inorganic rare-earth aluminate persistent phosphor. The core rationale for selecting this material—rather than the organic fluorophores or semiconductor quantum dots used in prior work—is its intrinsic chemical inertness toward the oxidative breaker system (ammonium persulfate). The Eu²⁺ luminescent center is protected within a rigid aluminate crystal lattice, ensuring signal integrity throughout the entire fracturing and flowback process. The analytical objective is positioned as "fluid sweep verification through physical evidence"—luminescence on the core surface indicates where the fracturing fluid has been—rather than absolute quantitative measurement of fracture geometry. This engineering-oriented positioning grounds the method\'s reliability in direct empirical observation.')
p('Keywords: hydraulic fracturing; fracture monitoring; fluorescent tracer; rare-earth aluminate phosphor; guar fracturing fluid; surface modification', indent=False)
doc.add_page_break()

# ======== 第一章 绪论 ========
h1('第一章  绪  论')
h2('1.1  研究背景与意义')
p('水力压裂是实现页岩油气、致密砂岩气及煤层气等非常规油气资源商业化开发的核心技术。自1949年首次商业应用以来，全球已完成逾250万次压裂施工，该技术累计为美国新增可采石油储量约30%、天然气储量约90%[1]。在压裂作业中，水力裂缝的延伸长度、高度、宽度、方位及其网络复杂程度直接决定了储层改造体积（SRV）的大小，而SRV被认为是影响非常规储层产能的关键控制参数[2-4]。准确获取压裂裂缝的空间展布特征，对优化压裂设计、评估增产效果及合理部署井网均具有重要的工程价值。')
p('当前应用于裂缝监测的技术手段主要包括微地震监测、分布式光纤传感（DAS/DTS）、放射性示踪剂、化学示踪剂、井中倾斜仪及电磁法等[5-6]。各类方法在裂缝形态识别方面各有优势与局限。Warpinski（1996）在其裂缝诊断经典综述中指出，裂缝的精确诊断是改善增产效果和降低成本的关键前提[7]。然而，de Melo和Carballo Cabrera（2025）系统评价多种裂缝成像技术后指出：没有任何单一技术能提供完整的裂缝几何信息[8]。更为根本的是，这些技术均属于间接测量范畴——微地震依赖速度模型反演声发射信号[9-11]，DAS依赖光纤-地层耦合应变响应[12-14]，示踪剂依赖返排浓度曲线反演[15-17]——所获取的裂缝信息均缺乏原位、直观的物理证据。尤其对于压后取心这一最直接的验证途径，现有技术难以在岩心尺度上留下可辨识的压裂液波及标记，裂缝诊断结果始终缺乏"可取证"的实物依据。')
p('荧光示踪技术为弥补上述"信号-实物"鸿沟提供了可行路径：向压裂液中引入荧光标记材料，使其随携砂液运移至裂缝中并在壁面附着，压后取心在紫外光下直接观察荧光分布，由此获得压裂液波及范围的实物证据。已有研究从不同技术路径探索了这一方向。Ishida等（Takeuchi等，2025）在露头尺度压裂实验中采用可凝固荧光树脂作为压裂介质，通过同轴套取岩心并紫外成像，清晰识别了主裂缝与分支裂缝的延伸轨迹，并发现部分荧光渗透区域未产生声发射事件——直接证明了间接监测确实遗漏了部分裂缝信息[18-19]。Guryanov等（2019）开发的GeoSplit量子点荧光聚合物微球系统已在分段压裂产液剖面监测中实现现场应用[20]。')
p('然而，上述已有方案与常规水力压裂工艺之间存在兼容性鸿沟。Ishida等的荧光树脂属于可凝固体系，其固化机制与常规压裂液的破胶返排工艺不兼容，本质上是一种岩石力学实验方法。Guryanov等的量子点微球依赖返排液间接分析，无法在裂缝壁面上留下可被取心验证的实物标记，且量子点在过硫酸铵氧化环境中存在降解风险。两种方案均未解决一个核心工程问题：如何在常规胍胶压裂液体系内，将一种对破胶环境稳定、可随携砂液泵注、且能在裂缝壁面牢固锚定的荧光示踪材料高效输送至裂缝各部位，使取心在紫外光下直接呈现压裂液的实际波及范围。')
p('本研究以SrAl₂O₄:Eu²⁺,Dy³⁺无机长余辉荧光粉为示踪介质、以HPG胍胶冻胶为压裂液载体，构建与现有压裂工艺完全兼容、以取心实物为直接证据的压裂液波及范围可视化方法。该技术定位为微地震和DAS/DTS等间接监测方法的辅助校准工具——通过压后取心的实物证据验证间接反演结果的可靠性，而非取代已有监测技术。')

h2('1.2  国内外研究现状')
h3('1.2.1  水力压裂裂缝监测技术')
p('裂缝监测技术在过去二十年间经历了从单井温度测井到多井光纤传感、从地面倾斜仪到井中微地震阵列的快速发展，但各类方法围绕的核心问题始终是一个：如何准确获取裂缝的延伸长度、高度和方位。')
p('分布式光纤传感是近十年进展最快的方向。Molenaar等（2012）首次将DAS技术部署于压裂井井下，验证了声学光纤对桥塞坐封、射孔和泵注等井下事件的实时识别能力[21]。这标志着光纤从温度传感（DTS）向声学/应变传感（DAS）的关键跨越。Sierra等（2008）建立了基于DTS温度暖回（Warm-back）分析的裂缝流体分配解释方法[22]，该方法至今仍是压裂段产液诊断的常用手段。Ugueto等（2016）将DAS与DTS联合部署于同一口井，通过低频DAS信号揭示了多簇射孔完井中各簇进液量的显著非均匀性——部分射孔簇几乎不进液——这一发现直接推动了"限流压裂"向"均匀布孔"的完井理念转变[23]。Karrenbach等（2017）在同一光纤缆线中同步采集了DAS微地震信号和低频应变信号，从原理上证明了单根光纤实现多参数监测的可行性[24]。在裂缝几何参数反演方面，Jin和Roy（2017）的工作具有里程碑意义：他们建立了低频DAS（<0.05 Hz）的应变信号与裂缝长度、高度、宽度及密度之间的定量对应关系，并通过数值模拟和现场数据双重验证了该方法的可靠性[12]。Ugueto等（2019）进一步将DAS应变锋面（Strain Front）概念引入压裂诊断——应变锋面的传播速度和空间展布可直接反映裂缝的延伸范围、段间封隔有效性以及最优井距[25]。')
p('微地震监测是裂缝诊断中应用最广泛、研究积累最深厚的方法。Maxwell等（2002）在Barnett页岩中首次利用井中微地震阵列揭示了水力压裂形成的复杂裂缝网络——观察到裂缝在优势方向（NE-SW）和次级方向（NW-SE）上均有扩展，打破了传统单平面双翼裂缝的理论框架[26]。Fisher等（2004）在同一盆地利用微地震裂缝测绘优化了水平井完井参数，并建立了累积产量与裂缝网络总长度/总面积——而非单裂缝半长——之间的统计关联，这在当时是一个颠覆性的认识[3]。Mayerhofer等（2010）在此基础上正式提出并系统定义了储层改造体积（SRV）的概念，即利用微地震事件云的三维外包络体积来近似表征裂缝网络的宏观空间范围，为非常规储层的压裂效果评价提供了一个可操作的量化指标[2]。十余年间，SRV已成为业内使用最广泛的概念之一。Cipolla和Maxwell等（2011）发表了微地震解释实用指南，系统构建了从数据质量控制、事件定位、不确定性量化到工程应用的完整工作流[9]；同年，Cipolla等整合了微地震测绘与半解析/非连续裂缝扩展模型，实现了裂缝复杂度参数的联合约束反演[27]。然而，Cipolla和Wallace（2014）对SRV概念进行了批判性重新审视：SRV本质上对裂缝的实际导流面积和导流能力缺乏直接约束——这两个参数才是控制产能的根本——建议将微地震体积（MV）作为更贴近物理基础的替代指标[4]。Maxwell（2014）在SEG专著中系统梳理了微地震在裂缝几何反演、储层表征和压裂工程优化中的理论基础和工程实践[28]。')
p('示踪剂技术的发展则呈现从放射性向非放射性、从单一标记向多编码、从定性向半定量演进的趋势。Salman等（2014）通过分析非常规储层化学示踪剂返排数据，建立了示踪剂回收率与裂缝有效性之间的关联：回收率高的压裂段对应高裂缝密度区，回收率低则提示裂缝网络稀疏或与井筒连通性差[15]。Tian等（2019）通过精细的分序注入实验——在单个Wolfcamp B压裂段中依次注入12种水相示踪剂和12种油相示踪剂——发现所有24种示踪剂在返排中几乎同时且显著混合，证明裂缝网络内的流体流动远非活塞式驱替，被污染的示踪剂信号可能导致裂缝有效性的定量解释产生重大误差[29]。在纳米示踪剂前沿方向，Kosynkin和Alaskar（2016）报道了碳点基纳米示踪剂（A-Dots）在碳酸盐岩油藏中的首次井间现场试验——约1吨纳米颗粒注入后约10个月在500 m外的产油井中检出，验证了纳米颗粒在高温（100°C）高盐（150,000 ppm）储层中的长距离运移能力[31]。Hu等（2019）利用碳量子点在砂岩岩心驱替中展示了类示踪剂的理想穿透行为——低吸附、低弥散、高回收率——并实现了定量化的油饱和度检测[32]。Arshad等（2024）报道了新型化学示踪剂在裂缝段产能剖面构建方面的首次全井段现场应用[30]。这些工作为本研究采用无机荧光微米粉作为固体示踪介质提供了重要的方法学基础。')
p('倾斜仪与电磁法作为辅助手段，各有其应用场景。Wright等（1998）引入的井下倾斜仪裂缝测绘（DTFM）可通过邻井阵列实现裂缝几何的约束反演[33]；Warpinski等（2006）发展了微地震-倾斜仪联合反演以提高裂缝几何估计的可靠性[34]。LaBrecque等（2016）对导电支撑剂电磁成像进行了取心+槽探双重地面真值验证[35]；Ahmadian等（2023）展示了导电支撑剂电磁法的裂缝实时动态监测潜力[36]。Cipolla和Wright（2002）在其裂缝诊断综述中对上述技术进行了系统比较[5]。')
p('综合来看，这些进展虽然显著，但任何单一技术都面临一个共性的工程约束：获取的信号需要从"物理量"转化为"裂缝几何"，这个转化过程不可避免地引入多解性。正是这种间接反演的结构性局限，促使研究者开始探索能够提供实物证据的互补性方案。')

h3('1.2.2  荧光示踪裂缝可视化技术')
p('在"直接可视化"路线方面，Chen等（2014）率先在实验室采用含荧光涂料的MMA热固性树脂实现裂缝的可视化[19]。Takeuchi等（2025）在露头尺度进行了系统验证：在76 mm钻孔中注入可凝固荧光树脂后，通过205 mm同轴套取岩心并紫外成像，不仅清晰识别了裂缝延伸轨迹，更通过与AE事件分布的对比揭示了声发射监测遗漏的裂缝区域——直接证明间接监测确实遗漏了部分裂缝信息[18]。Flury和Wai（2003）在Reviews of Geophysics发表的综述系统总结了荧光染料在地下水流可视化中的百年应用历史，为压裂裂缝的荧光示踪提供了方法学基础[37]。')
p('在"返排分析"路线方面，Guryanov等的GeoSplit量子点微球系统是该方向最具工程成熟度的方案[20]。Kang等（2015）制备了兼具调剖和示踪功能的荧光聚丙烯酰胺微球[38]。Hu等（2019）首次展示了碳量子点作为储层示踪剂的潜力[32]。Spitzmuller等（2024）报道了荧光染料包埋介孔硅纳米颗粒示踪剂，表面修饰后可在160°C、高盐条件下稳定示踪[39]。Kosynkin和Alaskar（2016）的A-Dots纳米示踪剂在500 m井距的碳酸盐岩油藏中实现了突破——将约1吨纳米颗粒注入后，约10个月后在产油井中检测到示踪信号[31]。')
p('上述两条路线的共同局限在于：直接可视化路线使用的可凝固树脂与常规压裂工艺不兼容，返排分析路线依赖间接信号而无法在裂缝壁面上留下可被取心验证的实物标记。两条路线均未解决"在常规胍胶压裂液体系内直接标记裂缝壁面"这一核心工程问题。')

h3('1.2.3  稀土铝酸盐荧光材料及其表面改性')
p('Matsuzawa等（1996）在Journal of the Electrochemical Society首次报道了SrAl₂O₄:Eu²⁺,Dy³⁺的长余辉特性——余辉亮度较传统ZnS:Cu,Co高出10倍以上，余辉持续数小时，无需放射性元素掺杂[40]。Clabau等（2005）基于DFT计算对Matsuzawa模型进行了机制修正[41]。Dorenbos（2005）从镧系元素能级热力学角度解释了陷阱深度来源[42]。Li等（2019）在Journal of Luminescence发表了从Bologna石（1602年）到Cr³⁺掺杂自旋体的完整历史综述[43]。Rojas-Hernandez等（2018）系统比较了不同合成方法（固相法、溶胶-凝胶、水热、燃烧）对荧光粉粒度和发光性能的影响，指出固相法（1300~1400°C）是最广泛使用的工业化路线[44]。')
p('表面包覆是解决荧光粉水性环境水解问题的工程手段。Guo等（2007）发现未处理SrAl₂O₄在水中12小时内pH从7升至12.4，水解产物包括SrAl₄O₇等晶相[45]。Karacaoglu等（2020）报道未包覆荧光粉浸水48小时即丧失磷光[46]。在无机包覆方面，Qi等（2017）通过尿素水解原位实现了致密SiO₂层均匀包覆（≥5 wt%）[47]。Karacaoglu等（2020）采用ALD沉积~10 nm Al₂O₃或~12 nm TiO₂薄膜，包覆后48小时浸水仍保持磷光[46]。Deng等（2013）通过燃烧法原位形成MAl₂B₂O₇致密层[48]。在有机改性方面，Urakawa等（2020）证明硅烷偶联剂可使荧光粉发光强度提升10~20%且不影响光学特性[49]。Lyu等（2020）以KH570为桥联剂构建了SiO₂-聚合物复合壳，实现了耐水性和有机相容性的平衡[50]。')
p('然而，上述包覆方案均面向涂料、塑料等行业开发，针对压裂液工程环境（高矿化度、高剪切、宽温度、氧化介质）的系统验证尚属空白。更关键的是，压裂工程对颗粒表面提出了独特的功能时序需求——注入期间需稳定悬浮，关井破胶期间需向锚定态切换，返排期间需牢固附着——如何在单一改性方案中协调这种时序功能切换，是兼具工程和科学价值的问题。')

h2('1.3  存在的主要问题')
p('综合国内外研究现状，当前存在以下三个层次的技术缺口：')
p('第一，材料层面——高密度微米荧光粉（3.6~4.0 g/cm³）在HPG基液（~1.0 g/cm³）中的悬浮稳定性问题未解决，缺乏面向压裂液工程的表面改性方案。荧光粉表面溶出的多价阳离子（Al³⁺、Ca²⁺）可能干扰胍胶交联反应，影响冻胶携砂和造缝性能。')
p('第二，工艺层面——注入（稳定悬浮）和关井后（壁面锚定）对颗粒表面化学的需求截然相反，缺乏利用压裂施工自身时序（注入→关井→返排）驱动功能切换的工程化方案。')
p('第三，评价层面——已有研究多采用静态浸泡或简单注入后直接成像评价示踪效果，跳过了注入剪切、关井破胶和返排冲刷三个工程必经环节。在动态驱替模拟压裂-返排全过程的条件下，荧光示踪效果是否依然成立，尚未被验证。')

h2('1.4  研究目标、内容与技术路线')
p('本研究以SrAl₂O₄:Eu²⁺,Dy³⁺无机长余辉荧光粉为示踪介质，以HPG压裂液为载体，以压后取心紫外直接观察为检测手段，构建一种与现有压裂工艺完全兼容的压裂液波及范围实物验证方法。')
p('研究内容包括以下六个方面：（1）荧光粉基础物性表征与储层环境适应性评价，明确井下工况的适用边界。（2）双层表面改性工艺优化与协同分散机理研究，实现高密度微米粉在HPG基液中的长期均匀悬浮。（3）高浓度荧光母液研制与压裂液体系标准化性能评价，确保对压裂液工程性能无显著损害。（4）"分散-吸附"功能切换机制与砂岩壁面吸附规律研究，揭示破胶前后表面化学变化的界面机制。（5）动态驱替模拟压裂-返排全过程与荧光信号半定量分析，验证可注入性、附着牢固度和空间对应性。（6）现场施工工艺方案设计与技术经济评估。')
p('技术路线以1000目（~13 μm）商用SrAl₂O₄:Eu²⁺,Dy³⁺荧光粉为示踪核心，以工业级HPG为稠化剂，有机硼为交联剂，过硫酸铵为破胶剂。各研究内容按照"材料改性→体系构建→动态验证→工程转化"的主线递进展开，最终形成一套经动态驱替验证和工程可行性评估的荧光压裂液示踪方案。')
doc.add_page_break()

# ======== 第二章 ========
h1('第二章  荧光粉基础物性及表面改性')
p('本章围绕荧光粉在压裂液体系中的适用性，依次解决三个基础问题：商用SrAl₂O₄:Eu²⁺,Dy³⁺荧光粉的基础物性特征与发光性能基线，储层温度-矿化度-氧化环境联合作用下的性能衰减规律与适用边界，以及通过双层表面改性和协同分散实现高密度微米粉在HPG基液中的长期稳定悬浮。这三个问题构成后续压裂液体系构建的材料科学基础。')
h2('2.1  荧光粉基础物性表征')
p('在开展任何改性工作之前，首先需要建立改性效果和环境适应性评价的参照基线。本章实验所用的SrAl₂O₄:Eu²⁺,Dy³⁺长余辉荧光粉为商用产品（固相法合成，1000目，D50≈13 μm）。基础物性表征包括：X射线衍射（XRD）分析晶相组成（使用Cu Kα辐射源，2θ扫描范围10°~80°）；扫描电子显微镜（SEM-EDS）观察颗粒形貌、粒径分布和元素组成；激光粒度分析仪测定粒径分布（D10、D50、D90及跨距Span值）；荧光光谱仪（稳态激发和发射光谱、余辉衰减曲线）测定发光特性。')
p('[实验数据待补充：XRD图谱、SEM照片、粒度分布曲线、激发/发射光谱、余辉衰减曲线]')

h2('2.2  储层环境适应性评价')
p('该荧光粉能否在压裂施工的时间窗口内——从注入到关井再到破胶返排，通常持续数小时至数天——保持足够的荧光信号，是整个技术方案能否成立的先决性问题。为此，本节系统考察荧光粉在模拟储层条件下的性能衰减规律，明确其在井下工况中的适用边界。实验设计如下：')
p('（1）热稳定性：将荧光粉在60、90、120、150°C的烘箱中分别放置0、24、72、168小时，冷却至室温后测定发射光谱积分面积（激发波长365 nm），计算相对发光保持率（%）。每个条件3个平行样。')
p('（2）化学稳定性-矿化度：将荧光粉浸泡于不同NaCl浓度（0、10、50、100 g/L）的溶液中，90°C恒温，在0、24、72、168小时取样测定发射光谱积分面积。同步测定浸泡液的pH值以监测水解程度。')
p('（3）化学稳定性-氧化环境：将荧光粉浸泡于不同浓度过硫酸铵溶液（0、0.05%、0.1%、0.2% w/v）中，90°C恒温，在0、6、12、24、48小时取样测定发射光谱积分面积。')
p('（4）化学稳定性-pH：将荧光粉浸泡于不同pH值缓冲溶液（pH 3、5、7、9、11），90°C恒温，按上述时间点取样测定。')
p('每项实验拟合相对发光强度-时间的一级衰减动力学：ln(I/I₀) = -kt，获取速率常数k和半衰期t₁/₂。')
p('[实验数据待补充]')

h2('2.3  双层表面改性工艺优化')
p('表面改性的目标是在荧光粉表面构建一个兼具三重功能的双层结构——内层耐水解的化学屏障、外层的空间位阻分散层、以及抑制多价阳离子溶出的离子屏蔽层。本节通过正交实验系统优化KH550偶联剂和PEG4000的用量、分子量及工艺条件，实现"内层KH550化学锚固+外层PEG4000物理屏蔽"，实现耐水性（化学保护层）、分散性（空间位阻层）和离子屏蔽（螯合层）三重功能。')
p('改性步骤：（1）KH550预处理——将荧光粉分散于无水乙醇-水混合溶剂（体积比95:5）中，加入不同用量KH550（相对于荧光粉质量0、0.5、1.0、2.0、3.0 wt%），乙酸调节pH至4~5以促进硅烷水解，室温搅拌2 h后60°C干燥固化1 h。（2）PEG包覆——将KH550预处理粉末分散于去离子水中，加入不同分子量（PEG2000、PEG4000、PEG6000）和浓度（0、1、3、5 wt%）的PEG，室温搅拌1 h后离心、去离子水洗涤、60°C干燥。')
p('工艺参数优化采用正交实验设计（L9(3⁴)）：考察KH550用量（1.0/2.0/3.0 wt%）、PEG分子量（2000/4000/6000）、PEG浓度（1/3/5 wt%）和搅拌时间（30/60/90 min）四个因素三个水平。评价指标以改性粉在HPG基液（0.5 wt% HP 温度25°C）中静置2小时后的相对浊度保持率（%）为主指标，Zeta电位绝对值为辅助指标。')
p('改性效果表征：傅里叶变换红外光谱（FTIR，KBr压片法，确认Si-O-Al键和C-H伸缩振动峰）、X射线光电子能谱（XPS，确认N 1s和Si 2p峰的变化）、热重分析（TGA，N₂气氛，室温至800°C，10°C/min，测定有机物包覆量）和SEM对比改性前后形貌。')
p('[实验数据待补充：正交实验结果表、FTIR/XPS/TGA图谱]')

h2('2.4  改性粉在胍胶基液中的分散稳定性')
p('双层改性是分散的基础，但仅靠双层包覆可能不足以解决高密度微米粉在低密度胍胶基液中的长期悬浮问题。本节旨在通过六组对照实验，系统量化柠檬酸（螯合剂）、非离子表面活性剂（润湿剂）和外加PEG（强化空间位阻）三种添加剂对分散稳定性的独立贡献，并检验三者之间是否存在超加和性的协同效应——即组合效果是否显著大于各组分单独效果之和。')
p('组1：双层改性粉 + 0.5 wt% HPG基液（空白对照）; 组2：双层改性粉 + 0.5 wt% HPG + 0.1 wt% 柠檬酸; 组3：双层改性粉 + 0.5 wt% HPG + 0.05 wt% Triton X-100; 组4：双层改性粉 + 0.5 wt% HPG + 5 wt% PEG4000（额外添加）; 组5：双层改性粉 + 0.5 wt% HPG + 柠檬酸 + Triton X-100（两两组合）; 组6：双层改性粉 + 0.5 wt% HPG + 柠檬酸 + Triton X-100 + PEG4000（三者全加）')
p('每组测定：（a）静置沉降曲线——将各分散液置于25 mL量筒中，间隔一定时间（0、15、30、60、120 min）记录沉降界面高度，计算相对浊度保持率；（b）激光粒度——D50和Span值随时间的变化（反映团聚动力学）；（c）Zeta电位-pH曲线——反映颗粒表面电荷随pH的变化和分散稳定性。')
p('通过组间对比，量化各组分的独立贡献和协同效应：柠檬酸的贡献=组2-组1，Triton的贡献=组3-组1，PEG的贡献=组4-组1，协同效应=组6的分散效果-(组2+组3+组4-3×组1)。若协同效应项为正且统计显著，则证明三者之间存在超加和性的协同分散机理。')
p('[实验数据待补充]')

h2('2.5  改性对胍胶交联的影响')
p('改性荧光粉表面可能残余微量的可溶出Al³⁺等多价阳离子——这些离子可与HPG分子链上的顺式邻二醇发生配位交联，干扰有机硼与HPG的正常交联反应。交联时间延长或冻胶强度下降是工程上不可接受的。为此，本节对比含/不含改性荧光粉的HPG冻胶的交联性能（1）交联时间测定：采用旋转粘度计，记录加入交联剂后粘度开始上升至最大值的90%所需时间。（2）粘弹性测定：采用高温高压流变仪（振荡模式，频率1 Hz，应变1%），在25~120°C温度范围内测定储能模量G\'和损耗模量G"的温度扫描曲线。（3）SEM观察冻胶微观结构：冷冻干燥后SEM观察含/不含荧光粉的冻胶网络形貌。')
p('[实验数据待补充]')

h2('2.6  本章小结')
p('[待实验完成后撰写]')

doc.add_page_break()

# ======== 第三章 ========
h1('第三章  荧光压裂液体系构建与性能评价')
p('在确认改性荧光粉的材料可行性之后，本章将其配制成可工程化应用的压裂液体系，并依据中国石油天然气行业标准SY/T 6376-2008和SY/T 5107-2016，对基液性能、交联冻胶性能、悬砂性能、破胶性能及地层伤害程度进行逐项检验。核心关切是：荧光粉的引入是否导致压裂液工程性能出现不可接受的劣化。')
h2('3.1  高浓度荧光母液制备与储存稳定性')
p('现场施工中，干粉直接加入压裂液不仅粉尘污染严重，而且难以保证荧光粉在储罐中的均匀分散——这正是"母液预配"工艺设计的初衷。本节确定母液的最优配制浓度和工艺参数，建立储存稳定性的评价方法和合格标准。母液制备步骤：将最优配方改性荧光粉（来自第二章优化结果）按目标浓度（40、50、60 g/L）分散于去离子水中，依次加入柠檬酸和Triton X-100，高速均质机（10000~15000 rpm×5 min）剪切分散，超声辅助（40 kHz×10 min）脱气泡并促进分散。')
p('母液质量评价：（1）静态稳定性——将母液静置于25 mL量筒中，间隔24 h记录分层界面高度，目标：7天分层高度比≤5%。（2）动态稳定性——低速离心沉降法（2000 rpm×10 min），测定离心前后浊度比（Turbidity Ratio，TR=T/T₀），目标TR≥0.90。（3）荧光强度稳定性——取离心前后母液上清液在365 nm激发下测定520 nm发射峰积分面积，计算保持率。')
p('[实验数据待补充]')

h2('3.2  基液性能评价')
p('荧光母液引入HPG基液后，首先需要确认基液的基础物理性质——粘度、pH和密度——是否因荧光母液的添加而发生显著偏离。若偏离超出常规HPG压裂液的范围，后续的交联配方和破胶设计将失去工程参考基准。为此，将荧光母液按设计体积比（0.5% v/v）与HPG基液（0.5 wt% HPG水溶液）复配，得荧光压裂液终液。依据SY/T 5107-2016《水基压裂液性能评价方法》测定基液表观粘度（六速旋转粘度计，170 s⁻¹，25°C）、pH值和密度，并与不含荧光粉的空白HPG基液进行对比。')
p('[实验数据待补充]')

h2('3.3  交联冻胶性能评价')
p('冻胶的粘弹性和热稳定性直接决定了裂缝的造缝宽度和支撑剂的输送距离——这是压裂液最基本、也是最不可妥协的工程功能。引入荧光粉后，如果冻胶的耐温耐剪切能力显著劣化，无论荧光示踪效果多好，该方案都不具备工程价值。为此，向荧光压裂液终液中加入有机硼延缓交联剂（0.3% v/v），依据SY/T 5107-2016测定以下性能：')
p('（1）交联时间：以旋转粘度计在恒定剪切速率下观察粘度剧烈上升的时间点。')
p('（2）耐温耐剪切流变特性：采用高温高压流变仪（密闭测量系统，170 s⁻¹恒定剪切），按温度程序（从25°C以3°C/min升温至120°C，恒温60 min或至粘度降至50 mPa·s以下）记录表观粘度-温度/时间曲线。评价标准：在目标储层温度下，冻胶表观粘度应≥50 mPa·s维持60 min以上（SY/T 6376-2008）。')
p('（3）粘弹性：振荡频率扫描（0.1~10 Hz，应变1%，25~120°C），记录G\'和G"随频率的变化，确定线性粘弹区范围和凝胶强度。参照Gomaa等（2015）[51]和Goel与Shah（2001）[52]提出的弹性主导准则（G\'>G"）评估支撑剂悬浮能力。')
p('[实验数据待补充]')

h2('3.4  悬砂性能评价')
p('支撑剂在冻胶中的沉降速度决定了裂缝内支撑剂的最终铺置剖面：沉降过快将导致裂缝上部无支撑剂充填，改造效果大打折扣。依据SY/T 6376-2008，采用静态悬砂法评价荧光冻胶的支撑剂悬浮能力在含荧光冻胶的量筒中放入20/40目陶粒支撑剂（密度~2.7 g/cm³，浓度480 kg/m³），在室温/储层温度下静置，每隔一定时间记录支撑剂沉降界面高度，计算支撑剂沉降速度（mm/min），并与空白HPG冻胶对比。评价标准：在储层温度下支撑剂沉降速度应≤0.5 mm/min（参照行业经验值）。')
p('[实验数据待补充]')

h2('3.5  破胶性能评价')
p('压裂液破胶的彻底性直接影响裂缝导流能力的恢复和投产进度：破胶不彻底将残留高粘度冻胶堵塞裂缝，残渣过高则堵塞支撑剂充填层的孔隙吼道。荧光粉作为不可降解的无机固体颗粒，其与瓜尔胶破胶残渣的叠加效应可能加剧支撑剂充填层的伤害程度。本节依据SY/T 5107-2016，向荧光冻胶中加入过硫酸铵破胶剂（0.05~0.2% w/v浓度梯度），在不同温度（60~120°C）下密闭破胶定时取样，冷却至室温后测定破胶液表观粘度（毛细管粘度计），绘制粘度-破胶时间曲线，确定破胶时间（粘度降至10 mPa·s以下所需时间）。完全破胶后，依据SY/T 5107-2016测定破胶液残渣含量（过滤、干燥、称重法，mg/L），并与空白HPG冻胶对比。')
p('[实验数据待补充]')

h2('3.6  与胍胶体系的配伍性及地层伤害评价')
p('压裂液功能添加剂的工程接受门槛不在于其功能的强弱，而在于对储层是否造成不可接受的伤害。13 μm的荧光粉颗粒属于不可溶、不可降解的惰性固体，在裂缝壁面和支撑剂充填层中的残留可能导致近井地带渗透率和导流能力的下降。本节通过滤饼伤害、岩心驱替伤害和导流室三组实验，将荧光粉独立贡献的伤害增量从HPG自身的瓜尔胶残渣伤害中定量剥离。需要指出，Al-Muntasheri（2014）[53]指出商品瓜尔胶中至少5 wt%为不溶性残渣——这是一个不可忽略的背景伤害值，荧光粉的引入是否会在此基础上产生显著的额外伤害，是本节实验设计的核心逻辑。')
p('荧光粉在裂缝中的残留伤害主要通过以下三种机制产生：（a）裂缝壁面滤饼形成——注入阶段荧光粉在滤失作用下于裂缝壁面堆积形成致密滤饼，降低裂缝壁面的有效渗透率；（b）破胶残渣协同伤害——荧光粉颗粒与瓜尔胶不溶性残渣混合后可能形成比单独残渣更致密的复合伤害层；（c）支撑剂充填层孔隙堵塞——13 μm荧光粉相对于20/40目陶粒充填层的平均孔隙直径（约30~50 μm）处于亚微米至微米级，若大量截留于充填层中，将导致导流能力下降。')
p('实验设计：')
p('（1）滤饼伤害实验：采用高温高压静态滤失仪（API滤失仪），在3.5 MPa压差、90°C条件下，测定荧光压裂液终液在天然砂岩滤饼（取自须家河组，直径2.5 cm，厚度0.5 cm）上的滤失量和滤饼厚度，与空白HPG冻胶（不含荧光粉）对比。滤失实验结束后取出滤饼，SEM观察滤饼微观形貌并EDS分析荧光粉在滤饼中的分布，判断荧光粉是否显著增厚滤饼或改变滤饼的致密度。')
p('（2）岩心驱替伤害实验：采用天然砂岩岩心柱（直径2.5 cm，长度5~7 cm），按照SY/T 6540-2002《钻井液完井液损害油层室内评价方法》的流程：先正向测定岩心初始煤油渗透率（K₀）→反向注入荧光压裂液破胶液（含残渣+荧光粉）2 PV，90°C下密闭接触2 h→正向煤油驱替测定伤害后渗透率（K_d），计算渗透率伤害率=(1-K_d/K₀)×100%。以不含荧光粉的HPG空白破胶液为对照，剥离荧光粉单独贡献的伤害增量ΔD=D_荧光-D_空白。每组实验重复3次。')
p('（3）支撑剂充填层导流能力实验：参照SY/T 6302-2009《压裂液、破胶液损害支撑剂充填层导流能力试验推荐方法》，在API导流室中铺设20/40目陶粒支撑剂单层（铺置浓度10 kg/m²），分别注入空白破胶液和含荧光粉破胶液，90°C下密闭24 h后测定导流能力-闭合压力曲线（闭合压力10~60 MPa），对比两种条件下的导流能力差异。')
p('[实验数据待补充]')

h2('3.7  本章小结')
p('[待实验完成后撰写]')
doc.add_page_break()

# ======== 第四章 ========
h1('第四章  裂缝壁面吸附与动态驱替示踪验证')
p('前三章分别从材料、流体和体系层面验证了荧光压裂液各组分的可行性。本章在此基础上，通过裂缝模型中的静态吸附实验和动态驱替实验，对荧光示踪方案在实际裂缝条件下的有效性进行系统验证。验证的维度包括：荧光粉颗粒在裂缝受限空间内的运移与空间分布均匀性，破胶诱导的"分散态→锚定态"功能切换的化学证据与吸附定量参数，以及动态驱替模拟压裂-返排全流程条件下的可注入性、附着牢固度和紫外成像可检出性。上述维度构成从基础材料验证向工程条件验证过渡的完整实验证据链。')
h2('4.1  模拟裂缝可视化实验装置')
p('实验装置的核心是一个可调节裂缝宽度的岩心夹持系统。基质为天然砂岩岩心（取自四川盆地须家河组露头，切割为50 mm×50 mm×100 mm标准长方体，孔隙度5%~12%），通过巴西劈裂法在岩样中部预制一条贯穿性人工张拉裂缝。裂缝宽度通过不同厚度的精密金属垫片（0.1、0.5、1.0、2.0 mm）精确控制，每种宽度制备3块平行样，共12块。保留劈裂产生的自然粗糙断裂面（JRC范围约8~16），以模拟天然裂缝壁面的粗糙特征。')
p('系统其余组件包括：三套独立中间容器（500 mL耐压容器，分别盛装荧光HPG压裂液、有机硼交联剂溶液和模拟返排液）、一台高精度恒流泵（双柱塞式，流量范围0.01~50 mL/min，可编程多段注入）、一台回压阀（控制裂缝闭合压力5~30 MPa）、一套加热套（PID控温，60~120°C±1°C）和一套紫外暗室成像系统（365 nm紫外LED光源、高分辨率工业相机、520±10 nm带通滤光片）。所有成像参数（紫外灯功率、灯距、f/4.0光圈、2 s曝光、ISO 800）在全部实验中锁定不变。')

h2('4.2  裂缝内颗粒运移与分布规律')
p('荧光示踪的前提假设是：荧光粉的空间分布忠实反映了压裂液的实际波及范围。如果13 μm颗粒在动态注入过程中大量沉降堆积在裂缝入口附近、裂缝尖端基本无荧光粉到达，则取心岩样上的荧光分布将严重低估压裂液的实际波及范围——这是本方法最核心的工程风险之一。荧光粉在裂缝受限空间中的运移和分布并非简单的Stokes自由沉降，而是受剪切诱导迁移（Shear-Induced Migration）、惯性升力（Inertial Lift）、壁面效应（Wall Effect）和重力沉降多机制耦合的复杂行为。运移结果直接决定荧光粉的空间分布——若颗粒在裂缝入口附近大量沉降堆积，裂缝尖端将无荧光标记。这一现象在石油工程中与支撑剂在裂缝中的运移-沉降行为同属颗粒两相流范畴[51-52]，但目前针对亚支撑剂尺寸（~10 μm）、低浓度（<1% v/v）的示踪颗粒在裂缝中的运移规律缺乏研究数据。')
p('实验设计：利用4.1的模拟裂缝装置，在注入阶段（阶段一）过程中，沿裂缝长度方向在夹持器壁面预设至少3个压力传感器（分别位于入口端、裂缝中部和尖端附近），通过压力信号的时空演化判断颗粒在裂缝不同区域的堆积程度。注入结束后、破胶和返排之前，另取一组平行实验不进行返排——直接拆卸夹持器，取出砂岩块，沿裂缝延伸方向等分为入口段、中段、尖端段3个分区，分别用紫外成像和SEM对每段的裂缝壁面进行观察（注意成像参数统一）。对比三段壁面的荧光粉面密度（通过ImageJ灰度分析的半定量判定），判断是否存在入口段堆积、尖端段无荧光粉的空间非均匀分布。若三段灰度无显著差异（单因素ANOVA，p>0.05），则13 μm荧光粉在该裂缝宽度-流速组合下可实现相对均匀的空间分布。')
p('[实验数据待补充]')

h2('4.3  破胶前后荧光粉表面化学变化')
p('本研究的核心工作假说——破胶诱导PEG脱附、暴露KH550活性氨基从而实现功能切换——能否成立，首先需要分子层面的直接化学证据。在进入耗时的动态驱替实验之前，本节先通过静态条件下的FTIR、XPS和Zeta电位三组表征手段进行交叉验证。如果破胶处理后PEG特征峰（C-O-C，~1100 cm⁻¹）减弱、N 1s峰（~399 eV，归属-NH₂）增强、Zeta电位绝对值降低三者同时成立，则该假说获得化学层面的支持；三者均无变化则假说被否定，需启用备选触发机制。取最优双层改性荧光粉（来自第二章），分别进行以下两种处理：（a）未经破胶处理的原始改性粉；（b）在90°C、0.1%过硫酸铵溶液中浸泡12 h（模拟关井破胶阶段的条件）后取出、洗涤、干燥的破胶处理粉。对两种粉末进行FTIR（KBr压片）、XPS（C 1s、O 1s、N 1s、Si 2p窄扫描）和Zeta电位-pH曲线测定。重点关注：（i）PEG特征峰（C-O-C伸缩振动~1100 cm⁻¹，C-H伸缩振动~2880 cm⁻¹）是否减弱——指示PEG脱附；（ii）N 1s峰（~399 eV，归属于-NH₂）是否增强——指示KH550氨基暴露；（iii）Zeta电位在pH 7附近的绝对值是否降低——指示空间位阻减弱、有利于颗粒-壁面近距离接触。')
p('[实验数据待补充]')

h2('4.4  静态吸附实验')
p('4.3节的表面化学表征即便证实了PEG脱附和KH550氨基暴露，仍需回答一个动力学问题：暴露的氨基在关井破胶的有限时间窗口内（通常6~24小时）能否完成与砂岩壁面硅羟基的有效锚定。这一过程的可行性取决于两个维度：热力学上，吸附自由能ΔG°是否足够负以保证吸附的自发性和驱动力；动力学上，吸附速率常数k₂和有效扩散系数是否满足在窗口期内达到显著吸附覆盖率的时间要求。本节通过批吸附实验获取上述两个维度的定量参数。')
p('实验步骤：将预切割的天然砂岩薄片（10×10×2 mm，表面用800目砂纸统一抛光）分别置于含不同初始浓度（0.1~5.0 g/L）荧光粉悬浮液的玻璃瓶中（模拟地层水基质：50 g/L NaCl + 2 g/L CaCl₂），在25、50、80°C恒温振荡（120 rpm）至吸附平衡（预实验确定平衡时间），取样通过UV-Vis分光光度计（或荧光光谱仪）测定上清液中残余荧光粉浓度（基于浊度-浓度标准曲线），由初始浓度与平衡浓度之差计算单位面积砂岩表面上的荧光粉吸附量（μg/cm²）。')
p('数据处理：拟合Langmuir（q_e = q_max·K_L·C_e/(1+K_L·C_e)）、Freundlich（q_e = K_F·C_e^(1/n)）和Temkin等温吸附模型，以R²和AIC准则选择最优模型。由不同温度下的Langmuir吸附常数计算热力学参数：ΔG° = -RT·ln(K_L)，ln(K_L) vs 1/T作图获得ΔH°和ΔS°（Van\'t Hoff方程）。吸附动力学数据（q_t vs t）拟合拟一级（ln(q_e-q_t) vs t）、拟二级（t/q_t vs t）和颗粒内扩散模型（q_t vs t^(1/2)），识别速率控制步骤。')
p('[实验数据待补充]')

h2('4.5  动态驱替四阶段实验')
p('前述第二至四章的实验分别从材料、流体和界面三个维度回答了荧光示踪方案的要素可行性问题。但所有静态实验都不能替代一个根本问题：在动态、连续、多阶段的压裂施工模拟条件下，该方案是否依然成立？本节构建了"注入—关井破胶—返排—取心—紫外成像"五步串联流程，旨在用三组联动数据——注入压力曲线（可注入性）、返排液荧光浓度-CPV曲线（附着牢固度）和岩心紫外图像（空间对应性）——构成荧光示踪技术工程可行性的完整证据链。')
p('阶段一（注入阶段）：将已装入天然砂岩夹心裂缝模型的夹持器加热至目标储层温度（90°C），以恒定流量（根据裂缝宽度和设计剪切速率计算，τ_w目标值10~100 Pa）注入荧光HPG压裂液（含0.3% v/v有机硼交联剂，在线混配），注入总量为裂缝体积的3倍。全程在线记录注入端和产出端的压力（p_in和p_out），绘制p_in - 注入PV数曲线。若p_in随注入量单调上升并趋于平台值→无堵塞，正常；若p_in持续单调上升无平台→提示颗粒在裂缝中堵塞或架桥。')
p('阶段二（关井破胶阶段）：关闭注入端和产出端阀门，维持回压5 MPa（模拟裂缝闭合压力），90°C恒温密闭12 h，使硼交联体系在过硫酸铵作用下充分破胶。此阶段不采集数据——它是为"分散→吸附"功能切换提供静态窗口期。')
p('阶段三（返排阶段）：以恒定低流速（模拟现场返排，~0.1 mL/min）从产出端反向注入模拟地层水（50 g/L NaCl + 2 g/L CaCl₂），累计返排体积为裂缝体积的5~10倍。在产出端（即原注入端）逐裂缝体积（CPV）收集返排液样品，采用荧光光谱仪（或UV-Vis在520 nm特定吸收波长建立标准曲线）测定每CPV样品中的荧光粉浓度（μg/mL），绘制归一化荧光粉浓度-累计返排CPV的返排曲线。通过积分计算荧光粉累计返排量（μg），与注入总量对比得出净残留率（%）。若返排曲线在1~2 CPV内迅速降至本底水平（净残留率>90%），则验证了荧光粉的壁面锚定效果；若曲线呈长拖尾或多峰→指示存在弱吸附组分在返排中逐步释放。')
p('阶段四（取心与成像阶段）：拆卸岩心夹持器，将砂岩夹心岩样取出。在紫外暗室成像系统中，将所有成像参数（已在4.1中锁定）对裂缝壁面进行拍照。用ImageJ或自行编写的Python脚本（OpenCV），在裂缝区域的数字化图像中沿裂缝走向等间距圈定5个等面积ROI（200×200像素），提取每个ROI的平均灰度值（0~255）和标准差。')
p('[实验数据待补充：注入压力曲线、返排浓度曲线、紫外照片]')

h2('4.6  荧光信号与裂缝几何的半定量关系')
p('本研究的分析目标为两层次：（核心层次）验证压裂液波及范围与荧光分布的定性空间对应关系；（探索层次）在受控实验条件下建立荧光灰度值与裂缝宽度的半定量单调映射关系。')
p('分析方法：（1）定性验证——目视或阈值分割法判定四种裂缝宽度梯度下裂缝壁面荧光信号"可检出"（信噪比SNR≥3:1）的比例，若>90%的岩样可检出荧光，则定性验证成功。（2）半定量分析——以四种裂缝宽度（0.1、0.5、1.0、2.0 mm）为横坐标、5个ROI的平均灰度值为纵坐标，绘制散点图（n=12，每种宽度3个平行样），计算Spearman秩相关系数ρ（不预设线性关系）。若ρ≥0.85且p<0.01，则灰度值与裂缝宽度之间存在统计显著的单调关系，支持"在受控条件下灰度可作为裂缝宽度的半定量判据"。若ρ<0.6，则灰度值对裂缝宽度不具备有用的半定量区分能力，荧光信号仅保留"有/无"定性判读功能。该"半定量→定性"递进式分析框架确保无论实验结果落在哪个层次，结论都有实证支撑。')
p('须明确指出，该半定量关系受控于同一砂岩类型、同一成像条件和同一荧光粉批次的严格约束，不具备跨条件的绝对定量测量能力。这一分析边界的清晰界定避免了对方法能力范围的不当外推。')
p('[实验数据待补充]')

h2('4.7  本章小结')
p('[待实验完成后撰写]')
doc.add_page_break()

# ======== 第五章 ========
h1('第五章  现场施工工艺方案与经济性评估')
p('实验室验证的可行性向现场应用的转化，需要回答三个工程层面的问题：施工工艺流程如何设计以实现荧光粉的均匀添加和质量控制、单井材料用量和经济成本是否在可接受范围内、以及荧光粉对储层和环境的潜在影响是否可控。本章基于前述实验结果和工程经验，提出一套"母液预配+在线稀释"的现场施工工艺方案，完成经济性对比分析，并评估技术方案的环保合规性与工程适用条件。需要指出，本章的评估结论需要在后续先导性现场试验中进一步验证和修正。')
h2('5.1  "母液预配+在线稀释"工艺方案')
p('基于实验室实验结果，设计以下现场施工工艺方案：')
p('（1）母液预配：在压裂施工现场的配液区设置专用母液配制罐（容积5~10 m³，配备高速搅拌器和防尘通风装置）。将最优配方改性荧光粉（来自第二章）和分散剂（柠檬酸+Triton X-100）按设计配比加入到去离子水中，高速搅拌（500~1000 rpm）30 min，超声辅助脱泡（工业超声棒，40 kHz，连续处理），制备浓度为40 g/L的荧光母液。母液配制完成后静置检验——静置2 h无明显分层（目视判定，分层界面高度≤总液面高度的5%）方可放行使用。母液有效期为配制后7天，逾期需重新检测。')
p('（2）在线稀释：在压裂泵注系统的混砂车上游设置静态混合器（Static Mixer，SMX型，DN50，长度≥1 m），将荧光母液通过计量泵（隔膜计量泵，精度±1%）按照设计体积比（0.5% v/v）注入HPG基液主路，经过静态混合器的充分混合后进入混砂车，再与支撑剂（陶粒20/40目）汇合，最后通过高压泵组注入井筒。')
p('（3）质量控制：在静态混合器下游设置在线取样口，每隔30 min取样检测终液荧光粉浓度（便携式荧光分光光度计，520 nm发射峰强度 vs 实验室制备的标准曲线）和表观粘度（便携式粘度计，现场快速检测），确保浓度偏差≤±15%标准值、粘度偏差≤±10%空白HPG基液值。')
p('（4）关键工艺参数汇总见表5-1。')
p('表5-1  现场施工关键工艺参数', indent=False)
p('母液配制浓度：40 g/L; 母液储存有效期：7天（常温，避光）; 母液在线稀释比：0.5% v/v; 母液计量泵精度：±1%; 在线取样检测频率：1次/30 min; 浓度容许偏差：±15%; 粘度容许偏差：±10%', indent=False)
p('[注：表5-1将在论文中格式化为正式表格]')

h2('5.2  单井材料用量估算')
p('以四川盆地某典型页岩气水平井的压裂规模为算例，进行材料用量估算与成本分析。')
p('算例参数：水平段长度1500 m，分8段压裂，每段注入压裂液总量2000 m³，单井累计注入量16000 m³。荧光母液添加比例为0.5% v/v（即每立方米压裂液添加5 L母液）。单井母液总需求量=16000×5=80000 L=80 m³。按母液浓度40 g/L计算，单井荧光粉总用量=80 m³×40 kg/m³=3200 kg。每段用量=400 kg。')
p('改性试剂用量估算（基于最优配比）：KH550（假设最优用量2.0 wt%）=3200×0.02=64 kg; PEG4000（假设最优浓度3.0 wt%）=3200×0.03=96 kg; 柠檬酸（0.1 wt%基础）=3.2 kg; Triton X-100（0.05 wt%基础）=1.6 kg。')
p('最终用量以实验室确定的最优配方为准，以上仅为假设值用于算例。')

h2('5.3  经济性评估')
p('经济性分析的比较基准为目前压裂监测中常用的化学示踪剂方案。化学示踪剂（如含氟苯甲酸类、全氟烃类）的单井服务费用通常在3万~8万美元（包含示踪剂材料费、注入设备费和实验室分析费）[53]，放射性示踪剂稍高。荧光示踪方案的单井主要成本构成：')
p('（1）荧光粉材料费：SrAl₂O₄:Eu²⁺,Dy³⁺工业级产品（纯度≥98%，1000目）目前市场报价约200~500元/kg（根据纯度和购买量浮动）。按3200 kg/井、300元/kg均价估算，单井荧光粉材料费约96万元（~13.5万美元）。需指出，该估算假设每口井每一段均需加入荧光粉——这是最大用量方案。若仅需在关键段（如怀疑压窜或需验证间接反演结果的井段）添加，用量可减少至1/3~1/2。')
p('（2）改性试剂费：KH550（工业级约50~80元/kg）+ PEG4000（工业级约20~30元/kg）+ 柠檬酸（工业级约10元/kg）+ Triton X-100（工业级约30元/kg），合计单井约0.8万~1.5万元。')
p('（3）附加工艺费：母液配制罐租赁、计量泵、静态混合器等设备（平均到单井，按5~10口井分摊设备投入）预估单井1万~3万元。')
p('（4）总成本：单井荧光示踪方案总经济成本预估约100万元（约14万美元），与化学示踪剂方案（3~8万美元）相比偏高。但在以下场景中经济性有利：（i）仅需在1~2个关键段使用（用量降至1/4~1/2，成本约25万~50万元）；（ii）该方法提供的是取心实物证据——若能够通过直接验证避免一次错误的水力压裂设计决策（如井距优化、射孔位置选择），节省的成本远超示踪方案本身的费用；（iii）无机荧光粉不涉及放射性，省去了放射性示踪剂的特殊运输、许可和废物处理费用。')
p('[注：上述报价为公开市场参考价，论文最终版本的准确价格需根据实际采购询价确定。]')

h2('5.4  环保与安全性评估')
p('荧光示踪方案的环保评估聚焦稀土元素的环境迁移风险和返排液处理可行性。')
p('（1）稀土元素毒性：SrAl₂O₄基质中的碱土金属Sr（锶）和稀土Eu（铕）、Dy（镝）均属于低毒性元素。Sr²⁺的生物毒性类似Ca²⁺——其在体内的主要靶器官为骨骼，摄入高剂量才会产生慢性毒性。Eu和Dy属于轻稀土元素，急性毒性低（LD50>2000 mg/kg，大鼠经口，参照同类稀土氧化物数据），但长期环境暴露的慢性影响仍需关注。荧光粉的单井用量3200 kg，假设其中5%在返排液中以离子或微粒形式排出（约160 kg稀土元素释放至返排液中），需进行浸出实验（参照GB 5085.3-2007《危险废物鉴别标准 浸出毒性鉴别》）确定稀土浸出浓度是否超过标准限值。')
p('（2）返排液处理：荧光粉为微米级固体颗粒（~13 μm），可通过常规的絮凝-沉降-过滤工艺有效去除（尺寸远大于溶解性离子的纳滤/反渗透范围）。返排液中的稀土离子（Sr²⁺、Eu³⁺、Dy³⁺）可通过化学沉淀法（加碱调pH至9~10使其生成氢氧化物沉淀）或离子交换法去除。建议在现场返排液处理流程中增加沉淀池（投加石灰或NaOH）+ 板框压滤的组合工艺段。处理后水中残余稀土浓度应参照地方环保标准执行。')
p('（3）岩屑环保风险：含荧光粉附着层的钻井/取心岩屑属于一般工业固体废物。荧光粉的SrAl₂O₄基质化学性质稳定、不溶于水（经表面包覆后），环境浸出风险低。建议按常规工业固体废物处置。')

h2('5.5  技术定位与适用条件')
p('基于上述实验室验证和工程评估，本研究荧光示踪压裂液的技术定位和适用条件总结如下：')
p('技术定位：辅助校准工具。在压裂监测体系中，微地震和DAS/DTS测量间接信号（声发射、应变、温度），提供全井段的裂缝扩展趋势信息，但其反演结果缺乏实物校准。本荧光示踪方案通过压后取心提供"压裂液实际波及范围"的实物证据，可与间接监测结果形成互补。具体工作模式为：在微地震/DAS指示的裂缝核心区域附近的1~2口关键井中，在压裂的1~2个关键段注入荧光压裂液（而非全井段注入），压后取心进行紫外验证——若岩心荧光分布与间接方法预测的裂缝范围一致，则间接反演结果获得实物验证；若不一致，则为反演参数的修正提供约束。')
p('适用条件：（1）储层类型——砂岩储层（硅羟基Si-OH为氨基化荧光粉提供静电吸引和氢键锚定位点；碳酸盐岩或页岩储层壁面化学不同，锚定效果可能不同，需单独验证）。（2）压裂液体系——HPG或瓜尔胶基压裂液（其他稠化剂体系需验证化学兼容性）。（3）检测方式——需取心+紫外成像（若无法取心，该方法的实物证据优势无法发挥；井下紫外探头直接观测可作为远期替代方案，但其技术成熟度远低于取心）。（4）不适用于页岩基质微裂缝（裂缝宽度<0.1 mm可能低于可检出阈值，需实验确定）。')

h2('5.6  本章小结')
p('[待实验完成后撰写]')
doc.add_page_break()

# ======== 第六章 ========
h1('第六章  结论与展望')
h2('6.1  主要结论')
p('[待全部实验完成和数据分析后撰写]')
h2('6.2  创新点')
p('（1）面向氧化破胶环境的无机荧光示踪方案：选择无机铝酸盐晶体SrAl₂O₄:Eu²⁺,Dy³⁺为示踪介质——其发光中心受刚性晶格保护，对过硫酸铵氧化破胶剂具有本征化学惰性——构建了与常规HPG压裂液体系完全兼容的荧光示踪方案，区别于已有研究中使用的有机荧光树脂和半导体量子点。')
p('（2）"化学螯合—界面润湿—空间位阻"多层次协同分散体系：从工程分散需求出发，系统整合了柠檬酸的螯合屏蔽（抑制Al³⁺干扰交联）、非离子表面活性剂的界面润湿（促进PEG链伸展）与PEG长链的空间位阻（渗透排斥+构象熵效应）三个机制的协同效应，解决了高密度微米粉在低密度非牛顿压裂液中的悬浮稳定性问题。')
p('（3）利用压裂工程时序的功能切换设计：设计了双层改性结构并提出了破胶诱导PEG脱附-KH550暴露-砂岩硅羟基锚定的时序功能切换假说，尝试将"注入→关井→返排"的工程时间序列转化为颗粒表面化学的功能转化序列。')
p('（4）动态驱替全流程验证的实验方法：构建了模拟"注入—关井破胶—返排—取心—成像"五步串联流程的动态实验系统，以驱替压力曲线、返排浓度曲线和岩心成像三组数据构成完整证据链，区别于已有研究中仅依赖静态浸泡或简单注入后成像的验证方式。')
h2('6.3  不足与展望')
p('[待实验完成后如实撰写]')
doc.add_page_break()

# ======== 参考文献 ========
h1('参考文献')
refs = [
    '[1] Montgomery, C.T. & Smith, M.B. (2010). Hydraulic Fracturing: History of an Enduring Technology. JPT, 62(12): 26-40.',
    '[2] Mayerhofer, M.J., et al. (2010). What Is Stimulated Reservoir Volume? SPE Production & Operations, 25(1): 89-98.',
    '[3] Fisher, M.K., et al. (2004). Optimizing Horizontal Completion Techniques in the Barnett Shale Using Microseismic Fracture Mapping. SPE-90051-MS.',
    '[4] Cipolla, C. & Wallace, J. (2014). Stimulated Reservoir Volume: A Misapplied Concept? SPE-168596-MS.',
    '[5] Cipolla, C.L. & Wright, C.A. (2002). Diagnostic Techniques To Understand Hydraulic Fracturing. SPE Production & Facilities, 17(1): 23-35.',
    '[6] Warpinski, N.R., et al. (2009). Stimulating Unconventional Reservoirs. JCPT, 48(10): 39-51.',
    '[7] Warpinski, N.R. (1996). Hydraulic Fracture Diagnostics. JPT, 48(10): 907-910.',
    '[8] de Melo, R.C.B. & Carballo Cabrera, A.C. (2025). Advances in Hydraulic Fracturing Mapping. SPE-228867-MS.',
    '[9] Cipolla, C., et al. (2011). A Practical Guide to Interpreting Microseismic Measurements. SPE-144067-MS.',
    '[10] Maxwell, S.C. & Cipolla, C. (2011). What Does Microseismicity Tell Us About Hydraulic Fracturing? SPE-146932-MS.',
    '[11] Warpinski, N.R., et al. (2003). Improved Microseismic Fracture Mapping Using Perforation Timing Measurements. SPE-84488-MS.',
    '[12] Jin, G. & Roy, B. (2017). Hydraulic-fracture geometry characterization using low-frequency DAS signal. TLE, 36(12): 975-980.',
    '[13] Rasool, M.H., et al. (2025). Industrial Adoption of DAS in Petroleum Engineering and Geosciences. JPCE, 1(1).',
    '[14] Karrenbach, M., et al. (2017). Hydraulic-fracturing-induced strain and microseismic using distributed fiber-optic sensing. TLE, 36(10): 837-844.',
    '[15] Salman, A., et al. (2014). Analysis of Chemical Tracer Flowback in Unconventional Reservoirs. SPE-171656-MS.',
    '[16] Viig, S.O., et al. (2013). Application of a New Class of Chemical Tracers to Measure Oil Saturation. SPE-164059-MS.',
    '[17] Huseby, O., et al. (2015). Assessing EOR Potential from Partitioning Tracer Data. SPE-172575-MS.',
    '[18] Takeuchi, T., et al. (2025). Outcrop-Scale Hydraulic Fracturing Experiments with a Coagulable Resin. Geosciences, 15(3): 103.',
    '[19] Chen, Y., et al. (2014). Visualization of Fractures Induced by Hydraulic Fracturing. ISRM-ARMS8-2014-192.',
    '[20] Guryanov, A., et al. (2019). Application of Fluorescent Markers to Determine Formation Fluid Inflow After MFrac. SPE-196776-MS.',
    '[21] Molenaar, M.M., et al. (2012). First Downhole Application of DAS for Hydraulic-Fracturing Monitoring. SPE Drilling & Completion, 27(1): 32-38.',
    '[22] Sierra, J., et al. (2008). DTS Monitoring of Hydraulic Fracturing: Experiences and Lessons Learned. SPE-116182-MS.',
    '[23] Ugueto, G.A., et al. (2016). Perforation Cluster Efficiency from Fiber Optics Diagnostics. SPE-179124-MS.',
    '[24] Karrenbach, M., et al. (2019). Fiber-Optic Distributed Acoustic Sensing of Microseismicity. Geophysics, 84(6).',
    '[25] Ugueto, G.A., et al. (2019). Can You Feel the Strain? DAS Strain Fronts for Fracture Geometry. SPE-195943-MS.',
    '[26] Maxwell, S.C., et al. (2002). Microseismic Imaging of Hydraulic Fracture Complexity in the Barnett Shale. SPE-77440-MS.',
    '[27] Cipolla, C., et al. (2011). Integrating Microseismic Mapping and Complex Fracture Modeling. SPE-140185-MS.',
    '[28] Maxwell, S.C. (2014). Microseismic Imaging of Hydraulic Fracturing. SEG Distinguished Instructor Series No. 17.',
    '[29] Tian, W., et al. (2019). Understanding Frac Fluid Distribution of an Individual Frac Stage. SPE-194362-MS.',
    '[30] Arshad, W., et al. (2024). First Application of Novel Tracer Technology for Monitoring Fracture Stage Contribution. IPTC-23916-MS.',
    '[31] Kosynkin, D. & Alaskar, M. (2016). Oil Industry First Interwell Trial of Reservoir Nanoagent Tracers. SPE-181551-MS.',
    '[32] Hu, Z., et al. (2019). Carbon Quantum Dots with Tracer-like Breakthrough Ability. Science of the Total Environment, 669: 579-589.',
    '[33] Wright, C.A., et al. (1998). Downhole Tiltmeter Fracture Mapping. SPE-46194-MS.',
    '[34] Warpinski, N.R., et al. (2006). Improving Hydraulic Frac Diagnostics by Joint Inversion. SPE-102690-MS.',
    '[35] LaBrecque, D., et al. (2016). Remote Imaging of Proppants Using Electromagnetic Methods. SPE-179170-MS.',
    '[36] Ahmadian, M., et al. (2023). Real-Time Monitoring of Fracture Dynamics with Contrast Agent-Assisted EM. SPE-212376-MS.',
    '[37] Flury, M. & Wai, N.N. (2003). Dyes as Tracers for Vadose Zone Hydrology. Reviews of Geophysics, 41(1): 1002.',
    '[38] Kang, W., et al. (2015). Preparation and Performance of Fluorescent Polyacrylamide Microspheres. Petroleum Science, 12: 483-491.',
    '[39] Spitzmuller, L., et al. (2024). Temperature Stability and Enhanced Transport of Silica Nanoparticle Tracers. Scientific Reports, 14: 19222.',
    '[40] Matsuzawa, T., et al. (1996). A New Long Phosphorescent Phosphor SrAl₂O₄:Eu²⁺,Dy³⁺. J. Electrochemical Society, 143(8): 2670-2673.',
    '[41] Clabau, F., et al. (2005). Mechanism of Phosphorescence in Eu²⁺-Doped SrAl₂O₄. Chemistry of Materials, 17(15): 3904-3912.',
    '[42] Dorenbos, P. (2005). Mechanism of Persistent Luminescence in Eu²⁺ and Dy³⁺ Codoped Aluminates. J. Electrochemical Society, 152(7): H107-H110.',
    '[43] Li, Y., et al. (2019). Persistent luminescence: History, mechanism, and perspective. J. Luminescence, 205: 581-620.',
    '[44] Rojas-Hernandez, R.E., et al. (2018). SrAl₂O₄:Eu,Dy as the most studied material. Renewable and Sustainable Energy Reviews, 81(2): 2759-2770.',
    '[45] Guo, C., et al. (2007). Stability of SrAl₂O₄:Eu²⁺,Dy³⁺ in water. Materials Chemistry and Physics, 106(2-3): 268-272.',
    '[46] Karacaoglu, E., et al. (2020). ALD coatings on SrAl₂O₄-based phosphor powders. J. American Ceramic Society, 103(6): 3706-3715.',
    '[47] Qi, T., et al. (2017). Improved water resistance of SrAl₂O₄:Eu²⁺,Dy³⁺ phosphor. Solid State Sciences, 65: 88-94.',
    '[48] Deng, S., et al. (2013). Surface modification of MAl₂O₄:Eu²⁺,Dy³⁺ phosphors. Applied Surface Science, 282: 315-320.',
    '[49] Urakawa, K., et al. (2020). Emission increase in persistent inorganic phosphor. J. Luminescence, 217: 116772.',
    '[50] Lyu, L., et al. (2020). Silica-Polymer Hybrid Shell on SrAl₂O₄:Eu²⁺,Dy³⁺. Materials, 13(2): 426.',
    '[51] Gomaa, A.M., et al. (2015). Proppant Transport? Viscosity Is Not All. SPE-173323-MS.',
    '[52] Goel, N. & Shah, S. (2001). A Rheological Criterion for Fracturing Fluids to Transport Proppant. SPE-71663-MS.',
    '[53] Al-Muntasheri, G.A. (2014). A Critical Review of Hydraulic-Fracturing Fluids. SPE Production & Operations, 29(4): 243-260.',
    # Additional references
    '[54] Economides, M.J. & Nolte, K.G. (2000). Reservoir Stimulation (3rd ed.). Wiley.',
    '[55] Gandossi, L. & Von Estorff, U. (2015). Overview of Hydraulic Fracturing Technologies for Shale Gas. JRC EUR 26347 EN.',
    '[56] King, G.E. (2012). Hydraulic Fracturing 101. SPE-152596-MS.',
    '[57] Barati, R. & Liang, J.T. (2014). A review of fracturing fluid systems. J. Applied Polymer Science, 131(16): 40735.',
    '[58] Civan, F. (2015). Reservoir Formation Damage (3rd ed.). Gulf Professional Publishing.',
    '[59] Harris, P.C. (1993). Chemistry and Rheology of Borate-Crosslinked Fluids. JPT, 45(3): 264-269.',
    '[60] Kesavan, S. & Prud\'homme, R.K. (1992). Rheology of guar and HPG crosslinked by borate. Macromolecules, 25(7): 2026-2032.',
    '[61] Gardner, D.C. & Eikerts, J.V. (1982). Effects of Shear and Proppant on Viscosity of Crosslinked Fluids. SPE-11066-MS.',
    '[62] Hurnaus, T. & Plank, J. (2015). Crosslinking of Guar and HPG Using ZrO₂ Nanoparticles. SPE-173778-MS.',
    '[63] Putzig, D.E. & St. Clair, J.D. (2007). A New Delay Additive for Hydraulic Fracturing Fluids. SPE-105066-MS.',
    '[64] Parker, M.A., et al. (1999). Fracturing-Fluid Breakers and Proppant-Pack Conductivity. SPE-50735-MS.',
    '[65] Brannon, H.D. & Tjon-Joe-Pin, R.M. (1994). Biotechnological Breakthrough for High-Temperature Fracturing. SPE-28513-MS.',
    '[66] Parker, M.A. & Laramay, S.B. (1992). Properties of Delayed-Release Breakers. SPE-24300-MS.',
    '[67] Rae, P. & Lullo, G.D. (1996). Fracturing Fluids and Breaker Systems Review. SPE-37359-MS.',
    '[68] Acharya, A. (1988). Viscoelasticity of Crosslinked Fluids and Proppant Transport. SPE Production Engineering, 3(4): 695-704.',
    '[69] Samuel, M.M., et al. (1999). Polymer-Free Fluid for Fracturing Applications. SPE Drilling & Completion, 14(4): 240-246.',
    '[70] Palisch, T.T., et al. (2010). Slickwater Fracturing: Food for Thought. SPE Production & Operations, 25(3): 327-344.',
    '[71] Harris, P.C. (1988). Fracturing-Fluid Additives. JPT, 40(10): 1317-1319.',
    '[72] Van den Eeckhout, K., et al. (2010). Persistent Luminescence in Eu²⁺-Doped Compounds: A Review. Materials, 3(4): 2536-2566.',
    '[73] Van den Eeckhout, K., et al. (2013). Persistent Luminescence in Non-Eu²⁺-Doped Compounds. Materials, 6(7): 2789-2818.',
    '[74] Vitola, V., et al. (2019). Recent progress in persistent luminescence in SrAl₂O₄:Eu,Dy. Materials Science and Technology, 35(18): 2153-2160.',
    '[75] Kaur, J., et al. (2014). Persistent luminescence of SrAl₂O₄ phosphors: a review. Research on Chemical Intermediates, 40: 317-343.',
    '[76] Do, H.-S., et al. (2010). Improved moisture resistance of SrS:Eu²⁺ with SiO₂ coating. J. Luminescence, 130(8): 1400-1403.',
    '[77] Bem, D.B., et al. (2010). SrAl₂O₄ phosphor in LDPE and PMMA polymers. J. Applied Polymer Science, 117(5): 2635-2640.',
    '[78] Anesh, M.P., et al. (2014). Eu²⁺-Doped Strontium Aluminate and Polymer Composite. Advances in Polymer Technology, 33(S1): 21436.',
    '[79] Sharma, S., et al. (2016). SrAl₂O₄ nanocrystalline phosphor for fingerprint detection. Materials Research Express, 3(1): 015004.',
    '[80] Holsa, J., et al. (2001). Persistent luminescence of MAl₂O₄:Eu²⁺. J. Alloys and Compounds, 323-324: 326-330.',
    '[81] Warpinski, N.R., et al. (2012). Measurements of Hydraulic-Fracture-Induced Seismicity in Gas Shales. SPE Production & Operations, 27(3): 240-252.',
    '[82] Duncan, P.M. & Eisner, L. (2010). Reservoir characterization using surface microseismic monitoring. Geophysics, 75(5): 75A139-75A146.',
    '[83] Eisner, L., et al. (2010). Beyond the dots in the box. TLE, 29(3): 326-333.',
    '[84] Bazin, B., et al. (2010). Fracturing in Tight Gas Reservoirs: Formation-Damage Mechanisms. SPE Journal, 15(4): 969-976.',
    '[85] 郭建春, 何春明. (2012). 压裂液破胶过程伤害微观机理. 石油学报, 33(6): 1018-1022.',
    '[86] 翁定为, 等. (2024). 水力压裂裂缝监测技术综述. 世界石油工业, 31(6): 66-76.',
    '[87] 邸德家. (2025). 油气井压裂示踪监测技术现状及发展建议. 钻采工艺, 2025(2): 74-81.',
    '[88] 吕兴栋, 等. (2005). 碱土铝酸盐长余辉发光材料的有机包覆. 应用化学, 22(6): 638-642.',
    '[89] Mikutis, G., et al. (2018). Silica-Encapsulated DNA-Based Tracers for Aquifer Characterization. ES&T, 52(21): 12142-12152.',
    '[90] Molnar, I.L., et al. (2015). Predicting Colloid Transport through Saturated Porous Media. WRR, 51(9): 6804-6845.',
    '[91] Rodriguez, E., et al. (2009). Enhanced Migration of Surface-Treated Nanoparticles in Sedimentary Rocks. SPE-124418-MS.',
    '[92] Elimelech, M. & O\'Melia, C.R. (1990). Kinetics of Deposition of Colloidal Particles in Porous Media. ES&T, 24(10): 1528-1536.',
    '[93] Bhattacharjee, S., et al. (1998). DLVO Interaction between Rough Surfaces. Langmuir, 14(12): 3365-3375.',
    '[94] Scheurer, C., et al. (2022). Sorption of Nanomaterials to Sandstone Rock. Nanomaterials, 12(2): 200.',
    '[95] Scott, M.P., et al. (2010). Evaluating Hydraulic Fracture Geometry from Sonic Anisotropy and Radioactive Tracer Logs. SPE-133059-MS.',
    '[96] SY/T 6376-2008. 压裂液通用技术条件. 国家发展和改革委员会, 2008.',
    '[97] SY/T 5107-2016. 水基压裂液性能评价方法. 国家能源局, 2016.',
    '[98] Wang, X., et al. (2025). Rheological properties of HPG fracturing fluids: Effects of cross-linking agents. Carbohydrate Polymers, 366: 123837.',
    '[99] Tang, J.S. (1995). Partitioning Tracers and In-Situ Fluid-Saturation Measurements. SPE Formation Evaluation, 10(1): 33-39.',
    '[100] Fisher, K., et al. (1995). Analysis and Economic Benefits of Radioactive Tracer Engineered Stimulation. SPE-30794-MS.',
]
for ref_text in refs:
    pref(ref_text)

# Save
output_path = r'c:\Users\郝\Desktop\claude\荧光压裂液\论文初稿.docx'
doc.save(output_path)
print(f'Saved: {output_path}\nReferences: {len(refs)}')