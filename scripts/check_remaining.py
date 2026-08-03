#!/usr/bin/env python3
from docx import Document
doc = Document('四氧化三铁环氧树脂拟合/ESP-T_终稿.docx')
with open('四氧化三铁环氧树脂拟合/remaining_check.txt','w',encoding='utf-8') as f:
    f.write('--- Remaining numbered points ---\n')
    for i,p in enumerate(doc.paragraphs):
        t = p.text
        for kw in ['First,', 'Second,', 'Third,', 'Fourth,', 'Fifth,', 'Sixth,',
                    'Firstly', 'Secondly', 'Thirdly']:
            if kw in t:
                idx = t.find(kw)
                f.write(f'  [{i}] {kw}: ...{t[idx:idx+80]}...\n')

    f.write('\n--- Pe discussion locations ---\n')
    for i,p in enumerate(doc.paragraphs):
        if 'Peclet number' in p.text:
            idx = p.text.find('Peclet')
            f.write(f'  [{i}]: {p.text[idx:idx+80]}...\n')

    f.write('\n--- 47% discussion locations ---\n')
    for i,p in enumerate(doc.paragraphs):
        if '47%' in p.text and i < 166:
            idx = p.text.find('47%')
            f.write(f'  [{i}]: ...{p.text[max(0,idx-20):idx+80]}...\n')

    f.write('\n--- Section transitions ---\n')
    for i in [9,11,13,15,105,125,148,158]:
        if i < len(doc.paragraphs):
            f.write(f'  [{i}] first 120 chars: {doc.paragraphs[i].text[:120]}\n')

print('Done')