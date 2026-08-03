# -*- coding: utf-8 -*-
"""
Task 1/13 (ESP-T paper rewrite): Create the manuscript skeleton DOCX.

Creates the structural skeleton at
    四氧化三铁环氧树脂拟合/ESP-T_投稿文件/ESP-T_v2_manuscript.docx
containing:
  - Title placeholder, Abstract, Keywords
  - Main sections 1-7 as Heading 1, subsections as Heading 2
  - Figure markers [Fig. N placeholder — see Figure Captions] at the
    locations specified in the task-1 brief / design doc section 3
  - References section with 35 numbered placeholder slots (Task 11 fills)
  - Figure Captions section with 10 placeholder slots
  - Tables section with placeholder slots
  - A blank paragraph after every heading for future content

Subsequent tasks (2-13) will write content into these slots.
"""
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

BASE = os.path.dirname(os.path.abspath(__file__))
PROJECT = os.path.dirname(BASE)
OUT = os.path.join(
    PROJECT, "四氧化三铁环氧树脂拟合", "ESP-T_投稿文件", "ESP-T_v2_manuscript.docx"
)

TITLE = "Title: [TBD — written in Task 12 after all content is final]"

# Section tree. Each entry: (kind, text)
#   kind "h1"  -> Heading 1 (main section)
#   kind "h2"  -> Heading 2 (subsection)
#   kind "fig" -> italic figure-marker paragraph
#   kind "note" -> italic note paragraph
SECTIONS = [
    ("h1", "Abstract"),
    ("h1", "Keywords"),
    ("h1", "1. Introduction"),
    ("h1", "2. Coupled Release–Transport Model"),
    ("h2", "2.1 Physical Basis"),
    ("fig", "[Fig. 1 placeholder — see Figure Captions] — BTC generation schematic"),
    ("h2", "2.2 Governing Equations"),
    ("h2", "2.3 Parameter Set and Derived Quantities"),
    ("h2", "2.4 Validation Strategy: Self-Calibration as Decisive Test"),
    ("h1", "3. Experimental Methods"),
    ("h2", "3.1 ESP-T Synthesis"),
    ("h2", "3.2 Material Prerequisites for Model Validity"),
    ("fig", "[Fig. 2 placeholder — see Figure Captions] — ESP-T characterization summary, 4-panel"),
    ("h2", "3.3 K-P Batch Release Kinetics"),
    ("h2", "3.4 Core Displacement: BTC Generation"),
    ("h2", "3.5 Parameter Estimation Strategy"),
    ("h1", "4. Results and Discussion"),
    ("h2", "4.1 Model Selection: Single-Process Models Are Insufficient"),
    ("fig", "[Fig. 3 placeholder — see Figure Captions] — K-P kinetics + model selection overlay"),
    ("fig", "[Fig. 4 placeholder — see Figure Captions] — 5-model overlay + ΔAICc bar chart"),
    ("h2", "4.2 Physical Self-Calibration: The Flow-Rate Test"),
    ("fig", "[Fig. 5 placeholder — see Figure Captions] — CENTERPIECE — BTC decomposition + Q self-calibration bar"),
    ("h2", "4.3 Comparison with Time-of-Arrival Methods"),
    ("fig", "[Fig. 6 placeholder — see Figure Captions] — TOA method comparison"),
    ("h2", "4.4 Independent Corroboration: K-P Kinetics and ADE Peclet Number"),
    ("fig", "[Fig. 7 placeholder — see Figure Captions] — K-P ↔ Pe independent corroboration dual panel"),
    ("h2", "4.5 Signal Decomposition and Robustness"),
    ("fig", "[Fig. 8 placeholder — see Figure Captions] — σ sensitivity analysis"),
    ("h1", "5. Extension to Two-Phase Production Allocation"),
    ("h2", "5.1 Physical Context"),
    ("h2", "5.2 The Dilution Problem and the Flux Solution"),
    ("fig", "[Fig. 9 placeholder — see Figure Captions] — Two-phase flow 3-panel"),
    ("h2", "5.3 Flux-to-Production-Rate Calibration"),
    ("h2", "5.4 Coupling to the BTC Framework"),
    ("h1", "6. Field Deployment Pathway"),
    ("fig", "[Fig. 10 placeholder — see Figure Captions] — Field deployment pathway schematic"),
    ("h1", "7. Conclusions"),
    ("h1", "References"),
    ("note", "[numbered list, 35 slots — populated in Task 11]"),
    ("refs", None),
    ("h1", "Figure Captions"),
    ("note", "[10 figure caption slots — populated as figures are placed]"),
    ("captions", None),
    ("h1", "Tables"),
    ("note", "[Table slots — populated as content is written]"),
    ("tables", None),
]

N_REFERENCES = 35
N_FIGURES = 10


def main():
    doc = Document()

    # Base style: Times New Roman 12 pt (SCI-manuscript convention)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "Times New Roman")

    # Title placeholder (level-0 "Title" style, centered)
    tp = doc.add_heading(TITLE, level=0)
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    for kind, text in SECTIONS:
        if kind == "h1":
            doc.add_heading(text, level=1)
            doc.add_paragraph("")  # blank paragraph for future content
        elif kind == "h2":
            doc.add_heading(text, level=2)
            doc.add_paragraph("")  # blank paragraph for future content
        elif kind == "fig":
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.italic = True
        elif kind == "note":
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.italic = True
        elif kind == "refs":
            # Reference slots (populated in Task 11)
            for i in range(1, N_REFERENCES + 1):
                doc.add_paragraph("[{}]".format(i))
        elif kind == "captions":
            # Figure caption slots (populated as figures are placed)
            for i in range(1, N_FIGURES + 1):
                doc.add_paragraph(
                    "[Fig. {}] Caption placeholder — populated as figure is placed.".format(i)
                )
        elif kind == "tables":
            # Table slots (populated as content is written)
            doc.add_paragraph(
                "[Table 1] Placeholder — material prerequisites summary (see Section 3.2)."
            )
            doc.add_paragraph(
                "[Table 2] Placeholder — model comparison summary (see Section 4.1)."
            )

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print("Saved:", OUT)

    # ---- Verification (Step 3) ----
    check = Document(OUT)
    headings1 = [p.text for p in check.paragraphs if p.style.name == "Heading 1"]
    headings2 = [p.text for p in check.paragraphs if p.style.name == "Heading 2"]
    markers = [p.text for p in check.paragraphs
               if p.text.startswith("[Fig.") and "see Figure Captions" in p.text]
    refs = [p.text for p in check.paragraphs if p.style.name == "Normal" and p.text.strip().startswith("[") and p.text.strip().rstrip("]").strip("[").isdigit()]

    main = ["1. Introduction", "2. Coupled Release–Transport Model",
            "3. Experimental Methods", "4. Results and Discussion",
            "5. Extension to Two-Phase Production Allocation",
            "6. Field Deployment Pathway", "7. Conclusions"]
    missing = [m for m in main if m not in headings1]

    print("Heading 1 count:", len(headings1))
    print("Heading 2 count:", len(headings2))
    print("Figure markers:", len(markers))
    print("Reference slots:", len(refs))
    print("Missing main sections:", missing or "none")
    assert not missing, "Missing main sections: %s" % missing
    assert len(headings2) == 18, "Expected 18 subsections, got %d" % len(headings2)
    assert len(markers) == 10, "Expected 10 figure markers, got %d" % len(markers)
    assert len(refs) == 35, "Expected 35 reference slots, got %d" % len(refs)
    print("VERIFICATION PASSED")


if __name__ == "__main__":
    main()
