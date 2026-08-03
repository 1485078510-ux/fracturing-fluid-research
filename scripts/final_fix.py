"""
Final comprehensive fix to ESP_EN_final.docx for SCI submission.
"""
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import copy, re

DST = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP_EN_final.docx'
doc = Document(DST)

IMG = set()
for i, p in enumerate(doc.paragraphs):
    for run in p.runs:
        if (run._element.findall('.//'+qn('w:drawing')) or run._element.findall('.//'+qn('a:blip'))):
            IMG.add(i); break

def sp(idx, text):
    if idx in IMG: return
    p = doc.paragraphs[idx]
    for r in p.runs: r.text = ''
    if p.runs: p.runs[0].text = text
    else: p.add_run(text)

# ===== FIX 1: Restore TITLE as Heading 1 before abstract =====
body = doc.element.body
first_elem = body[0]  # current first element (abstract)
# Create title paragraph
title_p = copy.deepcopy(doc.paragraphs[5]._element)  # clone intro heading
for r in title_p.findall(qn('w:r')): title_p.remove(r)
r = etree.SubElement(title_p, qn('w:r'))
t = etree.SubElement(r, qn('w:t'))
t.text = "Fabrication and Performance of an Oleophilic Fe3O4-Doped Epoxy Resin Tracer Proppant for Hydraulic Fracturing Production Monitoring"
t.set(qn('xml:space'), 'preserve')
pPr = title_p.find(qn('w:pPr'))
pStyle = pPr.find(qn('w:pStyle'))
pStyle.set(qn('w:val'), 'Heading1')
body.insert(0, title_p)

# ===== FIX 2: Fix all "degrees C" -> "C" and standardize temperature notation =====
# Replace in all paragraphs
temp_fixes = {
    "80 degrees C": "80 C",
    "10 degrees C/min": "10 C/min",
    "50 to 800 degrees C": "50 to 800 C",
    "357.27 degrees C": "357.27 C",
}
for para in doc.paragraphs:
    for run in para.runs:
        if run.text:
            for old, new in temp_fixes.items():
                run.text = run.text.replace(old, new)

# ===== FIX 3: Add "where:" before variable list in Section 3.7 =====
# Currently [84] has c_b without "where:"
sp(84, "where:")
sp(85, "c_b = background tracer concentration under steady flow, reflecting the release-dilution equilibrium of nano-Fe3O4@SA in the flowing oil phase (dimensionless, normalized by the peak concentration);")
sp(86, "A, a = concentration coefficients for the shut-in accumulation slug (A) and the sustained matrix-diffusion-controlled release from ESP-T (a);")
sp(87, "alpha = longitudinal dispersivity of the flow line, characterizing axial spreading of the tracer slug during transit (cm);")
sp(88, "Q = effective oil-phase volumetric flow rate obtained from tracer breakthrough curve inversion (cm^3/min);")
sp(89, "t0 = time at which sustained matrix-diffusion release overtakes the shut-in slug as the primary contributor to the detected tracer signal (min);")
sp(90, "sigma = characteristic timescale of the transition from slug-dominated to release-dominated tracer transport (min). In the above expressions, x = 100 cm and d = 5 cm are the fixed length and inner diameter of the flow line.")

# ===== FIX 4: Fix conclusion formatting =====
sp(100,
    "An oleophilic tracer proppant (ESP-T) was developed by encapsulating stearic "
    "acid-modified nano-Fe3O4@SA within an epoxy resin matrix via emulsion polymerization. "
    "Nano-Fe3O4@SA disperses uniformly as nanoclusters, and ESP-T exhibits sphericity "
    "and roundness exceeding 0.9, meeting industrial standards. ESP-T demonstrates "
    "excellent thermal stability (decomposition at 357.27 C), low bulk density "
    "(0.646 g/cm3), and a water-resistant, oil-permeable characteristic that "
    "facilitates oil flow while inhibiting water breakthrough.")

sp(101,
    "Tracer release follows the Korsmeyer-Peppas model (R2 > 0.90, n = 0.45-0.85), "
    "governed by synergistic Fickian diffusion and Case-II relaxation. A piecewise "
    "ADE model with smooth tanh transition was developed to interpret the breakthrough "
    "curve, decomposing the signal into a Gaussian pulse (shut-in accumulation) and an "
    "erfc tail (sustained matrix-diffusion-controlled release). The model achieves "
    "R2 = 0.9939, with a fitted flow rate (0.46 mL/min) closely matching the pump-set "
    "rate (0.5 mL/min, error 8%) and a mean residence time (37.4 min) in agreement with "
    "the convective travel time (38.6 min, ratio 0.967). Deconvolution shows that 47% "
    "of the integrated signal originates from the erfc tailing component. Under two-phase "
    "flow, tracer flux effectively quantifies per-interval oil production rates across "
    "varying oil-water ratios.")

# ===== FIX 5: Abstract — fix "Abstract:" formatting =====
sp(1,  # Abstract para 1 (shifted by title insertion)
    "Abstract: Per-interval production monitoring is essential for evaluating "
    "stimulation effectiveness and optimizing development strategies in unconventional "
    "reservoirs. We report an oleophilic tracer proppant (ESP-T) fabricated by "
    "encapsulating stearic acid-modified nano-Fe3O4 within an epoxy resin matrix "
    "via emulsion polymerization, integrating fracture propping and production "
    "monitoring into a single material. The microstructure, mechanical properties, "
    "thermal stability, wettability, and temperature-dependent tracer release "
    "kinetics of ESP-T are systematically characterized, and its monitoring "
    "performance is validated through core displacement experiments.")

sp(2,
    "Nano-Fe3O4 disperses uniformly as nanoclusters in the epoxy matrix. ESP-T "
    "exhibits sphericity and roundness exceeding 0.9, bulk density of 0.646 g/cm3, "
    "acid solubility of 3.3% (<=5% industrial standard), and thermal decomposition "
    "temperature of 357.27 C, well above typical downhole conditions. Incorporation "
    "of nano-Fe3O4 raises the water contact angle from 72.3 to 104.6 degrees, "
    "yielding a water-resistant, oil-permeable characteristic (oil filtration time "
    "5 min 11 s, 66.1% shorter than pure epoxy microspheres). Tracer release follows "
    "the Korsmeyer-Peppas model (R2 > 0.90; n = 0.45-0.85), governed by synergistic "
    "Fickian diffusion and Case-II relaxation.")

sp(3,
    "A piecewise advection-dispersion model with smooth tanh transition is developed "
    "to interpret the single-phase breakthrough curve, decomposing the signal into a "
    "Gaussian pulse (shut-in accumulation) and an erfc tail (sustained matrix-diffusion-"
    "controlled release). The model achieves R2 = 0.9939, with the fitted flow rate "
    "(0.46 mL/min) closely matching the pump-set rate (0.5 mL/min, error 8%). Under "
    "two-phase flow, tracer flux effectively quantifies per-interval oil production "
    "rates. ESP-T thus provides a robust dual-function platform for long-term "
    "production monitoring in unconventional reservoirs.")

# ===== SAVE =====
doc.save(DST)
print("Final fixes applied: title restored, temperature units fixed, 'where:' added, conclusions formatted.")