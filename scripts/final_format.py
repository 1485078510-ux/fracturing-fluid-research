# -*- coding: utf-8 -*-
"""Final format pass: chem subscripts, ref superscripts, clean empty paras."""
from docx import Document
from docx.oxml.ns import qn
import re

SRC = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP_final_v3.docx'
DST = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP_final_v3.docx'
doc = Document(SRC)

IMG = set()
for i, p in enumerate(doc.paragraphs):
    for run in p.runs:
        if (run._element.findall('.//'+qn('w:drawing')) or
            run._element.findall('.//'+qn('a:blip'))):
            IMG.add(i); break

# Patterns
REF_PAT = re.compile(r'\[\d+(?:[-,]\d+)*\]')
CHEM_PAT = re.compile(
    r'\b(Fe\d+O\d+(?:@\w+)?|FeCl\d+|MnCl\d+|SiO\d*|CO\d*|TiO\d*|Al\d*O\d*|C\d+H\d+COOH)\b'
)

count = 0
for idx, para in enumerate(doc.paragraphs):
    if idx in IMG or not para.text.strip():
        continue
    if not (CHEM_PAT.search(para.text) or REF_PAT.search(para.text)):
        continue

    # Get original formatting
    orig_font = None; orig_size = None
    for run in para.runs:
        if run.text.strip():
            orig_font = run.font.name; orig_size = run.font.size; break

    # Find all matches
    matches = []
    for m in CHEM_PAT.finditer(para.text):
        matches.append((m.start(), m.end(), 'chem'))
    for m in REF_PAT.finditer(para.text):
        matches.append((m.start(), m.end(), 'ref'))
    if not matches: continue
    matches.sort(key=lambda x: (x[0], -x[1]))

    # Deduplicate
    filtered, last = [], 0
    for s, e, t in matches:
        if s >= last: filtered.append((s, e, t)); last = e

    # Build segments
    segs, pos = [], 0
    for s, e, t in filtered:
        if s > pos: segs.append((para.text[pos:s], 'normal'))
        segs.append((para.text[s:e], t)); pos = e
    if pos < len(para.text):
        segs.append((para.text[pos:], 'normal'))

    # Expand chem segments
    expanded = []
    for txt, typ in segs:
        if typ == 'chem':
            parts = re.split(r'(\d+)', txt)
            for part in parts:
                if not part: continue
                expanded.append((part, 'sub' if part.isdigit() else 'normal'))
        else:
            expanded.append((txt, typ))

    # Rebuild runs
    for run in list(para.runs):
        run._element.getparent().remove(run._element)
    for txt, typ in expanded:
        if not txt: continue
        r = para.add_run(txt)
        if orig_font: r.font.name = orig_font
        if orig_size: r.font.size = orig_size
        if typ == 'sub': r.font.subscript = True; count += 1
        elif typ == 'ref': r.font.superscript = True; count += 1

doc.save(DST)
print(f"Formatted {count} items. Saved: {DST}")
print(f"Images: {len(doc.inline_shapes)}")