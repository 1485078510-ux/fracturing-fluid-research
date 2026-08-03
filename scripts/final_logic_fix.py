#!/usr/bin/env python3
"""Fix all remaining logical issues in one pass."""
from docx import Document
import time
doc = Document('四氧化三铁环氧树脂拟合/ESP-T_final_1781617385.docx')

# ── FIX 1: [173] Conclusions — remove duplicated old text ──
p173 = doc.paragraphs[173]
t173 = p173.text
# Keep only the new text, remove the concatenated old fragment
# "These numbers are not curve-fitting artifacts; they are physical quantities recovered from the tracer signal alone."
# After that, remove everything starting from "This quantitatively establishes..."
cut = "This quantitatively establishes"
if cut in t173:
    idx = t173.index(cut)
    t173 = t173[:idx].strip()
    for r in p173.runs: r.text = ''
    p173.runs[0].text = t173
    print('[173] old text removed')

# ── FIX 2: Restore limitations paragraph ──
# Add limitations back into [160] (FO validation paragraph)
p160 = doc.paragraphs[160]
t160 = p160.text
lim_text = (
    "\n\n"
    "Several boundaries of the present study should be noted. The model "
    "assumes one fractured interval with a uniform proppant pack; real "
    "wells have multiple intervals whose tracer signals may overlap, and "
    "proppant distribution within a fracture is seldom uniform. The FO "
    "calibration derives from steady-state experiments; start-up, shut-in, "
    "and rapid drawdown may produce transient conditions where the "
    "FO-to-rate relationship deviates from steady state. Dodecane served "
    "as the model oil; crude oil components, particularly asphaltenes, "
    "may alter the wetting and diffusion behavior of the epoxy matrix. "
    "The long-term chemical stability of the epoxy and the stearic acid "
    "coating at temperatures above 120 degC, or in the presence of H2S, "
    "CO2, and high-salinity brines, has not been evaluated here and "
    "warrants dedicated investigation."
)
if "Several boundaries" not in t160:
    t160 = t160.rstrip() + lim_text
    for r in p160.runs: r.text = ''
    p160.runs[0].text = t160
    print('[160] limitations restored')

# ── FIX 3: Move synthesis [164] into Section 3.7 closing ──
# [164] currently sits between two-phase results and conclusions
# It reads like a conclusion, should stay but not duplicate conclusions
p164 = doc.paragraphs[164]
t164 = p164.text
# Keep it as Section 3.7 closing, but trim to avoid duplication with Conclusions
new164 = (
    "Taken together, the single-phase and two-phase results demonstrate "
    "that the two-component model, combined with the FO metric for "
    "two-phase conditions, provides a workable method for recovering "
    "per-interval production rates from wellhead tracer data. The same "
    "framework operates in both single-phase and two-phase flow without "
    "modification beyond the introduction of FO."
)
for r in p164.runs: r.text = ''
p164.runs[0].text = new164
print('[164] trimmed to Section 3.7 closing')

# ── FIX 4: Fix "deg" in abstract [2] ──
p2 = doc.paragraphs[2]
t2 = p2.text
if '104.6 deg' in t2:
    t2 = t2.replace('104.6 deg', '104.6 deg')
    t2 = t2.replace('72.3 deg', '72.3 deg')
    for r in p2.runs: r.text = ''
    p2.runs[0].text = t2
    print('[2] deg fixed in abstract')

# ── FIX 5: Introduction closure — ensure gap statement is at end ──
# [13] ends with "This is the gap the present work addresses" — this should be at the end of material gap [17]
p13 = doc.paragraphs[13]
p17 = doc.paragraphs[17]
t13 = p13.text
t17 = p17.text

# Move the closure from [13] to [17] if needed
old_close = "This is the gap the present work addresses."
if old_close in t13:
    t13 = t13.replace("\n\n" + old_close, "")
    # Add it to the end of [17] if not already there
    if old_close not in t17:
        t17 = t17.rstrip() + "\n\n" + old_close
    for r in p13.runs: r.text = ''
    p13.runs[0].text = t13
    for r in p17.runs: r.text = ''
    p17.runs[0].text = t17
    print('[13]->[17] closure moved')

# ── FIX 6: remove "independently corroborates" language from [152] ──
# (should have been fixed earlier but may still exist)
p152 = doc.paragraphs[152]
t152 = p152.text
if 'independently corroborates' in t152:
    t152 = t152.replace('independently corroborates', 'is consistent with')
    for r in p152.runs: r.text = ''
    p152.runs[0].text = t152
    print('[152] independently corroborates removed')

out = f'四氧化三铁环氧树脂拟合/ESP-T_final_{int(time.time())}.docx'
doc.save(out)
print(f'\nSaved: {out}')