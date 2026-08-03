#!/usr/bin/env python3
"""Correct narrative: model interprets sustained-release proppant breakthrough curves."""
from docx import Document
import time
doc = Document('四氧化三铁环氧树脂拟合/ESP-T_final_1781618309.docx')

# INTRODUCTION
p11 = doc.paragraphs[11]
for r in p11.runs: r.text = ''
p11.runs[0].text = (
    "Among available diagnostic techniques, tracer-based methods offer "
    "a practical route to per-interval production monitoring: they "
    "require no downhole instrumentation and can be deployed during "
    "routine fracturing operations [8-12]. Conventional oil-soluble "
    "tracers, however, cannot sustain long-term monitoring because they "
    "are rapidly cleared from the well. Sustained-release tracer "
    "proppants address this limitation by immobilizing the tracer within "
    "a solid carrier that releases it gradually over weeks to months. "
    "Zhao et al. [13], Zhou et al. [14], Li et al. [15], and Gong "
    "et al. [16] have demonstrated sustained-release designs based on "
    "coated ceramics and polystyrene microspheres."
    "\n\n"
    "A sustained-release proppant, however, creates an interpretation "
    "problem that an instantaneous tracer injection does not. The "
    "breakthrough curve recorded at the wellhead is a composite of two "
    "overlapping contributions: a concentration pulse from tracer that "
    "accumulated in the proppant pack during shut-in, and a sustained "
    "tail from tracer that continues to release from the carrier matrix "
    "throughout production. These contributions cannot be separated by "
    "inspecting the concentration curve, yet distinguishing between them "
    "is essential for extracting production-rate information. Current "
    "practice characterizes release using the empirical Korsmeyer-Peppas "
    "power law [17,18], which describes how much tracer leaves the "
    "carrier over time but provides no route to interpreting the "
    "resulting wellhead signal in terms of production rate. Efforts "
    "using ADE-based Gaussian fitting extract fracture geometry [19]; "
    "tracer mass-balance methods [20] allocate production without "
    "exploiting the shape of the breakthrough curve. A quantitative "
    "transport model that connects the sustained release of a tracer "
    "proppant to the observed breakthrough signal, and from there to "
    "the oil production rate, has not been developed [26]."
)
print('[11] Introduction rewritten')

p13 = doc.paragraphs[13]
for r in p13.runs: r.text = ''
p13.runs[0].text = (
    "To validate such a model experimentally, one needs a tracer "
    "proppant that exhibits sustained, oil-phase release under downhole "
    "conditions. Existing tracer proppants have material limitations. "
    "Coated designs lose their tracer function when the coating "
    "dissolves. Polystyrene microspheres lack thermal stability above "
    "approximately 200 degC [27-30]. Epoxy resin, with its highly "
    "cross-linked network, offers thermal stability exceeding 350 degC, "
    "exceptional chemical resistance, and tunable mechanical properties "
    "[31,32]. Direct emulsion polymerization of epoxy microspheres "
    "consolidates density control, nanoparticle modification, and "
    "tracer encapsulation into a single synthetic step. Li et al. [33] "
    "and Wei et al. [34] demonstrated that epoxy matrices can carry "
    "water-soluble tracers for aqueous inflow profiling. No oleophilic "
    "epoxy tracer proppant for oil-phase monitoring has been reported."
)
print('[13] rewritten')

p17 = doc.paragraphs[17]
for r in p17.runs: r.text = ''
p17.runs[0].text = (
    "Here we develop a transport model that interprets the breakthrough "
    "curve of a sustained-release tracer proppant by decomposing it "
    "into physically distinct components: a Gaussian pulse from the "
    "shut-in accumulation slug, and an erfc tail from the ongoing "
    "matrix-diffusion-controlled release, joined by a smooth tanh "
    "transition. The decomposition yields a direct relationship between "
    "the concentration history and the oil production rate. We "
    "synthesize an oleophilic epoxy/Fe3O4 sustained-release tracer "
    "proppant (ESP-T) by emulsion polymerization, using stearic acid-"
    "modified nano-Fe3O4@SA as the oleophilic tracer. We characterize "
    "its structure, thermal stability, wettability, transport "
    "selectivity, and mechanical integrity, and quantify its release "
    "kinetics via the Korsmeyer-Peppas model. We then validate the "
    "interpretation model in two flow configurations relevant to field "
    "operations: single-phase oil flow following a shut-in period, "
    "and steady-state oil-water two-phase flow."
)
print('[17] rewritten')

for r in doc.paragraphs[21].runs: r.text = ''

# SECTION 3.7
p134 = doc.paragraphs[134]
for r in p134.runs: r.text = ''
p134.runs[0].text = (
    "The breakthrough curve of a sustained-release tracer proppant is "
    "fundamentally different from that of an instantaneous tracer "
    "injection. When the well opens after shut-in, tracer that "
    "accumulated in the proppant pack is swept to the wellhead as a "
    "concentrated slug. At the same time, tracer continues to release "
    "from the carrier matrix, producing a persistent tail that overlaps "
    "with the later portion of the slug signal. The measured "
    "concentration curve is the sum of these two contributions, and "
    "neither can be isolated by inspection."
    "\n\n"
    "We separate them by constructing a transport model that explicitly "
    "represents both processes. The one-dimensional advection-dispersion "
    "equation, dC/dt + v dC/dx = D d2C/dx2, governs the evolution of "
    "tracer concentration along the production tubing, where "
    "v = 4Q/(pi d2) is the flow velocity and D = alpha v the "
    "dispersion coefficient. For the shut-in slug we apply the classical "
    "instantaneous-injection solution [25], which takes the form of a "
    "Gaussian pulse. For the sustained release we apply the continuous-"
    "source solution, an erfc function. The two solutions are joined "
    "through a hyperbolic-tangent weighting function, "
    "w(t) = 1/2 [1 + tanh((t-t0)/sigma)], which provides a smooth, "
    "physically continuous transition from the pulse-dominated early "
    "period to the diffusion-dominated tail. The full composite model, "
    "C(t) = cb + w(t) C_rise + [1-w(t)] C_fall, is given as Eq. (2). "
    "The parameter sigma is fixed at the sampling interval (4 min); "
    "the six free parameters are listed below."
)
print('[134] rewritten')

p148 = doc.paragraphs[148]
for r in p148.runs: r.text = ''
p148.runs[0].text = (
    "The model was fitted to the breakthrough curve obtained from "
    "the single-phase ESP-T displacement experiment (Figure 3-8b, "
    "Table 3-3). The fit is close (R2 = 0.9939, RMSE = 0.0210), but "
    "the fitted parameter values are more informative than the "
    "goodness of fit. The model recovers an effective flow rate "
    "Q = 0.46 mL/min, within 8% of the independently set pump rate "
    "of 0.50 mL/min. This means the production rate can be estimated "
    "from the tracer signal without independent flow measurement. "
    "The fitted mean residence time MRT = 37.4 min agrees with the "
    "convective travel time x/v = 38.6 min, confirming the physical "
    "consistency of the model time scale."
    "\n\n"
    "The Peclet number Pe = x/alpha = 0.934 places the transport "
    "regime at the advection-dispersion transition, consistent with "
    "a gradual-release source. Integrating the two components "
    "separately shows that 47% of the total integrated signal "
    "originates from the erfc tail. Nearly half the tracer detected "
    "at the wellhead during the sampling period comes from sustained "
    "release rather than from the shut-in slug, a finding with direct "
    "implications for monitoring strategy: a single shut-in period "
    "suffices to generate a persistent, interpretable signal."
)
print('[148] rewritten')

p152 = doc.paragraphs[152]
for r in p152.runs: r.text = ''
p152.runs[0].text = (
    "The fitted parameters carry practical meaning beyond the fit "
    "itself. Q, MRT, and Pe converge from independent directions onto "
    "values consistent with the known experimental conditions, "
    "indicating that the model recovers physical quantities rather "
    "than accommodating noise. The 47% tail contribution captures "
    "the defining characteristic of a sustained-release proppant: "
    "the signal does not return to baseline after the slug passes "
    "because release continues. The model quantifies this and "
    "connects it to the production rate. The tail-plateau "
    "concentration cb reflects the steady-state release flux. Because "
    "the release kinetics of ESP-T are thermally activated "
    "(K increasing from 0.055 at 30 degC to 0.196 at 120 degC, "
    "Section 3.6), cb is expected to be temperature-dependent, "
    "providing a route to calibrating the model for wells at "
    "different depths."
    "\n\n"
    "This interpretation approach differs from prior applications "
    "of ADE-based analysis to tracer data, which have focused on "
    "extracting fracture geometry [19], performing mass-balance "
    "allocation [20], or characterizing fracture networks [26]. "
    "These methods treat the transport step as a means to a "
    "characterization endpoint. The two-component decomposition "
    "developed here treats the transport itself as the source of "
    "the production metric, by resolving the breakthrough curve "
    "into the contributions that a sustained-release proppant "
    "naturally generates."
)
print('[152] rewritten')

p154 = doc.paragraphs[154]
for r in p154.runs: r.text = ''
p154.runs[0].text = (
    "The single-phase analysis demonstrates that the model can "
    "interpret the breakthrough curve when only oil flows through "
    "the proppant pack. In producing wells, however, formation "
    "water eventually breaks through. To test whether the "
    "interpretation remains valid under two-phase conditions, we "
    "conducted steady-state displacement experiments at three "
    "oil-water ratios (OWR = 4:1, 1:1, 1:4) and four total flow "
    "rates (0.1-0.4 mL/min). At OWR = 1:1, the early effluent "
    "showed a higher oil fraction, consistent with the oil-permeable "
    "character of ESP-T (Section 3.5)."
)
print('[154] rewritten')

out = f'四氧化三铁环氧树脂拟合/ESP-T_final_{int(time.time())}.docx'
doc.save(out)
print(f'Saved: {out}')