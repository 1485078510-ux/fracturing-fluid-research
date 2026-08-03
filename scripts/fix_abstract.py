#!/usr/bin/env python3
"""修改摘要，使其全面概括发明内容（产品+制备方法+压裂液体系+示踪方法）"""
from docx import Document
from docx.oxml.ns import qn

input_path = r"c:\Users\郝\Desktop\claude\荧光压裂液\专利申请文件_正式格式_修改版.docx"

doc = Document(input_path)

old_abstract = (
    "本发明公开了一种用于压裂裂缝荧光示踪的改性稀土铝酸盐荧光粉、含该荧光粉的压裂液体系及其应用方法。"
    "所述改性荧光粉以SrAl₂O₄:Eu²⁺,Dy³⁺为基体，表面包覆硅烷偶联剂化学键合内层和聚乙二醇物理屏蔽外层。"
    "注入阶段PEG外层通过空间位阻效应保障分散稳定性；关井破胶阶段PEG在破胶剂氧化环境与储层温度作用下"
    "脱附降解，暴露内层活性氨基官能团并与砂岩壁面硅羟基通过静电吸引、氢键和化学缩合实现牢固锚定。"
    "本发明利用压裂施工工序时序驱动功能切换，无需额外触发剂，与现有工艺良好兼容，可为压裂裂缝提供可实物验证的持久荧光标记。"
)

new_abstract = (
    "本发明公开了一种用于压裂裂缝荧光示踪的改性稀土铝酸盐荧光粉、其制备方法、含该荧光粉的压裂液体系"
    "及其应用方法。所述改性荧光粉以SrAl₂O₄:Eu²⁺,Dy³⁺等稀土铝酸盐长余辉材料为基体，表面依次包覆"
    "硅烷偶联剂化学键合内层和聚乙二醇（PEG）物理屏蔽外层，并辅以螯合剂与非离子表面活性剂构成的协同"
    "分散助剂体系；其制备采用先在醇-水混合溶剂中进行硅烷偶联剂化学接枝、再在水相中进行PEG物理包覆的"
    "两步法。所述荧光压裂液体系以羟丙基胍胶为稠化剂、有机硼为交联剂、过硫酸铵为破胶剂，采用母液预配"
    "与在线稀释相结合的方式配制。压裂施工中，注入阶段PEG外层保障荧光粉的分散稳定性；关井破胶阶段"
    "PEG在破胶剂氧化环境中脱附降解，暴露内层活性氨基官能团并与砂岩壁面硅羟基通过静电吸引、氢键和"
    "化学缩合实现牢固锚定；返排后取心，紫外照射即可直接观察裂缝壁面荧光分布。本发明以压裂施工工序"
    "时序驱动功能切换，无需额外触发剂，可为压裂裂缝提供可实物验证的持久荧光标记。"
)

for para in doc.paragraphs:
    if old_abstract in para.text:
        replace_paragraph_text = None
        # 保留第一个run格式，替换文本
        first_run = para.runs[0] if para.runs else None
        if first_run:
            for i in range(len(para.runs) - 1, 0, -1):
                para.runs[i]._element.getparent().remove(para.runs[i]._element)
            first_run.text = new_abstract
            print("摘要已更新为全文概括版本")
        break
else:
    print("未找到原摘要文本，尝试模糊匹配...")
    for para in doc.paragraphs:
        if "一种用于压裂裂缝荧光示踪的改性稀土铝酸盐荧光粉" in para.text and len(para.text) > 100:
            first_run = para.runs[0] if para.runs else None
            if first_run:
                for i in range(len(para.runs) - 1, 0, -1):
                    para.runs[i]._element.getparent().remove(para.runs[i]._element)
                first_run.text = new_abstract
                print("摘要已更新（模糊匹配）")
            break

doc.save(input_path)
print("完成！")