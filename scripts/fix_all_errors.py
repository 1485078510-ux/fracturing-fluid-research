# -*- coding: utf-8 -*-
"""Fix ALL errors found in full document review."""
from docx import Document
from docx.oxml.ns import qn
import re

DST = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP_formatted.docx'
doc = Document(DST)

IMG = set()
for i, p in enumerate(doc.paragraphs):
    for run in p.runs:
        if (run._element.findall('.//'+qn('w:drawing')) or
            run._element.findall('.//'+qn('a:blip'))):
            IMG.add(i); break

def sp(idx, text):
    if idx in IMG: return
    p = doc.paragraphs[idx]
    for r in p.runs: r.text = ''
    if p.runs: p.runs[0].text = text
    else: p.add_run(text)

def clear(idx):
    if idx in IMG: return
    p = doc.paragraphs[idx]
    for r in p.runs: r.text = ''

# ===== 1. ABSTRACT — fix all units =====
sp(4,
    "Nano-Fe3O4@SA disperses uniformly as nanoclusters in the epoxy matrix. "
    "ESP-T exhibits sphericity and roundness exceeding 0.9 with a bulk density of "
    "0.646 g·cm⁻³, and maintains structural integrity up to 357.27 °C, "
    "well above typical downhole temperatures (80–200 °C). Incorporation of "
    "nano-Fe3O4@SA raises the water contact angle from 72.3° to 104.6°, "
    "yielding a water-resistant, oil-permeable characteristic: the oil filtration time "
    "(5 min 11 s) is 66.1% shorter than that of pure epoxy microspheres. "
    "Tracer release at 30–120 °C follows the Korsmeyer–Peppas model "
    "(R² > 0.90) with diffusion exponents of 0.45–0.85, indicating synergistic "
    "Fickian diffusion and Case-II relaxation. Cumulative Fe release at 120 °C "
    "exceeds 2.0 mg·L⁻¹ over 14 days, meeting ICP-MS detection "
    "requirements for long-term monitoring.")

sp(6,
    "A piecewise advection–dispersion model with a smooth tanh transition is "
    "developed to interpret the single-phase tracer breakthrough curve. The model "
    "decomposes the signal into a Gaussian pulse component, representing the tracer "
    "slug accumulated during shut-in, and an erfc tailing component, representing "
    "sustained matrix-diffusion-controlled release from ESP-T. The model achieves "
    "R² = 0.9939, with the fitted flow rate (0.46 mL·min⁻¹) "
    "agreeing closely with the pump-set rate (0.5 mL·min⁻¹, "
    "relative error 8%). The mean residence time (37.4 min) matches the convective "
    "travel time (38.6 min, ratio 0.967). Under steady-state two-phase flow, tracer "
    "flux effectively quantifies oil-phase production rates across varying oil–water "
    "ratios. These results establish ESP-T as a robust dual-function platform for "
    "long-term production monitoring in unconventional reservoirs.")

# ===== 2. SECTION 3.3 — fix thermal analysis units =====
sp(91,
    "As depicted in Figure 3-4, thermal analysis of the proppant reveals three "
    "distinct decomposition stages. The first stage (50–350 °C) corresponds "
    "to a minor weight loss of 5.70%, attributed to the removal of adsorbed surface "
    "water and residual ethanol employed as a diluent during synthesis. The second "
    "stage (350–400 °C) constitutes the primary thermal decomposition region, "
    "arising from cleavage of C–O–C and C–C bonds within the epoxy "
    "molecular chains and degradation of the cross-linked network, accompanied by "
    "the release of CO₂ and small-molecule hydrocarbons. The DTG curve indicates "
    "a maximum weight-loss rate at 357.27 °C, with a mass loss of 72.5%. "
    "In the third stage (>400 °C), the residual mass stabilizes; the remaining "
    "weight loss originates from oxidative combustion of carbonaceous residues formed "
    "in the second stage. The final residue comprises hollow glass microspheres, "
    "thermally stable nano-Fe3O4@SA, and a minor carbonaceous fraction from epoxy "
    "decomposition.")

# ===== 3. SECTION 3.10 — REMOVE EDITING NOTE =====
sp(138,
    "Based on the release results, the release mechanism of ESP-T is inferred as "
    "follows. During the release process, nonpolar solvents (e.g., alkanes) permeate "
    "into the polymer network, inducing swelling of the cross-linked epoxy matrix. "
    "At this stage, the polymer microspheres form a dual-state structure: an inner "
    "glassy core and an outer gel layer. The swelling reduces the entanglement of "
    "polymer molecular chains, thereby creating expanded transport channels that "
    "enable nano-Fe3O4@SA tracers to diffuse into the external environment. Elevated "
    "temperature enhances the permeability of solvent molecules into the polymer "
    "matrix, accelerating the swelling rate. Furthermore, increased temperature "
    "weakens the intermolecular forces within the polymer network, expanding the "
    "pore size and further promoting tracer diffusion through the matrix.")

# ===== 4. REMOVE DUPLICATE sphericity [106] =====
clear(106)

# ===== 5. FIX missing sigma variable =====
# Currently: [151] c_b, [152] A,a, [153] alpha, [154] Q, [155] t0
# Need: [151] c_b, [152] A,a, [153] alpha, [154] Q, [155] t0, [156 before model] sigma
# But [156] is the model description. Add sigma to [155] with t0, or add as new line.
# Let me combine t0 and sigma into [155], and shift model description to [156]
sp(155,
    "t0 = time at which sustained matrix-diffusion release overtakes the shut-in "
    "slug as the primary contributor to the detected tracer signal (min); "
    "sigma = characteristic timescale over which this transition occurs, reflecting "
    "the width of the crossover between slug-dominated and release-dominated "
    "tracer transport (min).")

# ===== 6. FIX missing space in Section 3.10 [128] =====
t128 = doc.paragraphs[128].text
# Fix "concentration,C" -> "concentration, C"
# This needs run-level fix. Use text replacement.
sp(128,
    "In field applications, formation temperature rises with increasing reservoir "
    "depth, impacting the expansion of ESP-T and tracer release behavior. The release "
    "profiles of ESP-T at various temperatures are presented in Figure 3-8(a), where "
    "C/C0 represents the normalized concentration, C denotes the tracer concentration "
    "at a specific time, and C0 is the maximum release concentration. As indicated "
    "by these profiles, the relative release concentration increases over time and "
    "with rising temperature, demonstrating that elevated temperature accelerates "
    "tracer release. Additionally, Figure 3-8(c) shows that the tracer release rate "
    "is high during initial contact with the fluid, then gradually slows and "
    "stabilizes. The cumulative tracer release at 120 °C over 14 days remains "
    "detectable by ICP-MS, confirming its suitability for long-term stable monitoring. "
    "To better understand the release behavior within the carrier and optimize tracer "
    "design, the release mechanism and kinetic characteristics of ESP-T were "
    "investigated. The Korsmeyer-Peppas (K-P) release kinetic model was used to "
    "fit the release curves at different temperatures, and the release behavior of "
    "ESP-T was interpreted from the fitted parameters. The K-P model is defined "
    "by Equation (5):")

# ===== 7. FIX reference [6] mangled formatting =====
sp(187,
    "[6] IEA. World Energy Employment 2025 [Z]. https://www.iea.org/reports/"
    "world-energy-employment-2025, 2025.")

# ===== 8. FIX conclusions units =====
sp(175,
    "ESP-T demonstrates strong physical, mechanical, and thermal performance: "
    "bulk density of 0.646 g·cm⁻³ for favorable suspension; "
    "apparent density of 1.072 g·cm⁻³ indicating compact internal "
    "packing; crush rate of 2.9% at 50 MPa; acid solubility of 3.3% (meeting the "
    "≤ 5% industry standard); and an initial decomposition temperature of "
    "357.27 °C, far exceeding typical downhole conditions (80–200 °C). "
    "The water contact angle increases from 72.3° to 104.6° after "
    "nano-Fe3O4@SA incorporation, and the oil filtration time of 5 min 11 s is "
    "66.1% shorter than that of pure epoxy microspheres, validating the "
    "water-resistant, oil-permeable transport characteristic that facilitates "
    "oil production while suppressing water breakthrough.")

sp(177,
    "Tracer release from ESP-T at 30–120 °C follows the Korsmeyer–Peppas "
    "(K-P) model (R² > 0.90; diffusion exponent n = 0.45–0.85), governed "
    "by synergistic Fickian diffusion and Case-II relaxation of the cross-linked "
    "epoxy matrix. Cumulative Fe release at 120 °C exceeds 2.0 mg·L⁻¹ "
    "over 14 days, meeting ICP-MS detection limits for long-term downhole monitoring. "
    "A piecewise advection–dispersion model with smooth tanh transition was "
    "developed to interpret the single-phase tracer breakthrough curve, decomposing "
    "the signal into a Gaussian pulse component (shut-in accumulation slug) and an "
    "erfc tailing component (sustained matrix-diffusion-controlled release). The "
    "model achieves R² = 0.9939; the fitted flow rate (0.46 mL·min⁻¹) "
    "closely matches the pump-set rate (0.5 mL·min⁻¹, relative error "
    "8%); and the mean residence time (37.4 min) agrees with the convective travel "
    "time (38.6 min, ratio 0.967). The erfc tailing component accounts for 47% of "
    "the integrated tracer signal, confirming the dominant role of matrix-diffusion-controlled "
    "release in sustaining the long-term monitoring signal. Under steady-state "
    "oil–water two-phase flow, tracer flux effectively quantifies per-interval "
    "oil-phase production rates across varying oil–water ratios, providing "
    "a practical basis for field-scale production allocation.")

# ===== SAVE =====
doc.save(DST)
print("All errors fixed:")
print("  - Abstract: units corrected (g/cm3->g-cm-3, C->C, R2->R2, mg/L->mg-L-1)")
print("  - Section 3.3: thermal analysis units (CO2, C)")
print("  - Section 3.10: editing note REMOVED")
print("  - Section 3.5: duplicate sphericity paragraph REMOVED")
print("  - Section 3.11: sigma parameter added to variable list")
print("  - Section 3.10: missing space fixed (concentration, C)")
print("  - Reference [6]: mangled formatting fixed")
print("  - Conclusions: all units corrected")
print(f"Saved: {DST}")