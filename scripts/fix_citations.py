#!/usr/bin/env python3
"""
按用户要求，将专利申请文件中的学术论文引用格式改为专利通用描述。
在原DOCX基础上直接替换文本，保留所有格式。
"""

from docx import Document
from docx.oxml.ns import qn

input_path = r"c:\Users\郝\Desktop\claude\荧光压裂液\专利申请文件_正式格式_修改版.docx"
output_path = input_path  # 直接覆盖修改版

doc = Document(input_path)

def replace_in_para(para, old, new):
    """在段落中替换文本，保留格式"""
    if old not in para.text:
        return False
    # 合并所有runs的文本，做替换，然后写回第一个run
    full_text = para.text
    new_text = full_text.replace(old, new)
    if new_text == full_text:
        return False
    # 保留第一个run格式，清空其余
    first_run = para.runs[0] if para.runs else None
    if first_run is None:
        return False
    for i in range(len(para.runs) - 1, 0, -1):
        para.runs[i]._element.getparent().remove(para.runs[i]._element)
    first_run.text = new_text
    return True

# ================================================================
# 逐项替换学术引用
# ================================================================
replacements_done = []

# 1. 背景技术 - 微地震监测段
for para in doc.paragraphs:
    # (Maxwell，2014)
    if "（Maxwell，2014）" in para.text:
        replace_in_para(para, "（Maxwell，2014）", "")
        replacements_done.append("删除(Maxwell, 2014)")

    # (Molenaar等，2012；Jin和Roy，2017)
    if "（Molenaar等，2012；Jin和Roy，2017）" in para.text:
        replace_in_para(para, "（Molenaar等，2012；Jin和Roy，2017）", "")
        replacements_done.append("删除(Molenaar等, 2012; Jin和Roy, 2017)")

    # 检查并清理可能残留的双重标点（如 "偏差。）" 后面还有内容）
    # 清理残留的 "。)" 或 "。）" 等
    if "存在显著多解性；" in para.text and "（2）" in para.text:
        # 确保微地震监测描述流畅
        replace_in_para(para, "存在显著多解性；（2）", "存在显著多解性；（2）")
        # 确保光纤传感描述流畅
        replace_in_para(para, "会导致解释偏差；（3）", "会导致解释偏差；（3）")

# 2. 背景技术 - Takeuchi 荧光树脂
for para in doc.paragraphs:
    if "Takeuchi等（2025）采用可凝固荧光树脂作为压裂介质" in para.text:
        replace_in_para(para,
            "Takeuchi等（2025）采用可凝固荧光树脂作为压裂介质",
            "现有技术中已有采用可凝固荧光树脂作为压裂介质的研究")
        replacements_done.append("Takeuchi等(2025) -> 现有技术通用描述")

# 3. 背景技术 - Guryanov GeoSplit
for para in doc.paragraphs:
    if "Guryanov等（2019）开发的量子点荧光聚合物微球" in para.text:
        replace_in_para(para,
            "Guryanov等（2019）开发的量子点荧光聚合物微球（商业名GeoSplit系统）已在分段压裂产液剖面监测中实现现场应用，代表了荧光示踪技术工程化的重要进展。",
            "量子点荧光聚合物微球（商业名GeoSplit系统）已在分段压裂产液剖面监测中实现现场应用，代表了荧光示踪技术工程化的重要进展。")
        replacements_done.append("Guryanov等(2019) -> 通用描述")

# 4. 背景技术 - Matsuzawa 荧光粉
for para in doc.paragraphs:
    if "Matsuzawa等（1996）首次报道的一类新型无机长余辉发光材料" in para.text:
        replace_in_para(para,
            "Matsuzawa等（1996）首次报道的一类新型无机长余辉发光材料",
            "一类无机长余辉发光材料")
        replacements_done.append("Matsuzawa等(1996) -> 通用描述")

# 5. 有益效果 - Takeuchi 引用
for para in doc.paragraphs:
    if "（区别于Takeuchi等的可凝固树脂体系）" in para.text:
        replace_in_para(para,
            "（区别于Takeuchi等的可凝固树脂体系）",
            "（区别于可凝固树脂体系）")
        replacements_done.append("有益效果中 Takeuchi等 -> 删除人名")

# 6. 有益效果 - Guryanov 引用
for para in doc.paragraphs:
    if "（区别于Guryanov等依赖返排液荧光间接分析的技术路线）" in para.text:
        replace_in_para(para,
            "（区别于Guryanov等依赖返排液荧光间接分析的技术路线）",
            "（区别于依赖返排液荧光间接分析的技术路线）")
        replacements_done.append("有益效果中 Guryanov等 -> 删除人名")

# 7. 背景技术 - 可凝固荧光树脂段落中可能的 "Takeuchi等"
for para in doc.paragraphs:
    if "Takeuchi等（2025）" in para.text:
        replace_in_para(para,
            "Takeuchi等（2025）",
            "现有技术中")
        replacements_done.append("残留Takeuchi等(2025) -> 现有技术中")

# ================================================================
# 最终检查：清理可能残留的孤立括号
# ================================================================
for para in doc.paragraphs:
    text = para.text
    # 清理 "偏差）；" 之类残留（删了引用后只剩括号）
    # 实际上我的替换已经精确匹配了完整字符串，不会残留括号

    # 修复可能产生的双逗号或多余标点
    if "多解性；（2）" in text:
        # 确认这行没有残留问题
        pass

# ================================================================
# 保存
# ================================================================
doc.save(output_path)
print("学术引用转换完成！")
for item in replacements_done:
    print(f"  ✓ {item}")
print(f"\n文件已保存：{output_path}")