"""Fix Conclusions heading and clear stray old paragraph."""
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy

INPUT = r"c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP-T_投稿文件\ESP-T_Final_4-revised.docx"

doc = Document(INPUT)

# Find the paragraph that's currently Heading 1 style with our new conclusion text
# and find References
conc_body_as_heading = None
refs_idx = None
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if p.style.name.startswith('Heading') and 'jointly estimating release' in text:
        conc_body_as_heading = i
    if text == 'References' or text.startswith('References'):
        refs_idx = i

print(f"Conclusion body as heading: para {conc_body_as_heading}")
print(f"References: para {refs_idx}")

if conc_body_as_heading:
    # 1. Change this paragraph's style from Heading 1 to Normal
    para = doc.paragraphs[conc_body_as_heading]
    para.style = doc.styles['Normal']
    print(f"  Changed para {conc_body_as_heading} style to Normal")

    # 2. Insert a "4. Conclusions" heading before this paragraph
    # We need to insert a new paragraph element before this one
    from docx.oxml import OxmlElement
    new_heading = OxmlElement('w:p')

    # Add paragraph properties for Heading 1 style
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Heading1')  # or the actual heading style ID
    pPr.append(pStyle)
    new_heading.append(pPr)

    # Add run with "4. Conclusions" text
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = '4. Conclusions'
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    new_heading.append(r)

    # Insert before the body paragraph
    para._element.addprevious(new_heading)
    print(f"  Inserted '4. Conclusions' heading before para {conc_body_as_heading}")

# 3. Find and clear the stray old conclusion paragraph
# It starts with "The present validation is limited to single-interval laboratory experiments"
# but does NOT contain "co-precipitation synthesis" (our new version has this)
if refs_idx:
    for i in range(refs_idx - 5, refs_idx):
        text = doc.paragraphs[i].text.strip()
        if text.startswith('The present validation is limited to single-interval'):
            if 'co-precipitation synthesis' not in text:
                # This is the OLD stray paragraph
                for r in list(doc.paragraphs[i].runs):
                    r._element.getparent().remove(r._element)
                for child in list(doc.paragraphs[i]._element):
                    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if tag in ('r', 'drawing', 'pict'):
                        doc.paragraphs[i]._element.remove(child)
                print(f"  Cleared stray old conclusion at para {i}")

print("\nSaving...")
doc.save(INPUT)
print("Done!")
