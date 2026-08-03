# -*- coding: utf-8 -*-
"""Fix conclusions: remove duplicates, update to match Section 3.11 content."""
from docx import Document

DST = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP_polished.docx'
doc = Document(DST)

def set_text(para, text):
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)

# Fix heading
set_text(doc.paragraphs[170], "4. Conclusions")

# [171] Synthesis — keep as is (already correct)

# [172] Physical properties — update to polished version
set_text(doc.paragraphs[172],
    "ESP-T demonstrated superior physical, mechanical, and thermal stability: a bulk density "
    "of 0.646 g/cm³ ensured favorable suspension in water-based fracturing fluids; an apparent "
    "density of 1.072 g/cm³ indicated a compact internal structure; the breakage rate at 50 MPa "
    "was 2.9%; and the acid solubility of 3.3% met the industrial requirement (≤ 5%), confirming "
    "applicability to acidic reservoirs. The initial thermal decomposition temperature of "
    "357.27 °C far exceeded typical downhole temperatures (80–200 °C), eliminating concerns "
    "of resin degradation during long-term service. The water contact angle increased from "
    "72.3° to 104.6°, and the oil filtration time (5 min 11 s) was 66.1% shorter than that "
    "of pure epoxy microspheres, validating the 'water-resistant and oil-permeable' transport "
    "characteristic.")

# [174] — REMOVE (old duplicate of tracer performance)
p = doc.paragraphs[174]
p._element.getparent().remove(p._element)

# [174] is now the old "In summary" paragraph (shifted up from 175)
# Update [174] (was 175) — final summary
set_text(doc.paragraphs[174],
    "In summary, ESP-T integrates fracture propping and tracer monitoring functions, making it "
    "suitable for acid fracturing, deep wells, and high-pressure formations. It provides "
    "reliable technical support for fracturing parameter optimization and enhanced oil recovery, "
    "demonstrating significant industrial application potential in unconventional oil and gas "
    "development.")

# [173] Tracer — already has the updated text from polish script, verify
# Add the Section 3.11 key numbers that are currently missing
set_text(doc.paragraphs[173],
    "In terms of tracer performance, Fe release from ESP-T at 30–120 °C followed the "
    "Korsmeyer–Peppas (K-P) kinetic model. The fitted diffusion exponent n ranged from "
    "0.45 to 0.85 at all tested temperatures, indicating a synergistic release mechanism "
    "governed by both Fickian diffusion and Case-II relaxation, with R² > 0.90 for all "
    "fittings. The cumulative tracer release at 120 °C over 14 days exceeded 2.0 mg/L, "
    "meeting ICP-MS detection requirements for long-term monitoring. Under shut-in "
    "single-phase flow, a piecewise advection-dispersion model with smooth tanh transition "
    "was developed to interpret the tracer breakthrough curve. The model decomposes the "
    "signal into a Gaussian pulse component (tracer slug from shut-in accumulation) and an "
    "erfc tailing component (sustained matrix-diffusion-controlled release from ESP-T). "
    "Fitting yielded R² = 0.9939, with a derived flow velocity of 2.59 cm/min, dispersivity "
    "of 107.1 cm, and Péclet number Pe = 0.934. The fitted flow rate (0.46 mL/min) agreed "
    "closely with the pump-set rate (0.5 mL/min, relative error 8%), and the mean residence "
    "time (37.4 min) matched the theoretical convective travel time (38.6 min, ratio 0.967). "
    "Deconvolution showed that 47% of the tracer signal originated from the tailing component, "
    "confirming the significant contribution of matrix-diffusion-controlled release to the "
    "overall monitoring signal. Under steady-state two-phase flow, tracer concentration "
    "decreased monotonically with increasing total flow rate and was independent of the "
    "oil–water ratio. The tracer flux method effectively quantified oil-phase flow rates in "
    "individual intervals, providing a reliable basis for dynamic production monitoring "
    "during the steady-state stage.")

# Remove empty line after [176] if exists
# Check what's at 175 and 176 now
for i in [175, 176]:
    t = doc.paragraphs[i].text.strip()
    if not t:
        print(f"  Removing empty para [{i}]")
        doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)

doc.save(DST)
print(f"Fixed: {DST}")
print("Changes: removed duplicate [174], updated heading, unified conclusions with Section 3.11 content")