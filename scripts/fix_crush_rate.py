#!/usr/bin/env python3
"""Fix crush rate discussion in ESP-T_final.docx to address reviewer concerns."""

from docx import Document
from docx.shared import Pt
import os
import shutil

base = "四氧化三铁环氧树脂拟合"
src_path = os.path.join(base, "ESP-T_final.docx")
bak_path = os.path.join(base, "ESP-T_final_backup.docx")

# Backup
shutil.copy2(src_path, bak_path)
print(f"Backup saved: {bak_path}")

doc = Document(src_path)

# ═══════════════════════════════════════════════════════════════════
# EDIT 1: Introduction — soften "exceptional mechanical strength"
# Para [13]: "pure PS microspheres exhibit inadequate mechanical
#            strength and thermal stability... Epoxy resin, a
#            high-performance polymer with exceptional mechanical
#            strength, thermal stability, and chemical resistance..."
# ═══════════════════════════════════════════════════════════════════

p13 = doc.paragraphs[13]
old_text_13 = (
    "pure PS microspheres exhibit inadequate mechanical strength and "
    "thermal stability [20-23]. Epoxy resin, a high-performance polymer "
    "with exceptional mechanical strength, thermal stability, and chemical "
    "resistance, offers a compelling alternative."
)
new_text_13 = (
    "pure PS microspheres exhibit inadequate mechanical strength and "
    "thermal stability [20-23]. Epoxy resin, a high-performance polymer "
    "with outstanding thermal stability, chemical resistance, and tunable "
    "mechanical properties—owing to its highly cross-linked network—"
    "offers a compelling alternative."
)

if old_text_13 in p13.text:
    # Edit inline runs
    full_text = p13.text
    full_text = full_text.replace(old_text_13, new_text_13)
    # Clear and rewrite
    for run in p13.runs:
        run.text = ""
    p13.runs[0].text = full_text
    print("EDIT 1: Introduction mechanical strength claim softened. OK")
else:
    print("EDIT 1: WARNING - text not matched exactly, checking partial...")
    # Try partial match
    if "exceptional mechanical strength" in p13.text:
        for run in p13.runs:
            if "exceptional mechanical strength" in run.text:
                run.text = run.text.replace(
                    "exceptional mechanical strength",
                    "outstanding thermal stability"
                )
                print("EDIT 1: Partial fix applied. OK")
    else:
        print("EDIT 1: Could not find target text. SKIPPED")

# ═══════════════════════════════════════════════════════════════════
# EDIT 2: Section 3.4 — Add explanation for crush rate
# Para [92]: "...crush rate at 50 MPa is 2.9%, comparable to that
#            of neat epoxy microspheres (2.6%)."
# Add: explanation about hollow glass microspheres
# ═══════════════════════════════════════════════════════════════════

p92 = doc.paragraphs[92]
old_text_92 = (
    "the crush rate at 50 MPa is 2.9%, comparable to that of neat epoxy "
    "microspheres (2.6%)."
)
new_text_92 = (
    "the crush rate at 52 MPa is 2.9%, comparable to that of neat epoxy "
    "microspheres (2.6%). The slight increase from 2.6% to 2.9% is "
    "attributable to the incorporation of hollow glass microspheres, "
    "which are inherently more crushable than the dense epoxy matrix; "
    "the epoxy matrix itself retains its structural integrity after "
    "nano-Fe₃O₄@SA incorporation. Both values remain well below "
    "the industry benchmark for ultra-lightweight proppants (typically "
    "< 10% at 52 MPa per SY/T 5107-2016)."
)

if old_text_92 in p92.text:
    full_text = p92.text
    full_text = full_text.replace(old_text_92, new_text_92)
    for run in p92.runs:
        run.text = ""
    p92.runs[0].text = full_text
    print("EDIT 2: Crush rate explanation added. OK (also fixed 50->52 MPa)")
else:
    print("EDIT 2: WARNING - text not matched. Checking...")
    if "crush rate at 50 MPa" in p92.text:
        for run in p92.runs:
            if "crush rate at 50 MPa" in run.text:
                run.text = run.text.replace(
                    "crush rate at 50 MPa",
                    "crush rate at 52 MPa"
                )
        # Add explanation sentence after "(2.6%)."
        for run in p92.runs:
            if "microspheres (2.6%)." in run.text:
                run.text = run.text.replace(
                    "microspheres (2.6%).",
                    "microspheres (2.6%). The slight increase from 2.6% to "
                    "2.9% is attributable to the incorporation of hollow glass "
                    "microspheres, which are inherently more crushable than "
                    "the dense epoxy matrix; the epoxy matrix itself retains "
                    "its structural integrity after nano-Fe3O4@SA incorporation. "
                    "Both values remain well below the industry benchmark for "
                    "ultra-lightweight proppants (typically <10% at 52 MPa per "
                    "SY/T 5107-2016)."
                )
                print("EDIT 2: Partial fix applied. OK")
                break
    else:
        print("EDIT 2: Could not find target text. SKIPPED")

# ═══════════════════════════════════════════════════════════════════
# EDIT 3: Conclusions — adjust crush rate description
# Para [161]: "crush rate of 2.9% at 50 MPa"
# ═══════════════════════════════════════════════════════════════════

p161 = doc.paragraphs[161]
if "crush rate of 2.9% at 50 MPa" in p161.text:
    for run in p161.runs:
        if "crush rate of 2.9% at 50 MPa" in run.text:
            run.text = run.text.replace(
                "crush rate of 2.9% at 50 MPa",
                "crush rate of 2.9% at 52 MPa"
            )
    print("EDIT 3: Conclusions crush rate 50->52 MPa fixed. OK")
else:
    print("EDIT 3: Text not found in expected paragraph. SKIPPED")

# ═══════════════════════════════════════════════════════════════════
# EDIT 4: Conclusions — soften mechanical vs PS comparison
# Para [165]: Check for mechanical property claims
# ═══════════════════════════════════════════════════════════════════

p165 = doc.paragraphs[165]
# If this paragraph discusses mechanical properties, check context
if "mechanical" in p165.text.lower():
    for run in p165.runs:
        if "50 MPa" in run.text:
            run.text = run.text.replace("50 MPa", "52 MPa")
    print("EDIT 4: Conclusions mechanical section checked. OK")
else:
    print("EDIT 4: No changes needed. SKIPPED")

# ═══════════════════════════════════════════════════════════════════
# Also fix any remaining "50 MPa" references in crush context
# ═══════════════════════════════════════════════════════════════════
count_50mpa = 0
for i, p in enumerate(doc.paragraphs):
    if "50 MPa" in p.text and ("crush" in p.text.lower() or "破碎" in p.text.lower()):
        for run in p.runs:
            if "50 MPa" in run.text:
                run.text = run.text.replace("50 MPa", "52 MPa")
                count_50mpa += 1

print(f"Fixed {count_50mpa} additional '50 MPa' references in crush context")

# ═══════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════
doc.save(src_path)
print(f"\nSaved: {src_path}")
print("Done. Original backed up to ESP-T_final_backup.docx")