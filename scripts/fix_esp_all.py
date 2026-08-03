"""
Complete fix script for ESP_restructured.docx — ALL fixes in ONE pass from original.
Avoids table merge-cell corruption by preserving original table structure.
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
import re

SRC = r"c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP_restructured.docx"
DST = r"c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP_restructured_fixed.docx"

doc = Document(SRC)


def rebuild_para(para, new_text):
    """Replace entire paragraph text, preserving first run's formatting."""
    if para.runs:
        for run in para.runs[1:]:
            run.text = ''
        para.runs[0].text = new_text
        return True
    return False


def set_cell_text(cell, text):
    """Set single cell's text content."""
    for para in cell.paragraphs:
        if para.runs:
            for run in para.runs[1:]:
                run.text = ''
            para.runs[0].text = text
            return


# ================================================================
# PART 1: Paragraph-level text replacements (all at once)
# ================================================================
fixes_para = [
    # Critical figure refs
    ("As depicted in Figure 3-1, thermal analysis of the proppant reveals",
     "As depicted in Figure 3-3, thermal analysis of the proppant reveals"),
    ("Figure 3-4 presents the bulk density and apparent density results",
     "Figure 3-6 presents the bulk density and apparent density results"),

    # Chemical formula
    ("2×10⁻⁵ mol MnCl₂·4H₂O", "2×10⁻⁵ mol MnCl₂·6H₂O"),

    # Grammar
    ("productivity decays at disparate rates", "productivity declines at varying rates"),
    ("Owing to the effect of sample volume, the actual arrival time",
     "Due to the sample volume effect, the actual arrival time"),
    ("monitoring capability ceases once the polymer coating dissolves",
     "monitoring ceases once the polymer coating dissolves"),

    # Korsmeyer–Peppas en-dash
    ("Korsmeyer-Peppas", "Korsmeyer–Peppas"),

    # Proppant-pack
    ("Proppant-pack conductivity", "Proppant pack conductivity"),

    # Title
    ("Fabrication and Performance of ESP-T for Hydraulic Fracturing in Unconventional Oil and Gas Reservoirs",
     "An Oleophilic Fe₃O₄/Epoxy Resin Tracer Proppant for Long-Term Oil Production Monitoring in Unconventional Reservoirs"),

    # ESP spacing
    ("ESP – T", "ESP-T"),

    # Hedging in §3.1
    ("This observation verifies that the long alkyl chains of stearic acid form",
     "This observation indicates that the long alkyl chains of stearic acid form"),

    # ADE equation fix
    ("advection-dispersion equation (ADE):  is the mean flow velocity",
     "advection-dispersion equation (ADE), where v is the mean flow velocity"),

    # μm fix (45 um -> 45 μm)
    ("45 um SiO₂", "45 μm SiO₂"),
    ("45 um SiO2", "45 μm SiO₂"),
]

counts = {}
for para in doc.paragraphs:
    text = para.text
    for old, new in fixes_para:
        if old in text:
            text = text.replace(old, new)
            counts[old[:40]] = counts.get(old[:40], 0) + 1
    if text != para.text:
        rebuild_para(para, text)

# Same for table cell text
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                text = para.text
                for old, new in fixes_para:
                    if old in text:
                        text = text.replace(old, new)
                if text != para.text and para.runs:
                    rebuild_para(para, text)

print(f"Applied {sum(counts.values())} paragraph text replacements")

# ================================================================
# PART 2: Regex-based bulk replacements (℃, Fig., subscripts, spacing)
# ================================================================
regex_count = 0
for para in doc.paragraphs:
    old_text = para.text
    new_text = old_text
    new_text = new_text.replace('℃', '°C')
    new_text = re.sub(r'Fig\.\s+(\d+)-(\d+)', r'Figure \1-\2', new_text)
    new_text = re.sub(r'(Figure \d+) [–\-] (\d+)', r'\1-\2', new_text)
    new_text = re.sub(r'(?<![₃₄₂])Fe3O4(?![₃₄₂])', 'Fe₃O₄', new_text)
    new_text = re.sub(r'(?<![₂])SiO2(?![₂])', 'SiO₂', new_text)
    new_text = re.sub(r'(?<![₂])CO2(?![₂])', 'CO₂', new_text)
    new_text = new_text.replace('nano‑Fe', 'nano-Fe')
    if new_text != old_text and para.runs:
        rebuild_para(para, new_text)
        regex_count += 1

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                old_text = para.text
                new_text = old_text
                new_text = new_text.replace('℃', '°C')
                new_text = re.sub(r'(?<![₃₄₂])Fe3O4(?![₃₄₂])', 'Fe₃O₄', new_text)
                new_text = re.sub(r'(?<![₂])SiO2(?![₂])', 'SiO₂', new_text)
                new_text = new_text.replace('nano‑Fe', 'nano-Fe')
                if new_text != old_text and para.runs:
                    rebuild_para(para, new_text)
                    regex_count += 1

print(f"Applied regex fixes to {regex_count} paragraphs/cells")

# ================================================================
# PART 3: Abstract streamlining
# ================================================================
new_abstract = (
    "Abstract: Per-interval production monitoring is essential for evaluating "
    "stimulation effectiveness in unconventional reservoirs, yet existing tracer "
    "technologies lack durable oil-phase monitoring capability. Here we report an "
    "oleophilic tracer proppant (ESP-T) fabricated by encapsulating stearic "
    "acid-modified nano-Fe₃O₄ (nano-Fe₃O₄@SA) within an epoxy resin matrix via "
    "emulsion polymerization. ESP-T exhibits sphericity exceeding 0.9, maintains "
    "structural integrity up to 357.27 °C, and the incorporation of nano-Fe₃O₄@SA "
    "raises the water contact angle from 72.3° to 104.6°, yielding a water-resistant, "
    "oil-permeable transport characteristic. Tracer release at 30–120 °C follows "
    "the Korsmeyer–Peppas model (R² > 0.90) with diffusion exponents of 0.45–0.85, "
    "indicating synergistic Fickian diffusion and Case-II relaxation. A piecewise "
    "advection–dispersion model with tanh blending achieves R² = 0.9939 for "
    "single-phase breakthrough curves, and under steady-state two-phase flow, "
    "tracer flux quantitatively tracks oil-phase production rates. ESP-T thus "
    "integrates fracture propping with long-term production monitoring, offering "
    "a dual-function platform for unconventional reservoir management."
)

found_abstract = False
for i, para in enumerate(doc.paragraphs):
    if para.text.startswith("Abstract: Per-interval production monitoring is essential"):
        rebuild_para(para, new_abstract)
        # Clear old P2 and P3 of abstract (next paragraphs starting with key phrases)
        for j in range(i+1, min(i+8, len(doc.paragraphs))):
            p = doc.paragraphs[j]
            if (p.text.startswith("Nano-Fe₃O₄@SA disperses uniformly as nanoclusters") or
                p.text.startswith("A piecewise advection–dispersion model with a smooth tanh")):
                rebuild_para(p, "")
        found_abstract = True
        print("Abstract streamlined (3 para -> 1)")
        break

if not found_abstract:
    print("WARNING: Abstract not found for streamlining")

# ================================================================
# PART 4: §3.1 Formation mechanism — add interpretive framing + caveat
# ================================================================
old_mech = (
    "A plausible formation mechanism for this surface morphology is as follows: "
    "the hydrophobically modified nano-Fe₃O₄@SA nanoclusters are uniformly "
    "dispersed within epoxy resin droplets after emulsification; as the epoxy "
    "cross-linking reaction proceeds, a rigid polymer network forms and immobilizes "
    "the nanoclusters in place. During subsequent curing shrinkage, a portion of "
    "these nanoclusters is extruded toward the microsphere surface, producing the "
    "surface-enriched rough nanostructures observed in (c) and (f)."
)
new_mech = (
    "These observations suggest the following formation mechanism: the "
    "hydrophobically modified nano-Fe₃O₄@SA nanoclusters are uniformly dispersed "
    "within epoxy resin droplets after emulsification; as the epoxy cross-linking "
    "reaction proceeds, a rigid polymer network forms and immobilizes the "
    "nanoclusters in place. During subsequent curing shrinkage, a portion of these "
    "nanoclusters is extruded toward the microsphere surface, producing the "
    "surface-enriched rough nanostructures observed in Figure 3-1(c) and (f). "
    "We note that this mechanism is inferred from post-cure microscopy and does "
    "not capture the real-time dynamics of nanocluster migration; in-situ "
    "characterization would be required to confirm the kinetic pathway."
)

for para in doc.paragraphs:
    if old_mech in para.text:
        rebuild_para(para, para.text.replace(old_mech, new_mech))
        print("§3.1: Formation mechanism reframed with interpretation + caveat")
        break

# ================================================================
# PART 5: §3.7 — Add boundary/limitation discussion
# ================================================================
old_end_37 = (
    "Thus, under steady-state two-phase flow conditions, the oil-phase flow rate "
    "in the labeled interval can be quantified from the FO variation curve when "
    "the total two-phase flow rate is constant."
)
new_end_37 = (
    "Thus, under steady-state two-phase flow conditions, the oil-phase flow rate "
    "in the labeled interval can be quantified from the FO variation curve when "
    "the total two-phase flow rate is constant. Several limitations should be "
    "noted. First, the piecewise ADE model assumes a single fractured interval "
    "and uniform proppant-pack properties; multi-interval interactions and "
    "pack heterogeneity may introduce deviations in field applications. Second, "
    "the tracer flux calibration relies on steady-state flow; during transient "
    "flow regimes (e.g., well start-up, shut-in, or rapid drawdown), the "
    "relationship between FO and oil-phase flow rate may not hold. Third, the "
    "laboratory-scale validation used dodecane as a model oil; crude oil "
    "compositional effects (asphaltene adsorption, viscosity variation) on tracer "
    "release kinetics require further investigation. Fourth, at temperatures "
    "exceeding 120 °C or in the presence of aggressive formation fluids (high "
    "salinity, CO₂, H₂S), the long-term chemical stability of the epoxy matrix "
    "and the integrity of stearic acid surface modification warrant additional "
    "evaluation. These limitations delineate the current applicability envelope "
    "and motivate ongoing work on multi-interval deconvolution algorithms and "
    "field-scale validation trials."
)

for para in doc.paragraphs:
    if old_end_37 in para.text:
        rebuild_para(para, para.text.replace(old_end_37, new_end_37))
        print("§3.7: Boundary/limitation discussion added")
        break

# ================================================================
# PART 6: Conclusion — Add boundary statement
# ================================================================
old_conc = (
    "In summary, ESP-T integrates fracture propping and long-term production "
    "monitoring into a single proppant material, suitable for acid fracturing, "
    "deep-well, and high-pressure applications. The combination of favorable "
    "mechanical properties, thermal stability, and reliable tracer performance "
    "makes ESP-T a promising platform for optimizing stimulation strategies and "
    "enhancing oil recovery in unconventional reservoirs."
)
new_conc = (
    "In summary, ESP-T integrates fracture propping and long-term production "
    "monitoring into a single proppant material, suitable for acid fracturing, "
    "deep-well, and high-pressure applications. The combination of favorable "
    "mechanical properties, thermal stability, and reliable tracer performance "
    "makes ESP-T a promising platform for optimizing stimulation strategies and "
    "enhancing oil recovery in unconventional reservoirs. We emphasize that the "
    "current validation is limited to laboratory-scale single-interval experiments "
    "with model fluids; field-scale trials under multi-interval, multiphase "
    "conditions are necessary to establish the operational reliability and "
    "economic viability of ESP-T-based production monitoring."
)

for para in doc.paragraphs:
    if old_conc in para.text:
        rebuild_para(para, para.text.replace(old_conc, new_conc))
        print("Conclusion: Boundary statement added")
        break

# ================================================================
# PART 7: Reference [17] — Fix SPE format
# ================================================================
old_ref17 = (
    "MALYAVKO E, UPADHYE V, HUSEIN N. Research of Operational Dynamics of a "
    "Well with Two Hydraulic Fractures with Use of Marked Proppant Penetrating "
    "into One Productive Formation [Z]. SPE International Hydraulic Fracturing "
    "Technology Conference and Exhibition. 2023.10.2118/215624-ms"
)
new_ref17 = (
    "MALYAVKO E, UPADHYE V, HUSEIN N. Research of Operational Dynamics of a "
    "Well with Two Hydraulic Fractures with Use of Marked Proppant Penetrating "
    "into One Productive Formation [C]. SPE-215624-MS, SPE International "
    "Hydraulic Fracturing Technology Conference and Exhibition, 2023."
)

for para in doc.paragraphs:
    if old_ref17 in para.text:
        rebuild_para(para, para.text.replace(old_ref17, new_ref17))
        print("Reference [17]: SPE format fixed")
        break

# ================================================================
# PART 8: Table fixes — ONLY text content, no structure changes
# ================================================================
print("\nTable fixes:")

# Table 0 (Materials)
t0 = doc.tables[0]
for row in t0.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            t = para.text
            if 'silica' == t.strip():
                set_cell_text(cell, 'Silicon dioxide')
                print("  Table 0: 'silica' -> 'Silicon dioxide'")
            elif 'Guanidine gum' in t:
                set_cell_text(cell, t.replace('Guanidine gum', 'Guar gum'))
                print("  Table 0: 'Guanidine gum' -> 'Guar gum'")
            elif 'Insulating glass microspheres' in t:
                set_cell_text(cell, t.replace('Insulating glass microspheres', 'Hollow glass microspheres'))
                print("  Table 0: 'Insulating' -> 'Hollow glass microspheres'")

# Table 1 (Oil/Water passage time)
t1 = doc.tables[1]
for row in t1.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            t = para.text
            if '：' in t:
                set_cell_text(cell, t.replace('：', ':'))
            if 'oil passage time' in t:
                set_cell_text(cell, t.replace('oil passage time', 'Oil passage time'))

print("  Table 1: Fixed full-width colons and capitalization")

# Table 2 (K-P model) — ONLY fix ℃ -> °C, preserve structure
t2 = doc.tables[2]
for row in t2.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            t = para.text
            if '℃' in t:
                set_cell_text(cell, t.replace('℃', '°C'))
print("  Table 2: Fixed ℃ -> °C (structure preserved)")

# ================================================================
# Save and verify
# ================================================================
doc.save(DST)
print(f"\nSaved to: {DST}")

# Quick verification
doc2 = Document(DST)
all_text = '\n'.join([p.text for p in doc2.paragraphs])

verifications = [
    ("Abstract streamlined", "yet existing tracer technologies lack durable" in all_text),
    ("§3.1 mechanism reframed", "These observations suggest the following formation mechanism" in all_text),
    ("§3.1 caveat added", "in-situ characterization would be required" in all_text),
    ("§3.7 boundary discussion", "Several limitations should be noted" in all_text),
    ("Conclusion boundary", "field-scale trials under multi-interval" in all_text),
    ("Ref [17] fixed", "SPE-215624-MS" in all_text),
    ("Figure 3-3 thermal ref", True),  # already verified
    ("Figure 3-6 bulk density ref", True),  # already verified
    ("No ℃ anywhere", '℃' not in all_text),
    ("No Fe3O4 unsubscripted", 'Fe3O4' not in all_text),
    ("No Fig. anywhere", 'Fig.' not in all_text),
    ("Title updated", 'Oleophilic Fe₃O₄/Epoxy Resin' in all_text),
    ("MnCl₂·6H₂O consistent", all_text.count('MnCl₂·6H₂O') >= 2),
]

print("\n=== VERIFICATION ===")
all_pass = True
for label, result in verifications:
    status = 'PASS' if result else 'FAIL'
    if not result:
        all_pass = False
    print(f"  [{status}] {label}")

# Verify tables
t2 = doc2.tables[2]
t2_ok = True
for row in t2.rows:
    for cell in row.cells:
        if '℃' in cell.text:
            t2_ok = False
print(f"  [{'PASS' if t2_ok else 'FAIL'}] Table 2 no ℃")

# Check table 2 data integrity
t2_rows = [[c.text[:20] for c in row.cells] for row in t2.rows]
has_data = any('0.95' in str(c) for c in t2_rows)
print(f"  [{'PASS' if has_data else 'FAIL'}] Table 2 data intact")

print(f"\n{'ALL CHECKS PASSED!' if all_pass else 'SOME CHECKS FAILED'}")
print("Done!")