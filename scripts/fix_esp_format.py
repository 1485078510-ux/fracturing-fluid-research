"""
Clean format & content fixes based on SCI reviewer audit.
Handles: superscript refs, equation labels, temp units, Pe interpretation, etc.
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy
import re
from lxml import etree

DST = r"c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP_restructured_fixed.docx"
doc = Document(DST)

nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def rebuild_para(para, new_text):
    """Replace entire paragraph text via first run, clear others."""
    if not para.runs:
        return False
    for run in para.runs[1:]:
        run._element.getparent().remove(run._element)
    para.runs[0].text = new_text
    return True


def make_run_superscript(run):
    """Set a run's text to superscript via XML."""
    run.font.superscript = True


def split_and_superscript_refs(para):
    """
    For each run in paragraph: if it contains [N] mixed with regular text,
    split into text + superscript_ref segments.
    This is the key function for making inline refs superscript.
    """
    runs = list(para.runs)
    modified = False

    for run in runs:
        text = run.text
        # Skip runs that are already superscript
        if run.font.superscript:
            continue
        # Skip runs without reference markers
        if not re.search(r'\[\d+(?:[,\-]\d+)*\]', text):
            continue

        # Split this run's text at reference boundaries
        parts = re.split(r'(\[\d+(?:[,\-]\d+)*\])', text)
        if len(parts) <= 1:
            continue

        # We'll rebuild: first part stays in current run, rest become new runs
        run.text = parts[0]

        # Get the parent element to insert after
        parent = run._element.getparent()
        current_element = run._element

        for part in parts[1:]:
            is_ref = bool(re.match(r'\[\d+(?:[,\-]\d+)*\]', part))
            # Create new run element
            new_r = etree.SubElement(parent, qn('w:r'))
            # Copy run properties from original, then modify
            rPr = run._element.find(qn('w:rPr'))
            if rPr is not None:
                new_rPr = deepcopy(rPr)
                new_r.append(new_rPr)
            else:
                new_rPr = etree.SubElement(new_r, qn('w:rPr'))

            if is_ref:
                # Add superscript property
                vertAlign = etree.SubElement(new_rPr, qn('w:vertAlign'))
                vertAlign.set(qn('w:val'), 'superscript')

            # Add text
            new_t = etree.SubElement(new_r, qn('w:t'))
            new_t.text = part
            new_t.set(qn('xml:space'), 'preserve')

            # Insert after current element
            parent.insert(list(parent).index(current_element) + 1, new_r)
            current_element = new_r

        modified = True

    return modified


# ================================================================
# PART 1: Text-level replacements (clean and simple)
# ================================================================
print("PART 1: Text-level fixes")

fixes = [
    # Equation consistency
    ("Equation (5)", "Eq. (5)"),
    ("Equation(5)", "Eq. (5)"),

    # Temperature units in Methods
    ("at 50 C for", "at 50 °C for"),
    ("at 80 C for", "at 80 °C for"),
    ("dried at 80 C", "dried at 80 °C"),

    # Remove years from mixed citation style (keep clean "Author et al. [N]")
    ("et al. (2020) [15]", "et al. [15]"),
    ("et al. (2022) [16]", "et al. [16]"),
    ("et al. (2023) [17]", "et al. [17]"),
    ("et al. (2023) [18]", "et al. [18]"),
    ("et al. (2024) [19]", "et al. [19]"),
    ("et al. (2021) [26]", "et al. [26]"),
    ("et al. (2024) [27]", "et al. [27]"),

    # Pe interpretation fix (Reviewer 3)
    ("Pe < 1 confirms dispersion-dominated transport, consistent with the gradual",
     "Pe ≈ 1 indicates a transitional transport regime where advection and dispersion are comparable, consistent with the gradual"),

    # Reference [25]: [J] -> [C] for conference paper
    # (handled separately below since it's a long string)
]

counts = {}
for para in doc.paragraphs:
    t = para.text
    modified = False
    for old, new in fixes:
        if old in t:
            t = t.replace(old, new)
            modified = True
            counts[old[:50]] = counts.get(old[:50], 0) + 1
    if modified and para.runs:
        rebuild_para(para, t)

for k, v in counts.items():
    print(f"  '{k}...' ({v}x)")

# Reference [25] special fix
ref25_old = "as a Well Sealant [J]. SPE International Conference"
ref25_new = "as a Well Sealant [C]. SPE International Conference"
for para in doc.paragraphs:
    if ref25_old in para.text:
        rebuild_para(para, para.text.replace(ref25_old, ref25_new))
        print(f"  Reference [25] [J] -> [C]")
        break

# ================================================================
# PART 2: Superscript all reference citations
# ================================================================
print("\nPART 2: Superscript reference citations")

total_fixed = 0
for para in doc.paragraphs:
    if split_and_superscript_refs(para):
        total_fixed += 1

print(f"  Fixed {total_fixed} paragraphs with inline refs -> superscript")

# Also mark standalone ref runs (entire run is just [N]) as superscript
for para in doc.paragraphs:
    for run in para.runs:
        if re.match(r'^\s*\[\d+(?:[,\-]\d+)*\]\s*$', run.text):
            if not run.font.superscript:
                run.font.superscript = True
                total_fixed += 1

# ================================================================
# PART 3: Final sweep - any remaining temperature or µm issues
# ================================================================
print("\nPART 3: Final sweep")
sweep_count = 0
for para in doc.paragraphs:
    t = para.text
    new_t = t
    # Fix any " X C" where X is a digit (not preceded by °)
    new_t = re.sub(r'(?<![°])(\d+)\s*C\b(?!\s*[a-z])', r'\1 °C', new_t)
    # Fix "45 um" if still present
    new_t = re.sub(r'\b45\s*um\b', '45 μm', new_t)
    if new_t != t and para.runs:
        rebuild_para(para, new_t)
        sweep_count += 1
print(f"  Final sweep fixes: {sweep_count} paragraphs")

# ================================================================
# Save and verify
# ================================================================
doc.save(DST)
print(f"\nSaved: {DST}")

# Quick verification
doc2 = Document(DST)
all_text = '\n'.join([p.text for p in doc2.paragraphs])

print("\n=== VERIFICATION ===")
checks = [
    ("No 'Equation (5)'", 'Equation (5)' not in all_text),
    ("Eq. (5) present", 'Eq. (5)' in all_text),
    ("No '50 C for'", '50 C for' not in all_text),
    ("No '80 C for'", '80 C for' not in all_text),
    ("50 °C present", '50 °C' in all_text),
    ("No '(2020) [15]'", '(2020) [15]' not in all_text),
    ("'et al. [15]' present", 'et al. [15]' in all_text),
    ("Pe ≈ 1 fix", 'Pe ≈ 1 indicates a transitional' in all_text),
    ("Ref [25] [C]", '[C]. SPE International Conference' in all_text),
    ("No 'Nano - Fe' (spaces)", 'Nano - Fe' not in all_text),
]

all_ok = True
for label, result in checks:
    s = 'PASS' if result else 'FAIL'
    if not result: all_ok = False
    print(f"  [{s}] {label}")

# Check superscript status
sup_count = 0
inline_count = 0
for para in doc2.paragraphs:
    for run in para.runs:
        if re.search(r'\[\d', run.text):
            if run.font.superscript:
                sup_count += 1
            else:
                inline_count += 1

print(f"  [INFO] Superscript ref runs: {sup_count}, Inline ref runs: {inline_count}")

if all_ok:
    print("\nALL FORMAT CHECKS PASSED!")
else:
    print("\nSOME CHECKS FAILED - review above")

print("Done!")