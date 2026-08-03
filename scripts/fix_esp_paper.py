"""
Fix errors in ESP_restructured.docx — clean rewrite from original.
Avoids corrupting hyperlink elements. Applies all fixes in one pass.
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
import re

SRC = r"c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP_restructured.docx"
DST = r"c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP_restructured_fixed.docx"

doc = Document(SRC)

def fix_paragraph_text(para, replacements):
    """Apply multiple replacements to a paragraph, rebuilding runs.
    Each replacement is (old, new). Only modifies if all replacements match."""
    text = para.text
    new_text = text
    for old, new in replacements:
        if old not in new_text:
            return False
        new_text = new_text.replace(old, new)
    if new_text == text:
        return False
    # Apply to first run, clear others
    if para.runs:
        for run in para.runs[1:]:
            run.text = ''
        para.runs[0].text = new_text
        return True
    return False


counts = {}

# ================================================================
# Apply all fixes paragraph by paragraph (single pass)
# ================================================================

# Define all replacements
# Each fix: (old_string, new_string, label)

fixes = [
    # CRITICAL: Wrong figure references
    ("As depicted in Figure 3-1, thermal analysis of the proppant reveals",
     "As depicted in Figure 3-3, thermal analysis of the proppant reveals",
     "Figure 3-1→3-3 (thermal analysis)"),

    ("Figure 3-4 presents the bulk density and apparent density results",
     "Figure 3-6 presents the bulk density and apparent density results",
     "Figure 3-4→3-6 (bulk density)"),

    # Chemical formula: MnCl2 hydrate consistency (match procedure to materials)
    ("2×10⁻⁵ mol MnCl₂·4H₂O",
     "2×10⁻⁵ mol MnCl₂·6H₂O",
     "MnCl₂·4H₂O→MnCl₂·6H₂O"),

    # Grammar
    ("productivity decays at disparate rates",
     "productivity declines at varying rates",
     "decays→declines"),

    ("Owing to the effect of sample volume, the actual arrival time",
     "Due to the sample volume effect, the actual arrival time",
     "Owing to→Due to"),

    ("monitoring capability ceases once the polymer coating dissolves",
     "monitoring ceases once the polymer coating dissolves",
     "capability ceases→ceases"),

    # Standardize Korsmeyer–Peppas (en-dash)
    ("Korsmeyer-Peppas", "Korsmeyer–Peppas", "Korsmeyer en-dash"),

    # Proppant-pack → Proppant pack
    ("Proppant-pack conductivity", "Proppant pack conductivity", "Proppant-pack"),

    # Title
    ("Fabrication and Performance of ESP-T for Hydraulic Fracturing in Unconventional Oil and Gas Reservoirs",
     "An Oleophilic Fe₃O₄/Epoxy Resin Tracer Proppant for Long-Term Oil Production Monitoring in Unconventional Reservoirs",
     "Title polish"),

    # ESP – T → ESP-T (spaces around dash)
    ("ESP – T", "ESP-T", "ESP-T spacing"),
]

# Process all paragraphs
for para in doc.paragraphs:
    for old, new, label in fixes:
        if old in para.text:
            if fix_paragraph_text(para, [(old, new)]):
                counts[label] = counts.get(label, 0) + 1

# Process tables too
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for old, new, label in fixes:
                    if old in para.text:
                        if fix_paragraph_text(para, [(old, new)]):
                            counts[label] = counts.get(label, 0) + 1

print("Applied text fixes:", {k: v for k, v in counts.items() if v > 0})

# ================================================================
# Now apply regex-based fixes (℃, Fig., Figure spacing, subscripts)
# These modify first run of matching paragraphs
# ================================================================
regex_count = 0

for para in doc.paragraphs:
    old_text = para.text
    new_text = old_text

    # ℃ → °C
    new_text = new_text.replace('℃', '°C')

    # Fig.X-Y → Figure X-Y (only in figure captions, not in text)
    # Careful: 'Fig.' might appear in legitimate contexts
    new_text = re.sub(r'^Fig\.(\d+)-(\d+)', r'Figure \1-\2', new_text)

    # Figure X - Y → Figure X-Y (remove spaces around hyphen)
    new_text = re.sub(r'(Figure \d+) [–\-] (\d+)', r'\1-\2', new_text)

    # SiO2 → SiO₂ (when not already subscript)
    new_text = re.sub(r'(?<![₂])SiO2(?![₂])', 'SiO₂', new_text)

    # Fe3O4 → Fe₃O₄ (when not already subscript)
    new_text = re.sub(r'(?<![₃₄])Fe3O4(?![₃₄])', 'Fe₃O₄', new_text)

    # CO2 → CO₂
    new_text = re.sub(r'(?<![₂])CO2(?![₂])', 'CO₂', new_text)

    # nano‑Fe → nano-Fe (en-dash to regular hyphen)
    new_text = new_text.replace('nano‑Fe', 'nano-Fe')

    if new_text != old_text and para.runs:
        for run in para.runs[1:]:
            run.text = ''
        para.runs[0].text = new_text
        regex_count += 1

# Same for tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                old_text = para.text
                new_text = old_text
                new_text = new_text.replace('℃', '°C')
                new_text = re.sub(r'^Fig\.(\d+)-(\d+)', r'Figure \1-\2', new_text)
                new_text = re.sub(r'(Figure \d+) [–\-] (\d+)', r'\1-\2', new_text)
                new_text = re.sub(r'(?<![₂])SiO2(?![₂])', 'SiO₂', new_text)
                new_text = re.sub(r'(?<![₃₄])Fe3O4(?![₃₄])', 'Fe₃O₄', new_text)
                new_text = re.sub(r'(?<![₂])CO2(?![₂])', 'CO₂', new_text)
                new_text = new_text.replace('nano‑Fe', 'nano-Fe')
                if new_text != old_text and para.runs:
                    for run in para.runs[1:]:
                        run.text = ''
                    para.runs[0].text = new_text
                    regex_count += 1

print(f"Regex paragraphs fixed: {regex_count}")

# Save
doc.save(DST)
print(f"\nSaved to: {DST}")

# ================================================================
# Verification
# ================================================================
print("\n=== VERIFICATION ===")
doc2 = Document(DST)
all_text = '\n'.join([p.text for p in doc2.paragraphs])

checks = {
    'No Fe3O4 (unsubscripted)': 'Fe3O4' not in all_text,
    'No SiO2 (unsubscripted)': 'SiO2' not in all_text,
    'No ℃': '℃' not in all_text,
    'No Fig. format': 'Fig.' not in all_text,
    'No Figure X - Y spacing': not re.search(r'Figure \d+ [–\-] \d+', all_text),
    'No en-dash in nano-Fe': 'nano‑Fe' not in all_text,
    'No ESP – T': 'ESP – T' not in all_text,
    'Figure 3-3 refs thermal': 'Figure 3-3' in all_text and 'thermal analysis' in all_text,
    'Figure 3-6 refs bulk density': 'Figure 3-6' in all_text and 'bulk density' in all_text,
    'declines at varying rates': 'declines at varying rates' in all_text,
    'MnCl₂·6H₂O consistent': all_text.count('MnCl₂·6H₂O') == 2 and 'MnCl₂·4H₂O' not in all_text,
    'Korsmeyer–Peppas en-dash': 'Korsmeyer–Peppas' in all_text,
    'New title': 'Oleophilic Fe₃O₄/Epoxy Resin Tracer Proppant' in all_text,
}

for check, result in checks.items():
    status = 'PASS' if result else 'FAIL'
    print(f'  [{status}] {check}')

failures = sum(1 for v in checks.values() if not v)
print(f'\n{failures} failures out of {len(checks)} checks')
print('Done.')