"""
Structural and content fixes for ESP_restructured_fixed.docx
Addresses: Abstract, Results/Discussion mixing, boundaries, tables, ref [17]
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
import re

DST = r"c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP_restructured_fixed.docx"
doc = Document(DST)

def rebuild_para(para, new_text):
    """Replace entire paragraph text, preserving first run's formatting."""
    if para.runs:
        for run in para.runs[1:]:
            run.text = ''
        para.runs[0].text = new_text
        return True
    return False


# ================================================================
# FIX 1: Streamline Abstract (3 paragraphs -> 1 concise paragraph)
# ================================================================
new_abstract = (
    "Per-interval production monitoring is essential for evaluating stimulation "
    "effectiveness in unconventional reservoirs, yet existing tracer technologies "
    "lack durable oil-phase monitoring capability. Here we report an oleophilic "
    "tracer proppant (ESP-T) fabricated by encapsulating stearic acid-modified "
    "nano-Fe₃O₄ (nano-Fe₃O₄@SA) within an epoxy resin matrix via emulsion "
    "polymerization. ESP-T exhibits sphericity exceeding 0.9, maintains structural "
    "integrity up to 357.27 °C, and its nano-Fe₃O₄@SA incorporation raises the "
    "water contact angle from 72.3° to 104.6°, yielding a water-resistant, "
    "oil-permeable transport characteristic. Tracer release at 30–120 °C follows "
    "the Korsmeyer–Peppas model (R² > 0.90) with diffusion exponents of 0.45–0.85, "
    "indicating synergistic Fickian diffusion and Case-II relaxation. A piecewise "
    "advection–dispersion model with tanh blending achieves R² = 0.9939 for "
    "single-phase breakthrough curves, and under steady-state two-phase flow, "
    "tracer flux quantitatively tracks oil-phase production rates. ESP-T thus "
    "integrates fracture propping with long-term production monitoring, offering "
    "a dual-function platform for unconventional reservoir management."
)

for para in doc.paragraphs:
    if para.text.startswith("Per-interval production monitoring is essential for evaluating stimulation"):
        rebuild_para(para, new_abstract)
        print("FIX 1: Abstract streamlined (3 para -> 1 concise para)")
        # Clear the next two paragraphs (old P2 and P3 of abstract)
        next_paras = []
        for p in doc.paragraphs:
            if p.text.startswith("Nano-Fe₃O₄@SA disperses uniformly as nanoclusters"):
                next_paras.append(p)
            if p.text.startswith("A piecewise advection"):
                next_paras.append(p)
        for p in next_paras[:2]:
            rebuild_para(p, "")
        break

# ================================================================
# FIX 2: §3.1 — Add interpretive framing to formation mechanism
# ================================================================
old_mech = (
    "A plausible formation mechanism for this surface morphology is as follows: "
    "the hydrophobically modified nano-Fe₃O₄@SA nanoclusters are uniformly "
    "dispersed within epoxy resin droplets after emulsification; as the epoxy "
    "cross-linking reaction proceeds, a rigid polymer network forms and immobilizes "
    "the nanoclusters in place. During subsequent curing shrinkage, a portion of "
    "these nanoclusters is extruded toward the microsphere surface, producing the "
    "surface-enriched rough nanostructures observed in (c) and (f)."
)

new_mech = (
    "These observations suggest the following formation mechanism: the "
    "hydrophobically modified nano-Fe₃O₄@SA nanoclusters are uniformly dispersed "
    "within epoxy resin droplets after emulsification; as the epoxy cross-linking "
    "reaction proceeds, a rigid polymer network forms and immobilizes the "
    "nanoclusters in place. During subsequent curing shrinkage, a portion of these "
    "nanoclusters is extruded toward the microsphere surface, producing the "
    "surface-enriched rough nanostructures observed in Figure 3-1(c) and (f). "
    "We note that this mechanism is inferred from post-cure microscopy and does "
    "not capture the real-time dynamics of nanocluster migration; in-situ "
    "characterization would be required to confirm the kinetic pathway."
)

for para in doc.paragraphs:
    if old_mech in para.text:
        new_full = para.text.replace(old_mech, new_mech)
        rebuild_para(para, new_full)
        print("FIX 2: §3.1 — Added interpretive framing + boundary caveat to formation mechanism")
        break

# ================================================================
# FIX 3: Add boundary/limitation discussion at end of §3.7
# ================================================================
old_end_37 = (
    "Thus, under steady-state two-phase flow conditions, the oil-phase flow rate "
    "in the labeled interval can be quantified from the FO variation curve when "
    "the total two-phase flow rate is constant."
)

new_end_37 = (
    "Thus, under steady-state two-phase flow conditions, the oil-phase flow rate "
    "in the labeled interval can be quantified from the FO variation curve when "
    "the total two-phase flow rate is constant. Several limitations should be "
    "noted. First, the piecewise ADE model assumes a single fractured interval "
    "and uniform proppant-pack properties; multi-interval interactions and "
    "pack heterogeneity may introduce deviations in field applications. Second, "
    "the tracer flux calibration relies on steady-state flow; during transient "
    "flow regimes (e.g., well start-up, shut-in, or rapid drawdown), the "
    "relationship between FO and oil-phase flow rate may not hold. Third, the "
    "laboratory-scale validation used dodecane as a model oil; crude oil "
    "compositional effects (asphaltene adsorption, viscosity variation) on tracer "
    "release kinetics require further investigation. Fourth, at temperatures "
    "exceeding 120 °C or in the presence of aggressive formation fluids (high "
    "salinity, CO₂, H₂S), the long-term chemical stability of the epoxy matrix "
    "and the integrity of stearic acid surface modification warrant additional "
    "evaluation. These limitations delineate the current applicability envelope "
    "and motivate ongoing work on multi-interval deconvolution algorithms and "
    "field-scale validation trials."
)

for para in doc.paragraphs:
    if old_end_37 in para.text:
        new_full = para.text.replace(old_end_37, new_end_37)
        rebuild_para(para, new_full)
        print("FIX 3: §3.7 — Added boundary/limitation discussion")
        break

# ================================================================
# FIX 4: Conclusion — Add boundary statement
# ================================================================
old_conclusion_last = (
    "In summary, ESP-T integrates fracture propping and long-term production "
    "monitoring into a single proppant material, suitable for acid fracturing, "
    "deep-well, and high-pressure applications. The combination of favorable "
    "mechanical properties, thermal stability, and reliable tracer performance "
    "makes ESP-T a promising platform for optimizing stimulation strategies and "
    "enhancing oil recovery in unconventional reservoirs."
)

new_conclusion_last = (
    "In summary, ESP-T integrates fracture propping and long-term production "
    "monitoring into a single proppant material, suitable for acid fracturing, "
    "deep-well, and high-pressure applications. The combination of favorable "
    "mechanical properties, thermal stability, and reliable tracer performance "
    "makes ESP-T a promising platform for optimizing stimulation strategies and "
    "enhancing oil recovery in unconventional reservoirs. We emphasize that the "
    "current validation is limited to laboratory-scale single-interval experiments "
    "with model fluids; field-scale trials under multi-interval, multiphase "
    "conditions are necessary to establish the operational reliability and "
    "economic viability of ESP-T-based production monitoring."
)

for para in doc.paragraphs:
    if old_conclusion_last in para.text:
        new_full = para.text.replace(old_conclusion_last, new_conclusion_last)
        rebuild_para(para, new_full)
        print("FIX 4: Conclusion — Added boundary statement")
        break

# ================================================================
# FIX 5: Reference [17] — Fix SPE format
# ================================================================
old_ref17 = (
    "MALYAVKO E, UPADHYE V, HUSEIN N. Research of Operational Dynamics of a "
    "Well with Two Hydraulic Fractures with Use of Marked Proppant Penetrating "
    "into One Productive Formation [Z]. SPE International Hydraulic Fracturing "
    "Technology Conference and Exhibition. 2023.10.2118/215624-ms"
)

new_ref17 = (
    "MALYAVKO E, UPADHYE V, HUSEIN N. Research of Operational Dynamics of a "
    "Well with Two Hydraulic Fractures with Use of Marked Proppant Penetrating "
    "into One Productive Formation [C]. SPE-215624-MS, SPE International "
    "Hydraulic Fracturing Technology Conference and Exhibition, 2023."
)

for para in doc.paragraphs:
    if old_ref17 in para.text:
        new_full = para.text.replace(old_ref17, new_ref17)
        rebuild_para(para, new_full)
        print("FIX 5: Reference [17] — Fixed SPE format")
        break

# ================================================================
# FIX 6: Fix tables
# ================================================================
print("\nFIX 6: Table corrections")

# Table 0 (Materials) fixes
t0 = doc.tables[0]
fixes_t0 = {
    "silica": "Silicon dioxide",
    "Guanidine gum": "Guar gum",
    "Insulating glass microspheres": "Hollow glass microspheres",
}
for row in t0.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            for old, new in fixes_t0.items():
                if old in para.text:
                    rebuild_para(para, para.text.replace(old, new))
                    print(f"  Table 0: '{old}' -> '{new}'")

# Table 1 (Oil/Water passage time) fixes
t1 = doc.tables[1]
fixes_t1 = {
    "oil passage time": "Oil passage time",
    "：": ":",  # full-width colon -> regular colon
}
for row in t1.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            for old, new in fixes_t1.items():
                if old in para.text:
                    rebuild_para(para, para.text.replace(old, new))
                    print(f"  Table 1: '{old}' -> '{new}'")

# Table 2 (K-P model) fixes
t2 = doc.tables[2]
# Fix headers - the table has merged cells causing duplication
# Row 0: "K-P model" appears in cols 0 and 1 -> should be just col 0
# The structure should be: Model | C/C0=Kt^n | ... for each temperature
# Let's fix specific cells
try:
    # Row 0: Make col 1 empty (duplicate header)
    if "K-P model" in t2.rows[0].cells[1].text:
        for para in t2.rows[0].cells[1].paragraphs:
            rebuild_para(para, "")
        print("  Table 2: Removed duplicate 'K-P model' in header row")

    # Row 1: Make col 1 empty (duplicate "Fit parameters")
    if "Fit parameters" in t2.rows[1].cells[1].text:
        for para in t2.rows[1].cells[1].paragraphs:
            rebuild_para(para, "")
        print("  Table 2: Removed duplicate 'Fit parameters'")

    # Row 2: Make col 1 empty (duplicate "30°C")
    if "30°C" in t2.rows[2].cells[1].text:
        for para in t2.rows[2].cells[1].paragraphs:
            rebuild_para(para, "")
        print("  Table 2: Removed duplicate '30°C'")

    # Row 3 (R2): Make col 1 empty (duplicate "0.95489")
    if "0.95489" in t2.rows[3].cells[1].text:
        for para in t2.rows[3].cells[1].paragraphs:
            rebuild_para(para, "")
        print("  Table 2: Removed duplicate R2 value")

    # Row 4 (K): Make col 1 empty (duplicate "0.05538")
    if "0.05538" in t2.rows[4].cells[1].text:
        for para in t2.rows[4].cells[1].paragraphs:
            rebuild_para(para, "")
        print("  Table 2: Removed duplicate K value")

    # Row 5 (n): Make col 1 empty (duplicate "0.59827")
    if "0.59827" in t2.rows[5].cells[1].text:
        for para in t2.rows[5].cells[1].paragraphs:
            rebuild_para(para, "")
        print("  Table 2: Removed duplicate n value")

    # Clean up header row 0: "release temperature" should only appear once per column
    # This needs cell-level fix - let's adjust the header structure
    # The ideal structure for row 0 (merged header):
    # col 0: "Model" | col 1: empty | col 2: "30°C" | col 3: "60°C" | col 4: "90°C" | col 5: "120°C"
    for para in t2.rows[0].cells[0].paragraphs:
        if "K-P model" in para.text:
            rebuild_para(para, "Model")
    for para in t2.rows[0].cells[2].paragraphs:
        if "release temperature" in para.text or "C/C0" in para.text:
            rebuild_para(para, "30 °C")
    for para in t2.rows[0].cells[3].paragraphs:
        if "release temperature" in para.text or "C/C0" in para.text:
            rebuild_para(para, "60 °C")
    for para in t2.rows[0].cells[4].paragraphs:
        if "release temperature" in para.text or "C/C0" in para.text:
            rebuild_para(para, "90 °C")
    for para in t2.rows[0].cells[5].paragraphs:
        if "release temperature" in para.text or "C/C0" in para.text:
            rebuild_para(para, "120 °C")

    # Row 1: "Fit parameters" label
    for para in t2.rows[1].cells[0].paragraphs:
        if "Fit parameters" in para.text:
            rebuild_para(para, "Parameter")

    print("  Table 2: Header structure cleaned up")
except Exception as e:
    print(f"  Table 2: Warning - {e}")

# ================================================================
# FIX 7: §3.7 model description - fix missing equation context
# Line 134 has incomplete text: "is the mean flow velocity and D = alpha * v"
# This appears to be an equation variable definition with missing variable name
# ================================================================
old_ade = (
    "The interpretation model is derived from the one-dimensional "
    "advection-dispersion equation (ADE):  is the mean flow velocity and "
    "D = alpha * v is the longitudinal dispersion coefficient."
)

new_ade = (
    "The interpretation model is derived from the one-dimensional "
    "advection-dispersion equation (ADE), where v is the mean flow velocity and "
    "D = α·v is the longitudinal dispersion coefficient."
)

for para in doc.paragraphs:
    if old_ade in para.text:
        new_full = para.text.replace(old_ade, new_ade)
        rebuild_para(para, new_full)
        print("FIX 7: §3.7 — Fixed incomplete ADE variable definition")
        break

# ================================================================
# FIX 8: Add "Discussion" framing label at start of interpretation-heavy sections
# ================================================================
# The section is called "Results and Discussion" — add clearer transition signals
# In §3.1 high-mag paragraph, change "This observation verifies" to softer language
old_verify = (
    "This observation verifies that the long alkyl chains of stearic acid form "
    "strong physical entanglement and hydrophobic interactions with epoxy resin "
    "molecular chains, yielding robust interfacial adhesion between the "
    "nanofillers and the epoxy matrix"
)

new_verify = (
    "This observation indicates that the long alkyl chains of stearic acid form "
    "strong physical entanglement and hydrophobic interactions with epoxy resin "
    "molecular chains, yielding robust interfacial adhesion between the "
    "nanofillers and the epoxy matrix"
)

for para in doc.paragraphs:
    if old_verify in para.text:
        new_full = para.text.replace(old_verify, new_verify)
        rebuild_para(para, new_full)
        print("FIX 8: §3.1 — 'verifies' -> 'indicates' (appropriate hedging for microscopy observation)")
        break

# ================================================================
# FIX 9: Standardize "45 um" -> "45 μm" in methods
# ================================================================
for para in doc.paragraphs:
    if "45 um" in para.text:
        rebuild_para(para, para.text.replace("45 um", "45 μm"))
        print("FIX 9: '45 um' -> '45 μm'")

# Save
doc.save(DST)
print(f"\nAll structural fixes applied. Saved to: {DST}")
print("Done!")