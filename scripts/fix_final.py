#!/usr/bin/env python3
"""修复检查发现的3个问题：摘要超字数、唯一用语、最具创新性残留"""
from docx import Document
from docx.oxml.ns import qn

input_path = r"c:\Users\郝\Desktop\claude\荧光压裂液\专利申请文件_正式格式_修改版.docx"

doc = Document(input_path)

# ================================================================
# 修复1: 摘要精简到300字以内
# ================================================================
old_abstract = (
    "本发明公开了一种用于压裂裂缝荧光示踪的改性稀土铝酸盐荧光粉、其制备方法、含该荧光粉的压裂液体系"
    "及其应用方法。所述改性荧光粉以SrAl₂O₄:Eu²⁺,Dy³⁺等稀土铝酸盐长余辉材料为基体，表面依次包覆"
    "硅烷偶联剂化学键合内层和聚乙二醇（PEG）物理屏蔽外层，并辅以螯合剂与非离子表面活性剂构成的协同"
    "分散助剂体系；其制备采用先在醇-水混合溶剂中进行硅烷偶联剂化学接枝、再在水相中进行PEG物理包覆的"
    "两步法。所述荧光压裂液体系以羟丙基胍胶为稠化剂、有机硼为交联剂、过硫酸铵为破胶剂，采用母液预配"
    "与在线稀释相结合的方式配制。压裂施工中，注入阶段PEG外层保障荧光粉的分散稳定性；关井破胶阶段"
    "PEG在破胶剂氧化环境中脱附降解，暴露内层活性氨基官能团并与砂岩壁面硅羟基通过静电吸引、氢键和"
    "化学缩合实现牢固锚定；返排后取心，紫外照射即可直接观察裂缝壁面荧光分布。本发明以压裂施工工序"
    "时序驱动功能切换，无需额外触发剂，可为压裂裂缝提供可实物验证的持久荧光标记。"
)

new_abstract = (
    "本发明公开了一种用于压裂裂缝荧光示踪的改性稀土铝酸盐荧光粉、其制备方法、含该荧光粉的压裂液体系"
    "及其应用方法。所述改性荧光粉以SrAl₂O₄:Eu²⁺,Dy³⁺等稀土铝酸盐长余辉材料为基体，表面包覆"
    "硅烷偶联剂化学键合内层和聚乙二醇物理屏蔽外层；制备采用先化学接枝硅烷偶联剂、再物理包覆聚乙二醇"
    "的两步法。所述荧光压裂液体系以羟丙基胍胶为稠化剂、有机硼为交联剂、过硫酸铵为破胶剂，采用母液预配"
    "与在线稀释方式配制。注入阶段PEG外层保障分散稳定性；关井破胶阶段PEG脱附暴露活性氨基并与砂岩壁面"
    "锚定；返排后取心紫外照射即可观察裂缝荧光分布。本发明以压裂施工工序时序驱动功能切换，无需额外"
    "触发剂，可为压裂裂缝提供可实物验证的持久荧光标记。"
)

print(f"原摘要字数: {len(old_abstract)}")
print(f"新摘要字数: {len(new_abstract)}")

for para in doc.paragraphs:
    if old_abstract in para.text:
        first_run = para.runs[0] if para.runs else None
        if first_run:
            for i in range(len(para.runs) - 1, 0, -1):
                para.runs[i]._element.getparent().remove(para.runs[i]._element)
            first_run.text = new_abstract
            print("摘要已精简")
        break
else:
    # 模糊匹配
    for para in doc.paragraphs:
        if "一种用于压裂裂缝荧光示踪的改性稀土铝酸盐荧光粉" in para.text and "摘要附图" not in para.text and len(para.text) > 300:
            first_run = para.runs[0] if para.runs else None
            if first_run:
                for i in range(len(para.runs) - 1, 0, -1):
                    para.runs[i]._element.getparent().remove(para.runs[i]._element)
                first_run.text = new_abstract
                print("摘要已精简(模糊匹配)")
            break

# ================================================================
# 修复2: 去除"唯一的"（绝对化用语）
# ================================================================
for para in doc.paragraphs:
    if "唯一的化学触发条件" in para.text:
        new_text = para.text.replace("唯一的化学触发条件", "化学触发条件")
        first_run = para.runs[0] if para.runs else None
        if first_run:
            for i in range(len(para.runs) - 1, 0, -1):
                para.runs[i]._element.getparent().remove(para.runs[i]._element)
            first_run.text = new_text
            print("已删除'唯一的'")
        break

# ================================================================
# 修复3: 去除"最具创新性"残留
# ================================================================
for para in doc.paragraphs:
    if "最具创新性的设计在于" in para.text:
        new_text = para.text.replace(
            "本发明最具创新性的设计在于将",
            "本发明将"
        )
        first_run = para.runs[0] if para.runs else None
        if first_run:
            for i in range(len(para.runs) - 1, 0, -1):
                para.runs[i]._element.getparent().remove(para.runs[i]._element)
            first_run.text = new_text
            print("已修改'最具创新性'")
        break
else:
    # 也可能之前已被修改为"关键设计"，检查是否还有"最具创新"
    for para in doc.paragraphs:
        if "最具创新" in para.text:
            print(f"残留'最具创新'于: {para.text[max(0, para.text.find('最具创新')-20):para.text.find('最具创新')+40]}")
            new_text = para.text.replace("最具创新性的设计在于将", "本发明将").replace("最具创新", "关键")
            first_run = para.runs[0] if para.runs else None
            if first_run:
                for i in range(len(para.runs) - 1, 0, -1):
                    para.runs[i]._element.getparent().remove(para.runs[i]._element)
                first_run.text = new_text
                print("已模糊修复'最具创新'")
            break

doc.save(input_path)
print("\n全部修复完成！")