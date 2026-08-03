"""Revert numbered sub-headings to unnumbered bold in Section 3.3."""
from docx import Document

INPUT = r"c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP-T_投稿文件\ESP-T_Final_4-revised.docx"
doc = Document(INPUT)

replacements = {
    '3.3.1 Model formulation': 'Model formulation',
    '3.3.2 Model selection': 'Model selection',
    '3.3.3 Physical consistency': 'Physical consistency',
    '3.3.4 Comparison with time-of-arrival methods': 'Comparison with time-of-arrival methods',
    '3.3.5 Signal decomposition and robustness': 'Signal decomposition and robustness',
}

for p in doc.paragraphs:
    text = p.text.strip()
    if text in replacements:
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        for child in list(p._element):
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag in ('r', 'drawing', 'pict'):
                p._element.remove(child)
        run = p.add_run(replacements[text])
        run.bold = True
        print(f'  Changed: "{text}" -> "{replacements[text]}"')

doc.save(INPUT)
print('Done!')
