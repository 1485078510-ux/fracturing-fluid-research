#!/usr/bin/env python3
"""Fix Introduction: merge fragmented paragraphs back to 5."""
from docx import Document
import time
doc = Document('四氧化三铁环氧树脂拟合/ESP-T_final_1781616952.docx')

# Current: [7][9][11][13][15][17][19][21][23] = 9 paragraphs
# Target: [7][9][11][NEW13][NEW15] = 5 paragraphs
# Merge: [13]+[15] -> new modeling gap
# Merge: [17]+[19] -> new material gap
# Merge: [21]+[23] -> new this work

# Read current text
p13 = doc.paragraphs[13].text.strip()
p15 = doc.paragraphs[15].text.strip()
p17 = doc.paragraphs[17].text.strip()
p19 = doc.paragraphs[19].text.strip()
p21 = doc.paragraphs[21].text.strip()
p23 = doc.paragraphs[23].text.strip()

# Merge [13]+[15]: modeling gap + building blocks
new13 = p13 + "\n\n" + p15

# Merge [17]+[19]: material gap + epoxy
# [17] starts with "Compounding the modeling gap, the materials used..."
# [19] starts with "Li et al. [33] and Wei et al. [34] showed..."
new17 = p17 + "\n\n" + p19

# Merge [21]+[23]: this work + validation
new21 = p21 + "\n\n" + p23

# Write merged paragraphs
for i, text in [(13, new13), (17, new17), (21, new21)]:
    p = doc.paragraphs[i]
    for r in p.runs: r.text = ''
    p.runs[0].text = text

# Clear old split paragraphs [15], [19], [23]
for i in [15, 19, 23]:
    for r in doc.paragraphs[i].runs: r.text = ''

out = f'四氧化三铁环氧树脂拟合/ESP-T_final_{int(time.time())}.docx'
doc.save(out)
print(f'Saved: {out}')
print('Introduction merged back to 5 paragraphs')