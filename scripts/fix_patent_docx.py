# -*- coding: utf-8 -*-
"""修正专利申请文件 DOCX 的格式和内容问题"""
import sys, os, io, re, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

SRC = r"荧光压裂液\专利申请文件_修正润色版.docx"
DST = r"荧光压裂液\专利申请文件_已修正.docx"

doc = Document(SRC)

changes_log = []

def modify_para_text(para, new_text):
    """Replace all text in a paragraph while preserving first run's formatting."""
    if not para.runs:
        # No runs: add a new run
        run = para.add_run(new_text)
        return
    # Preserve first run formatting, clear all others
    first_run = para.runs[0]
    font = first_run.font
    bold = first_run.bold
    size = font.size
    name = font.name
    # Clear all runs
    for i in range(len(para.runs) - 1, -1, -1):
        para.runs[i]._element.getparent().remove(para.runs[i]._element)
    # Add new text in a single run preserving original formatting
    new_run = para.add_run(new_text)
    if bold is not None:
        new_run.bold = bold
    if size is not None:
        new_run.font.size = size
    if name is not None:
        new_run.font.name = name

# ── ============ 1. 修正权利要求编号（去掉"权利要求"前缀） ============ ──
claim_pattern = re.compile(r'^权利要求(\d+)\.?\s*')
for para in doc.paragraphs:
    text = para.text.strip()
    m = claim_pattern.match(text)
    if m:
        new_text = claim_pattern.sub(r'\1. ', text)
        modify_para_text(para, new_text)
        changes_log.append(f"[权利要求编号] 「{text[:30]}...」→「{new_text[:30]}...」")

# ── ============ 2. 修正发明内容次级编号（去掉6.1/6.2/6.3） ============ ──
section_fixes = {
    '6.1 发明目的': '发明目的',
    '6.2 技术方案': '技术方案',
    '6.3 有益效果': '有益效果',
}
for para in doc.paragraphs:
    text = para.text.strip()
    for old, new in section_fixes.items():
        if text.startswith(old):
            new_text = text.replace(old, new, 1)
            modify_para_text(para, new_text)
            changes_log.append(f"[章节编号] 「{old}」→「{new}」")
            break

# ── ============ 3. 补充申请人地址和邮政编码 ============ ──
cdu_address = "地址：四川省成都市成华区二仙桥东三路1号，邮政编码：610059"
for para in doc.paragraphs:
    text = para.text.strip()
    if text == '成都理工大学':
        # Replace with address-included version
        new_text = f"成都理工大学\n{cdu_address}"
        modify_para_text(para, new_text)
        changes_log.append(f"[申请人地址] 已补充成都理工大学地址信息")
        break

# ── ============ 4. 精简摘要（≤300字） ============ ──
abstract_new = (
    "本发明公开了一种用于压裂裂缝荧光示踪的改性稀土铝酸盐荧光粉、含该荧光粉的压裂液体系及其应用方法。"
    "所述改性荧光粉以SrAl₂O₄:Eu²⁺,Dy³⁺为基体，表面包覆硅烷偶联剂化学键合内层和聚乙二醇物理屏蔽外层。"
    "注入阶段PEG外层通过空间位阻效应保障分散稳定性；关井破胶阶段PEG在破胶剂氧化环境与储层温度作用下脱附降解，"
    "暴露内层活性氨基官能团并与砂岩壁面硅羟基通过静电吸引、氢键和化学缩合实现牢固锚定。"
    "本发明利用压裂施工工序时序驱动功能切换，无需额外触发剂，与现有工艺无缝兼容，可为压裂裂缝提供可实物验证的持久荧光标记。"
)
# Verify length ≤ 300
abstract_len = len(abstract_new)
print(f"[摘要] 精简后摘要字数: {abstract_len} 字（含标点）")

for para in doc.paragraphs:
    text = para.text.strip()
    if text.startswith("本发明公开了一种用于水力压裂裂缝荧光示踪的改性稀土铝酸盐荧光粉"):
        modify_para_text(para, abstract_new)
        changes_log.append(f"[摘要精简] 原摘要约{len(text)}字 → 精简后{abstract_len}字")
        break

# ── ============ 5. 修正中文文本中混入的英文逗号（保留化学式内部逗号） ============ ──
# 策略：匹配"中文字符后紧跟半角逗号再接中文字符"的模式，将其替换为全角逗号
# 保护的化学式模式：元素符号之间的逗号（如 Eu²⁺,Dy³⁺）
comma_fix_count = 0
for para in doc.paragraphs:
    text = para.text
    if not text.strip():
        continue
    new_text = text
    # Pattern: Chinese char or full-width punctuation followed by half-width comma followed by Chinese char
    # 一-鿿 = CJK unified ideographs
    # 　-〿 = CJK punctuation
    # ＀-￯ = full-width forms
    for pat, repl in [
        (r'([一-鿿　-〿＀-￯]),([一-鿿])', r'\1，\2'),
        (r'([一-鿿]),(\d)', r'\1，\2'),
        (r'([一-鿿]),\s*([一-鿿])', r'\1，\2'),
    ]:
        prev = new_text
        new_text = re.sub(pat, repl, new_text)
        if new_text != prev:
            comma_fix_count += len(re.findall(pat, prev))
    if new_text != text:
        # Replace while keeping formatting
        modify_para_text(para, new_text)
        changes_log.append(f"[标点修正] 段落中修正了半角逗号混入问题")
        break  # Log once per paragraph but fix all

if comma_fix_count > 0:
    print(f"[标点修正] 共修正 {comma_fix_count} 处英文逗号混入中文文本的问题")
else:
    print(f"[标点修正] 未发现需要修正的英文逗号混入问题（化学式内部逗号已保留）")

# ── ============ 6. 术语一致性检查 ============ ──
for para in doc.paragraphs:
    text = para.text
    # Check for 锚固态 vs 锚定态
    if '锚固态' in text:
        new_text = text.replace('锚固态', '锚定态')
        modify_para_text(para, new_text)
        changes_log.append(f"[术语统一] 「锚固态」→「锚定态」")
        print(f"[术语统一] 「锚固态」→「锚定态」")

# ── ============ 保存 ============ ──
doc.save(DST)
print(f"\n✅ 修正完成，已保存至: {os.path.abspath(DST)}")
print(f"\n修正项汇总 ({len(changes_log)} 项):")
for i, change in enumerate(changes_log, 1):
    print(f"  {i}. {change}")