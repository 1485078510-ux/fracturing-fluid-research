#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Fix issues in thesis DOCX: delete uncited refs, fix typos, fix citations."""
import sys
import re
from docx import Document

# Force UTF-8 output
sys.stdout.reconfigure(encoding='utf-8', errors='replace')

DOC_PATH = r'荧光压裂液/论文初稿_v2.docx'

doc = Document(DOC_PATH)

# ============================================================
# HELPER: Replace text within a paragraph (handles runs)
# ============================================================
def replace_in_para(para, old, new):
    """Replace old text with new in paragraph, handling run boundaries."""
    full = para.text
    if old not in full:
        return False

    # Try simple within-run replacement first
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True

    # Text spans runs - rebuild into first run
    pos = full.find(old)
    end = pos + len(old)
    new_full = full[:pos] + new + full[end:]

    # Clear all runs, put rebuilt text in first run
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = new_full
    else:
        para.add_run(new_full)
    return True


def find_para_containing(doc, snippet):
    """Find first paragraph containing snippet."""
    for i, para in enumerate(doc.paragraphs):
        if snippet in para.text:
            return i, para
    return None, None


# ============================================================
# FIX 1: Delete uncited reference paragraphs [54]-[95] and [98]-[100]
# ============================================================
paras_to_delete = []
ref_number_pattern = re.compile(r'^\s*\[(\d+)\]\s')

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    m = ref_number_pattern.match(text)
    if m:
        n = int(m.group(1))
        if (54 <= n <= 95) or (98 <= n <= 100):
            paras_to_delete.append(i)

print(f'Deleting {len(paras_to_delete)} uncited reference paragraphs...')

# Delete in reverse order to preserve earlier indices
for i in reversed(paras_to_delete):
    p_elem = doc.paragraphs[i]._element
    p_elem.getparent().remove(p_elem)

print('Done deleting references.')

# ============================================================
# FIX 2: Renumber [96]->[54] and [97]->[55] in reference list
# ============================================================
for para in doc.paragraphs:
    text = para.text.strip()
    if text.startswith('[96] '):
        replace_in_para(para, '[96] ', '[54] ')
        print('Renumbered [96] -> [54]')
    elif text.startswith('[97] '):
        replace_in_para(para, '[97] ', '[55] ')
        print('Renumbered [97] -> [55]')

# ============================================================
# FIX 3: English abstract - "主线" -> "framework"
# ============================================================
idx, para = find_para_containing(doc, '主线')
if para and 'Following the' in para.text:
    replace_in_para(para, '主线', 'framework')
    print('Fixed English abstract: 主线 -> framework')
else:
    print('WARNING: Could not find English abstract 主线')

# ============================================================
# FIX 4: Citation [24] -> [14] for Karrenbach 2017
# ============================================================
for para in doc.paragraphs:
    if 'Karrenbach' in para.text and '2017' in para.text and '[24]' in para.text:
        replace_in_para(para, '[24]', '[14]')
        print('Fixed citation [24] -> [14] for Karrenbach 2017')
        break

# ============================================================
# FIX 5: Section 3.6 garbled sentence
# ============================================================
idx, para = find_para_containing(doc, '剥离出来本研究的荧光粉属于')
if para:
    replace_in_para(para, '剥离出来本研究的荧光粉属于', '剥离出来。本研究的荧光粉属于')
    print('Fixed garbled sentence in Section 3.6')
else:
    print('WARNING: Could not find garbled sentence')

# ============================================================
# FIX 6: Reference [24] - fix volume/issue and add full title
# ============================================================
for para in doc.paragraphs:
    text = para.text.strip()
    if text.startswith('[24] ') and 'Karrenbach' in text and '2019' in text:
        # Build corrected text
        new_text = text.replace(
            'Fiber-Optic Distributed Acoustic Sensing of Microseismicity. Geophysics, 84(6).',
            'Fiber-Optic Distributed Acoustic Sensing of Microseismicity, Strain and Temperature During Hydraulic Fracturing. Geophysics, 84(1), D11-D23.'
        )
        for run in para.runs:
            run.text = ''
        if para.runs:
            para.runs[0].text = new_text
        else:
            para.add_run(new_text)
        print('Fixed reference [24] info (84(6) -> 84(1), D11-D23)')
        break

# ============================================================
# FIX 7: Add bracket citations for industry standards
# [54] = SY/T 6376-2008, [55] = SY/T 5107-2016
# ============================================================
fix_count = 0
for para in doc.paragraphs:
    for run in para.runs:
        rt = run.text
        if 'SY/T 6376-2008' in rt and '[54]' not in rt:
            run.text = rt.replace('SY/T 6376-2008', 'SY/T 6376-2008[54]')
            fix_count += 1
        if 'SY/T 5107-2016' in rt and '[55]' not in rt:
            run.text = rt.replace('SY/T 5107-2016', 'SY/T 5107-2016[55]')
            fix_count += 1

print(f'Added {fix_count} bracket citations for standards.')

# ============================================================
# FIX 8: Cost citation [53] -> industry estimate
# ============================================================
idx, para = find_para_containing(doc, '3万~8万美元')
if para and '[53]' in para.text:
    replace_in_para(para, '[53]', '（行业经验估算）')
    print('Fixed cost citation: [53] -> (行业经验估算)')
else:
    print('WARNING: Could not find cost citation paragraph')

# ============================================================
# SAVE
# ============================================================
doc.save(DOC_PATH)
print('\n=== All fixes applied and saved to', DOC_PATH, '===')