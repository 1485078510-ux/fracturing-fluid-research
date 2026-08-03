# -*- coding: utf-8 -*-
"""Safe reference superscript + chemical subscript formatting."""
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
import re

DST = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP_formatted.docx'
doc = Document(DST)

# Image paragraphs — never touch
IMG = set()
for i, p in enumerate(doc.paragraphs):
    for run in p.runs:
        if (run._element.findall('.//'+qn('w:drawing')) or
            run._element.findall('.//'+qn('a:blip'))):
            IMG.add(i); break

# Patterns for chemical subscripts: element followed by digits
# Match things like Fe3O4, SiO2, H2O, CO2, FeCl3, MnCl2, FeCl2, etc.
# But NOT match things like "pH 10", "50 ℃", "3.3%", etc.
CHEM_PATTERN = re.compile(
    r'\b('
    r'Fe\d+O\d+(?:@\w+)?'           # Fe3O4, Fe3O4@SA
    r'|FeCl\d+'                      # FeCl3, FeCl2
    r'|MnCl\d+'                      # MnCl2
    r'|SiO\d*'                       # SiO2, SiO
    r'|CO\d*'                        # CO2, CO
    r'|H\d*O\d*'                     # H2O
    r'|TiO\d*'                       # TiO2
    r'|Al\d*O\d*'                    # Al2O3
    r'|C\d*H\d*'                     # C17H35COOH etc (handled separately)
    r')\b'
)
# Reference patterns
REF_PATTERN = re.compile(r'\[\d+(?:[-,]\d+)*\]')

# Longer chemical formula: C17H35COOH
LONG_CHEM = re.compile(r'\b(C\d+H\d+COOH)\b')

count_chem = 0
count_ref = 0

for idx, para in enumerate(doc.paragraphs):
    if idx in IMG:
        continue
    if not para.text.strip():
        continue

    # Collect text from all runs
    text = para.text
    if not text:
        continue

    # Check if this paragraph has anything to fix
    has_chem = CHEM_PATTERN.search(text) or LONG_CHEM.search(text)
    has_ref = REF_PATTERN.search(text) and 'EndNote' not in para.style.name

    if not has_chem and not has_ref:
        continue

    # Get original formatting from first run
    orig_font = None
    orig_size = None
    orig_bold = None
    for run in para.runs:
        if run.text.strip():
            orig_font = run.font.name
            orig_size = run.font.size
            orig_bold = run.font.bold
            break

    # Strategy: rebuild the paragraph text with proper formatting
    # Parse the text into segments: normal, subscript, superscript
    segments = []  # list of (text, type) where type is 'normal', 'sub', 'sup'

    # Find all positions of chemical formulas and references
    all_matches = []
    for m in CHEM_PATTERN.finditer(text):
        all_matches.append((m.start(), m.end(), 'chem'))
    for m in LONG_CHEM.finditer(text):
        all_matches.append((m.start(), m.end(), 'chem'))
    for m in REF_PATTERN.finditer(text):
        all_matches.append((m.start(), m.end(), 'ref'))

    if not all_matches:
        continue

    # Sort and deduplicate matches
    all_matches.sort(key=lambda x: (x[0], -x[1]))
    # Remove overlapping matches
    filtered = []
    last_end = 0
    for start, end, mtype in all_matches:
        if start >= last_end:
            filtered.append((start, end, mtype))
            last_end = end
    all_matches = filtered

    # Build segments
    pos = 0
    for start, end, mtype in all_matches:
        if start > pos:
            segments.append((text[pos:start], 'normal'))
        if mtype == 'chem':
            segments.append((text[start:end], 'chem'))
        else:
            segments.append((text[start:end], 'ref'))
        pos = end
    if pos < len(text):
        segments.append((text[pos:], 'normal'))

    # Split chemical formulas into normal+subscript parts
    expanded_segments = []
    for seg_text, seg_type in segments:
        if seg_type == 'chem':
            # Parse chemical formula: letters normal, digits subscript
            parts = re.split(r'(\d+)', seg_text)
            for part in parts:
                if not part:
                    continue
                if part.isdigit():
                    expanded_segments.append((part, 'sub'))
                else:
                    expanded_segments.append((part, 'normal'))
        else:
            expanded_segments.append((seg_text, seg_type))

    # Clear paragraph and rebuild
    # Remove all existing runs
    for run in list(para.runs):
        run._element.getparent().remove(run._element)

    # Add new formatted runs
    for seg_text, seg_type in expanded_segments:
        if not seg_text:
            continue
        new_run = para.add_run(seg_text)
        # Copy original formatting
        if orig_font:
            new_run.font.name = orig_font
        if orig_size:
            new_run.font.size = orig_size
        if orig_bold is not None:
            new_run.font.bold = orig_bold

        if seg_type == 'sub':
            new_run.font.subscript = True
            count_chem += 1
        elif seg_type == 'ref':
            new_run.font.superscript = True
            count_ref += 1

print(f"Formatted: {count_chem} chemical subscripts, {count_ref} reference superscripts")
doc.save(DST)
print(f"Saved: {DST}")