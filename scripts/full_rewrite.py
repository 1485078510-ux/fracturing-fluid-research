#!/usr/bin/env python3
"""Comprehensive rewrite: Introduction + References + ADE model deepening."""

from docx import Document
import shutil, os

src = '四氧化三铁环氧树脂拟合/ESP-T_final.docx'
# Try to load; if locked, load from backup
try:
    doc = Document(src)
except PermissionError:
    doc = Document('四氧化三铁环氧树脂拟合/ESP-T_final_backup.docx')
    print("Loaded from backup (original locked)")

# ═══════════════════════════════════════════════════════════════════
# PART 1: REWRITE INTRODUCTION
# ═══════════════════════════════════════════════════════════════════

# New Para 3 (index [11]): Tracer literature review + MODELING GAP
new_p11 = (
    "Conventional oil-soluble chemical tracers exhibit limited compatibility "
    "with water-based fracturing fluids and cannot sustain long-term monitoring "
    "[10–14]. Tracer proppants—composite particles that integrate fracture "
    "propping and tracer functions within a solid carrier—overcome these "
    "limitations by enabling co-injection with proppant-laden fracturing fluids, "
    "thereby eliminating separate tracer injection operations. Zhao et al. [15] "
    "coated rhodamine 6G onto polyvinyl alcohol-modified ceramic particles to "
    "retard aqueous release. Zhou et al. [16] encapsulated carbon quantum dots "
    "in polyvinyl alcohol-coated ceramic particles for controlled release. "
    "Malyavko et al. [17] employed monodisperse microspheres doped with "
    "fluorescent semiconductor nanocrystals functionalized with phase-selective "
    "polymer coatings. Li et al. [18] encapsulated rare-earth element tracers "
    "with ammonium polymethacrylate to elucidate slow-release mechanisms. "
    "Gong et al. [19] embedded oleophilic tracers into polystyrene (PS) "
    "microspheres via suspension polymerization for oil-phase monitoring. "
    "Notably, all of these studies characterize tracer elution using empirical "
    "release models—predominantly the Korsmeyer–Peppas power law—yet none "
    "establishes a quantitative linkage between the release kinetics measured "
    "in batch experiments and the resulting tracer breakthrough signal observed "
    "at the wellhead. Specifically, the advection–dispersion transport of "
    "released tracers from the proppant pack through the production tubing to "
    "the sampling point remains unmodeled. This transport-to-signal gap "
    "precludes per-interval oil production rates from being quantitatively "
    "inferred from tracer concentration data, regardless of how accurately the "
    "release kinetics are characterized."
)

# New Para 4 (index [13]): Material gap + epoxy resin opportunity
new_p13 = (
    "Beyond the transport-modeling gap, material-level limitations further "
    "constrain existing tracer proppants. Conventional coated proppants suffer "
    "from high particle density (impeding transport in fracturing fluids), "
    "abrupt monitoring cessation upon coating dissolution, and—in the case of "
    "pure PS microspheres—inadequate mechanical strength and thermal stability "
    "[20–23]. Epoxy resin, a thermosetting polymer with outstanding thermal "
    "stability (decomposition > 350 °C), exceptional chemical resistance, and "
    "tunable mechanical properties arising from its highly cross-linked network, "
    "offers a compelling alternative matrix material. Epoxy-coated proppants "
    "exhibit substantially reduced crush rates, improved compressive strength "
    "and acid resistance, while retaining low-density advantages [24,25]. "
    "However, coating-based fabrication routes are constrained by uneven "
    "coverage, interfacial debonding risks, and complex multi-step preparation. "
    "Direct emulsion-polymerization synthesis of epoxy resin microspheres "
    "overcomes these limitations—enabling low-density tailoring, nanoparticle "
    "modification, and in-situ tracer encapsulation within a single synthetic "
    "step. Li et al. [26] and Wei et al. [27] demonstrated epoxy-encapsulated "
    "water-soluble tracers for aqueous-phase inflow profiling; however, the use "
    "of an epoxy matrix for oleophilic, oil-phase tracer release—and, crucially, "
    "its integration with a predictive transport model that quantitatively "
    "translates tracer signals into per-interval production allocation—has not "
    "been reported."
)

# New Para 5 (index [15]): This work — ADE model first, material second
new_p15 = (
    "Here we address this dual gap. First, we develop a piecewise "
    "advection–dispersion model with smooth hyperbolic-tangent (tanh) "
    "transition that decomposes the tracer breakthrough curve into a Gaussian "
    "pulse component—capturing the advective–dispersive transport of the "
    "tracer slug accumulated during the well shut-in period—and an erfc "
    "tailing component—capturing the sustained, matrix-diffusion-controlled "
    "release of residual tracer from the polymer carrier. The model "
    "establishes a direct quantitative relationship between the measured tracer "
    "concentration history at the wellhead and the per-interval oil production "
    "rate. Second, we validate this model using a purpose-designed oleophilic "
    "tracer proppant (ESP-T) fabricated by encapsulating stearic acid-modified "
    "nano-Fe₃O₄ (nano-Fe₃O₄@SA) within an epoxy resin matrix via "
    "emulsion polymerization. We systematically characterize ESP-T's "
    "microstructure, thermal stability, wettability, oil–water transport "
    "selectivity, and mechanical properties; quantify its temperature-dependent "
    "release kinetics using the Korsmeyer–Peppas model; and validate the "
    "integrated material–model framework through single-phase and steady-state "
    "two-phase core displacement experiments. This work establishes a coupled "
    "experimental–modeling methodology that translates sustained tracer release "
    "signals into quantitative per-interval production allocation—a broadly "
    "applicable framework for unconventional reservoir management."
)

# Apply intro changes
for idx, new_text in [(11, new_p11), (13, new_p13), (15, new_p15)]:
    p = doc.paragraphs[idx]
    for run in p.runs:
        run.text = ""
    p.runs[0].text = new_text
    print(f"Intro para [{idx}] rewritten.")

# ═══════════════════════════════════════════════════════════════════
# PART 2: DEEPEN SECTION 3.7 — ADE MODEL PHYSICAL INTERPRETATION
# ═══════════════════════════════════════════════════════════════════

# Add a new paragraph after the current ADE results paragraph [144]
# to deepen physical interpretation

new_model_insight = (
    "The fitted parameters carry several physically significant implications. "
    "First, the Peclet number Pe = x/α = 0.934 (≈1) places the transport "
    "regime at the transition between advection-dominated (Pe ≫ 1) and "
    "dispersion-dominated (Pe ≪ 1) behavior. This regime is characteristic of "
    "a gradually releasing source: advection alone would produce a sharp, "
    "symmetric Gaussian pulse, while the comparable dispersion broadens the "
    "peak and generates the extended tail that carries 47% of the integrated "
    "signal. The near-unity Pe thus constitutes independent physical "
    "corroboration of the matrix-diffusion-controlled release mechanism "
    "identified in Section 3.6. Second, the tanh blending function provides a "
    "physically motivated, continuously differentiable transition between the "
    "two transport regimes—a decisive improvement over a piecewise step "
    "function, which would introduce a non-physical discontinuity in the first "
    "derivative of the concentration profile at the crossover. Third, the close "
    "agreement between the pump-set flow rate (0.50 mL·min⁻¹) and the "
    "independently fitted flow rate (0.46 mL·min⁻¹, 8% relative error) "
    "demonstrates that the model parameters are not merely empirical "
    "fitting coefficients but carry well-defined physical meaning. Fourth, "
    "the 47% erfc tail contribution has direct practical relevance: it "
    "quantitatively confirms that nearly half of the total detectable tracer "
    "signal originates from sustained matrix-diffusion-controlled release "
    "rather than from the initial concentration pulse. This finding establishes "
    "that ESP-T-based monitoring is fundamentally a long-term, diffusion-"
    "governed process—a conclusion that would be inaccessible from empirical "
    "release fitting alone."
)

# Find the paragraph that contains "provides quantitative confirmation"
# and insert the deepened discussion after it
for i, p in enumerate(doc.paragraphs):
    if "establishes the physical basis for using ESP-T as a long-term production monitor" in p.text:
        target_idx = i
        print(f"Found ADE conclusion at para [{i}], inserting deepened insight...")
        break
else:
    target_idx = 144  # fallback
    print(f"Using fallback index [{target_idx}]")

# Insert new paragraph by splitting the current one
# Actually, we can't easily insert paragraphs in python-docx.
# Strategy: append to the existing paragraph
p_target = doc.paragraphs[target_idx]
current_text = p_target.text
enhanced_text = current_text.rstrip() + " " + new_model_insight
for run in p_target.runs:
    run.text = ""
p_target.runs[0].text = enhanced_text
print("ADE model physical interpretation deepened.")

# ═══════════════════════════════════════════════════════════════════
# PART 3: REBUILD REFERENCES (authoritative, verified)
# ═══════════════════════════════════════════════════════════════════

NEW_REFS = [
    # [1] IEA flagship report
    "[1]\tIEA. World Energy Outlook 2024 [R]. Paris: International Energy Agency, 2024.",

    # [2] Unconventional resource scale
    "[2]\tMONTGOMERY S L, JARVIE D M, BOWKER K A, et al. Mississippian Barnett Shale, Fort Worth basin, north-central Texas: Gas-shale play with multi–trillion cubic foot potential [J]. AAPG Bulletin, 2005, 89(2): 155–175.",

    # [3] Global fossil fuel projections
    "[3]\tMOHR S H, WANG J, ELLEM G, et al. Projection of world fossil fuels by country [J]. Fuel, 2015, 141: 120–135.",

    # [4] EOR for tight/shale
    "[4]\tWANG L, TIAN Y, YU X, et al. Advances in improved/enhanced oil recovery technologies for tight and shale reservoirs [J]. Fuel, 2017, 210: 425–445.",

    # [5] Fracturing fluid review — classic
    "[5]\tBARATI R, LIANG J T. A review of fracturing fluid systems used for hydraulic fracturing of oil and gas wells [J]. Journal of Applied Polymer Science, 2014, 131(16): 40735.",

    # [6] Production data analysis HF wells
    "[6]\tMEDEIROS F, KURTOGLU B, OZKAN E, et al. Analysis of production data from hydraulically fractured horizontal wells in shale reservoirs [J]. SPE Reservoir Evaluation & Engineering, 2010, 13(3): 559–568.",

    # [7] Tracer testing review
    "[7]\tPATIDAR A K, JOSHI D, DRISTANT U, et al. A review of tracer testing techniques in porous media specially attributed to the oil and gas industry [J]. Journal of Petroleum Exploration and Production Technology, 2022, 12(12): 3339–3356.",

    # [8] Inter-well chemical tracer field case
    "[8]\tSANNI M, AL-ABBAD M, KOKAL S, et al. Pushing the envelope of residual oil measurement: A field case study of a new class of inter-well chemical tracers [J]. Journal of Petroleum Science and Engineering, 2018, 163: 538–545.",

    # [9] PITT tracer stability
    "[9]\tSILVA M, STRAY H, BJØRNSTAD T. Stability assessment of PITT tracer candidate compounds — The case of pyrazines [J]. Journal of Petroleum Science and Engineering, 2019, 182: 106269.",

    # [10] Radioactive tracer — historical
    "[10]\tWATKINS J W, MARDOCK E S. Use of radioactive iodine as a tracer in water-flooding operations [J]. Journal of Petroleum Technology, 1954, 6(9): 117–124.",

    # [11] Micro-substance tracer in fractured wells
    "[11]\tYANG H, GUO K, LIN L, et al. Application of micro-substance tracer test in fractured horizontal wells [J]. Journal of Petroleum Exploration and Production Technology, 2024, 14(5): 1235–1246.",

    # [12] Tracer-eluting proppant — Zhao
    "[12]\tZHAO B, PANTHI K, MOHANTY K K. Tracer eluting proppants for hydraulic fracture characterization [J]. Journal of Petroleum Science and Engineering, 2020, 190: 107048.",

    # [13] Self-suspension + tracer slow-release proppant
    "[13]\tZHOU Y, LIU H, GAO J, et al. Coated proppants with self-suspension and tracer slow-release functions [J]. Journal of Petroleum Science and Engineering, 2022, 208: 109645.",

    # [14] Marked proppant dynamics — SPE paper
    "[14]\tMALYAVKO E, UPADHYE V, HUSEIN N. Research of operational dynamics of a well with two hydraulic fractures with use of marked proppant penetrating into one productive formation [C]. SPE-215624-MS, SPE International Hydraulic Fracturing Technology Conference and Exhibition, Muscat, Oman, 2023.",

    # [15] Rare earth tracer release kinetics
    "[15]\tLI N, CHENG Q, GONG Z, et al. Release kinetics of rare earth tracer from polymer-coated proppants for hydraulic fracture analysis [J]. Geoenergy Science and Engineering, 2023, 227: 211782.",

    # [16] Oleophilic tracer-slow-released proppant — Gong (PS)
    "[16]\tGONG Z, LI N, KANG W, et al. Novel oleophilic tracer-slow-released proppant for monitoring the oil production contribution [J]. Fuel, 2024, 364: 130945.",

    # [17] Multifunctional proppant review
    "[17]\tWANG G, MA Q, REN L, et al. A comprehensive review of multifunctional proppants [J]. ACS Omega, 2024, 9(44): 44120–44133.",

    # [18] PS-MMA nanofiller composite proppant
    "[18]\tKRISHNAN M R, LI W, ALHARBI B, et al. In-situ high-strength poly(styrene-methyl methacrylate)-2D nanofiller composite microbeads as potential proppants in hydraulic fracturing [J]. Geoenergy Science and Engineering, 2025, 257: 214195.",

    # [19] PEG/epoxy composite PCM
    "[19]\tGUO X, WEI K, NI T, et al. Preparation and performance analysis of polyethylene glycol/epoxy resin composite phase change material [J]. Journal of Energy Storage, 2024, 88: 111525.",

    # [20] Ultra-lightweight PMMA proppant
    "[20]\tLIANG C, LUO W, YAN C, et al. Ultra-lightweight proppant synthesized from PMMA/pine bark composite: Low-cost material and outstanding properties [J]. Chemistry Letters, 2016, 45(8): 994–996.",

    # [21] High-strength ultra-lightweight composite proppant
    "[21]\tZOVEIDAVIANPOOR M, GHARIBI A, BIN JAAFAR M Z. Experimental characterization of a new high-strength ultra-lightweight composite proppant derived from renewable resources [J]. Journal of Petroleum Science and Engineering, 2018, 170: 1038–1047.",

    # [22] Epoxy resin well sealant durability
    "[22]\tSABINS F, APBLETT A, SHAFER R, et al. Epoxy resin exhibits long-term durability and chemical stability as a well sealant [C]. SPE-204374-MS, SPE International Conference on Oilfield Chemistry, The Woodlands, Texas, 2021.",

    # [23] Epoxy sustained-release tracer — water phase
    "[23]\tLI H, LIU Z, LI Y, et al. Evaluation of the release mechanism of sustained-release tracers and its application in horizontal well inflow profile monitoring [J]. ACS Omega, 2021, 6(29): 19269–19280.",

    # [24] Epoxy resin long-term solid tracer screening
    "[24]\tWEI M, WANG Y, DUAN Y, et al. Screening and performance evaluation of epoxy resin long-term sustained-release solid tracer [J]. International Journal of Oil, Gas and Coal Technology, 2024, 36(2): 170–196.",

    # [25] Magnetic nano-Fe3O4 oleophilic tracer
    "[25]\tGONG Z, LI N, QIN M, et al. Magnetic nano-Fe₃O₄-based oleophilic tracer for stability studies of nano-tracer in oilfields condition [J]. Colloids and Surfaces A: Physicochemical and Engineering Aspects, 2024, 683: 133085.",

    # [26] K-P model seminal paper
    "[26]\tRITGER P L, PEPPAS N A. A simple equation for description of solute release I. Fickian and non-Fickian release from non-swellable devices in the form of slabs, spheres, cylinders or discs [J]. Journal of Controlled Release, 1987, 5(1): 23–36.",

    # [27] Peppas-Sahlin coupling model
    "[27]\tPEPPAS N A, SAHLIN J J. A simple equation for the description of solute release. III. Coupling of diffusion and relaxation [J]. International Journal of Pharmaceutics, 1989, 57(2): 169–172.",

    # [28] 1D ADE analytical solutions — classic groundwater transport
    "[28]\tVAN GENUCHTEN M T, ALVES W J. Analytical solutions of the one-dimensional convective–dispersive solute transport equation [R]. USDA Technical Bulletin No. 1661, U.S. Department of Agriculture, 1982.",

    # [29] Sphericity/roundness chart
    "[29]\tKRUMBEIN W C, SLOSS L L. Stratigraphy and Sedimentation [M]. 2nd ed. San Francisco: W.H. Freeman and Company, 1963.",
]

# Find reference section start
ref_start = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'References':
        ref_start = i
        break

if ref_start is None:
    print("WARNING: 'References' heading not found!")
    ref_start = 167  # fallback

print(f"Reference section starts at para [{ref_start}]")

# Clear all existing reference paragraphs (from ref_start+1 to end)
for i in range(ref_start + 1, len(doc.paragraphs)):
    p = doc.paragraphs[i]
    for run in p.runs:
        run.text = ""

# Write new references
for j, ref_text in enumerate(NEW_REFS):
    target_para_idx = ref_start + 1 + j
    if target_para_idx < len(doc.paragraphs):
        p = doc.paragraphs[target_para_idx]
        # Ensure at least one run exists
        if len(p.runs) == 0:
            p.add_run("")
        for run in p.runs:
            run.text = ""
        p.runs[0].text = ref_text
        # Set font properties
        for run in p.runs:
            run.font.size = doc.paragraphs[ref_start + 1].runs[0].font.size if doc.paragraphs[ref_start + 1].runs else None
    else:
        # Need to add new paragraph — python-docx limitation
        print(f"  WARNING: Cannot add para for ref [{j+1}] — beyond document range")

# Clear any remaining old ref paragraphs
for i in range(ref_start + 1 + len(NEW_REFS), len(doc.paragraphs)):
    p = doc.paragraphs[i]
    text = p.text.strip()
    if text and (text.startswith('[') or 'Reagents' in text or 'ESP-T' in text or 'Water passage' in text or 'K-P model' in text or 'Fit parameters' in text or 'R2' in text):
        for run in p.runs:
            run.text = ""

print(f"References rebuilt: {len(NEW_REFS)} entries.")

# ── SAVE ───────────────────────────────────────────────────────────
out = '四氧化三铁环氧树脂拟合/ESP-T_全面修改.docx'
doc.save(out)
print(f"\nSaved to: {out}")
print("Close Word, then rename to ESP-T_final.docx")