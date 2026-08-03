#!/usr/bin/env python3
"""生成168轻断食减脂计划 Word 文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_table_with_style(doc, headers, rows, col_widths=None, header_color="2B579A"):
    """添加带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'Microsoft YaHei'
        set_cell_shading(cell, header_color)

    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Microsoft YaHei'
            if r % 2 == 1:
                set_cell_shading(cell, "F2F2F2")

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table

def main():
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # 修改默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ==================== 封面 ====================
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("100天减脂塑形计划")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(43, 87, 154)
    run.font.name = 'Microsoft YaHei'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("16:8 轻断食 · 科学饮食 · 力量+有氧训练")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.font.name = 'Microsoft YaHei'

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("目标：75kg → 65kg  |  周期：100天  |  每周训练4天")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(80, 80, 80)
    run.font.name = 'Microsoft YaHei'

    doc.add_page_break()

    # ==================== 目录概览 ====================
    h = doc.add_heading('目录', level=1)
    toc_items = [
        "一、个人数据与目标拆解",
        "二、16:8 轻断食方案",
        "三、每日营养与热量分配",
        "四、一周饮食参考食谱",
        "五、运动训练计划（每周4天）",
        "六、每周进度追踪表",
        "七、关键注意事项",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_page_break()

    # ==================== 一、个人数据与目标拆解 ====================
    doc.add_heading('一、个人数据与目标拆解', level=1)

    doc.add_heading('1.1 基础数据', level=2)
    add_table_with_style(doc,
        ["指标", "数值", "说明"],
        [
            ["身高", "169 cm", "——"],
            ["起始体重", "75 kg", "Day 1"],
            ["目标体重", "65 kg", "Day 100"],
            ["总减重", "10 kg", "100天"],
            ["BMI (起始)", "26.3", "超重边缘"],
            ["BMI (目标)", "22.8", "正常范围"],
            ["年龄（估算）", "22-28 岁", "用于BMR计算"],
        ],
        col_widths=[4, 3, 6]
    )

    doc.add_heading('1.2 热量计算', level=2)
    doc.add_paragraph(
        '基础代谢率（Mifflin-St Jeor 公式）：\n'
        'BMR = 10 × 75 + 6.25 × 169 − 5 × 25 − 161 ≈ 1520 kcal/天\n\n'
        '每日总消耗（TDEE，中等活动系数 1.55）：\n'
        'TDEE = 1520 × 1.55 ≈ 2350 kcal/天\n\n'
        '目标热量缺口：500-800 kcal/天\n'
        '目标摄入：1500-1600 kcal/天\n\n'
        '预期减重速度：0.5-0.8 kg/周 → 100天可减 7-11 kg，10kg 目标合理可行。'
    )

    doc.add_heading('1.3 阶段划分', level=2)
    add_table_with_style(doc,
        ["阶段", "天数", "目标体重", "日均摄入", "重点"],
        [
            ["适应期", "Day 1-14", "75→73 kg", "1600 kcal", "适应断食节奏，建立习惯"],
            ["燃脂期1", "Day 15-40", "73→69 kg", "1550 kcal", "稳定掉秤，力量提升"],
            ["燃脂期2", "Day 41-75", "69→66 kg", "1500 kcal", "加速燃脂，维持肌肉"],
            ["巩固期", "Day 76-100", "66→65 kg", "1550 kcal", "平稳过渡，防止反弹"],
        ],
        col_widths=[3, 2.5, 3, 3, 4]
    )

    doc.add_page_break()

    # ==================== 二、16:8 轻断食方案 ====================
    doc.add_heading('二、16:8 轻断食方案', level=1)

    doc.add_heading('2.1 时间安排', level=2)
    doc.add_paragraph(
        '• 进食窗口：每日 11:00 - 19:00（8小时）\n'
        '• 断食窗口：每日 19:00 - 次日 11:00（16小时）\n'
        '• 如果上午训练，可调整为 9:00 - 17:00 或 10:00 - 18:00\n'
        '• 断食期间允许：白开水、黑咖啡（无糖无奶）、纯茶'
    )

    doc.add_heading('2.2 每日进食节奏', level=2)
    add_table_with_style(doc,
        ["时间", "餐次", "热量占比", "说明"],
        [
            ["11:00", "第一餐（开食）", "30% (~480 kcal)", "高蛋白+复合碳水，启动代谢"],
            ["13:30-14:30", "第二餐（主餐）", "40% (~620 kcal)", "最丰盛一餐，蛋白质+蔬菜+优质碳水"],
            ["18:00-18:30", "第三餐（收食）", "25% (~400 kcal)", "轻食为主，高蛋白+蔬菜，低碳水"],
            ["19:00后", "断食开始", "0 kcal", "只喝水/茶/黑咖啡"],
        ],
        col_widths=[3, 3, 4, 5.5]
    )

    doc.add_paragraph(
        '\n⚠️ 如果两餐制更适合你（进食窗口内吃两顿大餐），可以合并第一餐和第二餐，时间安排在 12:00-13:00 和 17:00-18:00。'
    )

    doc.add_page_break()

    # ==================== 三、每日营养与热量分配 ====================
    doc.add_heading('三、每日营养与热量分配（1550 kcal 基准）', level=1)

    doc.add_heading('3.1 三大营养素', level=2)
    add_table_with_style(doc,
        ["营养素", "每日摄入", "热量", "占比", "说明"],
        [
            ["蛋白质", "110-130g", "440-520 kcal", "28-33%", "保护肌肉，提高代谢；体重×1.5-1.7g"],
            ["碳水化合物", "130-150g", "520-600 kcal", "33-38%", "训练供能，选低GI为主"],
            ["脂肪", "40-50g", "360-450 kcal", "23-28%", "必需脂肪酸，激素平衡"],
        ],
        col_widths=[2.5, 2.5, 2.5, 2, 5]
    )

    doc.add_heading('3.2 食物选择指南', level=2)

    doc.add_heading('✅ 优先选择（绿灯食物）', level=3)
    doc.add_paragraph(
        '蛋白质：鸡胸肉、去皮鸡腿、瘦牛肉、鱼虾、鸡蛋、豆腐、无糖酸奶、牛奶\n'
        '碳水：燕麦、糙米、红薯/紫薯、全麦面包、荞麦面、玉米、南瓜\n'
        '蔬菜：绿叶菜（菠菜、生菜、西兰花）、黄瓜、番茄、彩椒、菌菇（不限量！）\n'
        '脂肪：橄榄油、牛油果、坚果（每天一小把≈15g）、亚麻籽\n'
        '饮品：白开水（每天2-3L）、黑咖啡、纯茶'
    )

    doc.add_heading('⚠️ 适量控制（黄灯食物）', level=3)
    doc.add_paragraph(
        '水果（每天1-2份，选低糖）：蓝莓、草莓、猕猴桃、苹果、柚子\n'
        '全脂奶制品、鸡胸肉以外的鸡肉部位、红肉（每周2-3次）'
    )

    doc.add_heading('❌ 避免（红灯食物）', level=3)
    doc.add_paragraph(
        '含糖饮料、奶茶、果汁、甜点饼干、油炸食品、加工肉制品（火腿肠）、\n'
        '精白米饭/白面包大量摄入、酱料（沙拉酱/千岛酱）、酒精'
    )

    doc.add_page_break()

    # ==================== 四、一周饮食参考食谱 ====================
    doc.add_heading('四、一周饮食参考食谱（~1550 kcal/天）', level=1)

    days_recipes = [
        ("周一（高蛋白日）", [
            ("11:00 开食", "燕麦50g + 牛奶200ml + 鸡蛋2个 + 蓝莓50g", "~480 kcal"),
            ("14:00 主餐", "糙米饭100g + 鸡胸肉150g + 西兰花200g + 橄榄油5ml", "~590 kcal"),
            ("18:00 收食", "无糖酸奶150g + 坚果15g + 黄瓜/番茄不限量", "~380 kcal"),
        ]),
        ("周二（鱼肉日）", [
            ("11:00 开食", "全麦面包2片 + 牛油果半个 + 煮鸡蛋2个 + 纯牛奶200ml", "~500 kcal"),
            ("14:00 主餐", "蒸红薯150g + 煎三文鱼120g + 菠菜200g（蒜蓉清炒）", "~580 kcal"),
            ("18:00 收食", "豆腐150g（凉拌/煮汤） + 菌菇蔬菜汤 + 杂粮饼1小块", "~370 kcal"),
        ]),
        ("周三（牛肉日）", [
            ("11:00 开食", "燕麦粥 + 鸡蛋2个 + 苹果半个 + 核桃10g", "~470 kcal"),
            ("14:00 主餐", "荞麦面（干重60g） + 瘦牛肉100g + 彩椒洋葱番茄炒", "~600 kcal"),
            ("18:00 收食", "希腊酸奶150g + 蓝莓100g + 鸡蛋白2个", "~380 kcal"),
        ]),
        ("周四（虾仁日）", [
            ("11:00 开食", "蒸紫薯100g + 牛奶200ml + 煮鸡蛋2个", "~460 kcal"),
            ("14:00 主餐", "糙米饭100g + 虾仁150g + 西兰花+胡萝卜清炒", "~570 kcal"),
            ("18:00 收食", "毛豆100g + 番茄2个 + 少量坚果", "~420 kcal"),
        ]),
        ("周五（鸡腿日）", [
            ("11:00 开食", "全麦三明治（面包2片+鸡蛋+生菜+番茄） + 牛奶200ml", "~490 kcal"),
            ("14:00 主餐", "杂粮饭100g + 去皮烤鸡腿1个 + 烤蔬菜（西葫芦/彩椒/洋葱）", "~610 kcal"),
            ("18:00 收食", "无糖酸奶150g + 蛋白粉15g（可选） + 猕猴桃1个", "~350 kcal"),
        ]),
        ("周六（自由搭配日）", [
            ("11:00 开食", "燕麦/全麦面包 + 鸡蛋 + 水果 + 牛奶/酸奶", "~480 kcal"),
            ("14:00 主餐", "碳水+优质蛋白+大量蔬菜，可少量放纵（如少量意面/寿司）", "~620 kcal"),
            ("18:00 收食", "蔬菜汤 + 鸡蛋白 + 水果1份", "~350 kcal"),
        ]),
        ("周日（轻断食日）", [
            ("12:00 第一餐", "鸡蛋2个 + 牛奶250ml + 大份蔬菜沙拉", "~430 kcal"),
            ("17:00 第二餐", "鱼肉/鸡胸100g + 蒸蔬菜 + 红薯100g", "~480 kcal"),
            ("", "总摄入约 1200-1300 kcal，帮助制造更大缺口", ""),
        ]),
    ]

    for day_title, meals in days_recipes:
        doc.add_heading(day_title, level=2)
        rows = []
        for time, content, cal in meals:
            if time:
                rows.append([time, content, cal])
        add_table_with_style(doc,
            ["时间", "内容", "热量"],
            rows,
            col_widths=[3, 8.5, 3],
            header_color="2E7D32" if "周日" not in day_title else "E65100"
        )
        doc.add_paragraph()  # 间距

    doc.add_page_break()

    # ==================== 五、运动训练计划 ====================
    doc.add_heading('五、运动训练计划（每周4天）', level=1)

    doc.add_heading('5.1 训练安排', level=2)
    add_table_with_style(doc,
        ["星期", "训练内容", "时长", "重点"],
        [
            ["周一", "下肢力量 + 核心", "60-70 min", "臀腿大肌群（耗能大户）"],
            ["周二", "上肢力量 + 核心", "60-70 min", "背部+肩部塑形"],
            ["周三", "休息 / 散步 / 拉伸", "30 min", "主动恢复"],
            ["周四", "下肢力量 + 核心", "60-70 min", "臀腿（与周一不同动作）"],
            ["周五", "上肢 + HIIT", "60 min", "上肢+高强度间歇燃脂"],
            ["周六", "休息 / 户外活动", "——", "散步/骑行/爬山"],
            ["周日", "休息 / 瑜伽拉伸", "20-30 min", "完全恢复"],
        ],
        col_widths=[2, 4, 3, 4.5]
    )

    doc.add_heading('5.2 训练结构（每次60-70分钟）', level=2)
    doc.add_paragraph(
        '1️⃣ 热身（8-10分钟）\n'
        '   • 泡沫轴放松 + 动态拉伸 + 开合跳/跳绳3分钟\n\n'
        '2️⃣ 力量训练（35-40分钟）\n'
        '   • 4-5个动作，每个动作 3-4组 × 10-15次\n'
        '   • 组间休息 45-60秒\n'
        '   • 重点：复合动作为主（深蹲、硬拉、划船、推举）\n\n'
        '3️⃣ 有氧训练（15-20分钟）\n'
        '   • 选择：跑步机坡度快走（坡度12/速度5）或 HIIT（30秒冲刺+30秒休息）\n'
        '   • 心率目标：最大心率的 65-75%（约 130-150 bpm）\n\n'
        '4️⃣ 拉伸放松（5-10分钟）\n'
        '   • 静态拉伸训练肌群，每个动作保持20-30秒'
    )

    doc.add_heading('5.3 下肢训练日动作库（周一/周四交替）', level=2)
    add_table_with_style(doc,
        ["动作", "组数×次数", "目标肌群", "备注"],
        [
            ["哑铃/壶铃深蹲", "4×12", "股四头肌+臀大肌", "核心动作，保证幅度"],
            ["罗马尼亚硬拉", "3×12", "腘绳肌+臀大肌", "保持背部平直"],
            ["保加利亚分腿蹲", "3×10（每侧）", "臀腿+稳定性", "手持哑铃增加强度"],
            ["臀推/臀桥", "4×15", "臀大肌", "顶峰收缩2秒"],
            ["哑铃侧弓步", "3×12（每侧）", "大腿内外侧", "——"],
            ["站姿提踵", "3×20", "小腿", "收尾动作"],
        ],
        col_widths=[3.5, 3, 3.5, 3.5]
    )

    doc.add_heading('5.4 上肢训练日动作库（周二/周五交替）', level=2)
    add_table_with_style(doc,
        ["动作", "组数×次数", "目标肌群", "备注"],
        [
            ["哑铃划船", "4×12（每侧）", "背阔肌", "俯身保持稳定"],
            ["高位下拉/引体向上", "4×10-12", "背部宽度", "可用弹力带辅助"],
            ["哑铃肩推", "3×12", "三角肌", "坐姿更有控制"],
            ["哑铃卧推/俯卧撑", "3×12", "胸肌+肱三头肌", "确保全幅度"],
            ["侧平举", "3×15", "三角肌中束", "小重量控制为主"],
            ["哑铃弯举", "3×12", "肱二头肌", "控制离心"],
            ["绳索/弹力带面拉", "3×15", "后束+肩袖", "改善圆肩"],
        ],
        col_widths=[3.5, 3, 3.5, 3.5]
    )

    doc.add_page_break()

    # ==================== 六、每周进度追踪表 ====================
    doc.add_heading('六、每周进度追踪表', level=1)
    doc.add_paragraph('每周一早上空腹称重并记录，填入下表追踪进度。')

    weeks_data = [
        ["第1周", "Day 1-7", "75.0", "", "适应断食，可能有轻微不适"],
        ["第2周", "Day 8-14", "74.3", "", "体重开始松动"],
        ["第3周", "Day 15-21", "73.6", "", "进入燃脂期"],
        ["第4周", "Day 22-28", "72.9", "", "月度小结"],
        ["第5周", "Day 29-35", "72.2", "", ""],
        ["第6周", "Day 36-42", "71.5", "", "40天节点"],
        ["第7周", "Day 43-49", "70.8", "", ""],
        ["第8周", "Day 50-56", "70.0", "", "50天，过半！"],
        ["第9周", "Day 57-63", "69.2", "", ""],
        ["第10周", "Day 64-70", "68.4", "", ""],
        ["第11周", "Day 71-77", "67.6", "", "75天节点"],
        ["第12周", "Day 78-84", "66.8", "", ""],
        ["第13周", "Day 85-91", "66.0", "", ""],
        ["第14周", "Day 92-98", "65.3", "", "接近目标"],
        ["第100天", "Day 99-100", "65.0", "", "🎯 目标达成！"],
    ]
    add_table_with_style(doc,
        ["周期", "天数", "预期体重(kg)", "实际体重(kg)", "本周记录"],
        weeks_data,
        col_widths=[2, 2.5, 3, 3, 5],
        header_color="C62828"
    )

    doc.add_page_break()

    # ==================== 七、关键注意事项 ====================
    doc.add_heading('七、关键注意事项', level=1)

    doc.add_heading('7.1 16:8执行要点', level=2)
    tips_fasting = [
        "前3-5天会有饥饿感，属于正常反应；多喝水、黑咖啡可以抑制饥饿。",
        "断食期间绝对不能摄入任何热量，包括含糖口香糖、牛奶咖啡、果汁等。",
        "如果低血糖症状明显（头晕、出冷汗），立刻吃一小份水果或全麦饼干，不要硬撑。",
        "生理期可以适当放宽进食窗口（10:8），听从身体信号。",
        "周末可以稍微灵活，但进食窗口不要超过10小时。",
    ]
    for tip in tips_fasting:
        doc.add_paragraph(tip, style='List Bullet')

    doc.add_heading('7.2 运动注意事项', level=2)
    tips_exercise = [
        "新手前两周先熟悉动作模式，重量以能做满12-15次为准，不要追求大重量。",
        "每次训练前必须热身，训练后必须拉伸。",
        "训练日和休息日穿插安排，不要连续4天训练（恢复和训练同样重要）。",
        "如果当天感到极度疲劳，可以只做力量、不做有氧，或者换成散步。",
        "体重不掉时，优先增加每日步数（目标8000-10000步），比加更多有氧更可持续。",
    ]
    for tip in tips_exercise:
        doc.add_paragraph(tip, style='List Bullet')

    doc.add_heading('7.3 心态与习惯', level=2)
    tips_mind = [
        "体重波动1-2kg是正常的（水分、生理周期），看趋势不要看单日数字。",
        "每周拍一张正面/侧面照片，比体重秤更能反映变化。",
        "如果某天吃多了，不要自责，第二天恢复正常就好，不要报复性节食。",
        "睡眠非常重要！每天保证7-8小时睡眠，睡眠不足会降低代谢、增加饥饿感。",
        "找一个伙伴互相监督，或加入减脂社群，社会支持能大大提高成功率。",
        "100天是一段旅程，享受过程，不要只盯着终点。每减1kg都是胜利！",
    ]
    for tip in tips_mind:
        doc.add_paragraph(tip, style='List Bullet')

    doc.add_heading('7.4 平台期应对策略', level=2)
    doc.add_paragraph(
        '如果连续2周体重不下降（进入平台期），按顺序尝试以下调整：\n\n'
        '1. 增加每日步数到 10,000-12,000 步\n'
        '2. 将进食窗口缩短至 6-7 小时（如 11:00-17:00）\n'
        '3. 碳水循环：训练日吃150g碳水，休息日降至80-100g\n'
        '4. 调整训练：增加力量训练重量/组数，或有氧改为HIIT\n'
        '5. 检查隐藏热量：酱料、坚果、外卖油量往往被低估'
    )

    doc.add_paragraph()

    # ==================== 结尾 ====================
    doc.add_paragraph()
    end_p = doc.add_paragraph()
    end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end_p.add_run("💪 坚持100天，遇见更好的自己！加油！💪")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(43, 87, 154)
    run.font.name = 'Microsoft YaHei'

    # 保存
    output_path = os.path.expanduser("~/Desktop/claude/100天减脂计划_168轻断食.docx")
    doc.save(output_path)
    print(f"文档已保存至: {output_path}")

if __name__ == "__main__":
    main()