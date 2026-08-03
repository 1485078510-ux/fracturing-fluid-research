#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成页岩气藏地应力演化与裂缝扩展耦合机制文献综述（中英文各一份）
SCI期刊风格，Word格式保存至向晶文件夹
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = r"c:\Users\郝\Desktop\claude\向晶"

# ============================================================
# 中文文献综述内容
# ============================================================

CN_TITLE = "页岩气藏地应力演化与裂缝扩展耦合机制研究进展与展望"

CN_ABSTRACT = (
    "地应力场的动态演化与水力裂缝扩展之间存在强烈的双向耦合——应力重分布控制裂缝扩展路径，"
    "裂缝生长又反过来重塑应力场格局。这一耦合机制是页岩气藏体积压裂理论的核心科学问题，"
    "直接影响储层改造效果与井间干扰程度。本文系统梳理了近十年来该领域的研究进展，以三条"
    "主线组织文献。其一，从Biot经典孔隙弹性理论出发，追溯至现代四维（4D）动态地应力模型，"
    "涵盖应力阴影效应、衰竭诱导应力反转及多段压裂应力累积扰动等关键机制。其二，剖析页岩"
    "各向异性介质中的裂缝扩展行为，聚焦层理面相互作用、天然裂缝网络激活判据及多簇裂缝竞争"
    "扩展规律。其三，综述研究方法体系——纵向上涵盖有限元、离散元、扩展有限元、相场法及多尺度"
    "耦合等数值手段，真三轴压裂-AE/CT/DIC联合表征等室内实验技术，以及微地震-分布式声波/温度"
    "传感（DAS/DTS）协同监测等现场方法。在此基础上，凝练出当前研究面临的四类核心瓶颈：缺乏"
    "从纳米孔隙力学到储层尺度应力动力学的跨尺度桥接机制；热-力-化多场耦合的完整性不足；数据"
    "驱动方法的物理可解释性薄弱；以及时变蠕变-损伤对裂缝长期导流能力影响的系统研究缺失。"
    "最后，提出以层次化多尺度建模、物理信息神经网络、全耦合THMC模拟、数字孪生驱动的地质工程"
    "一体化和多井平台应力协同管理为核心的未来研究路线图，以期为该领域的理论深化与工程实践"
    "提供系统性的参考框架。"
)

CN_KEYWORDS = "页岩气藏；地应力演化；水力压裂；裂缝扩展；流固耦合；数值模拟；相场法；微地震监测"

CN_SECTIONS = [
    ("1 引言", [
        "页岩气的商业化开发依赖于水平井多段水力压裂技术对储层的大规模体积改造。压裂过程中，"
        "数以万方计的高压流体注入储层，不仅直接驱动裂缝的起裂与延伸，更深刻地扰动乃至重塑了"
        "储层原始地应力场的分布格局。而地应力场的每一次调整，又反过来决定着后续裂缝的扩展路径、"
        "几何形态及复杂程度——由此构成了地层变形与流体驱动之间典型的双向耦合反馈回路[1-3]。"
        "深入揭示这一耦合机制的内在规律，对于优化压裂设计、预测缝网形态、评估井间干扰风险以及"
        "指导加密井科学部署，具有不可替代的理论价值与工程意义。",

        "近年来，随着页岩气勘探开发向深层（>3500 m）、超深层领域持续推进，储层高温高压环境与"
        "复杂天然裂缝系统的叠加效应，使得地应力-裂缝耦合问题愈加突出[4-5]。国内外学者围绕这一"
        "主题开展了系统性的研究工作，在孔隙弹性理论拓展、多场耦合数值算法、真三轴物理模拟及光纤"
        "传感现场监测等多个维度取得了长足进展。然而，页岩储层固有的强非均质性、裂缝系统的多尺度"
        "跨级特征以及多物理场耦合过程的高度非线性，使得该领域在理论与方法层面仍面临诸多根本性"
        "挑战[6-8]。",

        "本文旨在系统梳理页岩气藏地应力演化与裂缝扩展耦合机制的研究脉络。与已有综述不同，本工作"
        "以「理论框架-扩展机理-耦合机制-研究方法-现存不足-未来方向」为逻辑主线，力求从多维度呈现"
        "该领域的全貌：既追溯Biot孔隙弹性理论在页岩储层中的适用性边界，也审视水力裂缝穿越层理-"
        "天然裂缝交互的三模式判据；既梳理有限元/离散元/相场法等数值手段的演进谱系，也分析真三轴"
        "实验-光纤监测等多源数据融合的技术瓶颈；最终凝练出当前研究的核心困境，提出具有可操作性"
        "的未来研究路线图，以期为该领域的理论突破与工程实践提供系统性的参考框架。"
    ]),

    ("2 地应力演化理论基础与研究现状", [
        ("2.1 Biot孔隙弹性理论与应力演化", [
            "地应力演化的理论基础源于Biot（1941）提出的孔隙弹性理论[9]。该理论建立了流体饱和多孔介质中"
            "应力-应变-孔隙压力之间的耦合关系，其核心有效应力原理为：σ' = σ − α·p，其中α为Biot系数，"
            "表征孔隙压力对总应力的抵消效率。对于富含有机质的页岩储层，由于矿物组成和孔隙结构的强各向异性，"
            "需要采用横观各向同性孔隙弹性模型来描述其力学行为[10]。Suarez-Rivera和Fjær（2013）系统评估了"
            "各向异性富有机质泥页岩系统的孔隙弹性效应，发现平行和垂直于层理方向的Biot系数存在显著差异，"
            "这对准确预测压裂过程中的有效应力变化至关重要[11]。",

            "在此基础上，Wang等（2025）基于Biot理论、嵌入式离散裂缝模型（EDFM）和有限体积法，建立了"
            "全耦合流-固力学模型，模拟了裂缝性页岩气藏40年生产期内的应力演化规律。研究发现，孔隙压力梯度"
            "的时空分布是驱动应力变化的主导因素，各应力分量（σ_xx, σ_yy, σ_xy）、差应力和应力反转区域"
            "在不同位置和时间达到各自的极值——这意味着不存在适用于整个储层的单一最优加密时机[12]。",

            "从更广泛的物理过程来看，页岩储层中的应力演化还受到多重机制的共同控制。吸附/解吸作用引起"
            "的基质收缩/膨胀效应、扩散和气体滑脱效应等微观机制同样对应力场的演变产生不可忽略的影响。"
            "近年来，考虑多物理场耦合——包括温度场（THM耦合）和化学场（THMC耦合）——的地应力演化模型"
            "成为新的研究热点[13-14]。"
        ]),

        ("2.2 应力阴影效应", [
            "应力阴影（Stress Shadow）效应是指水力裂缝张开时在裂缝周围产生的诱导应力场，该诱导应力"
            "会改变相邻裂缝的应力状态，从而影响多裂缝的同步或顺序扩展[15-17]。应力阴影效应是理解多簇"
            "压裂裂缝竞争扩展机制的核心概念。",

            "Ping等（2024）基于XFEM建立了二维多裂缝同步扩展模型，系统研究了应力阴影效应的影响因素。"
            "结果表明，增大簇间距或注入低黏度压裂液可有效减弱应力阴影效应；地应力差较小时裂缝偏转更为"
            "显著，而随着地应力差增大，裂缝逐渐转向最大水平主应力方向[18]。Han等（2023）进一步发现，"
            "当裂缝长度扩展至射孔间距的约2.5倍时，应力阴影效应达到最强；小间距条件下裂缝会向后扩展偏转"
            "[19]。Ju和Wang等（2024）采用自适应FEM/DEM方法，对比了5种簇间距（12.5–100 m）和3种压裂"
            "工艺（顺序、同时、平行）下的应力阴影影响区域、裂缝长度和微震震级[20]。",

            "Chang等（2024）通过全耦合DEM模拟，区分了黏度主导区和韧性主导区两种扩展机制下的应力阴影"
            "表现差异：黏度主导区裂缝趋于均匀扩展，而韧性主导区裂缝倾向于向相邻裂缝的反方向扩展以相互"
            "避开[21]。这些研究共同揭示，应力阴影效应的强弱受簇间距、地应力差、压裂液黏度、注入速率和"
            "岩石断裂韧性等多因素综合调控。",

            "值得注意的是，Zhu等（2021）和唐煊赫等（2021）指出，应力阴影效应不仅影响单段内的多簇裂缝"
            "扩展，还会在多段压裂的段间产生累积效应，引起「Frac-hit」现象和「微地震屏障」效应[22-23]。"
            "因此，在压裂设计中需要从段间距、簇间距、注入顺序和时间等多个维度综合优化。"
        ]),

        ("2.3 四维地应力动态演化", [
            "四维（4D）地应力演化是指在三维空间基础上增加时间维度，系统追踪应力场在压裂-生产全生命周期"
            "中的动态变化过程。朱海燕和唐煊赫等（2021）系统总结了渗流-地质力学耦合的四种基本框架："
            "全耦合、顺序耦合（单向和双向迭代）、单向耦合和拟耦合，并指出四维地应力演化模型主要分为"
            "基于全耦合的连续介质模型、离散裂缝模型和迭代耦合模型三大类[22]。",

            "在深层页岩气领域，Xu等（2025）针对四川盆地深层页岩气水平井平台，研究了压裂过程中的四维"
            "动态应力场演化规律。结果显示，压裂过程中最小水平主应力上升1.8–6.4 MPa，最大水平主应力上升"
            "1.1–3.2 MPa；随压裂段数的增加，最小水平主应力呈现明显的累积增长趋势，应力差从15 MPa降至"
            "11 MPa[24]。Ruan等（2026）进一步研究了天然弱面对四维应力演化的影响，发现低刚度弱面易发生"
            "形变并降低内部应力，最大水平主应力受弱面扰动的程度随弱面与主应力方向夹角增大而增强，且最小"
            "水平主应力表现为先减弱后增强的非单调响应特征[25]。",

            "在应力反转方面，Agrawal和Sharma（2018）利用基于近场动力学（Peridynamics）的孔隙弹性压裂"
            "模拟器，发现生产过程中母井裂缝之间会形成拉伸区域（应力拱效应），子井裂缝会不对称地向低压"
            "区域生长，优先排驱已生产区域而不是原始储层[26]。Wang等（2025）指出，若初始水平应力差超过"
            "某一临界值，应力反转（σ_h > σ_H）永远不会发生；仅当初始应力差较小时，衰竭才能驱动最大水平"
            "主应力方向的完全重定向[12]。这些发现对加密井部署和重复压裂时机选择具有重要指导意义。"
        ])
    ]),

    ("3 水力裂缝扩展机理研究现状", [
        ("3.1 裂缝起裂与扩展理论", [
            "水力裂缝的起裂与扩展是一个涉及岩石断裂力学、流体力学和孔隙弹性力学的多学科交叉问题。"
            "经典的线弹性断裂力学（LEFM）基于应力强度因子准则（K_I ≥ K_IC）判断裂缝起裂，能够较好地"
            "描述均质脆性岩石中的单一裂缝扩展行为。然而，页岩因其层理发育、矿物非均质性强和各向异性显著，"
            "裂缝扩展行为远比经典理论预测的复杂[27-28]。",

            "Gong等（2022）在Minerals期刊发表综述，系统梳理了页岩储层中影响水力裂缝扩展的地质因素"
            "（矿物组成、天然裂缝、地应力状态）和工程因素（压裂液类型、注入速率、射孔参数），指出单一"
            "数值方法难以完整描述页岩压裂的复杂物理过程，推荐通过整合多种数值算法来提高模型精度[29]。",

            "在水力裂缝与天然裂缝的交互方面，研究普遍认为存在三种基本模式：裂缝被天然裂缝捕获（arrest）、"
            "穿越天然裂缝（crossing）和沿天然裂缝偏移（dilation/deflection）[30-31]。决定交互模式的"
            "关键控制因素包括：天然裂缝的逼近角、地应力差、天然裂缝的黏聚力和摩擦系数、压裂液黏度和"
            "注入速率等。四川盆地南部深层页岩储层的研究表明，高应力差条件抑制天然裂缝激活及复杂缝网形成，"
            "天然裂缝逼近角的增加促使水力裂缝由单一捕获模式转变为捕获、穿越和阻挡的复合模式[32]。"
        ]),

        ("3.2 层理面与各向异性效应", [
            "页岩的层理结构是其最显著的特征之一，层理面通常为力学弱面，显著影响水力裂缝的垂向扩展行为。"
            "层状岩体中水力裂缝扩展的研究表明，当层理面与水力裂缝的逼近角较大、水平主应力差较大且胶结"
            "强度较强时，裂缝穿透层理面；反之裂缝沿层理面方向扩展[33]。层理倾角对裂缝网络复杂度的控制"
            "同样不可忽视——Zhao等（2022）通过真三轴实验结合X射线CT扫描发现，当层理角接近45°时，最有利"
            "于形成复杂的水力裂缝网络[34]。",

            "Saber等（2023）基于XFEM和内聚力模型（CZM）模拟了横观各向同性页岩地层中的多裂缝扩展，"
            "发现横观各向同性模型预测的裂缝比各向同性模型更窄、更长；层理倾角由0°增至90°时，缝宽增加，"
            "缝长呈先增后减的趋势变化[35]。Yang等（2023）通过页岩-砂岩互层储层的真三轴压裂实验，利用"
            "声发射监测和实时水压监测揭示了裂缝穿越岩性界面的行为特征：大的垂向应力差是裂缝穿透多岩性"
            "层位的关键条件[36]。"
        ]),

        ("3.3 天然裂缝网络与多裂缝竞争扩展", [
            "页岩储层中普遍发育的天然裂缝系统是影响水力裂缝网络形成的关键内在因素。Wu等（2022）利用"
            "真三轴实验系统结合声发射监测，系统研究了页岩油储层中多裂缝扩展行为，发现裂缝干扰在起裂阶段"
            "即已开始并严重影响多裂缝扩展；层理界面和纹理是裂缝干扰的主要原因，阻碍了裂缝高度增长；"
            "剪切型AE事件标志着层理滑移和纹理诱导裂缝偏转引起的裂缝干扰[37]。",

            "在天然裂缝稳定性演化方面，最新的研究建立了三维地质力学-离散裂缝网络（DFN）耦合评价方法，"
            "发现流体注入通过「有效正应力降低」与「摩擦强度下降」两条机制诱发天然裂缝失稳，而长期生产导致"
            "的孔隙压力下降可以增强有效正应力，整体提升裂缝稳定性[38]。Wang等（2024）基于相场法建立了"
            "水力-力学-损伤（HMD）耦合模型，系统研究了多簇压裂中水力裂缝与天然裂缝的交互关系，提出了"
            "三类接触关系的分类：θ=α、θ<α和θ>α，并发现±45°天然裂缝角度时可获得最高的水力裂缝分形维数"
            "（2.1280）[39]。",

            "在多裂缝竞争扩展方面，多簇压裂中各裂缝因应力阴影和流量分配不均而产生非均匀扩展。Ping等（2024）"
            "的研究表明，排量增加导致优势裂缝迅速发展，裂缝长度显著增加，而外围裂缝受到抑制[18]。"
            "这要求在压裂设计中平衡簇间距、注入参数和射孔设计，以实现各簇裂缝的均匀扩展，最大化储层改造"
            "体积（SRV）。"
        ])
    ]),

    ("4 地应力演化与裂缝扩展耦合机制研究", [
        ("4.1 流-固-力学耦合的理论框架", [
            "地应力演化与裂缝扩展的耦合本质上是流体流动、固体变形和裂缝损伤/扩展三个物理过程的交互作用。"
            "该耦合机制的核心在于：流体压力驱动裂缝扩展并改变孔隙压力场，裂缝的张开/闭合改变岩石的渗透率"
            "和刚度，应力场的重分布反过来影响裂缝的扩展方向和孔隙流体的流动路径，形成闭合的反馈回路[6,40]。",

            "《工程地质学报》2024年发表的综述系统分析了页岩储层压裂的流固耦合力学过程，将数值模拟方法"
            "分为两大维度：复杂压裂缝网动态演化模拟和储层尺度压裂模拟。在耦合策略上，全耦合方法同时求解"
            "所有控制方程，理论上最为精确但计算成本极高；顺序耦合（迭代耦合）依次求解各物理场并以迭代方式"
            "收敛，在计算效率和精度间取得较好平衡，是目前应用最广泛的框架[6]。",

            "在耦合机制的深化理解方面，中深层页岩储层天然裂缝稳定性演化研究揭示了时间依赖性的重要作用。"
            "压裂液注入引起的地层压力扩散并非瞬时完成，而是在时间和空间上逐步发展，这导致了有效应力和"
            "天然裂缝稳定性的时空动态变化[38]。此外，裂隙渗透率随有效应力的动态变化（应力敏感效应）构建了"
            "力学场与渗流场之间的关键反馈通道，是耦合机制中的重要非线性环节[41]。"
        ]),

        ("4.2 生产诱导应力变化与重复压裂", [
            "生产过程中的孔隙压力衰竭会引起储层地应力场的显著变化，这对重复压裂和加密井的裂缝扩展产生"
            "决定性影响。Agrawal和Sharma（2018）的研究表明，生产过程中母井之间形成拉伸区域，导致子井"
            "裂缝不对称地向已衰竭（低压）区域生长，而非均匀进入未开发的原始储层[26]。Yu等（2025）通过"
            "多尺度裂缝耦合流动-地质力学模拟发现，近井裂缝开度在生产过程中可降低达25.1%，渗透率下降达"
            "48.9%，考虑应力耦合效应的产量预测比非耦合模拟低20.3%[41]。",

            "Coupled Wellbore-Reservoir-Geomechanical Modelling（2024）在非均质水力裂缝条件下揭示了"
            "优势裂缝与劣势裂缝之间因应力演化和衰竭不均衡而产生的产量竞争效应，这种竞争进一步加剧了流动"
            "分布的不均匀性；增大簇间距则有助于促进更均匀的应力状态变化[42]。",

            "对于重复压裂，Guo等（2024）基于CT技术的真三轴重复压裂实验表明，天然裂缝和应力状态是影响"
            "重复压裂裂缝扩展的首要因素；利用暂堵剂形成大尺寸「片状」封堵体比端部封堵更有利于起裂新裂缝"
            "[43]。低水平应力差条件下，重复压裂更容易产生复杂裂缝网络。"
        ]),

        ("4.3 断层活化与套变问题中的应力-裂缝耦合", [
            "近年来，页岩气压裂过程中的断层活化和套管变形问题引起了广泛关注，这些问题本质上也是应力演化"
            "与裂缝（断层）扩展耦合的体现。基于流固耦合方法的研究表明，压裂注水过程中孔隙压力扩散区不断"
            "扩展，断层附近的有效应力降低，断层活化风险升高；断层两侧的体积应变呈现压缩和拉伸的相反趋势"
            "[44]。",

            "Scientific Reports（2024）发表的动态机制研究将断层滑移诱发的套管变形划分为三个阶段：触发活化→"
            "加速滑移→减速滑移。断层滑移行为严重受断层走向与最大地应力方向夹角的控制；动态模型考虑了应力降、"
            "摩擦系数变化以及套管/水泥环的阻力效应[45]。这为水平井压裂设计的断层避让和套损防治提供了理论依据。"
        ])
    ]),

    ("5 主要研究方法", [
        ("5.1 数值模拟方法", [
            "数值模拟是研究地应力-裂缝耦合机制最核心的手段。当前主要的数值方法包括以下六类：",

            "（1）有限元法（FEM）与内聚力模型（CZM）：FEM是应用最广泛的连续介质力学方法，通过内聚力"
            "单元模拟裂缝的起裂和扩展，能够较好地处理裂缝尖端的非线性过程区。Zhu等（2024）利用基于FEM的"
            "内聚力模型模拟了三维裂缝竞争扩展行为，揭示了裂缝延伸方向和长度受最大水平主应力方向的显著控制"
            "[46]。",

            "（2）扩展有限元法（XFEM）：通过富集函数在常规有限元形函数中引入不连续性，避免了裂缝扩展"
            "时的网格重构需求，在处理任意路径裂缝扩展方面具有独特优势。Ping等（2024）和Han等（2023）"
            "利用XFEM广泛研究了应力阴影效应和多裂缝同步扩展问题[18-19]。",

            "（3）离散元法（DEM）：基于颗粒接触模型模拟岩石的断裂和破碎过程，天然适合处理多裂缝起裂、"
            "交汇和网络化问题。Chang等（2024）建立了全耦合水力压裂DEM模型，成功模拟了从韧性主导到黏度"
            "主导不同机制下的多裂缝同时扩展[21]。",

            "（4）相场法（PFM）：基于变分原理和总能量最小化，通过连续的相场变量自动追踪裂缝路径，无需"
            "额外的断裂准则。Lu等（2023）系统综述了相场法在水力压裂模拟中的应用现状，指出该方法的天然"
            "优势在于多物理场耦合问题的处理[47]。Wang等（2024）进一步建立了HMD耦合相场模型，实现了多簇"
            "裂缝与天然裂缝交互的全过程模拟[39]。Yi等（2024）将相场法拓展至多孔黏弹性介质，首次考虑了"
            "页岩的蠕变行为对裂缝扩展的影响[48]。",

            "（5）边界元法（BEM）与位移不连续法（DDM）：仅需在边界上离散，降维处理，适合处理无限域"
            "和半无限域问题，在应力阴影计算和裂缝相互作用分析中应用广泛。",

            "（6）自适应FEM/DEM耦合方法：结合两种方法的优势，在裂缝附近采用DEM捕捉断裂过程，在远离"
            "裂缝区域采用FEM保证计算效率。Song等（2024）和Ju等（2024）使用基于ELFEN TGR软件的自适应"
            "FEM/DEM方法，考虑了水力裂缝、天然裂缝、孔隙渗流、支撑剂运移的完整物理过程[20,49]。",

            "（7）多尺度耦合策略：将不同尺度的裂缝系统采用不同方法描述。宏观水力裂缝用EDFM或XFEM显式"
            "表征，中小尺度天然裂缝用多重相互作用连续体（MINC）或双孔双渗（DPDK）等效，通过计算均匀化"
            "或灵活性叠加方法实现跨尺度的双向耦合[41,50]。"
        ]),

        ("5.2 室内物理实验方法", [
            "室内实验是验证理论模型和揭示物理机制不可或缺的手段。当前主要的实验方法包括：",

            "（1）真三轴水力压裂实验：是最接近现场条件的实验室模拟方法。通过在立方体岩样（通常300 mm × "
            "300 mm × 300 mm）的三个方向独立加载不同应力，模拟真实的三向地应力状态，同时注入压裂液驱动"
            "裂缝扩展[34,36-37]。Wu等（2022）利用真三轴系统和AE监测技术，成功捕捉了多簇压裂中裂缝干扰"
            "的起始和发展过程[37]。Zhao等（2022）开发了高应力真三轴实验系统，研究了静水压力条件下不同"
            "层理角度的页岩压裂特性[34]。",

            "（2）声发射（AE）监测：通过记录岩石断裂时释放的弹性波信号，对裂缝起裂、扩展和交汇过程进行"
            "实时定位和机制判别。AE信号参数（b值、RA-AF值）可用于区分拉伸型和剪切型裂缝事件[34,37]。",

            "（3）X射线CT扫描与三维重构：在压裂前后对岩样进行CT扫描，通过三维重构和分形维数分析定量表征"
            "裂缝网络的几何形态和复杂度[34,43]。Guo等（2024）利用CT技术研究了重复压裂中暂堵剂对裂缝扩展"
            "的影响[43]。",

            "（4）数字图像相关（DIC）技术：通过追踪岩样表面散斑图案的变形，获取全场位移和应变分布，"
            "可用于分析裂缝尖端的应变集中和损伤演化过程。",

            "（5）分布式声波传感（DAS）实验室应用：KAUST的研究团队（2023）在50 cm³立方体岩块六个表面"
            "分布DAS光纤，在真三轴应力条件下进行压裂实验，实现了比常规DAS高约10倍的采样频率，并通过CT"
            "成像对DAS反演的微震定位结果进行了标定[51]。"
        ]),

        ("5.3 现场监测方法", [
            "现场监测是获取储层尺度地应力演化和裂缝扩展信息的最直接手段。当前主要技术包括：",

            "（1）微地震监测：是目前水力压裂裂缝表征最成熟和广泛应用的现场技术。通过在地面或邻井中布设"
            "检波器阵列，记录压裂过程中岩石破裂产生的微地震事件，反演裂缝的空间展布和激活序列。最新的"
            "混合井中DAS与三分量（3C）检波器联合微地震成像技术（2023，IEEE TGRS），利用几何平均逆时偏移"
            "方法对Montney地层的41个微地震事件进行成像，并探索了利用DAS+3C数据联合约束震源机制[52]。",

            "（2）分布式声波传感（DAS）：利用光纤本身作为连续传感器阵列，实现沿整个井筒的实时声波/应变"
            "测量。Correa等（2024）在Austin Chalk/Eagle Ford现场实验室，首次将地面轨道振动器（SOV）与"
            "DAS结合进行连续时移VSP采集，以每小时的时间分辨率捕捉了裂缝的打开和闭合过程[53]。"
            "Glubokovskikh等（2024）进一步利用连续井中DAS数据追踪了个别天然裂缝的激活和生长演化[54]。",

            "（3）低频DAS（<0.05 Hz）：ConocoPhillips专利（2023）提出了利用低频DAS信号约束远场裂缝"
            "长度、高度、宽度和密度的方法，可从水平邻井中检测应力阴影引起的应变扰动，以及直井中裂缝"
            "碰撞引起的热信号[55]。",

            "（4）分布式温度传感（DTS）：通过监测压裂和生产过程中的温度变化，反演各簇的注入量分配和"
            "生产贡献。ESG Solutions（2024）在Marcellus页岩案例中整合了DAS注入分配、DTS温度瞬态分析和"
            "微地震形变/应力状态分析，实现了裂缝行为的时空综合描述[56]。"
        ])
    ]),

    ("6 主要研究成果与认识", [
        "综合以上文献分析，当前研究在地应力演化与裂缝扩展耦合机制方面已形成四个层面的系统性成果：",

        "（1）理论层面：以Biot孔隙弹性理论为基石，构建了涵盖渗流-应力-损伤多场交互的地应力演化理论"
        "框架。明确了有效应力原理在裂缝性页岩储层中的适用边界，揭示了Biot系数各向异性（平行vs.垂直"
        "层理方向）对耦合行为的定量影响，并初步建立了考虑吸附/解吸诱导基质变形的非等温孔隙弹性拓展"
        "模型[9-12]。",

        "（2）机理层面：系统揭示了应力阴影效应的多因素（簇间距、应力差、流体黏度、注入速率、断裂韧性）"
        "协同调控规律；确立了水力裂缝-天然裂缝交互的三种基本模式及其在逼近角-应力差-界面强度三维参数"
        "空间中的相态分布；识别了生产诱导应力反转的发生条件——即初始水平应力差低于临界阈值是应力方向"
        "完全重定向的必要条件，这一发现对加密井部署和重复压裂时机选择具有直接的工程指导意义[15-21,26,30-32]。",

        "（3）方法层面：初步形成了「数值模拟为核心引擎、室内实验为物理验证、现场监测为实况约束」的"
        "三位一体研究方法论。数值模拟正从单一方法向多方法耦合（FEM-DEM、XFEM-EDFM）和多尺度层次化"
        "（微观DEM→宏观EDFM）方向演进；室内实验已确立真三轴-AE-CT-DIC多模态联合表征的标准范式；"
        "现场监测正经历从单一微地震向DAS+DTS+微地震多物理量协同的技术换代[22-56]。",

        "（4）工程层面：提出了基于应力演化轨迹的加密井部署优化原则、基于应力反转临界条件的重复压裂"
        "时机选择方法，以及以增大簇间距和优化注入顺序为核心的应力阴影缓解策略。在地质工程一体化方面，"
        "压裂-生产一体化模拟器的概念验证已初步完成，正迈向工程化应用[22,41-43]。"
    ]),

    ("7 现存问题与不足", [
        "尽管已取得上述丰硕成果，当前研究在以下七个维度仍面临实质性瓶颈，且各瓶颈之间存在深度关联，"
        "单一维度的局部突破难以引发系统性的研究范式变革：",

        "（1）多尺度跨级桥接的理论缺失：从纳米级有机孔隙到千米级储层，空间尺度跨越逾10个数量级。"
        "现有的跨尺度信息传递仍高度依赖经验性的岩石物理关系（如经验渗透率-应力耦合公式），缺乏基于"
        "严格计算均匀化理论的桥接框架。微观实验（纳米压痕、分子动力学模拟）揭示的孔隙尺度物理机制，"
        "尚无法以自洽的方式约束宏观连续介质模型的本构参数[6,50]。",

        "（2）多物理场耦合的完备性不足：超过80%的现有模型局限于HM（水力-力学）双向耦合，温度场和"
        "化学场的集成尚处于概念验证阶段。页岩黏土矿物与水基压裂液的化学相互作用——包括黏土水化膨胀"
        "应力、半透膜渗透效应、离子交换诱导的岩石结构弱化——对裂缝面力学性质和近缝应力场的定量影响"
        "仍未得到系统表征，这在高含黏土矿物的深层页岩中尤为关键[13-14]。",

        "（3）天然裂缝系统的表征与参数化困境：天然裂缝的几何属性（密度、产状、迹长、开度）和力学"
        "属性（黏聚力、摩擦系数、法向/切向刚度）在储层三维空间中的分布具有强非均质性，而现有技术"
        "手段（岩心、测井、地震）对裂缝属性的约束能力非常有限。这一「数据荒漠」导致DFN模型的预测"
        "结果高度依赖先验假设，不同实现之间的裂缝扩展预测差异可达数倍[32,38]。",

        "（4）高保真模拟的计算可行性壁垒：全耦合、高分辨率（亚米级网格）的三维数值模型单次运行"
        "耗时可达数周，与工程尺度的快速迭代优化需求形成尖锐矛盾。相场法在描述复杂裂缝拓扑（分叉、"
        "交汇、非平面扩展）方面具有理论完备性优势，但其损伤带宽度与网格尺寸的强关联性导致计算成本"
        "随模型尺寸呈超线性增长，成为制约其工程化应用的首要瓶颈[47-48]。",

        "（5）时变效应的系统研究空白：页岩的黏弹塑性蠕变行为可导致裂缝长期有效闭合、支撑剂嵌入深度"
        "在数月至数年时间尺度上的持续增长，以及压裂液破胶后裂缝面物理化学性质的时变演化——这些时间"
        "依赖性过程共同构成了应力-裂缝耦合的「长期项」，但其定量规律与作用机制的研究几乎空白[48]。",

        "（6）数据驱动方法的「黑箱」困境：机器学习（特别是深度学习）在裂缝扩展预测、微地震事件定位"
        "和压裂参数反演中展现出超越传统方法的潜力，但物理可解释性的缺失严重制约了其在工程决策中的"
        "可信度。将物理约束（Biot理论、断裂力学准则）嵌入网络架构的物理信息神经网络（PINN）仍处于"
        "早期探索阶段，在复杂几何和强非线性条件下的收敛性与泛化能力有待验证[57-58]。",

        "（7）实验室成果向储层尺度的外推困境：真三轴实验受限于岩样尺寸（通常≤300 mm立方体）和加载"
        "时间（通常≤数小时），储层尺度的空间非均质性和时间尺度效应（生产周期长达数十年）无法在实验室"
        "条件下复现。相似性准则的缺失使得从厘米级实验直接外推至千米级储层存在原理性偏差[34,37]。"
    ]),

    ("8 未来研究方向与展望", [
        "针对上述问题，本文提出以下八项具有优先级的未来研究方向，兼顾基础理论的纵深突破与工程应用的"
        "横向拓展：",

        "（1）层次化多尺度一体化建模：构建「分子尺度（MD/DFT）→ 孔隙尺度（LBM/PNM）→ 岩心尺度"
        "（FEM/DEM）→ 储层尺度（EDFM/等效连续体）」的四级层次化耦合框架。关键在于发展基于计算均匀化"
        "理论和机器学习代理模型的跨尺度信息传递协议，使得微观物理机制能够以可计算的方式约束宏观模型"
        "的本构参数与状态方程，最终实现「微观机理可知、宏观工程可用」的多尺度模拟范式[41,50-51]。",

        "（2）物理信息神经网络与机理-数据融合建模：以Biot孔隙弹性方程、断裂力学能量准则和非线性摩擦"
        "本构为物理约束内核，嵌入深度学习网络架构的损失函数与网络结构中，构建兼具第一性原理严格性和"
        "数据驱动效率的代理模型。重点攻克强非线性、多场耦合条件下PINN的训练收敛性难题，并借助SHAP、"
        "LIME等可解释性工具，从训练后的网络中逆向挖掘潜在的物理规律，实现数据对机理的反哺[57-58]。",

        "（3）全耦合THMC多场模拟平台研发：在现有HM耦合框架的基础上，逐级集成热传导-对流（T）、"
        "化学反应-输运-相变（C）过程，建立覆盖「注入-焖井-返排-生产」全周期的THMC全耦合模型。"
        "这一平台对于理解深层高温环境、复杂压裂液体系（如酸性压裂液、CO₂泡沫压裂液）以及页岩-流体"
        "化学不相容性引发的次生应力效应，具有不可替代的价值[13-14,48]。",

        "（4）数字孪生驱动的地质工程一体化闭环：以实时DAS/DTS/微地震监测流数据为输入，以高保真"
        "数值模型为引擎，以集合卡尔曼滤波和贝叶斯反演为同化手段，构建压裂过程的数字孪生系统。"
        "该系统的核心能力是实现「监测→反演→预测→优化→控制」的分钟级闭环，将当前「压后评估」的"
        "事后模式升级为「压中调控」的实时模式[53-56]。",

        "（5）时变本构理论与全生命周期稳定性评估：发展耦合蠕变损伤的黏弹塑性本构模型，系统量化"
        "页岩蠕变速率、应力松弛时标、支撑剂嵌入动力学和裂缝导流能力衰减之间的定量关系。将实验室"
        "短时（小时级）蠕变实验与外推至生产时间尺度（年级）的理论桥接作为优先突破方向，为页岩气井"
        "全生命周期的地质力学稳定性评估和重复压裂时机决策提供理论支撑[48]。",

        "（6）缝网系统的统计力学描述：跳出单一裂缝的确定性描述范式，借鉴逾渗理论、自组织临界性和"
        "复杂网络科学等统计物理学方法，在系统层面建立应力场-裂缝网络协同演化的统计表征框架。重点"
        "探索裂缝网络的连通性相变阈值、分形维数的演化标度律以及注入能量-缝网复杂度的涌现关系。",

        "（7）多井平台-区块尺度的应力场协同管理：将研究视野从单井/井组拓展至多井平台乃至整个开发区块，"
        "研究多井异步压裂、同步压裂和联合生产条件下的区域应力场协同演化规律。核心目标是发展以最小化"
        "井间应力干扰、最大化区块最终采收率为导向的井网井距-压裂顺序-生产制度的协同优化方法[24-25,42]。",

        "（8）面向碳中和的低碳压裂应力-裂缝耦合研究：在「双碳」战略背景下，系统开展CO₂干法压裂、"
        "超临界CO₂压裂和CO₂泡沫压裂中的应力-裂缝耦合机制研究。重点揭示CO₂与页岩的化学作用（碳酸盐"
        "溶解-沉淀、有机质溶胀）对局部应力场和裂缝扩展路径的影响，探索CO₂地质封存与页岩气增产协同"
        "优化的可行性边界，为低碳化页岩气开发提供理论储备。"
    ]),

    ("9 结论", [
        "地应力演化与裂缝扩展的耦合机制是页岩气水力压裂理论体系中最为核心的科学问题之一，"
        "涉及岩石力学、流体力学、断裂力学与传热传质学等多学科深度交叉。本文系统梳理了该领域"
        "的研究现状、方法体系与发展脉络，形成以下主要认识：",

        "（1）以Biot孔隙弹性理论为基石的地应力演化研究已构建起较为完备的理论框架。应力阴影效应、"
        "衰竭诱导应力反转以及四维动态应力场演化，共同构成了该领域的三大核心研究命题，三者之间"
        "存在紧密的因果关联与时空接力关系。",

        "（2）水力裂缝扩展研究已实现从均质单一裂缝模型向考虑层理面、天然裂缝网络和横观各向同性"
        "的复杂缝网模型的范式跃迁。水力裂缝-天然裂缝交互的三模式判据（捕获-穿越-偏转）得到了广泛"
        "的实验验证与数值再现，但逼近角-应力差-界面强度的三维相图仍有待完善。",

        "（3）数值模拟、室内实验与现场监测三大方法体系已初步形成互补闭环，但三者间的一体化协同"
        "水平仍不足以支撑全过程的「模拟-标定-验证-预测」工作流。相场法、自适应FEM/DEM耦合及多尺度"
        "层次化策略代表了数值方法的前沿方向；真三轴-AE-CT-DIC多模态联合表征已成为室内实验的标准"
        "范式；DAS-DTS-微地震的多物理量协同监测正在重塑现场裂缝诊断的技术格局。",

        "（4）当前研究在多尺度跨级桥接、多物理场（特别是化学场）耦合完备性、天然裂缝空间分布的"
        "表征不确定性、高保真模型的计算可行性、时变蠕变-损伤效应以及数据驱动方法的物理可解释性"
        "等六个维度仍面临实质性瓶颈，这些瓶颈相互关联，单一突破难以产生系统性进展。",

        "（5）未来研究应着力构建「微观机制约束-介观均匀化桥接-宏观工程响应」的层次化多尺度建模"
        "范式，大力发展物理信息神经网络以实现数据-机理的双向赋能，推进全耦合THMC多场模拟以逼近"
        "真实的储层物理化学环境，建立数字孪生系统以实现压裂过程的实时反演与动态优化，并深化时变"
        "本构理论与多井平台应力协同管理研究，共同推动页岩气压裂从「经验设计」向「机理驱动、数据"
        "赋能」的精准调控转型。"
    ])
]

CN_REFERENCES = [
    "[1] Maxwell S C. Microseismic Imaging of Hydraulic Fracturing: Improved Engineering of Unconventional Shale Reservoirs[M]. Tulsa: Society of Exploration Geophysicists, 2014.",
    "[2] Molenaar M M, Hill D J, Webster P, et al. First downhole application of distributed acoustic sensing for hydraulic-fracturing monitoring and diagnostics[J]. SPE Drilling & Completion, 2012, 27(1): 32-38.",
    "[3] Jin L, Roy B. Hydraulic fracturing induced seismicity[J]. Earth-Science Reviews, 2017, 174: 52-68.",
    "[4] 赵金洲, 任岚, 沈骋, 等. 页岩气储层缝网压裂理论与技术研究进展[J]. 天然气工业, 2018, 38(3): 1-14.",
    "[5] 邹才能, 董大忠, 王玉满, 等. 中国页岩气特征、挑战及前景[J]. 石油勘探与开发, 2015, 42(6): 689-701.",
    "[6] 唐煊赫, 朱海燕, 等. 页岩储层压裂流固耦合数值模拟及裂缝扩展研究进展[J]. 工程地质学报, 2024, 32(4): 1381-1396.",
    "[7] 郭建春, 苟波, 任山, 等. 深层页岩气储层压裂改造技术进展与展望[J]. 天然气工业, 2022, 42(6): 67-82.",
    "[8] 陈勉, 金衍, 张广清. 石油工程岩石力学[M]. 北京: 科学出版社, 2008.",
    "[9] Biot M A. General theory of three-dimensional consolidation[J]. Journal of Applied Physics, 1941, 12(2): 155-164.",
    "[10] Cheng A H D. Poroelasticity[M]. Cham: Springer International Publishing, 2016.",
    "[11] Suarez-Rivera R, Fjær E. Evaluating the poroelastic effect on anisotropic, organic-rich, mudstone systems[J]. Rock Mechanics and Rock Engineering, 2013, 46(3): 569-580.",
    "[12] Wang Q, Wang Y, et al. Evolution law of stress induced by pressure depletion in fractured shale reservoirs: implications for subsequent refracturing and infill well development[J]. Petroleum, 2025, 11(1): 71-83.",
    "[13] Yi D, Yi L P, Yang Z Z, et al. Coupled thermo-hydro-mechanical-phase field modelling for hydraulic fracturing in thermo-poroelastic media[J]. Computers and Geotechnics, 2024, 172: 106421.",
    "[14] Yi D, Yang Z Z, Yi L P, et al. Hydraulic fracturing phase-field model in porous viscoelastic media[J]. International Journal of Mechanical Sciences, 2024, 276: 109374.",
    "[15] Roussel N P, Sharma M M. Optimizing fracture spacing and sequencing in horizontal-well fracturing[J]. SPE Production & Operations, 2011, 26(2): 173-184.",
    "[16] Olson J E. Multi-fracture propagation modeling: applications to hydraulic fracturing in shales and tight gas sands[C]// The 42nd U.S. Rock Mechanics Symposium. San Francisco: ARMA, 2008.",
    "[17] Bunger A P, Zhang X, Jeffrey R G. Parameters affecting the interaction among closely spaced hydraulic fractures[J]. SPE Journal, 2012, 17(1): 292-306.",
    "[18] Ping Y, Wang H, Zhang J, et al. Numerical simulation of the simultaneous development of multiple fractures in horizontal wells based on the extended finite element method[J]. Energies, 2024, 17(5): 1057.",
    "[19] Han G D, Cui Z, Zhu H Y. The effect of perforation spacing on the variation of stress shadow[J]. LAPSE, 2023: 2023.33468.",
    "[20] Ju Y, Wang Y L, et al. Stress shadow effects in multistage horizontal hydrofracturing of tight reservoirs: a numerical analysis considering perforation cluster spacings and fracturing sequences[J]. Geomechanics and Geophysics for Geo-Energy and Geo-Resources, 2024, 10: 183.",
    "[21] Chang X, Hou B, Ding Y. DEM modeling of simultaneous propagation of multiple hydraulic fractures across different regimes, from toughness- to viscosity-dominated[J]. Rock Mechanics and Rock Engineering, 2024, 57: 481-503.",
    "[22] 朱海燕, 唐煊赫, 赵金洲, 等. 页岩气储层四维地应力演化及加密井复杂裂缝扩展研究进展[J]. 石油科学通报, 2021, 6(3): 455-475.",
    "[23] 郭建春, 路千里, 等. 页岩气压裂关键问题与探索[J]. 天然气工业B辑, 2023, 10(2): 183-197.",
    "[24] Xu E, Yu T, Chen L Q, et al. Study on the evolution law of four-dimensional dynamic stress fields in fracturing of deep shale gas platform wells[J]. Processes, 2025, 13(9): 2709.",
    "[25] 阮奇, 张烈辉, 赵玉龙, 等. 天然弱面对页岩气四维应力及加密井裂缝扰动规律研究[J]. 地质力学学报, 2026, 32(1).",
    "[26] Agrawal S, Sharma M M. Impact of pore pressure depletion on stress reorientation and its implications on the growth of child well fractures[C]// SPE/AAPG/SEG Unconventional Resources Technology Conference. Houston: URTeC, 2018: 2875375.",
    "[27] Adachi J, Siebrits E, Peirce A, et al. Computer simulation of hydraulic fractures[J]. International Journal of Rock Mechanics and Mining Sciences, 2007, 44(5): 739-757.",
    "[28] Detournay E. Mechanics of hydraulic fractures[J]. Annual Review of Fluid Mechanics, 2016, 48: 311-339.",
    "[29] Gong X, Ma X H, Liu Y Y, et al. Advances in hydraulic fracture propagation research in shale reservoirs[J]. Minerals, 2022, 12(11): 1438.",
    "[30] Blanton T L. An experimental study of interaction between hydraulically induced and pre-existing fractures[C]// SPE/DOE Unconventional Gas Recovery Symposium. Pittsburgh: SPE, 1982: 10847.",
    "[31] Warpinski N R, Teufel L W. Influence of geologic discontinuities on hydraulic fracture propagation[J]. Journal of Petroleum Technology, 1987, 39(2): 209-220.",
    "[32] 四川盆地南部深层页岩储层天然裂缝对体积压裂裂缝网络的影响[J]. 大庆石油地质与开发, 2025, 44(3).",
    "[33] 层状岩体中水力裂缝扩展规律研究[J]. 中国煤炭地质, 2025, 37(8).",
    "[34] Zhao H, Liang B, Sun W J, et al. Effects of hydrostatic pressure on hydraulic fracturing properties of shale using X-ray computed tomography and acoustic emission[J]. Journal of Petroleum Science and Engineering, 2022, 219: 110684.",
    "[35] Saber A, et al. Propagation of multiple hydraulic fractures in a transversely isotropic shale formation[C]// SSRN Preprint, 2023: 4334460.",
    "[36] Yang L, Sheng X C, Zhang B, et al. Propagation behavior of hydraulic fractures in shale under triaxial compression considering the influence of sandstone layers[J]. Gas Science and Engineering, 2023, 118: 205091.",
    "[37] Wu S, Gao K, Wang X Q, et al. Investigating the propagation of multiple hydraulic fractures in shale oil rocks using acoustic emission[J]. Rock Mechanics and Rock Engineering, 2022, 55: 6015-6032.",
    "[38] 中深层页岩储层天然裂缝稳定性演化规律研究[J]. 石油科学通报, 2026, 11(1).",
    "[39] Wang Y J, Wang B, Su H, et al. Mechanisms of fracture propagation from multi-cluster using a phase field based HMD coupling model in fractured reservoir[J]. Petroleum Science, 2024, 21(3): 1829-1851.",
    "[40] 张丹, 易良平, 杨兆中, 等. 裂缝性储层水力压裂混合模式相场模型[J]. 应用数学和力学(英文版), 2024, 45: 911-930.",
    "[41] Yu D, Kang W, et al. A novel coupled flow and geomechanics simulation method employing multi-scale fracture model[C]// International Field Exploration and Development Conference. IFEDC, 2024.",
    "[42] Coupled wellbore-reservoir-geomechanical modelling of uneven depletion and induced stress responses in shale oil reservoirs with heterogeneous hydraulic fractures[C]// International Geomechanics Symposium. Kuala Lumpur: ARMA-IGS, 2024: 0343.",
    "[43] Guo P, Li X, Zheng B, et al. Study on fracture propagation rules of shale refracturing based on CT technology[J]. Processes, 2024, 12(1): 131.",
    "[44] 基于流固耦合方法的压裂断层失稳规律研究[J]. 工程地质学报, 2024, 32(4): 1439-1446.",
    "[45] Zhang Y, Li X, Wang H, et al. Investigation on dynamic mechanism of fault slip and casing deformation during multi-fracturing in shale gas wells[J]. Scientific Reports, 2024, 14: 13164.",
    "[46] 朱海燕, 唐煊赫, 等. 压裂裂缝系统动态演化机制: 水力压裂物理实验与有限元数值模拟的启示[J]. 石油科学(英文版), 2024, 21(6): 3839-3866.",
    "[47] 卢千里, 张航, 郭建春, 等. 基于相场法的水力裂缝扩展模拟技术现状及展望[J]. 天然气工业, 2023, 43(3): 59-68.",
    "[48] 易多, 杨兆中, 易良平, 等. 多孔黏弹性介质中水力压裂相场模型[J]. 国际机械科学学报, 2024.",
    "[49] 宋宪章, 等. 多级水力压裂应力阴影效应的数值分析[J]. 矿业科学学报, 2024, 9(4): 507-518.",
    "[50] Wang C. A multi-scale, multi-continuum and multi-physics model to simulate coupled fluid flow and geomechanics in shale gas reservoirs[D]. Golden: Colorado School of Mines, 2018.",
    "[51] KAUST. Distributed acoustic sensing (DAS) for hydraulic fracture monitoring in laboratory scale[C]// 3rd EAGE Workshop on Fiber Optic Sensing for Energy Applications. Chengdu: EAGE, 2023.",
    "[52] Characterizing microearthquakes induced by hydraulic fracturing with hybrid borehole DAS and three-component geophone data[J]. IEEE Transactions on Geoscience and Remote Sensing, 2023, 61: 5904615.",
    "[53] Correa J, Glubokovskikh S, Nayak A, et al. Revealing complex subsurface dynamics with continuous seismic monitoring: observations using distributed acoustic sensing and surface orbital vibrators during hydraulic fracturing[J]. Geophysics, 2024, 89(6): M1-M16.",
    "[54] Glubokovskikh S, Correa J, Ajo-Franklin J, et al. Continuous surface-to-distributed acoustic sensor snapshots explain reactivation of individual natural fractures during an unconventional reservoir stimulation[J]. Geophysics, 2024, 89(6): M17-M32.",
    "[55] ConocoPhillips Company. Low frequency distributed acoustic sensing hydraulic fracture geometry[P]. US Patent: 2023/0003119 A1, 2023-01-05.",
    "[56] ESG Solutions. Fibre-optic sensing and microseismic monitoring evaluate and enhance hydraulic fracturing via real-time and post-treatment analysis[R]. Marcellus Shale Case Study, 2024.",
    "[57] Qu H Y, Zhang J L, Zhou F J, et al. Evaluation of hydraulic fracturing based on deep neural network with physical constraints[J]. Petroleum Science, 2023, 20(2): 1129-1141.",
    "[58] Shentu J J, Lin B T, Jin Y, et al. Interpretable machine learning for hydraulic fracture propagation in conglomerate rock based on discrete element method[J]. Acta Geotechnica, 2024, 19: 3581-3602."
]

# ============================================================
# 英文文献综述内容
# ============================================================

EN_TITLE = "Coupling Mechanisms Between In-Situ Stress Evolution and Hydraulic Fracture Propagation in Shale Gas Reservoirs: A Comprehensive Review"

EN_ABSTRACT = (
    "The coupling between in-situ stress evolution and hydraulic fracture propagation—a bidirectional "
    "feedback loop in which stress redistribution governs fracture paths while fracture growth remodels "
    "the stress field—remains a central unresolved challenge in shale gas development, directly controlling "
    "stimulation effectiveness and inter-well interference risks. This review systematically synthesizes "
    "advances in understanding this coupling mechanism published over the past decade. We organize the "
    "literature around three conceptual pillars. First, we trace the theoretical evolution from Biot's "
    "classical poroelasticity framework to modern four-dimensional (4D) dynamic stress models that capture "
    "stress shadow effects, depletion-induced stress reversal, and cumulative inter-stage stress "
    "perturbation. Second, we examine fracture propagation mechanisms in anisotropic shale formations, "
    "with emphasis on bedding-plane interactions, natural fracture network activation, and multi-cluster "
    "competitive growth. Third, we survey the methodological toolkit—spanning numerical approaches (FEM, "
    "DEM, XFEM, phase-field, and multi-scale coupled methods), laboratory techniques (true triaxial "
    "testing with integrated AE/CT/DIC characterization), and field monitoring technologies (microseismic, "
    "distributed acoustic sensing, and distributed temperature sensing). Our synthesis reveals four "
    "persistent knowledge gaps: the absence of rigorous cross-scale bridging mechanisms linking nanoscale "
    "poromechanics to reservoir-scale stress dynamics; incomplete integration of thermal and chemical "
    "fields within hydro-mechanical coupling frameworks; limited physical interpretability of emerging "
    "machine-learning-based approaches; and inadequate treatment of time-dependent creep and damage effects "
    "on long-term fracture conductivity. We close by outlining a research roadmap that prioritizes "
    "hierarchical multi-scale modeling, physics-informed neural networks, fully coupled THMC simulation, "
    "digital twin-driven geo-engineering integration, and multi-well platform-scale stress management "
    "strategies. This review provides a unified reference for researchers and practitioners seeking to "
    "bridge the gap between mechanistic understanding and engineering application of stress-fracture "
    "coupling in unconventional reservoirs."
)

EN_KEYWORDS = "shale gas reservoir; in-situ stress evolution; hydraulic fracturing; fracture propagation; fluid-solid coupling; numerical simulation; phase-field method; microseismic monitoring"

EN_SECTIONS = [
    ("1 Introduction", [
        "Shale gas, as a critical unconventional natural gas resource, relies on multi-stage hydraulic "
        "fracturing in horizontal wells for commercial production. During hydraulic stimulation, the "
        "injection of large volumes of high-pressure fluid not only directly drives fracture initiation "
        "and propagation but also substantially perturbs the original in-situ stress field. The evolving "
        "stress field, in turn, governs the propagation paths, geometry, and complexity of subsequent "
        "fractures, forming a strongly coupled, bidirectional feedback loop [1-3]. A thorough understanding "
        "of this coupling mechanism is of profound theoretical significance and practical value for "
        "optimizing fracturing designs, predicting fracture network morphology, assessing inter-well "
        "interference risks, and guiding infill well deployment.",

        "In recent years, as shale gas development has expanded into deep and ultra-deep reservoirs, "
        "the high-temperature, high-pressure environment and complex natural fracture systems have made "
        "the stress-fracture coupling problem increasingly prominent [4-5]. Researchers worldwide have "
        "devoted substantial effort to this topic, yielding significant advances across theoretical "
        "models, numerical methods, experimental techniques, and field monitoring technologies. Nevertheless, "
        "owing to the strong heterogeneity of shale formations, the multi-scale complexity of fracture "
        "systems, and the highly nonlinear nature of multi-physics coupling, numerous challenges persist [6-8].",

        "This review aims to systematically synthesize the current state of research on the coupling "
        "mechanisms between in-situ stress evolution and hydraulic fracture propagation in shale gas "
        "reservoirs. We provide a comprehensive overview of theoretical foundations, research methodologies, "
        "major findings, and existing challenges, and conclude by proposing promising future research "
        "directions to serve as a reference for advancing this field."
    ]),

    ("2 Fundamentals and Current Status of In-Situ Stress Evolution", [
        ("2.1 Biot's Poroelasticity and Stress Evolution", [
            "The theoretical foundation of in-situ stress evolution rests on Biot's theory of poroelasticity, "
            "originally formulated in 1941 [9]. This framework establishes the coupled stress-strain-pore "
            "pressure relationships in fluid-saturated porous media, with the core effective stress principle: "
            "σ′ = σ − α·p, where α is the Biot coefficient characterizing the efficiency of pore pressure "
            "in counteracting total stress. For organic-rich shale formations, the strong anisotropy in "
            "mineral composition and pore structure necessitates transversely isotropic poroelastic models "
            "to adequately describe their mechanical behavior [10]. Suarez-Rivera and Fjær (2013) "
            "systematically evaluated poroelastic effects in anisotropic, organic-rich mudstone systems, "
            "demonstrating that Biot coefficients measured parallel and perpendicular to bedding differ "
            "significantly—a critical consideration for accurately predicting effective stress changes "
            "during hydraulic fracturing operations [11].",

            "Building on this framework, Wang et al. (2025) developed a fully coupled fluid-flow/geomechanics "
            "model based on Biot's theory, the embedded discrete fracture model (EDFM), and the finite "
            "volume method to simulate stress evolution in fractured shale gas reservoirs over a 40-year "
            "production period. Their results revealed that the spatiotemporal distribution of pore pressure "
            "gradients is the dominant driver of stress changes: individual stress components (σ_xx, σ_yy, "
            "σ_xy), differential stress (Δσ), and stress reversal zones each attain their extrema at "
            "different locations and times—implying that no single optimal timing for infill operations "
            "exists for the entire reservoir [12].",

            "Beyond poroelasticity, multiple coupled mechanisms contribute to stress evolution in shale "
            "reservoirs. Matrix shrinkage/swelling driven by adsorption/desorption, diffusive transport, "
            "and gas slippage effects at the microscale all exert non-negligible influences on stress field "
            "dynamics. More recently, models incorporating thermo-hydro-mechanical (THM) and thermo-hydro-"
            "mechanical-chemical (THMC) couplings have emerged as new frontiers [13-14]."
        ]),

        ("2.2 Stress Shadow Effects", [
            "The stress shadow effect refers to the induced stress field generated around an opening "
            "hydraulic fracture, which modifies the stress state of adjacent fractures and thereby "
            "influences the synchronous or sequential propagation of multiple fractures [15-17]. This "
            "concept is central to understanding competitive fracture growth in multi-cluster stimulation.",

            "Ping et al. (2024) developed a two-dimensional XFEM-based model for simultaneous multi-fracture "
            "propagation, systematically investigating the controlling factors of stress shadow effects. "
            "Their results demonstrated that increasing cluster spacing or injecting low-viscosity fracturing "
            "fluid effectively mitigates stress shadow effects; fracture deflection is more pronounced under "
            "low horizontal stress differences, with fractures gradually reorienting toward the maximum "
            "horizontal stress direction as the stress difference increases [18]. Han et al. (2023) further "
            "revealed that stress shadow intensity peaks when fracture length reaches approximately 2.5 times "
            "the perforation spacing, with small spacing causing backward fracture deflection [19]. Ju and "
            "Wang et al. (2024) employed an adaptive FEM/DEM method to compare stress shadow influence zones, "
            "fracture lengths, and microseismic magnitudes across five cluster spacings (12.5–100 m) and "
            "three fracturing sequences (sequential, simultaneous, and parallel) [20].",

            "Chang et al. (2024), through fully coupled DEM simulations, distinguished stress shadow behavior "
            "under two propagation regimes: viscosity-dominated regimes promote uniform fracture growth, "
            "while toughness-dominated regimes cause fractures to propagate away from adjacent fractures to "
            "mutually avoid each other [21]. Collectively, these studies reveal that stress shadow intensity "
            "is governed by the complex interplay of cluster spacing, in-situ stress difference, fracturing "
            "fluid viscosity, injection rate, and rock fracture toughness.",

            "Notably, Zhu et al. (2021) and Tang et al. (2021) emphasized that stress shadow effects not "
            "only influence intra-stage multi-cluster fracture propagation but also produce cumulative "
            "inter-stage effects across multi-stage fracturing, giving rise to 'frac-hit' phenomena and "
            "'microseismic barrier' effects [22-23]. Comprehensive optimization across stage spacing, "
            "cluster spacing, injection sequence, and timing is therefore essential in fracturing design."
        ]),

        ("2.3 Four-Dimensional Dynamic Stress Evolution", [
            "Four-dimensional (4D) stress evolution refers to systematically tracking the spatiotemporal "
            "dynamics of the stress field (three spatial dimensions plus time) throughout the entire "
            "fracturing-production lifecycle. Zhu and Tang et al. (2021) provided a systematic review of "
            "four basic frameworks for coupled fluid-flow/geomechanics modeling: fully coupled, sequential "
            "coupled (one-way and two-way iterative), one-way coupled, and pseudo-coupled approaches, "
            "categorizing 4D stress evolution models into continuum-based, discrete fracture-based, and "
            "iteratively coupled models [22].",

            "In the context of deep shale gas development, Xu et al. (2025) investigated the 4D dynamic "
            "stress field evolution during multi-stage fracturing of deep shale gas platform wells in the "
            "Sichuan Basin. Their results showed that the minimum horizontal principal stress increased by "
            "1.8–6.4 MPa and the maximum horizontal principal stress by 1.1–3.2 MPa during fracturing; "
            "the minimum horizontal stress exhibited a clear cumulative growth trend with increasing "
            "fracturing stages, reducing the stress difference from 15 MPa to 11 MPa [24]. Ruan et al. "
            "(2026) further examined the influence of natural weak planes on 4D stress evolution, finding "
            "that low-stiffness weak planes readily deform with reduced internal stress, that maximum "
            "horizontal stress perturbation by weak planes intensifies with increasing angle between the "
            "weak plane and the principal stress direction, and that the minimum horizontal stress exhibits "
            "a non-monotonic response pattern of initial weakening followed by strengthening [25].",

            "Regarding stress reversal, Agrawal and Sharma (2018), using a peridynamics-based poroelastic "
            "fracturing simulator, demonstrated that tensile regions develop between parent well fractures "
            "during production due to poroelastic stress arching effects, causing child well fractures to "
            "grow asymmetrically toward depleted (low-pressure) regions rather than accessing virgin "
            "reservoir rock [26]. Wang et al. (2025) quantified a critical initial horizontal stress "
            "difference threshold: stress reversal (σ_h > σ_H) never occurs if the initial stress difference "
            "exceeds this critical value; only when the initial difference is sufficiently small can "
            "depletion drive complete reorientation of the maximum horizontal stress direction [12]. These "
            "findings carry profound implications for infill well deployment and refracturing timing decisions."
        ])
    ]),

    ("3 Hydraulic Fracture Propagation Mechanisms", [
        ("3.1 Fracture Initiation and Propagation Theories", [
            "Hydraulic fracture initiation and propagation represent a multidisciplinary problem integrating "
            "rock fracture mechanics, fluid mechanics, and poroelasticity. Classical linear elastic fracture "
            "mechanics (LEFM), based on the stress intensity factor criterion (K_I ≥ K_IC), adequately "
            "describes single fracture propagation in homogeneous brittle rocks. However, shale formations, "
            "characterized by well-developed bedding, strong mineral heterogeneity, and pronounced anisotropy, "
            "exhibit fracture propagation behavior far more complex than classical predictions [27-28].",

            "Gong et al. (2022) provided a comprehensive review in Minerals, systematically cataloging "
            "the geological factors (mineral composition, natural fractures, in-situ stress state) and "
            "engineering factors (fracturing fluid type, injection rate, perforation parameters) influencing "
            "hydraulic fracture propagation in shale reservoirs, and recommending the integration of multiple "
            "numerical algorithms to improve model accuracy given that no single method can fully capture "
            "the complex physics of shale fracturing [29].",

            "Regarding hydraulic fracture-natural fracture (HF-NF) interaction, a broad consensus recognizes "
            "three fundamental modes: arrest (HF captured by NF), crossing (HF penetrates through NF), and "
            "dilation/deflection (HF diverts along NF) [30-31]. The dominant interaction mode is governed "
            "by the approach angle, in-situ stress difference, NF cohesion and friction coefficient, "
            "fracturing fluid viscosity, and injection rate. Studies on deep shale reservoirs in the "
            "southern Sichuan Basin revealed that high stress difference conditions suppress natural "
            "fracture activation and complex fracture network formation, while increasing NF approach "
            "angles drive a transition from single arrest mode to a composite arrest-crossing-blocking "
            "regime [32]."
        ]),

        ("3.2 Bedding Plane and Anisotropy Effects", [
            "Bedding planes, as the most prominent structural feature of shale, act as mechanical weaknesses "
            "that significantly influence vertical hydraulic fracture propagation. Studies on fracture "
            "propagation in layered rock masses indicate that fractures penetrate bedding planes when the "
            "approach angle is large, the horizontal stress difference is high, and the bedding cementation "
            "strength is strong; conversely, fractures propagate along bedding planes [33]. Zhao et al. "
            "(2022), through true triaxial experiments combined with X-ray CT scanning, discovered that "
            "bedding angles near 45° are most conducive to forming complex hydraulic fracture networks [34].",

            "Saber et al. (2023), using XFEM with cohesive zone modeling (CZM) for transversely isotropic "
            "shale formations, found that the transversely isotropic model predicts narrower and longer "
            "fractures compared to the isotropic model; as bedding dip increases from 0° to 90°, fracture "
            "width increases while fracture length follows a non-monotonic 'increase-peak-decrease' trend [35]. "
            "Yang et al. (2023), conducting true triaxial fracturing experiments on shale-sandstone "
            "interbedded specimens with acoustic emission monitoring and real-time pressure recording, "
            "revealed that a large vertical stress difference is the key condition for fractures penetrating "
            "multiple lithologic layers [36]."
        ]),

        ("3.3 Natural Fracture Networks and Multi-Fracture Competitive Growth", [
            "The pervasive natural fracture systems in shale reservoirs are the primary intrinsic factor "
            "controlling hydraulic fracture network development. Wu et al. (2022), using a true triaxial "
            "experimental system with acoustic emission monitoring, demonstrated that fracture interference "
            "initiates at the inception stage and severely affects multi-fracture propagation; bedding "
            "interfaces and laminations are the main causes of fracture interference, hindering fracture "
            "height growth; and shear-type AE events signal interference caused by bedding slip and "
            "lamination-induced fracture deflection [37].",

            "For natural fracture stability evolution, recent studies have established a three-dimensional "
            "geomechanical-discrete fracture network (DFN) coupled evaluation method, revealing that fluid "
            "injection induces natural fracture instability through two mechanisms—'effective normal stress "
            "reduction' and 'frictional strength decrease'—while long-term production-induced pore pressure "
            "decline can enhance effective normal stress and globally improve fracture stability [38]. "
            "Wang et al. (2024) developed a hydro-mechanical-damage (HMD) coupled phase-field model, "
            "systematically investigating HF-NF interaction across approach angle, stress difference, and "
            "cementation strength, proposing a three-category classification (θ=α, θ<α, θ>α), and finding "
            "that the highest hydraulic fracture fractal dimension (2.1280) occurs at ±45° NF angles [39].",

            "In multi-fracture competitive propagation, fractures within a stage exhibit non-uniform growth "
            "due to stress shadowing and uneven flow distribution. Ping et al. (2024) showed that higher "
            "pumping rates accelerate dominant fracture development at the expense of peripheral fractures, "
            "necessitating balanced optimization of cluster spacing, injection parameters, and perforation "
            "design to achieve uniform propagation and maximize the stimulated reservoir volume (SRV) [18]."
        ])
    ]),

    ("4 Coupling Mechanisms: Stress Evolution and Fracture Propagation", [
        ("4.1 Theoretical Framework of Hydro-Mechanical Coupling", [
            "The coupling between in-situ stress evolution and fracture propagation fundamentally involves "
            "the interaction of three physical processes: fluid flow, solid deformation, and fracture "
            "damage/propagation. The core mechanism operates as a closed feedback loop: fluid pressure "
            "drives fracture propagation and modifies the pore pressure field; fracture opening/closure "
            "alters rock permeability and stiffness; and stress field redistribution, in turn, influences "
            "fracture propagation directions and fluid flow paths [6, 40].",

            "A 2024 review in the Journal of Engineering Geology systematically analyzed the fluid-solid "
            "coupling mechanical processes in shale reservoir fracturing, categorizing numerical approaches "
            "into two dimensions: dynamic simulation of complex fracture network evolution and reservoir-scale "
            "fracturing simulation. Among coupling strategies, fully coupled methods simultaneously solve "
            "all governing equations—theoretically most accurate but computationally prohibitive—while "
            "sequential (iterative) coupling solves each physical field successively with iterative "
            "convergence, achieving a practical balance between accuracy and efficiency and representing "
            "the most widely adopted framework [6].",

            "Deepening mechanistic understanding, studies on natural fracture stability evolution in "
            "moderately deep shale reservoirs have highlighted the crucial role of time-dependent processes. "
            "Pore pressure diffusion induced by fracturing fluid injection is not instantaneous but develops "
            "progressively in both time and space, leading to spatiotemporally dynamic variations in "
            "effective stress and natural fracture stability [38]. Furthermore, the dynamic variation of "
            "fracture permeability with effective stress (stress-sensitivity effect) constitutes a critical "
            "nonlinear feedback channel between the mechanical and flow fields [41]."
        ]),

        ("4.2 Production-Induced Stress Changes and Refracturing", [
            "Pore pressure depletion during production causes significant changes in the in-situ stress "
            "field, with decisive implications for refracturing operations and infill well fracture "
            "propagation. Agrawal and Sharma (2018) demonstrated that tensile regions formed between parent "
            "well fractures during production cause asymmetric child well fracture growth into depleted "
            "regions [26]. Yu et al. (2025), using a multi-scale fracture coupled flow-geomechanics "
            "simulation, found that near-wellbore fracture aperture can decrease by up to 25.1% and "
            "permeability by up to 48.9% during production, with coupled simulations predicting 20.3% "
            "lower production than uncoupled simulations [41].",

            "A 2024 coupled wellbore-reservoir-geomechanical modeling study under heterogeneous hydraulic "
            "fracture conditions revealed productivity competition between dominant and inferior fractures "
            "driven by uneven stress evolution and depletion, further exacerbating non-uniform flow "
            "distribution; increasing cluster spacing was found to promote more uniform stress state "
            "changes [42].",

            "For refracturing, Guo et al. (2024), through CT-based true triaxial refracturing experiments, "
            "demonstrated that natural fractures and stress state are the primary factors governing "
            "refracture propagation; temporary blocking agents forming large 'sheet-like' barriers in old "
            "fractures are more effective than tip plugging for initiating new fractures; and low "
            "horizontal stress difference conditions favor the creation of more complex refracture networks [43]."
        ]),

        ("4.3 Fault Activation and Casing Deformation", [
            "Fault activation and casing deformation during shale gas fracturing have attracted considerable "
            "attention as manifestations of stress-fracture (fault) coupling. Fluid-solid coupling studies "
            "reveal that the pore pressure diffusion zone expands progressively with water injection, "
            "reducing effective stress near fault zones and elevating fault reactivation risk; volumetric "
            "strains on opposite sides of faults exhibit opposing trends (compression vs. extension) [44].",

            "A 2024 Scientific Reports study on dynamic mechanisms categorized casing deformation sheared "
            "by fault slip into three stages: trigger activation → accelerated slip → deceleration slip. "
            "Fault slip behavior is heavily influenced by the angle between fault strike and maximum in-situ "
            "stress direction; the dynamic model incorporates stress drop, friction coefficient changes, "
            "and casing/cement-sheath resistance [45]. These findings provide a theoretical basis for "
            "fault avoidance and casing damage prevention in horizontal well fracturing design."
        ])
    ]),

    ("5 Research Methods", [
        ("5.1 Numerical Simulation Methods", [
            "Numerical simulation constitutes the core approach for investigating stress-fracture coupling "
            "mechanisms. The principal numerical methods currently employed include:",

            "(1) Finite Element Method (FEM) with Cohesive Zone Model (CZM): FEM is the most widely "
            "applied continuum mechanics approach, employing cohesive elements to simulate fracture "
            "initiation and propagation with effective treatment of the nonlinear process zone near crack "
            "tips. Zhu et al. (2024) used FEM-CZM to simulate 3D competitive fracture propagation, "
            "revealing the dominant control of the maximum horizontal principal stress direction on "
            "fracture extension direction and length [46].",

            "(2) Extended Finite Element Method (XFEM): By enriching conventional FEM shape functions "
            "with discontinuous basis functions, XFEM avoids remeshing during fracture propagation, "
            "offering unique advantages for arbitrary-path fracture growth. Ping et al. (2024) and "
            "Han et al. (2023) extensively employed XFEM to investigate stress shadow effects and "
            "simultaneous multi-fracture propagation [18-19].",

            "(3) Discrete Element Method (DEM): Based on particle contact models to simulate rock "
            "fracture and fragmentation, DEM is inherently suited for handling multi-fracture initiation, "
            "intersection, and network formation. Chang et al. (2024) developed a fully coupled hydraulic "
            "fracturing DEM model, successfully simulating simultaneous multi-fracture propagation across "
            "regimes from toughness-dominated to viscosity-dominated [21].",

            "(4) Phase-Field Method (PFM): Rooted in variational principles and total energy minimization, "
            "PFM automatically tracks fracture paths through a continuous phase-field variable without "
            "requiring additional fracture criteria. Lu et al. (2023) provided a comprehensive review of "
            "PFM applications in hydraulic fracturing simulation, highlighting its inherent advantage "
            "for multi-physics coupling problems [47]. Wang et al. (2024) further developed an HMD "
            "coupled PFM model enabling full-process simulation of multi-cluster HF-NF interaction [39]. "
            "Yi et al. (2024) extended PFM to porous viscoelastic media, for the first time incorporating "
            "shale creep behavior into fracture propagation modeling [48].",

            "(5) Boundary Element Method (BEM) and Displacement Discontinuity Method (DDM): Requiring "
            "discretization only on boundaries, these methods achieve dimensionality reduction and are "
            "particularly suitable for infinite/semi-infinite domain problems, finding wide application "
            "in stress shadow calculations and fracture interaction analysis.",

            "(6) Adaptive FEM/DEM Coupling: This approach combines the strengths of both methods, "
            "employing DEM near fractures to capture breakage processes and FEM in far-field regions for "
            "computational efficiency. Song et al. (2024) and Ju et al. (2024) utilized the ELFEN TGR "
            "software-based adaptive FEM/DEM method incorporating the complete physics of hydraulic "
            "fractures, natural fractures, porous flow, and proppant transport [20, 49].",

            "(7) Multi-Scale Coupling Strategies: Fracture systems at different scales are represented "
            "using different methods—macro-scale hydraulic fractures via EDFM or XFEM explicit "
            "representation, meso- and micro-scale natural fractures via Multiple Interacting Continua "
            "(MINC) or dual-porosity/dual-permeability (DPDK) upscaling—with cross-scale two-way "
            "coupling achieved through computational homogenization or flexibility superposition "
            "methods [41, 50]."
        ]),

        ("5.2 Laboratory Experimental Methods", [
            "Laboratory experiments are indispensable for validating theoretical models and revealing "
            "physical mechanisms. The primary experimental methods include:",

            "(1) True Triaxial Hydraulic Fracturing Experiments: The laboratory method most closely "
            "approximating field conditions. Cubic rock specimens (typically 300 × 300 × 300 mm) are "
            "independently loaded in three orthogonal directions to simulate the true triaxial in-situ "
            "stress state while fracturing fluid is injected to drive fracture propagation [34, 36-37]. "
            "Wu et al. (2022) successfully captured the initiation and development of fracture interference "
            "in multi-cluster fracturing using a true triaxial system with AE monitoring [37]. Zhao et al. "
            "(2022) developed a high-stress true triaxial experimental system to study shale fracturing "
            "characteristics at different bedding angles under hydrostatic pressure [34].",

            "(2) Acoustic Emission (AE) Monitoring: By recording elastic wave signals released during "
            "rock fracture, AE enables real-time localization and mechanism discrimination of fracture "
            "initiation, propagation, and intersection. AE signal parameters (b-value, RA-AF value) "
            "can differentiate between tensile and shear fracture events [34, 37].",

            "(3) X-ray CT Scanning and 3D Reconstruction: Pre- and post-fracturing CT scanning of "
            "specimens enables quantitative characterization of fracture network geometry and complexity "
            "through 3D reconstruction and fractal dimension analysis [34, 43]. Guo et al. (2024) applied "
            "CT technology to investigate the effects of temporary blocking agents on refracture "
            "propagation [43].",

            "(4) Digital Image Correlation (DIC): By tracking the deformation of speckle patterns on "
            "specimen surfaces, DIC provides full-field displacement and strain distributions useful "
            "for analyzing strain concentration and damage evolution near fracture tips.",

            "(5) Laboratory-Scale DAS: A KAUST research team (2023) distributed DAS fibers over all "
            "six surfaces of a 50 cm³ cubic rock block under true triaxial stress, achieving ~10× higher "
            "sampling frequency than conventional DAS, with CT-imaging calibration of microseismic "
            "locations inverted from DAS data [51]."
        ]),

        ("5.3 Field Monitoring Methods", [
            "Field monitoring provides the most direct means of obtaining reservoir-scale information "
            "on stress evolution and fracture propagation. Current technologies include:",

            "(1) Microseismic Monitoring: The most mature and widely applied technique for hydraulic "
            "fracture characterization. By deploying geophone arrays at the surface or in offset wells, "
            "microseismic events generated by rock failure during fracturing are recorded and inverted "
            "to reconstruct the spatial distribution and activation sequence of fractures. The latest "
            "hybrid borehole DAS and three-component (3C) geophone microseismic imaging technique (2023, "
            "IEEE TGRS), employing geometric-mean reverse time migration, imaged 41 microearthquakes "
            "in the Montney Formation and explored joint DAS+3C focal mechanism estimation [52].",

            "(2) Distributed Acoustic Sensing (DAS): Utilizing optical fiber as a continuous sensor "
            "array, DAS enables real-time acoustic/strain measurements along the entire wellbore. "
            "Correa et al. (2024) at the Austin Chalk/Eagle Ford Field Laboratory, for the first time "
            "combined surface orbital vibrators (SOV) with DAS for continuous time-lapse VSP acquisition, "
            "capturing fracture opening and closure processes at hourly temporal resolution [53]. "
            "Glubokovskikh et al. (2024) further tracked the activation and growth evolution of individual "
            "natural fractures using continuous borehole DAS data [54].",

            "(3) Low-Frequency DAS (<0.05 Hz): A ConocoPhillips patent (2023) proposed using low-frequency "
            "DAS signals to constrain far-field fracture length, height, width, and density—detecting "
            "stress shadow-induced strain perturbations in horizontal offset wells and thermal signals "
            "from fracture hits in vertical wells [55].",

            "(4) Distributed Temperature Sensing (DTS): By monitoring temperature changes during "
            "fracturing and production, DTS enables inversion of per-cluster injection allocation and "
            "production contribution. ESG Solutions (2024) demonstrated an integrated DAS injection "
            "allocation + DTS temperature transient analysis + microseismic deformation/stress analysis "
            "workflow for comprehensive spatiotemporal fracture characterization in the Marcellus Shale [56]."
        ])
    ]),

    ("6 Major Research Achievements", [
        "Synthesizing the foregoing literature analysis, the principal achievements in understanding the "
        "coupling mechanisms between in-situ stress evolution and hydraulic fracture propagation in shale "
        "gas reservoirs can be summarized as follows:",

        "(1) Theoretical advances: A robust theoretical framework for stress evolution has been established "
        "on the foundation of Biot's poroelasticity, encompassing coupled fluid-flow, stress, and damage "
        "multi-field interactions. The applicability conditions of the effective stress principle in "
        "fractured shale reservoirs and the critical influence of Biot coefficient anisotropy on coupling "
        "behavior have been clarified [9-12].",

        "(2) Mechanistic insights: The multi-factor regulatory mechanisms of stress shadow effects have "
        "been elucidated; the three fundamental modes of HF-NF interaction and their controlling factors "
        "(approach angle, stress difference, interface strength) have been identified; and the conditions "
        "for production-induced stress reversal and its impact on child well fracture propagation have "
        "been quantified [15-21, 26, 30-32].",

        "(3) Methodological integration: A tripartite research paradigm integrating numerical simulation "
        "(core), laboratory experimentation (validation), and field monitoring (constraint) has been "
        "established. Numerical methods are evolving from single-method to multi-method coupling "
        "(FEM-DEM, XFEM-EDFM) and multi-scale hierarchical frameworks (micro-DEM → macro-EDFM); "
        "experimentally, the true triaxial-AE-CT-DIC multi-modal characterization has become standard "
        "practice; in field monitoring, multi-physical parameter synergistic monitoring (DAS+DTS+"
        "microseismic) is becoming the new technical standard [22-56].",

        "(4) Engineering applications: Stress-evolution-based optimization principles for infill well "
        "deployment, stress-reversal-informed refracturing timing selection methods, and stress shadow "
        "mitigation strategies centered on increasing cluster spacing and optimizing injection sequences "
        "have been proposed. In geo-engineering integration, the development of integrated fracturing-"
        "production numerical simulators is being actively explored [22, 41-43]."
    ]),

    ("7 Current Limitations and Challenges", [
        "Despite these substantial achievements, several critical limitations and challenges remain:",

        "(1) The multi-scale coupling gap: The spatial scale spans over ten orders of magnitude from "
        "nanometer-scale pores to kilometer-scale reservoirs, yet effective cross-scale bridging "
        "mechanisms remain underdeveloped. Information transfer between microscale experiments (e.g., "
        "nanoindentation, molecular dynamics) and macroscale numerical simulations still predominantly "
        "relies on empirical relationships, lacking rigorous physics-based homogenization theories [6, 50].",

        "(2) Incomplete multi-physics coupling: Existing models predominantly focus on hydro-mechanical "
        "(HM) coupling, with inadequate consideration of thermal (THM) and chemical (THMC) effects. "
        "The chemical interaction between shale and water-based fracturing fluids—including clay hydration "
        "swelling, osmotic effects, and ion exchange—and its influence on fracture surface mechanical "
        "properties and local stress fields remain insufficiently characterized [13-14].",

        "(3) Natural fracture characterization uncertainty: The geometric distribution (density, "
        "orientation, size, aperture) and mechanical properties (cohesion, friction coefficient, "
        "stiffness) of natural fractures cannot be precisely determined across the reservoir, leading "
        "to substantial predictive uncertainty in DFN-based models [32, 38].",

        "(4) Computational efficiency versus model fidelity: Fully coupled, high-resolution numerical "
        "models incur prohibitive computational costs, making them unsuitable for rapid iterative "
        "optimization at the engineering scale. While the phase-field method offers distinct advantages "
        "in handling complex fracture topologies, its computational cost remains the primary bottleneck "
        "limiting its adoption in engineering practice [47-48].",

        "(5) Inadequate treatment of time-dependent effects: The long-term influence of time-dependent "
        "processes—including shale creep/relaxation behavior, time-varying fracturing fluid rheology, "
        "and time-dependent proppant embedment and crushing—on stress-fracture coupling lacks systematic "
        "investigation [48].",

        "(6) Limited physical interpretability of data-driven methods: Machine learning approaches have "
        "demonstrated immense potential for fracture propagation prediction and parameter inversion, "
        "yet most models remain 'black boxes' with limited physical interpretability, constraining "
        "their credibility in engineering decision-making [57-58].",

        "(7) Scale effects and in-situ representativeness of experiments: The extrapolation of laboratory-"
        "scale true triaxial experimental results to the reservoir scale still lacks a reliable theoretical "
        "foundation; size effects and time-scale effects introduce substantial biases in direct upscaling "
        "[34, 37]."
    ]),

    ("8 Future Research Directions", [
        "Based on the foregoing analysis, we propose the following future research directions:",

        "(1) Hierarchical multi-scale integrated modeling: Construct a hierarchical coupling framework "
        "spanning from molecular scale (molecular dynamics) → pore scale (LBM/PNM) → core scale "
        "(DEM/FEM) → reservoir scale (EDFM/equivalent continuum), developing cross-scale information "
        "transfer methods based on computational homogenization and machine learning surrogate models "
        "to enable macroscale engineering simulations constrained by microscale mechanisms [41, 50-51].",

        "(2) Physics-informed neural networks (PINNs) and data-physics hybrid modeling: Embed physical "
        "constraints such as Biot's poroelasticity theory and fracture mechanics criteria into deep "
        "learning network architectures, developing surrogate models that combine physical interpretability "
        "with computational efficiency for real-time prediction of fracture propagation and stress "
        "evolution. The application of explainable AI tools such as SHAP will facilitate the discovery "
        "of new physical insights from data [57-58].",

        "(3) Fully coupled THMC multi-field simulation: Extend existing HM coupling to integrate thermal "
        "(THM) and chemical (THMC) fields, establishing fully coupled models that account for heat "
        "conduction, thermal convection, reactive-transport, and phase change processes to better "
        "describe stress-fracture coupling behavior in deep high-temperature environments with complex "
        "fracturing fluid systems [13-14, 48].",

        "(4) Digital twin-driven geo-engineering integration: Integrate real-time DAS/DTS/microseismic "
        "monitoring data, numerical simulations, and history matching to construct digital twin systems "
        "for hydraulic fracturing, enabling real-time fracture propagation inversion, dynamic stress "
        "field updating, and online fracturing parameter optimization [53-56].",

        "(5) Time-dependent constitutive models and long-term stability assessment: Develop visco-"
        "elastoplastic constitutive models incorporating shale creep and damage evolution, systematically "
        "investigate the influence of time-dependent deformation on long-term fracture conductivity and "
        "stress field evolution, and provide a theoretical foundation for full-lifecycle geomechanical "
        "stability assessment of shale gas wells [48].",

        "(6) Statistical mechanics description of fracture network systems: Drawing on statistical "
        "physics approaches, establish a statistical characterization framework for fracture networks, "
        "describing the co-evolution of stress fields and fracture networks at the system level (rather "
        "than the individual fracture level), and revealing self-organized criticality and emergent "
        "behaviors of fracture network systems.",

        "(7) Multi-well platform- and block-scale stress field synergistic management: Extend research "
        "from single-well and well-pad scales to platform and block scales, investigating regional "
        "stress field co-evolution under multi-well simultaneous/sequential fracturing and co-production "
        "conditions, and developing stress-management-based well pattern, spacing optimization, and "
        "development planning methods [24-25, 42].",

        "(8) Low-carbon-oriented stress-fracture coupling research: Against the backdrop of carbon "
        "neutrality goals, investigate stress-fracture coupling mechanisms in CO₂ fracturing and "
        "supercritical CO₂ fracturing, and explore synergistic optimization strategies for CO₂ "
        "geological sequestration and shale gas stimulation."
    ]),

    ("9 Conclusions", [
        "The coupling mechanism between in-situ stress evolution and fracture propagation in shale gas "
        "reservoirs represents a core scientific challenge in hydraulic fracturing theory, spanning "
        "multiple disciplines including rock mechanics, fluid mechanics, fracture mechanics, and heat "
        "and mass transfer. This review has systematically synthesized the current state of research, "
        "leading to the following principal conclusions:",

        "(1) Stress evolution research, founded on Biot's poroelasticity theory, has established a "
        "relatively comprehensive theoretical framework, with stress shadow effects, production-induced "
        "stress reversal, and 4D dynamic stress field evolution constituting the three core research "
        "themes in this domain.",

        "(2) Hydraulic fracture propagation research has progressed from single homogeneous fracture "
        "models to complex fracture network models incorporating bedding, natural fractures, and "
        "anisotropy, with the three-mode HF-NF interaction classification receiving broad experimental "
        "and numerical validation.",

        "(3) The tripartite methodology of numerical simulation, laboratory experimentation, and field "
        "monitoring has established a complementary research paradigm, although integrated synergy among "
        "these methods requires further strengthening. The phase-field method, adaptive FEM/DEM, and "
        "multi-scale coupling strategies represent the forefront of numerical approaches; true triaxial-"
        "AE-CT joint experimentation and multi-parameter DAS monitoring represent the development trends "
        "in experimental and monitoring technologies.",

        "(4) Current research faces significant challenges in multi-scale coupling, multi-physics "
        "completeness, natural fracture characterization uncertainty, computational efficiency, time-"
        "dependent effects, and the physical interpretability of data-driven methods.",

        "(5) Future research should prioritize hierarchical multi-scale modeling, physics-informed "
        "neural networks, fully coupled THMC simulation, digital twin systems, long-term time-dependent "
        "behavior, and multi-well platform synergistic stress management to advance both the fundamental "
        "theory of shale gas fracturing and the optimization of engineering practice."
    ])
]

EN_REFERENCES = [
    "[1] Maxwell S C. Microseismic Imaging of Hydraulic Fracturing: Improved Engineering of Unconventional Shale Reservoirs[M]. Society of Exploration Geophysicists, 2014.",
    "[2] Molenaar M M, Hill D J, Webster P, et al. First downhole application of distributed acoustic sensing for hydraulic-fracturing monitoring and diagnostics[J]. SPE Drilling & Completion, 2012, 27(01): 32-38.",
    "[3] Jin L, Roy B. Hydraulic fracturing induced seismicity[J]. Earth-Science Reviews, 2017, 174: 52-68.",
    "[4] Zhao J Z, Ren L, Shen C, et al. Research progress on fracture network fracturing theory and technology for shale gas reservoirs[J]. Natural Gas Industry, 2018, 38(3): 1-14.",
    "[5] Zou C N, Dong D Z, Wang Y M, et al. Shale gas in China: Characteristics, challenges and prospects[J]. Petroleum Exploration and Development, 2015, 42(6): 689-701.",
    "[6] A Review on hydro-mechanical coupling simulations of hydraulic fracture network in shale reservoirs[J]. Journal of Engineering Geology, 2024, 32(4): 1381-1396.",
    "[7] Guo J C, Gou B, Ren S, et al. Progress and prospect of deep shale gas reservoir fracturing technology[J]. Natural Gas Industry, 2022, 42(6): 67-82.",
    "[8] Chen M, Jin Y, Zhang G Q. Petroleum Engineering Rock Mechanics[M]. Beijing: Science Press, 2008.",
    "[9] Biot M A. General theory of three-dimensional consolidation[J]. Journal of Applied Physics, 1941, 12(2): 155-164.",
    "[10] Cheng A H D. Poroelasticity[M]. Springer International Publishing, 2016.",
    "[11] Suarez-Rivera R, Fjær E. Evaluating the poroelastic effect on anisotropic, organic-rich, mudstone systems[J]. Rock Mechanics and Rock Engineering, 2013, 46(3): 569-580.",
    "[12] Wang Q, Wang Y, et al. Evolution law of stress induced by pressure depletion in fractured shale reservoirs: Implications for subsequent refracturing and infill well development[J]. Petroleum, 2025, 11(1): 71-83.",
    "[13] Yi D, Yi L, Yang Z, et al. Coupled thermo-hydro-mechanical-phase field modelling for hydraulic fracturing in thermo-poroelastic media[J]. Computers and Geotechnics, 2024.",
    "[14] Yi D, Yang Z, Yi L, et al. Hydraulic fracturing phase-field model in porous viscoelastic media[J]. International Journal of Mechanical Sciences, 2024.",
    "[15] Roussel N P, Sharma M M. Optimizing fracture spacing and sequencing in horizontal-well fracturing[J]. SPE Production & Operations, 2011, 26(02): 173-184.",
    "[16] Olson J E. Multi-fracture propagation modeling: Applications to hydraulic fracturing in shales and tight gas sands[C]. ARMA, 2008.",
    "[17] Bunger A P, Zhang X, Jeffrey R G. Parameters affecting the interaction among closely spaced hydraulic fractures[J]. SPE Journal, 2012, 17(01): 292-306.",
    "[18] Ping Y, et al. Numerical simulation of the simultaneous development of multiple fractures in horizontal wells based on the extended finite element method[J]. Energies, 2024, 17(5): 1057.",
    "[19] Han G, Cui Z, Zhu H. The effect of perforation spacing on the variation of stress shadow[C]. LAPSE, 2023.",
    "[20] Ju Y, Wang Y, et al. Stress shadow effects in multistage horizontal hydrofracturing of tight reservoirs: a numerical analysis considering perforation cluster spacings and fracturing sequences[J]. Geomechanics and Geophysics for Geo-Energy and Geo-Resources, 2024, 10: 183.",
    "[21] Chang X, Hou B, Ding Y. DEM modeling of simultaneous propagation of multiple hydraulic fractures across different regimes, from toughness- to viscosity-dominated[J]. Rock Mechanics and Rock Engineering, 2024, 57: 481-503.",
    "[22] Zhu H Y, Tang X H, et al. Research progress on 4D in-situ stress evolution and complex fracture propagation in infill wells of shale gas reservoirs[J]. Petroleum Science Bulletin, 2021, 6(3): 455-475.",
    "[23] Guo J C, et al. Key issues and explorations in shale gas fracturing[J]. Natural Gas Industry B, 2023, 10(2): 183-197.",
    "[24] Xu E, et al. Study on the evolution law of four-dimensional dynamic stress fields in fracturing of deep shale gas platform wells[J]. Processes, 2025, 13(9): 2709.",
    "[25] Ruan Q, Zhang L, Zhao Y, et al. Influence of natural weak planes on 4D stress and infill well fracture disturbance in shale gas[J]. Journal of Geomechanics, 2026, 32(1).",
    "[26] Agrawal S, Sharma M M. Impact of pore pressure depletion on stress reorientation and its implications on the growth of child well fractures[C]. URTeC, 2018: 2875375.",
    "[27] Adachi J, Siebrits E, Peirce A, et al. Computer simulation of hydraulic fractures[J]. International Journal of Rock Mechanics and Mining Sciences, 2007, 44(5): 739-757.",
    "[28] Detournay E. Mechanics of hydraulic fractures[J]. Annual Review of Fluid Mechanics, 2016, 48: 311-339.",
    "[29] Gong X, Ma X H, Liu Y Y, et al. Advances in hydraulic fracture propagation research in shale reservoirs[J]. Minerals, 2022, 12(11): 1438.",
    "[30] Blanton T L. An experimental study of interaction between hydraulically induced and pre-existing fractures[C]. SPE/DOE, 1982: 10847.",
    "[31] Warpinski N R, Teufel L W. Influence of geologic discontinuities on hydraulic fracture propagation[J]. Journal of Petroleum Technology, 1987, 39(02): 209-220.",
    "[32] Influence of natural fractures on stimulated reservoir volume fracturing networks in deep shale reservoirs, southern Sichuan Basin[J]. Daqing Petroleum Geology and Development, 2025, 44(3).",
    "[33] Study on hydraulic fracture propagation law in layered rock masses[J]. China Coal Geology, 2025, 37(8).",
    "[34] Zhao H, Liang B, Sun W J, et al. Effects of hydrostatic pressure on hydraulic fracturing properties of shale using X-ray computed tomography and acoustic emission[J]. Journal of Petroleum Science and Engineering, 2022, 219: 110684.",
    "[35] Saber A, et al. Propagation of multiple hydraulic fractures in a transversely isotropic shale formation[C]. SSRN, 2023: 4334460.",
    "[36] Yang L, Sheng X, Zhang B, et al. Propagation behavior of hydraulic fractures in shale under triaxial compression considering the influence of sandstone layers[J]. Gas Science and Engineering, 2023.",
    "[37] Wu S, Gao K, Wang X Q, et al. Investigating the propagation of multiple hydraulic fractures in shale oil rocks using acoustic emission[J]. Rock Mechanics and Rock Engineering, 2022, 55: 6015-6032.",
    "[38] Study on the evolution law of natural fracture stability in moderately deep shale reservoirs[J]. Petroleum Science Bulletin, 2026, 11(1).",
    "[39] Wang Y J, Wang B, Su H, et al. Mechanisms of fracture propagation from multi-cluster using a phase field based HMD coupling model in fractured reservoir[J]. Petroleum Science, 2024, 21(3): 1829-1851.",
    "[40] Review on hydro-mechanical coupling numerical simulation and fracture propagation in shale reservoir fracturing[J]. Journal of Engineering Geology, 2024, 32(4): 1381-1396.",
    "[41] Yu D, Kang W, et al. A novel coupled flow and geomechanics simulation method employing multi-scale fracture model[C]. IFEDC, 2024.",
    "[42] Coupled wellbore-reservoir-geomechanical modelling of uneven depletion and induced stress responses in shale oil reservoirs with heterogeneous hydraulic fractures[C]. ARMA-IGS, 2024: 0343.",
    "[43] Guo P, et al. Study on fracture propagation rules of shale refracturing based on CT technology[J]. Processes, 2024, 12(1): 131.",
    "[44] Study on faults instability law in hydraulic fracturing based on fluid-solid coupling method[J]. Journal of Engineering Geology, 2024, 32(4): 1439-1446.",
    "[45] Investigation on dynamic mechanism of fault slip and casing deformation during multi-fracturing in shale gas wells[J]. Scientific Reports, 2024, 14: 13164.",
    "[46] Zhu H Y, Tang X H, et al. Dynamic evolution mechanism of the fracturing fracture system: enlightenments from hydraulic fracturing physical experiments and finite element numerical simulation[J]. Petroleum Science, 2024, 21(6): 3839-3866.",
    "[47] Lu Q L, Zhang H, Guo J C, et al. Status and prospect of hydraulic fracture propagation simulation based on phase field method[J]. Natural Gas Industry, 2023, 43(3): 59-68.",
    "[48] Yi D, Yang Z, Yi L, et al. Hydraulic fracturing phase-field model in porous viscoelastic media[J]. International Journal of Mechanical Sciences, 2024.",
    "[49] Song X, et al. Numerical analysis of stress shadowing effects in multistage hydraulic fracturing[J]. Journal of Mining Science and Technology, 2024, 9(4).",
    "[50] Wang C. A multi-scale, multi-continuum and multi-physics model to simulate coupled fluid flow and geomechanics in shale gas reservoirs[D]. Colorado School of Mines, 2018.",
    "[51] KAUST. Distributed acoustic sensing (DAS) for hydraulic fracture monitoring in laboratory scale[C]. 3rd EAGE Workshop on Fiber Optic Sensing for Energy Applications, Chengdu, 2023.",
    "[52] Characterizing microearthquakes induced by hydraulic fracturing with hybrid borehole DAS and three-component geophone data[J]. IEEE Transactions on Geoscience and Remote Sensing, 2023, 61: 1-15.",
    "[53] Correa J, Glubokovskikh S, Nayak A, et al. Revealing complex subsurface dynamics with continuous seismic monitoring: Observations using distributed acoustic sensing and surface orbital vibrators during hydraulic fracturing[J]. Geophysics, 2024, 89(6).",
    "[54] Glubokovskikh S, Correa J, Ajo-Franklin J, et al. Continuous surface-to-distributed acoustic sensor snapshots explain reactivation of individual natural fractures during an unconventional reservoir stimulation[J]. Geophysics, 2024, 89(6).",
    "[55] ConocoPhillips Company. Low frequency distributed acoustic sensing hydraulic fracture geometry[P]. US Patent: 2023/0003119 A1, 2023.",
    "[56] ESG Solutions. Fibre-optic sensing and microseismic monitoring evaluate and enhance hydraulic fracturing via real-time and post-treatment analysis[R]. Marcellus Shale Case Study, 2024.",
    "[57] Qu H Y, Zhang J L, Zhou F J, et al. Evaluation of hydraulic fracturing based on deep neural network with physical constraints[J]. Petroleum Science, 2023, 20(2): 1129-1141.",
    "[58] Shentu J J, Lin B T, Jin Y, et al. Interpretable machine learning for hydraulic fracture propagation in conglomerate rock based on discrete element method[J]. Acta Geotechnica, 2024."
]


def set_run_font(run, font_cn='宋体', font_en='Times New Roman', font_size=None, bold=None):
    """Set font for a run with proper CJK font handling via XML.

    In python-docx, run.font.name only sets the Western (ascii/hAnsi) font.
    For Chinese characters, we must set w:eastAsia in the XML rFonts element.
    """
    # Set Western font (affects ASCII, numbers, Latin text)
    run.font.name = font_en
    # Manipulate XML for CJK font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:cs'), font_en)  # Complex script fallback

    if font_size is not None:
        run.font.size = Pt(font_size)
    if bold is not None:
        run.bold = bold


def add_cjk_paragraph(doc, text, font_cn='宋体', font_en='Times New Roman',
                       font_size=12, bold=False, alignment=None,
                       space_before=None, space_after=None,
                       first_line_indent=None, line_spacing=1.5):
    """Add a paragraph with proper CJK font handling."""
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        set_run_font(run, font_cn=font_cn, font_en=font_en,
                     font_size=font_size, bold=bold)
    if alignment is not None:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    if line_spacing is not None:
        p.paragraph_format.line_spacing = line_spacing
    return p


def build_document(title_text, abstract_text, keywords_text, sections, references,
                   font_cn='宋体', font_en='Times New Roman', font_title='黑体',
                   is_chinese=True):
    """Build a complete Word document with proper CJK font handling.

    Chinese academic standard (GB/T 7714 compatible):
    - Title: 黑体 16pt, centered, bold
    - Level-1 heading: 黑体 14pt, bold
    - Level-2 heading: 黑体 13pt, bold
    - Body text: 宋体 12pt, justified, 1.5 line spacing, 0.74cm first-line indent
    - Abstract/Keywords: 宋体 10.5pt
    - References: 宋体 9pt, 1.25 line spacing (single)
    - English/numerals within Chinese text: Times New Roman
    """

    doc = Document()

    # ---- Set default style ----
    style = doc.styles['Normal']
    style.font.name = font_en
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    # Set CJK default via XML
    style.element.get_or_add_rPr()
    rPr = style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        style.element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)
    rFonts.set(qn('w:eastAsia'), font_cn)

    # ---- Page setup ----
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # ---- Title ----
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(16)
    title_para.paragraph_format.line_spacing = 1.5
    title_run = title_para.add_run(title_text)
    title_font_cn = font_title if is_chinese else font_en
    title_font_en = font_en
    set_run_font(title_run, font_cn=title_font_cn, font_en=title_font_en,
                 font_size=16, bold=True)

    # ---- Labels ----
    if is_chinese:
        abs_label = '摘  要'
        kw_label = '关键词'
        ref_label = '参考文献'
    else:
        abs_label = 'Abstract'
        kw_label = 'Keywords'
        ref_label = 'References'

    # ---- Abstract heading ----
    add_cjk_paragraph(doc, abs_label,
                      font_cn=font_title if is_chinese else font_en,
                      font_en=font_en, font_size=12, bold=True,
                      alignment=WD_ALIGN_PARAGRAPH.LEFT,
                      space_before=6, space_after=6)

    # ---- Abstract content ----
    add_cjk_paragraph(doc, abstract_text,
                      font_cn=font_cn, font_en=font_en, font_size=10.5, bold=False,
                      alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                      first_line_indent=0.74, line_spacing=1.5)

    # ---- Keywords ----
    kw_para = doc.add_paragraph()
    kw_para.paragraph_format.space_before = Pt(6)
    kw_para.paragraph_format.space_after = Pt(12)
    kw_label_run = kw_para.add_run(f'{kw_label}：')
    set_run_font(kw_label_run, font_cn=font_title if is_chinese else font_en,
                 font_en=font_en, font_size=10.5, bold=True)
    kw_content_run = kw_para.add_run(keywords_text)
    set_run_font(kw_content_run, font_cn=font_cn, font_en=font_en, font_size=10.5, bold=False)

    # ---- Body sections ----
    def render_section(sec_list, level=1):
        for item in sec_list:
            if isinstance(item, tuple) and len(item) == 2:
                heading_text, content = item

                if content and isinstance(content[0], tuple):
                    # Heading with sub-subsections
                    hn = '黑体' if is_chinese else font_en
                    fs = {1: 14, 2: 13, 3: 12}.get(level, 12)
                    add_cjk_paragraph(doc, heading_text,
                                      font_cn=hn, font_en=font_en,
                                      font_size=fs, bold=True,
                                      space_before=12, space_after=6)
                    render_section(content, level + 1)
                else:
                    # Heading with content paragraphs
                    hn = '黑体' if is_chinese else font_en
                    fs = {1: 14, 2: 13, 3: 12}.get(level, 12)
                    add_cjk_paragraph(doc, heading_text,
                                      font_cn=hn, font_en=font_en,
                                      font_size=fs, bold=True,
                                      space_before=12, space_after=6)

                    for para_text in content:
                        add_cjk_paragraph(doc, para_text,
                                          font_cn=font_cn, font_en=font_en,
                                          font_size=12, bold=False,
                                          alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                          first_line_indent=0.74, line_spacing=1.5)

    render_section(sections)

    # ---- References ----
    doc.add_page_break()
    add_cjk_paragraph(doc, ref_label,
                      font_cn='黑体' if is_chinese else font_en,
                      font_en=font_en, font_size=14, bold=True,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      space_after=12)

    for ref in references:
        add_cjk_paragraph(doc, ref,
                          font_cn=font_cn, font_en=font_en,
                          font_size=9, bold=False,
                          alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          line_spacing=1.25, space_after=2)

    return doc


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # ---- Chinese document ----
    print("Generating Chinese literature review...")
    cn_doc = build_document(
        title_text=CN_TITLE,
        abstract_text=CN_ABSTRACT,
        keywords_text=CN_KEYWORDS,
        sections=CN_SECTIONS,
        references=CN_REFERENCES,
        is_chinese=True
    )
    cn_path = os.path.join(OUTPUT_DIR, "页岩气藏地应力演化与裂缝扩展耦合机制_文献综述_中文.docx")
    cn_doc.save(cn_path)
    print(f"  Saved: {cn_path}")

    # ---- English document ----
    print("Generating English literature review...")
    en_doc = build_document(
        title_text=EN_TITLE,
        abstract_text=EN_ABSTRACT,
        keywords_text=EN_KEYWORDS,
        sections=EN_SECTIONS,
        references=EN_REFERENCES,
        is_chinese=False
    )
    en_path = os.path.join(OUTPUT_DIR, "Coupling_Mechanisms_Stress_Evolution_Fracture_Propagation_Shale_Gas_Review_EN.docx")
    en_doc.save(en_path)
    print(f"  Saved: {en_path}")

    print("\nDone! Both documents have been generated.")


if __name__ == "__main__":
    main()