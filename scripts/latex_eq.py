# -*- coding: utf-8 -*-
"""Insert LaTeX-formatted equations for Word Alt+= conversion."""
from docx import Document

DST = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP_EN_final.docx'
doc = Document(DST)

def sp(idx, text):
    p = doc.paragraphs[idx]
    for r in p.runs: r.text = ''
    if p.runs: p.runs[0].text = text
    else: p.add_run(text)

# [82] ADE derivation with LaTeX
sp(82,
    "To quantitatively interpret the breakthrough curve, a piecewise advection-dispersion "
    "model with smooth transition was developed. The model is based on the one-dimensional ADE:\n"
    "    [LaTeX: \\partial C/\\partial t + v \\cdot \\partial C/\\partial x = D \\cdot \\partial^2 C/\\partial x^2]\n"
    "where v = 4Q/(\\pi d^2) is the mean flow velocity and D = \\alpha v is the longitudinal "
    "dispersion coefficient. The rising phase adopts the ADE pulse solution (Gaussian component), "
    "while the tailing phase uses the continuous-source boundary solution (erfc component). "
    "The two regimes are blended via a hyperbolic tangent weighting function "
    "w(t) = 0.5[1 + \\tanh((t_0 - t)/\\sigma)], producing a smooth, physically continuous "
    "breakthrough curve defined by Eq. (6):")

# [83] Equation (6a-6c) in LaTeX
sp(83,
    "[LaTeX Eq.6:  C(t)=w(t)C_{\\rm rise}(t)+[1-w(t)]C_{\\rm fall}(t)]\n"
    "\n"
    "[LaTeX Eq.6a: C_{\\rm rise}(t)=c_b+\\frac{A d}{\\sqrt{16\\alpha Q t}}"
    "\\exp\\!\\left[-\\frac{(x\\pi d^2-4Qt)^2}{16\\alpha Q t\\pi d^2}\\right]]\n"
    "\n"
    "[LaTeX Eq.6b: C_{\\rm fall}(t)=c_b+\\frac{a}{2}"
    "\\mathrm{erfc}\\!\\left[\\frac{4Qt-x\\pi d^2}{\\sqrt{16\\alpha Q t\\pi d^2}}\\right]]\n"
    "\n"
    "[LaTeX Eq.6c: w(t)=\\frac{1}{2}\\!\\left[1+\\tanh\\!\\left(\\frac{t_0-t}{\\sigma}\\right)\\right]]")

doc.save(DST)
print("LaTeX equations inserted. In Word, select each [LaTeX ...] block and press Alt+= to convert.")