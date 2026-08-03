# -*- coding: utf-8 -*-
"""
将 Markdown 实验报告转换为格式化的 DOCX 文件，适合导师和甲方汇报。
"""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def add_heading_cn(doc, text, level):
    """添加中文标题并设置字体"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(doc, text, bold=False, size=11, alignment=None, color=None):
    """添加段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    return p

def add_bullet(doc, text, level=0):
    """添加列表项"""
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(10.5)
    return p

def add_table_from_data(doc, headers, rows, col_widths=None):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(10)
        run.bold = True

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(10)

    doc.add_paragraph()  # spacing after table
    return table

def main():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # === Title ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('荧光压裂液体系构建与性能研究')
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('实验进展报告')
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('汇报对象：导师 / 甲方项目负责人　　　汇报日期：2026年6月8日\n研究生：郝乐乐　　　导师：李娜　　　成都理工大学能源学院')
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # === 1. Project Overview ===
    add_heading_cn(doc, '一、项目概要与技术定位', level=1)

    add_heading_cn(doc, '1.1 项目背景与核心矛盾', level=2)
    add_para(doc, '水力压裂是页岩气商业开发的核心技术，压裂裂缝的延伸范围直接决定储层改造效果（SRV）。然而，现有裂缝监测技术存在根本性的方法论局限：')
    add_bullet(doc, '微地震监测：依赖于速度模型对岩石破裂声发射信号的反演，深层定位误差通常数十米量级')
    add_bullet(doc, '分布式光纤传感（DAS/DTS）：感知范围局限于光纤近场区域，对远场裂缝几何缺乏直接约束')
    add_bullet(doc, '化学示踪剂：依赖返排液浓度-时间曲线反推裂缝贡献，定量分析假设在非均质裂缝网络中难以严格成立')
    add_para(doc, '上述技术均属于间接测量范畴——测量的是裂缝产生的物理效应，而非裂缝本身的空间坐标。从物理效应反推实体几何，必然需要借助反演模型与先验假设，多解性问题无法在间接测量范式内得到根本性消除。', bold=True)
    add_para(doc, '核心矛盾：压裂工程界需要"看到"裂缝，而现有方法只能"推断裂缝"。压后取心是唯一能提供实物验证的手段，但常规压裂液在岩心上不留下可辨识的宏观标记——间接监测与实物验证之间存在结构性的"实物证据断层"。')

    add_heading_cn(doc, '1.2 本项目解决方案', level=2)
    add_para(doc, '本项目研制一种与现有压裂工艺完全兼容的荧光压裂液体系，使压后取心岩样在紫外光照射下能够直接呈现压裂液的实际波及范围。')

    add_table_from_data(doc,
        ['要素', '方案'],
        [
            ['示踪介质', '无机稀土铝酸盐长余辉荧光粉 SrAl₂O₄:Eu²⁺,Dy³⁺'],
            ['载体流体', '羟丙基胍胶（HPG）压裂液'],
            ['检测手段', '压后取心 + 紫外光（365 nm）直接观察'],
            ['技术核心', '无机荧光粉对破胶剂过硫酸铵具有本征化学惰性——示踪信号在破胶返排全过程中保持完整性'],
            ['技术定位', '为微地震/DAS等间接监测方法提供独立的实物校准基准'],
        ])

    add_heading_cn(doc, '1.3 与已有方案的对比优势', level=2)
    add_table_from_data(doc,
        ['对比维度', '有机荧光树脂路线', '量子点返排分析路线', '本项目方案'],
        [
            ['可视化方式', '直接（取心紫外）', '间接（返排液分析）', '✅ 直接（取心紫外）'],
            ['工艺兼容性', '❌ 不可破胶返排', '✅ 兼容', '✅ 兼容'],
            ['氧化破胶稳定性', '不适用（不可逆固化）', '❌ 有降解风险', '✅ 本征化学惰性'],
            ['实物证据提供', '有', '无', '✅ 有'],
        ])

    add_para(doc, '本项目是已知方案中唯一同时满足"直接可视化"与"完全工艺兼容"两项要求的技术路线。', bold=True, color=(0, 100, 0))

    # === 2. Technical Route ===
    add_heading_cn(doc, '二、总体技术路线', level=1)
    add_para(doc, '采用"材料改性 → 体系构建 → 动态验证 → 工程转化"四模块递进式研究框架。')

    add_table_from_data(doc,
        ['模块', '核心任务', '关键输出'],
        [
            ['模块一：材料改性', '荧光粉基础物性 → 储层环境适应性 → 双层表面改性 → 协同分散', '改性荧光粉（分散稳定、可响应破胶）'],
            ['模块二：体系构建', '高浓度母液 → 基液/冻胶/悬砂/破胶/地层伤害全评价', '荧光压裂液终液（性能合格）'],
            ['模块三：动态验证', '动态驱替五步串联 → 三组数据交叉验证', '工程可行性证据链'],
            ['模块四：工程转化', '现场工艺方案 → 经济性评估 → 环保评估', '施工方案 + 成本模型'],
        ])

    add_para(doc, '实验方法论三原则：对照原则（含/不含荧光粉平行对比）、梯次原则（从简单到复杂逐级验证）、统计原则（每组≥3平行样，均值±SD，ANOVA检验α=0.05）。', bold=True)

    doc.add_page_break()

    # === 3. Module 1: Material Modification ===
    add_heading_cn(doc, '三、模块一：荧光粉基础物性表征与表面改性', level=1)

    add_heading_cn(doc, '3.1 基础物性表征与储层环境适应性评价', level=2)
    add_para(doc, '实验材料：市售1000目 SrAl₂O₄:Eu²⁺,Dy³⁺ 荧光粉（固相法合成，D50≈13 μm）。')

    add_table_from_data(doc,
        ['表征/实验项目', '方法', '关键参数/条件'],
        [
            ['晶相纯度', 'XRD（Cu Kα, 10°~80°）', '与JCPDS No.34-0379对比'],
            ['颗粒形貌与粒径', 'SEM + 激光粒度仪', 'D10, D50, D90, Span值'],
            ['光学特性', '荧光光谱仪', '激发/发射光谱（~520 nm）, 余辉衰减曲线'],
            ['热稳定性', '60~150°C, 24/72/168 h', '相对发光保持率, 一级衰减k, t₁/₂'],
            ['化学稳定性-盐度', 'NaCl 0~100 g/L', '同上'],
            ['化学稳定性-氧化', 'APS 0~0.2% w/v', '同上'],
            ['化学稳定性-pH', 'pH 3~11 缓冲体系', '同上'],
        ])
    add_para(doc, '目的：绘制荧光粉在井下工况中的定量"适用边界图"——在何种温度/矿化度/pH/氧化条件下荧光信号能保持足够的强度和时间窗口。', bold=True)

    add_heading_cn(doc, '3.2 双层表面改性——"KH550 + PEG4000"方案', level=2)
    add_para(doc, '核心科学问题：如何在单一改性方案中协调注入阶段"分散"与破胶后"锚定"两种截然相反的功能需求？')
    add_para(doc, '答案：构建一个具有"时序响应"特性的双层表面化学结构——利用压裂施工自身的破胶环节作为功能切换的化学触发器。')

    add_table_from_data(doc,
        ['层次', '材料', '功能', '在注入阶段', '在破胶后'],
        [
            ['内层', 'KH550（硅烷偶联剂）', '化学锚固层', '被PEG屏蔽', '暴露-NH₂（锚定基团）'],
            ['外层', 'PEG4000（聚乙二醇）', '物理屏蔽层', '提供空间位阻→稳定悬浮', '被APS氧化脱附↓退出'],
        ])
    add_para(doc, '技术巧思：不需要增加额外的"触发"操作步骤——压裂施工本身就包含破胶环节，过硫酸铵的氧化作用恰好充当PEG脱附的化学触发器。整个方案在室温水相中完成，工艺温和，具备工业放大潜力。', bold=True, color=(0, 100, 0))

    add_heading_cn(doc, '3.3 工艺优化：L9(3⁴)正交实验', level=2)
    add_table_from_data(doc,
        ['因素', '水平1', '水平2', '水平3'],
        [
            ['A: KH550用量 (wt%)', '1.0', '2.0', '3.0'],
            ['B: PEG分子量', '2000', '4000', '6000'],
            ['C: PEG浓度 (wt%)', '1', '3', '5'],
            ['D: 搅拌时间 (min)', '30', '60', '90'],
        ])
    add_para(doc, '评价指标：静置2 h相对浊度保持率（主指标）+ Zeta电位绝对值（辅指标）。极差分析和方差分析确定各因素显著性排序及最优水平组合。')
    add_para(doc, '表征验证：FTIR（Si-O-Al ~980 cm⁻¹）+ XPS（N 1s, Si 2p）+ TGA（有机物质量百分比）+ SEM（改性前后形貌对比）。')

    add_heading_cn(doc, '3.4 多层次协同分散体系', level=2)
    add_para(doc, '在双层包覆基础上引入三种分散助剂，通过六组递进对照实验定量解析各组分的独立贡献与协同效应。')

    add_table_from_data(doc,
        ['组别', '组成', '目的'],
        [
            ['组1（空白）', '双层改性粉 + 0.5 wt% HPG', '基线'],
            ['组2', '组1 + 0.1 wt% 柠檬酸', '螯合Al³⁺，防交联干扰'],
            ['组3', '组1 + 0.05 wt% Triton X-100', '界面润湿，促PEG伸展'],
            ['组4', '组1 + 5 wt% 游离PEG4000', '强化空间位阻'],
            ['组5', '组1 + 柠檬酸 + Triton', '检验螯合-润湿二元协同'],
            ['组6（全配方）', '组1 + 柠檬酸 + Triton + PEG', '检验三机制超加和效应'],
        ])
    add_para(doc, '核心分析逻辑：协同效应 = 组6 − (组2+组3+组4−3×组1)。若该项为正且统计显著（ANOVA, p<0.05），则证明"螯合-润湿-位阻"三种机制之间存在超越简单加和的协同分散效应。', bold=True)

    doc.add_page_break()

    # === 4. Module 2: System Formulation ===
    add_heading_cn(doc, '四、模块二：荧光压裂液体系构建与性能评价', level=1)
    add_para(doc, '核心目标：验证荧光母液的引入不显著损害压裂液的核心工程性能——这是决定荧光示踪方案能否嵌入现有压裂工艺的"及格线"。', bold=True)

    add_heading_cn(doc, '4.1 高浓度荧光母液', level=2)
    add_para(doc, '制备工艺：改性荧光粉 + 去离子水 + 柠檬酸 + Triton X-100 → 高速均质（10000~15000 rpm, 5 min）→ 超声脱泡（40 kHz, 10 min）。目标浓度≥40 g/L。')

    add_table_from_data(doc,
        ['质量指标', '标准', '检测方法'],
        [
            ['静态稳定性', '7天分层高度比 ≤ 5%', '25 mL量筒静置，每24 h记录'],
            ['动态稳定性', '离心浊度比 TR ≥ 0.90', '2000 rpm × 10 min'],
            ['荧光保持率', '离心前后520 nm峰面积比 ≥ 90%', '荧光光谱仪'],
        ])

    add_heading_cn(doc, '4.2 压裂液全体系性能评价清单', level=2)
    add_table_from_data(doc,
        ['评价项目', '实验方法', '依据标准', '评判标准'],
        [
            ['基液性能', '六速旋转粘度计（170 s⁻¹, 25°C）', 'SY/T 5107-2016', '粘度/pH/密度与空白组无显著偏差'],
            ['交联时间', '旋转粘度计监测', 'SY/T 6376-2008', '与空白组无显著差异'],
            ['耐温耐剪切', 'HTHP流变仪（170 s⁻¹, 25→120°C）', 'SY/T 6376-2008', '≥50 mPa·s @ 目标温度, ≥60 min'],
            ['粘弹性', '频率扫描（0.1~10 Hz, 1%应变）', '行业惯例', 'G\' > G"（弹性主导）'],
            ['悬砂性能', '静态沉降法（20/40目陶粒, 480 kg/m³）', '行业惯例', '储层温度下沉降速度 ≤ 0.5 mm/min'],
            ['破胶性能', 'APS 0.05~0.2%, 60~120°C密闭破胶', 'SY/T 5107-2016', '粘度 ≤ 10 mPa·s, 残渣不显著增加'],
            ['地层伤害', '滤饼伤害 + 岩心驱替 + 导流能力', 'SY/T 6540-2002', '荧光粉独立伤害增量ΔD可接受'],
        ])

    add_heading_cn(doc, '4.3 关键风险识别', level=2)
    add_bullet(doc, '交联干扰风险：改性粉表面残余可溶出Al³⁺可能竞争消耗硼交联剂 → 通过含/不含荧光粉冻胶G\'值对比直接检验')
    add_bullet(doc, '破胶残渣风险：13 μm不可降解固体颗粒可能增加破胶液残渣 → 通过过滤称重法定量比较')
    add_bullet(doc, '导流能力损害风险：颗粒堵塞支撑剂充填层 → 通过API导流室10~60 MPa全压力范围测试')

    doc.add_page_break()

    # === 5. Module 3: Dynamic Validation ===
    add_heading_cn(doc, '五、模块三：裂缝壁面吸附与动态驱替示踪验证', level=1)
    add_para(doc, '核心目标：在完整模拟压裂施工全流程的条件下，以三组联动数据构成荧光示踪工程可行性的完整证据链。', bold=True)

    add_heading_cn(doc, '5.1 模拟裂缝可视化实验装置', level=2)
    add_table_from_data(doc,
        ['组件', '规格/参数'],
        [
            ['岩心模型', '须家河组天然砂岩（φ 5%~12%）, 50×50×100 mm, 巴西劈裂法预制裂缝'],
            ['裂缝宽度梯度', '0.1 / 0.5 / 1.0 / 2.0 mm（精密金属垫片控制）'],
            ['恒流泵', '双柱塞式, 0.01~50 mL/min, 可编程多段注入'],
            ['回压阀', '闭合压力 5~30 MPa'],
            ['加热套', 'PID控温, 60~120°C±1°C'],
            ['紫外成像系统', '365 nm LED + 工业相机 + 520±10 nm带通滤光片'],
        ])
    add_para(doc, '所有成像参数（f/4.0, 2 s曝光, ISO 800）全实验锁定不变，确保不同岩样之间图像可比。')

    add_heading_cn(doc, '5.2 四阶段动态驱替实验', level=2)
    add_para(doc, '完整复现压裂施工全过程，中间不跳过任何一个工程环节。')

    add_table_from_data(doc,
        ['阶段', '操作', '条件', '采集数据', '判据'],
        [
            ['① 注入', '注入荧光HPG冻胶3 PV', '90°C, 恒定流量', '注入压力-累计PV曲线', '可注入性：压力趋平台，无异常攀升'],
            ['② 关井破胶', '关闭阀门, 密闭放置', '90°C, 5 MPa回压, 12 h', '（不采集, 提供功能切换窗口）', '—'],
            ['③ 返排', '反向注模拟地层水 5~10 CPV', '0.1 mL/min', '返排液归一化浓度-累计CPV曲线', '附着牢固度：1~2 CPV内降至本底, 净残留率>90%'],
            ['④ 取心成像', '拆卸夹持器, 紫外拍照', '锁定成像参数', '5 ROI灰度值 (mean±SD)', '空间对应性：裂缝壁面可见荧光标记'],
        ])

    add_heading_cn(doc, '5.3 破胶前后表面化学变化（三重交叉验证）', level=2)
    add_para(doc, '这是全文最核心的科学问题——如果破胶并没有诱导PEG脱附，整个功能切换假说就失去了化学基础。')

    add_table_from_data(doc,
        ['表征手段', '检测目标', '证据类型', '预期变化'],
        [
            ['FTIR', 'PEG C-O-C (~1100 cm⁻¹), C-H (~2880 cm⁻¹)', '振动光谱证据', '破胶后峰强度显著减弱'],
            ['XPS N 1s', '-NH₂ (~399 eV) 峰面积', '元素化学态证据', '破胶后峰面积显著增强'],
            ['Zeta电位-pH', '表面电荷密度', '胶体界面证据', '破胶后电位曲线明显偏移'],
        ])
    add_para(doc, '完备性设计：对"三项均支持脱附""部分支持""均不支持"三种可能结果均预设了对应的机制解释和替代方案，确保假说检验的严谨性。', bold=True)

    add_heading_cn(doc, '5.4 静态吸附实验', level=2)
    add_para(doc, '从热力学（吸附是否自发？ΔG°<0？）和动力学（关井6~24 h内能否完成？）两个维度获取定量锚定参数。')
    add_bullet(doc, '条件：天然砂岩薄片（10×10×2 mm）, 初始浓度0.1~5.0 g/L, 三温度（25/50/80°C）, 模拟地层水（50 g/L NaCl + 2 g/L CaCl₂）')
    add_bullet(doc, '等温吸附：Langmuir / Freundlich / Temkin 三模型拟合, R²+AIC选优')
    add_bullet(doc, '热力学：Van\'t Hoff方程 → ΔG°, ΔH°, ΔS°')
    add_bullet(doc, '动力学：拟一级 / 拟二级 / 颗粒内扩散模型 → 速率控制步骤判定')

    add_heading_cn(doc, '5.5 荧光信号与裂缝几何的半定量关系', level=2)
    add_para(doc, '终极目标：使紫外图像从"看到裂缝"升级到"读出裂缝宽度"。')
    add_table_from_data(doc,
        ['层次', '方法', '条件', '结论'],
        [
            ['定性验证', 'SNR >= 3:1, 检出率 > 90%', '四种裂缝宽度', '确认能看见'],
            ['半定量探索', 'Spearman秩相关（不预设线性）', 'n=12（4×3）', 'ρ≥0.85, p<0.01 → 单调关系成立'],
        ])
    add_para(doc, '必须明确的方法论边界：该半定量关系受控于同一砂岩类型、同一成像条件和同一荧光粉批次——不具备跨条件绝对定量测量能力。以下条件可导致关系失效：粘土矿物>10~15%、JRC>18、批次间差异>20~30%、多次重复改造。', color=(180, 0, 0))

    doc.add_page_break()

    # === 6. Module 4: Engineering Translation ===
    add_heading_cn(doc, '六、模块四：现场施工工艺方案与经济性评估', level=1)

    add_heading_cn(doc, '6.1 "母液预配+在线稀释"工艺方案', level=2)
    add_para(doc, '工艺流程：母液配制罐（5~10 m³, 高速搅拌500~1000 rpm × 30 min, 超声脱泡）→ 计量泵（±1%精度, 0.5% v/v）→ 静态混合器（SMX型, DN50）→ 混砂车（+20/40目陶粒）→ 高压泵组 → 井筒。')
    add_para(doc, '质量控制：每30 min在线取样检测520 nm发射峰强度，浓度偏差≤±15%，粘度偏差≤±10%。')

    add_heading_cn(doc, '6.2 经济性评估', level=2)
    add_para(doc, '算例：四川盆地典型页岩气水平井（水平段1500 m, 8段, 单段2000 m³, 总注入16000 m³, 母液添加0.5% v/v）。')

    add_table_from_data(doc,
        ['费用项目', '用量', '单价', '金额（万元）'],
        [
            ['荧光粉 SrAl₂O₄:Eu,Dy', '3200 kg', '200~500 元/kg', '~96'],
            ['KH550（硅烷偶联剂）', '64 kg', '50~80 元/kg', '~0.4'],
            ['PEG4000', '96 kg', '20~30 元/kg', '~0.2'],
            ['柠檬酸 + Triton X-100', '~5 kg', '10~30 元/kg', '~0.1'],
            ['设备租赁（母液罐/计量泵/混合器, 摊销）', '—', '—', '1~3'],
            ['单井合计', '—', '—', '~100（约14万美元）'],
        ])
    add_para(doc, '与化学示踪剂方案（3~8万美元/井）对比分析：', bold=True)
    add_bullet(doc, '全井段使用偏高，但仅需在1~2个关键段使用 → 用量降至1/4~1/2，成本约25~50万元')
    add_bullet(doc, '核心价值在于提供取心实物证据——若通过直接验证避免一次错误压裂设计决策（井距优化/射孔位置选择），节省的成本远超示踪方案本身费用')
    add_bullet(doc, '无机材料不涉及放射性，省去特殊运输、许可和放射性废物处理费用')

    add_heading_cn(doc, '6.3 环保安全性评估', level=2)
    add_table_from_data(doc,
        ['方面', '评估结论'],
        [
            ['稀土元素毒性', 'Sr/Eu/Dy为低毒性（LD50>2000 mg/kg, 大鼠经口），需浸出实验（GB 5085.3-2007）确认'],
            ['返排液处理', '13 μm微粒可通过常规絮凝-沉降-过滤工艺有效去除；稀土离子可用化学沉淀法（加碱pH 9~10）或离子交换法去除'],
            ['含荧光粉岩屑', 'SrAl₂O₄基质化学稳定、不溶于水，建议按一般工业固废处置'],
        ])

    add_heading_cn(doc, '6.4 技术定位与适用条件', level=2)
    add_para(doc, '定位：辅助校准工具。在微地震/DAS指示的裂缝核心区域附近的1~2口关键井的1~2个关键段注入荧光压裂液，压后取心紫外验证——若岩心荧光分布与间接方法预测一致，则间接反演结果获实物验证；若不一致，则为反演参数修正提供约束。', bold=True)
    add_bullet(doc, '✅ 适用：砂岩储层（Si-OH提供锚定位点）、HPG/瓜尔胶基压裂液、需取心+紫外成像的井')
    add_bullet(doc, '❌ 不适用：碳酸盐岩/页岩储层（锚定化学不同，需重新验证）、裂缝宽度<0.1 mm（低于检出阈值）')

    doc.add_page_break()

    # === 7. Innovation ===
    add_heading_cn(doc, '七、研究创新点', level=1)
    add_table_from_data(doc,
        ['编号', '创新点', '简要说明'],
        [
            ['①', '无机荧光示踪策略', '利用SrAl₂O₄:Eu,Dy对APS氧化环境的本征化学惰性，解决有机/量子点示踪剂的信号衰减问题'],
            ['②', '"螯合-润湿-位阻"协同分散', '首次系统整合三种机制实现高密度微米粉（3.6~4.0 g/cm³）在低密度非牛顿流体中的长期悬浮'],
            ['③', '时序功能切换表面设计', '利用压裂施工自身时序（注入→关井→返排）驱动颗粒表面从"分散态"到"锚定态"的转化，无需额外操作'],
            ['④', '动态驱替全流程验证方法', '构建"注入-关井-返排-取心-成像"五步串联实验，以三组联动数据构成工程可行性完整证据链'],
        ])

    # === 8. Progress ===
    add_heading_cn(doc, '八、当前进度总览', level=1)
    add_table_from_data(doc,
        ['模块', '研究内容', '进度', '说明'],
        [
            ['模块一', '基础物性 + 环境适应性 + 表面改性 + 协同分散', '██░░ 方案已设计', '待材料采购后执行'],
            ['模块二', '母液制备 + 基液/冻胶/悬砂/破胶/伤害评价', '██░░ 方案已设计', '待模块一完成后执行'],
            ['模块三', '动态驱替装置 + 吸附实验 + 五步串联验证', '██░░ 方案已设计', '装置待搭建'],
            ['模块四', '工艺方案 + 经济性 + 环保评估', '████ 初步完成', '待实验确定最优参数后精化'],
        ])
    add_para(doc, '总体进度：约30%——实验方案设计阶段已基本完成，已输出6章完整论文初稿框架，核心实验数据待采集。')

    # === 9. Next Steps ===
    add_heading_cn(doc, '九、下一步工作计划', level=1)

    add_para(doc, '近期（1~2个月）：', bold=True)
    add_bullet(doc, '采购实验材料：1000目荧光粉、KH550、PEG4000、HPG、有机硼交联剂、APS等')
    add_bullet(doc, '执行模块一实验：XRD/SEM/荧光光谱表征 → 4组环境适应性 → L9正交改性优化')
    add_bullet(doc, '搭建模拟裂缝可视化实验装置')

    add_para(doc, '中期（3~4个月）：', bold=True)
    add_bullet(doc, '执行模块二实验：母液制备 → 基液/冻胶/悬砂/破胶/地层伤害全系列评价')
    add_bullet(doc, '执行模块三实验：静态吸附 → 动态驱替四阶段串联验证')

    add_para(doc, '远期（5~6个月）：', bold=True)
    add_bullet(doc, '数据分析与论文正文撰写')
    add_bullet(doc, '经济性评估数据更新（实际采购价格）')
    add_bullet(doc, '毕业答辩准备')

    # === 10. Support Needed ===
    add_heading_cn(doc, '十、需要甲方确认/支持的事项', level=1)
    add_bullet(doc, '目标储层参数确认：目标井的温度、矿化度、pH范围是否在本文实验设计覆盖范围内（60~150°C, 0~100 g/L, pH 3~11）？')
    add_bullet(doc, '荧光粉批量采购渠道：甲方是否有推荐的商用SrAl₂O₄:Eu,Dy供应商？需确认到货周期和批次间发光性能稳定性。')
    add_bullet(doc, '现场先导性试验意向：若实验室验证顺利，是否可在甲方某口井的1~2个非关键段开展先导性验证？')
    add_bullet(doc, '取心计划协调：荧光示踪验证依赖压后取心，需与甲方钻井-完井计划协调取心窗口。')

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('本报告基于论文初稿（2026年6月版）撰写。标注"待补充"的项目将在实验完成后更新。')
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(150, 150, 150)

    # Save
    output_path = r'C:\Users\郝\Desktop\claude\output\荧光压裂液实验报告_汇报版.docx'
    doc.save(output_path)
    print(f'报告已保存至: {output_path}')

if __name__ == '__main__':
    main()