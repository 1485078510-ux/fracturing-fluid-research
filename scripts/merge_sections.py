# -*- coding: utf-8 -*-
"""Merge 2.3+2.4, delete Fig.2-2, renumber subsequent figures."""
from docx import Document
from docx.oxml.ns import qn
import re

DST = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP_formatted.docx'
doc = Document(DST)

def sp(idx, text):
    p = doc.paragraphs[idx]
    for r in p.runs: r.text = ''
    if p.runs: p.runs[0].text = text
    else: p.add_run(text)

def clear(idx):
    p = doc.paragraphs[idx]
    for r in p.runs: r.text = ''

# ===== 1. Merge 2.3 heading =====
sp(31, "2.3 Preparation of Epoxy Resin Microspheres and ESP-T")

# ===== 2. Remove Fig.2-2 caption [34] and image para [35] =====
clear(34)   # Fig.2-2 caption
clear(35)   # empty (image is here — need to check)
# Actually [35] was shown as "(empty)" — the image is likely in [35] as an inline shape
# Let me check if [35] has an image
has_img_35 = False
for run in doc.paragraphs[35].runs:
    if (run._element.findall('.//'+qn('w:drawing')) or
        run._element.findall('.//'+qn('a:blip'))):
        has_img_35 = True; break
if has_img_35:
    # Remove the image run but keep paragraph
    for run in doc.paragraphs[35].runs:
        drawings = run._element.findall('.//'+qn('w:drawing'))
        for d in drawings:
            d.getparent().remove(d)
        blips = run._element.findall('.//'+qn('a:blip'))
        for b in blips:
            # Remove the blip's parent elements going up to drawing
            parent = b.getparent()
            while parent is not None and parent.tag != qn('w:drawing'):
                parent = parent.getparent()
            if parent is not None:
                parent.getparent().remove(parent)

# ===== 3. Remove 2.4 heading [36] =====
clear(36)

# ===== 4. Renumber Fig.2-3 -> Fig.2-2 at [39] =====
sp(39, "Fig.2-2 Schematic illustration for the preparation of Fe3O4-encapsulated epoxy resin proppants")

# ===== 5. Renumber all subsequent figure references in body text =====
# Fig. 2-5 -> Fig. 2-3, Fig. 2-6 -> Fig. 2-4, Fig. 2-7 -> Fig. 2-5
# Figure 2-6 -> Figure 2-4, Figure 2-7 -> Figure 2-5
renumber_map = {
    'Fig. 2-5': 'Fig. 2-3',
    'Fig. 2-6': 'Fig. 2-4',
    'Fig. 2-7': 'Fig. 2-5',
    'Figure 2-6': 'Figure 2-4',
    'Figure 2-7': 'Figure 2-5',
}

# Also update all paragraphs with figure references
for para in doc.paragraphs:
    # Check if paragraph contains figure references
    # Use regex to replace in run text
    for run in para.runs:
        if run.text:
            for old, new in renumber_map.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)

# ===== 6. Also update text references to Section 2.3/2.4 =====
# Section 2.3 references are now "Section 2.3"
# Section 2.4 references should become "Section 2.3" (merged)
for para in doc.paragraphs:
    for run in para.runs:
        if run.text and 'Section 2.4' in run.text:
            run.text = run.text.replace('Section 2.4', 'Section 2.3')

# ===== SAVE =====
doc.save(DST)
print("Merged 2.3+2.4, removed Fig.2-2, renumbered Figs 2-3/2-4/2-5")
print(f"Saved: {DST}")