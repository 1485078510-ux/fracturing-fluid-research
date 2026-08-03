#!/usr/bin/env python3
"""Rewrite entire paper in a natural, human academic voice."""
from docx import Document
doc = Document('四氧化三铁环氧树脂拟合/ESP-T_终稿.docx')

# ═══════════════════════════════════════════════════════════════════
# SECTION 3.7 — complete natural voice
# ═══════════════════════════════════════════════════════════════════

# [126] Derivation
p = doc.paragraphs[126]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "A tracer breakthrough curve carries more information than a single "
    "concentration value. Two processes contribute to the shape observed "
    "at the wellhead. The first is straightforward: tracer that accumulated "
    "in the proppant pack during shut-in is swept out when the well opens, "
    "producing a pulse that rises and falls as the slug passes the sampling "
    "point. The second is slower and less obvious: residual tracer continues "
    "to diffuse out of the epoxy matrix long after the main slug has passed, "
    "feeding a sustained, low-level signal into the produced oil. These two "
    "contributions overlap in time, and neither can be isolated by simply "
    "reading the concentration curve."
    "\n\n"
    "We separate them by modeling the transport from first principles. "
    "The one-dimensional advection-dispersion equation, "
    "dC/dt + v dC/dx = D d2C/dx2, "
    "governs the evolution of tracer concentration along the production "
    "tubing, with v = 4Q/(pi d2) and D = alpha v. For the shut-in slug "
    "the classical instantaneous-injection solution [29] gives a Gaussian "
    "pulse: C_rise = (M/A)/sqrt(4 pi D t) exp[-(x - v t)2 / (4 D t)]. "
    "For the sustained release the continuous-source solution gives an "
    "erfc form: C_fall = (C0/2) erfc[(x - v t) / sqrt(4 D t)]. "
    "We blend the two with a tanh weight, "
    "w(t) = 1/2 [1 + tanh((t - t0)/sigma)], "
    "which shifts smoothly from the pulse-dominated early period to the "
    "diffusion-dominated tail. The full model, "
    "C(t) = cb + w(t) C_rise + [1 - w(t)] C_fall, "
    "is Eq. (2). The parameter sigma is fixed at the sampling interval "
    "(4 min); the remaining six parameters are discussed below."
)
print('[126]')

# [142] Fitting
p = doc.paragraphs[142]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "Figure 3-8(b) shows the fit to a single-phase breakthrough curve; "
    "Table 3-3 lists the fitted values. The match is close (R2 = 0.9939, "
    "RMSE = 0.0210), but what matters more is that the recovered parameters "
    "make physical sense. The model returns an effective flow rate "
    "Q = 0.46 mL/min; the pump was set to 0.50 mL/min, an 8% difference. "
    "The mean residence time from the fit, MRT = 37.4 min, is within 3% "
    "of the convective travel time x/v = 38.6 min calculated from the "
    "known tubing length and flow velocity. The Peclet number "
    "Pe = x/alpha = 0.934 sits near unity, the regime expected for a "
    "gradual-release source where neither advection nor dispersion "
    "dominates. Integrating C_rise and C_fall separately reveals that "
    "47% of the total signal area belongs to the erfc tail. Nearly half "
    "the tracer detected at the wellhead comes not from the shut-in slug "
    "but from the slow, sustained release that continues throughout "
    "production."
)
print('[142]')

# [146] Parameter meaning
p = doc.paragraphs[146]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "These numbers tell a consistent story. The flow rate, the residence "
    "time, and the Peclet number converge from three independent directions "
    "on values that match what is known from the experimental setup. This "
    "is not a trivial outcome of curve-fitting; it indicates that the model "
    "recovers relationships rather than inventing them. The 47% tail "
    "contribution deserves emphasis because of what it implies for "
    "monitoring strategy. If half the signal comes from sustained release, "
    "a short shut-in followed by continuous sampling captures the essential "
    "information; there is no need for repeated shut-in cycles. The "
    "tail-plateau concentration cb, fitted as a constant in this single-"
    "temperature experiment, reflects the steady-state diffusion flux "
    "that Section 3.6 showed to be thermally activated (K rising from "
    "0.055 at 30 degC to 0.196 at 120 degC). One would expect cb to "
    "be higher in hotter wells, a prediction testable with multi-"
    "temperature field data."
)
print('[146]')

# [148] Bridge
p = doc.paragraphs[148]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "A fractured well does not produce single-phase oil indefinitely. "
    "Formation water eventually breaks through, and the produced fluid "
    "becomes a mixture whose oil fraction varies from stage to stage "
    "and from month to month. The question is whether the analysis "
    "developed above, tested so far on single-phase flow, remains useful "
    "when both oil and water are moving through the proppant pack. We "
    "examined this with steady-state core displacement experiments at "
    "three oil-water ratios (OWR = 4:1, 1:1, 1:4) and four total flow "
    "rates (0.1-0.4 mL/min). Direct observation of the effluent at "
    "OWR = 1:1 showed a higher oil fraction early in the displacement, "
    "consistent with the oil-permeable character of ESP-T documented "
    "in Section 3.5."
)
print('[148]')

# [149] Concentration
p = doc.paragraphs[149]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "Figure 3-9(a) plots the tracer concentration against total flow "
    "rate for each OWR. Concentration drops as flow rate rises, a "
    "straightforward consequence of dilution: faster flow means less "
    "contact time per unit volume of fluid passing through the pack. "
    "The concentration shows no meaningful dependence on OWR, suggesting "
    "that oil and water compete for access to the proppant surface in "
    "rough proportion to their volume fractions. If concentration tracks "
    "dilution more than it tracks the actual tracer release rate, then "
    "concentration alone is a poor proxy for the production rate of a "
    "given stage."
)
print('[149]')

# [151] FO
p = doc.paragraphs[151]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "The way around this is to work with tracer flux rather than "
    "concentration. We define FO as the mass of tracer passing "
    "the sampling point per unit time, the product of the oil-phase "
    "tracer concentration and the oil volumetric flow rate. Because "
    "FO equals the release rate from the ESP-T pack [8], it strips "
    "out the dilution effect and reflects only what the proppant "
    "is actually releasing. Figure 3-9(b) confirms that the logic "
    "holds: FO increases with OWR, as expected when more oil contacts "
    "the proppant surface, but it does not depend on total flow rate. "
    "The proppant releases tracer at a rate set by the oil-wetted area, "
    "not by how fast fluid sweeps past it."
)
print('[151]')

# [153] Validation + limits
p = doc.paragraphs[153]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "The critical test is whether FO can be converted into a production "
    "rate. Figure 3-9(c) normalizes FO from the two-phase experiments "
    "by the steady-state FO measured during single-phase oil displacement "
    "(3.187 micro-g/min) and compares the result with the known oil flow "
    "rate at each OWR. The agreement is close, confirming that the oil "
    "production rate of a labeled interval can be recovered from FO "
    "when the total flow rate is known, even with a water phase present."
    "\n\n"
    "Several boundaries of the present study should be noted. The model "
    "assumes one fractured interval with a uniform proppant pack; real "
    "wells have multiple intervals whose tracer signals may overlap, "
    "and proppant distribution within a fracture is seldom uniform. "
    "The FO calibration derives from steady-state experiments; start-up, "
    "shut-in, and rapid drawdown may produce transient conditions where "
    "the FO-to-rate relationship deviates from steady state. Dodecane "
    "is a clean model oil; crude oil contains asphaltenes and other "
    "surface-active components that could alter the wetting and diffusion "
    "behavior of the epoxy matrix. The long-term stability of the epoxy "
    "and the stearic acid coating at temperatures above 120 degC, or "
    "in the presence of H2S, CO2, and high-salinity brines, has not "
    "been evaluated here and warrants a dedicated study."
)
print('[153]')

# [154] Synthesis
p = doc.paragraphs[154]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "Taken together, the single-phase and two-phase results outline "
    "a workable approach. The model takes a wellhead concentration "
    "history, separates the shut-in pulse from the sustained tail, "
    "and returns an effective flow rate that can be checked against "
    "independent measurements. That the same framework works for both "
    "single-phase and two-phase flow, with no modification beyond "
    "introducing the FO metric, suggests a degree of robustness "
    "that encourages further development toward field application."
)
print('[154]')

# [156] Field
p = doc.paragraphs[156]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "In the field, each fracture stage would receive ESP-T particles "
    "doped with a distinct metal (Mn, Zn, Cu, or others), creating a "
    "chemically coded tracer that survives for the life of the well. "
    "After a single shut-in, produced fluid is sampled at the wellhead "
    "at intervals guided by the expected pulse arrival time, which the "
    "model predicts from the wellbore geometry and the estimated flow "
    "velocity. The fitted Q for each stage flags underperforming "
    "intervals. Because FO is independent of total flow rate, periodic "
    "sampling over months or years tracks how each stage's contribution "
    "evolves. A stage whose Q trends downward faster than its neighbors "
    "may be closing or scaling; a stage that remains flat while others "
    "decline may warrant refracturing. None of this requires downhole "
    "tools or additional shut-ins. The tracer is already in the fracture, "
    "and the model provides the interpretation."
)
print('[156]')

doc.save('四氧化三铁环氧树脂拟合/ESP-T_终稿.docx')
print('\nSection 3.7 complete.')