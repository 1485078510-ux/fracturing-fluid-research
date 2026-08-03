# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def SF(run, cn='宋体', en='Times New Roman', sz=12, b=False):
    run.font.size = Pt(sz); run.bold = b; run.font.name = en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cn)

def H(doc, t, fs=14):  # heading
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(t); SF(r, '黑体', 'Times New Roman', fs, True)

def B(doc, t, fs=12):  # body
    p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(t); SF(r, '宋体', 'Times New Roman', fs)

def S(doc, t, fs=12):  # sub-heading
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(t); SF(r, '黑体', 'Times New Roman', fs, True)

def C(doc, t, fs=12):  # claim
    p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(t); SF(r, '宋体', 'Times New Roman', fs)

print('Functions defined OK')
