# -*- coding: utf-8 -*-
"""基于原文档精准优化：保留图片、保留原有格式，仅修改内容和结构问题。"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"c:\Users\郝\Desktop\claude\荧光压裂液\荧光压裂液专利申请书.docx"
OUT = r"c:\Users\郝\Desktop\claude\荧光压裂液\荧光压裂液专利申请书_优化版.docx"

doc = Document(SRC)

def paragraph_has_image(p):
    for run in p.runs:
        drawings = run._element.findall('.//' + qn('w:drawing'))
        if drawings:
            return True
        blips = run._element.findall('.//' + qn('a:blip'))
        if blips:
            return True
    return False

def replace_paragraph_text(p, new_text):
    runs_to_clear = []
    has_drawing = False
    for run in p.runs:
        drawings = run._element.findall('.//' + qn('w:drawing'))
        if drawings:
            has_drawing = True
            continue
        runs_to_clear.append(run)
    if has_drawing:
        if runs_to_clear:
            runs_to_clear[0].text = new_text
            for run in runs_to_clear[1:]:
                run.text = ""
        else:
            new_r = OxmlElement('w:r')
            new_rPr = OxmlElement('w:rPr')
            new_r.append(new_rPr)
            new_t = OxmlElement('w:t')
            new_t.text = new_text
            new_t.set(qn('xml:space'), 'preserve')
            new_r.append(new_t)
            p._element.insert(0, new_r)
    else:
        if p.runs:
            p.runs[0].text = new_text
            for run in p.runs[1:]:
                run.text = ""
        else:
            p.add_run(new_text)

def clear_paragraph_text(p):
    for run in p.runs:
        drawings = run._element.findall('.//' + qn('w:drawing'))
        if not drawings:
            run.text = ""

paragraphs = doc.paragraphs

# Find section indices
section_map = {}
for i, p in enumerate(paragraphs):
    text = p.text.strip()
    if text in ["技术领域", "背景技术", "发明内容", "发明目的", "技术方案",
                "有益效果", "附图说明", "具体实施方式", "权利要求书", "摘要"]:
        section_map[text] = i
    elif text == "说明书":
        section_map["说明书"] = i
    elif text.startswith("摘要附图"):
        section_map["摘要附图"] = i

print("Sections found:")
for k, v in sorted(section_map.items(), key=lambda x: x[1]):
    has_img = paragraph_has_image(paragraphs[v])
    print(f"  P{v}: [{k}] {'[IMG]' if has_img else ''}")

# Count images
img_paras = [i for i, p in enumerate(paragraphs) if paragraph_has_image(p)]
print(f"Image paragraphs: {img_paras}")

original_img_count = 0
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        original_img_count += 1
print(f"Total images in doc: {original_img_count}")

# --- MOD 1: Replace "发明目的" heading + merge purpose paragraphs ---
purpose_idx = section_map.get("发明目的")
if purpose_idx is not None:
    clear_paragraph_text(paragraphs[purpose_idx])
    print(f"MOD1: Cleared P{purpose_idx} (发明目的 heading)")

    purpose_paras = []
    for i in range(purpose_idx + 1, min(purpose_idx + 10, len(paragraphs))):
        t = paragraphs[i].text.strip()
        if t.startswith("本发明") and ("目的在于" in t[:15] or "另一目的" in t[:15] or "又一目的" in t[:15]):
            purpose_paras.append(i)
        else:
            break

    merged = ("本发明要解决的技术问题在于克服现有技术的上述不足，"
              "提供一种与常规水力压裂工艺兼容、能够在压裂液破胶返排全过程中保持荧光信号完整性、"
              "且可响应破胶化学环境变化实现从分散态向锚定态功能切换的改性稀土铝酸盐荧光粉。"
              "同时，本发明还提供含上述改性荧光粉的荧光压裂液体系，"
              "以及利用该荧光压裂液体系进行压裂裂缝荧光示踪的应用方法，"
              "该方法可为压裂液波及范围提供可实物验证的直接证据。")

    if purpose_paras:
        replace_paragraph_text(paragraphs[purpose_paras[0]], merged)
        for idx in purpose_paras[1:]:
            clear_paragraph_text(paragraphs[idx])
        print(f"MOD1: Merged purpose into P{purpose_paras[0]}, cleared {purpose_paras[1:]}")

# --- MOD 2: Clean blank lines in 附图说明 ---
tujie_idx = section_map.get("附图说明")
shishi_idx = section_map.get("具体实施方式")
if tujie_idx and shishi_idx:
    for i in range(tujie_idx + 1, shishi_idx):
        p = paragraphs[i]
        if paragraph_has_image(p):
            continue
        if not p.text.strip():
            clear_paragraph_text(p)
    print(f"MOD2: Cleaned blanks in 附图说明")

# --- MOD 3: Fix terminology ---
term_count = 0
for p in paragraphs:
    for run in p.runs:
        if "分散状态" in run.text:
            run.text = run.text.replace("分散状态", "分散态")
            term_count += 1
        if "锚定状态" in run.text:
            run.text = run.text.replace("锚定状态", "锚定态")
            term_count += 1
print(f"MOD3: Fixed {term_count} terminology instances")

# --- MOD 4: Add industrial applicability note ---
claims_idx = section_map.get("权利要求书")
if claims_idx:
    last_body_idx = claims_idx - 1
    while last_body_idx > 0 and not paragraphs[last_body_idx].text.strip():
        last_body_idx -= 1
    last_text = paragraphs[last_body_idx].text.strip()
    if "以上实施例仅为本发明" in last_text:
        new_p = OxmlElement('w:p')
        new_pPr = OxmlElement('w:pPr')
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:line'), '360')
        spacing.set(qn('w:lineRule'), 'auto')
        new_pPr.append(spacing)
        ind = OxmlElement('w:ind')
        ind.set(qn('w:firstLine'), '480')
        new_pPr.append(ind)
        new_p.append(new_pPr)
        new_r = OxmlElement('w:r')
        new_rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), '宋体')
        new_rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '24')
        new_rPr.append(sz)
        new_r.append(new_rPr)
        new_t = OxmlElement('w:t')
        new_t.set(qn('xml:space'), 'preserve')
        new_t.text = ("此外，本发明所述改性稀土铝酸盐荧光粉及其荧光压裂液体系具有良好的工业应用前景，"
                      "可适用于各类需要进行水力压裂裂缝监测的油气井，包括页岩气、致密油、煤层气等"
                      "非常规储层的压裂施工作业，并与现有压裂泵注设备和井下工具实现全流程兼容。")
        new_r.append(new_t)
        new_p.append(new_r)
        parent = paragraphs[last_body_idx]._element.getparent()
        idx_in_parent = list(parent).index(paragraphs[last_body_idx]._element)
        parent.insert(idx_in_parent + 1, new_p)
        print(f"MOD4: Added industrial note after P{last_body_idx}")

# Save
doc.save(OUT)

# Verify
doc2 = Document(OUT)
img_count2 = 0
for rel in doc2.part.rels.values():
    if "image" in rel.reltype:
        img_count2 += 1
total_p = len(doc2.paragraphs)
img_p_count = sum(1 for p in doc2.paragraphs if paragraph_has_image(p))

print(f"\n{'='*60}")
print(f"OPTIMIZED DOCUMENT SAVED: {OUT}")
print(f"  Paragraphs: {total_p}")
print(f"  Images: {img_count2} preserved in {img_p_count} locations")
print(f"  Original had: {original_img_count} images")
print(f"{'='*60}")