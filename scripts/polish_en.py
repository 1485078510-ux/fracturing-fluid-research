# -*- coding: utf-8 -*-
"""Final polish: fix all issues + chem/ref formatting for ESP_EN_final.docx"""
from docx import Document
from docx.oxml.ns import qn
import re

DST = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP_EN_final.docx'
doc = Document(DST)

IMG = set()
for i, p in enumerate(doc.paragraphs):
    for run in p.runs:
        if (run._element.findall('.//'+qn('w:drawing')) or run._element.findall('.//'+qn('a:blip'))):
            IMG.add(i); break

def sp(idx, text):
    if idx in IMG: return
    p = doc.paragraphs[idx]
    for r in p.runs: r.text = ''
    if p.runs: p.runs[0].text = text
    else: p.add_run(text)

# Fix specific text issues
sp(12, # synthesis - fix temperature units
    "A clean three-necked flask was mounted in a thermostatic water bath and connected "
    "to nitrogen purge and mechanical stirring. The flask was charged with 100 mL "
    "deionized water and heated to 80 degrees C. Upon reaching temperature, 2.703 g "
    "(0.01 mol) FeCl3 was added and the system purged with nitrogen for 15 min, "
    "followed by addition of 1.15 g (0.058 mol) FeCl2-4H2O and 2x10^-5 mol MnCl2-6H2O. "
    "After stirring for 10 min, 5.5 mL ammonia solution was rapidly added and the "
    "reaction proceeded at pH 10 for 2 h. The resulting black suspension was magnetically "
    "separated and the supernatant discarded. The precipitate was washed with anhydrous "
    "ethanol under sonication until neutral pH. Finally, the precipitate was blended "
    "with an ethanolic stearic acid solution and sonicated for oleophilic modification. "
    "After washing, the product was diluted to 100 mL and stored [28].")

sp(20, # characterization
    "Microstructure and elemental distribution were analyzed by scanning electron "
    "microscopy (SEM, ZEISS-Sigma 500) with energy-dispersive X-ray spectroscopy (EDS) "
    "at magnifications up to 50,000x. Surface morphology was examined by optical "
    "microscopy (Leica DM2700P). Thermal stability was evaluated by thermogravimetric "
    "analysis (TGA, TA Instruments Q500) in air at a heating rate of 10 degrees C/min "
    "from 50 to 800 degrees C. Water contact angle (WCA) was measured using a video "
    "optical contact angle analyzer (OCA20, Germany) with 5 uL droplets; five replicate "
    "measurements were performed per sample. Physical and mechanical properties including "
    "roundness, sphericity, bulk density, apparent density, acid solubility, and crush "
    "rate were evaluated according to industry standards. Proppant pack conductivity "
    "was indirectly assessed via oil-water filtration time through a 200-mesh screen "
    "using 2.0 g of proppant and 20 mL of deionized water or dodecane.")

sp(51, # density
    "Bulk density and apparent density results are presented in Figure 3-5. Pure epoxy "
    "microspheres have an average bulk density of 0.6179 g/cm3 and apparent density of "
    "1.02 g/cm3, while ESP-T has values of 0.646 and 1.072 g/cm3, respectively. The "
    "higher density of ESP-T is attributed to nano-Fe3O4@SA (approximately 5.18 g/cm3) "
    "relative to epoxy resin (approximately 1.1 g/cm3). Both proppants have bulk "
    "densities below that of water, ensuring favorable suspension in water-based "
    "fracturing fluids. The increase in apparent density reflects nano-Fe3O4@SA filling "
    "internal pores, yielding a denser structure that contributes to low crush rate.")

sp(68, # K-P
    "The K-P model describes two key mechanisms: Fickian diffusion and Case-II "
    "relaxation, where polymer swelling increases internal porosity. For spherical "
    "carriers, n <= 0.45 indicates Fickian-diffusion-dominated release; "
    "0.45 < n < 0.85 indicates anomalous transport co-governed by both mechanisms; "
    "and n >= 0.85 indicates Case-II-relaxation-dominated release. Fitting results "
    "(Table 3-2) show n values in the range 0.45-0.85 across all tested temperatures "
    "(R2 > 0.90), confirming a synergistic Fickian/Case-II mechanism.")

sp(86, # fitting results
    "The fitted curve is shown in Figure 3-8(b). The model achieves R2 = 0.9939 and "
    "RMSE = 0.0210, with residuals randomly distributed within +/- 2 RMSE. Key derived "
    "parameters are: v = 4Q/(pi d2) = 2.59 cm/min; alpha = 107.1 cm; Pe = x/alpha = 0.934. "
    "The fitted flow rate (0.46 mL/min) agrees closely with the pump-set rate (0.5 mL/min, "
    "relative error 8%; Figure 3-8c), validating the model accuracy. The mean residence "
    "time (37.4 min) matches the convective travel time (38.6 min, ratio 0.967). "
    "Deconvolution shows that 47% of the integrated signal originates from the erfc "
    "tailing component, confirming the dominant role of matrix-diffusion-controlled "
    "release, consistent with the non-Fickian mechanism in Section 3.6.")

sp(95, # conclusions
    "An oleophilic tracer proppant (ESP-T) was developed by encapsulating stearic "
    "acid-modified nano-Fe3O4@SA within an epoxy resin matrix via emulsion polymerization. "
    "Nano-Fe3O4@SA disperses uniformly as nanoclusters, and ESP-T exhibits sphericity "
    "and roundness exceeding 0.9, meeting industrial standards. ESP-T demonstrates "
    "excellent thermal stability (decomposition at 357.27 degrees C), low bulk density "
    "(0.646 g/cm3), and a water-resistant, oil-permeable characteristic that facilitates "
    "oil flow while inhibiting water breakthrough. Tracer release follows the Korsmeyer-"
    "Peppas model (R2 > 0.90, n = 0.45-0.85), governed by synergistic Fickian diffusion "
    "and Case-II relaxation. A piecewise ADE model with smooth tanh transition was "
    "developed to interpret the breakthrough curve, achieving R2 = 0.9939 with a fitted "
    "flow rate (0.46 mL/min) closely matching the pump-set rate (0.5 mL/min, error 8%). "
    "Under two-phase flow, tracer flux effectively quantifies per-interval oil production "
    "rates across varying oil-water ratios.")

# Chemical subscript + reference superscript formatting
REF_PAT = re.compile(r'\[\d+(?:[-,]\d+)*\]')
CHEM_PAT = re.compile(r'\b(Fe\d+O\d+(?:@\w+)?|FeCl\d+|MnCl\d+|SiO\d*|CO\d*|TiO\d*|Al\d*O\d*|C\d+H\d+COOH)\b')
count = 0
for idx, para in enumerate(doc.paragraphs):
    if idx in IMG or not para.text.strip(): continue
    if not (CHEM_PAT.search(para.text) or REF_PAT.search(para.text)): continue
    orig_font = None; orig_size = None
    for run in para.runs:
        if run.text.strip(): orig_font = run.font.name; orig_size = run.font.size; break
    matches = []
    for m in CHEM_PAT.finditer(para.text): matches.append((m.start(), m.end(), 'chem'))
    for m in REF_PAT.finditer(para.text): matches.append((m.start(), m.end(), 'ref'))
    if not matches: continue
    matches.sort(key=lambda x: (x[0], -x[1]))
    filtered, last = [], 0
    for s, e, t in matches:
        if s >= last: filtered.append((s, e, t)); last = e
    segs, pos = [], 0
    for s, e, t in filtered:
        if s > pos: segs.append((para.text[pos:s], 'normal'))
        segs.append((para.text[s:e], t)); pos = e
    if pos < len(para.text): segs.append((para.text[pos:], 'normal'))
    expanded = []
    for txt, typ in segs:
        if typ == 'chem':
            for part in re.split(r'(\d+)', txt):
                if not part: continue
                expanded.append((part, 'sub' if part.isdigit() else 'normal'))
        else: expanded.append((txt, typ))
    for run in list(para.runs): run._element.getparent().remove(run._element)
    for txt, typ in expanded:
        if not txt: continue
        r = para.add_run(txt)
        if orig_font: r.font.name = orig_font
        if orig_size: r.font.size = orig_size
        if typ == 'sub': r.font.subscript = True; count += 1
        elif typ == 'ref': r.font.superscript = True; count += 1

doc.save(DST)
print(f"Fixed 6 paragraphs. {count} chem/ref formatted.")
print(f"Images: {len(doc.inline_shapes)}")
print(f"Saved: {DST}")