# -*- coding: utf-8 -*-
"""Full-paper Nature-level polish for ESP_final_v2.docx."""
from docx import Document
from docx.oxml.ns import qn

DST = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP_final_v2.docx'
doc = Document(DST)

IMG = set()
for i, p in enumerate(doc.paragraphs):
    for run in p.runs:
        if (run._element.findall('.//'+qn('w:drawing')) or
            run._element.findall('.//'+qn('a:blip'))):
            IMG.add(i); break

def sp(idx, text):
    if idx in IMG:
        print(f"  SKIP [{idx}]"); return
    p = doc.paragraphs[idx]
    for r in p.runs: r.text = ''
    if p.runs: p.runs[0].text = text
    else: p.add_run(text)

# ===== ABSTRACT =====
sp(2, "Abstract: Accurate, per-interval monitoring of oil production contribution is "
    "essential for optimizing unconventional reservoir development. Here we design and "
    "fabricate an oleophilic Fe3O4-doped epoxy resin tracer proppant (ESP-T) via emulsion "
    "polymerization, using stearic acid-modified nano-Fe3O4 (nano-Fe3O4@SA) as the lipophilic "
    "tracer and epoxy resin as the carrier matrix, integrating fracture propping and production "
    "monitoring into a single material. We systematically characterize the microstructure, "
    "mechanical properties, thermal stability, wettability, oil-water permeability, and "
    "temperature-dependent tracer release kinetics of ESP-T, and validate its production "
    "monitoring performance through single-phase and two-phase core displacement experiments.")

sp(4, "Nano-Fe3O4@SA disperses uniformly as nanoclusters within the epoxy matrix. ESP-T "
    "achieves sphericity and roundness exceeding 0.9, a bulk density of 0.646 g/cm3, "
    "acid solubility of 3.3% (meeting the <=5% industrial standard), and an initial "
    "thermal decomposition temperature of 357.27 C, well above typical downhole conditions "
    "(80-200 C). Doping with nano-Fe3O4@SA raises the water contact angle from 72.3 degrees to "
    "104.6 degrees, imparting a water-resistant, oil-permeable transport characteristic: the oil "
    "filtration time (5 min 11 s) is 66.1% shorter than that of pure epoxy microspheres. "
    "Tracer release at 30-120 C follows the Korsmeyer-Peppas model (R2 > 0.90) with "
    "diffusion exponents of 0.45-0.85, indicating co-governance by Fickian diffusion and "
    "Case-II relaxation. Cumulative Fe release at 120 C over 14 days exceeds 2.0 mg/L, "
    "meeting ICP-MS detection requirements for long-term monitoring.")

sp(6, "To interpret the single-phase tracer breakthrough curve, we develop a piecewise "
    "advection-dispersion model with a smooth tanh transition, decomposing the signal "
    "into a Gaussian pulse component and an erfc tailing component. The model achieves "
    "R2 = 0.9939; the fitted flow rate (0.46 mL/min) closely matches the pump-set "
    "rate (0.5 mL/min, relative error 8%). Under steady-state two-phase flow, tracer "
    "flux effectively quantifies oil-phase flow rates across varying oil-water ratios. "
    "ESP-T thus offers a robust dual-function solution for long-term production monitoring "
    "in unconventional reservoirs.")
sp(8, "")

# ===== INTRODUCTION =====
sp(13, "Sustained growth in global energy demand has positioned unconventional oil and gas "
    "resources at the center of contemporary petroleum development [1-5]. These resources "
    "now account for more than 53% of the world total hydrocarbon reserves [6,7], and "
    "their efficient development is critical to global energy security. Hydraulic fracturing, "
    "the core stimulation technology for unconventional reservoirs, creates artificial "
    "fractures that enhance hydrocarbon seepage pathways and substantially boost productivity "
    "in low-permeability and tight formations [8]. However, the oil production contribution "
    "of each fractured interval varies markedly post-stimulation, and interval productivity "
    "decays at disparate rates as field development proceeds [9]. Accurate, per-interval "
    "production monitoring is therefore essential: it supplies key data for reservoir "
    "productivity assessment and guides the optimization of subsequent development strategies, "
    "ultimately raising the overall hydrocarbon recovery factor.")

sp(15, "Tracer technology has become the mainstream approach for monitoring per-interval "
    "production contribution, owing to its low cost, operational simplicity, and independence "
    "from sophisticated downhole equipment [10]. Unlike conventional methods such as well-test "
    "analysis, distributed fiber-optic sensing, and microseismic monitoring, tracer technology "
    "directly reflects fluid flow dynamics through concentration variations in produced fluids "
    "while drastically reducing equipment and maintenance costs. Originally developed for "
    "hydrological research, tracer methods were subsequently extended to characterize inter-well "
    "fluid flow in water-flooded reservoirs and evaluate reservoir connectivity [11-13]. With "
    "the widespread adoption of staged fracturing, tracers have become a standard tool for "
    "quantifying per-interval production contribution [14]. Compared with water-soluble tracers, "
    "oil-soluble variants yield more reliable production allocation data by tracking dynamic "
    "concentration profiles in the crude oil phase.")

sp(17, "Conventional oil-soluble tracers suffer from poor compatibility with fracturing fluid "
    "injection and limited long-term monitoring capability. Tracer proppants, composite "
    "materials that integrate fracture propping and tracer functions, address these limitations "
    "by immobilizing tracers within a solid carrier, enabling long-term dynamic monitoring "
    "upon co-injection with fracturing fluids and eliminating the need for separate tracer "
    "injection. Zhao et al. (2020) [15] coated rhodamine 6G onto polyvinyl alcohol-modified "
    "ceramic particles via solvent evaporation to retard aqueous release. Zhou et al. (2022) "
    "[16] encapsulated carbon quantum dots in polyvinyl alcohol-coated ceramic particles for "
    "controlled release. Malyavko et al. (2023) [17] employed monodisperse microspheres doped "
    "with fluorescent semiconductor nanocrystals functionalized with gas-philic, oleophilic, "
    "and hydrophilic polymer coatings to acquire phase-specific production data. Li et al. "
    "(2023) [18] encapsulated rare-earth element tracers with ammonium polymethacrylate to "
    "elucidate slow-release mechanisms. Gong et al. (2024) [19] embedded tracers into "
    "polystyrene (PS) microspheres via suspension polymerization and investigated their "
    "controlled-release behavior.")

sp(19, "Despite these advances, conventional coated tracer proppants face persistent challenges: "
    "high density impedes transport in fracturing fluids; monitoring capability ceases once "
    "the polymer coating dissolves; and pure PS microspheres exhibit inadequate mechanical "
    "strength and thermal stability [20-23]. Epoxy resin, a high-performance polymer with "
    "exceptional mechanical strength, thermal stability, and chemical resistance, offers a "
    "compelling alternative. Coating proppants with epoxy resin substantially reduces crush "
    "rate, improves compressive strength and acid resistance, and retains low-density "
    "advantages [24,25]. However, coating-based approaches remain constrained by uneven "
    "coverage, interfacial debonding risks, and complex multi-step preparation. Direct "
    "synthesis of epoxy resin microspheres as proppant matrices overcomes these limitations "
    "by enabling low-density tailoring, nanoparticle modification, and in-situ tracer "
    "encapsulation within a single synthetic step. Li et al. (2021) [26] explored the release "
    "of water-soluble tracers from epoxy resin in aqueous media. Wei et al. (2024) [27] "
    "encapsulated water-soluble tracers in epoxy matrices to locate water-producing intervals. "
    "However, the use of epoxy resin as an oleophilic release matrix for oil-phase production "
    "monitoring remains unexplored.")

sp(21, "Here we address this gap by designing and synthesizing an oleophilic Fe3O4/epoxy resin "
    "tracer proppant (ESP-T). Using epoxy resin as the matrix and stearic acid-modified "
    "nano-Fe3O4@SA as the oleophilic tracer, we achieve stable tracer-carrier integration "
    "via emulsion polymerization. We systematically characterize the microstructure, "
    "mechanical properties, thermal stability, wettability, and oil-water permeability of "
    "ESP-T; investigate its temperature-dependent tracer release behavior and kinetic "
    "mechanisms; and validate its oil production monitoring accuracy through single-phase "
    "and two-phase core displacement experiments. This work aims to establish a multifunctional "
    "proppant platform that combines low-density suspendability, high compressive resistance, "
    "and reliable tracing performance for hydraulic fracturing stimulation and long-term "
    "productivity monitoring in unconventional reservoirs.")

# ===== RESULTS =====
sp(69, "Figure 3-1 presents SEM micrographs of pure epoxy resin microspheres and ESP-T. "
    "At low magnification (a, d), both samples show excellent sphericity and high "
    "monodispersity, with uniform particle sizes and no inter-particle agglomeration, "
    "demonstrating the controllability of the suspension polymerization process. Well-formed "
    "spherical microspheres are stably obtained both with and without nano-Fe3O4@SA "
    "incorporation. Comparing (a) and (d), pure epoxy microspheres feature a relatively "
    "smooth surface with minor particulate debris, whereas ESP-T microspheres exhibit a "
    "uniformly rough surface topography, providing direct evidence for the successful "
    "incorporation of nano-Fe3O4@SA and its pronounced effect on surface morphology.")

sp(71, "At medium magnification (b), pure epoxy microspheres display subtle surface wrinkles "
    "and depressions, typical of thermosetting resins, arising primarily from volume shrinkage "
    "during epoxy curing. In contrast, ESP-T microspheres (e) are uniformly covered with "
    "micro- and submicron-scale convex protrusions. These protrusions exist as well-dispersed "
    "island-like structures rather than large agglomerates, confirming that nano-Fe3O4@SA "
    "disperses as nanoclusters, not as individual nanoparticles, within the epoxy matrix. "
    "This morphology demonstrates that stearic acid surface modification effectively suppresses "
    "uncontrolled macroscopic phase separation, enabling nanoscale agglomeration with "
    "microscale uniformity.")

sp(75, "A plausible formation mechanism for this surface morphology is as follows: the "
    "hydrophobically modified nano-Fe3O4@SA nanoclusters are uniformly dispersed within "
    "epoxy resin droplets after emulsification; as the epoxy cross-linking reaction "
    "proceeds, a rigid polymer network forms and immobilizes the nanoclusters in place. "
    "During subsequent curing shrinkage, a portion of these nanoclusters is extruded "
    "toward the microsphere surface, producing the surface-enriched rough nanostructures "
    "observed in (e) and (f).")

sp(107, "Roundness and sphericity were determined by comparison with the Krumbien-Sloss "
    "chart [29]. Both values exceed 0.9, satisfying industrial specifications "
    "(roundness >= 0.6, sphericity >= 0.6) and ensuring favorable fracture conductivity "
    "and pumpability. The formation of regular spherical morphology is attributed to "
    "optimized stirring: a rate of 380 RPM imparts uniform shear force to epoxy resin "
    "droplets within the nano-SiO2 aqueous dispersion, favoring regularly spherical "
    "droplets. The subsequently introduced guar gum (0.9 g for neat epoxy microspheres; "
    "1.1 g for ESP-T) elevates system viscosity, suppressing droplet deformation and "
    "sedimentation before curing. Hollow glass microspheres act as internal reinforcing "
    "supports, further preserving structural integrity and ideal sphericity.")

sp(117, "The acid solubility of ESP-T is 3.3%, compared with 2.5% for neat epoxy resin "
    "proppants, a negligible difference. Although nano-Fe3O4@SA is soluble in dilute "
    "HCl and HF, the stearic acid-derived hydrophobic film shields the particle surface "
    "from acid attack, mitigating dissolution. Meanwhile, uniformly dispersed "
    "nano-Fe3O4@SA nanoclusters act as an inorganic barrier within the epoxy matrix, "
    "slowing acid penetration and reducing resin degradation. Both proppants meet the "
    "industrial standard (acid solubility <= 5%), making them suitable for acid fracturing "
    "operations. The acid solubility standard deviation is below 0.2% for both materials, "
    "demonstrating stable dissolution behavior and excellent preparation reproducibility.")

sp(119, "The crush rate of neat epoxy microspheres at 50 MPa is 2.6%, compared with 2.9% "
    "for ESP-T, a negligible difference. The incorporation of nano-Fe3O4@SA thus exerts "
    "minimal influence on crush resistance. The slight increase observed for ESP-T likely "
    "arises from lattice defects introduced by metal doping of Fe3O4 and marginally weakened "
    "interfacial adhesion between the nanofiller and epoxy matrix, which together intensify "
    "stress concentration under high pressure. Nevertheless, the uniformly dispersed "
    "nano-Fe3O4@SA nanoclusters help distribute applied load and inhibit crack propagation, "
    "producing crush performance comparable to that of neat epoxy microspheres.")

sp(122, "Proppant-pack conductivity can be indirectly assessed via oil-water filtration time: "
    "shorter filtration times correspond to lower flow resistance and higher conductivity. "
    "The water filtration time of neat epoxy microspheres is 2 min 53 s, whereas that of "
    "ESP-T is 28 min 41 s. The hydrophobic surface of ESP-T impedes aqueous-phase spreading "
    "and flow through inter-particle pores; water molecules must overcome the repulsive "
    "force of the hydrophobic surface, increasing flow resistance and prolonging filtration. "
    "Conversely, the oil filtration time of neat epoxy microspheres is 15 min 11 s versus "
    "5 min 11 s for ESP-T. The hydrophobic surface is compatible with the oil phase, enabling "
    "rapid spreading within proppant pores and markedly reducing flow resistance. ESP-T thus "
    "exhibits a water-resistant, oil-permeable characteristic: in formations containing "
    "both phases, the proppant pack provides enhanced conductivity to oil, promoting rapid "
    "oil flow toward the wellbore, while the increased resistance to the aqueous phase "
    "helps mitigate water channeling and improve oil recovery.")

# 3.10 release mechanism
sp(140, "Based on the release results, the release mechanism of ESP-T is inferred as follows. "
    "During the release process, nonpolar solvents (e.g., alkanes) permeate into the polymer "
    "network, inducing swelling of the cross-linked epoxy matrix. At this stage, the polymer "
    "microspheres form a dual-state structure: an inner glassy core and an outer gel layer. "
    "The swelling reduces the entanglement of polymer molecular chains, thereby creating "
    "expanded transport channels that enable nano-Fe3O4@SA tracers to diffuse into the "
    "external environment. Elevated temperature enhances the permeability of solvent molecules "
    "into the polymer matrix, accelerating the swelling rate. Furthermore, increased temperature "
    "weakens intermolecular forces within the polymer network, expanding the pore size and "
    "further promoting tracer diffusion through the matrix.")

sp(133, "C = tracer concentration at time t (mg/L);")
sp(134, "C0 = maximum (equilibrium) release concentration (mg/L);")
sp(135, "K = kinetic rate constant;")
sp(136, "t = release time (h);")
sp(137, "n = diffusion exponent (dimensionless).")

# ===== SECTION 3.11 =====
sp(146, "We derive the interpretation model from the one-dimensional advection-dispersion "
    "equation (ADE): dC/dt + v dC/dx = D d2C/dx2, where v = 4Q/(pi d2) is the mean "
    "flow velocity and D = alpha v is the longitudinal dispersion coefficient. The tracer "
    "pulse released from ESP-T into the oil phase during the shut-in period is described "
    "by the classical ADE analytical solution for an instantaneous slug injection, which "
    "yields the Gaussian-form rising component in Eq. (6). The sustained slow release of "
    "residual tracer from the epoxy matrix after the main pulse is captured by solving "
    "the ADE with a continuous-source boundary condition, yielding the erfc tailing "
    "component. The two regimes are blended via a hyperbolic tangent weighting function "
    "w(t) = 0.5[1 + tanh((t0 - t)/sigma)], producing a smooth, physically continuous "
    "breakthrough curve.")

sp(149, "where:")
sp(150, "c_b = baseline concentration after the tracer signal stabilizes; "
    "A = peak amplitude coefficient, proportional to the tracer mass accumulated in the oil phase during shut-in;")
sp(151, "a = tailing amplitude coefficient, proportional to the sustained tracer release rate from the ESP-T matrix;")
sp(152, "alpha = longitudinal dispersivity, characterizing axial spreading of the tracer slug along the flow line (cm);")
sp(153, "Q = effective volumetric flow rate through the flow line (cm3/min);")
sp(154, "t0 = time at which sustained matrix-diffusion release overtakes the initial pulse as the dominant signal source (min);")
sp(155, "sigma = time span of the transition from pulse-dominated to tailing-dominated regime (min).")

sp(156, "In Eq. (6), x = 100 cm and d = 5 cm are the fixed length and inner diameter "
    "of the flow line. The rising component C_rise captures the advective-dispersive "
    "transport of the tracer slug from the proppant-packed section to the sampling "
    "point; the tailing component C_fall describes the sustained release of residual "
    "nano-Fe3O4@SA from the epoxy matrix into the flowing oil. This structure reflects "
    "the superposition of two processes: a short-duration pulse from tracer accumulation "
    "during shut-in, and a prolonged tail governed by the matrix-diffusion-controlled "
    "release kinetics established in Section 3.10.")

sp(158, "The fitted curve is shown in Figure 3-9(b), and the fitted parameters with derived "
    "transport properties are listed in Table 3-3. The model yields R2 = 0.9939 and "
    "RMSE = 0.0210, with residuals randomly distributed within +/- 2 RMSE across the "
    "full time range. Key derived parameters are: mean flow velocity "
    "v = 4Q/(pi d2) = 2.59 cm/min; longitudinal dispersivity alpha = 107.1 cm; and Peclet "
    "number Pe = x/alpha = 0.934. Pe < 1 confirms dispersion-dominated transport, consistent "
    "with the gradual, sustained release of nano-Fe3O4@SA from the ESP-T matrix. The "
    "fitted flow rate (0.46 mL/min) agrees closely with the pump-set rate (0.5 mL/min; "
    "relative error 8%; Figure 3-9c), validating the model accuracy for flow rate "
    "determination from tracer breakthrough data. The mean residence time "
    "(MRT = 37.4 min) matches the theoretical convective travel time (x/v = 38.6 min; "
    "ratio 0.967). Deconvolution of the fitted curve reveals that 47% of the integrated "
    "tracer signal originates from the erfc tailing component, confirming the dominant "
    "contribution of matrix-diffusion-controlled release, in agreement with the "
    "non-Fickian mechanism identified in Section 3.10.")

# ===== CONCLUSIONS =====
sp(170, "We synthesized oleophilic nano-Fe3O4@SA nanoparticles via coprecipitation with "
    "stearic acid modification and fabricated ESP-T via emulsion polymerization. SEM, "
    "elemental mapping, and optical microscopy confirmed uniform dispersion of "
    "nano-Fe3O4@SA as nanoclusters within the epoxy matrix. ESP-T exhibits sphericity "
    "and roundness exceeding 0.9, meeting industrial proppant molding standards.")

sp(172, "ESP-T demonstrates strong physical, mechanical, and thermal performance: bulk "
    "density of 0.646 g/cm3 for favorable suspension; apparent density of 1.072 g/cm3 "
    "indicating a compact structure; crush rate of 2.9% at 50 MPa; acid solubility "
    "of 3.3% (meeting the <=5% standard); and an initial decomposition temperature of "
    "357.27 C, far above downhole conditions (80-200 C). The water contact angle "
    "increases from 72.3 degrees to 104.6 degrees, and the oil filtration time (5 min 11 s) "
    "is 66.1% shorter than that of pure epoxy microspheres, validating the water-resistant, "
    "oil-permeable characteristic that facilitates oil flow while inhibiting water "
    "breakthrough.")

sp(174, "Tracer release from ESP-T at 30-120 C follows the Korsmeyer-Peppas model "
    "(R2 > 0.90; n = 0.45-0.85), governed by synergistic Fickian diffusion and Case-II "
    "relaxation. Cumulative Fe release at 120 C over 14 days exceeds 2.0 mg/L, meeting "
    "ICP-MS detection requirements. A piecewise ADE model with smooth tanh transition "
    "was developed to interpret the single-phase tracer breakthrough curve, decomposing "
    "the signal into a Gaussian pulse (shut-in accumulation) and an erfc tail (sustained "
    "matrix-diffusion-controlled release). The model achieves R2 = 0.9939; the fitted "
    "flow rate (0.46 mL/min) closely matches the pump-set rate (0.5 mL/min; error 8%); "
    "and the mean residence time (37.4 min) agrees with the theoretical travel time "
    "(38.6 min; ratio 0.967). Tailing accounts for 47% of the integrated signal, "
    "underscoring the critical role of matrix-diffusion-controlled release. Under "
    "steady-state two-phase flow, tracer flux effectively quantifies per-interval "
    "oil-phase flow rates across varying oil-water ratios.")

sp(176, "ESP-T integrates fracture propping and tracer monitoring into a single material, "
    "suitable for acid fracturing, deep wells, and high-pressure formations. It provides "
    "a quantitative basis for field-scale production prediction and fracturing parameter "
    "optimization, with broad industrial application potential in unconventional oil "
    "and gas development.")

# ===== SAVE =====
OUT = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP_polished_final.docx'
doc.save(OUT)
print(f"Saved: {OUT}")