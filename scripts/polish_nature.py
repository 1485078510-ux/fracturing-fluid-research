# -*- coding: utf-8 -*-
"""Nature-level academic polish of ESP_polished.docx."""
from docx import Document
from docx.oxml.ns import qn

DST = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP_polished.docx'
doc = Document(DST)

# Build image paragraph set
IMG = set()
for i, p in enumerate(doc.paragraphs):
    for run in p.runs:
        if (run._element.findall('.//' + qn('w:drawing')) or
            run._element.findall('.//' + qn('a:blip'))):
            IMG.add(i); break

def sp(para, text):
    """Safe polish: skip image paragraphs."""
    i = doc.paragraphs.index(para)
    if i in IMG:
        print(f"  SKIP [{i}] — contains image")
        return
    for run in para.runs: run.text = ''
    if para.runs: para.runs[0].text = text
    else: para.add_run(text)

# ============================================================
# ABSTRACT — tighten, claim-first
# ============================================================
sp(doc.paragraphs[2],
    "Abstract: Accurate monitoring of per-interval oil production contribution is "
    "critical for optimizing unconventional reservoir development. Here we design and "
    "fabricate an oleophilic Fe₃O₄-doped epoxy resin tracer proppant (ESP-T) via "
    "emulsion polymerization, using stearic acid-modified nano-Fe₃O₄ (nano-Fe₃O₄@SA) "
    "as the lipophilic tracer and epoxy resin as the carrier matrix—integrating "
    "fracture propping and production monitoring into a single material. We systematically "
    "characterize the microstructure, mechanical properties, thermal stability, wettability, "
    "oil–water permeability, and temperature-dependent tracer release kinetics of ESP-T, "
    "and validate its production monitoring performance through single-phase and two-phase "
    "core displacement experiments.")

sp(doc.paragraphs[4],
    "Nano-Fe₃O₄@SA disperses uniformly as nanoclusters within the epoxy matrix. ESP-T "
    "achieves sphericity and roundness exceeding 0.9, a bulk density of 0.646 g·cm⁻³, "
    "acid solubility of 3.3% (meeting the ≤ 5% industrial standard), and an initial "
    "thermal decomposition temperature of 357.27 °C—far above typical downhole conditions "
    "(80–200 °C). Doping with nano-Fe₃O₄@SA raises the water contact angle from 72.3° to "
    "104.6°, imparting a 'water-resistant, oil-permeable' transport characteristic: the "
    "oil filtration time (5 min 11 s) is 66.1% shorter than that of pure epoxy microspheres. "
    "Tracer release at 30–120 °C follows the Korsmeyer–Peppas model (R² > 0.90) with "
    "diffusion exponents of 0.45–0.85, indicating co-governance by Fickian diffusion and "
    "Case-II relaxation. Cumulative Fe release at 120 °C over 14 days exceeds 2.0 mg·L⁻¹, "
    "meeting ICP-MS detection limits for long-term monitoring.")

sp(doc.paragraphs[6],
    "To interpret the single-phase tracer breakthrough curve, we develop a piecewise "
    "advection-dispersion model with a smooth tanh transition, decomposing the signal "
    "into a Gaussian pulse component (tracer slug released during shut-in) and an erfc "
    "tailing component (sustained matrix-diffusion-controlled release). The model achieves "
    "R² = 0.9939; the fitted flow rate (0.46 mL·min⁻¹) closely matches the pump-set rate "
    "(0.5 mL·min⁻¹, relative error 8%). Under steady-state two-phase flow, tracer flux "
    "effectively quantifies oil-phase flow rates across varying oil–water ratios. ESP-T "
    "thus offers a robust dual-function solution for long-term production monitoring in "
    "unconventional reservoirs.")

# ============================================================
# INTRODUCTION — remove Chinese-influenced phrasing, tighten
# ============================================================

# Para 13 (Introduction para 1)
sp(doc.paragraphs[13],
    "Sustained growth in global energy demand has positioned unconventional oil and gas "
    "resources at the center of contemporary petroleum development [1–5]. These resources "
    "now account for more than 53% of the world's total hydrocarbon reserves [6,7], and "
    "their efficient development is critical to global energy security. Hydraulic fracturing, "
    "the core stimulation technology for unconventional reservoirs, creates artificial "
    "fractures that enhance hydrocarbon seepage pathways and substantially boost productivity "
    "in low-permeability and tight formations [8]. However, the oil production contribution "
    "of each fractured interval varies markedly post-stimulation, and interval productivity "
    "decays at disparate rates as field development proceeds [9]. Accurate, per-interval "
    "production monitoring is therefore essential: it supplies the data needed for reservoir "
    "productivity assessment and guides the optimization of subsequent development strategies, "
    "ultimately raising the overall hydrocarbon recovery factor.")

# Para 15 (Introduction para 2 — tracer technology)
# Not modifying — already acceptable, no major issues
sp(doc.paragraphs[15],
    "Tracer technology has become the mainstream approach for monitoring per-interval "
    "production contribution, owing to its low cost, operational simplicity, and independence "
    "from sophisticated downhole equipment [10]. Unlike conventional methods such as well-test "
    "analysis, distributed fiber-optic sensing, and microseismic monitoring, tracer technology "
    "directly reflects fluid flow dynamics through concentration variations in produced fluids "
    "while drastically reducing equipment and maintenance costs. Originally developed for "
    "hydrological research, tracer methods were subsequently extended to characterize inter-well "
    "fluid flow in water-flooded reservoirs and evaluate reservoir connectivity [11–13]. With "
    "the widespread adoption of staged fracturing, tracers have become a standard tool for "
    "quantifying per-interval production contribution [14]. Compared with water-soluble tracers, "
    "oil-soluble variants yield more reliable production allocation data by tracking dynamic "
    "concentration profiles in the crude oil phase.")

# Para 17 (Introduction para 3 — tracer proppants)
sp(doc.paragraphs[17],
    "Conventional oil-soluble tracers suffer from poor compatibility with fracturing fluid "
    "injection and limited long-term monitoring capability. Tracer proppants—composite "
    "materials that integrate fracture propping and tracer functions—address these limitations "
    "by immobilizing tracers within a solid carrier, enabling long-term dynamic monitoring "
    "upon co-injection with fracturing fluids and eliminating the need for separate tracer "
    "injection. Zhao et al. (2020) [15] coated rhodamine 6G onto polyvinyl alcohol-modified "
    "ceramic particles via solvent evaporation to retard aqueous release. Zhou et al. (2022) "
    "[16] encapsulated carbon quantum dots in polyvinyl alcohol-coated ceramic particles for "
    "controlled release. Malyavko et al. (2023) [17] employed monodisperse microspheres doped "
    "with fluorescent semiconductor nanocrystals functionalized with gas-philic, oleophilic, "
    "and hydrophilic polymer coatings to acquire phase-specific production data. Li et al. "
    "(2023) [18] encapsulated rare-earth element tracers with ammonium polymethacrylate to "
    "elucidate slow-release mechanisms. Gong et al. (2024) [19] embedded tracers into "
    "polystyrene (PS) microspheres via suspension polymerization and investigated their "
    "controlled-release behavior.")

# Para 19 (Introduction para 4 — limitations + epoxy resin gap)
sp(doc.paragraphs[19],
    "Despite these advances, conventional coated tracer proppants face persistent challenges: "
    "high density impedes transport in fracturing fluids; monitoring capability ceases once "
    "the polymer coating dissolves; and pure PS microspheres exhibit inadequate mechanical "
    "strength and thermal stability [20–23]. Epoxy resin—a high-performance polymer with "
    "exceptional mechanical strength, thermal stability, and chemical resistance—offers a "
    "compelling alternative. Coating proppants with epoxy resin substantially reduces crush "
    "rate, improves compressive strength and acid resistance, and retains low-density "
    "advantages [24,25]. However, coating-based approaches remain constrained by uneven "
    "coverage, interfacial debonding risks, and complex multi-step preparation. Direct "
    "synthesis of epoxy resin microspheres as proppant matrices overcomes these limitations "
    "by enabling low-density tailoring, nanoparticle modification, and in-situ tracer "
    "encapsulation within a single synthetic step. Li et al. (2021) [26] explored the release "
    "of water-soluble tracers from epoxy resin in aqueous media. Wei et al. (2024) [27] "
    "encapsulated water-soluble tracers in epoxy matrices to locate water-producing intervals. "
    "However, the use of epoxy resin as an oleophilic release matrix for oil-phase production "
    "monitoring remains unexplored.")

# Para 21 (Introduction para 5 — this work)
sp(doc.paragraphs[21],
    "Here we address this gap by designing and synthesizing an oleophilic Fe₃O₄/epoxy resin "
    "tracer proppant (ESP-T). Using epoxy resin as the matrix and stearic acid-modified "
    "nano-Fe₃O₄@SA as the oleophilic tracer, we achieve stable tracer-carrier integration "
    "via emulsion polymerization. We systematically characterize the microstructure, "
    "mechanical properties, thermal stability, wettability, and oil–water permeability of "
    "ESP-T; investigate its temperature-dependent tracer release behavior and kinetic "
    "mechanisms; and validate its oil production monitoring accuracy through single-phase "
    "and two-phase core displacement experiments. This work aims to establish a multifunctional "
    "proppant platform that combines low-density suspendability, high compressive resistance, "
    "and reliable tracing performance for hydraulic fracturing stimulation and long-term "
    "productivity monitoring in unconventional reservoirs.")

# ============================================================
# RESULTS — key fixes
# ============================================================

# Para 69 (3.1 SEM body) — minor polish
sp(doc.paragraphs[69],
    "Figure 3-1 presents SEM micrographs of pure epoxy resin microspheres and ESP-T. "
    "At low magnification (a, d), both samples show excellent sphericity and high "
    "monodispersity, with uniform particle sizes and no inter-particle agglomeration, "
    "demonstrating the controllability of the suspension polymerization process and the "
    "uniformity of curing. Well-formed spherical microspheres are stably obtained both "
    "with and without nano-Fe₃O₄@SA incorporation. Comparing (a) and (d), pure epoxy "
    "microspheres feature a relatively smooth surface with minor particulate debris, "
    "whereas ESP-T microspheres exhibit a uniformly rough surface topography. This "
    "provides direct evidence for the successful incorporation of nano-Fe₃O₄@SA and "
    "its pronounced effect on surface morphology.")

# Para 117 (3.8 Breakage Rate) — make claims more definitive
sp(doc.paragraphs[117],
    "The crush rate of neat epoxy resin microspheres at 50 MPa is 2.6%, compared with "
    "2.9% for ESP-T—a negligible difference. The incorporation of nano-Fe₃O₄@SA thus "
    "exerts minimal influence on crush resistance. The slight increase observed for "
    "ESP-T likely arises from two factors: lattice defects introduced by metal doping "
    "of Fe₃O₄, and marginally weakened interfacial adhesion between the nanofiller and "
    "the epoxy matrix, which together intensify stress concentration under high pressure. "
    "Nevertheless, the uniformly dispersed nano-Fe₃O₄@SA nanoclusters help distribute "
    "applied load and inhibit crack propagation, resulting in crush performance comparable "
    "to that of neat epoxy microspheres.")

# Para 122 (3.9 Conductivity — filtration time body)
sp(doc.paragraphs[122],
    "Proppant-pack conductivity can be indirectly assessed via oil–water filtration time: "
    "shorter filtration times correspond to lower flow resistance and higher conductivity. "
    "The measured water and oil filtration times are shown in Figure 3-7 and Table 3-1. "
    "The water filtration time of neat epoxy microspheres is 2 min 53 s, whereas that of "
    "ESP-T is 28 min 41 s. The hydrophobic surface of ESP-T impedes aqueous-phase spreading "
    "and flow through the inter-particle pores; water molecules must overcome the repulsive "
    "force of the hydrophobic surface, increasing flow resistance and prolonging filtration. "
    "Conversely, the oil filtration time of neat epoxy microspheres is 15 min 11 s versus "
    "5 min 11 s for ESP-T. The hydrophobic surface is compatible with the oil phase, enabling "
    "rapid spreading within the proppant pores and markedly reducing flow resistance. ESP-T "
    "thus exhibits a 'water-resistant, oil-permeable' characteristic: in formations containing "
    "both phases, the proppant pack provides enhanced conductivity to oil, promoting rapid "
    "oil flow toward the wellbore, while the increased resistance to the aqueous phase helps "
    "mitigate water channeling and improve oil recovery.")

# ============================================================
# SECTION 3.11 — polish the model description and results
# ============================================================

sp(doc.paragraphs[146],
    "We derive the interpretation model from the one-dimensional advection-dispersion "
    "equation (ADE): ∂C/∂t + v·∂C/∂x = D·∂²C/∂x², where v = 4Q/(πd²) is the mean "
    "flow velocity and D = αv is the longitudinal dispersion coefficient. The tracer "
    "pulse released from ESP-T into the oil phase during the shut-in period is described "
    "by the classical ADE analytical solution for an instantaneous slug injection, which "
    "yields the Gaussian-form rising component in Eq. (6). The sustained slow release of "
    "residual tracer from the epoxy matrix after the main pulse is captured by solving "
    "the ADE with a continuous-source boundary condition, yielding the erfc tailing "
    "component. The two regimes are blended via a hyperbolic tangent weighting function "
    "w(t) = ½[1 + tanh((t₀ − t)/σ)], producing a smooth, physically continuous "
    "breakthrough curve.")

sp(doc.paragraphs[157],
    "In Eq. (6), x = 100 cm and d = 5 cm are the fixed length and inner diameter of the "
    "flow line. The rising component C_rise captures the advective-dispersive transport "
    "of the tracer slug from the proppant-packed section to the sampling point; the "
    "tailing component C_fall describes the sustained release of residual nano-Fe₃O₄@SA "
    "from the epoxy matrix into the flowing oil. The parameter t₀ marks the transition "
    "center between the two regimes, and σ controls the transition width. This two-component "
    "structure captures the essential physics: the recorded tracer signal is the superposition "
    "of a short-duration concentration pulse generated by tracer accumulation during shut-in "
    "and a prolonged tail governed by the matrix-diffusion-controlled release kinetics "
    "established in Section 3.10.")

sp(doc.paragraphs[159],
    "The fitted curve is shown in Figure 3-9(b), and the fitted parameters with derived "
    "transport properties are listed in Table 3-3. The model yields R² = 0.9939 and "
    "RMSE = 0.0210, with residuals randomly distributed within ± 2 × RMSE across the "
    "full time range. Key derived parameters are: mean flow velocity "
    "v = 4Q/(πd²) = 2.59 cm/min; longitudinal dispersivity α = 107.1 cm; and Péclet "
    "number Pe = x/α = 0.934. Pe < 1 confirms dispersion-dominated transport, consistent "
    "with the gradual, sustained release of nano-Fe₃O₄@SA from the ESP-T matrix. The "
    "fitted flow rate (0.46 mL/min) agrees closely with the pump-set rate (0.5 mL/min; "
    "relative error 8%; Figure 3-9c), validating the model's accuracy for flow rate "
    "determination from tracer breakthrough data. The mean residence time (MRT = 37.4 min) "
    "matches the theoretical convective travel time (x/v = 38.6 min; ratio 0.967). "
    "Deconvolution of the fitted curve reveals that 47% of the integrated tracer signal "
    "originates from the erfc tailing component, confirming the dominant contribution of "
    "matrix-diffusion-controlled release, in agreement with the non-Fickian mechanism "
    "identified in Section 3.10.")

# ============================================================
# CONCLUSIONS — tighten, remove redundancy, strengthen closing
# ============================================================

sp(doc.paragraphs[171],
    "We synthesized oleophilic nano-Fe₃O₄@SA nanoparticles via coprecipitation with stearic "
    "acid modification and fabricated ESP-T via emulsion polymerization. SEM, elemental "
    "mapping, and optical microscopy confirmed uniform dispersion of nano-Fe₃O₄@SA as "
    "nanoclusters within the epoxy matrix. ESP-T exhibits sphericity and roundness exceeding "
    "0.9, meeting industrial proppant molding standards.")

sp(doc.paragraphs[173],
    "ESP-T demonstrates strong physical, mechanical, and thermal performance: bulk density "
    "of 0.646 g/cm³ for favorable suspension; apparent density of 1.072 g/cm³ indicating "
    "a compact structure; crush rate of 2.9% at 50 MPa; acid solubility of 3.3% (≤ 5% "
    "standard); and an initial decomposition temperature of 357.27 °C, far above downhole "
    "conditions. The water contact angle increases from 72.3° to 104.6°, and the oil "
    "filtration time (5 min 11 s) is 66.1% shorter than that of pure epoxy microspheres, "
    "validating the 'water-resistant, oil-permeable' characteristic that facilitates oil "
    "flow while inhibiting water breakthrough.")

sp(doc.paragraphs[175],
    "Tracer release from ESP-T at 30–120 °C follows the Korsmeyer–Peppas model "
    "(R² > 0.90; n = 0.45–0.85), governed by synergistic Fickian diffusion and Case-II "
    "relaxation. Cumulative Fe release at 120 °C over 14 days exceeds 2.0 mg/L, meeting "
    "ICP-MS detection requirements. A piecewise ADE model with smooth tanh transition "
    "was developed to interpret the single-phase tracer breakthrough curve, decomposing "
    "the signal into a Gaussian pulse (shut-in accumulation) and an erfc tail (sustained "
    "matrix-diffusion-controlled release). The model achieves R² = 0.9939; the fitted "
    "flow rate (0.46 mL/min) closely matches the pump-set rate (0.5 mL/min; error 8%); "
    "and the mean residence time (37.4 min) agrees with the theoretical travel time "
    "(38.6 min; ratio 0.967). Tailing accounts for 47% of the integrated signal, "
    "underscoring the critical role of matrix-diffusion-controlled release in the "
    "monitoring signal. Under steady-state two-phase flow, tracer flux effectively "
    "quantifies per-interval oil-phase flow rates across varying oil–water ratios.")

sp(doc.paragraphs[177],
    "ESP-T integrates fracture propping and tracer monitoring into a single material, "
    "suitable for acid fracturing, deep wells, and high-pressure formations. It provides "
    "a quantitative basis for field-scale production prediction and fracturing parameter "
    "optimization, with broad industrial application potential in unconventional oil "
    "and gas development.")

# ============================================================
doc.save(DST)
print("Nature-level polish complete. ESP_polished.docx updated.")