#!/usr/bin/env python3
"""Polish: title, abstract, keywords, dedup, DSC removal, conclusions."""
from docx import Document
import shutil

src = "四氧化三铁环氧树脂拟合/ESP-T_完整版.docx"
shutil.copy2(src, src.replace(".docx", "_polish_backup.docx"))
doc = Document(src)

def rewrite(idx, new_text):
    p = doc.paragraphs[idx]
    for r in p.runs: r.text = ""
    p.runs[0].text = new_text

# 1. TITLE
rewrite(0,
    "A Piecewise Advection-Dispersion Model Integrated with an Oleophilic "
    "Epoxy/Fe3O4 Tracer Proppant for Per-Interval Production Allocation"
)
print("1/8 Title")

# 2. ABSTRACT
rewrite(2,
    "Abstract: Per-interval oil production monitoring is essential for "
    "evaluating stimulation effectiveness in unconventional reservoirs, yet "
    "existing tracer technologies lack both a quantitative transport model "
    "linking release kinetics to production rates and a durable oil-phase "
    "monitoring capability. Here we address this dual gap by (i) developing a "
    "piecewise advection-dispersion model with smooth tanh transition that "
    "decomposes tracer breakthrough curves into a Gaussian pulse component "
    "(shut-in accumulation slug) and an erfc tailing component "
    "(matrix-diffusion-controlled sustained release), and (ii) validating "
    "this model using an oleophilic tracer proppant (ESP-T) fabricated by "
    "encapsulating stearic acid-modified nano-Fe3O4 (nano-Fe3O4@SA) within "
    "an epoxy resin matrix via emulsion polymerization. The model achieves "
    "R2 = 0.9939 with the erfc tail accounting for 47% of the integrated "
    "tracer signal, demonstrating that sustained matrix-diffusion-controlled "
    "release dominates long-term monitoring. The fitted flow rate "
    "(0.46 mL/min) agrees closely with the pump-set rate (0.50 mL/min, "
    "8% relative error), and the Peclet number (Pe = 0.934) independently "
    "corroborates the non-Fickian release mechanism identified via "
    "Korsmeyer-Peppas kinetics (n = 0.45-0.85, R2 > 0.90). ESP-T exhibits "
    "sphericity exceeding 0.9, thermal stability up to 357.27 degC, and a "
    "water contact angle of 104.6 deg (versus 72.3 deg for neat epoxy), "
    "yielding a water-resistant, oil-permeable transport characteristic "
    "with oil filtration time reduced by 66.1%. Under steady-state two-phase "
    "flow, tracer flux quantitatively tracks oil-phase production rates "
    "across varying oil-water ratios. This coupled experimental-modeling "
    "framework integrates fracture propping with long-term production "
    "allocation, offering a broadly applicable platform for unconventional "
    "reservoir management."
)
print("2/8 Abstract")

# 3. KEYWORDS
rewrite(4,
    "Key words: Unconventional oil and gas reservoirs; Hydraulic fracturing; "
    "Tracer proppant; Epoxy resin; Fe3O4 nanoparticles; Advection-dispersion "
    "model; Piecewise modeling; Production allocation; Release kinetics"
)
print("3/8 Keywords")

# 4. Fix grammar in para [7]
p7 = doc.paragraphs[7]
if "a more than 50%" in p7.text:
    full = p7.text.replace("a more than 50%", "more than 50%")
    for r in p7.runs: r.text = ""
    p7.runs[0].text = full
print("4/8 Grammar")

# 5. Remove DSC [77]
for r in doc.paragraphs[77].runs: r.text = ""
print("5/8 DSC removed")

# 6. Remove duplicate ADE [128]
for r in doc.paragraphs[128].runs: r.text = ""
print("6/8 Duplicate ADE removed")

# 7. Rename 3.7
rewrite(125, "3.7 Quantitative Production Allocation via Piecewise ADE Modeling")
print("7/8 Section 3.7")

# 8a. Conclusions [159] - model first
rewrite(159,
    "This study addressed the dual challenge of quantitative tracer-signal "
    "interpretation and durable oil-phase tracer delivery by (i) developing a "
    "piecewise advection-dispersion model with tanh blending that decomposes "
    "tracer breakthrough curves into physically meaningful Gaussian-pulse and "
    "erfc-tail components, and (ii) validating this model using an oleophilic "
    "epoxy/Fe3O4 tracer proppant (ESP-T). The model achieves R2 = 0.9939, "
    "with the erfc tailing component accounting for 47% of the integrated "
    "tracer signal, and the independently fitted flow rate (0.46 mL/min) "
    "agrees within 8% of the pump-set value (0.50 mL/min). The Peclet number "
    "(Pe = 0.934) independently corroborates the non-Fickian "
    "diffusion-relaxation release mechanism identified from Korsmeyer-Peppas "
    "kinetics. Nano-Fe3O4@SA nanoparticles were synthesized via "
    "coprecipitation with stearic acid surface modification and uniformly "
    "encapsulated within an epoxy resin matrix via emulsion polymerization; "
    "SEM and elemental mapping confirmed nanocluster-level dispersion with "
    "sphericity and roundness exceeding 0.9."
)

# 8b. [161] - material performance
rewrite(161,
    "ESP-T delivers a balanced performance profile suited to downhole "
    "conditions: bulk density of 0.646 g/cm3 (below water, enabling "
    "self-suspension in fracturing fluids); apparent density of 1.072 g/cm3 "
    "(indicative of compact internal packing); crush rate of 2.9% at 52 MPa; "
    "acid solubility of 3.3% (meeting the <=5% industry standard); and an "
    "initial decomposition temperature of 357.27 degC, far exceeding typical "
    "downhole temperatures (80-200 degC). The incorporation of "
    "nano-Fe3O4@SA elevates the water contact angle from 72.3 deg to "
    "104.6 deg, and the oil filtration time of 5 min 11 s represents "
    "a 66.1% reduction relative to pure epoxy microspheres, validating the "
    "designed water-resistant, oil-permeable transport characteristic."
)

# 8c. [163] - release kinetics + two-phase
rewrite(163,
    "Tracer release from ESP-T at 30-120 degC follows the Korsmeyer-Peppas "
    "model (R2 > 0.90) with diffusion exponents n = 0.45-0.85 across all "
    "tested temperatures, confirming an anomalous transport mechanism "
    "co-governed by Fickian diffusion and Case-II relaxation. Cumulative Fe "
    "release at 120 degC exceeds 2.0 mg/L over 14 days, meeting ICP-MS "
    "detection requirements for long-term downhole monitoring. The piecewise "
    "ADE model further reveals that 47% of the total tracer signal originates "
    "from the erfc tail component, quantitatively establishing "
    "matrix-diffusion-controlled sustained release as the dominant mechanism "
    "under flow conditions. Under steady-state oil-water two-phase flow, "
    "tracer flux (FO) is independent of total flow rate and scales with the "
    "oil-water ratio, enabling quantitative per-interval oil production "
    "allocation after single-phase calibration."
)

# 8d. [165] - summary
rewrite(165,
    "In summary, this work establishes a coupled experimental-modeling "
    "framework in which a physically grounded piecewise ADE model, validated "
    "by an oleophilic epoxy/Fe3O4 tracer proppant, translates sustained "
    "tracer release signals into quantitative per-interval production "
    "allocation. The framework is applicable to acid fracturing, deep-well, "
    "and high-pressure scenarios where per-interval production monitoring is "
    "critical. We emphasize that the current validation is limited to "
    "laboratory-scale single-interval experiments with dodecane as a model "
    "oil; field-scale trials under multi-interval, multiphase conditions "
    "with crude oil are necessary to establish the operational reliability "
    "and economic viability of this approach."
)
print("8/8 Conclusions")

out = src.replace("完整版", "终稿")
doc.save(out)
print("Saved: " + out)
