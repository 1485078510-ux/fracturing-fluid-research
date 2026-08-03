#!/usr/bin/env python3
"""Polish pass: split long sentences, remove em dashes, smooth flow."""
from docx import Document

doc = Document('四氧化三铁环氧树脂拟合/ESP-T_终稿.docx')

# ── EM DASH FIXES ──
em_dash_fixes = {
    85: [  # WCA paragraph
        ('surface—an increment', 'surface, an increment'),
        ('transition. This transformation', 'transition, driven by'),
    ],
    97: [  # Conductivity closing
        ('proppant pores and markedly reducing flow resistance. ESP-T thus exhibits a water-resistant, oil-permeable characteristic:',
         'proppant pores, markedly reducing flow resistance. ESP-T thus exhibits a water-resistant, oil-permeable characteristic:'),
    ],
    146: [  # Deepened discussion
        ('explanation—a conclusion', 'explanation, a conclusion'),
        ('exercise itself—the independently', 'exercise itself. The independently'),
        ('signal alone—an essential', 'signal alone, an essential'),
        ('extended-duration process—a conclusion', 'extended-duration process, a conclusion'),
        ('coefficients—they correspond', 'coefficients: they correspond'),
    ],
    158: [  # Conclusions opening
        ('kinetics—measured', 'kinetics, measured'),
        ('models—and', 'models, and'),
    ],
    160: [  # Conclusions model
        ('release—rather', 'release, rather'),
    ],
}

for para_idx, replacements in em_dash_fixes.items():
    p = doc.paragraphs[para_idx]
    t = p.text
    changed = False
    for old, new in replacements:
        if old in t:
            t = t.replace(old, new)
            changed = True
    if changed:
        for r in p.runs: r.text = ''
        p.runs[0].text = t
        print(f'[{para_idx}] em dashes fixed')

# ── Replace remaining em dashes (—) with commas in non-math paragraphs ──
# Skip derivation [126] and figure captions
skip_idx = {126, 68, 72, 81, 88, 94, 99, 121, 145, 155}
for i, p in enumerate(doc.paragraphs):
    if i in skip_idx: continue
    t = p.text
    if '—' in t:
        # Only fix em dashes used as punctuation (not in chemical formulas)
        # Replace em dash with comma or period
        t = t.replace('—', ', ')
        # Clean up double commas
        t = t.replace(', ,', ',')
        for r in p.runs: r.text = ''
        p.runs[0].text = t
        print(f'[{i}] general em dash cleanup')

# ── SPLIT LONG SENTENCES ──
# [2] Abstract: split the 64-word sentence
p2 = doc.paragraphs[2]
t2 = p2.text
old2 = (
    'Here we address this dual gap by (i) developing a piecewise '
    'advection-dispersion model with smooth tanh transition that decomposes '
    'tracer breakthrough curves into a Gaussian pulse component (shut-in '
    'accumulation slug) and an erfc tailing component (matrix-diffusion-'
    'controlled sustained release), and (ii) validating this model using an '
    'oleophilic tracer proppant (ESP-T) fabricated by encapsulating stearic '
    'acid-modified nano-Fe3O4 (nano-Fe3O4@SA) within an epoxy resin matrix '
    'via emulsion polymerization.'
)
new2 = (
    'Here we address this dual gap. We develop a piecewise '
    'advection-dispersion model with smooth tanh transition that decomposes '
    'tracer breakthrough curves into a Gaussian pulse component, representing '
    'the shut-in accumulation slug, and an erfc tailing component, representing '
    'matrix-diffusion-controlled sustained release. We validate this model '
    'using an oleophilic tracer proppant (ESP-T) fabricated by encapsulating '
    'stearic acid-modified nano-Fe3O4 (nano-Fe3O4@SA) within an epoxy resin '
    'matrix via emulsion polymerization.'
)
if old2 in t2:
    t2 = t2.replace(old2, new2)
    for r in p2.runs: r.text = ''
    p2.runs[0].text = t2
    print('[2] long sentence split')

# [85] WCA mechanism - split 54w sentence
p85 = doc.paragraphs[85]
t85 = p85.text
old85 = (
    'This transformation, driven by the stearic acid modification of '
    'nano-Fe3O4@SA: the carboxyl groups (-COOH) of stearic acid '
    '(C17H35COOH) undergo a coordination reaction with surface hydroxyl '
    'groups of nano-Fe3O4@SA to form stable chemical bonds, while the long '
    'alkyl chains (-C17H35) orient outward from the epoxy matrix, '
    'constructing a hydrophobic film on the proppant surface.'
)
new85 = (
    'This transformation, driven by the stearic acid modification of '
    'nano-Fe3O4@SA. The carboxyl groups (-COOH) of stearic acid '
    '(C17H35COOH) coordinate with surface hydroxyl groups of '
    'nano-Fe3O4@SA to form stable chemical bonds, while the long alkyl '
    'chains (-C17H35) orient outward from the epoxy matrix, '
    'constructing a hydrophobic film on the proppant surface.'
)
if old85 in t85:
    t85 = t85.replace(old85, new85)
    for r in p85.runs: r.text = ''
    p85.runs[0].text = t85
    print('[85] long sentence split')

# [154] engineering significance - split the long middle sentence
p154 = doc.paragraphs[154]
t154 = p154.text
old154 = (
    'The decomposition into Gaussian-pulse and erfc-tail components further '
    'distinguishes two information channels: the shut-in accumulation slug, '
    'sensitive to shut-in duration and tracer loading, and the sustained '
    'tail, which reflects the intrinsic release characteristics of the '
    'proppant and persists without repeated shut-in operations.'
)
new154 = (
    'The decomposition into Gaussian-pulse and erfc-tail components further '
    'distinguishes two information channels. The shut-in accumulation slug '
    'is sensitive to shut-in duration and tracer loading; the sustained '
    'tail reflects the intrinsic release characteristics of the proppant '
    'and persists without repeated shut-in operations.'
)
if old154 in t154:
    t154 = t154.replace(old154, new154)
    for r in p154.runs: r.text = ''
    p154.runs[0].text = t154
    print('[154] long sentence split')

doc.save('四氧化三铁环氧树脂拟合/ESP-T_终稿.docx')
print('\nPolishing pass complete.')