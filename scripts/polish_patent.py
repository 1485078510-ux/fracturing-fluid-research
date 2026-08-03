#!/usr/bin/env python3
"""
修复专利权利要求编号和内容问题。
采用稳健的逐条识别方法，基于行首编号模式匹配。
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
import re

input_path = r"c:\Users\郝\Desktop\claude\荧光压裂液\专利申请文件_正式格式.docx"
output_path = r"c:\Users\郝\Desktop\claude\荧光压裂液\专利申请文件_正式格式_修改版.docx"

doc = Document(input_path)

def replace_paragraph_text(para, new_text):
    """替换段落文本，保留第一个run的格式"""
    if not para.runs:
        run = para.add_run(new_text)
        return
    first_run = para.runs[0]
    for i in range(len(para.runs) - 1, 0, -1):
        para.runs[i]._element.getparent().remove(para.runs[i]._element)
    first_run.text = new_text

def insert_paragraph_after(para, text, doc_obj):
    """在指定段落后插入新段落，复制其格式"""
    new_para = doc_obj.add_paragraph()
    new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
    new_para.paragraph_format.line_spacing = para.paragraph_format.line_spacing
    new_para.alignment = para.alignment
    para._element.addnext(new_para._element)
    run = new_para.add_run(text)
    if para.runs:
        ref_run = para.runs[0]
        run.font.name = ref_run.font.name
        run.font.size = ref_run.font.size
        run.font.bold = ref_run.font.bold
        try:
            run.element.rPr.rFonts.set(qn('w:eastAsia'),
                ref_run.element.rPr.rFonts.get(qn('w:eastAsia')) if ref_run.element.rPr is not None else '宋体')
        except:
            pass
    return new_para

def insert_paragraph_before(para, text, doc_obj):
    """在指定段落前插入新段落，复制其格式"""
    new_para = doc_obj.add_paragraph()
    new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
    new_para.paragraph_format.line_spacing = para.paragraph_format.line_spacing
    new_para.alignment = para.alignment
    para._element.addprevious(new_para._element)
    run = new_para.add_run(text)
    if para.runs:
        ref_run = para.runs[0]
        run.font.name = ref_run.font.name
        run.font.size = ref_run.font.size
        run.font.bold = ref_run.font.bold
        try:
            run.element.rPr.rFonts.set(qn('w:eastAsia'),
                ref_run.element.rPr.rFonts.get(qn('w:eastAsia')) if ref_run.element.rPr is not None else '宋体')
        except:
            pass
    return new_para

# ================================================================
# Step 1: 修复附图说明重复
# ================================================================
print("=== 修复附图说明重复 ===")
# 查找所有包含重复图题的段落（仅有图题没有详细描述的简短版本）
for para in doc.paragraphs:
    text = para.text.strip()
    # 简洁图题（后面紧跟的是空行或下一张图），在详细描述之后
    if text == "图1 改性稀土铝酸盐荧光粉的结构示意图":
        replace_paragraph_text(para, "")
        print("已删除图1重复图题")
    elif text == "图2 荧光压裂液体系现场施工工艺流程图":
        replace_paragraph_text(para, "")
        print("已删除图2重复图题")
    elif text == "图3 压裂裂缝荧光示踪方法的流程框图":
        replace_paragraph_text(para, "")
        print("已删除图3重复图题")

# ================================================================
# Step 2: 修复摘要附图说明
# ================================================================
print("\n=== 修复摘要附图 ===")
for para in doc.paragraphs:
    if "摘要附图为本发明说明书的图1" in para.text:
        new_text = ("摘要附图为说明书图1（改性稀土铝酸盐荧光粉的结构示意图），"
                   "该图展示了荧光粉基体-硅烷偶联剂化学键合内层-PEG物理屏蔽外层"
                   "的三层核壳结构关系，是反映本发明技术方案核心构思的代表性图示。")
        replace_paragraph_text(para, new_text)
        print("已精简摘要附图说明")

# ================================================================
# Step 3: 修复权利要求书
# ================================================================
print("\n=== 修复权利要求书 ===")

# 先定位"权利要求书"标题后的所有权利要求段落
claim_started = False
claim_paras = []  # (paragraph_object, claim_number_as_int_or_None)

for para in doc.paragraphs:
    text = para.text.strip()
    if "权利要求书" == text or "权利要求书" in text:
        claim_started = True
        continue
    if claim_started:
        # 检查是否已经进入摘要（权利要求书结束）
        if text.startswith("摘要") and len(text) <= 4:
            break
        # 检查是否是空段落且在权利要求区域末尾
        # 识别权利要求段落：以数字.开头的
        m = re.match(r'^(\d+)\.\s', text)
        if m:
            claim_paras.append((para, int(m.group(1))))
        elif claim_paras and text == "":
            # 空段落，可能是分隔
            pass
        elif claim_paras:
            # 非空非编号段落，可能是上一权利要求的续行（多段落权利要求）
            claim_paras.append((para, None))  # None表示续行

print(f"识别到 {len(claim_paras)} 个权利要求相关段落")

# 打印识别到的权利要求
for para, num in claim_paras:
    if num is not None:
        print(f"  权利要求 {num}: {para.text[:80]}...")

# 现在根据内容重新规划权利要求结构
# 目标结构：
# 1. 产品独立权利要求（改性荧光粉基体+双层）
# 2. 硅烷偶联剂选择（从属1）
# 3. PEG参数（从属1）
# 4. PEG优选（从属3）
# 5. 荧光粉基体优选（从属1）
# 6. 硅烷偶联剂优选（从属2）
# 7. 分散助剂体系（从属1-6）
# 8. 分散助剂优选（从属7）
# 9. 制备方法独立权利要求
# 10. 制备方法优选（从属9）
# 11. 制备方法步骤3（从属9）
# 12. 压裂液体系独立权利要求
# 13. 压裂液组分选择（从属12）
# 14. 母液预配（从属12或13）
# 15. 示踪方法独立权利要求
# 16. 添加段选择（从属15）
# 17. 功能切换特征（从属15）
# 18. 用途独立权利要求

# 构建新权利要求的完整文本
new_claims = []

# 首先提取原权利要求中的实质内容
# 原权利要求1: 产品（基体+双层）
# 需要补全 Sr₂MgSi₂O₇:Eu²⁺,Dy³⁺ 和粒度范围
new_claims.append(
    "1. 一种改性稀土铝酸盐荧光粉，其特征在于，包括稀土铝酸盐长余辉荧光粉基体和包覆于所述基体表面的双层改性层；"
    "所述稀土铝酸盐长余辉荧光粉基体为SrAl₂O₄:Eu²⁺,Dy³⁺、Sr₄Al₁₄O₂₅:Eu²⁺,Dy³⁺、CaAl₂O₄:Eu²⁺,Nd³⁺"
    "或Sr₂MgSi₂O₇:Eu²⁺,Dy³⁺中的一种或多种，粒度为400~2000目；"
    "所述双层改性层包括内层和外层，所述内层为硅烷偶联剂化学键合层，所述硅烷偶联剂通过Si-O-Al共价键锚固于荧光粉基体表面；"
    "所述外层为聚乙二醇（PEG）物理屏蔽层，所述聚乙二醇通过物理吸附沉积于内层硅烷偶联剂表面。"
)

# 原权利要求2: 硅烷偶联剂选择（去除"优选"）
new_claims.append(
    "2. 根据权利要求1所述的改性稀土铝酸盐荧光粉，其特征在于，所述硅烷偶联剂选自"
    "3-氨基丙基三乙氧基硅烷（KH550）、3-氨基丙基三甲氧基硅烷（APTMS）、"
    "γ-甲基丙烯酰氧基丙基三甲氧基硅烷（KH570）或乙烯基三甲氧基硅烷中的一种或多种；"
    "所述硅烷偶联剂的用量为荧光粉基体质量的0.5~5.0 wt%。"
)

# 原权利要求3: PEG参数（去除"优选"）
new_claims.append(
    "3. 根据权利要求1所述的改性稀土铝酸盐荧光粉，其特征在于，所述聚乙二醇的数均分子量Mn为1000~10000，"
    "用量为荧光粉基体质量的1.0~10.0 wt%。"
)

# 新权利要求4: PEG优选范围
new_claims.append(
    "4. 根据权利要求3所述的改性稀土铝酸盐荧光粉，其特征在于，所述聚乙二醇的数均分子量Mn为2000~6000，"
    "用量为荧光粉基体质量的2.0~5.0 wt%。"
)

# 原权利要求4内容（粒度优选）+ 改为权利要求5
new_claims.append(
    "5. 根据权利要求1所述的改性稀土铝酸盐荧光粉，其特征在于，所述荧光粉基体为SrAl₂O₄:Eu²⁺,Dy³⁺，"
    "粒度为800~1200目。"
)

# 新权利要求6: 硅烷偶联剂优选
new_claims.append(
    "6. 根据权利要求2所述的改性稀土铝酸盐荧光粉，其特征在于，所述硅烷偶联剂为KH550，"
    "用量为荧光粉基体质量的1.0~3.0 wt%。"
)

# 原权利要求5内容（分散助剂体系）
new_claims.append(
    "7. 根据权利要求1至6中任一项所述的改性稀土铝酸盐荧光粉，其特征在于，还包括协同分散助剂体系，"
    "所述分散助剂体系包括螯合剂和非离子表面活性剂；"
    "所述螯合剂选自柠檬酸、乙二胺四乙酸（EDTA）或其钠盐、酒石酸中的一种或多种，用量为荧光粉基体"
    "质量的0.05~0.5 wt%；"
    "所述非离子表面活性剂选自烷基酚聚氧乙烯醚系列、脂肪醇聚氧乙烯醚系列或聚山梨酯系列中的一种或多种，"
    "用量为荧光粉基体质量的0.01~0.2 wt%。"
)

# 新权利要求8: 分散助剂优选
new_claims.append(
    "8. 根据权利要求7所述的改性稀土铝酸盐荧光粉，其特征在于，所述螯合剂为柠檬酸，"
    "所述非离子表面活性剂为Triton X-100。"
)

# 原权利要求6内容（制备方法独立权利要求）
new_claims.append(
    "9. 一种权利要求1至8中任一项所述改性稀土铝酸盐荧光粉的制备方法，其特征在于，包括以下步骤："
    "步骤（1）：将稀土铝酸盐荧光粉基体分散于无水乙醇-去离子水混合溶剂（体积比90:10~98:2）中，加入占"
    "荧光粉基体质量0.5~5.0 wt%的硅烷偶联剂，用乙酸调节pH至4~6，在15~40°C下搅拌反应1~4小时，50~80°C"
    "干燥固化，得到硅烷偶联剂化学键合荧光粉；"
    "步骤（2）：将步骤（1）所得产物分散于去离子水中，加入占荧光粉基体质量1.0~10.0 wt%的聚乙二醇"
    "（Mn=1000~10000），在15~40°C下搅拌反应0.5~2小时，分离、洗涤、干燥，得到所述改性稀土铝酸盐荧光粉。"
)

# 原权利要求7内容（制备方法优选）
new_claims.append(
    "10. 根据权利要求9所述的制备方法，其特征在于，步骤（1）中所述混合溶剂的体积比为95:5（无水乙醇:"
    "去离子水）；步骤（2）中所述聚乙二醇的数均分子量Mn为2000~6000。"
)

# 新权利要求11: 制备方法步骤（3）
new_claims.append(
    "11. 根据权利要求9所述的制备方法，其特征在于，还包括步骤（3）：将步骤（2）所得产物与螯合剂和"
    "非离子表面活性剂按比例混合，或将螯合剂和非离子表面活性剂预先溶解于分散介质中，以便在配制荧光"
    "母液时实现协同分散效果。"
)

# 原权利要求8内容（压裂液体系）
new_claims.append(
    "12. 一种荧光压裂液体系，其特征在于，包括以下组分："
    "（a）权利要求1至8中任一项所述的改性稀土铝酸盐荧光粉；"
    "（b）稠化剂；"
    "（c）交联剂；"
    "（d）破胶剂；"
    "（e）水。"
)

# 原权利要求9内容（组分选择）
new_claims.append(
    "13. 根据权利要求12所述的荧光压裂液体系，其特征在于，所述稠化剂选自羟丙基胍胶（HPG）、瓜尔胶、"
    "羧甲基羟丙基胍胶（CMHPG）中的一种或多种，在压裂液终液中的浓度为0.3~1.0 wt%；"
    "所述交联剂选自有机硼交联剂、有机锆交联剂或有机钛交联剂中的一种或多种，在压裂液终液中的浓度为"
    "0.1~0.5 vol%；"
    "所述破胶剂选自过硫酸铵、过硫酸钾、胶囊包裹过硫酸盐或酶破胶剂中的一种或多种，在压裂液终液中的"
    "浓度为0.02~0.3 wt%。"
)

# 原权利要求10内容（母液预配）
new_claims.append(
    "14. 根据权利要求12或13所述的荧光压裂液体系，其特征在于，采用母液预配与在线稀释相结合的方式制备："
    "首先，将所述改性稀土铝酸盐荧光粉按20~80 g/L的浓度分散于去离子水中，加入分散助剂，通过高速剪切"
    "和/或超声辅助制成荧光悬浮母液；"
    "然后，将所述母液按体积比0.1~2.0%通过在线混合装置注入所述稠化剂的基液主流中，经静态混合器混合"
    "得到荧光压裂液终液。"
)

# 原权利要求11内容（示踪方法）
new_claims.append(
    "15. 一种利用权利要求12至14中任一项所述荧光压裂液体系的压裂裂缝荧光示踪方法，其特征在于，包括"
    "以下步骤："
    "注入步骤：将所述荧光压裂液终液与支撑剂混合后，通过压裂泵注系统泵入目标压裂层段；"
    "关井破胶步骤：停泵关井，在储层温度60~150°C下密闭维持6~48小时，使交联冻胶中的交联键在破胶剂"
    "作用下断裂，同时改性荧光粉外层的聚乙二醇在破胶剂氧化环境与储层温度的协同作用下脱附，暴露内层"
    "硅烷偶联剂的活性官能团，所述活性官能团与裂缝壁面岩石表面发生锚定；"
    "返排步骤：开井返排，携带破胶残渣和未锚定的游离荧光粉颗粒排出井筒；"
    "检测步骤：在压后取心获得的含裂缝岩心上，以紫外光源照射，观察裂缝壁面的荧光分布，荧光分布区域"
    "即为压裂液的波及区域。"
)

# 原权利要求12内容（添加段）
new_claims.append(
    "16. 根据权利要求15所述的方法，其特征在于，在所述注入步骤中，荧光母液的在线添加在前置液段、"
    "携砂液段或顶替液段中的一段或多段进行。"
)

# 原权利要求13内容（功能切换）
new_claims.append(
    "17. 根据权利要求15所述的方法，其特征在于，在所述关井破胶步骤中，所述改性荧光粉从分散态向锚定态"
    "的功能切换完全利用压裂施工自身的注入-关井-返排工序时序驱动，以压裂液配方中固有的过硫酸铵破胶剂"
    "为脱附触发条件，无需额外添加专用触发剂。"
)

# 原权利要求14内容（用途）
new_claims.append(
    "18. 权利要求1至8中任一项所述的改性稀土铝酸盐荧光粉在水力压裂裂缝监测中的应用。"
)

# 现在替换权利要求段落
# 策略：收集所有claim_paras中的"主段落"（编号不为None的），用新内容替换
main_claim_paras = [(p, n) for p, n in claim_paras if n is not None]

print(f"\n原权利要求主段落数: {len(main_claim_paras)}")
print(f"新权利要求数: {len(new_claims)}")

# 找出所有续行段落（编号为None的）
continuation_paras = [(p, n) for p, n in claim_paras if n is None]
print(f"续行段落数: {len(continuation_paras)}")

# 删除所有续行段落（它们的内容将被合并到主段落中）
for para, _ in continuation_paras:
    para._element.getparent().remove(para._element)
    print(f"  已删除续行段落")

# 现在用新内容替换主段落
for i, (para, old_num) in enumerate(main_claim_paras):
    if i < len(new_claims):
        replace_paragraph_text(para, new_claims[i])
        new_num = i + 1
        print(f"  原权利要求{old_num} -> 新权利要求{new_num}")
    else:
        # 多余的旧权利要求段落，删除
        para._element.getparent().remove(para._element)
        print(f"  删除多余的原权利要求{old_num}")

# 如果新权利要求比旧的多，在最后插入额外的
if len(new_claims) > len(main_claim_paras):
    last_para = main_claim_paras[-1][0]
    for i in range(len(main_claim_paras), len(new_claims)):
        insert_paragraph_after(last_para, new_claims[i], doc)
        last_para = last_para._element.getnext()
        # 获取对应的paragraph对象
        for p in doc.paragraphs:
            if p._element is last_para:
                last_para = p
                break
        print(f"  插入新权利要求{i+1}")

# ================================================================
# Step 4: 修复说明书中的语言问题
# ================================================================
print("\n=== 修复语言问题 ===")

for para in doc.paragraphs:
    text = para.text

    # 修复自我评价性语言
    if "本发明最具创新性的设计在于" in text:
        new_text = text.replace("本发明最具创新性的设计在于将", "本发明的关键设计在于将")
        replace_paragraph_text(para, new_text)
        print("已修改: 最具创新性 -> 关键设计")

    # 修复"无缝兼容" (专利中避免绝对化用语)
    if "无缝兼容" in text:
        new_text = text.replace("无缝兼容", "良好兼容")
        replace_paragraph_text(para, new_text)
        print("已修改: 无缝兼容 -> 良好兼容")

    # 修复"创造性地转化"
    if "创造性地转化为" in text:
        new_text = text.replace("创造性地转化为", "转化为")
        replace_paragraph_text(para, new_text)
        print("已修改: 创造性地转化为 -> 转化为")

    # 修复过于具体的自由基机理描述
    if "硫酸根自由基（SO₄•⁻）攻击PEG醚键" in text:
        new_text = text.replace(
            "过硫酸铵热分解产生的硫酸根自由基（SO₄•⁻）攻击PEG醚键（C-O-C），引发氧化链断裂，生成低分子量PEG碎片和含氧低聚物",
            "过硫酸铵热分解产生的氧化性物种攻击PEG醚键（C-O-C），引发链断裂，生成低分子量PEG碎片和含氧低聚物"
        )
        replace_paragraph_text(para, new_text)
        print("已修改: 自由基细节 -> 通用表述")

# ================================================================
# Step 5: 修改摘要中的"无缝兼容"
# ================================================================
for para in doc.paragraphs:
    if "与现有工艺无缝兼容" in para.text:
        new_text = para.text.replace("与现有工艺无缝兼容", "与现有工艺兼容")
        replace_paragraph_text(para, new_text)
        print("已修改摘要: 无缝兼容 -> 兼容")

# ================================================================
# 保存
# ================================================================
doc.save(output_path)
print(f"\n✅ 文件已保存至：{output_path}")
print("修改完成！")