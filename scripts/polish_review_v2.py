# -*- coding: utf-8 -*-
"""Polish the review paper - reads original DOCX, creates polished version with highlights"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json, os

SRC = r"c:\Users\郝\Desktop\claude\向晶\页岩气藏地应力演化与裂缝扩展耦合机制研究进展与展望-V2.docx"
DST = r"c:\Users\郝\Desktop\claude\向晶\页岩气藏地应力演化与裂缝扩展耦合机制研究进展与展望-润色版.docx"

def add_highlight(run, color="yellow"):
    rPr = run._element.get_or_add_rPr()
    h = OxmlElement("w:highlight")
    h.set(qn("w:val"), color)
    rPr.append(h)

def get_font_info(para):
    if para.runs:
        r = para.runs[0]
        return {"size": r.font.size, "bold": r.font.bold, "name": r.font.name}
    return {"size": Pt(12), "bold": None, "name": "SimSun"}

def apply_font(run, fi):
    if fi["size"]: run.font.size = fi["size"]
    if fi["bold"] is not None: run.font.bold = fi["bold"]
    if fi["name"]: run.font.name = fi["name"]

# Load replacements from JSON
json_path = r"c:\Users\郝\Desktop\claude\向晶\replacements.json"
with open(json_path, "r", encoding="utf-8") as f:
    ALL_CHANGES = json.load(f)

doc = Document(SRC)
total = 0
matched = 0

for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue

    # Find all applicable replacements for this paragraph
    applicable = []
    for entry in ALL_CHANGES:
        keyword = entry["keyword"]
        reps = entry["replacements"]
        if keyword in para.text:
            for r in reps:
                old = r["old"]
                new = r["new"]
                if old in para.text:
                    applicable.append((old, new))

    if not applicable:
        continue

    # Deduplicate
    seen = set()
    unique = []
    for old, new in applicable:
        if old not in seen:
            unique.append((old, new))
            seen.add(old)

    # Sort by length (longest first to avoid partial replacements)
    unique.sort(key=lambda x: len(x[0]), reverse=True)

    # Find all replacement positions in original text
    full = para.text
    # Apply all replacements to get the final text
    positions = []
    temp_text = full
    offset = 0  # Track cumulative offset from previous replacements
    for old, new in unique:
        idx = temp_text.find(old)
        if idx >= 0:
            # Record position in the evolving text
            positions.append((idx, len(old), len(new)))
            temp_text = temp_text[:idx] + new + temp_text[idx+len(old):]

    if not positions:
        continue

    # Now build segments from the original text with proper highlighting
    # Strategy: apply replacements sequentially, tracking which chars are new
    # Simple approach: split text into segments based on matches
    # Find matches in original text
    matches = []
    for old, new in unique:
        start = full.find(old)
        if start >= 0:
            matches.append((start, start + len(old), new))

    matches.sort(key=lambda x: x[0])

    # Merge overlapping matches
    merged = []
    for m in matches:
        if merged and m[0] < merged[-1][1]:
            prev = merged.pop()
            # Keep the new text from the last match covering this region
            merged.append((prev[0], max(prev[1], m[1]), m[2]))
        else:
            merged.append(m)

    # Build run segments
    segments = []  # (text, is_highlighted)
    cursor = 0
    for start, end, new_txt in merged:
        if cursor < start:
            segments.append((full[cursor:start], False))
        segments.append((new_txt, True))
        cursor = end
    if cursor < len(full):
        segments.append((full[cursor:], False))

    # Replace paragraph runs
    fi = get_font_info(para)
    for run in list(para.runs):
        run._element.getparent().remove(run._element)

    for seg_text, highlight in segments:
        if not seg_text:
            continue
        r = para.add_run(seg_text)
        apply_font(r, fi)
        if highlight:
            add_highlight(r, "yellow")

    matched += 1
    total += len(unique)

doc.save(DST)
print(f"Done! Matched {matched} paragraphs, {total} replacements. Saved to DST.")
