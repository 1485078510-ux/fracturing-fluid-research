# -*- coding: utf-8 -*-
"""
Task 12: Final polish — title, keywords, global consistency.
Edits ESP-T_v2_manuscript.docx in place. Idempotent: each edit is guarded by
an exact-match check on the current paragraph text; already-applied edits are
skipped. Run: python scripts/polish_task12.py
"""
import copy
import sys

import docx
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

PATH = r"四氧化三铁环氧树脂拟合/ESP-T_投稿文件/ESP-T_v2_manuscript.docx"
TITLE = ("Flow-Rate Self-Calibration from Tracer-Proppant Breakthrough Curves: "
         "A Coupled Release–Transport Model for Per-Stage Production Allocation")
KEYWORDS = ("Tracer proppant; Breakthrough curve; Advection–dispersion equation; "
            "Release kinetics; Self-calibration; Production allocation; "
            "Two-phase flow; Epoxy resin")

CAPTIONS = {
    1: "Schematic of breakthrough-curve (BTC) generation in the core displacement experiment. "
       "During the 96-h shut-in, tracer diffuses from the epoxy matrix of the ESP-T pack and "
       "accumulates in the pore space of the pack; on flowback at the pump setting, the "
       "accumulated slug is swept to the sampling point as a Gaussian pulse (Source I), followed "
       "by the sustained diffusion-controlled release tail (Source II).",
    2: "ESP-T characterization summary (four panels) verifying the material prerequisites of the "
       "coupled model: SEM-EDS cross-section showing Fe distributed throughout the particle; "
       "TGA/DTG thermogram with decomposition onset 357 °C; K-P batch release fits at "
       "30–120 °C (n = 0.45–0.85); water contact angle (104.6°) and "
       "oil/water filtration time ratio 5.53 demonstrating oil-phase selectivity.",
    3: "K-P batch release kinetics (30–120 °C, Mt/M∞ < 0.6) and overlay of the "
       "single-process candidate models on the measured BTC: each single-process model reproduces "
       "only one feature of the curve — the peak or the tail — and none captures both.",
    4: "Model selection. (a) Fits of the five candidate models — dual-component "
       "tanh-blended, single Gaussian, single erfc, exponential decay, and Korsmeyer–Peppas "
       "power law — to the 21-point single-phase BTC. (b) ΔAICc ranking: only the "
       "dual-component model is supported (ΔAICc = 0; every alternative ≥ 32.66).",
    5: "Physical self-calibration. (a) Fitted dual-component model (Eq. 1) with component "
       "decomposition: the shaded Gaussian pulse carries 53% of the integrated signal and the "
       "shaded erfc tail carries 47%, joined across the tanh transition centered on the fitted "
       "crossover t₀ = 25.66 min; residuals beneath the fit show no systematic structure. "
       "(b) Fitted flow rate Q = 0.46 ± 0.02 mL/min from four independent global searches "
       "versus the pump setting of 0.50 mL/min (8% deviation).",
    6: "Time-of-arrival comparison. (a) Measured BTC with peak-time (t_peak = 15 min), half-peak "
       "(t_half ≈ 5 min), and first-moment markers. (b) Flow-rate estimates from the "
       "peak-time (1.31 mL/min, +162%), half-peak (3.93 mL/min, +685%), and first-moment "
       "(0.53 mL/min, +5.8%) methods and the coupled model (0.46 mL/min, −8%) relative to "
       "the pump setting (0.50 mL/min). (c) Qualitative comparison of method capabilities.",
    7: "Independent corroboration of the two-process picture. Left panel: K-P fits to the static "
       "batch release data (n = 0.45–0.85, non-Fickian release). Right panel: fitted "
       "single-phase BTC annotated with the recovered Peclet number (Pe = x/α = 0.934, "
       "dispersion-dominated transport).",
    8: "Transition-width sensitivity. (a) erfc tail fraction versus σ over a six-fold range "
       "(0.5×–3.0× of the fitted value 3.96 min); the fraction stays within "
       "46.7–47.5%. (b) Component decomposition at the fitted σ. (c) Fitted BTCs at "
       "five σ values.",
    9: "Two-phase production allocation. (a) Oil-phase tracer concentration C_oil versus total "
       "flow rate at OWR 4:1, 1:1, and 1:4 (dilution effect). (b) Oil-phase tracer mass flux "
       "F_O = C_oil × Q_oil versus total flow rate: flat within each OWR, confirming the "
       "mass balance. (c) Normalized flux F_O/F_O,ref versus oil flow rate Q_oil, collapsing the "
       "three OWR families onto a single trend (Pearson r = 0.97, RMSD = 8.3%).",
    10: "Field deployment pathway: (1) stage tagging with distinct metal/rare-earth dopants; "
        "(2) single shut-in after fracturing; (3) flowback sampling at intervals guided by "
        "Eq. (1); (4) per-stage BTC fitting to recover stage flow rates Q_i; (5) production "
        "allocation via the flux method during steady production; (6) well-spacing and completion "
        "optimization from per-stage rate trends.",
}


def para_texts(d):
    return [(i, p.text) for i, p in enumerate(d.paragraphs)]


def replace_in_para(p, old, new):
    """Substring replacement across the paragraph's full text, rebuilding runs with
    the first run's formatting. Returns True if applied."""
    text = p.text
    if old not in text:
        return False
    new_text = text.replace(old, new, 1)
    _rebuild(p, [(new_text, None)])
    return True


def set_para_text(p, new_text):
    _rebuild(p, [(new_text, None)])


def _rebuild(p, parts):
    """Rebuild paragraph runs: parts = [(text, bold_or_None)]. Preserves paragraph
    style (pPr) and copies the original first run's rPr formatting."""
    rpr = None
    if p.runs and p.runs[0]._r.rPr is not None:
        rpr = copy.deepcopy(p.runs[0]._r.rPr)
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    for text, bold in parts:
        run = p.add_run(text)
        if rpr is not None:
            run._r.insert(0, copy.deepcopy(rpr))
        if bold:
            run.bold = True


def remove_paragraph(p):
    p._p.getparent().remove(p._p)


def insert_paragraph_after(p, text):
    new_p = p._p.makeelement(qn("w:p"), {})
    if p._p.pPr is not None:
        new_p.append(copy.deepcopy(p._p.pPr))
    p._p.addnext(new_p)
    np = Paragraph(new_p, p._parent)
    np.add_run(text)
    return np


def main():
    d = docx.Document(PATH)
    paras = d.paragraphs
    stats = []

    def applied(name, ok):
        stats.append((name, ok))
        print(("OK  " if ok else "SKIP") + f" {name}")

    # ---- 1. Title ----
    done = False
    for p in paras:
        if p.text == TITLE:
            done = True  # already applied
            break
        if p.text.startswith("Title:"):
            set_para_text(p, TITLE)
            done = p.text == TITLE
            break
    applied("title", done)

    # ---- 2. Keywords (fill empty paragraph after 'Keywords' heading) ----
    for i, p in enumerate(paras):
        if p.text.strip() == "Keywords":
            nxt = paras[i + 1]
            if nxt.text.strip() == "":
                set_para_text(nxt, KEYWORDS)
                applied("keywords", True)
            else:
                applied("keywords (already present)", nxt.text.strip() == KEYWORDS)
            break

    # ---- 3. Symbol consistency: A-area collision in classical solution (2.2) ----
    old = ("For an instantaneous pulse of mass M released at x = 0 and t = 0, the ADE admits "
           "the classical solution C(x, t) = (M/A)·(1/√(4πDt))·exp(−(x "
           "− vt)²/(4Dt)), in which the concentration depends")
    new = ("For an instantaneous pulse released at x = 0 and t = 0, the ADE admits the classical "
           "solution C(x, t) = (M/√(4πDt))·exp(−(x − vt)²/(4Dt)), "
           "with M the released mass per unit cross-sectional area, in which the concentration "
           "depends")
    done = False
    for p in paras:
        if "the classical solution C(x, t)" in p.text:
            done = replace_in_para(p, old, new)
            applied("2.2 remove A-as-area collision", done)
            break

    # ---- 4. c_b / C_rise / C_fall subscript unification ----
    sub_map = [
        ("Crise(t) = cb + (A·d)/√(16παQt·d²) × exp(−z²)",
         "C_rise(t) = c_b + (A·d)/√(16παQt·d²) × exp(−z²)", "eq Crise"),
        ("Cfall(t) = cb + (a/2) × erfc(−z)",
         "C_fall(t) = c_b + (a/2) × erfc(−z)", "eq Cfall"),
        ("C(t) = cb + w(t)·Crise(t) + [1 − w(t)]·Cfall(t)   (Eq. 1)",
         "C(t) = c_b + w(t)·C_rise(t) + [1 − w(t)]·C_fall(t)   (Eq. 1)", "eq blended"),
        ("seven parameters — cb, A, a, α, Q, t₀, and σ",
         "seven parameters — c_b, A, a, α, Q, t₀, and σ", "2.3 parameter list"),
    ]
    for old_s, new_s, name in sub_map:
        found = False
        for p in paras:
            if old_s in p.text:
                found = replace_in_para(p, old_s, new_s)
                applied(name, found)
                break
        if not found:
            applied(name, False)

    # Table 1 'cb' cell -> 'c_b'
    t1 = d.tables[0]
    done = False
    for r in t1.rows:
        if r.cells[1].text.strip() == "cb":
            for para in r.cells[1].paragraphs:
                set_para_text(para, "c_b")
            done = True
    applied("Table 1 cb -> c_b", done)

    # Definition prose for Source I: 'where cb is the baseline concentration'
    done = False
    for p in paras:
        if "where cb is the baseline concentration" in p.text:
            done = replace_in_para(p, "where cb is the baseline concentration",
                                   "where c_b is the baseline concentration")
            break
    applied("2.2 c_b definition prose", done)

    # Prose after Eq. (1) also references Crise(t)/Cfall(t)
    done = done2 = False
    for p in paras:
        if "C(t) → Crise(t)" in p.text:
            done = replace_in_para(p, "C(t) → Crise(t); as t ≫ t₀, w → 0 and "
                                      "C(t) → Cfall(t)",
                                   "C(t) → C_rise(t); as t ≫ t₀, w → 0 and "
                                   "C(t) → C_fall(t)")
            break
    applied("2.2 prose C_rise/C_fall", done)
    for p in paras:
        if "for estimating the per-interval flow rate Q" in p.text:
            done2 = replace_in_para(p, "per-interval flow rate Q", "per-stage flow rate Q")
            break
    applied("4.3 per-interval -> per-stage", done2)

    # ---- 5. 2.4: literal 'Q is unconstrained' ----
    old = ("The model is given no access to this value — Q enters the objective function "
           "without penalty, prior, or constraint toward the pump setting, bounded only by the "
           "[10, 5000] mL/min search envelope (Section 2.3).")
    new = ("The model is given no access to this value: Q is unconstrained in the objective "
           "function — no penalty, prior, or constraint directs it toward the pump setting, "
           "and it is bounded only by the wide [10, 5000] mL/min search envelope (Section 2.3).")
    done = False
    for p in paras:
        if "The model is given no access to this value" in p.text:
            done = replace_in_para(p, old, new)
            applied("2.4 Q is unconstrained", done)
            break

    # ---- 6. 4.1: order of magnitude -> several-fold ----
    done = False
    for p in paras:
        if "an order of magnitude below every single-process alternative" in p.text:
            done = replace_in_para(p, "an order of magnitude below every single-process "
                                      "alternative", "several-fold below every single-process "
                                      "alternative")
            applied("4.1 several-fold", done)
            break

    # ---- 7. 4.2: explicit cross-refs to 2.4 and 3.4 ----
    old = ("Q itself was unconstrained in the objective function — no penalty, prior, or "
           "constraint directed it toward the pump setting of 0.50 mL/min; it was bounded only "
           "by the wide search envelope [10, 5000] mL/min.")
    new = ("Q itself was unconstrained in the objective function (Section 2.4) — no penalty, "
           "prior, or constraint directed it toward the pump setting of 0.50 mL/min "
           "(Section 3.4); it was bounded only by the wide search envelope [10, 5000] mL/min.")
    done = False
    for p in paras:
        if "Q itself was unconstrained in the objective function" in p.text:
            done = replace_in_para(p, old, new)
            applied("4.2 cross-refs 2.4/3.4", done)
            break

    # ---- 8. 4.3: MRT 37.1 vs 37.4 reconciliation line ----
    old = " However, the first moment conflates the shut-in slug and the sustained tail into a single number:"
    new = (" This raw first moment (37.1 min) differs slightly from the model-derived mean "
           "residence time of 37.4 min (Section 4.2), which is computed from the fitted transport "
           "parameters rather than directly from the measured curve." + old)
    done = False
    for p in paras:
        if "This raw first moment (37.1 min) differs slightly" in p.text:
            done = True  # already applied
            break
        if "MRT = ∫t·C dt / ∫C dt = 37.1 min" in p.text:
            done = replace_in_para(p, old, new)
            break
    applied("4.3 MRT reconciliation", done)

    # ---- 9. per-interval -> per-stage (Abstract, 5.3) ----
    n1 = n2 = 0
    for p in paras:
        if "enables per-interval production allocation from wellhead samples alone" in p.text:
            n1 += replace_in_para(p, "per-interval production allocation",
                                  "per-stage production allocation")
        if "confirming that per-interval oil production rates can be recovered" in p.text:
            n2 += replace_in_para(p, "per-interval oil production rates",
                                  "per-stage oil production rates")
    applied("per-interval -> per-stage (Abstract)", n1 == 1)
    applied("per-interval -> per-stage (5.3)", n2 == 1)

    # ---- 10. 3.1: cite Fig. S3 (optical micrographs) ----
    old = "Pure epoxy microspheres were prepared identically without nano-Fe₃O₄@SA as a reference."
    new = old + " Optical micrographs of both materials are shown in Fig. S3."
    done = False
    for p in paras:
        if "Optical micrographs of both materials are shown in Fig. S3" in p.text:
            done = True  # already applied
            break
        if old in p.text:
            done = replace_in_para(p, old, new)
            break
    applied("3.1 Fig. S3 citation", done)

    # ---- 11. Figure Captions section: slot line removed, 10 captions written ----
    for p in paras:
        if p.text.strip() == "[10 figure caption slots — populated as figures are placed]":
            remove_paragraph(p)
            applied("remove caption slot line", True)
            break
    n_cap = 0
    for p in list(paras):
        if p.text.strip() in [f"[Fig. {n}] Caption placeholder — populated as figure is "
                              "placed." for n in range(1, 11)]:
            num = int(p.text.strip()[5:p.text.index("]")])
            _rebuild(p, [("Fig. %d.  " % num, True), (CAPTIONS[num], None)])
            n_cap += 1
    applied("10 figure captions", n_cap == 10)

    # ---- 12. Tables section: slot line removed, entries renumbered ----
    for p in paras:
        if p.text.strip() == "[Table slots — populated as content is written]":
            remove_paragraph(p)
            applied("remove table slot line", True)
            break
    done1 = done2 = done3 = False
    for p in list(paras):
        if p.text.strip() == "[Table 1] Placeholder — material prerequisites summary (see Section 3.2).":
            set_para_text(p, "[Table 1] Parameter set of the coupled release–transport model (Eq. 1) — Section 2.3.")
            done1 = True
        elif p.text.strip() == "[Table 2] Placeholder — model comparison summary (see Section 4.1).":
            set_para_text(p, "[Table 2] Material prerequisites of the coupled release–transport model and the evidence supporting each — Section 3.2.")
            np = insert_paragraph_after(p, "[Table 3] Model-selection statistics for the five candidate models fitted to the 21-point single-phase BTC — Section 4.1.")
            done2 = done3 = True
    applied("Tables renumbered (1/2/3)", done1 and done2 and done3)

    d.save(PATH)
    n_ok = sum(1 for _, ok in stats if ok)
    print(f"\n{len(stats)} checks, {n_ok} applied/already-present, {len(stats) - n_ok} skipped.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
