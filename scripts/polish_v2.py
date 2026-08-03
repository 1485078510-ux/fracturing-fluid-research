# -*- coding: utf-8 -*-
"""
V2: Safe polish — only modifies pure-text paragraphs, preserves images.
Identifies paragraphs with inline shapes and skips them.
"""
from docx import Document
from docx.oxml.ns import qn
import copy

SRC = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\Fabrication and Performance of ESP.docx'
DST = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP_polished.docx'

doc = Document(SRC)

# Build set of indices that contain inline images (DO NOT TOUCH)
IMAGE_PARAS = set()
for i, p in enumerate(doc.paragraphs):
    # Check for drawings (images) in the paragraph XML
    drawings = p._element.findall('.//' + qn('w:drawing'))
    if drawings:
        IMAGE_PARAS.add(i)
    # Also check for inline shapes via the runs
    for run in p.runs:
        if run._element.findall('.//' + qn('w:drawing')):
            IMAGE_PARAS.add(i)
            break
        if run._element.findall('.//' + qn('wp:inline')):
            IMAGE_PARAS.add(i)
            break
        if run._element.findall('.//' + qn('a:blip')):
            IMAGE_PARAS.add(i)
            break

print(f"Found {len(IMAGE_PARAS)} paragraphs with images: {sorted(IMAGE_PARAS)}")

def safe_set_text(para, text, idx):
    """Only modify paragraphs that don't contain images."""
    if idx in IMAGE_PARAS:
        print(f"  SKIP para [{idx}] — contains image")
        return False
    # Clear all runs, set text in first run
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)
    return True

# ============================================================
# ABSTRACT
# ============================================================
safe_set_text(doc.paragraphs[2],
    "Abstract: Accurate monitoring of oil production contribution from individual fracturing segments "
    "is critical for optimizing unconventional reservoir development. In this study, an oleophilic "
    "Fe₃O₄-doped epoxy resin tracer proppant (ESP-T) was designed and fabricated via emulsion "
    "polymerization, with stearic acid-modified nano-Fe₃O₄ (nano-Fe₃O₄@SA) as the lipophilic tracer "
    "and epoxy resin as the carrier matrix, integrating fracture propping and production monitoring "
    "functions. The microstructure, mechanical properties, thermal stability, wettability, oil–water "
    "permeability, and temperature-dependent tracer release kinetics of ESP-T were systematically "
    "investigated, and its oil production monitoring performance was validated through single-phase "
    "and two-phase core displacement experiments.", 2)

safe_set_text(doc.paragraphs[4],
    "Results show that nano-Fe₃O₄@SA is uniformly dispersed as nanoclusters in the epoxy matrix. "
    "ESP-T exhibits sphericity and roundness exceeding 0.9, bulk density of 0.646 g·cm⁻³, acid "
    "solubility of 3.3% (≤ 5% standard), and an initial thermal decomposition temperature of "
    "357.27 °C, far above typical downhole temperatures (80–200 °C). The water contact angle "
    "increases from 72.3° to 104.6° after nano-Fe₃O₄@SA doping, yielding a 'water-resistant "
    "and oil-permeable' characteristic: the oil filtration time (5 min 11 s) is 66.1% shorter "
    "than that of pure epoxy microspheres. Tracer release at 30–120 °C follows the Korsmeyer–Peppas "
    "model with R² > 0.90 and diffusion exponents of 0.45–0.85, indicating a synergistic "
    "Fickian diffusion and Case-II relaxation mechanism. The cumulative release at 120 °C over "
    "14 days exceeds 2.0 mg·L⁻¹, meeting ICP-MS detection requirements.", 4)

safe_set_text(doc.paragraphs[6],
    "A piecewise advection-dispersion model with smooth tanh transition was developed to interpret "
    "the single-phase tracer breakthrough curve, decomposing the signal into a Gaussian pulse "
    "component and an erfc tailing component. The model achieves R² = 0.9939, with the fitted "
    "flow rate (0.46 mL·min⁻¹) closely matching the pump-set rate (0.5 mL·min⁻¹, error 8%). "
    "Under steady-state two-phase flow, tracer flux effectively quantifies oil-phase flow rates "
    "across different oil–water ratios. The developed ESP-T demonstrates considerable potential "
    "for long-term production monitoring in unconventional oil and gas reservoirs.", 6)

# Clear para [8] (old abstract conclusion) — merged into [6], keep index stable
safe_set_text(doc.paragraphs[8], "", 8)

# ============================================================
# INTRODUCTION — fix informal expressions
# ============================================================
safe_set_text(doc.paragraphs[13],
    "Against the backdrop of sustained growth in global energy demand, unconventional oil and gas "
    "resources have emerged as the central focus of contemporary petroleum development owing to "
    "their immense reserve potential [1–5]. To date, unconventional resources account for more than "
    "53% of the world's total hydrocarbon reserves [6,7], and their efficient development constitutes "
    "a critical pillar of global energy security. As the core technology for stimulating unconventional "
    "reservoirs, hydraulic fracturing establishes artificial fractures to optimize hydrocarbon seepage "
    "pathways, thereby substantially boosting the productivity of low-permeability and tight formations "
    "[8]. Nevertheless, the oil production contribution of each fractured interval varies markedly "
    "post-fracturing, and interval productivity decays at disparate rates as field development proceeds "
    "[9]. Accurate monitoring of the production contribution rate of individual fractured intervals "
    "not only supplies key data for reservoir productivity assessment but also guides the optimization "
    "of subsequent development strategies, ultimately elevating the overall hydrocarbon recovery factor.", 13)

safe_set_text(doc.paragraphs[17],
    "Despite the inherent drawbacks of conventional oil-soluble tracers—including poor compatibility "
    "with fracturing fluid injection and limited long-term monitoring capability—tracer proppants, "
    "as composite materials integrating fracture propping and tracer monitoring functions, have "
    "effectively addressed these limitations. By immobilizing tracers within proppant carriers, "
    "tracer proppants enable long-term dynamic monitoring of fractured intervals upon co-injection "
    "with fracturing fluids, eliminating the cumbersome procedure of separate tracer injection. "
    "Zhao et al. (2020) [15] coated rhodamine 6G onto polyvinyl alcohol-modified ceramic particles "
    "via solvent evaporation to retard tracer release in aqueous media. Zhou et al. (2022) [16] "
    "encapsulated carbon quantum dots (CQDs) in polyvinyl alcohol-coated ceramic particles to achieve "
    "controllable release. Malyavko et al. (2023) [17] employed monodisperse microspheres doped with "
    "fluorescent semiconductor nanocrystals as tracers, coating them with gas-philic, oleophilic, and "
    "hydrophilic polymers to acquire phase distribution data for each fractured interval. Li et al. "
    "(2023) [18] utilized rare-earth elements as tracers and encapsulated them with ammonium "
    "polymethacrylate to elucidate the slow-release mechanism. Gong et al. (2024) [19] fabricated "
    "slow-release tracer proppants by embedding tracers into polystyrene (PS) microspheres via "
    "suspension polymerization and investigated their controlled-release behavior.", 17)

safe_set_text(doc.paragraphs[19],
    "However, conventional coated tracer proppants suffer from high density, which impedes transport "
    "in fracturing fluids, and lose monitoring functionality once the polymer coating is fully "
    "dissolved. Pure PS microspheres exhibit inferior mechanical strength and thermal stability "
    "[20–23]. Epoxy resin—a high-performance polymer with exceptional mechanical strength, thermal "
    "stability, and chemical resistance—has therefore been integrated into proppant design, typically "
    "as a coating or composite matrix. Coating proppants with epoxy resin significantly reduces the "
    "crush rate, improves compressive strength and acid resistance, while retaining the low-density "
    "advantage [24,25]. Despite notable advances in coating technology, it remains constrained by "
    "challenges including uneven coating, interfacial debonding risks, and complex preparation "
    "procedures. As a further optimized alternative, direct synthesis of epoxy resin microspheres "
    "as proppants enables low-density tailoring via synthetic regulation and facilitates nanoparticle "
    "modification and tracer encapsulation. Li et al. (2021) [26] experimentally explored the release "
    "mechanism of water-soluble tracers from epoxy resin in aqueous environments. Wei et al. (2024) "
    "[27] encapsulated water-soluble tracers in an epoxy resin matrix to locate water-producing "
    "intervals in oil wells. Nevertheless, research on epoxy resin as an oleophilic release "
    "matrix remains inadequate.", 19)

# ============================================================
# RESULTS — minor text fixes (skip all caption paragraphs)
# ============================================================

# 3.1 SEM body text
safe_set_text(doc.paragraphs[69],
    "Figure 3-1 presents the SEM micrographs of pure epoxy resin microspheres and ESP-T proppants. "
    "At low magnification (a, d), both samples exhibit excellent sphericity and high monodispersity, "
    "with uniform particle sizes and no inter-particle agglomeration, demonstrating the excellent "
    "controllability of the suspension polymerization process and the uniformity of the curing "
    "reaction. Notably, well-formed spherical epoxy resin microspheres can be stably fabricated "
    "with or without the incorporation of nano-Fe₃O₄@SA. A direct comparison of (a) and (d) reveals "
    "that pure epoxy microspheres (a) feature a relatively smooth surface with minor particulate "
    "debris, whereas ESP-T microspheres (d) exhibit a uniformly rough surface topography. This "
    "observation provides preliminary evidence for the successful incorporation of nano-Fe₃O₄@SA "
    "and its significant modulation of the microsphere surface morphology.", 69)

# 3.5 Sphericity — fix abrupt start
safe_set_text(doc.paragraphs[107],
    "The formation of regular spherical morphology is attributed to the optimized stirring "
    "conditions. A stirring rate of 380 RPM imparts uniform shear force to epoxy resin droplets "
    "within the nano-SiO₂ aqueous dispersion, favoring the formation of regularly spherical "
    "droplets. The subsequently introduced guar gum (0.9 g for neat epoxy microspheres and "
    "1.1 g for ESP-T) elevates the system viscosity, thereby suppressing droplet deformation "
    "and sedimentation prior to curing. Hollow glass microspheres act as internal reinforcing "
    "supports, further preserving the structural integrity and ideal sphericity of the microspheres.", 107)

# 3.7 Acid Solubility
safe_set_text(doc.paragraphs[117],
    "The acid solubility of ESP-T was determined to be 3.3%, compared with 2.5% for neat epoxy "
    "resin proppants—a negligible difference. Although nano-Fe₃O₄@SA is soluble in dilute HCl "
    "and HF, the stearic acid-derived hydrophobic film shields the particle surface from acid "
    "attack, mitigating dissolution. Meanwhile, uniformly dispersed nano-Fe₃O₄@SA nanoclusters "
    "act as an inorganic barrier within the epoxy matrix, slowing acid penetration and reducing "
    "resin degradation. Lower acid solubility indicates superior stability in acidic environments, "
    "ensuring long-term fracture propping integrity. Both proppants meet the industrial standard "
    "(acid solubility ≤ 5%), making them suitable for acid fracturing operations. The acid "
    "solubility standard deviation is below 0.2% for both materials, demonstrating stable "
    "dissolution behavior and excellent preparation reproducibility.", 117)

# ============================================================
# SECTION 3.10 — REMOVE editing note
# ============================================================
# Para 140: contains an image in the last run. Replace only text runs, preserve image.
p140 = doc.paragraphs[140]
new_text_140 = (
    "Based on the release results, the release mechanism of ESP-T is inferred as follows. "
    "During the release process, nonpolar solvents (e.g., alkanes) permeate into the polymer "
    "network, inducing swelling of the cross-linked epoxy matrix. At this stage, the polymer "
    "microspheres form a dual-state structure: an inner glassy core and an outer gel layer. "
    "The swelling reduces the entanglement of polymer molecular chains, thereby creating "
    "expanded transport channels that enable nano-Fe₃O₄@SA tracers to diffuse into the "
    "external environment. Elevated temperature enhances the permeability of solvent molecules "
    "into the polymer matrix, accelerating the swelling rate. Furthermore, increased temperature "
    "weakens the intermolecular forces within the polymer network, expanding the pore size "
    "and further promoting tracer diffusion through the matrix."
)
# Clear text runs but keep image runs
for run in p140.runs:
    has_img = (len(run._element.findall('.//' + qn('w:drawing'))) > 0 or
               len(run._element.findall('.//' + qn('a:blip'))) > 0)
    if has_img:
        continue  # preserve image run
    run.text = ''
# Put new text in first non-image run
for run in p140.runs:
    has_img = (len(run._element.findall('.//' + qn('w:drawing'))) > 0 or
               len(run._element.findall('.//' + qn('a:blip'))) > 0)
    if not has_img:
        run.text = new_text_140
        break
print("  Fixed para [140] — editing note removed, image preserved")

# K-P equation variables — fix formatting (these are safe, no images)
safe_set_text(doc.paragraphs[133], "C = tracer concentration at time t (mg/L);", 133)
safe_set_text(doc.paragraphs[134], "C₀ = maximum (equilibrium) release concentration (mg/L);", 134)
safe_set_text(doc.paragraphs[135], "K = kinetic rate constant;", 135)
safe_set_text(doc.paragraphs[136], "t = release time (h);", 136)
safe_set_text(doc.paragraphs[137], "n = diffusion exponent (dimensionless).", 137)

# ============================================================
# SECTION 3.11 — ADE derivation + model + results
# ============================================================
safe_set_text(doc.paragraphs[146],
    "The interpretation model is derived from the one-dimensional advection-dispersion "
    "equation (ADE): ∂C/∂t + v·∂C/∂x = D·∂²C/∂x², where v = 4Q/(πd²) is the mean "
    "flow velocity and D = αv is the dispersion coefficient. For the tracer pulse released "
    "from ESP-T into the oil phase during the shut-in period, the analytical solution of "
    "the ADE for an instantaneous slug injection yields the Gaussian-form rising component "
    "in Eq. (6). For the sustained slow release of residual tracer from the epoxy matrix "
    "after the main pulse, solving the ADE with a continuous-source boundary condition "
    "yields the erfc tailing component. The two regimes are blended via a hyperbolic "
    "tangent weighting function, producing a smooth, physically continuous breakthrough curve.", 146)

safe_set_text(doc.paragraphs[147],
    "Accordingly, the tracer breakthrough curve was interpreted using the following "
    "piecewise advection-dispersion model with smooth transition:", 147)

safe_set_text(doc.paragraphs[149], "where:", 149)

new_vars = [
    "c_b = baseline concentration after the tracer signal has stabilized (dimensionless);",
    "A = peak amplitude coefficient, proportional to the tracer mass accumulated in the oil phase during shut-in;",
    "a = tailing amplitude coefficient, proportional to the sustained tracer release rate from the ESP-T matrix;",
    "α = longitudinal dispersivity, characterizing axial spreading of the tracer slug along the flow line, cm;",
    "Q = effective volumetric flow rate through the flow line, cm³/min;",
    "t₀ = time at which sustained matrix-diffusion release overtakes the initial pulse as the dominant signal source, min;",
    "σ = time span of the transition from pulse-dominated to tailing-dominated regime, min.",
]
for para_idx, text in zip(range(150, 157), new_vars):
    safe_set_text(doc.paragraphs[para_idx], text, para_idx)

safe_set_text(doc.paragraphs[157],
    "In the above expressions, x = 100 cm and d = 5 cm are the fixed length and inner "
    "diameter of the flow line, respectively. In Eq. (6), the rising component C_rise "
    "captures the advective-dispersive arrival of the tracer slug at the sampling point, "
    "while the tailing component C_fall describes the slow, sustained release of residual "
    "nano-Fe₃O₄@SA from the epoxy matrix into the flowing oil. The parameter t₀ marks "
    "the center of the transition between the two regimes, and σ controls the width of "
    "the transition zone. This two-component structure reflects the physical reality that "
    "the tracer signal originates from two superimposed processes: a short-duration "
    "concentration pulse generated by tracer accumulation during shut-in, and a prolonged "
    "tail governed by the matrix-diffusion-controlled release kinetics established in "
    "Section 3.10.", 157)

safe_set_text(doc.paragraphs[159],
    "The fitting results are presented in Figure 3-9(b), and the fitted parameters with "
    "derived transport properties are summarized in Table 3-3. The model achieves "
    "R² = 0.9939 and RMSE = 0.0210, with residuals distributed randomly within "
    "± 2 × RMSE across the entire time range. Key derived parameters include: mean flow "
    "velocity v = 4Q/(πd²) = 2.59 cm/min; longitudinal dispersivity α = 107.1 cm; "
    "and Péclet number Pe = x/α = 0.934. The Pe < 1 indicates dispersion-dominated "
    "transport, consistent with the gradual, sustained release of nano-Fe₃O₄@SA from "
    "the ESP-T matrix. A comparison between the fitted flow rate (0.46 mL/min) and the "
    "experimental pump-set flow rate (0.5 mL/min) yields a relative error of 8% "
    "(Figure 3-9c), verifying the accuracy of the model for flow rate determination "
    "from tracer breakthrough data. The mean residence time MRT = 37.4 min agrees "
    "closely with the convective travel time x/v = 38.6 min (ratio 0.967). "
    "Deconvolution of the fitted curve shows that 47% of the total tracer signal "
    "originates from the erfc tailing component, confirming the significant contribution "
    "of matrix-diffusion-controlled release and its consistency with the non-Fickian "
    "mechanism identified in Section 3.10.", 159)

# ============================================================
# CONCLUSIONS (actual indices: 171, 173, 175, 177)
# ============================================================
safe_set_text(doc.paragraphs[171],
    "Oleophilic nano-Fe₃O₄@SA nanoparticles were synthesized via coprecipitation and "
    "modified with stearic acid, and ESP-T was subsequently fabricated via emulsion "
    "polymerization. SEM, elemental mapping, and optical microscopy confirmed that "
    "nano-Fe₃O₄@SA was uniformly dispersed in the epoxy matrix as nanoclusters without "
    "significant agglomeration, and ESP-T exhibited sphericity and roundness exceeding "
    "0.9, satisfying industrial molding standards.", 171)

safe_set_text(doc.paragraphs[173],
    "ESP-T demonstrated superior physical, mechanical, and thermal stability: a bulk "
    "density of 0.646 g/cm³ ensured favorable suspension in water-based fracturing "
    "fluids; an apparent density of 1.072 g/cm³ indicated a compact internal structure; "
    "the breakage rate at 50 MPa was 2.9%; and the acid solubility of 3.3% met the "
    "industrial requirement (≤ 5%), confirming applicability to acidic reservoirs. "
    "The initial thermal decomposition temperature of 357.27 °C far exceeded typical "
    "downhole temperatures (80–200 °C), eliminating concerns of resin degradation "
    "during long-term service. The water contact angle increased from 72.3° to 104.6°, "
    "and the oil filtration time (5 min 11 s) was 66.1% shorter than that of pure "
    "epoxy microspheres, validating the 'water-resistant and oil-permeable' transport "
    "characteristic.", 173)

safe_set_text(doc.paragraphs[175],
    "In terms of tracer performance, Fe release from ESP-T at 30–120 °C followed the "
    "Korsmeyer–Peppas (K-P) kinetic model. The fitted diffusion exponent n ranged from "
    "0.45 to 0.85 at all tested temperatures, indicating a synergistic release mechanism "
    "governed by both Fickian diffusion and Case-II relaxation, with R² > 0.90 for all "
    "fittings. The cumulative tracer release at 120 °C over 14 days exceeded 2.0 mg/L, "
    "meeting ICP-MS detection requirements for long-term monitoring. Under shut-in "
    "single-phase flow, a piecewise advection-dispersion model with smooth tanh transition "
    "was developed to interpret the tracer breakthrough curve. The model decomposes the "
    "signal into a Gaussian pulse component (tracer slug from shut-in accumulation) and an "
    "erfc tailing component (sustained matrix-diffusion-controlled release from ESP-T). "
    "Fitting yielded R² = 0.9939, with a derived flow velocity of 2.59 cm/min, dispersivity "
    "of 107.1 cm, and Péclet number Pe = 0.934. The fitted flow rate (0.46 mL/min) agreed "
    "closely with the pump-set rate (0.5 mL/min, relative error 8%), and the mean residence "
    "time (37.4 min) matched the theoretical convective travel time (38.6 min, ratio 0.967). "
    "Deconvolution showed that 47% of the tracer signal originated from the tailing component, "
    "confirming the significant contribution of matrix-diffusion-controlled release. "
    "Under steady-state two-phase flow, tracer concentration decreased monotonically with "
    "increasing total flow rate and was independent of the oil–water ratio. The tracer flux "
    "method effectively quantified oil-phase flow rates in individual intervals, providing a "
    "reliable basis for dynamic production monitoring during the steady-state stage.", 175)

safe_set_text(doc.paragraphs[177],
    "In summary, ESP-T integrates fracture propping and tracer monitoring functions, "
    "making it suitable for acid fracturing, deep wells, and high-pressure formations. "
    "It provides reliable technical support for fracturing parameter optimization and "
    "enhanced oil recovery, demonstrating significant industrial application potential "
    "in unconventional oil and gas development.", 177)

# ============================================================
# SAVE
# ============================================================
doc.save(DST)
print(f"\nSaved: {DST}")
print("All image paragraphs preserved. No figure captions modified.")