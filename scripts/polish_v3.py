# -*- coding: utf-8 -*-
"""
Polishing script that reads the original DOCX and creates a revised version
with yellow highlighting on changed text portions only.
Reads replacement pairs from a JSON file.
"""
import json
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"c:\Users\郝\Desktop\claude\向晶\页岩气藏地应力演化与裂缝扩展耦合机制研究进展与展望-V2.docx"
DST = r"c:\Users\郝\Desktop\claude\向晶\页岩气藏地应力演化与裂缝扩展耦合机制研究进展与展望-润色版.docx"
JSON_PATH = r"c:\Users\郝\Desktop\claude\向晶\replacements.json"

def add_highlight(run):
    """Add yellow highlight to a run."""
    rPr = run._element.get_or_add_rPr()
    h = OxmlElement('w:highlight')
    h.set(qn('w:val'), 'yellow')
    rPr.append(h)

def get_font_info(para):
    """Extract font info from first run of paragraph."""
    if para.runs:
        r = para.runs[0]
        return {
            'size': r.font.size,
            'bold': r.font.bold,
            'name': r.font.name
        }
    return {'size': Pt(12), 'bold': None, 'name': '宋体'}

def apply_font(run, fi):
    """Apply font info to a run."""
    if fi['size']:
        run.font.size = fi['size']
    if fi['bold'] is not None:
        run.font.bold = fi['bold']
    if fi['name']:
        run.font.name = fi['name']

def process_paragraph(para, changes):
    """
    Apply multiple text replacements to a paragraph with highlighting.
    changes: list of (old_text, new_text) tuples
    Only the new_text portions get yellow highlight.
    """
    if not changes:
        return 0

    full_text = para.text
    fi = get_font_info(para)

    # Find all match positions in the ORIGINAL text
    matches = []
    for old, new in changes:
        idx = full_text.find(old)
        if idx >= 0:
            matches.append((idx, idx + len(old), new))

    if not matches:
        return 0

    # Sort by start position
    matches.sort(key=lambda x: x[0])

    # Handle overlapping matches (keep the later one which has priority)
    resolved = []
    for m in matches:
        if resolved and m[0] < resolved[-1][1]:
            # Overlap: keep the one with larger end position (more comprehensive)
            prev = resolved.pop()
            if m[1] > prev[1]:
                resolved.append(m)
            else:
                resolved.append(prev)
        else:
            resolved.append(m)

    # Build text segments: (text, highlight_flag)
    segments = []
    cursor = 0
    for start, end, new_txt in resolved:
        if cursor < start:
            segments.append((full_text[cursor:start], False))
        segments.append((new_txt, True))
        cursor = end
    if cursor < len(full_text):
        segments.append((full_text[cursor:], False))

    # Clear existing runs
    for run in list(para.runs):
        run._element.getparent().remove(run._element)

    # Create new runs
    for seg_text, is_highlighted in segments:
        if not seg_text:
            continue
        r = para.add_run(seg_text)
        apply_font(r, fi)
        if is_highlighted:
            add_highlight(r)

    return len(resolved)

# Main processing
print("Loading document...")
doc = Document(SRC)
print(f"Document has {len(doc.paragraphs)} paragraphs")

print("Loading replacements...")
with open(JSON_PATH, 'r', encoding='utf-8') as f:
    all_entries = json.load(f)
print(f"Loaded {len(all_entries)} replacement groups")

total_paras = 0
total_changes = 0

for para in doc.paragraphs:
    if not para.text.strip():
        continue

    # Collect all applicable replacements
    applicable = []
    for entry in all_entries:
        if entry['keyword'] in para.text:
            for rep in entry['replacements']:
                if rep['old'] in para.text:
                    applicable.append((rep['old'], rep['new']))

    if not applicable:
        continue

    # Remove duplicates (by old text)
    seen = set()
    unique = []
    for old, new in applicable:
        if old not in seen:
            unique.append((old, new))
            seen.add(old)

    # Sort by length (longest first) to avoid partial match issues
    unique.sort(key=lambda x: len(x[0]), reverse=True)

    n = process_paragraph(para, unique)
    if n > 0:
        total_paras += 1
        total_changes += n

doc.save(DST)
print(f"Done! Modified {total_paras} paragraphs with {total_changes} changes.")
print(f"Saved to: {DST}")
