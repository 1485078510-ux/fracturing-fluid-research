#!/usr/bin/env python3
"""Renumber all references by first-appearance order + add missing citations."""
from docx import Document
import re

doc = Document("四氧化三铁环氧树脂拟合/ESP-T_终稿.docx")

# Step 1: Add missing citations in body text
# [26] Ritger & Peppas 1987 -> Section 3.6, K-P model introduction
# [27] Peppas & Sahlin 1989 -> same section
# [28] van Genuchten & Alves 1982 -> Section 3.7, ADE derivation

for i, p in enumerate(doc.paragraphs):
    t = p.text

    # Add [26] after "developed by Korsmeyer and Peppas"
    if "developed by Korsmeyer and Peppas, widely employed" in t and "[26]" not in t:
        t = t.replace(
            "developed by Korsmeyer and Peppas, widely employed",
            "developed by Korsmeyer and Peppas [26], widely employed"
        )
        for r in p.runs: r.text = ""
        p.runs[0].text = t
        print(f"Added [26] at para [{i}]")
        break

for i, p in enumerate(doc.paragraphs):
    t = p.text
    # Add [27] after the K-P threshold explanation
    if "Case-II-relaxation-controlled release" in t and "[27]" not in t:
        t = t.replace(
            "Case-II-relaxation-controlled release.",
            "Case-II-relaxation-controlled release [27]."
        )
        for r in p.runs: r.text = ""
        p.runs[0].text = t
        print(f"Added [27] at para [{i}]")
        break

for i, p in enumerate(doc.paragraphs):
    t = p.text
    # Add [28] after "classical ADE analytical solution"
    if "the classical ADE analytical solution yields" in t and "[28]" not in t:
        t = t.replace(
            "the classical ADE analytical solution yields",
            "the classical ADE analytical solution [28] yields"
        )
        for r in p.runs: r.text = ""
        p.runs[0].text = t
        print(f"Added [28] at para [{i}]")
        break

# Step 2: Remap old->new based on first-appearance order
OLD2NEW = {
    2:1, 23:2, 24:3, 25:4, 1:5, 3:6, 4:7, 5:8, 6:9, 7:10,
    8:11, 9:12, 10:13, 11:14, 12:15, 13:16, 14:17, 15:18,
    16:19, 17:20, 18:21, 19:22, 20:23, 21:24, 22:25,
    29:26, 26:27, 27:28, 28:29
}

def remap_citation(match):
    """Remap a citation bracket like [2,23-25] to new numbering."""
    content = match.group(1)
    # Skip if this doesn't look like citation numbers (e.g., has letters)
    if re.search(r'[a-zA-Z]', content):
        return f"[{content}]"
    parts = []
    for part in content.split(','):
        part = part.strip()
        if '-' in part:
            pieces = part.split('-')
            if len(pieces) == 2:
                try:
                    new_a = OLD2NEW.get(int(pieces[0]), int(pieces[0]))
                    new_b = OLD2NEW.get(int(pieces[1]), int(pieces[1]))
                    if new_a == new_b:
                        parts.append(str(new_a))
                    elif new_b == new_a + 1:
                        parts.append(f"{new_a},{new_b}")
                    else:
                        parts.append(f"{new_a}-{new_b}")
                except:
                    parts.append(part)
            else:
                parts.append(part)
        else:
            try:
                parts.append(str(OLD2NEW.get(int(part), int(part))))
            except:
                parts.append(part)
    # Collapse: if we have simple list, try to form ranges
    nums = []
    for p in parts:
        try:
            nums.append(int(p))
        except:
            nums.append(None)
    if all(n is not None for n in nums) and len(nums) > 2:
        nums.sort()
        # Simple collapse: consecutive numbers
        result_parts = []
        i = 0
        while i < len(nums):
            start = nums[i]
            while i+1 < len(nums) and nums[i+1] == nums[i]+1:
                i += 1
            end = nums[i]
            if start == end:
                result_parts.append(str(start))
            elif end == start + 1:
                result_parts.append(f"{start},{end}")
            else:
                result_parts.append(f"{start}-{end}")
            i += 1
        return f"[{','.join(result_parts)}]"
    return f"[{','.join(parts)}]"

# Find reference section
ref_start = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "References":
        ref_start = i
        break

# Remap in-text citations (before reference section)
count = 0
for i in range(ref_start):
    p = doc.paragraphs[i]
    old_text = p.text
    new_text = re.sub(r'\[([^\]]+)\]', remap_citation, old_text)
    if new_text != old_text:
        for r in p.runs:
            r.text = ""
        p.runs[0].text = new_text
        count += 1
print(f"Remapped citations in {count} paragraphs")

# Step 3: Reorder reference list entries
# Current refs at positions ref_start+1 to ref_start+29
# Need to reorder so that new position N contains the ref that OLD2NEW maps to N

# Build inverse map: new_num -> old_num
NEW2OLD = {v: k for k, v in OLD2NEW.items()}

# Read current refs
old_refs = {}
for j in range(29):
    idx = ref_start + 1 + j
    if idx < len(doc.paragraphs):
        old_refs[j + 1] = doc.paragraphs[idx].text

# Build new ordered list
new_refs = []
for n in range(1, 30):
    old_num = NEW2OLD.get(n, n)
    new_refs.append(old_refs.get(old_num, f"[{n}] MISSING"))

# Write back
for j in range(29):
    idx = ref_start + 1 + j
    if idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]
        for r in p.runs:
            r.text = ""
        p.runs[0].text = new_refs[j]

print(f"Reference list reordered by first-appearance")
print("First 5 refs:", [new_refs[j][:60] for j in range(5)])

doc.save("四氧化三铁环氧树脂拟合/ESP-T_终稿.docx")
print("Saved.")