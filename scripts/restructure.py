# -*- coding: utf-8 -*-
"""Restructure ESP_formatted.docx to match Chinese draft structure."""
from docx import Document
from docx.oxml.ns import qn

DST = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP_formatted.docx'
doc = Document(DST)

IMG = set()
for i, p in enumerate(doc.paragraphs):
    for run in p.runs:
        if (run._element.findall('.//'+qn('w:drawing')) or
            run._element.findall('.//'+qn('a:blip'))):
            IMG.add(i); break

def sp(idx, text):
    if idx in IMG: return
    p = doc.paragraphs[idx]
    for r in p.runs: r.text = ''
    if p.runs: p.runs[0].text = text
    else: p.add_run(text)

def demote(idx):
    """Change Heading style to Normal."""
    p = doc.paragraphs[idx]
    # Change style to Normal
    p.style = doc.styles['Normal']

def clear(idx):
    if idx in IMG: return
    p = doc.paragraphs[idx]
    for r in p.runs: r.text = ''

# ===== PART 1: Reorganize Results sections =====
# Map: old section -> new section
# 3.1 SEM [66] -> 3.1 SEM [keep]
# 3.2 Optical Microscopy [82] -> MERGE into 3.1 as body text, remove as section
# 3.3 Thermal Stability [90] -> 3.2 Thermal Stability
# 3.4 WCA [97] -> 3.3 Water Contact Angle
# 3.5 Sphericity [104] -> INTEGRATE into 3.1, remove as section
# 3.6 Density [108] -> 3.4 Bulk Density and Apparent Density
# 3.7 Acid Solubility [114] -> INTEGRATE, remove as section
# 3.8 Breakage [116] -> INTEGRATE, remove as section
# 3.9 Conductivity [119] -> 3.5 Proppant Pack Conductivity
# 3.10 Release [127] -> 3.6 Tracer Release Behavior
# 3.11 Production [142] -> 3.7 Oil Production Monitoring

# Step 1: Demote sections to be removed (change heading to Normal)
# 3.2 Optical Microscopy [82] -> demote heading, keep content as body
demote(82)
# Keep the content paragraphs [83],[84],[85],[87],[88] as body text under 3.1

# 3.5 Sphericity [104] -> demote
demote(104)
# Keep [105] as body text (merge into SEM or Density section)

# 3.7 Acid Solubility [114] -> demote heading
demote(114)
# Keep content [115] as body text

# 3.8 Breakage [116] -> demote heading
demote(116)
# Keep content [117] as body text

# Step 2: Renumber remaining section headings
# 3.3 Thermal [90] -> 3.2
sp(90, "3.2 Thermal Stability")
# 3.4 WCA [97] -> 3.3
sp(97, "3.3 Water Contact Angle (WCA)")
# 3.6 Density [108] -> 3.4
sp(108, "3.4 Bulk Density and Apparent Density")
# 3.9 Conductivity [119] -> 3.5
sp(119, "3.5 Proppant Pack Conductivity")
# 3.10 Release [127] -> 3.6
sp(127, "3.6 Tracer Release Behavior at Different Temperatures")
# 3.11 Production [142] -> 3.7
sp(142, "3.7 Oil Production Monitoring of ESP-T")

# Step 3: Merge sphericity content into 3.1 or 3.4
# Demote the sphericity body para heading by integrating into a note
sp(105,
    "The sphericity and roundness of both proppant types were determined by "
    "comparison with the Krumbien-Sloss chart [29]. Both values exceed 0.9, "
    "satisfying industrial specifications (roundness >= 0.6, sphericity >= 0.6) "
    "and ensuring favorable fracture conductivity and pumpability. A stirring "
    "rate of 380 RPM imparts uniform shear force to epoxy resin droplets within "
    "the nano-SiO2 aqueous dispersion, favoring regularly spherical droplets. "
    "Guar gum (0.9 g for neat epoxy microspheres; 1.1 g for ESP-T) elevates "
    "system viscosity, suppressing droplet deformation and sedimentation before "
    "curing. Hollow glass microspheres act as internal reinforcing supports, "
    "further preserving structural integrity and ideal sphericity.")

# Step 4: Add demoted section titles as bold text
sp(82, "Morphological Characterization via Optical Microscopy")
sp(104, "Sphericity and Roundness")
sp(114, "Acid Solubility")
sp(116, "Breakage Rate")

# Step 5: Update figure references in captions
# Fig.3-4 -> Fig.3-2 (thermal), Fig.3-5 -> Fig.3-3 (WCA), etc.
# This is complex - let me search for figure references and update them
figure_renumber = {
    'Figure 3-4': 'Figure 3-2',  # thermal
    'Figure 3-5': 'Figure 3-3',  # WCA
    'Figure 3-6': 'Figure 3-4',  # density
    'Figure 3-7': 'Figure 3-5',  # conductivity
    'Figure 3-8': 'Figure 3-6',  # release
    'Figure 3-9': 'Figure 3-7',  # production - single phase
    'Figure 3-10': 'Figure 3-8', # two-phase
    'Figure 3-1': 'Figure 3-1',  # SEM - unchanged
    'Figure 3-2': 'Figure 3-1',  # mapping -> part of 3.1
    'Figure 3-3': 'Figure 3-1',  # optical -> part of 3.1
    'Fig. 3-4': 'Fig. 3-2',
    'Fig. 3-5': 'Fig. 3-3',
    'Fig. 3-6': 'Fig. 3-4',
    'Fig. 3-7': 'Fig. 3-5',
    'Fig. 3-8': 'Fig. 3-6',
    'Fig. 3-9': 'Fig. 3-7',
    'Fig. 3-10': 'Fig. 3-8',
}

for para in doc.paragraphs:
    for run in para.runs:
        if run.text:
            for old, new in figure_renumber.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)

# Also update section cross-references
section_renumber = {
    'Section 3.3': 'Section 3.2',
    'Section 3.4': 'Section 3.3',
    'Section 3.5': 'Section 3.1',  # sphericity merged
    'Section 3.6': 'Section 3.4',
    'Section 3.9': 'Section 3.5',
    'Section 3.10': 'Section 3.6',
    'Section 3.11': 'Section 3.7',
}

for para in doc.paragraphs:
    for run in para.runs:
        if run.text:
            for old, new in section_renumber.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)

# ===== PART 2: Update Methods section numbering =====
# Current: 2.1-2.7
# Target: 2.1-2.6 (merge characterization into 2.4)
sp(41, "2.4 Characterization Methods")  # was 2.5
sp(49, "2.5 Tracer Release Behavior at Different Temperatures")  # was 2.6
sp(54, "2.6 Oil Production Monitoring with ESP-T")  # was 2.7

# ===== SAVE =====
doc.save(DST)
print("Restructured to match Chinese draft format.")
print("Changes: merged 4 sections, renumbered figures and sections, updated cross-references.")