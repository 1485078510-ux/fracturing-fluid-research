#!/usr/bin/env python3
"""Restructure Section 3.7 for logical flow."""
from docx import Document
doc = Document('四氧化三铁环氧树脂拟合/ESP-T_终稿.docx')

# [126] Motivation first, then equations
p126 = doc.paragraphs[126]
new126 = (
    "The tracer breakthrough curve recorded at the wellhead reflects the "
    "superposition of two distinct transport processes: the advective-dispersive "
    "migration of the tracer slug accumulated in the proppant pack during the "
    "shut-in period, and the sustained, matrix-diffusion-controlled release of "
    "residual tracer from the polymer carrier after the main pulse has passed. "
    "To capture both processes within a single physically grounded framework, "
    "we construct a composite model from the one-dimensional "
    "advection-dispersion equation (ADE), "
    "dC/dt + v dC/dx = D d2C/dx2, "
    "where v = 4Q/(pi d2) is the mean flow velocity and "
    "D = alpha v is the longitudinal dispersion coefficient, with Q the effective "
    "oil-phase volumetric flow rate, d the tubing inner diameter, and alpha the "
    "longitudinal dispersivity."
    "\n\n"
    "Two classical ADE analytical solutions form the building blocks of the "
    "model. The first describes an instantaneous slug of tracer mass M injected "
    "at x = 0, t = 0 in a semi-infinite domain, yielding a Gaussian pulse [29]: "
    "C_rise(x,t) = (M/A)/sqrt(4 pi D t) exp[-(x - v t)2 / (4 D t)]. "
    "This component captures the advective-dispersive transport of the tracer "
    "accumulated during shut-in. The second describes a continuous source "
    "maintaining a constant concentration C0 at the inlet, yielding a "
    "complementary-error-function form: "
    "C_fall(x,t) = (C0/2) erfc[(x - v t) / sqrt(4 D t)], "
    "which captures the sustained release from the epoxy matrix. The two regimes "
    "operate simultaneously during production; we blend them via a "
    "hyperbolic-tangent weighting function "
    "w(t) = 1/2 [1 + tanh((t - t0)/sigma)], where t0 is the crossover center "
    "and sigma is fixed at the sampling interval (4 min). The full composite "
    "model is "
    "C(t) = cb + w(t) C_rise + [1 - w(t)] C_fall, expressed as Eq. (2), "
    "with the tanh weight smoothly shifting dominance from the Gaussian pulse "
    "at early times to the erfc tail at late times."
)
for r in p126.runs: r.text = ''
p126.runs[0].text = new126
print('[126] restructured')

# [140] Remove - redundant
p140 = doc.paragraphs[140]
for r in p140.runs: r.text = ''
print('[140] removed')

# [142] Fitting results
p142 = doc.paragraphs[142]
new142 = (
    "The model was fitted to the single-phase tracer breakthrough curve "
    "(Figure 3-8b) and the fitted parameters are summarized in Table 3-3. "
    "The model achieves R2 = 0.9939 with RMSE = 0.0210 and residuals "
    "randomly distributed within +/- 2 sigma across the full time range. "
    "The fitted effective flow rate Q = 0.46 mL/min agrees closely with the "
    "independently set pump flow rate of 0.50 mL/min (relative error 8%). "
    "The fitted mean residence time MRT = 37.4 min matches the convective "
    "travel time x/v = 38.6 min (ratio 0.967). The Peclet number "
    "Pe = x/alpha = 0.934 places the transport at the transition between "
    "advection-dominated and dispersion-dominated regimes, consistent with "
    "a matrix-diffusion-controlled source. Integration of the fitted "
    "C_rise and C_fall components separately shows that 47% of the total "
    "integrated tracer signal originates from the erfc tail. This confirms "
    "that the non-Fickian release mechanism identified in Section 3.6 remains "
    "the dominant transport mode under flow conditions, establishing the "
    "physical basis for long-term monitoring with ESP-T."
)
for r in p142.runs: r.text = ''
p142.runs[0].text = new142
print('[142] rewritten')

# [146] Compact physical interpretation
p146 = doc.paragraphs[146]
new146 = (
    "The fitted parameters carry physical significance beyond the quality "
    "of the fit. The independently recovered Q, MRT, and Pe all agree with "
    "values predicted from the known pump setting, flow-path geometry, and "
    "the non-Fickian release kinetics characterized in Section 3.6. This "
    "multi-parameter consistency indicates that the model parameters "
    "correspond to physically measurable quantities recoverable from the "
    "tracer signal alone. The 47% contribution of the erfc tail is "
    "particularly instructive for field operations: nearly half of the "
    "detectable tracer signal originates from sustained release rather than "
    "from the initial shut-in pulse, meaning that long-term monitoring does "
    "not depend on repeated shut-in cycles. The tail-plateau concentration "
    "cb, though fitted as a constant in the present single-temperature "
    "experiment, reflects the steady-state diffusion flux documented in "
    "Section 3.6. Given that the K-P rate constant K increases systematically "
    "from 0.055 (30 degC) to 0.196 (120 degC), cb is expected to scale "
    "with reservoir temperature, providing a pathway for temperature-dependent "
    "model parameterization."
)
for r in p146.runs: r.text = ''
p146.runs[0].text = new146
print('[146] rewritten')

# [148] Bridge to two-phase
p148 = doc.paragraphs[148]
new148 = (
    "The single-phase analysis establishes the model's ability to recover "
    "the oil production rate from the tracer signal. To evaluate whether "
    "this capability extends to two-phase flow, steady-state core displacement "
    "experiments were conducted at three oil-water ratios (OWR = 4:1, 1:1, "
    "1:4) and four total flow rates (0.1-0.4 mL/min). Direct observation "
    "of the effluent at OWR = 1:1 confirmed a slightly higher oil fraction "
    "in the early-stage effluent, consistent with the oil-permeable "
    "characteristic of ESP-T established in Section 3.5."
)
for r in p148.runs: r.text = ''
p148.runs[0].text = new148
print('[148] rewritten')

# [151] FO results
p151 = doc.paragraphs[151]
new151 = (
    "Under steady-state production, tracer concentrations are inherently "
    "low, and the breakthrough curve conveys information primarily through "
    "the magnitude and trend of the concentration signal. To extract "
    "quantitative rate information, we introduce the tracer flux "
    "FO, defined as the mass of tracer passing through the wellhead per "
    "unit time, which equals the tracer release rate of ESP-T [8]. "
    "Figure 3-9(b) shows that FO increases with OWR but is independent "
    "of total flow rate. This behavior is physically consistent with the "
    "oil-wetting character of ESP-T: at constant OWR the oil-proppant "
    "contact area is fixed, so FO remains constant regardless of total "
    "flow; increasing OWR expands the contact area, raising FO."
)
for r in p151.runs: r.text = ''
p151.runs[0].text = new151
print('[151] rewritten')

# [153] Validation + limitations
p153 = doc.paragraphs[153]
new153 = (
    "Figure 3-9(c) compares the normalized FO with the actual oil-phase "
    "flow rate across the three OWR values at a fixed total flow rate of "
    "0.1 mL/min. Normalization used the steady-state FO from single-phase "
    "oil displacement (3.187 micro-g/min) as the reference. The close "
    "agreement between normalized FO and actual oil-phase flow rate "
    "demonstrates that, under steady-state two-phase flow, the oil "
    "production rate of a labeled interval can be quantified from the "
    "FO variation curve when the total flow rate is known."
    "\n\n"
    "Several limitations of the current validation should be noted. The "
    "ADE model assumes a single fractured interval with uniform proppant "
    "packing; multi-interval interactions and pack heterogeneity may "
    "introduce deviations in the field. The FO calibration relies on "
    "steady-state flow, and the FO-to-rate relationship may not hold "
    "during transient regimes such as well start-up, shut-in, or rapid "
    "drawdown. Dodecane was used as the oil-phase model fluid; the effects "
    "of crude oil composition, including asphaltene adsorption and "
    "viscosity-dependent transport, remain to be evaluated. The long-term "
    "chemical stability of the epoxy matrix and the integrity of the "
    "stearic acid surface modification at temperatures above 120 degC, "
    "or in the presence of aggressive formation fluids (high salinity, "
    "CO2, H2S), require additional investigation."
)
for r in p153.runs: r.text = ''
p153.runs[0].text = new153
print('[153] rewritten')

# [154] Synthesis
p154 = doc.paragraphs[154]
new154 = (
    "Taken together, the single-phase and two-phase results demonstrate "
    "a practical workflow for tracer-based production monitoring. The "
    "effective flow rate Q, independently validated against the pump "
    "setting, serves as a physically grounded proxy for the per-interval "
    "oil production rate. Because the model parameters are physically "
    "constrained rather than purely empirical, they can be cross-checked "
    "against operational data, pump-set flow rate, wellbore geometry, "
    "transport time scales, providing an internal consistency verification "
    "that strengthens confidence in the interpreted results."
)
for r in p154.runs: r.text = ''
p154.runs[0].text = new154
print('[154] rewritten')

# [156] Field guidance
p156 = doc.paragraphs[156]
new156 = (
    "In field practice, ESP-T particles doped with distinct metal elements "
    "(e.g., Mn, Zn, Cu) for each fracture stage are co-injected with the "
    "proppant during hydraulic fracturing, enabling simultaneous multi-stage "
    "deployment in a single well. The fitted crossover time t0 and pulse-peak "
    "position guide shut-in duration design: the objective is a clearly "
    "identifiable Gaussian peak, beyond which additional shut-in time yields "
    "diminishing returns, since 47% of the signal resides in the sustained "
    "tail. Comparing Q values across stages identifies underperforming "
    "intervals, informing refracturing candidate selection and stimulation "
    "design adjustments for subsequent wells. Because FO is independent of "
    "total flow rate and scales predictably with OWR, periodic wellhead "
    "sampling during routine production, without additional shut-ins or "
    "downhole intervention, enables per-interval production tracking over "
    "the full well lifecycle."
)
for r in p156.runs: r.text = ''
p156.runs[0].text = new156
print('[156] rewritten')

doc.save('四氧化三铁环氧树脂拟合/ESP-T_终稿.docx')
print('\nSection 3.7 restructured and saved.')