#!/usr/bin/env python3
"""Sync Section 3.7 restructure to Chinese version."""
from docx import Document
doc = Document('四氧化三铁环氧树脂拟合/ESP-T_中文版.docx')

# [126] Restructure
p126 = doc.paragraphs[126]
new126 = (
    "井口记录的示踪剂突破曲线反映了两个不同传输过程的叠加：关井期间积蓄在支撑剂"
    "充填层中的示踪剂段塞的对流-弥散迁移，以及主脉冲通过后残余示踪剂从聚合物载体中"
    "的持续基质扩散控制释放。为在单一物理框架内捕捉这两个过程，我们从一维对流-弥散"
    "方程（ADE）出发构建复合模型：dC/dt + v dC/dx = D d2C/dx2，其中v = 4Q/(pi d2)"
    "为平均流速，D = alpha v为纵向弥散系数，Q为有效油相体积流量，d为管柱内径，"
    "alpha为纵向弥散度。"
    "\n\n"
    "两个经典ADE解析解构成模型的基础组件。第一个描述半无限域中在x = 0、t = 0处"
    "注入示踪剂质量M的瞬时段塞，产生高斯脉冲形式[29]："
    "C_rise(x,t) = (M/A)/sqrt(4 pi D t) exp[-(x - v t)2 / (4 D t)]，"
    "该分量捕捉关井期间积蓄示踪剂的对流-弥散传输。第二个描述入口处维持恒定浓度C0"
    "的连续源，产生互补误差函数形式："
    "C_fall(x,t) = (C0/2) erfc[(x - v t) / sqrt(4 D t)]，"
    "该分量捕捉环氧基体中示踪剂的持续释放。两种传输模式在生产过程中同时作用；通过"
    "双曲正切权重函数w(t) = 1/2 [1 + tanh((t - t0)/sigma)]将二者平滑拼接，其中t0"
    "为交叉中心时间，sigma固定为采样间隔（4 min）。完整复合模型为"
    "C(t) = cb + w(t) C_rise + [1 - w(t)] C_fall，如式（2）所示，"
    "tanh权重在早期使高斯脉冲占主导，后期平滑过渡至erfc拖尾。"
)
for r in p126.runs: r.text = ''
p126.runs[0].text = new126
print('[126] restructured')

# [140] Remove (redundant)
if len(doc.paragraphs) > 140:
    for r in doc.paragraphs[140].runs: r.text = ''
print('[140] removed')

# [142] Fitting results
p142 = doc.paragraphs[142]
new142 = (
    "模型拟合至单相示踪剂突破曲线（图3-8b），拟合参数汇总于表3-3。模型"
    "R2 = 0.9939，RMSE = 0.0210，全时间范围内残差随机分布于+/- 2 sigma以内。"
    "拟合有效流量Q = 0.46 mL/min与独立设定的泵流量0.50 mL/min高度吻合（相对误差"
    "8%）。拟合平均停留时间MRT = 37.4 min与对流传输时间x/v = 38.6 min一致"
    "（比值0.967）。Peclet数Pe = x/alpha = 0.934将传输状态置于对流主导与弥散"
    "主导之间的过渡区，与基质扩散控制源的预期一致。对拟合的C_rise和C_fall分量"
    "分别积分显示，47%的总积分示踪信号源于erfc拖尾。这证实了第3.6节识别的非Fick"
    "释放机制在流动条件下仍是主导传输模式，为ESP-T长期监测奠定了物理基础。"
)
for r in p142.runs: r.text = ''
p142.runs[0].text = new142
print('[142] rewritten')

# [146] Compact physical interpretation (find the right index)
for i in range(143, 150):
    if i < len(doc.paragraphs):
        t = doc.paragraphs[i].text
        if '拟合参数值得深入审视' in t or '拟合参数具有多个超越经验曲线拟合的物理意义' in t:
            new146 = (
                "拟合参数具有超越拟合质量的物理意义。独立恢复的Q、MRT和Pe均与已知泵设定、"
                "流道几何和第3.6节表征的非Fick释放动力学预测值一致。这种多参数一致性表明"
                "模型参数对应物理可测量量，可仅从示踪信号中恢复。47%的erfc拖尾贡献对现场"
                "作业具有特别的指导意义：近一半的可检测示踪信号源于持续释放而非初始关井"
                "脉冲，意味着长期监测不依赖反复关井循环。尾端平台浓度c_b虽在本单温度实验"
                "中作为常数拟合，但反映了第3.6节记录的稳态扩散通量。鉴于K-P速率常数K从"
                "30 degC的0.055系统性增至120 degC的0.196，c_b预期将随储层温度缩放，为"
                "温度依赖性模型参数化提供了途径。"
            )
            for r in doc.paragraphs[i].runs: r.text = ''
            doc.paragraphs[i].runs[0].text = new146
            print(f'[{i}] rewritten')
            break

# [148] Bridge
for i in range(145, 152):
    if i < len(doc.paragraphs):
        t = doc.paragraphs[i].text
        if '为评估ESP-T在两相流条件下的表现' in t or '为探讨油水两相流' in t:
            new148 = (
                "单相分析确立了模型从示踪信号恢复产油速率的能力。为评估该能力是否可"
                "拓展至两相流动条件，在三种油水比（OWR = 4:1、1:1、1:4）和四种总流量"
                "（0.1-0.4 mL/min）下进行了稳态岩心驱替实验。直接观察OWR = 1:1时流出液，"
                "确认早期流出液中油相分数略高，与第3.5节建立的ESP-T亲油特性一致。"
            )
            for r in doc.paragraphs[i].runs: r.text = ''
            doc.paragraphs[i].runs[0].text = new148
            print(f'[{i}] rewritten')
            break

# [151] FO results
for i in range(148, 154):
    if i < len(doc.paragraphs):
        t = doc.paragraphs[i].text
        if '示踪剂通量（FO）' in t and '定义为单位时间' in t:
            new151 = (
                "稳态生产过程中，示踪剂浓度本身较低，突破曲线主要通过浓度信号的幅值和"
                "趋势传递信息。为提取定量速率信息，引入示踪剂通量FO，其定义为单位时间"
                "通过井口的示踪剂质量，等于ESP-T的示踪剂释放速率[8]。图3-9(b)显示FO随"
                "OWR增大而升高，但与总流量无关。这一行为与ESP-T的亲油特性物理一致：恒定"
                "OWR下油相与支撑剂接触面积固定，因此无论总流量如何FO保持恒定；增大OWR"
                "则扩大了接触面积，使FO升高。"
            )
            for r in doc.paragraphs[i].runs: r.text = ''
            doc.paragraphs[i].runs[0].text = new151
            print(f'[{i}] rewritten')
            break

# [153] Validation + limitations
for i in range(150, 158):
    if i < len(doc.paragraphs):
        t = doc.paragraphs[i].text
        if '归一化FO' in t and '3.187' in t:
            new153 = (
                "图3-9(c)比较了在总流量0.1 mL/min下三种油水比的归一化FO与实际油相流量。"
                "归一化以单相油驱替稳态FO（3.187 micro-g/min）为基准。归一化FO与实际"
                "油相流量的高度吻合表明，在稳态两相流条件下，当总流量已知时，可从FO变化"
                "曲线定量得出标记层段的产油速率。"
                "\n\n"
                "当前验证存在若干局限需要注意。ADE模型假设单裂缝段和均匀支撑剂充填；"
                "现场多段相互作用和充填非均质性可能引入偏差。FO标定依赖稳态流动；在开井、"
                "关井或快速降压等瞬变流态下，FO与产油速率的关系可能不成立。实验使用十二烷"
                "作为油相模型流体；原油组分效应（包括沥青质吸附和粘度依赖性传输）有待评估。"
                "环氧基体在120 degC以上温度或强腐蚀性流体（高矿化度、CO2、H2S）环境中的"
                "长期化学稳定性及硬脂酸表面改性的完整性需要进一步研究。"
            )
            for r in doc.paragraphs[i].runs: r.text = ''
            doc.paragraphs[i].runs[0].text = new153
            print(f'[{i}] rewritten')
            break

# [154] Synthesis
for i in range(152, 158):
    if i < len(doc.paragraphs):
        t = doc.paragraphs[i].text
        if '工程价值' in t and '闭合' in t:
            new154 = (
                "综合单相和两相结果，本文展示了一套实用的示踪基产量监测工作流。拟合"
                "有效流量Q经泵设定值独立验证后，可作为各段产油速率的物理代理指标。由于"
                "模型参数受物理约束而非纯经验性，它们可与作业数据——泵设流量、井筒几何"
                "尺寸、传输时间尺度——交叉核对，提供内部一致性检验，增强所解释结果的"
                "置信度。"
            )
            for r in doc.paragraphs[i].runs: r.text = ''
            doc.paragraphs[i].runs[0].text = new154
            print(f'[{i}] rewritten')
            break

# [156] Field guidance
for i in range(154, 160):
    if i < len(doc.paragraphs):
        t = doc.paragraphs[i].text
        if '掺杂不同金属元素' in t or '各压裂段的Fe3O4核心' in t:
            new156 = (
                "在现场实践中，各压裂段掺杂不同金属元素（如Mn、Zn、Cu）的ESP-T颗粒"
                "随支撑剂在水力压裂过程中共注入，实现单井多段同时部署。拟合交叉时间t0"
                "和脉冲峰位指导关井时长设计：目标是产生清晰可辨的高斯峰，超过此点额外"
                "关井时间收益递减，因为47%的信号位于持续拖尾中。比较各段Q值可识别低产"
                "层段，为重复压裂候选井筛选和后续井增产设计调整提供依据。由于FO与总流量"
                "无关且随OWR可预测缩放，常规生产过程中的定期井口采样——无需额外关井或"
                "井下作业——即可在全井生命周期内跟踪各段产量贡献。"
            )
            for r in doc.paragraphs[i].runs: r.text = ''
            doc.paragraphs[i].runs[0].text = new156
            print(f'[{i}] rewritten')
            break

doc.save('四氧化三铁环氧树脂拟合/ESP-T_中文版.docx')
print('\nCN Section 3.7 synced.')