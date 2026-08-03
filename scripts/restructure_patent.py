# -*- coding: utf-8 -*-
"""将专利申请文档重组为正式提交格式：
说明书（无连续编号）| 权利要求书 | 摘要 | 摘要附图
移除误入说明书的申请人/发明人信息（应放入请求书表格）
"""
import sys, io, os, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"荧光压裂液\专利申请文件_已修正.docx"
DST = r"荧光压裂液\专利申请文件_正式格式.docx"

src_doc = Document(SRC)

# ── Collect all paragraphs from source ──
src_paras = [(p.text.strip(), p.text) for p in src_doc.paragraphs]

# ── Helper: create a clean new document ──
doc = Document()

# Set default font for the document
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Page margins (standard A4 for Chinese patents)
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

def add_heading_para(text, font_name='黑体', size=Pt(14), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """Add a centered heading paragraph."""
    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = size
    run.bold = bold
    return para

def add_section_heading(text):
    """Add a bold section heading (left-aligned)."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(12)
    run.bold = True
    return para

def add_body_para(text):
    """Add a body paragraph with first-line indent."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.first_line_indent = Cm(0.74)  # ~2 Chinese chars
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    return para

def add_page_break():
    """Add a page break."""
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)

def add_blank_line():
    doc.add_paragraph()

# ── ========== Parse source document into sections ========== ──

# Classify each source paragraph
HEADING_PATTERNS = {
    'invention_name_label': re.compile(r'^一、\s*发明名称'),
    'applicant_label': re.compile(r'^二、\s*申请人'),
    'inventor_label': re.compile(r'^三、\s*发明人'),
    'technical_field': re.compile(r'^四、\s*技术领域'),
    'background': re.compile(r'^五、\s*背景技术'),
    'invention_content': re.compile(r'^六、\s*发明内容'),
    'brief_description': re.compile(r'^七、\s*附图说明'),
    'detailed_description': re.compile(r'^八、\s*具体实施方式'),
    'claims_title': re.compile(r'^九、\s*权利要求书'),
    'abstract_title': re.compile(r'^十、\s*摘要'),
    'abstract_fig_title': re.compile(r'^十一、\s*摘要附图'),
}

# Parse into labeled segments
segments = []
current_label = None
current_paras = []
invention_name = ""

for stripped, raw in src_paras:
    if not stripped:
        if current_paras:
            current_paras.append("")
        continue

    matched = None
    for label, pattern in HEADING_PATTERNS.items():
        if pattern.match(stripped):
            if current_label and current_paras:
                segments.append((current_label, current_paras))
            current_label = label
            current_paras = []
            matched = label
            break

    if matched:
        continue
    else:
        current_paras.append(stripped)

# Don't forget the last segment
if current_label and current_paras:
    segments.append((current_label, current_paras))

# Build a lookup: label -> paragraphs
seg_dict = {}
for label, paras in segments:
    if label not in seg_dict:
        seg_dict[label] = []
    seg_dict[label].extend(paras)

# The invention name is the first non-empty paragraph after "一、发明名称"
# Actually, let's extract it directly: it's the first body paragraph
# Find the invention name from original paragraphs
found_title = False
for stripped, raw in src_paras:
    if HEADING_PATTERNS['invention_name_label'].match(stripped):
        found_title = True
        continue
    if found_title and stripped:
        invention_name = stripped
        break

if not invention_name:
    invention_name = "一种用于压裂裂缝荧光示踪的改性稀土铝酸盐荧光粉、荧光压裂液体系及其应用方法"

print(f"发明名称: {invention_name}")
print(f"解析到的章节: {list(seg_dict.keys())}")

# ── ========== BUILD DOCUMENT ========== ──

# ═══════════════ PART 1: 说明书 ═══════════════

# 1.0 发明名称 (no label, centered)
add_heading_para(invention_name, font_name='黑体', size=Pt(16), bold=True)
add_blank_line()

# 1.1 技术领域
add_section_heading("技术领域")
for p in seg_dict.get('technical_field', []):
    if p.strip():
        add_body_para(p)

# 1.2 背景技术
add_section_heading("背景技术")
for p in seg_dict.get('background', []):
    if p.strip():
        add_body_para(p)

# 1.3 发明内容
add_section_heading("发明内容")

# 发明内容内部: 把 发明目的/技术方案/有益效果 作为子标题
# Parse the invention_content paragraphs and find sub-sections
ic_paras = seg_dict.get('invention_content', [])

# Sub-section markers within 发明内容
sub_markers = {
    '发明目的': '发明目的',
    '技术方案': '技术方案',
    '有益效果': '有益效果',
}

# Parse invention_content into sub-parts
ic_sub = {}
current_sub = None
current_sub_paras = []

for p in ic_paras:
    matched_sub = None
    for marker_key, marker_text in sub_markers.items():
        if p.strip().startswith(marker_text) and len(p.strip()) <= 10:
            if current_sub and current_sub_paras:
                ic_sub[current_sub] = current_sub_paras
            current_sub = marker_key
            current_sub_paras = []
            matched_sub = marker_key
            break
    if matched_sub:
        continue
    else:
        current_sub_paras.append(p)

if current_sub and current_sub_paras:
    ic_sub[current_sub] = current_sub_paras

# Write sub-sections
for sub_name in ['发明目的', '技术方案', '有益效果']:
    paras = ic_sub.get(sub_name, [])
    if paras:
        # Sub-heading in bold
        add_section_heading(sub_name)
        for p in paras:
            if p.strip():
                add_body_para(p)

# 1.4 附图说明
add_section_heading("附图说明")
for p in seg_dict.get('brief_description', []):
    if p.strip():
        add_body_para(p)

# 1.5 具体实施方式
add_section_heading("具体实施方式")
for p in seg_dict.get('detailed_description', []):
    if p.strip():
        add_body_para(p)

# ═══════════════ PART 2: 权利要求书 (new page) ═══════════════
add_page_break()
add_heading_para("权利要求书", font_name='黑体', size=Pt(16), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_blank_line()

for p in seg_dict.get('claims_title', []):
    if p.strip():
        # Claims already fixed to "1." format from previous fix
        add_body_para(p)

# ═══════════════ PART 3: 说明书摘要 (new page) ═══════════════
add_page_break()
add_heading_para("摘要", font_name='黑体', size=Pt(16), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_blank_line()

for p in seg_dict.get('abstract_title', []):
    if p.strip():
        add_body_para(p)

# ═══════════════ PART 4: 摘要附图 ═══════════════
add_page_break()
add_heading_para("摘要附图", font_name='黑体', size=Pt(16), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_blank_line()

for p in seg_dict.get('abstract_fig_title', []):
    if p.strip():
        add_body_para(p)

# ── ========== SAVE ========== ──
doc.save(DST)
print(f"\n✅ 正式格式文件已保存至: {os.path.abspath(DST)}")
print(f"\n文档结构:")
print(f"  说明书: 发明名称 → 技术领域 → 背景技术 → 发明内容 → 附图说明 → 具体实施方式")
print(f"  权利要求书 (独立分页)")
print(f"  摘要 (独立分页)")
print(f"  摘要附图 (独立分页)")
print(f"\n已移除: 二、申请人 和 三、发明人 (应填入发明专利请求书表格)")