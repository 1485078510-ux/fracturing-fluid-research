#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Restructure thesis:
1. Move 2.5 (crosslinking impact) -> merge into 3.3
2. Add "方案选择依据" to 2.3
3. Add "假说否证" to 4.3
4. Add "失效边界" to 4.6
5. Compress Chapter 5
6. Add "其他岩性拓展" to 6.3
"""
import sys
from copy import deepcopy
from docx import Document
from docx.shared import Pt, RGBColor
from lxml import etree

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

DOC = r'荧光压裂液/论文初稿_v2.docx'
doc = Document(DOC)

# ============================================================
# HELPERS
# ============================================================
def find_para_by_text(doc, snippet, skip=0):
    """Find paragraph index containing snippet (with skip for nth match)."""
    found = 0
    for i, para in enumerate(doc.paragraphs):
        if snippet in para.text:
            if found >= skip:
                return i, para
            found += 1
    return None, None

def clone_element(elem):
    """Deep-copy an XML element."""
    return deepcopy(elem)

def insert_after(target_para, new_element):
    """Insert new_element XML node after target_para."""
    target_para._element.addnext(new_element)

def delete_para(para):
    """Remove paragraph from document."""
    para._element.getparent().remove(para._element)

def add_text_para_after(doc, target_para, text):
    """Create a new paragraph with text and insert after target_para."""
    new_para = doc.add_paragraph(text)
    target_para._element.addnext(new_para._element)
    return new_para

def replace_heading_text(para, new_text):
    """Replace heading text in paragraph."""
    # Replace in the first run
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)

def para_text(para):
    """Get paragraph text safely."""
    return para.text.strip()


# ============================================================
# CHANGE 1: Move 2.5 (P90-P92) -> merge into 3.3 (after P108)
# ============================================================
print("=" * 60)
print("CHANGE 1: Move 2.5 -> 3.3")

# Find key paragraphs by text matching
# 2.5 heading: contains "2.5" and "改性对胍胶交联"
idx_25_heading, p_25_heading = find_para_by_text(doc, '2.5  改性对胍胶交联')
# 2.5 content: contains "残余微量的可溶出Al"
idx_25_content, p_25_content = find_para_by_text(doc, '残余微量的可溶出Al')
# 2.5 data placeholder
idx_25_data, p_25_data = find_para_by_text(doc, '实验数据待补充', skip=2)
# 2.6 heading
idx_26_heading, p_26_heading = find_para_by_text(doc, '2.6  本章小结')

# Target in 3.3: find paragraph with Gomaa and Goel citations (P108)
idx_33_target, p_33_target = find_para_by_text(doc, 'Gomaa等（2015）[51]和Goel')

print(f'2.5 heading: P{idx_25_heading}')
print(f'2.5 content: P{idx_25_content}')
print(f'2.6 heading: P{idx_26_heading}')
print(f'3.3 target: P{idx_33_target}')

# Clone 2.5 content and insert after 3.3's last content paragraph
cloned_content = clone_element(p_25_content._element)
insert_after(doc.paragraphs[idx_33_target], cloned_content)
print('Inserted 2.5 content into 3.3')

# Also insert a transition sentence before the moved content
# Find the cloned paragraph in its new position
new_text = (
    '此外，改性荧光粉表面可能残余的Al³⁺等多价阳离子对交联反应的干扰亦不容忽视——'
    'Al³⁺可与HPG分子链上的顺式邻二醇发生配位交联，干扰有机硼交联剂与HPG的正常交联反应。'
    '为此，在评价冻胶流变性能的同时，需对比含/不含改性荧光粉的HPG冻胶的交联行为：'
    '（1）交联时间测定——采用旋转粘度计，记录加入交联剂后粘度上升至最大值90%所需时间；'
    '（2）粘弹性温度扫描——与前述方法相同，对比含/不含荧光粉的冻胶体系；'
    '（3）冻胶微观结构SEM——冷冻干燥后观察含/不含荧光粉的冻胶网络形貌差异。'
    '通过含与不含荧光粉的平行对比，定量评估荧光粉引入对交联过程的干扰程度。'
)
# Actually, since we moved the content, let's just modify its text slightly
# to add more context at the beginning. The original paragraph (P91) is now deleted
# but we inserted the clone. Let me find the clone by content.
for i, para in enumerate(doc.paragraphs):
    if i > idx_33_target and i < idx_33_target + 5:
        if '残余微量的可溶出Al' in para.text:
            # This is our inserted clone - prepend a brief transition
            new_full_text = (
                '在上述流变性能评价的基础上，还需单独考察荧光粉对交联化学的潜在干扰。'
                + para.text
            )
            for run in para.runs:
                run.text = ''
            if para.runs:
                para.runs[0].text = new_full_text
            else:
                para.add_run(new_full_text)
            print('Enhanced moved content with transition sentence')
            break

# Delete 2.5 heading, content, and data placeholder from Chapter 2
# Need to re-find since indices may have shifted
_, p25h = find_para_by_text(doc, '2.5  改性对胍胶交联')
_, p25c = find_para_by_text(doc, '残余微量的可溶出Al')
_, p25d = find_para_by_text(doc, '实验数据待补充', skip=2)
delete_para(p25h)
delete_para(p25c)
delete_para(p25d)
print('Deleted 2.5 from Chapter 2')

# Renumber 2.6 -> 2.5
_, p26h = find_para_by_text(doc, '2.6  本章小结')
if p26h:
    if p26h.runs:
        for run in p26h.runs:
            if '2.6' in run.text:
                run.text = run.text.replace('2.6', '2.5')
                break
    print('Renumbered 2.6 -> 2.5')


# ============================================================
# CHANGE 2: Add "方案选择依据" to 2.3 (after P79)
# ============================================================
print("\n" + "=" * 60)
print("CHANGE 2: Add 方案选择依据 to 2.3")

idx_23, p_23 = find_para_by_text(doc, '表面改性的目标是在荧光粉表面构建一个兼具三重功能的双层结构')
print(f'2.3 first content para: P{idx_23}')

scheme_text = (
    '改性方案的选型基于以下考量。在无机包覆路线中，SiO₂包覆（Qi等，2017）虽可有效阻隔水解，'
    '但致密SiO₂层不具备分散功能，包覆后仍需额外的表面活性剂修饰；Al₂O₃/TiO₂ ALD包覆（Karacaoglu等，2020）'
    '虽可实现纳米级精确厚度控制，但ALD工艺成本高昂、批次处理量小，不适用于压裂现场所需的公斤级以上改性规模。'
    '在有机-无机复合路线中，Lyu等（2020）的SiO₂-KH570-聚合物方案虽兼顾了耐水性和有机相容性，'
    '但未考虑破胶环境下的表面功能切换需求。经综合比较，本文选择"硅烷偶联剂化学锚固+PEG物理屏蔽"的纯有机方案：'
    'KH550的Si-OH端与荧光粉表面羟基缩合提供化学锚固，另一端-NH₂在PEG包覆期间被屏蔽、待破胶脱附后暴露以驱动壁面锚定；'
    'PEG4000作为外层提供空间位阻分散，同时在水解环境中作为牺牲层保护内层KH550。该方案不涉及无机前驱体水解的复杂动力学控制，'
    '工艺条件温和（室温、水相），适合工业放大。'
)
add_text_para_after(doc, doc.paragraphs[idx_23], scheme_text)
print('Added 方案选择依据 paragraph')


# ============================================================
# CHANGE 3: Add "假说否证" to 4.3 (after P135)
# ============================================================
print("\n" + "=" * 60)
print("CHANGE 3: Add 假说否证 to 4.3")

idx_43, p_43 = find_para_by_text(doc, '本研究的核心工作假说')
idx_43_data, p_43_data = find_para_by_text(doc, '实验数据待补充', skip=3)
print(f'4.3 first para: P{idx_43}, data placeholder: P{idx_43_data}')

falsify_text = (
    '为保证论证的完备性，需考虑上述三重交叉验证的可能结果及其对应的机制解释。'
    '若FTIR中PEG特征峰无减弱、XPS中N 1s无增强、Zeta电位无降低——即三者均不支持PEG脱附——'
    '则表明过硫酸铵氧化环境不足以诱导PEG链的断裂脱附，"破胶诱导功能切换"假说被否定。'
    '在此情形下，需考虑以下替代锚定机制：（a）PEG链虽未脱附，但其在高温（90°C）和破胶液离子强度下'
    '可能发生构象坍缩，使PEG刷厚度降低，颗粒与壁面间距缩小至范德华力作用范围内，实现非特异性物理附着；'
    '（b）石英砂表面的硅羟基与PEG醚氧原子之间的氢键作用可能在PEG未完全脱附时即已贡献部分锚定力；'
    '（c）破胶后冻胶网络坍缩形成的聚合物残渣可能在颗粒与壁面之间起"胶黏桥接"作用。'
    '若三项指标中仅部分支持脱附（如FTIR支持但Zeta电位不支持），则提示存在不完全脱附或空间异质性，'
    '需结合XPS深度剖析和TOF-SIMS成像进一步表征表面的面内化学不均匀性。'
    '以上备选机制和部分支持情形的讨论确保本研究对假说检验的严谨性——无论假说被证实、部分证实还是被否证，'
    '均能从分子层面给出颗粒表面化学在破胶前后变化的合理解释。'
)
add_text_para_after(doc, doc.paragraphs[idx_43], falsify_text)
print('Added 假说否证 paragraph')


# ============================================================
# CHANGE 4: Add "失效边界" to 4.6 (after P152)
# ============================================================
print("\n" + "=" * 60)
print("CHANGE 4: Add 失效边界 to 4.6")

idx_46_boundary, p_46_boundary = find_para_by_text(doc, '该半定量关系受控于同一砂岩类型')
print(f'4.6 boundary statement: P{idx_46_boundary}')

failure_text = (
    '为明确该半定量映射关系的工程适用边界，还需系统识别可能导致灰度-宽度关系失效的关键因素。'
    '以下情形下，灰度值与裂缝宽度之间的单调映射预期将退化或失效：（1）裂缝壁面矿物组分显著差异——'
    '当砂岩中粘土矿物（如蒙脱石、伊利石）含量超过10%~15%时，粘土对荧光粉的非选择性吸附将产生'
    '与裂缝宽度无关的高背景荧光，掩埋宽度信号；（2）荧光粉批次间差异——不同批次的商用荧光粉在粒度分布、'
    '表面缺陷密度和发光量子效率方面存在波动，若未对每批次建立独立的灰度-浓度标定曲线，批次效应可导致'
    '灰度值系统性偏移20%~30%；（3）裂缝壁面粗糙度极端值——当JRC>18时，粗糙壁面的"阴影效应"使紫外'
    '激发光和荧光发射均受到几何遮挡，灰度值不再仅反映荧光粉面密度；（4）多次注入或重复改造——当同一井段'
    '经历两次以上压裂施工时，后续注入的荧光粉可能与前期残留叠加，产生非加和性的灰度响应。'
    '在上述情形下，灰度值仅保留"有/无"定性判读功能，不可用于裂缝宽度推断。'
    '此外，灰度与裂缝宽度的定量对应关系受控于注入阶段荧光粉浓度和流速——'
    '此二参数在本文中被锁定，但在现场尺度下并非恒定。因此，在工程推广中，本法始终定位为定性-半定量工具，'
    '不替代微地震等间接反演技术的定量输出。'
)
add_text_para_after(doc, doc.paragraphs[idx_46_boundary], failure_text)
print('Added 失效边界 paragraph')


# ============================================================
# CHANGE 5: Compress Chapter 5
# ============================================================
print("\n" + "=" * 60)
print("CHANGE 5: Compress Chapter 5")

# 5.1 heading - keep as is
_, p51h = find_para_by_text(doc, '5.1  "母液预配+在线稀释"工艺方案')

# 5.2 heading - change to merge 5.2+5.3+5.4
_, p52h = find_para_by_text(doc, '5.2  单井材料用量估算')
print(f'5.2 heading: {para_text(p52h)[:60]}')

# Replace 5.2 heading with new merged heading
if p52h.runs:
    p52h.runs[0].text = '5.2  经济技术评估'
print('Changed 5.2 heading -> "5.2 经济技术评估"')

# Delete 5.3 heading
_, p53h = find_para_by_text(doc, '5.3  经济性评估')
if p53h:
    delete_para(p53h)
    print('Deleted 5.3 heading')

# Delete 5.4 heading
_, p54h = find_para_by_text(doc, '5.4  环保与安全性评估')
if p54h:
    delete_para(p54h)
    print('Deleted 5.4 heading')

# Renumber 5.5 -> 5.3
_, p55h = find_para_by_text(doc, '5.5  技术定位与适用条件')
if p55h:
    for run in p55h.runs:
        if '5.5' in run.text:
            run.text = run.text.replace('5.5', '5.3')
            break
    print('Renumbered 5.5 -> 5.3')

# Delete 5.6 heading and its placeholder
_, p56h = find_para_by_text(doc, '5.6  本章小结')
_, p56p = find_para_by_text(doc, '待实验完成后撰写', skip=4)
if p56h:
    delete_para(p56h)
if p56p:
    delete_para(p56p)
print('Deleted 5.6 heading and placeholder')


# ============================================================
# CHANGE 6: Add "其他岩性拓展" to 6.3 (after P199)
# ============================================================
print("\n" + "=" * 60)
print("CHANGE 6: Add 其他岩性拓展 to 6.3")

idx_63, p_63 = find_para_by_text(doc, '6.3  不足与展望')
print(f'6.3 heading: P{idx_63}')

lithology_text = (
    '（1）储层岩性拓展。本研究的锚定机制依赖砂岩表面的硅羟基（Si-OH），方法适用性受储层矿物学约束。'
    '对于碳酸盐岩储层（方解石/白云石，表面以Ca²⁺/Mg²⁺位点为主），可探索以膦酸基团（-PO₃H₂）替代氨基（-NH₂）'
    '作为锚定端基——膦酸根与Ca²⁺的配位键强度远高于静电吸引，有望实现碳酸盐岩壁面的化学特异性锚定。'
    '对于页岩储层（富含有机质和粘土矿物），表面的化学异质性对单一锚定策略构成挑战，可考虑引入两性离子修饰'
    '或疏水-亲水平衡策略以适应页岩表面的复杂化学组成。上述拓展需重新设计偶联剂端基化学，并通过岩心驱替实验逐类验证。'
)
add_text_para_after(doc, doc.paragraphs[idx_63], lithology_text)

# Also add another paragraph about scale-up and comparison experiments
scaleup_text = (
    '（2）尺度验证与对比实验。本研究的动态驱替验证局限于实验室岩心尺度（~10 cm），'
    '从实验室到现场的跨越中，需开展大型岩样（米级）验证实验以评估裂缝尺寸放大对颗粒运移-锚定行为的影响。'
    '同时，本研究的核心优势——无机荧光粉对氧化破胶剂的本征化学惰性——尚缺乏与有机荧光示踪材料的直接对比实验数据。'
    '建议在后续工作中以相同条件对比SrAl₂O₄与荧光素钠、罗丹明B或量子点在过硫酸铵溶液中的发光保持率，'
    '将这一关键比较优势从理论推论转化为实验事实。'
)
# Insert after the lithology paragraph we just added
# Find it
for i, para in enumerate(doc.paragraphs):
    if '储层岩性拓展' in para.text:
        add_text_para_after(doc, para, scaleup_text)
        print('Added 岩性拓展 + 尺度验证 paragraphs')
        break

print("\n" + "=" * 60)
print("All restructuring changes applied.")

# ============================================================
# SAVE
# ============================================================
doc.save(DOC)
print(f'Saved to {DOC}')