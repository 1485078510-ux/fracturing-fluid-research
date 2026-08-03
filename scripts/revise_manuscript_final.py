"""
Revise ESP-T manuscript: Introduction, Conclusions, Abstract, and Section 3.3 structure.
Saves to ESP-T_Final_4-revised.docx
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
import copy

INPUT = r"c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP-T_投稿文件\ESP-T_Final_3-revised.docx"
OUTPUT = r"c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP-T_投稿文件\ESP-T_Final_4-revised.docx"

doc = Document(INPUT)

# === HELPERS ===

def clear_para(p):
    """Remove all content from a paragraph element."""
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    for child in list(p._element):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('r', 'drawing', 'pict'):
            p._element.remove(child)

def write_para(p, text):
    """Replace paragraph content with new text."""
    clear_para(p)
    p.add_run(text)

def find_para_containing(doc, text, start=0):
    """Find first paragraph whose text contains the given string."""
    for i, p in enumerate(doc.paragraphs):
        if i >= start and text in p.text:
            return i, p
    return None, None

# === 1. REVISED ABSTRACT ===

REVISED_ABSTRACT = (
    "Tracer proppants enable long-term per-interval monitoring without downhole hardware, "
    "but interpreting their wellhead breakthrough curves (BTCs) requires jointly accounting "
    "for sustained release from the polymer matrix and transport through the production tubing. "
    "Current practice relies on the zero-dimensional Korsmeyer-Peppas (K-P) model to characterize "
    "batch release kinetics; the advection-dispersion framework widely used for transport "
    "interpretation has not been coupled to the sustained source term that these materials generate. "
    "We develop a piecewise advection-dispersion model that decomposes the BTC into a Gaussian pulse "
    "(shut-in accumulation slug) and an erfc tail (sustained matrix-diffusion-controlled release), "
    "linked by a smooth tanh transition. The six-parameter model is validated with an oleophilic "
    "epoxy/Fe₃O₄ tracer proppant (ESP-T) in core displacement experiments. The dual-component "
    "formulation is statistically decisive over four single-component alternatives (ΔAICc = 32.7, "
    "F-test p < 10⁻⁶, R² = 0.9939). The fitted flow rate (0.46 mL/min) agrees with "
    "the independent pump setting (0.50 mL/min) within 8% without being constrained in the objective "
    "function. The Peclet number (Pe = 0.934) independently confirms the non-Fickian transport "
    "mechanism identified via K-P kinetics (n = 0.45–0.85). The erfc tail accounts for 47% of "
    "the integrated signal, a result stable to within ±0.8% across a six-fold variation in the "
    "transition width. Under two-phase flow, the oil-phase tracer mass flux eliminates water-dilution "
    "artifacts and tracks oil production rates (Pearson r = 0.97, RMSD = 8.3%). The framework enables "
    "per-interval production allocation from wellhead samples alone, requiring no downhole tools and "
    "only a single shut-in."
)

# === 2. REVISED INTRODUCTION ===

INTRO_PARAS = [
    # Para 1: Problem + existing tech + tracers
    (
        "Determining per-stage production contributions in multi-stage fractured horizontal "
        "wells remains an unsolved operational challenge [5,6]. Without per-stage data, "
        "underperforming intervals go undiagnosed, well spacing cannot be optimized, and "
        "completion designs cannot be validated against production outcomes [7,8]. Existing "
        "downhole diagnostic methods each involve trade-offs: production logging requires well "
        "intervention and provides only snapshot measurements [9]; distributed fiber-optic "
        "sensing demands permanent cable installation at costs difficult to justify for marginal "
        "wells [10,11]; microseismic monitoring characterizes fracture geometry rather than "
        "production contribution [12]. Tracers offer an alternative that avoids downhole "
        "hardware—chemical or radioactive species injected with the fracturing fluid produce "
        "time-resolved wellhead concentration signals that have been used to infer inter-well "
        "connectivity [15,16], diagnose fracture interference [17], and qualitatively assess "
        "stage-level inflow [18–20]."
    ),
    # Para 2: Dissolved tracers -> tracer proppants
    (
        "A limitation of dissolved tracers is their finite monitoring window: as the tracer is "
        "progressively diluted and produced, the signal decays. Tracer proppants address this by "
        "immobilizing the tracer agent within a solid carrier that is co-injected with the proppant "
        "pack, so the tracer remains in the fracture after placement and enables long-term monitoring "
        "without repeated injection [21]. A range of material designs have been reported—ceramic "
        "carriers with organic dye coatings [21,22], rare-earth-doped polymer matrices for "
        "multi-element coding [23], oleophilic Fe₃O₄/polystyrene microspheres [24], and "
        "dual-zone polymer-coated proppants [25,26]."
    ),
    # Para 3: K-P model and its limitation
    (
        "To date, tracer proppant performance has been evaluated through a single experimental "
        "paradigm: batch release measurements in a well-mixed vessel, interpreted with the "
        "Korsmeyer-Peppas (K-P) power law, C/C₀ = K·tⁿ [28–30]. The K-P model "
        "identifies the release mechanism (Fickian diffusion for n ≤ 0.43, anomalous transport "
        "for 0.43 < n < 0.85, Case-II relaxation for n ≥ 0.85) and provides the "
        "temperature-dependent rate constant K. However, K-P is a zero-dimensional batch model: it "
        "describes release into a fixed-volume vessel, with no spatial coordinate, no flow field, "
        "and no pathway from a release rate to a concentration measured at a distant sampling point "
        "[28–30]. It can characterize how fast the tracer leaves the proppant in a beaker; it "
        "cannot predict what concentration will be observed at the wellhead, when it will peak, or "
        "how it will decay—the features of a breakthrough curve that encode production information."
    ),
    # Para 4: ADE-based BTC interpretation and its assumption
    (
        "Independently, a mature methodology exists for interpreting BTCs on the transport side. "
        "The one-dimensional advection-dispersion equation (ADE), ∂C/∂t + v·∂C/∂x "
        "= D·∂²C/∂x², provides analytical solutions that relate the shape "
        "of a BTC to transport parameters [31]. For a pulse injection of known mass, the temporal "
        "moments of the BTC yield the mean residence time—hence the effective flow "
        "velocity—and the variance, which gives the dispersivity [32]. Full-curve ADE inversion "
        "has been applied to extract reservoir and fracture properties from BTC shapes [16,33] and "
        "to model tracer transport in multi-stage fractured wells with phase partitioning [34]. "
        "The common premise across these applications is that the tracer source is known: an injection "
        "pulse of specified mass, duration, and location [31,32]. The BTC is interpreted as the "
        "system’s transport response to that known input."
    ),
    # Para 5: The joint estimation problem
    (
        "A tracer proppant BTC departs from this premise. The tracer source is not a single, "
        "operator-controlled injection but a sustained release governed by matrix diffusion, "
        "solvent-driven swelling, and the time-varying concentration gradient at the proppant-fluid "
        "interface—processes that continue throughout production and whose rate parameters are "
        "not known a priori. The observed BTC is therefore shaped jointly by an unknown release "
        "function and unknown transport parameters. Extracting the per-interval flow rate—the "
        "quantity required for production allocation—calls for estimating both sets of unknowns "
        "simultaneously from a single curve. This joint estimation problem has received little "
        "attention in the tracer proppant literature, where the release (K-P) and transport (ADE) "
        "frameworks have been applied separately. The practical consequence is that current tracer "
        "proppants can confirm which stages produce but cannot quantify how much each stage produces "
        "from the BTC alone."
    ),
    # Para 6: This work
    (
        "In this work, we develop a method for jointly estimating release and transport parameters "
        "from a single tracer proppant BTC. The BTC is decomposed into two additive components linked "
        "by a smooth hyperbolic-tangent transition: a Gaussian pulse representing tracer that "
        "accumulated in the near-wellbore region during shut-in and is swept to the sampling point "
        "as a coherent slug upon flowback, and an erfc tail representing sustained "
        "matrix-diffusion-controlled release that continues after the main slug has passed. The "
        "model contains six free parameters, estimated simultaneously by nonlinear least squares; "
        "no parameter is independently constrained in the objective function. The fit yields both "
        "the release partition (Gaussian vs.\ erfc integrated contributions) and the transport "
        "parameters (effective flow rate Q, dispersivity α, Peclet number Pe). The method is "
        "validated in single-phase and two-phase core displacement experiments using an oleophilic "
        "epoxy/Fe₃O₄ tracer proppant (ESP-T). The epoxy matrix provides the sustained, "
        "diffusion-controlled release that the erfc component of the model represents, while stearic "
        "acid surface modification (water contact angle 104.6°) directs tracer release to the "
        "oil phase, isolating the oil production signal. We compare the model against four "
        "single-component alternatives using information-theoretic criteria, test whether the fitted "
        "flow rate agrees with the independent pump setting, and verify that the decomposition into "
        "Gaussian and erfc fractions is stable under variation of the transition-width parameter. "
        "Under two-phase flow, we use the oil-phase tracer mass flux—rather than "
        "concentration—to track oil production rates across varying oil-water ratios, "
        "eliminating the dilution artifact inherent in concentration-based interpretation."
    ),
]

# === 3. REVISED CONCLUSIONS ===

CONCLUSION_PARAS = [
    # Para 1: Method summary + key modeling findings
    (
        "A method for jointly estimating release and transport parameters from a single tracer "
        "proppant breakthrough curve was developed and validated at the laboratory scale. The BTC "
        "is decomposed into a Gaussian pulse (shut-in accumulation slug) and an erfc tail (sustained "
        "matrix-diffusion-controlled release), connected by a smooth tanh transition; six parameters "
        "are estimated simultaneously by nonlinear least squares. Applied to single-phase core "
        "displacement data from an oleophilic epoxy/Fe₃O₄ tracer proppant (ESP-T), the "
        "two-component model is strongly preferred over four single-component alternatives "
        "(ΔAICc = 32.7, F-test p < 10⁻⁶, R² = 0.9939). The fitted flow rate "
        "(0.46 mL/min) agrees with the independent pump setting (0.50 mL/min) within 8% without "
        "being constrained in the objective function, confirming that the model extracts physically "
        "meaningful transport parameters from the BTC shape. The Peclet number (Pe = 0.934) "
        "independently corroborates the non-Fickian transport regime identified via K-P kinetics "
        "(n = 0.45–0.85). The erfc tail accounts for 47% of the integrated tracer signal, a "
        "result stable to within ±0.8% across a six-fold variation in the transition width."
    ),
    # Para 2: Two-phase + TOA
    (
        "Under steady-state two-phase flow, the oil-phase tracer mass flux eliminates the dilution "
        "artifact inherent in concentration-based interpretation and tracks oil production rates "
        "across oil-water ratios from 4:1 to 1:4 (Pearson r = 0.97, RMSD = 8.3%). Simple "
        "time-of-arrival methods are inadequate for this class of BTCs (errors +162% to +685%)."
    ),
    # Para 3: Limitations + outlook
    (
        "The present study is limited to single-interval, laboratory-scale experiments with "
        "dodecane as the model oil. Extending the framework to field conditions requires validation "
        "under multi-interval configurations, crude oil, transient flow, elevated temperature, and "
        "aggressive fluid environments (H₂S, CO₂, high-salinity brines). The "
        "co-precipitation synthesis accommodates a range of transition metals and rare earth elements "
        "for multi-stage tracer coding; combined with the transport analysis framework, this provides "
        "a pathway to per-stage production allocation in multi-fractured horizontal "
        "wells—requiring no downhole tools and only a single shut-in."
    ),
]

# ====================================================================
# APPLY EDITS
# ====================================================================

# --- Abstract ---
print("Editing Abstract...")
# Abstract heading is para 1, body is para 2 (and sometimes para 3)
write_para(doc.paragraphs[2], REVISED_ABSTRACT)
# If para 3 is also part of abstract body, clear it
if doc.paragraphs[3].text.strip() and 'Keywords' not in doc.paragraphs[3].text:
    clear_para(doc.paragraphs[3])
print("  Done.")

# --- Introduction ---
print("Editing Introduction...")
intro_body_start = 7   # First body paragraph after "1. Introduction" heading (para 6)
intro_body_end = 15    # "2. Experimental Section" heading is para 15
old_intro_count = intro_body_end - intro_body_start  # 8 paragraphs

for i, new_text in enumerate(INTRO_PARAS):
    para_idx = intro_body_start + i
    write_para(doc.paragraphs[para_idx], new_text)
    print(f"  Replaced para {para_idx}")

# Clear any remaining old introduction paragraphs
for i in range(intro_body_start + len(INTRO_PARAS), intro_body_end):
    clear_para(doc.paragraphs[i])
    print(f"  Cleared para {i}")
print("  Done.")

# --- Section 3.3: Add numbered sub-headings ---
print("Editing Section 3.3 sub-headings...")
# Find sub-heading paragraphs by their text
sub_headings = {
    "Physical origin of the two-component BTC": "3.3.1 Model formulation",
    "Model selection and physical validation": "3.3.2 Model selection",
    # "Physical consistency" actually comes after "Model selection and physical validation"
    # Need to handle: the original combines model selection + physical validation into one subsection
    "Time-of-Arrival Comparison": "3.3.4 Comparison with time-of-arrival methods",
    "Signal decomposition and robustness": "3.3.5 Signal decomposition and robustness",
}

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()

    # "Physical origin of the two-component BTC." -> "3.3.1 Model formulation"
    if text.startswith("Physical origin of the two-component BTC"):
        clear_para(p)
        p.add_run("3.3.1 Model formulation").bold = True
        print(f"  Para {i}: 'Physical origin...' -> '3.3.1 Model formulation'")

    # "Governing equation and analytical solutions." -> remove as separate heading,
    # merge into 3.3.1 by changing to bold run indicating sub-topic
    elif text.startswith("Governing equation and analytical solutions"):
        clear_para(p)
        p.add_run("Governing equations.").bold = True
        print(f"  Para {i}: Merged into 3.3.1")

    # "Two-component piecewise model with smooth transition." -> remove as separate heading
    elif text.startswith("Two-component piecewise model with smooth transition"):
        clear_para(p)
        p.add_run("Two-component model.").bold = True
        print(f"  Para {i}: Merged into 3.3.1")

    # "Model selection and physical validation." -> "3.3.2 Model selection"
    elif text.startswith("Model selection and physical validation"):
        clear_para(p)
        p.add_run("3.3.2 Model selection").bold = True
        print(f"  Para {i}: -> '3.3.2 Model selection'")

    # Add "3.3.3 Physical consistency" before the physical self-calibration paragraph
    # Find: "Beyond statistical fit quality, the model is validated through physical self-calibration"
    elif text.startswith("Beyond statistical fit quality"):
        # Check if the previous paragraph is already our new heading
        prev_text = doc.paragraphs[i-1].text.strip()
        if not prev_text.startswith("3.3.3"):
            # Insert a heading before this paragraph
            heading_p = doc.paragraphs[i-1]
            new_p_elem = doc.add_paragraph()
            new_p_elem.add_run("3.3.3 Physical consistency").bold = True
            heading_p._element.addnext(new_p_elem._element)
            print(f"  Inserted '3.3.3 Physical consistency' before para {i}")

    # "Time-of-Arrival Comparison" -> "3.3.4 Comparison with time-of-arrival methods"
    elif text.startswith("Time-of-Arrival Comparison"):
        clear_para(p)
        p.add_run("3.3.4 Comparison with time-of-arrival methods").bold = True
        print(f"  Para {i}: -> '3.3.4 Comparison with time-of-arrival methods'")

    # "Signal decomposition and robustness" -> "3.3.5 Signal decomposition and robustness"
    elif text.startswith("Signal decomposition and robustness"):
        clear_para(p)
        p.add_run("3.3.5 Signal decomposition and robustness").bold = True
        print(f"  Para {i}: -> '3.3.5 Signal decomposition and robustness'")

    # Find the "parameter estimation" sentence and add fitting method note
    elif text.startswith("Equation (1) contains seven parameters"):
        # Append fitting method info
        original = p.text
        clear_para(p)
        p.add_run(original + " Parameters were estimated by differential evolution followed by L-BFGS-B refinement (SciPy 1.11), minimizing the sum of squared residuals. The same protocol was applied to all candidate models.")
        print(f"  Para {i}: Added fitting method description")

print("  Done.")

# --- Conclusions ---
print("Editing Conclusions...")
conc_body_start = 114  # First body paragraph after "4. Conclusions" heading (para 113)
conc_body_end = 118    # "References" heading is para 118

for i, new_text in enumerate(CONCLUSION_PARAS):
    para_idx = conc_body_start + i
    write_para(doc.paragraphs[para_idx], new_text)
    print(f"  Replaced para {para_idx}")

# Clear any remaining old conclusion paragraphs
for i in range(conc_body_start + len(CONCLUSION_PARAS), conc_body_end):
    clear_para(doc.paragraphs[i])
    print(f"  Cleared para {i}")
print("  Done.")

# --- Save ---
print(f"\nSaving to {OUTPUT}...")
doc.save(OUTPUT)
print("Done! Revised manuscript saved.")
