#!/usr/bin/env python3
"""修改专利申请文件_正式格式.docx"""

import docx

doc = docx.Document(r'荧光压裂液/专利申请文件_正式格式.docx')
paras = doc.paragraphs

# Helper: replace paragraph text, preserving first-run formatting
def set_para_text(para, new_text):
    """Replace all text in a paragraph, keeping formatting from first run if any"""
    if para.runs:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = new_text
    else:
        run = para.add_run(new_text)

# Helper: replace substring within a paragraph (across runs)
def replace_in_para(para, old, new):
    """Replace old text with new text in a paragraph, working across runs"""
    full = para.text
    if old in full:
        new_full = full.replace(old, new)
        for i, run in enumerate(para.runs):
            if i == 0:
                run.text = new_full
            else:
                run.text = ''


# ============================================================
# STEP 1: DELETE paragraphs (bottom to top to preserve indices)
# ============================================================

# These are the paragraph indices to DELETE (original indices before any changes)
# P149 = old claim 14
# P138 = old claim 9 (merged into new claim 7)
# P31  = description: non-ionic surfactant
# P30  = description: chelating agent
# P29  = description: dispersant system intro

to_delete = sorted([149, 138, 31, 30, 29], reverse=True)
for idx in to_delete:
    p = paras[idx]
    p._element.getparent().remove(p._element)
    print(f"Deleted paragraph {idx}")

print(f"\n--- Deletions complete ---\n")

# Re-fetch paragraphs since the document XML has changed
paras = doc.paragraphs

# ============================================================
# STEP 2: MODIFY description paragraphs (search by content)
# ============================================================

# --- P27: Silane list in description - remove KH570 and vinyl silane ---
for i, p in enumerate(paras):
    if 'KH570' in p.text and '乙烯基三甲氧基硅烷' in p.text and 'APTMS' in p.text:
        replace_in_para(p,
            '所述硅烷偶联剂选自3-氨基丙基三乙氧基硅烷（KH550/APTES）、3-氨基丙基三甲氧基硅烷（APTMS）、γ-甲基丙烯酰氧基丙基三甲氧基硅烷（KH570）或乙烯基三甲氧基硅烷中的一种或多种，优选含氨基的硅烷偶联剂，更优选KH550',
            '所述硅烷偶联剂为含氨基的硅烷偶联剂，选自3-氨基丙基三乙氧基硅烷（KH550/APTES）、3-氨基丙基三甲氧基硅烷（APTMS）或N-（2-氨基乙基）-3-氨基丙基三甲氧基硅烷（KH792）中的一种或多种，优选KH550')
        print(f"P{i}: Description silane list updated")
        break

# --- P36 area: Preparation method step (3) - add dosage amounts ---
for i, p in enumerate(paras):
    if '在使用时将步骤（2）得到的改性稀土铝酸盐荧光粉与螯合剂和非离子表面活性剂按所述比例混合' in p.text:
        replace_in_para(p,
            '在使用时将步骤（2）得到的改性稀土铝酸盐荧光粉与螯合剂和非离子表面活性剂按所述比例混合，或者将螯合剂和非离子表面活性剂预先溶解于分散介质中，以便在配制荧光母液时实现协同分散效果。',
            '在使用时将步骤（2）得到的改性稀土铝酸盐荧光粉与螯合剂和非离子表面活性剂混合——螯合剂（选自柠檬酸、EDTA或其钠盐、酒石酸中的一种或多种，优选柠檬酸）用量为荧光粉基体质量的0.05~0.5 wt%，非离子表面活性剂（选自烷基酚聚氧乙烯醚系列、脂肪醇聚氧乙烯醚系列或聚山梨酯系列中的一种或多种，优选Triton X-100）用量为荧光粉基体质量的0.01~0.2 wt%；或者将螯合剂和非离子表面活性剂预先溶解于分散介质中，以便在配制荧光母液时实现协同分散效果。')
        print(f"P{i}: Preparation method step (3) updated with dosages")
        break

# --- PEG desorption mechanism: change primary driver to thermal dissolution ---

# P51 (now shifted): 关井破胶阶段 detailed mechanism
for i, p in enumerate(paras):
    if '硫酸根自由基（SO' in p.text and 'PEG醚键' in p.text:
        old_text = p.text
        replace_in_para(p,
            '在此阶段，交联冻胶中的有机硼交联键在破胶剂（过硫酸铵）作用下发生氧化断裂，聚合物网络解体；同时，改性荧光粉外层的PEG物理屏蔽层在破胶剂氧化环境与储层温度的协同作用下，经历以下三个过程而逐步脱附降解：（i）过硫酸铵热分解产生的硫酸根自由基（SO₄•⁻）攻击PEG醚键（C-O-C），引发氧化链断裂，生成低分子量PEG碎片和含氧低聚物；（ii）低分子量产物的水溶性显著增大，倾向于从荧光粉表面向水相溶解释放；（iii）破胶剂分解产生的酸性微环境（局部pH下降）促使PEG醚氧原子的质子化程度增加，削弱PEG与内层硅烷偶联剂的物理吸附作用。',
            '在此阶段，交联冻胶中的有机硼交联键在破胶剂（过硫酸铵）作用下发生氧化断裂，聚合物网络解体；同时，改性荧光粉外层的PEG物理屏蔽层在储层高温（60~150°C）与破胶剂氧化环境的协同作用下逐步脱附。其脱附机制以高温驱动的物理溶解为主、氧化降解为辅：（i）PEG与水分子之间的氢键作用随温度升高而减弱，同时PEG分子链热运动加剧，物理吸附于硅烷偶联剂层表面的PEG链段在高温地层水中的溶解度显著增大，倾向于从颗粒表面向水相溶解释放——这是PEG外层脱附的主要驱动力；（ii）作为辅助机制，过硫酸铵热分解产生的硫酸根自由基（SO₄•⁻）可攻击PEG醚键（C-O-C），引发部分氧化链断裂，生成低分子量PEG碎片和含氧低聚物，进一步促进溶解脱附。')
        print(f"P{i}: PEG desorption mechanism updated (thermal dissolution primary)")
        break

# Beneficial effect point (2): update mechanism reference
for i, p in enumerate(paras):
    if '在关井破胶阶段，PEG外层作为牺牲响应层，在过硫酸铵氧化环境与储层温度的协同作用下发生脱附降解' in p.text:
        replace_in_para(p,
            '在关井破胶阶段，PEG外层作为牺牲响应层，在过硫酸铵氧化环境与储层温度的协同作用下发生脱附降解',
            '在关井破胶阶段，PEG外层作为牺牲响应层，在储层高温驱动下发生物理溶解脱附（辅以过硫酸铵氧化降解），')
        print(f"P{i}: Beneficial effect point (2) updated")
        break

# Beneficial effect point (3): update mechanism reference
for i, p in enumerate(paras):
    if 'PEG外层在氧化降解、溶解度增大和酸性微环境的三重协同作用下发生脱附' in p.text:
        replace_in_para(p,
            'PEG外层在氧化降解、溶解度增大和酸性微环境的三重协同作用下发生脱附',
            'PEG外层在高温物理溶解（主）和氧化降解（辅）的协同作用下发生脱附')
        print(f"P{i}: Beneficial effect point (3) updated")
        break

print(f"\n--- Description modifications complete ---\n")

# ============================================================
# STEP 3: MODIFY claims (search by content, since indices shifted)
# ============================================================

# --- Claim 1 part 3: limit silane to amino-containing ---
for i, p in enumerate(paras):
    if '所述内层为硅烷偶联剂化学键合层，所述硅烷偶联剂通过Si-O-Al共价键锚固于荧光粉基体表面；所述外层为聚乙二醇' in p.text:
        replace_in_para(p,
            '所述双层改性层包括内层和外层，所述内层为硅烷偶联剂化学键合层，所述硅烷偶联剂通过Si-O-Al共价键锚固于荧光粉基体表面；所述外层为聚乙二醇（PEG）物理屏蔽层，所述聚乙二醇通过物理吸附沉积于内层硅烷偶联剂表面。',
            '所述双层改性层包括内层和外层，所述内层为含氨基的硅烷偶联剂化学键合层，所述含氨基的硅烷偶联剂通过Si-O-Al共价键锚固于荧光粉基体表面，其末端氨基在PEG外层脱附后暴露并提供与岩石表面锚定的活性位点；所述外层为聚乙二醇（PEG）物理屏蔽层，所述聚乙二醇通过物理吸附沉积于内层硅烷偶联剂表面。')
        print(f"P{i}: Claim 1 updated (amino silane limitation)")
        break

# --- Claim 2: remove KH570 and vinyl silane ---
for i, p in enumerate(paras):
    if '2. 根据权利要求1所述的改性稀土铝酸盐荧光粉' in p.text and 'KH570' in p.text:
        replace_in_para(p,
            '所述硅烷偶联剂选自3-氨基丙基三乙氧基硅烷（KH550）、3-氨基丙基三甲氧基硅烷（APTMS）、γ-甲基丙烯酰氧基丙基三甲氧基硅烷（KH570）或乙烯基三甲氧基硅烷中的一种或多种，优选含氨基的硅烷偶联剂',
            '所述含氨基的硅烷偶联剂选自3-氨基丙基三乙氧基硅烷（KH550）、3-氨基丙基三甲氧基硅烷（APTMS）或N-（2-氨基乙基）-3-氨基丙基三甲氧基硅烷（KH792）中的一种或多种，优选KH550')
        print(f"P{i}: Claim 2 updated (KH570 & vinyl silane removed)")
        break

# --- Old Claim 6 -> New Claim 5 ---
for i, p in enumerate(paras):
    if '6. 一种权利要求1至5中任一项所述改性稀土铝酸盐荧光粉的制备方法' in p.text:
        set_para_text(p, '5. 一种权利要求1至4中任一项所述改性稀土铝酸盐荧光粉的制备方法，其特征在于，包括以下步骤：')
        print(f"P{i}: Old claim 6 -> New claim 5")
        break

# --- Old Claim 7 -> New Claim 6 ---
for i, p in enumerate(paras):
    if '7. 根据权利要求6所述的制备方法' in p.text:
        replace_in_para(p, '7. 根据权利要求6所述的制备方法', '6. 根据权利要求5所述的制备方法')
        print(f"P{i}: Old claim 7 -> New claim 6")
        break

# --- Old Claim 8 -> New Claim 7 (title only, content already updated below) ---
for i, p in enumerate(paras):
    if '8. 一种荧光压裂液体系，其特征在于，包括以下组分：' in p.text:
        set_para_text(p, '7. 一种荧光压裂液体系，其特征在于，包括以下组分：')
        print(f"P{i}: Old claim 8 -> New claim 7 (title)")
        break

# --- New Claim 7 (a): update reference from 1-5 to 1-4 ---
for i, p in enumerate(paras):
    txt = p.text.strip()
    if txt.startswith('（a）') and '1至5' in txt:
        set_para_text(p, '（a）权利要求1至4中任一项所述的改性稀土铝酸盐荧光粉；')
        print(f"P{i}: Claim 7(a) updated")
        break

# --- New Claim 7 (b): add specifics from old claim 9 ---
for i, p in enumerate(paras):
    if p.text.strip() == '（b）稠化剂；':
        set_para_text(p, '（b）稠化剂，选自羟丙基胍胶（HPG）、瓜尔胶、羧甲基羟丙基胍胶（CMHPG）中的一种或多种，优选HPG，在压裂液终液中的浓度为0.3~1.0 wt%；')
        print(f"P{i}: Claim 7(b) updated with specifics")
        break

# --- New Claim 7 (c): add specifics ---
for i, p in enumerate(paras):
    if p.text.strip() == '（c）交联剂；':
        set_para_text(p, '（c）交联剂，选自有机硼交联剂、有机锆交联剂或有机钛交联剂中的一种或多种，优选有机硼延缓交联剂，在压裂液终液中的浓度为0.1~0.5 vol%；')
        print(f"P{i}: Claim 7(c) updated with specifics")
        break

# --- New Claim 7 (d): add specifics ---
for i, p in enumerate(paras):
    if p.text.strip() == '（d）破胶剂；':
        set_para_text(p, '（d）破胶剂，选自过硫酸铵、过硫酸钾、胶囊包裹过硫酸盐或酶破胶剂中的一种或多种，优选过硫酸铵，在压裂液终液中的浓度为0.02~0.3 wt%；')
        print(f"P{i}: Claim 7(d) updated with specifics")
        break

# --- Old Claim 10 -> New Claim 8 ---
for i, p in enumerate(paras):
    if '10. 根据权利要求8或9所述的荧光压裂液体系' in p.text:
        set_para_text(p, '8. 根据权利要求7所述的荧光压裂液体系，其特征在于，采用母液预配与在线稀释相结合的方式制备：')
        print(f"P{i}: Old claim 10 -> New claim 8")
        break

# --- Old Claim 11 -> New Claim 9 (title) ---
for i, p in enumerate(paras):
    if '11. 一种利用权利要求8至10中任一项所述荧光压裂液体系的压裂裂缝荧光示踪方法' in p.text:
        set_para_text(p, '9. 一种利用权利要求7至8中任一项所述荧光压裂液体系的压裂裂缝荧光示踪方法，其特征在于，包括以下步骤：')
        print(f"P{i}: Old claim 11 -> New claim 9 (title)")
        break

# --- New Claim 9 (old 11) 关井破胶步骤: update PEG mechanism ---
for i, p in enumerate(paras):
    if '关井破胶步骤：停泵关井，在储层温度60~150' in p.text and '同时改性荧光粉外层的聚乙二醇在破胶剂氧化环境与储层温度的协同作用下脱附' in p.text:
        replace_in_para(p,
            '同时改性荧光粉外层的聚乙二醇在破胶剂氧化环境与储层温度的协同作用下脱附，暴露内层硅烷偶联剂的活性官能团',
            '同时改性荧光粉外层的聚乙二醇在储层高温驱动下发生物理溶解脱附（辅以破胶剂氧化降解），暴露内层含氨基硅烷偶联剂的活性氨基官能团')
        print(f"P{i}: Claim 9 (old 11) mechanism updated")
        break

# --- Old Claim 12 -> New Claim 10 ---
for i, p in enumerate(paras):
    if '12. 根据权利要求11所述的方法' in p.text:
        replace_in_para(p, '12. 根据权利要求11所述的方法', '10. 根据权利要求9所述的方法')
        print(f"P{i}: Old claim 12 -> New claim 10")
        break

# --- Old Claim 13 -> New Claim 11 ---
for i, p in enumerate(paras):
    if '13. 根据权利要求11所述的方法' in p.text:
        replace_in_para(p, '13. 根据权利要求11所述的方法', '11. 根据权利要求9所述的方法')
        print(f"P{i}: Old claim 13 -> New claim 11")
        break

print(f"\n--- Claims modifications complete ---\n")

# ============================================================
# STEP 4: Fix figure description duplicates
# ============================================================
for i, p in enumerate(paras):
    txt = p.text.strip()
    if txt == '图1 改性稀土铝酸盐荧光粉的结构示意图':
        set_para_text(p, '')
        print(f"P{i}: Removed duplicate Fig 1 title")
        break

for i, p in enumerate(paras):
    txt = p.text.strip()
    if txt == '图2 荧光压裂液体系现场施工工艺流程图':
        set_para_text(p, '')
        print(f"P{i}: Removed duplicate Fig 2 title")
        break

for i, p in enumerate(paras):
    txt = p.text.strip()
    if txt == '图3 压裂裂缝荧光示踪方法的流程框图':
        set_para_text(p, '')
        print(f"P{i}: Removed duplicate Fig 3 title")
        break

# ============================================================
# STEP 5: Update 摘要 - PEG mechanism
# ============================================================
for i, p in enumerate(paras):
    if '关井破胶阶段PEG在破胶剂氧化环境与储层温度作用下脱附降解' in p.text:
        replace_in_para(p,
            '关井破胶阶段PEG在破胶剂氧化环境与储层温度作用下脱附降解',
            '关井破胶阶段PEG在储层高温作用下发生物理溶解脱附（辅以破胶剂氧化降解）')
        print(f"P{i}: Abstract PEG mechanism updated")
        break

# ============================================================
# SAVE
# ============================================================
output_path = r'荧光压裂液/专利申请文件_正式格式_修改版.docx'
doc.save(output_path)
print(f"\n{'='*60}")
print(f"All modifications complete. Saved to: {output_path}")
print(f"{'='*60}")