#!/usr/bin/env python3
"""Rewrite Section 3.7: model as centerpiece, anchored to ESP-T material properties."""
from docx import Document
import time
doc = Document('四氧化三铁环氧树脂拟合/ESP-T_final_1781618084.docx')

# [134] — Derivation: every step references ESP-T
p = doc.paragraphs[134]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "The characterization in Sections 3.1-3.6 establishes that ESP-T "
    "is oleophilic, thermally stable, and releases tracer through a "
    "diffusion-controlled mechanism that persists for weeks at reservoir "
    "temperature. These properties make ESP-T suitable as a monitoring "
    "tool, but they do not by themselves answer the question that "
    "motivated its development: how much oil is each fractured interval "
    "producing? Answering that question requires a model that translates "
    "the tracer concentration history recorded at the wellhead into a "
    "production rate—a model grounded in the specific release behavior "
    "of ESP-T."
    "\n\n"
    "Two processes, both rooted in how ESP-T functions, contribute to "
    "the breakthrough curve. During the shut-in period, tracer released "
    "from ESP-T accumulates in the oil phase within the proppant pack. "
    "When the well opens, this accumulated tracer is swept out as a "
    "concentrated slug, producing a pulse that rises and falls as it "
    "passes the sampling point. Simultaneously, residual tracer within "
    "the epoxy matrix continues to diffuse outward—the sustained release "
    "quantified by the K-P kinetics in Section 3.6—feeding a persistent "
    "low-level signal into the produced oil long after the main slug "
    "has passed. These two contributions overlap in time. One cannot "
    "inspect a concentration curve and separate the shut-in slug from "
    "the sustained tail by eye."
    "\n\n"
    "We resolve this overlap by modeling the transport explicitly. The "
    "one-dimensional advection-dispersion equation, dC/dt + v dC/dx = "
    "D d2C/dx2, governs how tracer concentration evolves along the "
    "production tubing. The flow velocity v = 4Q/(pi d2) connects the "
    "model directly to the production rate Q—this is the quantity we "
    "ultimately wish to recover. The dispersion coefficient D = alpha v "
    "accounts for mixing along the flow path. For the shut-in slug we "
    "use the classical instantaneous-injection solution [25], which "
    "gives a Gaussian pulse whose amplitude reflects the tracer mass "
    "accumulated during shut-in—a quantity that depends on the shut-in "
    "duration and on ESP-T's oleophilic surface (Section 3.3), which "
    "ensures that released tracer partitions into the oil phase rather "
    "than being lost to water. For the sustained release we use the "
    "continuous-source solution, an erfc function whose characteristic "
    "decay time is governed by the matrix-diffusion kinetics documented "
    "in Section 3.6. We join the two components with a tanh weighting "
    "function, w(t) = 1/2 [1 + tanh((t-t0)/sigma)], producing a smooth "
    "transition from pulse-dominated early times to diffusion-dominated "
    "late times. The full composite model, C(t) = cb + w(t) C_rise + "
    "[1-w(t)] C_fall, is Eq. (2). The parameter sigma is fixed at the "
    "sampling interval (4 min); the remaining six parameters are "
    "discussed below."
)
print('[134] rewritten')

# [148] — Fitting results: connect every number to ESP-T
p = doc.paragraphs[148]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "Fitted to the single-phase breakthrough curve obtained with ESP-T "
    "(Figure 3-8b, Table 3-3), the model achieves R2 = 0.9939 with "
    "RMSE = 0.0210. The quality of the fit, however, is less important "
    "than what the fitted parameters reveal. The model recovers an "
    "effective flow rate Q = 0.46 mL/min—within 8% of the pump setting "
    "of 0.50 mL/min—without being told the true value. In field terms, "
    "this means the model can estimate the production rate of a labeled "
    "interval from the tracer signal alone. The mean residence time "
    "MRT = 37.4 min matches the convective travel time x/v = 38.6 min "
    "calculated from the known tubing length and the fitted velocity, "
    "confirming that the model's time scale is physically correct. The "
    "Peclet number Pe = x/alpha = 0.934 sits near unity, the regime "
    "expected when the tracer source is gradual rather than "
    "instantaneous—precisely what one would predict from the non-Fickian "
    "kinetics of ESP-T documented in Section 3.6."
    "\n\n"
    "Integrating the two components separately reveals that 47% of the "
    "total integrated signal originates from the erfc tail. In practical "
    "terms, nearly half the tracer detected at the wellhead during the "
    "sampling period comes not from the shut-in slug but from the "
    "sustained, diffusion-controlled release that ESP-T's epoxy matrix "
    "provides."
)
print('[148] rewritten')

# [152] — Physical meaning: anchored to ESP-T throughout
p = doc.paragraphs[152]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "Several features of these results reinforce each other. The flow "
    "rate Q, residence time MRT, and Peclet number converge from three "
    "independent directions onto values consistent with the known "
    "experimental inputs and the independently measured material "
    "behavior of ESP-T. This convergence is not guaranteed by curve-"
    "fitting; it indicates that the model is recovering physically "
    "meaningful quantities rather than accommodating noise. The 47% "
    "tail contribution carries a direct operational consequence. Because "
    "ESP-T sustains release through matrix diffusion rather than through "
    "coating dissolution, a single shut-in period suffices to generate "
    "a persistent monitoring signal. There is no need to interrupt "
    "production for repeated shut-in cycles, and no risk that the "
    "tracer signal will vanish when a coating layer is exhausted. The "
    "tail-plateau concentration cb reflects the steady-state diffusion "
    "flux maintained by ESP-T at the experimental temperature. The K-P "
    "rate constant K rises from 0.055 at 30 degC to 0.196 at 120 degC "
    "(Table 3-2), directly implying that cb will be higher in hotter "
    "wells—a testable prediction that connects the batch release "
    "experiments of Section 3.6 to the transport model."
    "\n\n"
    "These characteristics distinguish the present approach from "
    "earlier work. Liu et al. [19] applied ADE-based Gaussian fitting "
    "to tracer flowback curves, extracting fracture geometry rather "
    "than production rate. Tian et al. [20] used tracer mass balance "
    "for production allocation without exploiting breakthrough curve "
    "shape. Li et al. [26] employed discrete fracture models for "
    "tracer transport simulation while remaining focused on fracture "
    "characterization. In each case the transport step serves reservoir "
    "description. Here, because ESP-T's release mechanism produces "
    "physically distinct pulse and tail components, the transport itself "
    "becomes the source of a production metric. The material and the "
    "model are not independent contributions; the model works because "
    "ESP-T's sustained, matrix-diffusion-controlled release creates "
    "the two-component signal structure that the model decomposes."
)
print('[152] rewritten')

# [154] — Two-phase bridge
p = doc.paragraphs[154]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "The single-phase analysis validates the model under conditions "
    "where only oil flows through the proppant pack. In a producing "
    "well, however, formation water eventually breaks through and the "
    "produced fluid becomes an oil-water mixture. The question is "
    "whether the analysis remains informative when both phases are "
    "present. We tested this using steady-state core displacement "
    "experiments at three oil-water ratios (OWR = 4:1, 1:1, 1:4) "
    "and four total flow rates (0.1-0.4 mL/min). Direct observation "
    "of the effluent at OWR = 1:1 showed a higher oil fraction early "
    "in the displacement, consistent with the oil-permeable character "
    "of ESP-T (Section 3.5)."
)
print('[154] rewritten')

# [156] — Concentration data
p = doc.paragraphs[156]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "Figure 3-9(a) shows the tracer concentration as a function of "
    "total flow rate for each OWR. Concentration drops as flow rate "
    "rises—a straightforward dilution effect—and shows negligible "
    "dependence on OWR. This is physically consistent with ESP-T's "
    "wetting behavior (Section 3.5): oil and water access the "
    "proppant surface in proportion to their volume fractions. The "
    "practical problem is that concentration alone is a poor proxy "
    "for the oil production rate, because it confounds the tracer "
    "release rate with the total flow-driven dilution."
)
print('[156] rewritten')

# [158] — FO definition + results
p = doc.paragraphs[158]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "We resolve this by defining the tracer flux FO as the mass of "
    "tracer passing the sampling point per unit time—the product of "
    "the oil-phase tracer concentration and the oil volumetric flow "
    "rate. FO equals the release rate from the ESP-T pack [8] and is "
    "therefore independent of dilution. Figure 3-9(b) confirms that "
    "FO behaves as expected from ESP-T's oleophilic surface: it "
    "increases with OWR as more oil contacts the proppant, but remains "
    "constant when the total flow rate changes at fixed OWR. ESP-T "
    "releases tracer at a rate set by the oil-wetted area (Section "
    "3.3), not by the flow velocity through the pack."
)
print('[158] rewritten')

# [160] — Validation + limitations
p = doc.paragraphs[160]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "Figure 3-9(c) normalizes FO from the two-phase experiments by "
    "the steady-state FO obtained during single-phase oil displacement "
    "(3.187 micro-g/min) and compares the result with the known oil "
    "flow rate at each OWR. The agreement is close, demonstrating "
    "that the oil production rate of a labeled interval can be "
    "recovered from FO when the total flow rate is known, even with "
    "a water phase present. This result closes the loop from material "
    "design (an oleophilic, diffusion-controlled-release proppant) "
    "through transport modeling (a two-component ADE decomposition) "
    "to a quantitative production metric."
    "\n\n"
    "Several boundaries should be noted. The model assumes one "
    "fractured interval with a uniform proppant pack; real wells "
    "produce from multiple intervals whose tracer signals overlap. "
    "The FO calibration derives from steady-state experiments; "
    "transient conditions during start-up, shut-in, or rapid drawdown "
    "may cause deviations. Dodecane served as the model oil; crude "
    "oil components, particularly asphaltenes, could alter ESP-T's "
    "wetting and diffusion behavior. The long-term chemical stability "
    "of the epoxy matrix and the stearic acid coating above 120 degC, "
    "or in the presence of H2S, CO2, and high-salinity brines, "
    "remains to be evaluated."
)
print('[160] rewritten')

out = f'四氧化三铁环氧树脂拟合/ESP-T_final_{int(time.time())}.docx'
doc.save(out)
print(f'Saved: {out}')