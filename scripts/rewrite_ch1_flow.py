#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Rewrite Chapter 1 to be more fluid and engaging.
- Replace laundry-list literature review with thematic synthesis
- Fix 1.2.4 structural placement
- Improve narrative flow of 1.1
"""
import sys
from docx import Document

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

DOC = r'荧光压裂液/论文初稿_v2.docx'
doc = Document(DOC)

def find_para(doc, snippet, skip=0):
    found = 0
    for i, para in enumerate(doc.paragraphs):
        if snippet in para.text:
            if found >= skip:
                return i, para
            found += 1
    return None, None

def replace_para_text(para, new_text):
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)

def add_para_after(doc, target_para, text):
    new_para = doc.add_paragraph(text)
    target_para._element.addnext(new_para._element)
    return new_para

def delete_para(para):
    para._element.getparent().remove(para._element)


# ================================================================
# 1. REWRITE 1.1 — add narrative arc
# ================================================================
print("1. Rewriting 1.1 to build narrative tension...")

# Keep P36 (opening paragraph) but make it flow better
_, p36 = find_para(doc, '水力压裂是实现页岩油气、致密砂岩气及煤层气等非常规油气资源商业化开发的核心技术')
if p36:
    new_p36 = (
        '水力压裂技术诞生于1949年，迄今全球已完成逾250万次压裂施工，为美国累计贡献了约30%的'
        '新增可采石油储量和约90%的天然气储量[1]。这项技术的核心逻辑简洁而有力：向储层注入高压流体，'
        '制造人工裂缝网络，将油气从纳米-微米级基质孔隙中"释放"出来。裂缝延伸到哪里，'
        '产量就从哪里来——储层改造体积（SRV）直接决定了非常规储层的产能上限[2-4]。'
        '然而，裂缝在地下的实际延伸范围究竟如何？这一看似基础的问题，至今仍未得到圆满回答。'
    )
    replace_para_text(p36, new_p36)
    print('  Rewrote P36 (opening)')

# Rewrite P37 — the diagnostic technology problem
_, p37 = find_para(doc, '当前应用于裂缝监测的技术手段主要包括微地震监测')
if p37:
    new_p37 = (
        '回答这一问题的难度在于：裂缝深埋于数千米之下，无法被直接观测。业界发展了一系列'
        '间接监测手段——微地震监测通过接收岩石破裂的声发射信号反演裂缝位置[9-11]，'
        '分布式光纤传感（DAS/DTS）通过光纤的应变和温度响应推断裂缝扩展动态[12-14]，'
        '化学示踪剂和放射性示踪剂通过返排液的浓度曲线回溯裂缝贡献[15-17]。'
        '这些技术各有所长，但共享一个方法论局限：它们测量的是裂缝的"影子"而非裂缝本身——'
        '微地震监测到的是岩石破裂事件，而非压裂液的实际波及范围；DAS感知的是光纤附近的'
        '应变变化，而非裂缝壁面的空间位置；示踪剂的返排信号反映的是裂缝网络与井筒的'
        '连通性，而非裂缝几何的直接影像。正如Warpinski（1996）在其经典综述中所言，'
        '裂缝的精确诊断是改善增产效果和降低成本的关键前提[7]；而de Melo和Carballo Cabrera'
        '（2025）对多种裂缝成像技术的系统评价则给出了冷静的结论：没有任何单一技术能提供'
        '完整的裂缝几何信息[8]。'
    )
    replace_para_text(p37, new_p37)
    print('  Rewrote P37 (diagnostic problem)')

# P38 — fluorescent tracer as a bridge
_, p38 = find_para(doc, '荧光示踪技术为弥补上述"信号-实物"鸿沟提供了可行路径')
if p38:
    new_p38 = (
        '在这个技术背景下，一个朴素的思路浮现出来：如果能让压裂液本身在裂缝壁面上留下'
        '可辨识的标记，那么压后取心——裂缝诊断中最直接、最具说服力的验证手段——'
        '就可以在紫外光下直接"看见"压裂液曾经到达的位置。这就是荧光示踪压裂液的基本构想。'
        '这一构想并非空中楼阁。Ishida研究团队（Takeuchi等，2025）在日本神冈矿山的露头尺度'
        '压裂实验中，以可凝固荧光树脂作为压裂介质，通过同轴套取205 mm岩心并在紫外光下成像，'
        '清晰识别了主裂缝与分支裂缝的延伸轨迹[18-19]。更有意思的是，他们发现部分被荧光树脂'
        '填充的裂缝区域并未产生声发射事件——换句话说，微地震监测"漏看"了一部分裂缝。'
        '这个发现为荧光直接可视化方法的价值提供了最直接的注脚。另一条技术路线来自Guryanov等'
        '（2019）开发的GeoSplit量子点微球系统，通过6种量子点编码实现63种独立标记信号，'
        '利用流式细胞术分析返排液，已实现分段产液剖面的现场监测[20]。'
    )
    replace_para_text(p38, new_p38)
    print('  Rewrote P38 (fluorescent tracer bridge)')

# P39 — the compatibility gap
_, p39 = find_para(doc, '上述已有方案与常规水力压裂工艺之间存在兼容性鸿沟')
if p39:
    new_p39 = (
        '然而，上述两种方案均无法嵌入常规水力压裂的工程流程。Ishida团队的荧光树脂是可凝固体系——'
        '树脂注入后在地下固化成型，然后通过取心观察裂缝形貌。这一方法的科学价值毋庸置疑，'
        '但它本质上是一种岩石力学实验方法：可凝固树脂的固化机制与常规压裂液"先携砂造缝、'
        '再破胶返排"的工艺逻辑根本不同，无法应用于真实的压裂施工。Guryanov团队的量子点微球'
        '虽然与常规压裂液兼容，但其检测方式依赖返排液的间接分析——量子点随压裂液返排至地面'
        '后才被检测，无法在裂缝壁面上留下可供取心验证的实物标记。更为棘手的是，半导体量子点'
        '在过硫酸铵的强氧化破胶环境中存在降解风险——示踪信号可能在破胶阶段就已衰减甚至消失。'
        '两条技术路线分别解决了问题的一半：一个做到了"可视化"但不兼容压裂工艺，'
        '另一个兼容了压裂工艺但做不到"壁面标记"。核心工程问题仍然悬置：'
        '如何在一个常规胍胶压裂液体系内，将一种化学稳定、可随携砂液泵注、且能在裂缝壁面上'
        '牢固锚定的荧光示踪材料输送至裂缝的各部位？'
    )
    replace_para_text(p39, new_p39)
    print('  Rewrote P39 (compatibility gap)')

# P40 — this study's approach
_, p40 = find_para(doc, '本研究以SrAl₂O₄:Eu²⁺,Dy³⁺无机长余辉荧光粉为示踪介质')
if p40:
    new_p40 = (
        '本研究给出的答案是：以无机稀土铝酸盐长余辉荧光粉SrAl₂O₄:Eu²⁺,Dy³⁺为示踪介质，'
        '以羟丙基胍胶（HPG）冻胶为压裂液载体，构建一种与现有压裂工艺完全兼容的荧光示踪方案。'
        '选择这种无机荧光粉的技术逻辑基于一个简单但关键的判断：其Eu²⁺发光中心受SrAl₂O₄刚性'
        '晶格的严密保护，对过硫酸铵氧化破胶剂具有本征的化学惰性——这意味着示踪信号可以在'
        '破胶返排全过程中保持完整，这是有机荧光染料和半导体量子点难以实现的。'
        '同时，该材料的长余辉特性（激发光撤除后持续发光数小时）使紫外检测无需严格的时间窗口控制。'
        '本研究的技术定位是辅助校准工具——不替代微地震或DAS，而是通过压后取心的实物证据，'
        '为这些间接监测方法的反演结果提供地面实况（ground truth）校准。'
    )
    replace_para_text(p40, new_p40)
    print('  Rewrote P40 (this study)')

# Keep P41-P42 (expanded background on China) as is — they're good

# ================================================================
# 2. REWRITE 1.2.1 — from laundry list to thematic synthesis
# ================================================================
print("\n2. Rewriting 1.2.1...")

# Delete old P45-P49 (the laundry list paragraphs)
for snippet in [
    '在DAS监测方面，Molenaar等（2012）',
    '在微地震监测方面，Maxwell等（2002）',
    '在示踪剂方面，Salman等（2014）',
    '在倾斜仪与电磁法方面，Wright等（1998）',
    '在国内研究方面，中国石油集团和成都理工大学'
]:
    _, para = find_para(doc, snippet)
    if para:
        delete_para(para)
        print(f'  Deleted old paragraph: {snippet[:30]}...')

# Now insert new synthesized paragraphs before P50 (the summary paragraph)
_, p_keep = find_para(doc, '共性局限在于：间接反演多解性强')
# We'll insert before the old summary, then rewrite the summary

# Find the 1.2.1 heading
_, p_121h = find_para(doc, '1.2.1  水力压裂裂缝监测技术')

# Insert after heading
new_das = (
    '分布式光纤传感（DAS/DTS）是近十年发展最为迅速的裂缝监测技术。Molenaar等（2012）'
    '首次将DAS应用于井下压裂监测，验证了光纤对各类井下事件——包括射孔、封隔器坐封和'
    '流体注入——的实时追踪能力[21]，标志着DAS从地震勘探领域进入压裂诊断领域。'
    'Sierra等（2008）从温度维度切入，利用DTS的温度暖回效应建立了各射孔簇的流体分配'
    '解释方法[22]。Ugueto等（2016）将DAS与DTS联合应用，通过应变和温度的双重信号'
    '首次以定量数据揭示了多簇射孔中普遍存在的簇间非均匀进液问题[23]——'
    '这一问题在随后的数年中成为压裂优化领域的研究热点。在裂缝几何定量表征方面，'
    'Jin和Roy（2017）的工作具有里程碑意义：他们建立了低频DAS信号（<0.05 Hz）的'
    '应变梯度与裂缝几何参数的定量对应关系[12]，使光纤从"定性感知"走向"定量测量"。'
    'Karrenbach等（2017, 2019）进一步展示了在同一光纤安装中同步采集微地震、应变和'
    '温度三种信号的可行性，为多物理场联合诊断提供了硬件基础[14,24]。'
)
add_para_after(doc, p_121h, new_das)
print('  Added new DAS synthesis paragraph')

new_microseismic = (
    '微地震监测是压裂裂缝诊断中应用历史最长、现场经验最丰富的方法。Maxwell等（2002）'
    '在Barnett页岩中首次利用微地震成像揭示了水力裂缝的复杂网络特征[26]，'
    '改变了人们对"压裂形成对称双翼裂缝"的经典认知。Fisher等（2004）将微地震裂缝测绘'
    '与水平井完井优化相结合，展示了诊断数据指导工程决策的实际价值[3]。'
    'Mayerhofer等（2010）在这一基础上提出了储层改造体积（SRV）概念——以微地震事件云'
    '的三维外包络体积表征裂缝网络的宏观范围[2]，这一概念迅速成为非常规储层压裂效果'
    '评价的行业标准。然而，SRV概念的内在局限也很快暴露：Cipolla和Wallace（2014）指出，'
    'SRV对裂缝面积和导流能力这两个决定产能的关键参数缺乏约束[4]；微地震事件密度与'
    '支撑剂分布之间并无确定的物理对应关系。Cipolla等（2011）发表的微地震解释实用指南[9]'
    '以及其与复杂裂缝扩展模型的整合工作[27]，代表了对微地震方法论本身的批判性反思。'
    'Maxwell（2014）在其SEG专著中系统总结了微地震在裂缝几何反演中的应用与局限[28]。'
)
# Insert after the DAS paragraph we just added
_, p_insert_after = find_para(doc, '为多物理场联合诊断提供了硬件基础')
add_para_after(doc, p_insert_after, new_microseismic)
print('  Added new microseismic synthesis paragraph')

new_tracers = (
    '化学示踪剂技术沿着另一条逻辑展开：不试图"看"裂缝，而是通过注入与返排的物质'
    '守恒关系推断裂缝的有效体积和连通性。Salman等（2014）系统分析了非常规储层化学示踪剂'
    '的返排数据，建立了示踪剂回收率与裂缝有效贡献的定量关联[15]。Tian等（2019）通过'
    '分序注入12种水相和12种油相化学示踪剂，揭示了单段Wolfcamp B压裂中不同示踪剂的'
    '显著混合——证明裂缝网络内的流体交换远超简单的活塞驱替假设[29]。Arshad等（2024）'
    '实现了新型化学示踪剂在裂缝段产能剖面构建方面的首次现场应用[30]。在荧光示踪分支，'
    'Kosynkin和Alaskar（2016）的碳点基纳米示踪剂（A-Dots）在500 m井距的碳酸盐岩油藏中'
    '取得了令人瞩目的成果——约1吨纳米颗粒注入后，约10个月后在产油井中检测到了示踪信号[31]。'
    'Hu等（2019）进一步展示了碳量子点作为储层示踪剂的潜力[32]。'
)
_, p_insert_after = find_para(doc, '对微地震方法论本身的批判性反思')
add_para_after(doc, p_insert_after, new_tracers)
print('  Added new tracer synthesis paragraph')

new_tiltmeter = (
    '倾斜仪和电磁法提供了与前两类方法互补的测量维度。Wright等（1998）将井下倾斜仪'
    '引入裂缝测绘，通过测量压裂诱发的地层微倾斜反演裂缝方位和几何参数[33]。'
    'Warpinski等（2006）发展了微地震-倾斜仪联合反演方法，利用两类数据的互补约束'
    '降低反演的多解性[34]。在电磁法方向，LaBrecque等（2016）开展了导电支撑剂电磁成像'
    '的小规模现场验证[35]；Ahmadian等（2023）进一步展示了导电支撑剂辅助电磁法'
    '对裂缝扩展动态的实时监测潜力[36]。Cipolla和Wright（2002）在其综述中对'
    '上述各类诊断技术进行了系统比较，结论至今仍有参考价值[5]。'
)
_, p_insert_after = find_para(doc, '碳量子点作为储层示踪剂的潜力')
add_para_after(doc, p_insert_after, new_tiltmeter)
print('  Added new tiltmeter/EM synthesis paragraph')

new_domestic = (
    '在国内，裂缝监测技术的研究与应用近年来也取得了显著进展。中国石油集团和成都理工大学'
    '等单位的学者在这一领域开展了大量工作。翁定为等（2024）[65]系统综述了水力压裂裂缝监测'
    '技术的国内外进展，指出微地震监测在中国页岩气开发中虽已实现规模化应用，但仍面临'
    '布井阵列几何分辨率受限和速度模型不确定性的双重制约。邸德家（2025）[66]梳理了'
    '油气井压裂示踪监测技术的现状，认为化学示踪剂的分段监测精度受返排液采样频率和'
    '示踪剂在裂缝网络中的非均匀分配影响，单一示踪剂技术难以独立完成裂缝几何的定量表征。'
    '在DAS监测方面，中国石油大学（北京）和西南石油大学等单位近年来开展了井下光纤传感'
    '的现场试验，但在低频DAS信号裂缝几何反演方面仍处于方法探索阶段。'
)
_, p_insert_after = find_para(doc, '结论至今仍有参考价值')
add_para_after(doc, p_insert_after, new_domestic)
print('  Added new domestic research paragraph')

# Rewrite the summary paragraph (old P50)
_, p_summary = find_para(doc, '共性局限在于：间接反演多解性强')
if p_summary:
    new_summary = (
        '综观上述技术谱系，一个结构性的认知缺口逐渐清晰：所有现有裂缝监测技术本质上都'
        '属于间接测量——它们测量的是裂缝产生的某种物理效应（声发射、应变、温度、示踪剂浓度），'
        '而非裂缝本身的空间坐标。从物理效应反推裂缝几何，必然涉及反演模型和先验假设，'
        '多解性问题无法从根本上消除。更关键的是，压后取心——裂缝诊断谱系中最直接、'
        '最具终极验证力的手段——与上述间接方法之间存在一个"实物证据断层"：'
        '现有技术无法让取出的岩心"开口说话"，无法在岩心尺度上留下任何可辨识的压裂液波及标记。'
        '这构成了本研究的第一个逻辑起点：能否让取心岩样提供裂缝波及范围的直接实物证据？'
    )
    replace_para_text(p_summary, new_summary)
    print('  Rewrote summary paragraph')


# ================================================================
# 3. REWRITE 1.2.2 — two routes with critical comparison
# ================================================================
print("\n3. Rewriting 1.2.2...")

# Delete old P52-P54
for snippet in [
    '在"直接可视化"路线方面，Chen等（2014）',
    '在"返排分析"路线方面，Guryanov等的GeoSplit',
    '上述两条路线的共同局限在于',
    '从示踪材料选择的角度审视上述两条路线'
]:
    _, para = find_para(doc, snippet)
    if para:
        delete_para(para)
        print(f'  Deleted: {snippet[:30]}...')

# Find 1.2.2 heading
_, p_122h = find_para(doc, '1.2.2  荧光示踪裂缝可视化技术')

new_122_open = (
    '荧光示踪为跨越上述"信号-实物"鸿沟提供了具有可行性的技术路径。其核心思想简明直接：'
    '将荧光标记材料引入压裂液，随携砂液运移至裂缝并在壁面附着，压后取心在紫外光下即可'
    '直接观察荧光分布——由此获得压裂液波及范围的实物证据。围绕这一思想，'
    '已有研究形成了"直接可视化"和"返排分析"两条技术路线。'
)
add_para_after(doc, p_122h, new_122_open)
print('  Added 1.2.2 opening')

new_direct_vis = (
    '直接可视化路线追求的目标是让裂缝壁面本身"发光"。Chen等（2014）率先在实验室尺度'
    '采用含荧光涂料的MMA热固性树脂实现了水力裂缝的可视化[19]。这一思路在Takeuchi等（2025）'
    '的露头尺度实验中得到了迄今最系统的验证：在岐阜县神冈矿山的76 mm钻孔中注入可凝固'
    '荧光树脂后，通过205 mm同轴套取岩心并在紫外光下成像，不仅清晰识别了主裂缝与分支'
    '裂缝的三维延伸轨迹，更通过与声发射（AE）事件分布的对比揭示了一个发人深省的发现：'
    '部分荧光渗透区域未产生任何AE事件。这一发现直接证明了间接监测方法确实存在"盲区"——'
    '有些裂缝形成了，但未伴随可被检测到的声发射信号[18]。Flury和Wai（2003）在Reviews of '
    'Geophysics上发表的综述系统总结了荧光染料在地下水流可视化中近百年的应用历史[37]，'
    '为压裂裂缝的荧光示踪提供了坚实的方法学基础。然而，这条路线面临一个根本性的工程兼容'
    '性问题：可凝固树脂的固化是不可逆的——它无法像常规压裂液那样在完成造缝后破胶返排。'
)
add_para_after(doc, p_122h, new_direct_vis)
print('  Added direct visualization paragraph')

new_flowback = (
    '返排分析路线选择了另一条路径：不追求壁面标记，而是通过检测返排液中的示踪信号间接'
    '推断裂缝信息。这实际上是传统化学示踪剂方法的荧光升级版。Guryanov等（2019）的GeoSplit'
    '量子点微球系统是这一路线最具工程成熟度的代表——利用6种量子点编码组合实现63种独立标记，'
    '通过流式细胞术对返排样品进行高通量分析，已成功应用于分段压裂产液剖面监测[20]。'
    '以纳米颗粒为载体的示踪方案同样发展迅速：Kang等（2015）制备了兼具调剖和示踪双重功能'
    '的荧光聚丙烯酰胺微球[38]；Spitzmuller等（2024）开发的荧光染料包埋介孔硅纳米颗粒示踪剂，'
    '经表面修饰后可在160°C、高盐条件下稳定示踪[39]；Hu等（2019）首次展示了碳量子点在砂岩'
    '岩心驱替中类示踪剂的穿透能力[32]。这些方案的共同设计逻辑是"注入-返排-检测"——'
    '示踪材料最终需要返回地面才能被检测，因此在裂缝壁面上的残留反而是不希望发生的"损失"。'
)
# Insert after the direct visualization paragraph
_, p_insert_after = find_para(doc, '为压裂裂缝的荧光示踪提供了坚实的方法学基础')
add_para_after(doc, p_insert_after, new_flowback)
print('  Added flowback analysis paragraph')

new_gap_122 = (
    '上述两条路线各自向前迈出了重要一步，但均未能完全跨越"信号-实物"鸿沟。直接可视化'
    '路线使用的可凝固树脂与常规压裂工艺不兼容，本质上是一种岩石力学实验方法而非压裂工程'
    '技术；返排分析路线的检测依赖返排液间接信号，无法在裂缝壁面上留下可被取心验证的实物'
    '标记，本质上仍是传统示踪剂方法的延伸。从示踪材料设计的角度审视，一条有效的压裂荧光'
    '示踪方案需同时满足四项基本要求——化学稳定性（耐受过硫酸铵氧化破胶环境）、悬浮输送性'
    '（与压裂液流变学匹配，注入阶段不沉降）、壁面锚定性（关井破胶后从分散态切换至锚定态）'
    '和检测便捷性（紫外灯直接观察，无需复杂前处理）。这四项要求构成了本论文荧光示踪材料'
    '筛选和体系设计的评价框架，也将后续的文献综述和实验设计串联为一个有机整体。'
)
_, p_insert_after = find_para(doc, '反而是不希望发生的"损失"')
add_para_after(doc, p_insert_after, new_gap_122)
print('  Added 1.2.2 gap synthesis')


# ================================================================
# 4. ENHANCE 1.2.3 — add critical commentary
# ================================================================
print("\n4. Enhancing 1.2.3 with critical synthesis...")

# The 1.2.3 section (phosphor materials) is already decent but add a synthesis paragraph
_, p_123_last = find_para(doc, '兼具工程和科学价值的问题')
if p_123_last:
    # Replace the ending with a stronger synthesis that connects to this study
    new_123_end = (
        '然而，上述包覆方案均面向涂料、塑料等行业开发，其性能评价标准（如常温水浸48小时'
        '不丧失磷光）与压裂液的工程环境（60~150°C、高矿化度、高剪切、氧化介质）存在根本性差异，'
        '面向压裂液工程的系统验证尚属空白。更为深层的挑战在于：压裂施工对颗粒表面化学提出了'
        '一种独特的功能时序需求——注入期间需要稳定悬浮（颗粒-颗粒排斥主导），关井破胶期间需要'
        '向锚定态切换（颗粒-壁面吸引主导），返排期间需要牢固附着（抵抗流体剪切）。这三种需求'
        '对表面化学的要求各不相同甚至相互矛盾，如何在单一改性方案中协调这种时序功能切换，'
        '是一个兼具工程价值和基础科学意义的问题。这也是本论文第二章表面改性方案设计的核心'
        '科学问题——不仅要对荧光粉进行"保护"，更要在保护层中植入"响应开关"，使其能够在'
        '特定的工程时间节点（破胶）触发功能切换。'
    )
    replace_para_text(p_123_last, new_123_end)
    print('  Enhanced 1.2.3 ending')
else:
    # Search for alternative
    _, p_123_last = find_para(doc, '如何在单一改性方案中协调这种时序功能切换')
    if p_123_last:
        new_123_end2 = (
            '然而，上述包覆方案均面向涂料、塑料等行业开发，其性能评价标准与压裂液的工程环境'
            '（60~150°C、高矿化度、高剪切、氧化介质）存在根本性差异，面向压裂液工程的系统验证'
            '尚属空白。更为深层的挑战在于：压裂施工对颗粒表面化学提出了一种独特的功能时序需求——'
            '注入期间需要稳定悬浮（颗粒-颗粒排斥主导），关井破胶期间需要向锚定态切换'
            '（颗粒-壁面吸引主导），返排期间需要牢固附着（抵抗流体剪切）。这三种需求对表面化学'
            '的要求各不相同甚至相互矛盾，如何在单一改性方案中协调这种时序功能切换，是一个兼具'
            '工程价值和基础科学意义的问题。这也是本论文第二章表面改性方案设计的核心科学问题。'
        )
        replace_para_text(p_123_last, new_123_end2)
        print('  Enhanced 1.2.3 ending (alternative)')


# ================================================================
# 5. FIX 1.2.4 STRUCTURAL PLACEMENT
# ================================================================
print("\n5. Fixing 1.2.4 structure...")

# The 1.2.4 heading is at P65 but content (P60-P64) was placed BEFORE it.
# Need to: find the heading, then the content paragraphs that should be after it

# Find 1.2.4 heading
_, p_124h = find_para(doc, '1.2.4  胍胶压裂液体系研究现状')
# Find 1.3 heading
_, p_13h = find_para(doc, '1.3  存在的主要问题')

# Find the misplaced content paragraphs: P60-P64
# "综合上述分析，荧光示踪材料的引入对胍胶压裂液体系"
# "在压裂液性能评价方面，中国已建立了较为完善的标准体系"
# "压裂液破胶是决定裂缝导流能力恢复的关键步骤"
# "胍胶压裂液的交联化学核心是硼酸根离子"
# "胍胶（Guar Gum）是从瓜尔豆胚乳中提取的天然半乳甘露聚糖多糖"

# Strategy: Clone these paragraphs and insert them after 1.2.4 heading, then delete originals
paras_to_move = []
for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    # Check if this is one of the misplaced content paragraphs
    if t.startswith('综合上述分析，荧光示踪材料的引入对胍胶压裂液体系') or \
       t.startswith('在压裂液性能评价方面，中国已建立了') or \
       t.startswith('压裂液破胶是决定裂缝导流能力恢复') or \
       t.startswith('胍胶压裂液的交联化学核心是硼酸根离子') or \
       t.startswith('胍胶（Guar Gum）是从瓜尔豆胚乳中提取'):
        paras_to_move.append((i, para))

print(f'  Found {len(paras_to_move)} misplaced content paragraphs to move')

if paras_to_move and p_124h and p_13h:
    from copy import deepcopy
    from lxml import etree

    # Clone and insert after 1.2.4 heading, in correct logical order
    # Correct order: Guar basics → HPG → crosslinking → breaker → standards → compatibility

    # We need to reorder: the paragraphs are in wrong order
    # Current order in text: compatibility(P60) → standards(P61) → breaker(P62) → crosslinking(P63) → guar(P64)
    # Correct order: guar → HPG → crosslinking → breaker → standards → compatibility

    # First, delete all misplaced paragraphs (in reverse order)
    for i, para in reversed(paras_to_move):
        delete_para(para)
    print('  Deleted misplaced paragraphs')

    # Now insert them in correct order after 1.2.4 heading
    # 1. Guar basics (was P64)
    guar_text = (
        '胍胶（Guar Gum）是从瓜尔豆胚乳中提取的天然半乳甘露聚糖多糖，其主链由β-1,4-糖苷键'
        '连接的甘露糖单元组成，侧链为α-1,6-糖苷键连接的半乳糖单元。甘露糖与半乳糖的比值'
        '（M/G比）是决定胍胶水溶性和交联活性的关键结构参数——天然胍胶的M/G比约为1.6~1.8:1，'
        '其分子链上密集分布的顺式邻二醇基团构成了硼酸根离子的交联位点。羟丙基胍胶（HPG）'
        '是胍胶经环氧丙烷醚化改性的衍生物，通过在甘露糖C-6位引入羟丙基取代基，削弱了分子间'
        '氢键网络，显著提升了水合速率，并将水不溶物含量从天然胍胶的5~10 wt%降至3~5 wt%[53,61]。'
        'HPG因其优异的增稠效率、良好的交联活性和较低的残渣含量，已成为水基压裂液的主流稠化剂。'
    )
    add_para_after(doc, p_124h, guar_text)

    # 2. Crosslinking chemistry
    xlink_text = (
        '胍胶压裂液的交联化学围绕硼酸根离子B(OH)₄⁻与HPG顺式邻二醇基团的可逆配位反应展开。'
        'Harris（1993）[58]系统表征了硼交联HPG冻胶的化学与流变学行为，阐明硼酸根与邻二醇'
        '形成1:1和2:1两种配合物，其中2:1配合物构成分子间交联点——这是冻胶获得粘弹性和支撑剂'
        '悬浮能力的化学基础。Kesavan和Prud\'homme（1992）[59]基于小振幅振荡剪切实验建立了'
        '硼交联瓜尔胶的凝胶化动力学模型，定量揭示了交联密度与硼浓度、pH和温度之间的函数关系。'
        '在工程实践中，为延缓交联速度以避免井筒中过早交联导致的高泵注摩阻，普遍采用有机硼'
        '延缓交联剂（如硼酸三乙醇胺酯）。其延迟机制基于配体交换：三乙醇胺配体在常温下占据'
        '硼的配位位点，阻止其与HPG邻二醇结合；升温后配体被HPG邻二醇逐步置换，游离硼酸根'
        '释放，交联反应启动。这种温度触发的延迟机制使交联时间具有可控性。'
    )
    add_para_after(doc, p_124h, xlink_text)

    # 3. Breaker chemistry
    breaker_text = (
        '破胶是压裂液完成造缝和携砂使命后的关键步骤，直接影响裂缝导流能力的恢复程度。'
        '过硫酸铵（(NH₄)₂S₂O₈，APS）是60~120°C储层中最广泛使用的氧化破胶剂。其破胶机理为：'
        'APS在水中热分解产生硫酸根自由基SO₄•⁻，自由基攻击HPG主链的糖苷键和缩醛键，'
        '导致聚合物链断裂、分子量急剧下降，冻胶三维网络坍缩为低粘度液体。Brannon和'
        'Tjon-Joe-Pin（1994）[60]报道了酶破胶剂在高温（>120°C）条件下比APS具有更好的'
        '控释特性和破胶效率。Al-Muntasheri（2014）[53]在其综述中系统比较了氧化破胶剂、'
        '酶破胶剂和酸破胶剂的适用温度范围和破胶效率，指出氧化破胶剂仍是中低温储层的主流选择。'
        '破胶的彻底性是工程上不可妥协的指标：破胶液残渣（包括未完全降解的聚合物碎片、'
        '不溶性植物纤维和残留化学添加剂）堵塞支撑剂充填层的孔隙吼道，可使导流能力降低30%~60%。'
        '因此，引入压裂液体系的任何功能添加剂——包括本研究的荧光示踪颗粒——都必须在不显著增加'
        '残渣含量、不显著延迟破胶时间的前提下实现其设计功能。这一约束构成了第三章破胶性能'
        '评价和地层伤害评价的核心关切。'
    )
    add_para_after(doc, p_124h, breaker_text)

    # 4. Standards
    standards_text = (
        '中国已建立了较为完善的压裂液性能评价标准体系。SY/T 6376-2008《压裂液通用技术条件》[54]'
        '规定了压裂液基液、冻胶和破胶液的各项性能指标及试验方法；SY/T 5107-2016《水基压裂液'
        '性能评价方法》[55]提供了基液表观粘度、交联时间、耐温耐剪切能力、静态悬砂和破胶性能的'
        '标准化测试流程；SY/T 6540-2002和SY/T 6302-2009分别规范了钻井液-完井液对储层的伤害评价'
        '和压裂液对支撑剂充填层导流能力的损害评价方法。本研究严格依据上述标准评价荧光压裂液'
        '体系的各项性能，以常规HPG压裂液的性能基线为参照，定量评估荧光粉引入对各项指标的'
        '独立影响——这是判断荧光示踪方案是否具备工程可行性的基本前提。'
    )
    add_para_after(doc, p_124h, standards_text)

    # 5. Chemical compatibility requirements
    compat_text = (
        '综合上述胍胶化学基础、交联机理、破胶机制和评价标准的分析，荧光示踪材料对胍胶压裂液'
        '体系提出了三个方面的化学相容性要求：（1）示踪颗粒表面不得暴露可参与硼交联的多价阳离子'
        '（如Al³⁺），否则将竞争消耗交联剂，导致冻胶强度下降或交联时间失控；'
        '（2）示踪颗粒在过硫酸铵强氧化环境中的化学稳定性须经过实验验证——表面包覆层既要'
        '在注入阶段充当保护屏障，又要在破胶阶段有控制地响应氧化环境以实现功能切换；'
        '（3）颗粒的引入不得显著增加破胶液残渣含量，否则将抵消荧光示踪带来的诊断收益。'
        '上述要求构成了第二章（表面改性设计）和第三章（压裂液性能评价）的核心工程约束。'
    )
    add_para_after(doc, p_124h, compat_text)

    print('  Re-inserted content in correct logical order after 1.2.4 heading')


# ================================================================
# 6. SAVE
# ================================================================
doc.save(DOC)
print(f'\nAll flow improvements saved to {DOC}')