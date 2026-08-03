#!/usr/bin/env python3
"""Complete rewrite of ESP-T paper."""
from docx import Document
import time
doc = Document('四氧化三铁环氧树脂拟合/ESP-T_final_1781618646.docx')

def S(idx, text):
    p = doc.paragraphs[idx]
    for r in p.runs: r.text = ''
    if len(p.runs) == 0: p.add_run('')
    p.runs[0].text = text

# TITLE
S(0, "Sustained-Release Tracer Proppant with Transport-Based Breakthrough Curve Interpretation for Production Monitoring in Fractured Wells")

# ABSTRACT
S(2,
    "Abstract: Per-interval production monitoring in fractured unconventional "
    "wells relies on tracer-based methods, yet interpreting the breakthrough "
    "curve of a sustained-release tracer proppant to extract quantitative "
    "production information remains an unresolved challenge. Here we address "
    "this by developing a two-component transport model that decomposes the "
    "breakthrough curve into a Gaussian pulse, representing the shut-in "
    "accumulation slug, and an erfc tail, representing sustained matrix-"
    "diffusion-controlled release, joined by a smooth tanh weighting function. "
    "The model was validated using an oleophilic epoxy/Fe3O4 sustained-release "
    "tracer proppant (ESP-T) synthesized by emulsion polymerization of stearic "
    "acid-modified nano-Fe3O4@SA. Fitted to single-phase displacement data, "
    "the model achieves R2 = 0.9939 and recovers the effective production flow "
    "rate Q within 8% of the independently set pump value (0.46 vs. 0.50 mL/min). "
    "The erfc tail accounts for 47% of the integrated tracer signal. ESP-T "
    "exhibits thermal stability to 357 degC, a water contact angle of 104.6 deg "
    "(vs. 72.3 deg for neat epoxy), and a 66% reduction in oil filtration time, "
    "consistent with a water-resistant, oil-permeable transport character. "
    "Under steady-state two-phase flow, tracer flux tracks the oil production "
    "rate across oil-water ratios from 4:1 to 1:4. The combined experimental-"
    "modeling approach provides a framework for per-interval production "
    "monitoring without downhole instrumentation."
)

# INTRODUCTION
S(7,
    "Unconventional oil and gas resources account for more than half of global "
    "hydrocarbon reserves [1-4]. Their development depends on hydraulic "
    "fracturing, which creates artificial fractures in low-permeability "
    "formations to enhance hydrocarbon flow paths and well productivity [5,6]."
)
S(9,
    "After stimulation, the contribution of each fractured interval to total "
    "production varies considerably, and individual intervals decline at "
    "different rates [7]. Without interval-level production data, reservoir "
    "management and stimulation design rely on inference. A method that "
    "quantifies how much oil each interval produces, without permanent "
    "downhole instrumentation, would directly inform refracturing and infill "
    "drilling decisions."
)
S(11,
    "Tracer-based methods provide such a capability. A chemically distinct "
    "tracer added to each fracturing stage appears in the produced fluid, "
    "and its concentration history carries information about that stage's "
    "production. Sustained-release tracer proppants [13-16] extend the "
    "monitoring period from hours to weeks by immobilizing the tracer "
    "within a solid carrier co-injected with the proppant. However, the "
    "breakthrough curve of a sustained-release proppant contains two "
    "overlapping contributions: a concentration pulse from tracer that "
    "accumulates during shut-in, and a long tail from tracer that continues "
    "to release throughout production. These cannot be separated by "
    "inspecting the curve, yet distinguishing them is essential because "
    "the tail carries direct information about the ongoing production "
    "rate. Current practice uses the empirical Korsmeyer-Peppas power law "
    "to characterize release [17,18]; ADE-based fitting extracts fracture "
    "geometry [19]; mass-balance methods allocate production without "
    "exploiting breakthrough curve shape [20]. A transport-based "
    "interpretation framework that resolves the breakthrough curve into "
    "its component contributions and recovers a production rate from the "
    "fitted parameters has received limited attention [21-26]."
)
S(13,
    "Validating such a framework requires a proppant whose release is "
    "dominated by matrix diffusion, so that the sustained tail is well "
    "developed, and whose surface is oleophilic, so that tracer partitions "
    "into the oil phase. Coated proppants lose function upon coating "
    "dissolution; polystyrene microspheres lack thermal stability above "
    "200 degC [27-30]. Epoxy resin, with its cross-linked network, offers "
    "thermal stability above 350 degC, chemical resistance, and tunable "
    "mechanics [31,32]. Emulsion polymerization consolidates nanoparticle "
    "modification and tracer encapsulation into a single step. Epoxy-"
    "encapsulated water-soluble tracers have been demonstrated [33,34]; "
    "an oleophilic variant for oil-phase monitoring has not been reported."
)
S(17,
    "Here we develop a transport model that interprets the breakthrough "
    "curve of a sustained-release tracer proppant by decomposing it into "
    "a Gaussian pulse and an erfc tail connected by a tanh transition. "
    "The decomposition recovers the effective production flow rate Q from "
    "the tracer signal. We synthesize an oleophilic epoxy/Fe3O4 sustained-"
    "release proppant (ESP-T) as the experimental platform, characterize "
    "its structure, stability, wettability, transport, and release kinetics, "
    "and validate the model in single-phase post-shut-in flow and steady-"
    "state two-phase flow."
)

# RESULTS - SEM
S(67,
    "SEM imaging at low magnification (Figure 3-1, panels a and d) shows "
    "well-formed microspheres with excellent sphericity and no inter-particle "
    "agglomeration in both pure epoxy and ESP-T samples. Pure epoxy surfaces "
    "are smooth; ESP-T surfaces are uniformly rough, indicating incorporation "
    "of the nanofiller."
)
S(69,
    "At higher magnification (panels b and e), pure epoxy displays wrinkles "
    "and depressions from curing shrinkage. ESP-T is covered with micron- "
    "and submicron-scale protrusions distributed as discrete islands, "
    "indicating that nano-Fe3O4@SA disperses as nanoclusters and that "
    "stearic acid modification prevents macroscopic phase separation."
)
S(71,
    "At the highest magnification (panels c and f), pure epoxy exhibits the "
    "dense texture of a cross-linked network. ESP-T shows stearic acid-"
    "wrapped nanoclusters embedded in the matrix, with no interfacial cracks "
    "or gaps. This integration indicates physical entanglement between the "
    "stearic acid alkyl chains and the epoxy chains."
)
S(73,
    "The images suggest that hydrophobic nano-Fe3O4@SA clusters disperse "
    "within emulsified epoxy droplets, become immobilized during cross-"
    "linking, and are partially extruded toward the surface during curing "
    "shrinkage. This mechanism is inferred from post-cure microscopy; "
    "in-situ characterization would confirm the pathway directly."
)
S(78,
    "EDS mapping (Figure 3-2) confirms that Fe is distributed throughout "
    "the particle cross-section. Minor signal inhomogeneities reflect "
    "phase domains consistent with nanocluster dispersion."
)

# THERMAL
S(83,
    "Thermal analysis (Figure 3-3) identifies three stages. Stage one "
    "(50-350 degC, 5.70% loss) involves adsorbed water and ethanol. Stage "
    "two (350-400 degC, 72.5% loss, DTG peak at 357.27 degC) is the primary "
    "decomposition from epoxy network degradation. Above 400 degC the "
    "residue stabilizes. A brief DSC endotherm accompanies decomposition."
)
S(87,
    "The decomposition temperature of 357.27 degC far exceeds downhole "
    "temperatures (80-150 degC, up to 200 degC in deep wells). Below "
    "200 degC only trace moisture is lost."
)

# WCA
S(92,
    "Pure epoxy microspheres have a WCA of 72.3 deg (Figure 3-4). ESP-T "
    "yields 104.6 deg, a 32.3 deg increase crossing the hydrophobic "
    "threshold. Stearic acid carboxyl groups coordinate with Fe3O4 hydroxyls "
    "while C17 alkyl chains orient outward, forming a hydrophobic film."
)

# DENSITY
S(99,
    "Both proppants are lighter than water (Figure 3-5): bulk densities "
    "of 0.6179 (pure epoxy) and 0.646 g/cm3 (ESP-T). The increase reflects "
    "the higher density of nano-Fe3O4@SA. Both can be suspended in water-"
    "based fracturing fluids without additional agents."
)

# CONDUCTIVITY
S(105,
    "Filtration tests (Figure 3-6, Table 3-1) show that water passes through "
    "pure epoxy in under 3 min but requires nearly 30 min for ESP-T. "
    "Dodecane shows the reverse: 15 min for pure epoxy, 5 min for ESP-T "
    "(66% reduction). The hydrophobic surface repels water while allowing "
    "oil to spread, favoring oil flow in two-phase production."
)

# RELEASE KINETICS
S(113,
    "Tracer release is diffusion-controlled and temperature-dependent. "
    "Figure 3-7(a) shows release profiles at 30-120 degC; release "
    "accelerates with temperature, is fastest in the first 24 h, and "
    "remains above ICP-MS detection after 14 days at 120 degC."
)
S(122,
    "The Korsmeyer-Peppas model, C/C0 = K tn [23], distinguishes release "
    "mechanisms by the diffusion exponent n. For spheres: n <= 0.43 Fickian "
    "diffusion, 0.43 < n < 0.85 anomalous transport, n >= 0.85 Case-II "
    "relaxation [24]."
)
S(124,
    "Fitted to the release data (Figure 3-7b, Table 3-2), K increases from "
    "0.055 (30 degC) to 0.196 (120 degC). All n values fall within "
    "0.45-0.85 (anomalous transport). R2 exceeds 0.94. The K-P model "
    "lacks an upper asymptote and is strictly valid for Mt/Minf < 0.6."
)
S(126,
    "The mechanism involves solvent-driven swelling: dodecane permeates "
    "the epoxy network, creating a glassy core and gel layer. Swelling "
    "reduces chain entanglement and opens transport channels. Higher "
    "temperature enhances permeability and accelerates diffusion."
)

# SECTION 3.7
S(134,
    "The breakthrough curve of a sustained-release proppant contains two "
    "overlapping signals: a slug from tracer accumulated during shut-in, "
    "and a tail from tracer continuing to release. Because they overlap, "
    "neither can be isolated by inspection."
    "\n\n"
    "We separate them using the one-dimensional advection-dispersion "
    "equation, dC/dt + v dC/dx = D d2C/dx2, where v = 4Q/(pi d2) connects "
    "the model to the production rate Q, and D = alpha v accounts for "
    "mixing. The shut-in slug is described by the instantaneous-injection "
    "solution [25] (a Gaussian pulse). The sustained release is described "
    "by the continuous-source solution (an erfc function). The two are "
    "joined by a tanh weight w(t) = 1/2 [1 + tanh((t0-t)/sigma)], "
    "producing a smooth transition. The model, "
    "C(t) = cb + w(t) C_rise + [1-w(t)] C_fall, is Eq. (2). "
    "Sigma is fixed at the sampling interval (4 min)."
    "\n\n"
    "Q is recovered because v appears in both components: it determines "
    "the pulse arrival time and width, and the tail decay rate. As the "
    "only parameter shared by both components, Q is constrained by two "
    "independent portions of the data. The model was not given the pump "
    "rate as input; it recovered Q from the curve shape."
)
S(148,
    "Fitted to the single-phase breakthrough curve (Figure 3-8b, Table 3-3), "
    "the model achieves R2 = 0.9939 with RMSE = 0.0210. Q = 0.46 mL/min "
    "agrees with the pump setting of 0.50 mL/min within 8%. MRT = 37.4 min "
    "matches x/v = 38.6 min. Pe = x/alpha = 0.934 is consistent with a "
    "gradual-release source. Integrating the components shows 47% of the "
    "signal originates from the erfc tail."
)
S(152,
    "Q, MRT, and Pe are each consistent with independently known values, "
    "indicating physical recovery rather than over-fitting. The 47% tail "
    "contribution means the signal persists without repeated shut-ins. "
    "The tail-plateau concentration cb reflects steady-state release flux; "
    "since ESP-T release is thermally activated (K: 0.055 to 0.196 over "
    "30-120 degC), cb should scale with reservoir temperature."
    "\n\n"
    "This differs from prior ADE-based tracer analyses that extract "
    "fracture geometry [19], perform mass-balance allocation [20], or "
    "characterize fracture networks [26]. Here, the transport decomposition "
    "itself yields the production metric."
)
S(154,
    "To test the model under two-phase conditions, we conducted steady-"
    "state displacement at OWR = 4:1, 1:1, 1:4 and 0.1-0.4 mL/min. "
    "Early effluent at OWR = 1:1 showed higher oil fraction, consistent "
    "with ESP-T's oil-permeable character (Section 3.5)."
)
S(156,
    "Figure 3-9(a): concentration decreases with total flow rate (dilution) "
    "and is independent of OWR. Concentration alone confounds release "
    "rate with dilution."
)
S(158,
    "We therefore define tracer flux FO = (oil concentration) x (oil flow "
    "rate), the mass of tracer per unit time. FO equals the ESP-T release "
    "rate [8] and is independent of dilution. Figure 3-9(b): FO increases "
    "with OWR but remains constant across flow rates at fixed OWR, "
    "consistent with release governed by oil-wetted area."
)
S(160,
    "Figure 3-9(c) normalizes FO by the single-phase steady-state value "
    "(3.187 micro-g/min). Normalized FO agrees with the known oil flow "
    "rate at each OWR, confirming that the production rate can be recovered "
    "from FO. Boundaries: single-interval assumption, steady-state "
    "calibration, dodecane as model oil, untested long-term stability "
    "above 120 degC or in aggressive fluids."
)

# CONCLUSIONS
S(170,
    "A two-component transport model was developed to interpret "
    "breakthrough curves of sustained-release tracer proppants by "
    "decomposing them into a Gaussian pulse and an erfc tail. An "
    "oleophilic epoxy/Fe3O4 proppant (ESP-T) was synthesized for "
    "validation in single-phase post-shut-in and steady-state two-phase flow."
)
S(172,
    "The model recovers the production flow rate Q within 8% of the known "
    "pump setting. The erfc tail accounts for 47% of the signal, confirming "
    "that sustained release provides a persistent monitoring signal from "
    "a single shut-in. Fitted parameters are internally consistent with "
    "independent experimental values."
)
S(174,
    "ESP-T provides thermal stability to 357 degC, oleophilic surface "
    "(WCA 104.6 deg, 66% oil filtration reduction), and self-suspension "
    "(bulk density 0.646 g/cm3). Tracer flux tracks oil production rate "
    "across OWR from 4:1 to 1:4."
)
S(176,
    "Validation is limited to laboratory single-interval experiments "
    "with dodecane. Field application requires multi-interval, multiphase "
    "testing with crude oil. The interpretation framework and experimental "
    "techniques established here provide the foundation for such testing."
)

out = f'四氧化三铁环氧树脂拟合/ESP-T_rewrite_{int(time.time())}.docx'
doc.save(out)
print(f'Saved: {out}')
print('Complete rewrite done.')