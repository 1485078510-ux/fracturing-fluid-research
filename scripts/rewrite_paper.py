# -*- coding: utf-8 -*-
"""Rewrite ESP paper following Chinese reference structure."""
from docx import Document
from docx.oxml.ns import qn

SRC_REF = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\环氧+四氧化三铁改4.10.docx'
SRC_EN  = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP_restructured.docx'
DST     = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP_final_v2.docx'

doc_ref = Document(SRC_REF)
doc_en  = Document(SRC_EN)

# Read Chinese reference structure to understand paragraph mapping
print("=== Chinese reference structure ===")
for i, p in enumerate(doc_ref.paragraphs):
    t = p.text.strip()
    s = p.style.name
    if t:
        print(f'[{i}] [{s}] {t[:120]}')

# The approach: build a new document by copying paragraphs from doc_en
# in the order defined by the Chinese reference structure.
# Sections in order match:
# 1. 引言 [1-5] -> use our polished Introduction
# 2. 实验 [7, 11, 15, 19, 21, 25] -> match our Methods
# 3. 结果 [34, 41, 46, 50, 54, 59, 74] -> match our Results
# 4. 结论 [94-96] -> use our polished Conclusions
"""
print()
print("Done reading structures.")
print("Will use python-docx to copy and reorganize paragraphs.")
"""