#!/usr/bin/env python3
"""SCI-level sentence-by-sentence polish following nature-polishing rules."""
from docx import Document
doc = Document('四氧化三铁环氧树脂拟合/ESP-T_终稿.docx')

# ── Calibrated verb replacements ──
VERB_FIXES = [
    # Overclaim -> calibrated
    ('demonstrating that','indicating that'),# for non-definitive results
    ('clearly shows','shows'),
    ('proves','indicates'),
    # Weak -> stronger where evidence supports
    ('is able to','can'),
    ('was found to',''),
    # Article fixes
    ('the Fig.','Fig.'),
]

def polish_para(p, text):
    """Replace paragraph text with polished version."""
    for r in p.runs: r.text = ''
    p.runs[0].text = text

# ═══════════════════════════════════════════════════════════════════
# ABSTRACT [2] — tighten, calibrate verbs
# ═══════════════════════════════════════════════════════════════════
polish_para(doc.paragraphs[2],
    "Abstract: Per-interval oil production monitoring is essential for "
    "evaluating stimulation effectiveness in unconventional reservoirs, "
    "yet existing tracer technologies lack both a quantitative transport "
    "model linking release kinetics to production rates and a durable "
    "oil-phase monitoring capability. Here we address this dual gap. "
    "We develop a piecewise advection-dispersion model with smooth tanh "
    "transition that decomposes tracer breakthrough curves into a "
    "Gaussian pulse component, representing the shut-in accumulation "
    "slug, and an erfc tailing component, representing matrix-diffusion-"
    "controlled sustained release. We validate this model using an "
    "oleophilic tracer proppant (ESP-T) fabricated by encapsulating "
    "stearic acid-modified nano-Fe3O4 (nano-Fe3O4@SA) within an epoxy "
    "resin matrix via emulsion polymerization. The model achieves "
    "R2 = 0.9939 with the erfc tail accounting for 47% of the integrated "
    "tracer signal, indicating that sustained matrix-diffusion-controlled "
    "release dominates long-term monitoring. The fitted flow rate "
    "(0.46 mL/min) agrees with the pump-set rate (0.50 mL/min) within "
    "8%, and the Peclet number (Pe = 0.934) independently corroborates "
    "the non-Fickian release mechanism identified via Korsmeyer-Peppas "
    "kinetics (n = 0.45-0.85, R2 > 0.90). ESP-T exhibits sphericity "
    "exceeding 0.9, thermal stability to 357.27 degC, and a water "
    "contact angle of 104.6 deg (versus 72.3 deg for neat epoxy), "
    "yielding a water-resistant, oil-permeable transport characteristic "
    "with oil filtration time reduced by 66%. Under steady-state "
    "two-phase flow, tracer flux quantitatively tracks oil-phase "
    "production rates across varying oil-water ratios. This coupled "
    "experimental-modeling framework integrates fracture propping "
    "with long-term production monitoring, offering a broadly "
    "applicable platform for unconventional reservoir management."
)
print('[2] Abstract polished')

# ═══════════════════════════════════════════════════════════════════
# SECTION 3.6 — K-P model description [114]
# ═══════════════════════════════════════════════════════════════════
polish_para(doc.paragraphs[114],
    "The K-P model, a power-law formulation developed by Korsmeyer and "
    "Peppas [27], is widely employed to characterize release kinetics "
    "from controlled-release systems. It captures two mechanisms: "
    "Fickian diffusion, in which species migrate down a concentration "
    "gradient according to Fick's laws, and Case-II relaxation, in "
    "which polymer swelling upon contact with the release medium "
    "increases internal porosity and facilitates escape of entrapped "
    "species. For spherical carriers, the diffusion exponent n "
    "identifies the dominant mechanism: n <= 0.43 corresponds to "
    "Fickian-diffusion-controlled release, 0.43 < n < 0.85 indicates "
    "anomalous transport co-governed by both mechanisms, and n >= 0.85 "
    "corresponds to Case-II-relaxation-controlled release [28]."
)
print('[114] K-P polished')

# ═══════════════════════════════════════════════════════════════════
# SECTION 3.3 [84-86] WCA — merge, tighten
# ═══════════════════════════════════════════════════════════════════
polish_para(doc.paragraphs[84],
    "Water contact angle measurements (Figure 3-4) quantify the change "
    "in surface character produced by the nanofiller. Pure epoxy "
    "microspheres give an average WCA of 72.3 deg, a weakly hydrophilic "
    "surface consistent with the hydroxyl groups present on epoxy chains. "
    "ESP-T yields an average WCA of 104.6 deg, a 32.3 deg increase that "
    "crosses the hydrophobic threshold. This transition is driven by "
    "the stearic acid modification: the carboxyl groups coordinate with "
    "surface hydroxyls on the nano-Fe3O4@SA, while the long alkyl chains "
    "orient outward from the epoxy matrix, forming a hydrophobic film. "
    "The surface enrichment of nanoclusters observed in SEM amplifies "
    "the effect by concentrating these alkyl chains at the proppant "
    "surface, collectively reducing the surface free energy and "
    "suppressing water spreading."
)
# Clear [85] and [86] since content merged into [84]
for idx in [85, 86]:
    for r in doc.paragraphs[idx].runs: r.text = ''
print('[84-86] WCA merged and tightened')

# ═══════════════════════════════════════════════════════════════════
# SECTION 3.2 [75] — TGA tightening
# ═══════════════════════════════════════════════════════════════════
polish_para(doc.paragraphs[75],
    "Thermal analysis (Figure 3-3) identifies three decomposition "
    "stages. The first, 50-350 degC, involves a minor weight loss of "
    "5.70%, attributable to adsorbed water and residual ethanol. The "
    "second stage, 350-400 degC, constitutes the primary decomposition: "
    "cleavage of C-O-C and C-C bonds in the epoxy network releases CO2 "
    "and small hydrocarbons, with the DTG curve peaking at 357.27 degC "
    "and accounting for 72.5% of the total mass loss. Above 400 degC "
    "the residual mass stabilizes; further weight loss derives from "
    "oxidative combustion of the carbonaceous residue. The final residue "
    "comprises hollow glass microspheres, thermally stable "
    "nano-Fe3O4@SA, and a minor carbonaceous fraction. A brief DSC "
    "endotherm accompanies the decomposition, consistent with the "
    "DTG profile."
)
print('[75] TGA tightened')

# ═══════════════════════════════════════════════════════════════════
# Title — check and polish
# ═══════════════════════════════════════════════════════════════════
polish_para(doc.paragraphs[0],
    "An Oleophilic Epoxy/Fe3O4 Tracer Proppant and Tracer Breakthrough "
    "Analysis for Production Monitoring in Fractured Wells"
)
print('[0] Title')

doc.save('四氧化三铁环氧树脂拟合/ESP-T_终稿.docx')
print('\nSCI polish complete.')