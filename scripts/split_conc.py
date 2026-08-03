from docx import Document
import time
doc = Document('四氧化三铁环氧树脂拟合/ESP-T_final_1781616748.docx')
p172 = doc.paragraphs[172]
full = p172.text
split_at = 'These numbers are not curve-fitting artifacts'
if split_at in full:
    idx = full.index(split_at)
    part1 = full[:idx].strip()
    part2 = full[idx:].strip()
    for r in p172.runs: r.text = ''
    p172.runs[0].text = part1
    p173 = doc.paragraphs[173]
    for r in p173.runs: r.text = ''
    if len(p173.runs) == 0:
        p173.add_run('')
    p173.runs[0].text = part2
    out = f'四氧化三铁环氧树脂拟合/ESP-T_final_{int(time.time())}.docx'
    doc.save(out)
    print(f'Saved: {out}')
else:
    print('Split point not found')