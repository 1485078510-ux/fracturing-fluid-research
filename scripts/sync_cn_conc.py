#!/usr/bin/env python3
"""Sync CN conclusions with EN rewrite."""
from docx import Document
doc = Document('四氧化三铁环氧树脂拟合/ESP-T_中文版.docx')

# Find "4. 结论" or "4 结论"
concl_start = None
for i,p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t in ['4. 结论', '4 结论']:
        concl_start = i
        break

if not concl_start:
    print('Conclusions not found')
    exit()

print(f'Conclusions at [{concl_start}]')

# Write para 1 (after header, skipping empty lines)
p1 = doc.paragraphs[concl_start + 1]
for r in p1.runs: r.text = ''
if len(p1.runs) == 0: p1.add_run('')
p1.runs[0].text = (
    "本工作始于一个简单的观察：压裂井井口记录的示踪剂突破曲线所承载的信息，"
    "远超过当前解释方法所能提取。业界已有成熟的材料将示踪剂送入裂缝，也有成熟"
    "的动力学模型描述示踪剂如何释放。缺失的是一个将两者定量连接的环节——一种"
    "从浓度历史中恢复产出该信号的产油速率的方法。本文证明了这样的连接可以从"
    "第一原理传输建模中建立起来，并且行之有效。"
)
print('P1 written')

# Find next non-empty paragraph for para 2
p2 = None
for j in range(concl_start + 3, concl_start + 8):
    if j < len(doc.paragraphs):
        t = doc.paragraphs[j].text.strip()
        if t and len(t) > 10:
            if p2 is None:
                p2 = doc.paragraphs[j]
                break

if p2:
    for r in p2.runs: r.text = ''
    if len(p2.runs) == 0: p2.add_run('')
    p2.runs[0].text = (
        "关键的使能步骤是将突破曲线分解为两个分量——高斯脉冲和erfc拖尾——并通过"
        "平滑的tanh过渡连接。这不是纯粹的数学练习。两个分量对应可识别的物理过程："
        "关井积蓄段塞和持续基质扩散控制释放，其相对贡献可以测量。在本文报道的"
        "单相实验中，erfc拖尾占总积分示踪信号的47%。拟合有效流量Q与独立设定的泵"
        "流量误差在8%以内，拟合停留时间与计算传输时间的偏差在3%以内。这些数字不是"
        "曲线拟合的产物，而是仅从示踪信号中恢复的物理量。"
    )
    print('P2 written')

# Find next non-empty for para 3
p3 = None
for j in range(concl_start + 5, concl_start + 10):
    if j < len(doc.paragraphs):
        t = doc.paragraphs[j].text.strip()
        if t and len(t) > 10 and j != concl_start + 1 and j != concl_start + 3:
            if p3 is None:
                p3 = doc.paragraphs[j]
                break

if p3:
    for r in p3.runs: r.text = ''
    if len(p3.runs) == 0: p3.add_run('')
    p3.runs[0].text = (
        "验证该模型的材料平台ESP-T本身即是一种实用的示踪支撑剂。环氧基体提供了"
        "远超井下需求的357 degC热稳定性。对Fe3O4示踪剂的硬脂酸改性赋予表面亲油性，"
        "水接触角达104.6 deg，油过滤时间较纯环氧缩短66%。在稳态两相流动条件下，"
        "示踪剂通量在从油主导到水主导的全范围油水比下均可追踪产油速率。"
    )
    print('P3 written')

# Para 4
p4 = None
for j in range(concl_start + 7, concl_start + 12):
    if j < len(doc.paragraphs):
        t = doc.paragraphs[j].text.strip()
        if t and len(t) > 10 and j not in [concl_start + 1, concl_start + 3, concl_start + 5]:
            if p4 is None:
                p4 = doc.paragraphs[j]
                break

if p4:
    for r in p4.runs: r.text = ''
    if len(p4.runs) == 0: p4.add_run('')
    p4.runs[0].text = (
        "当前验证限于以十二烷为模型油的实验室规模单段实验。将方法拓展至现场条件——"
        "多个层段同时生产且原油为复杂混合物而非纯烷烃——需要专门的现场试验。本工作"
        "确立的是：核心物理推理是成立的，必要的测量技术——掺杂金属示踪剂的ICP-MS分析"
        "和基于传输的信号分解——已准备就绪。下一步是在规模上检验该框架。"
    )
    print('P4 written')

doc.save('四氧化三铁环氧树脂拟合/ESP-T_中文版.docx')
print('CN conclusions synced')