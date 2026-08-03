#!/usr/bin/env python3
"""Sync CN with all final EN language fixes."""
from docx import Document
doc = Document('四氧化三铁环氧树脂拟合/ESP-T_中文版.docx')

def set_text(para, text):
    for r in para.runs: r.text = ''
    if len(para.runs) == 0: para.add_run('')
    para.runs[0].text = text

# Abstract [2] - remove "dual gap"
for i, p in enumerate(doc.paragraphs):
    if '双重空白开展工作' in p.text:
        t = p.text.replace('本文针对这一双重空白开展工作。建立了',
                          '本文建立了')
        set_text(p, t)
        print(f'[2] CN abstract fixed at [{i}]')
        break

# [67] SEM - "textural"
for i, p in enumerate(doc.paragraphs):
    if '纹理上' in p.text and '纯环氧表面光滑' in p.text:
        t = p.text.replace('纹理上：纯环氧表面光滑',
                          '表面纹理上：纯环氧表面光滑')
        set_text(p, t)
        print(f'[67] CN textural at [{i}]')
        break

# [71] "critical observation"
for i, p in enumerate(doc.paragraphs):
    if '关键的观察在界面处' in p.text:
        t = p.text.replace('关键的观察在界面处',
                          '关键的界面特征是')
        set_text(p, t)
        print(f'[71] CN fixed at [{i}]')
        break

# [83] Remove informal transition
for i, p in enumerate(doc.paragraphs):
    if '建立了ESP-T的结构之后' in p.text or ('ESP-T' in p.text and '是否能够在井下条件下存活' in p.text):
        t = p.text
        t = t.replace('建立了ESP-T的结构之后，我们转向它是否能够在井下条件下存活的问题。',
                      '')
        set_text(p, t)
        print(f'[83] CN transition removed at [{i}]')
        break

# [87] "confirms" -> "indicates"
for i, p in enumerate(doc.paragraphs):
    if '这一热裕度确保' in p.text and '全作业温度范围内保持结构完整性' in p.text:
        t = p.text.replace('确保环氧基体在油气井全作业温度范围内保持结构完整性',
                          '表明环氧基体在油气井全作业温度范围内可保持结构完整性')
        set_text(p, t)
        print(f'[87] CN fixed at [{i}]')
        break

# [105] "barely" -> remove
for i, p in enumerate(doc.paragraphs):
    if '仅需' in p.text and '5 min' in p.text and '缩短66%' in p.text:
        t = p.text.replace('ESP-T仅需5 min',
                          'ESP-T需5 min')
        set_text(p, t)
        print(f'[105] CN barely removed at [{i}]')
        break

# [113] Remove long transition
for i, p in enumerate(doc.paragraphs):
    if '储层温度随深度增加而升高，由于示踪剂从环氧基体中的释放是扩散控制过程' in p.text:
        t = p.text.replace('在建立了材料的结构、热稳定性、润湿性和传输选择性之后，我们考察其在储层相关条件下如何释放示踪剂。',
                          '')
        set_text(p, t)
        print(f'[113] CN transition trimmed at [{i}]')
        break

# [126] informal opening
for i, p in enumerate(doc.paragraphs):
    if '形成的物理图景是溶剂驱动的溶胀' in p.text:
        t = p.text.replace('形成的物理图景是溶剂驱动的溶胀。',
                          '这些结果与溶剂驱动的溶胀机制一致。')
        set_text(p, t)
        print(f'[126] CN fixed at [{i}]')
        break

# [134] informal phrase - Q recovery
for i, p in enumerate(doc.paragraphs):
    if 'Q从突破曲线中恢复的机制简单明了但值得明确说明' in p.text or ('恢复Q' in p.text and '突破曲线' in p.text and '简单' in p.text):
        t = p.text.replace('Q从突破曲线中恢复的机制简单明了但值得明确说明。',
                          'Q从突破曲线中的恢复可理解如下。')
        set_text(p, t)
        print(f'[134] CN Q phrase fixed at [{i}]')
        break

# [152] "defining characteristic"
for i, p in enumerate(doc.paragraphs):
    if '捕捉了缓释型支撑剂的决定性特征' in p.text:
        t = p.text.replace('捕捉了缓释型支撑剂的决定性特征',
                          '捕捉了缓释型支撑剂的特征行为')
        set_text(p, t)
        print(f'[152] CN fixed at [{i}]')
        break

# [160] "closes the loop"
for i, p in enumerate(doc.paragraphs):
    if '闭合了从材料设计' in p.text and '到定量产量指标的回路' in p.text:
        t = p.text.replace('闭合了从材料设计（亲油型扩散控制释放支撑剂）经传输建模（双分量ADE分解）到定量产量指标的回路',
                          '将材料设计（亲油型扩散控制释放支撑剂）与传输建模（双分量ADE分解）连接至定量产量指标')
        set_text(p, t)
        print(f'[160] CN fixed at [{i}]')
        break

# [164] Trim redundant synthesis
for i, p in enumerate(doc.paragraphs):
    if '综合来看，单相和两相结果表明双分量模型与两相条件的FO指标相结合' in p.text:
        set_text(p,
            "综上，双分量模型与两相条件的FO指标相结合，可在单相和两相流中从井口"
            "示踪数据恢复分段产油速率，除引入FO外无需其他修改。"
        )
        print(f'[164] CN trimmed at [{i}]')
        break

doc.save('四氧化三铁环氧树脂拟合/ESP-T_中文版.docx')
print('CN fully synced.')