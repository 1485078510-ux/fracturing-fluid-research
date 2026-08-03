#!/usr/bin/env python3
"""Sync CN version with final English narrative - key paragraphs only."""
from docx import Document
doc = Document('四氧化三铁环氧树脂拟合/ESP-T_中文版.docx')

def set_text(para, text):
    for r in para.runs: r.text = ''
    if len(para.runs) == 0: para.add_run('')
    para.runs[0].text = text

# Introduction [11] - search by content
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if 'Zhao等' in t and 'Zhou等' in t:
        set_text(p,
            "在现有诊断技术中，示踪剂方法为分段产量监测提供了一条实用的路径："
            "无需井下仪器，可随常规压裂作业部署[8-12]。然而，传统油溶性示踪剂"
            "因被快速清除而无法维持长期监测。缓释型示踪支撑剂通过将示踪剂固定于"
            "固体载体中、使其在数周至数月内逐步释放，克服了这一局限。Zhao等[13]、"
            "Zhou等[14]、Li等[15]和Gong等[16]已展示了基于涂层陶瓷和聚苯乙烯微球"
            "的缓释设计。"
            "\n\n"
            "然而，缓释型示踪支撑剂产生了一个瞬时注入示踪剂所没有的解释难题。"
            "井口记录的突破曲线是两个重叠贡献的叠加：关井期间积蓄在支撑剂充填层中"
            "的示踪剂形成的浓度脉冲，以及整个生产期间持续从载体基体中释放的示踪剂"
            "形成的持续拖尾。仅靠观察浓度曲线无法分离这两项贡献，但区分它们对于"
            "提取产量信息是必要的。当前实践采用经验的Korsmeyer-Peppas幂律模型描述"
            "释放[17,18]，这描述了示踪剂从载体中释放的速率，但未提供将井口信号"
            "解释为产量速率的途径。基于ADE的高斯拟合方法提取裂缝几何参数[19]；"
            "示踪剂质量平衡方法[20]在不利用突破曲线形态的情况下进行产量分配。一个"
            "系统性将缓释型示踪支撑剂的持续释放与观测到的突破信号——进而与产油速率"
            "——连接起来的传输模型，尚未得到充分关注[26]。"
        )
        print(f'[11] CN synced at [{i}]')
        break

# Introduction [13] - material gap
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if '环氧树脂为突破这些限制提供了途径' in t or ('涂覆型支撑剂' in t and '环氧树脂' in t and 'Li等' in t):
        set_text(p,
            "为通过实验验证这样的模型，需要一个在井下条件下展现持续油相释放的"
            "示踪支撑剂。现有示踪支撑剂存在材料局限。涂覆型设计在涂层溶解后即失去"
            "示踪功能。聚苯乙烯微球缺乏约200 degC以上的热稳定性[27-30]。环氧树脂"
            "以其高度交联的网络，提供了超过350 degC的热稳定性、卓越的耐化学性和"
            "可调控的力学性能[31,32]。环氧微球的直接乳液聚合将密度控制、纳米粒子"
            "改性和示踪剂包覆整合于单步合成中。Li等[33]和Wei等[34]已证明环氧基体"
            "可作为水溶性示踪剂的载体用于水相流入剖面监测。用于油相监测的亲油型"
            "环氧示踪支撑剂尚未见报道。"
        )
        print(f'[13] CN synced at [{i}]')
        break

# Introduction [17] - this work
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if '本文针对这两项空白开展工作' in t or ('分段对流-弥散' in t and 'ESP-T' in t and '关井' in t):
        set_text(p,
            "本文建立了一个传输模型，通过将缓释型示踪支撑剂的突破曲线分解为"
            "物理上不同的分量——关井积蓄段塞的高斯脉冲和持续基质扩散控制释放的"
            "erfc拖尾，以平滑的tanh过渡连接——来解释该曲线。该分解建立了浓度历史"
            "与产油速率之间的直接关系。我们通过乳液聚合合成了亲油型环氧/Fe3O4"
            "缓释示踪支撑剂（ESP-T），以硬脂酸改性的nano-Fe3O4@SA为亲油示踪剂。"
            "我们表征了其结构、热稳定性、润湿性、传输选择性和力学完整性，并通过"
            "Korsmeyer-Peppas模型量化了其释放动力学。随后在两种与现场作业相关的"
            "流动构型中验证了该解释模型：关井后的单相油流和稳态油水两相流。"
        )
        print(f'[17] CN synced at [{i}]')
        break

# Section 3.7 [134] - derivation
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if '两种过程共同塑造了井口观测到的曲线形态' in t or ('关井期间积蓄' in t and '残余示踪剂' in t and '扩散出来' in t):
        set_text(p,
            "缓释型示踪支撑剂的突破曲线与瞬时示踪剂注入的突破曲线有本质不同。"
            "关井后开井时，积蓄在支撑剂充填层中的示踪剂作为浓缩段塞被扫向井口。"
            "与此同时，示踪剂继续从载体基体中释放，产生持续的拖尾，与段塞信号的"
            "后段重叠。实测浓度曲线是这两项贡献之和，仅靠观察无法分离。"
            "\n\n"
            "我们通过构建一个明确表示两种过程的传输模型来分离它们。一维对流-弥散"
            "方程dC/dt + v dC/dx = D d2C/dx2控制示踪剂浓度沿生产管柱的演化，其中"
            "v = 4Q/(pi d2)为流速，D = alpha v为弥散系数。对关井段塞采用经典的"
            "瞬时段塞注入解[25]，形式为高斯脉冲。对持续释放采用连续源解，形式为"
            "erfc函数。两种解通过双曲正切权重函数w(t) = 1/2 [1 + tanh((t-t0)/sigma)]"
            "拼接，提供从脉冲主导的早期到扩散主导的拖尾期的平滑、物理连续的过渡。"
            "完整复合模型C(t) = cb + w(t) C_rise + [1-w(t)] C_fall如式（2）所示。"
            "参数sigma固定为采样间隔（4 min）；六个自由参数如下所列。"
        )
        print(f'[134] CN synced at [{i}]')
        break

# Section 3.7 [148] - fitting
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if '图3-8(b)' in t and 'R2 = 0.9939' in t and 'Q = 0.46' in t:
        set_text(p,
            "模型拟合至ESP-T单相驱替实验获得的突破曲线（图3-8b，表3-3）。拟合"
            "吻合度很高（R2 = 0.9939，RMSE = 0.0210），但拟合参数值比拟合优度更"
            "具信息量。模型恢复有效流量Q = 0.46 mL/min，与独立设定的泵流量"
            "0.50 mL/min误差在8%以内。这表明产油速率可从示踪信号中估算，无需独立"
            "的流量测量。拟合平均停留时间MRT = 37.4 min与对流传输时间x/v = 38.6 min"
            "一致，支持了模型时间尺度的物理一致性。"
            "\n\n"
            "Peclet数Pe = x/alpha = 0.934将传输状态置于对流-弥散过渡区，与渐进"
            "释放源一致。分别积分两个分量显示，47%的总积分信号源于erfc拖尾。采样"
            "期间井口检测到的近一半示踪剂来自持续释放而非关井段塞，这提示单次关井"
            "即可产生持续的、可解释的信号。"
        )
        print(f'[148] CN synced at [{i}]')
        break

# Section 3.7 [152] - significance
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if 'Q、MRT和Pe从三个独立方向' in t or ('Q' in t and 'MRT' in t and 'Pe' in t and '收敛' in t):
        set_text(p,
            "拟合参数具有超越拟合本身的实用意义。Q、MRT和Pe从独立方向收敛到与"
            "已知实验条件一致的值，提示模型恢复的是物理量而非噪声。47%的拖尾贡献"
            "捕捉了缓释型支撑剂的决定性特征：由于释放持续进行，信号在段塞通过后不"
            "会返回基线。模型捕捉了这一行为并将其与产量速率关联。尾端平台浓度cb反映"
            "了稳态释放通量。由于ESP-T的释放动力学是热激活的（K从30 degC的0.055"
            "升至120 degC的0.196，第3.6节），cb预期具有温度依赖性，为不同深度油井"
            "的模型标定提供了途径。"
            "\n\n"
            "这一解释方法与先前将ADE分析应用于示踪剂数据的工作有所不同，后者聚焦于"
            "提取裂缝几何[19]、执行质量平衡分配[20]或表征裂缝网络[26]。这些方法将"
            "传输步骤视为达到表征终点的工具。本文建立的双分量分解将传输本身作为产量"
            "指标的来源，通过将突破曲线分解为缓释型支撑剂天然产生的可分离贡献。"
        )
        print(f'[152] CN synced at [{i}]')
        break

# Section 3.7 [154] - two-phase
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if '压裂井不会无限期地产出单相油' in t or ('单相分析' in t and '两相' in t and '稳态' in t):
        set_text(p,
            "单相分析表明该模型可在仅有油流过支撑剂充填层时解释突破曲线。然而"
            "在生产井中，地层水最终会突破。为检验该解释在两相条件下是否仍然有效，"
            "我们在三种油水比（OWR = 4:1、1:1、1:4）和四种总流量（0.1-0.4 mL/min）"
            "下进行了稳态驱替实验。在OWR = 1:1时，早期流出液中油相分数更高，与"
            "ESP-T的亲油特性一致（第3.5节）。"
        )
        print(f'[154] CN synced at [{i}]')
        break

# Conclusions [170-176]
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if '本工作始于一个简单的观察' in t:
        set_text(p,
            "本工作回应了示踪基产量监测中的一个实际挑战：如何解释缓释型示踪支撑剂"
            "的突破曲线以提取定量产量信息。我们建立了一个双分量传输模型，将曲线分解"
            "为高斯脉冲（关井积蓄段塞）和erfc拖尾（持续基质扩散控制释放），以平滑"
            "tanh过渡连接。我们合成了亲油型环氧/Fe3O4缓释示踪支撑剂（ESP-T），并在"
            "关井后单相流和稳态两相流实验中验证了该模型。"
        )
        print(f'CN conclusions para 1 synced at [{i}]')
        break

for i, p in enumerate(doc.paragraphs):
    t = p.text
    if '关键的使能步骤是将突破曲线分解' in t and '高斯脉冲和erfc拖尾' in t:
        set_text(p,
            "模型从示踪信号中恢复产量速率（拟合Q = 0.46 mL/min vs 泵设0.50 mL/min，"
            "误差8%），并揭示47%的积分信号源于持续释放。这一发现具有直接的操作意义："
            "单次关井即可产生持续的、可解释的监测信号。拟合参数物理自洽，停留时间和"
            "Peclet数独立匹配实验几何和批实验释放动力学预测值。"
        )
        print(f'CN conclusions para 2 synced at [{i}]')
        break

for i, p in enumerate(doc.paragraphs):
    t = p.text
    if 'ESP-T本身即是一种实用的示踪支撑剂' in t or ('ESP-T' in t and '热稳定性' in t and '缩短66%' in t):
        set_text(p,
            "ESP-T本身即是一种实用的缓释型支撑剂：环氧基体提供357 degC的热稳定性，"
            "硬脂酸改性赋予亲油性（WCA 104.6 deg，油过滤时间缩短66%），体积密度"
            "0.646 g/cm3可在压裂液中自悬浮。在稳态两相流条件下，示踪剂通量在4:1至"
            "1:4的油水比范围内追踪产油速率。"
        )
        print(f'CN conclusions para 3 synced at [{i}]')
        break

for i, p in enumerate(doc.paragraphs):
    t = p.text
    if '当前验证限于' in t and '十二烷' in t and '下一步' in t:
        set_text(p,
            "当前验证限于以十二烷为模型油的实验室规模单段实验。将方法拓展至现场条件"
            "——多个层段同时生产、原油为复杂混合物——需要专门的现场试验。本工作确立"
            "的是核心解释框架的合理性和必要实验技术的可行性。"
        )
        print(f'CN conclusions para 4 synced at [{i}]')
        break

doc.save('四氧化三铁环氧树脂拟合/ESP-T_中文版.docx')
print('CN fully synced.')