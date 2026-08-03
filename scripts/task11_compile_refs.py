# -*- coding: utf-8 -*-
"""
Task 11: Compile, deduplicate, and format all references in ESP-T_v2_manuscript.docx

- Collect in-text citations from the manuscript body (paras 7-9 contain all of them).
- Map each in-text citation number to the source reference list (ESP-T_Final_4-revised.docx [1]-[41]).
- Renumber citations by order of first appearance into a compact sequential list (no orphans).
- Replace the 35 placeholder slots with the 17 formatted reference entries; delete unused slots
  and the scaffolding note paragraph.
"""
import re
import os
from docx import Document
from docx.shared import Pt

TARGET = r"四氧化三铁环氧树脂拟合/ESP-T_投稿文件/ESP-T_v2_manuscript.docx"
SOURCE = r"四氧化三铁环氧树脂拟合/ESP-T_投稿文件/ESP-T_Final_4-revised.docx"

# ---------- 1. Load source references [1]-[41] ----------
src = Document(SOURCE)
source_refs = {}
for p in src.paragraphs:
    m = re.match(r"^\[(\d+)\]", p.text.strip())
    if m:
        n = int(m.group(1))
        if 1 <= n <= 41:
            # strip the leading "[N] " and any trailing whitespace
            body = p.text.strip()
            body = re.sub(r"^\[\d+\]\s*", "", body)
            source_refs[n] = body
assert len(source_refs) == 41, f"expected 41 source refs, got {len(source_refs)}"

# ---------- 2. In-text citation -> source reference mapping ----------
# Citations as they appear in the manuscript body (source numbering):
#   [5,6], [9], [10,11], [12], [28,29], [31,32], [21], [22], [23], [24], [16], [33], [34]
# Final numbering = order of first appearance in the text:
first_appearance = [5, 6, 9, 10, 11, 12, 28, 29, 31, 32, 21, 22, 23, 24, 16, 33, 34]
final_num = {src_no: i + 1 for i, src_no in enumerate(first_appearance)}
print("Mapping source# -> final#:", final_num)

# In-text token replacement map (source token -> final token), longest tokens first
token_map = {
    "[5,6]": "[1,2]",
    "[10,11]": "[4,5]",
    "[28,29]": "[7,8]",
    "[31,32]": "[9,10]",
    "[9]": "[3]",
    "[12]": "[6]",
    "[21]": "[11]",
    "[22]": "[12]",
    "[23]": "[13]",
    "[24]": "[14]",
    "[16]": "[15]",
    "[33]": "[16]",
    "[34]": "[17]",
}
tokens_sorted = sorted(token_map.keys(), key=len, reverse=True)
token_re = re.compile("|".join(re.escape(t) for t in tokens_sorted))

# ---------- 3. Edit target document ----------
doc = Document(TARGET)

# Find the References heading; only paragraphs before it are body text.
refs_heading_idx = next(
    i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "References"
)

# 3a. Rewrite in-text citations in body paragraphs (7, 8, 9 are the only ones with citations)
changed = []
for i, p in enumerate(doc.paragraphs):
    if i >= refs_heading_idx:
        break
    if not token_re.search(p.text):
        continue
    if p.runs:
        new_text = token_re.sub(lambda m: token_map[m.group(0)], p.text)
        p.runs[0].text = new_text
        for extra in p.runs[1:]:
            extra.text = ""
        changed.append(i)
        print(f"rewrote citations in paragraph {i}")
assert changed == [7, 8, 9], f"unexpected citation paragraphs: {changed}"

# 3b. Rebuild the references section
note_removed = False
placeholder_nums = []
for p in list(doc.paragraphs):
    t = p.text.strip()
    if t.startswith("[numbered list, 35 slots"):
        p._element.getparent().remove(p._element)
        note_removed = True
    elif re.fullmatch(r"\[\d+\]", t):
        placeholder_nums.append((int(t[1:-1]), p))

assert note_removed, "scaffolding note paragraph not found"
assert len(placeholder_nums) == 35, f"expected 35 placeholder paragraphs, got {len(placeholder_nums)}"

FINAL_ENTRY_FOR_SLOT = {  # slot number -> source ref number
    1: 5, 2: 6, 3: 9, 4: 10, 5: 11, 6: 12, 7: 28, 8: 29, 9: 31, 10: 32,
    11: 21, 12: 22, 13: 23, 14: 24, 15: 16, 16: 33, 17: 34,
}

for slot_no, para in sorted(placeholder_nums):
    if slot_no in FINAL_ENTRY_FOR_SLOT:
        src_no = FINAL_ENTRY_FOR_SLOT[slot_no]
        text = f"[{final_num[src_no]}] {source_refs[src_no]}"
        if para.runs:
            para.runs[0].text = text
            for extra in para.runs[1:]:
                extra.text = ""
        else:
            para.add_run(text)
        para.paragraph_format.space_after = Pt(2)
        print(f"slot [{slot_no}] -> entry: {text[:70]}...")
    else:
        para._element.getparent().remove(para._element)
        print(f"slot [{slot_no}] -> deleted (orphan / unused)")

doc.save(TARGET)
print("saved", TARGET)
