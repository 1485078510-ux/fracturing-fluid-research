# -*- coding: utf-8 -*-
"""Verify that only intended paragraphs changed between backup and edited manuscript."""
from docx import Document

orig = Document("四氧化三铁环氧树脂拟合/ESP-T_投稿文件/ESP-T_v2_manuscript.docx.bak_task11")
new = Document("四氧化三铁环氧树脂拟合/ESP-T_投稿文件/ESP-T_v2_manuscript.docx")

o = [p.text for p in orig.paragraphs]
n = [p.text for p in new.paragraphs]
print("orig paragraphs:", len(o), "| new paragraphs:", len(n))

changed = [i for i in range(min(len(o), len(n))) if o[i] != n[i]]
print("paragraph indices with text differences:", changed)

body_same = all(o[i] == n[i] for i in range(0, 115))
print("body paras 0-114 identical:", body_same)
print("orig refs section paras:", len(o) - 115, "| new refs section paras:", len(n) - 115)
if len(n) < len(o):
    print("orig paragraphs beyond new end:", [t[:40] for t in o[len(n):]])
if len(n) > len(o):
    print("new paragraphs beyond orig end:", [t[:40] for t in n[len(o):]])
