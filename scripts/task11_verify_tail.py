# -*- coding: utf-8 -*-
"""Confirm the Figure Captions / Tables sections (after References) survived the edit."""
from docx import Document

orig = Document("四氧化三铁环氧树脂拟合/ESP-T_投稿文件/ESP-T_v2_manuscript.docx.bak_task11")
new = Document("四氧化三铁环氧树脂拟合/ESP-T_投稿文件/ESP-T_v2_manuscript.docx")

o = [p.text for p in orig.paragraphs]
n = [p.text for p in new.paragraphs]

# New doc: paras 134..end should equal orig paras 153..end (Figure Captions onward)
tail_same = o[153:] == n[134:]
print("Figure Captions/Tables sections identical after shift:", tail_same)
if not tail_same:
    for i in range(max(len(o) - 153, len(n) - 134)):
        a = o[153 + i] if 153 + i < len(o) else "<end>"
        b = n[134 + i] if 134 + i < len(n) else "<end>"
        if a != b:
            print("first mismatch:", repr(a[:50]), "vs", repr(b[:50]))
            break

print()
print("=== NEW REFERENCES SECTION ===")
for i in range(115, len(n)):
    t = n[i]
    if not t.strip():
        print(f"[{i}] <empty>")
    else:
        print(f"[{i}] {t[:90]}")
    if t.strip() == "References":
        continue
    if t.strip().startswith("Figure Captions"):
        break
