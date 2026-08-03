# -*- coding: utf-8 -*-
"""Task 13: Build Supplementary_Material_v2.docx from source Supplementary_Material.docx.

- Keeps S1-S5 structure; updates main-text cross-references to the v2 manuscript
  numbering (Section 2.1->3.1, 2.3->3.2, 3.1->3.2, 3.2->3.3, 3.3->4.2, 3.4->5,
  Eq. 1 / Section 3.3 -> Eq. (1) / Section 2.2).
- Renumbers SM figures so they do not collide with the new main-text Fig. S1
  (synthesis schematic) and Fig. S2 (apparatus schematic):
  optical micrographs S3 (unchanged), release profiles S1->S4, two-phase S2->S5.
- Adds the figure-numbering note and an S5.3 lead-in linking to main-text Section 3.5.
- Reuses embedded figure images (release profiles, two-phase) from the source file.
"""
import copy
import os
import zipfile
import shutil
from docx import Document
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC = os.path.join(BASE, "四氧化三铁环氧树脂拟合", "Supplementary_Material.docx")
DST = os.path.join(BASE, "四氧化三铁环氧树脂拟合", "ESP-T_投稿文件", "Supplementary_Material_v2.docx")
TMP_MEDIA = os.path.join(BASE, "_task13_media")

# Old main-text section -> new v2 manuscript section.
# Order matters: 3.2->3.3 must run before 2.3->3.2 and 3.1->3.2 create new
# "Section 3.2" strings; 3.3->4.2 must run before 3.2->3.3 creates "Section 3.3".
REPLACEMENTS = [
    ("Sections 3.2–3.4 of the main text", "Sections 3.3, 4.2, and 5 of the main text"),
    ("Section 3.4 of the main text", "Section 5 of the main text"),
    ("Section 3.3 of the main text", "Section 4.2 of the main text"),
    ("Section 3.2 of the main text", "Section 3.3 of the main text"),
    ("Section 3.1 of the main text", "Section 3.2 of the main text"),
    ("Section 2.3 of the main text", "Section 3.2 of the main text"),
    ("Section 2.1 of the main text", "Section 3.1 of the main text"),
    ("for all model fitting in Section 3.3", "for all model fitting in Section 4.2"),
    ("supporting the K-P analysis in Section 3.2", "supporting the K-P analysis in Section 3.3"),
    ("supporting the production allocation analysis in Section 3.4",
     "supporting the production allocation analysis in Section 5"),
    ("(Eq. 1, main text, Section 3.3)", "(Eq. (1), main text, Section 2.2)"),
]

# Figure caption renumbering (old caption prefix -> new caption prefix).
FIG_RENUMBER = {
    "Figure S1.": "Figure S4.",   # release profiles (S4.2)
    "Figure S2.": "Figure S5.",   # two-phase data (S4.3)
    "Figure S3.": "Figure S3.",   # optical micrographs (S3.1) - unchanged
}

# Media files to extract from the source package: old Fig. S1 / S2 images.
MEDIA = [("word/media/image1.png", "release_profiles.png"),
         ("word/media/image2.png", "twophase.png")]

# Paragraph markers for special handling.
S53_STAGE1 = "Stage 1 -- Differential Evolution"
PAPER_TITLE = ("An Oleophilic Epoxy/Fe₃O₄ Tracer Proppant with Piecewise ADE Modeling "
               "for Production Monitoring in Multi-Stage Fractured Horizontal Wells")
NOTE = ("Note: Fig. S1 (ESP-T synthesis schematic) and Fig. S2 (displacement apparatus schematic) "
        "appear in the main text (Sections 3.1 and 3.4) and are not reproduced here; supplementary "
        "figures are therefore numbered from Fig. S3 onward.")
S53_LEAD = ("The numerical implementation details of the two-pass parameter estimation protocol "
            "summarized in Section 3.5 of the main text are as follows.")


def transform_text(text):
    for old, new in REPLACEMENTS:
        text = text.replace(old, new)
    return text


def main():
    os.makedirs(TMP_MEDIA, exist_ok=True)
    # Extract embedded figure images from the source package.
    with zipfile.ZipFile(SRC) as z:
        for internal, name in MEDIA:
            with open(os.path.join(TMP_MEDIA, name), "wb") as f:
                f.write(z.read(internal))
    # Read the inline image widths (EMU) used in the source to keep the same scale.
    src_doc = Document(SRC)
    ns_ext = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
    extents = [int(e.get("cx")) for e in src_doc.element.body.iter(ns_ext + "extent")][:2]
    img_cx = {"release_profiles.png": extents[0] if len(extents) > 0 else Emu(int(5.9 * 914400)),
              "twophase.png": extents[1] if len(extents) > 1 else Emu(int(6.2 * 914400))}

    # New document.
    out = Document()
    normal = out.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    # Title block.
    t = out.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Supplementary Material")
    r.bold = True
    r.font.size = Pt(16)
    t2 = out.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run(PAPER_TITLE)
    r2.italic = True
    r2.font.size = Pt(12)
    note = out.add_paragraph()
    nr = note.add_run(NOTE)
    nr.italic = True
    nr.font.size = Pt(10)

    # Copy source body elements in order.
    for child in src_doc.element.body.iterchildren():
        tag = child.tag
        if tag.endswith("}p"):
            p = Paragraph(child, src_doc)
            txt = p.text.strip()
            if not txt:
                continue
            # Skip the two title paragraphs (re-added above).
            if txt == "Supplementary Material" or txt.startswith("An Oleophilic Epoxy/Fe₃O₄"):
                continue
            # Insert the figure images just before their (renumbered) captions.
            if txt.startswith("Figure S1."):
                img_p = out.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_p.add_run().add_picture(os.path.join(TMP_MEDIA, "release_profiles.png"),
                                            width=Emu(int(img_cx["release_profiles.png"])))
            if txt.startswith("Figure S2."):
                img_p = out.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_p.add_run().add_picture(os.path.join(TMP_MEDIA, "twophase.png"),
                                            width=Emu(int(img_cx["twophase.png"])))
            # S5.3 lead-in linking the protocol to main-text Section 3.5.
            if txt.startswith(S53_STAGE1):
                out.add_paragraph(S53_LEAD)
            # Apply caption renumbering and cross-reference updates.
            new_txt = txt
            if txt.startswith("Figure S"):
                prefix, rest = txt.split(".", 1)
                new_txt = FIG_RENUMBER.get(prefix + ".", prefix + ".") + rest
            new_txt = transform_text(new_txt)
            # Rebuild as a single-run paragraph (source paragraphs are single-run).
            np = out.add_paragraph()
            run = np.add_run(new_txt)
            if p.runs:
                src_run = p.runs[0]
                if src_run.bold:
                    run.bold = True
                if src_run.italic:
                    run.italic = True
                if src_run.font.size:
                    run.font.size = src_run.font.size
        elif tag.endswith("}tbl"):
            from docx.table import Table
            tbl = Table(copy.deepcopy(child), out)
            out.element.body.append(tbl._tbl)

    out.save(DST)
    shutil.rmtree(TMP_MEDIA, ignore_errors=True)
    print("saved:", DST)


if __name__ == "__main__":
    main()
