#!/usr/bin/env python3
"""Unify Section 3.7 language and paragraph transitions."""
from docx import Document
doc = Document('四氧化三铁环氧树脂拟合/ESP-T_终稿.docx')

# [142] Smooth lead-in from model description
p142 = doc.paragraphs[142]
new142 = (
    "Fitting this model to the single-phase tracer breakthrough curve "
    "(Figure 3-8b; parameters in Table 3-3) yields R2 = 0.9939 with "
    "RMSE = 0.0210 and residuals randomly distributed within +/- 2 sigma "
    "across the full time range. The fitted effective flow rate "
    "Q = 0.46 mL/min agrees with the independently set pump flow rate of "
    "0.50 mL/min within 8%, and the fitted mean residence time "
    "MRT = 37.4 min matches the convective travel time x/v = 38.6 min "
    "(ratio 0.967). The Peclet number Pe = x/alpha = 0.934 places the "
    "transport at the advection-dispersion transition, consistent with "
    "a matrix-diffusion-controlled source. Integrating the C_rise and "
    "C_fall components separately reveals that 47% of the total tracer "
    "signal originates from the erfc tail, confirming that the non-Fickian "
    "release mechanism identified in Section 3.6 governs tracer transport "
    "under flow conditions."
)
for r in p142.runs: r.text = ''
p142.runs[0].text = new142
print('[142] OK')

# [146] Bridge: fitting -> parameter meaning
p146 = doc.paragraphs[146]
new146 = (
    "Beyond the quality of the fit, the recovered parameter values carry "
    "physical meaning that reinforces confidence in the model. The flow "
    "rate Q, the residence time MRT, and the Peclet number Pe are each "
    "independently consistent with the known pump setting, the flow-path "
    "geometry, and the non-Fickian kinetics from Section 3.6, respectively. "
    "This multi-parameter agreement confirms that the model parameters "
    "correspond to measurable physical quantities recoverable from the "
    "tracer signal alone. Among these, the 47% erfc tail contribution "
    "carries particular operational significance: sustained matrix-diffusion-"
    "controlled release, not the initial shut-in pulse, dominates the "
    "long-term monitoring signal, implying that a single shut-in period "
    "suffices for extended production surveillance. The tail-plateau "
    "concentration cb reflects the steady-state diffusion flux sustained "
    "by the thermally activated release documented in Section 3.6 "
    "(K increasing from 0.055 at 30 degC to 0.196 at 120 degC), "
    "suggesting that cb should scale with reservoir temperature."
)
for r in p146.runs: r.text = ''
p146.runs[0].text = new146
print('[146] OK')

# [148] Bridge: single-phase -> two-phase
p148 = doc.paragraphs[148]
new148 = (
    "The single-phase results validate the model's core premise: the "
    "tracer breakthrough curve, properly decomposed, yields a quantitative "
    "estimate of the oil production rate. To determine whether this "
    "capability persists under the more realistic condition of two-phase "
    "flow, we conducted steady-state core displacement experiments at "
    "three oil-water ratios (OWR = 4:1, 1:1, 1:4) and four total flow "
    "rates (0.1-0.4 mL/min). Direct observation of the effluent at "
    "OWR = 1:1 confirmed a higher oil fraction in the early-stage "
    "effluent, consistent with the oil-permeable characteristic of ESP-T "
    "established in Section 3.5."
)
for r in p148.runs: r.text = ''
p148.runs[0].text = new148
print('[148] OK')

# [151] Connected to [149] dilution observation
p151 = doc.paragraphs[151]
new151 = (
    "The dilution-driven concentration trend in Figure 3-9(a) highlights "
    "a practical complication: concentration alone mixes the tracer release "
    "rate with the total flow rate, obscuring the underlying production "
    "signal. To separate these effects, we define the tracer flux "
    "FO as the mass of tracer passing the wellhead per unit time, "
    "equal to the product of the oil-phase concentration and the oil "
    "volumetric flow rate. FO directly reflects the tracer release rate "
    "from the ESP-T pack [8], independent of dilution. Figure 3-9(b) "
    "shows that FO increases with OWR but remains constant across total "
    "flow rates. This behavior is physically consistent with Section 3.5: "
    "at fixed OWR the oil-proppant contact area is unchanged, so FO stays "
    "constant regardless of how fast fluid moves through the pack; "
    "increasing OWR expands the oil-wetted area, proportionally raising FO."
)
for r in p151.runs: r.text = ''
p151.runs[0].text = new151
print('[151] OK')

# [153] Validation -> acknowledge boundaries before synthesis
p153 = doc.paragraphs[153]
new153 = (
    "Figure 3-9(c) compares the normalized FO with the actual oil-phase "
    "flow rate for all three OWR values at a fixed total flow rate of "
    "0.1 mL/min, using the steady-state FO from single-phase oil "
    "displacement (3.187 micro-g/min) as the normalization reference. "
    "The agreement between normalized FO and the set oil flow rate "
    "confirms that, under steady-state two-phase conditions, the oil "
    "production rate of a labeled interval can be recovered from the "
    "FO measurement when the total flow rate is known."
    "\n\n"
    "Several boundaries of the present validation should be kept in view. "
    "The ADE model assumes a single fractured interval with uniform "
    "proppant packing; multi-interval interactions and pack heterogeneity "
    "in the field may introduce systematic deviations. The FO calibration "
    "derives from steady-state experiments and may not hold during "
    "transient well operations such as start-up, shut-in, or rapid "
    "drawdown. Dodecane served as the model oil; crude oil components, "
    "particularly asphaltenes, may alter the wetting and diffusion "
    "behavior of the epoxy matrix. The long-term chemical stability of "
    "the epoxy matrix and the integrity of the stearic acid modification "
    "at temperatures above 120 degC, or in the presence of aggressive "
    "formation fluids (high salinity, CO2, H2S), warrant dedicated "
    "investigation."
)
for r in p153.runs: r.text = ''
p153.runs[0].text = new153
print('[153] OK')

# [154] Synthesis - acknowledge limits, then conclude
p154 = doc.paragraphs[154]
new154 = (
    "Within these boundaries, the single-phase and two-phase results "
    "together outline a coherent workflow. The model decomposes a "
    "wellhead tracer concentration history into physically meaningful "
    "components, yielding an effective flow rate Q that serves as a "
    "verifiable proxy for the per-interval oil production rate. Because "
    "the key parameters (Q, MRT, Pe) are physically constrained, they "
    "can be cross-checked against independent operational data, providing "
    "an internal consistency verification absent from purely empirical "
    "interpretations of tracer data."
)
for r in p154.runs: r.text = ''
p154.runs[0].text = new154
print('[154] OK')

# [156] Field guidance - pick up from workflow
p156 = doc.paragraphs[156]
new156 = (
    "Translating this workflow to field operations, ESP-T particles "
    "doped with distinct metal elements (e.g., Mn, Zn, Cu) for each "
    "fracture stage are co-injected with the proppant during hydraulic "
    "fracturing, enabling simultaneous multi-stage deployment in a single "
    "well. The model then informs several operational decisions. The "
    "fitted crossover time t0 and the Gaussian pulse-peak position guide "
    "the shut-in duration: the objective is a clearly identifiable peak, "
    "and the 47% tail contribution indicates that extending the shut-in "
    "beyond this point yields diminishing returns. The time to peak, "
    "predictable from wellbore geometry and the fitted flow velocity, "
    "also tells the operator when to expect the maximum concentration "
    "at the wellhead, enabling an optimized sampling schedule. Comparing "
    "the fitted Q values across stages identifies underperforming "
    "intervals for refracturing candidate selection or stimulation design "
    "adjustments on subsequent wells. Because FO is independent of total "
    "flow rate, routine production sampling, without additional shut-ins "
    "or downhole intervention, can track per-stage contributions over "
    "the full well lifecycle. A declining total production rate does not "
    "degrade the measurement; periodic ICP-MS analysis processed through "
    "the model yields a time series of per-interval production rates, "
    "revealing differential stage decline and informing decisions on "
    "infill drilling or artificial lift adjustments."
)
for r in p156.runs: r.text = ''
p156.runs[0].text = new156
print('[156] OK')

doc.save('四氧化三铁环氧树脂拟合/ESP-T_终稿.docx')
print('\nSection 3.7 unified.')