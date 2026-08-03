#!/usr/bin/env python3
"""Unify Introduction with body structure + add section transitions."""
from docx import Document
import time
doc = Document('四氧化三铁环氧树脂拟合/ESP-T_final_1781617611.docx')

# ═══════════════════════════════════════════════════════════════════
# REWRITE INTRODUCTION to match body structure
# Body order: Material→Release→Modeling→Field
# Intro should preview: material→release kinetics→transport model→production allocation
# ═══════════════════════════════════════════════════════════════════

# Keep [7] and [9] as-is (energy context + monitoring problem)

# [11] Tracers + problem statement (shortened)
p11 = doc.paragraphs[11]
for r in p11.runs: r.text = ''
p11.runs[0].text = (
    "Among available diagnostic techniques, tracer-based methods offer a "
    "practical route to per-interval production monitoring: they require "
    "no downhole instrumentation, can be deployed during routine fracturing "
    "operations, and provide direct chemical evidence of contributions from "
    "each labeled interval. Conventional oil-soluble chemical tracers, "
    "however, exhibit limited compatibility with water-based fracturing "
    "fluids and cannot sustain long-term monitoring [8-12]. Tracer proppants, "
    "which immobilize tracer agents within a solid carrier co-injected with "
    "the proppant, overcome these limitations by eliminating separate tracer-"
    "injection operations. Zhao et al. [13], Zhou et al. [14], Li et al. [15], "
    "and Gong et al. [16] have demonstrated tracer proppants based on ceramic "
    "and polystyrene carriers with various coating and encapsulation "
    "strategies. Yet across all of these studies, tracer signals are "
    "interpreted through empirical kinetic models, predominantly the "
    "Korsmeyer-Peppas power law [17,18]. The transport of released tracer "
    "from the proppant pack through the production tubing to the sampling "
    "point has not been quantitatively linked to per-interval production "
    "rates [19,20,26]. The Korsmeyer-Peppas framework [23,24] and classical "
    "ADE analytical solutions [25] provide well-established mathematical "
    "tools, and oleophilic nano-tracers [21] together with proppant standards "
    "[22] supply the material foundations, but these elements have not been "
    "integrated into a production-allocation methodology."
)

# [13] Material motivation (new, connects to body 3.1-3.5)
p13 = doc.paragraphs[13]
for r in p13.runs: r.text = ''
p13.runs[0].text = (
    "Addressing this gap requires a tracer proppant that delivers two "
    "capabilities simultaneously: durable, oil-phase tracer release under "
    "downhole conditions, and a material structure compatible with the "
    "transport model used to interpret the resulting signal. Existing "
    "materials fall short. Coated proppants are dense and lose their tracer "
    "function once the coating dissolves. Polystyrene microspheres lack "
    "the thermal stability and mechanical strength needed at depth [27-30]. "
    "Epoxy resin, with its highly cross-linked network, offers thermal "
    "stability above 350 degC, chemical resistance, and tunable mechanics "
    "[31,32]. Unlike multi-step coating processes, direct emulsion "
    "polymerization of epoxy microspheres consolidates density control, "
    "nanoparticle modification, and tracer encapsulation into a single "
    "synthetic step. Li et al. [33] and Wei et al. [34] demonstrated that "
    "epoxy matrices can carry water-soluble tracers for aqueous inflow "
    "profiling. Extending this concept to oil-phase monitoring requires an "
    "oleophilic matrix that releases its payload selectively into oil rather "
    "than water. No such material, coupled with a transport-based "
    "interpretation model, has been reported."
)

# [17] This work (merged from old [17] and [21])
p17 = doc.paragraphs[17]
for r in p17.runs: r.text = ''
p17.runs[0].text = (
    "Here we address both requirements. We synthesize an oleophilic "
    "epoxy/Fe3O4 tracer proppant (ESP-T) by emulsion polymerization, using "
    "stearic acid-modified nano-Fe3O4@SA as the oleophilic tracer. We "
    "characterize its microstructure, thermal stability, wettability, oil-"
    "water transport selectivity, and mechanical integrity. We quantify its "
    "temperature-dependent tracer release kinetics through the "
    "Korsmeyer-Peppas model. We then develop a piecewise advection-dispersion "
    "model that decomposes the tracer breakthrough curve into a Gaussian "
    "pulse, representing the shut-in accumulation slug, and an erfc tail, "
    "representing sustained matrix-diffusion-controlled release. These "
    "components are joined by a smooth tanh transition, yielding a direct "
    "quantitative relationship between the concentration history and the "
    "per-interval production rate. We validate the integrated approach "
    "through single-phase and steady-state two-phase core displacement "
    "experiments. The study proceeds from material design through release "
    "characterization to transport modeling and production-rate "
    "quantification, establishing a framework in which each element "
    "supports the next."
)

# Clear old [21] (merged into [17])
for r in doc.paragraphs[21].runs: r.text = ''

# ═══════════════════════════════════════════════════════════════════
# ADD SECTION TRANSITIONS in Results
# ═══════════════════════════════════════════════════════════════════

# [67] SEM opening: lead with purpose
p67 = doc.paragraphs[67]
t67 = p67.text
if not t67.startswith('The structure of ESP-T'):
    for r in p67.runs: r.text = ''
    p67.runs[0].text = (
        "The structure of ESP-T determines its function. SEM imaging "
        "at low magnification (Figure 3-1, panels a and d) reveals "
        "well-formed microspheres with excellent sphericity and no "
        "inter-particle agglomeration in both pure epoxy and ESP-T "
        "samples, confirming that the emulsion polymerization process "
        "accommodates the nano-Fe3O4@SA filler. The most visible "
        "difference is textural: pure epoxy surfaces are smooth, whereas "
        "ESP-T surfaces are uniformly rough, an early indication that "
        "the nanofiller has been incorporated throughout the particle."
    )

# [82] Thermal: connect to SEM
p82 = doc.paragraphs[82]  # "3.2 Thermal Stability"
# Add transition note at start of [83]
p83 = doc.paragraphs[83]
t83 = p83.text
if not t83.startswith('Having established'):
    for r in p83.runs: r.text = ''
    p83.runs[0].text = (
        "Having established the structure of ESP-T, we turn to the "
        "question of whether it can survive downhole conditions. Thermal "
        "analysis (Figure 3-3) identifies three decomposition stages. "
        "The first, 50-350 degC, involves a minor weight loss of 5.70%, "
        "attributable to adsorbed water and residual ethanol. The second "
        "stage, 350-400 degC, constitutes the primary decomposition: "
        "cleavage of C-O-C and C-C bonds in the epoxy network releases "
        "CO2 and small hydrocarbons, with the DTG curve peaking at "
        "357.27 degC and accounting for 72.5% of the total mass loss. "
        "Above 400 degC the residual mass stabilizes; further weight "
        "loss derives from oxidative combustion of the carbonaceous "
        "residue. The final residue comprises hollow glass microspheres, "
        "thermally stable nano-Fe3O4@SA, and a minor carbonaceous "
        "fraction. A brief DSC endotherm accompanies the decomposition, "
        "consistent with the DTG profile."
    )

# [91] WCA: connect to SEM nanocluster extrusion
p92 = doc.paragraphs[92]
t92 = p92.text
if not t92.startswith('The rough, nanocluster-rich surface'):
    for r in p92.runs: r.text = ''
    p92.runs[0].text = (
        "The rough, nanocluster-rich surface documented by SEM produces "
        "a pronounced shift in wettability. Water contact angle "
        "measurements (Figure 3-4) give an average of 72.3 deg for pure "
        "epoxy microspheres, consistent with the hydroxyl groups on "
        "epoxy chains. ESP-T yields 104.6 deg, a 32.3 deg increase "
        "that crosses the hydrophobic threshold. The stearic acid "
        "modification drives this transition: carboxyl groups coordinate "
        "with Fe3O4 surface hydroxyls, while the C17 alkyl chains orient "
        "outward, forming a low-energy hydrophobic film. The surface "
        "enrichment of these chains, driven by the nanocluster extrusion "
        "observed in SEM, amplifies the effect."
    )

# [104] Conductivity: connect to wettability
p105 = doc.paragraphs[105]
t105 = p105.text
if not t105.startswith('This wettability reversal'):
    for r in p105.runs: r.text = ''
    p105.runs[0].text = (
        "This wettability reversal produces a corresponding reversal in "
        "transport behavior. Filtration tests (Figure 3-6, Table 3-1) "
        "quantify the contrast: water passes through a pure epoxy pack "
        "in under three minutes, but requires nearly half an hour to "
        "percolate through ESP-T under identical conditions. For "
        "dodecane the situation inverts: pure epoxy passes oil in just "
        "over 15 min; ESP-T does so in barely 5 min, a 66% reduction. "
        "The hydrophobic surface repels water from inter-particle pores "
        "while allowing oil to spread freely. In a formation producing "
        "both phases, this selectivity favors oil flow toward the "
        "wellbore while suppressing water breakthrough."
    )

# [112] Release: connect to material properties
p113 = doc.paragraphs[113]
t113 = p113.text
prefix = (
    "With the material's structure, thermal stability, wettability, "
    "and transport selectivity established, we examine how it releases "
    "tracer under reservoir-relevant conditions. "
)
if not t113.startswith(prefix[:20]):
    for r in p113.runs: r.text = ''
    p113.runs[0].text = prefix + (
        "Reservoir temperature increases with depth, and because tracer "
        "release from the epoxy matrix is diffusion-controlled, "
        "temperature directly affects the release rate. Figure 3-7(a) "
        "shows release profiles of ESP-T in dodecane at 30, 60, 90, "
        "and 120 degC, plotted as C/C0 versus time. Release accelerates "
        "with temperature at all time points. The release rate is "
        "highest during the first 24 hours and decays gradually "
        "(Figure 3-7c). At 120 degC, cumulative release over 14 days "
        "remains well above the ICP-MS detection threshold, confirming "
        "that ESP-T sustains a measurable tracer signal over timescales "
        "relevant to long-term monitoring."
    )

# [133] Modeling: connect to release kinetics
p134 = doc.paragraphs[134]
t134 = p134.text
prefix2 = (
    "The release kinetics characterized above describe how tracer leaves "
    "the proppant. The next question is how that released tracer travels "
    "to the wellhead and what its concentration history reveals about "
    "the production rate. "
)
if not t134.startswith(prefix2[:20]):
    for r in p134.runs: r.text = ''
    p134.runs[0].text = prefix2 + (
        "Two processes contribute to the shape observed at the wellhead. "
        "The first is straightforward: tracer accumulated in the proppant "
        "pack during shut-in is swept out when the well opens, producing "
        "a pulse that rises and falls as the slug passes the sampling "
        "point. The second is slower: residual tracer continues to "
        "diffuse out of the epoxy matrix long after the main slug has "
        "passed, feeding a sustained, low-level signal into the produced "
        "oil. These contributions overlap in time and cannot be separated "
        "by inspection of the concentration curve alone."
        "\n\n"
        "We separate them by modeling transport from first principles. "
        "The one-dimensional advection-dispersion equation, "
        "dC/dt + v dC/dx = D d2C/dx2, governs tracer concentration "
        "along the production tubing, with v = 4Q/(pi d2) and D = alpha v. "
        "For the shut-in slug the classical instantaneous-injection "
        "solution [25] gives a Gaussian pulse. For sustained release "
        "the continuous-source solution gives an erfc form. We blend "
        "them with a tanh weight, w(t) = 1/2 [1 + tanh((t-t0)/sigma)], "
        "which shifts smoothly from the pulse-dominated early period "
        "to the diffusion-dominated tail. The full model, "
        "C(t) = cb + w(t) C_rise + [1-w(t)] C_fall, is Eq. (2). "
        "The parameter sigma is fixed at the sampling interval (4 min); "
        "the remaining six parameters are discussed below."
    )

# [154] Two-phase: connect to single-phase
p154 = doc.paragraphs[154]
t154 = p154.text
if not t154.startswith('The single-phase results'):
    for r in p154.runs: r.text = ''
    p154.runs[0].text = (
        "The single-phase results validate the model under idealized "
        "conditions. A fractured well, however, does not produce single-"
        "phase oil indefinitely. Formation water eventually breaks "
        "through, and the produced fluid becomes a mixture. We therefore "
        "examined whether the analysis remains valid when both oil and "
        "water flow through the proppant pack, using steady-state core "
        "displacement experiments at three oil-water ratios "
        "(OWR = 4:1, 1:1, 1:4) and four total flow rates "
        "(0.1-0.4 mL/min). Direct observation of the effluent at "
        "OWR = 1:1 showed a higher oil fraction early in the "
        "displacement, consistent with the oil-permeable character "
        "of ESP-T documented in Section 3.5."
    )

out = f'四氧化三铁环氧树脂拟合/ESP-T_final_{int(time.time())}.docx'
doc.save(out)
print(f'Saved: {out}')
print('Introduction rewritten, section transitions added.')