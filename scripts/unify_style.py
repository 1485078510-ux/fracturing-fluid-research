#!/usr/bin/env python3
"""Unify natural voice across remaining sections."""
from docx import Document
doc = Document('四氧化三铁环氧树脂拟合/ESP-T_终稿.docx')

# ═══════════════════════════════════════════════════════════════════
# SECTION 3.1-3.6: Results — naturalize language
# ═══════════════════════════════════════════════════════════════════

# [59] SEM low mag
p = doc.paragraphs[59]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "Figure 3-1 presents SEM micrographs of pure epoxy microspheres and "
    "ESP-T. At low magnification (panels a and d), both samples show "
    "excellent sphericity and uniformity, with no inter-particle "
    "agglomeration. Well-formed microspheres are obtained reproducibly "
    "with and without nano-Fe3O4@SA, confirming that the emulsion "
    "polymerization process tolerates the nanofiller without disruption. "
    "The pure epoxy surface appears relatively smooth; ESP-T, in "
    "contrast, exhibits a uniformly rough texture, the first indication "
    "that the nanofiller has been incorporated and has altered the "
    "surface morphology."
)
print('[59]')

# [61] SEM medium mag
p = doc.paragraphs[61]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "At medium magnification (panels b and e), the differences sharpen. "
    "Pure epoxy microspheres show subtle wrinkles and depressions, "
    "features typical of volume shrinkage during thermoset curing. "
    "ESP-T microspheres are covered with micron- and submicron-scale "
    "protrusions distributed as discrete island-like structures rather "
    "than large agglomerates. This morphology indicates that "
    "nano-Fe3O4@SA disperses as nanoclusters, not as individual "
    "nanoparticles, within the epoxy matrix, and that the stearic acid "
    "modification has successfully suppressed uncontrolled macroscopic "
    "phase separation."
)
print('[61]')

# [63] SEM high mag
p = doc.paragraphs[63]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "At high magnification (panels c and f), the fine structure becomes "
    "clear. Pure epoxy displays a dense, undulating lava-like surface "
    "characteristic of cross-linked polymer networks. ESP-T shows a "
    "debris-like morphology corresponding to stearic acid-wrapped "
    "nanoclusters. Critically, no interfacial cracks are visible between "
    "these nanoclusters and the underlying epoxy. The nanofillers appear "
    "embedded rather than merely attached to the surface. This tight "
    "integration is direct evidence that the long alkyl chains of stearic "
    "acid form physical entanglements and hydrophobic interactions with "
    "the epoxy molecular chains, producing robust interfacial adhesion."
)
print('[63]')

# [65] SEM mechanism
p = doc.paragraphs[65]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "The sequence of images suggests the following formation pathway. "
    "Hydrophobically modified nano-Fe3O4@SA nanoclusters disperse "
    "uniformly within epoxy droplets after emulsification. As cross-"
    "linking proceeds, the rigidifying polymer network immobilizes them "
    "in place. During curing shrinkage, a fraction of the nanoclusters "
    "is extruded toward the surface, producing the surface-enriched "
    "rough nanostructures visible in panels c and f. We note that this "
    "mechanism is inferred from post-cure microscopy; real-time in-situ "
    "characterization would be required to confirm the kinetic pathway "
    "directly."
)
print('[65]')

# [70] SEM mapping
p = doc.paragraphs[70]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "Elemental mapping (Figure 3-2) corroborates the structural evidence. "
    "The Fe signal is distributed throughout the entire particle cross-"
    "section, confirming that nano-Fe3O4@SA resides within the epoxy "
    "matrix, not merely on the surface. Minor inhomogeneities in signal "
    "intensity suggest the presence of distinct phase domains, consistent "
    "with the nanocluster dispersion observed in SEM. Together, the "
    "imaging and elemental data confirm successful synthesis of the "
    "ESP-T proppant."
)
print('[70]')

# [75] TGA - keep technical but flowing
p = doc.paragraphs[75]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "Thermal analysis (Figure 3-3) reveals three decomposition stages. "
    "The first, from 50 to 350 degC, involves a modest weight loss of "
    "5.70%, attributable to adsorbed water and residual ethanol from "
    "synthesis. The second stage, 350-400 degC, is the primary "
    "decomposition event: cleavage of C-O-C and C-C bonds in the epoxy "
    "network releases CO2 and small hydrocarbons, with the DTG curve "
    "peaking at 357.27 degC and accounting for 72.5% of the total mass "
    "loss. Above 400 degC the residual mass stabilizes; the remaining "
    "weight loss comes from oxidative combustion of the carbonaceous "
    "residue. The final residue consists of hollow glass microspheres, "
    "thermally stable nano-Fe3O4@SA, and a minor carbonaceous fraction. "
    "A brief DSC endotherm accompanies the decomposition, initiating "
    "at 332.41 degC and peaking at 357.27 degC, consistent with the "
    "DTG profile."
)
print('[75]')

# [79] TGA significance
p = doc.paragraphs[79]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "The initial decomposition temperature of 357.27 degC is far above "
    "the temperature range encountered in oil and gas wells (typically "
    "80-150 degC, reaching 200 degC in deep wells). Below 200 degC the "
    "material loses only trace moisture, with no detectable degradation "
    "of the epoxy matrix. This thermal margin confirms that the epoxy "
    "matrix remains structurally intact across the full operational "
    "temperature range of oil and gas wells."
)
print('[79]')

# [84] WCA
p = doc.paragraphs[84]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "Water contact angle measurements (Figure 3-4) quantify the change "
    "in surface character produced by the nanofiller. Pure epoxy "
    "microspheres give an average WCA of 72.3 deg, a weakly hydrophilic "
    "surface consistent with the hydroxyl groups present on epoxy chains. "
    "ESP-T yields an average WCA of 104.6 deg, a 32.3 deg increase that "
    "crosses the hydrophobic threshold. The stearic acid modification "
    "drives this transition: the carboxyl groups coordinate with surface "
    "hydroxyls on the nano-Fe3O4@SA, while the long alkyl chains orient "
    "outward from the epoxy matrix, constructing a hydrophobic film. "
    "The surface enrichment of nanoclusters observed in SEM amplifies "
    "the effect by concentrating these alkyl chains at the proppant "
    "surface."
)
print('[84]')

# [91-92] Density
p = doc.paragraphs[91]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "Figure 3-5 compares the bulk and apparent densities of the two "
    "materials. Pure epoxy microspheres have a bulk density of "
    "0.6179 g/cm3 and an apparent density of 1.02 g/cm3; ESP-T gives "
    "0.646 and 1.072 g/cm3, respectively. The increase reflects the "
    "higher density of nano-Fe3O4@SA (approximately 5.18 g/cm3) "
    "relative to epoxy (roughly 1.1 g/cm3). Both proppants have a bulk "
    "density below that of water, meaning they can be suspended in "
    "water-based fracturing fluids without additional suspending agents. "
    "The 0.05 g/cm3 increase in apparent density is consistent with "
    "nanoclusters filling internal pores, yielding a denser particle "
    "structure. ESP-T meets key industry specifications: sphericity "
    "and roundness exceed 0.9 (Krumbien-Sloss chart [26]), acid "
    "solubility is 3.3% (well below the 5% standard), and the crush "
    "rate at 52 MPa is 2.9%, comparable to pure epoxy microspheres "
    "(2.6%). The slight increase in crush rate is attributable to "
    "the hollow glass microspheres, which are inherently more crushable "
    "than the dense epoxy matrix; the epoxy itself retains its "
    "structural integrity after nano-Fe3O4@SA incorporation."
)
# Clear [92] since content merged into [91]
p92 = doc.paragraphs[92]
for r in p92.runs: r.text = ''
print('[91-92] merged')

# [97] Conductivity
p = doc.paragraphs[97]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "The oil and water transport behavior of the proppant pack was "
    "assessed indirectly by measuring the time required for each fluid "
    "to pass through a packed bed (Figure 3-6, Table 3-1). Pure epoxy "
    "microspheres pass water in 2 min 53 s; ESP-T requires 28 min 41 s, "
    "a nearly tenfold increase. The hydrophobic surface impedes water "
    "from spreading through the inter-particle pores. The trend reverses "
    "for oil: pure epoxy passes dodecane in 15 min 11 s, while ESP-T "
    "does so in 5 min 11 s, a 66% reduction. The hydrophobic surface "
    "is compatible with the oil phase, allowing rapid spreading and "
    "low flow resistance. ESP-T thus functions as a water-resistant, "
    "oil-permeable medium: in a formation containing both phases, the "
    "proppant pack favors oil flow toward the wellbore while restricting "
    "water, a characteristic that helps mitigate water channeling."
)
print('[97]')

# [105] Release intro
p = doc.paragraphs[105]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "Reservoir temperature increases with depth, and because the release "
    "of tracer from the epoxy matrix is a diffusion-controlled process, "
    "temperature directly affects the release rate. Figure 3-7(a) shows "
    "the release profiles of ESP-T in dodecane at 30, 60, 90, and "
    "120 degC, plotted as normalized concentration C/C0 versus time. "
    "Release accelerates with temperature at all time points. The "
    "release rate is highest during the first 24 hours of fluid contact "
    "and then decays gradually (Figure 3-7c). At 120 degC, cumulative "
    "release over 14 days remains well above the ICP-MS detection "
    "threshold, confirming that ESP-T can sustain a measurable tracer "
    "signal over timescales relevant to long-term monitoring."
)
print('[105]')

# [116] K-P results
p = doc.paragraphs[116]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "The K-P model was fitted to the release data at each temperature "
    "(Figure 3-7b, Table 3-2). The rate constant K increases "
    "systematically from 0.055 at 30 degC to 0.196 at 120 degC, "
    "consistent with the thermally activated nature of both diffusion "
    "and polymer relaxation. All fitted n values fall within 0.45-0.85, "
    "the range corresponding to anomalous transport co-governed by "
    "Fickian diffusion and Case-II relaxation. The R2 exceeds 0.94 at "
    "every temperature. We note that the K-P model, as a power-law "
    "formulation, lacks an upper asymptote and is strictly valid only "
    "for the early-to-intermediate release regime (Mt/Minf < 0.6). The "
    "extrapolated C/C0 values at 14 days exceed unity, reflecting this "
    "inherent limitation; the reported R2 values therefore pertain to "
    "the 0-14 day fitting range and should not be extrapolated further."
)
print('[116]')

# [118] Release mechanism
p = doc.paragraphs[118]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "The physical picture that emerges is one of solvent-driven swelling. "
    "Dodecane molecules permeate into the cross-linked epoxy network, "
    "generating a dual-state structure with an inner glassy core and an "
    "outer gel layer. Swelling reduces polymer chain entanglement, "
    "creating expanded transport channels through which the tracer "
    "diffuses into the external medium. Higher temperature enhances "
    "solvent permeability, accelerates swelling, and weakens "
    "intermolecular forces within the network, enlarging the pore "
    "dimensions available for tracer diffusion. This mechanism accounts "
    "for both the temperature dependence of K and the mixed "
    "Fickian/Case-II character reflected in the n values."
)
print('[118]')

# ═══════════════════════════════════════════════════════════════════
# CONCLUSIONS — unify with Section 3.7 voice
# ═══════════════════════════════════════════════════════════════════

p = doc.paragraphs[158]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "The central challenge addressed in this work is the disconnect "
    "between tracer release kinetics, measured in batch experiments "
    "and fitted with empirical models, and the quantitative "
    "interpretation of tracer breakthrough signals at the wellhead. "
    "We approached this by developing a two-component transport model "
    "that decomposes the breakthrough curve into a Gaussian pulse "
    "and an erfc tail, linked by a smooth tanh transition, and by "
    "designing an oleophilic epoxy/Fe3O4 tracer proppant (ESP-T) "
    "to serve as the validation platform."
)
print('[158]')

p = doc.paragraphs[160]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "On the modeling side, the two-component analysis achieves "
    "R2 = 0.9939 for single-phase breakthrough curves, with the "
    "erfc tail accounting for 47% of the integrated signal. This "
    "quantitatively establishes that sustained matrix-diffusion-"
    "controlled release, not the initial shut-in pulse, dominates "
    "the long-term monitoring signal. The fitted flow rate "
    "(0.46 mL/min) agrees within 8% of the independently set pump "
    "value (0.50 mL/min); the fitted residence time (37.4 min) "
    "matches the calculated travel time (38.6 min); and the Peclet "
    "number (0.934) independently corroborates the non-Fickian "
    "mechanism inferred from Korsmeyer-Peppas kinetics. Under steady-"
    "state two-phase flow, tracer flux scales predictably with the "
    "oil-water ratio and is independent of total flow rate, providing "
    "a pathway from tracer concentration data to per-interval oil "
    "production rates."
)
print('[160]')

p = doc.paragraphs[162]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "The ESP-T proppant, designed as the enabling platform for this "
    "model, exhibits a performance profile suited to downhole "
    "deployment. Stearic acid modification imparts oleophilic "
    "character: the water contact angle rises from 72.3 deg to "
    "104.6 deg, oil filtration time drops by 66%, and water filtration "
    "time increases nearly tenfold. Thermal analysis confirms "
    "structural integrity to 357 degC, with only trace moisture loss "
    "below 200 degC. The bulk density (0.646 g/cm3) remains below "
    "that of water, enabling self-suspension in fracturing fluids, "
    "while crush rate (2.9% at 52 MPa) and acid solubility (3.3%) "
    "meet established benchmarks for ultra-lightweight proppants."
)
print('[162]')

p = doc.paragraphs[164]
for r in p.runs: r.text = ''
p.runs[0].text = (
    "Taken together, the material and modeling contributions establish "
    "a framework in which sustained tracer release, governed by matrix-"
    "diffusion kinetics, is quantitatively linked to the observed "
    "breakthrough signal and, ultimately, to per-interval production "
    "rates. The framework is applicable to acid fracturing, deep-well, "
    "and high-pressure scenarios in unconventional reservoirs. The "
    "present validation is limited to laboratory-scale single-interval "
    "experiments with dodecane as the model oil. Extending the approach "
    "to field-scale, multi-interval, multiphase conditions with crude "
    "oil, and establishing the corresponding operational reliability, "
    "represents the necessary next step toward practical deployment."
)
print('[164]')

doc.save('四氧化三铁环氧树脂拟合/ESP-T_终稿.docx')
print('\nFull paper voice unified.')