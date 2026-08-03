# -*- coding: utf-8 -*-
"""Task 10: write the Abstract and Section 1 (Introduction) into the v2 manuscript.

Content per task-10 brief / design doc §Abstract and §1:

- Abstract (~250 words): 6-beat structure (problem -> gap -> approach ->
  decisive result -> supporting results -> significance). No citations.
- Introduction: 4 paragraphs (~1200 words total).
  P1 Engineering need (per-stage contribution, tool trade-offs, tracer
     proppants) - cites [5,6,9,10,11,12].
  P2 Core difficulty (BTC encodes release + transport; K-P and ADE applied
     separately; underdetermined inverse problem) - cites [28,29,31,32].
  P3 Prior work ceiling (Stream A materials [21-24], Stream B ADE
     interpretation [16,33,34], TOA errors +162% to +685%, the gap).
  P4 This work (coupled Gaussian + erfc + tanh model; predictive validation,
     Q unconstrained; ESP-T vehicle; preview of Q +/-8%, Pe <-> K-P n,
     two-phase r = 0.97; scope).

Conventions identical to Tasks 2-9 (write_section2.py et al.):
style Normal, explicit non-bold/non-italic runs, Unicode subscripts
(Fe3O4, t0), spaced em dashes, en-dash ranges (10-5000, 0.45-0.85),
literal-underscore subscripts (Q_fit, Q_pump), bracket citation markers
[5,6] in Elsevier no-space style.
"""
import sys

sys.stdout.reconfigure(encoding="utf-8")
from docx import Document

TGT = "四氧化三铁环氧树脂拟合/ESP-T_投稿文件/ESP-T_v2_manuscript.docx"

ABSTRACT = (
    "Interpreting tracer-proppant wellhead breakthrough curves (BTCs) for per-stage "
    "production allocation requires jointly accounting for sustained release from the "
    "polymer matrix and transport through the production tubing — a coupled inverse "
    "problem not addressed by current practice. Existing approaches apply the "
    "Korsmeyer–Peppas (K-P) model for batch release characterization and the "
    "advection–dispersion framework for transport interpretation separately; the "
    "sustained source term inherent to tracer proppants has not been coupled to the "
    "transport solution. We develop a coupled release–transport model that "
    "decomposes the BTC into a Gaussian pulse (shut-in accumulation slug) and an erfc "
    "tail (sustained matrix-diffusion-controlled release), linked by a smooth "
    "hyperbolic-tangent transition; six parameters are estimated simultaneously from a "
    "single BTC, and the flow rate Q receives no constraint in the objective function. "
    "Applied to an oleophilic epoxy/Fe₃O₄ tracer proppant (ESP-T) in core "
    "displacement experiments, the model yields Q_fit = 0.46 mL/min against the "
    "independent pump setting of 0.50 mL/min (deviation 8%), with Q confined only by "
    "wide search bounds (10–5000 mL/min) and no prior directing it toward the "
    "correct value. The Peclet number (Pe = 0.934) independently corroborates the "
    "non-Fickian transport mechanism identified via K-P kinetics (n = 0.45–0.85); "
    "under two-phase flow, the oil-phase tracer mass flux eliminates water-dilution "
    "artifacts and tracks oil production rates (Pearson r = 0.97, RMSD = 8.3%). The "
    "framework enables per-interval production allocation from wellhead samples alone, "
    "requiring no downhole tools and only a single shut-in."
)

INTRO_P1 = (
    "Multi-stage hydraulic fracturing has made horizontal wells the dominant "
    "production technology for tight and shale reservoirs — a single well may "
    "comprise dozens of fracture stages — but completion optimization and "
    "well-spacing decisions hinge on a quantity the wells do not directly report: "
    "the production contribution of each individual fracture stage, the basic input "
    "to completion and spacing decisions [5,6]. Downhole "
    "tools provide partial answers, each with characteristic trade-offs. Production "
    "logging quantifies stage-level inflow profiles but requires a wireline or "
    "coiled-tubing intervention, disturbs the near-wellbore flow field, and records a "
    "snapshot in time rather than a continuous record of a system that evolves over "
    "months [9]. Distributed fiber-optic sensing — both acoustic (DAS) and "
    "temperature (DTS) — monitors the well continuously, but it requires a "
    "permanently installed cable at substantial cost, and its diagnostics, such as "
    "stimulation-hit detection and fracture-geometry inference, do not directly measure "
    "produced flow per stage [10,11]. Microseismic imaging maps the induced fracture "
    "network during stimulation but says little about which fractures subsequently "
    "produce [12]. Tracer proppants avoid downhole hardware altogether: tracer-bearing "
    "particles are injected with the proppant pack, remain permanently placed in the "
    "fracture, and are sampled at the surface by collecting produced fluid at the "
    "wellhead. Their signal persists for months to years, enabling long-term per-stage "
    "surveillance without intervention — a scalable route to the stage-contribution "
    "data that completion and spacing decisions require [5,6]."
)

INTRO_P2 = (
    "The information needed to quantify stage contributions is encoded in the tracer "
    "breakthrough curve (BTC) — the time series of tracer concentration measured at "
    "the wellhead. A tracer-proppant BTC superimposes two processes: the kinetics with "
    "which tracer is released from the polymer matrix, and the advection–dispersion "
    "transport of the released tracer from the fracture through the wellbore to the "
    "sampling point. Both contribute parameters of interest, yet both are observed "
    "through a single concentration signal, and neither can be isolated without "
    "assumptions about the other. Current practice treats the two processes separately. "
    "Release characterization applies the zero-dimensional Korsmeyer–Peppas (K-P) "
    "power law to batch immersion tests, in which a proppant specimen is soaked in "
    "fluid and the cumulative release fraction is fitted as a function of time [28,29]. "
    "Transport interpretation applies the advection–dispersion equation (ADE), "
    "whose analytical solutions assume a known source term — an instantaneous "
    "pulse of specified mass released at a known time [31,32]. Applied sequentially, "
    "these two steps can identify which stages produce: signal presence and arrival "
    "order indicate producing intervals. They cannot quantify how much each stage "
    "produces, because the measured concentration reflects an unknown convolution of "
    "the release-rate history and the transport response; recovering two unknown "
    "functions from one observable C(t) is an underdetermined inverse problem. The "
    "missing step is a framework that couples the release source term to the transport "
    "solution and estimates both parameter sets jointly from the BTC itself."
)

INTRO_P3 = (
    "Two streams of prior work approach this problem from opposite ends, and neither "
    "closes the gap. The first stream develops tracer-proppant materials. Zhao et al. "
    "fabricated dye-eluting ceramic proppants and demonstrated tracer elution for "
    "fracture characterization [21]. Zhou et al. coated proppants with self-suspension "
    "and slow-release layers, tailoring the release rate through coating chemistry "
    "[22]. Li et al. prepared polymer-coated proppants doped with rare-earth tracers "
    "and characterized their release kinetics in batch experiments [23]. Gong et al. "
    "reported an oleophilic Fe₃O₄–polymer tracer proppant for monitoring "
    "oil production contribution [24]. All four evaluate release performance "
    "exclusively through batch immersion tests fitted with the K-P power law; none "
    "couples the measured release kinetics to a transport solution, and none "
    "interprets a flow-through BTC generated by the material itself. These batch "
    "assessments are also inherently static — the tracer accumulates in a "
    "quiescent fluid — whereas in a producing well the released tracer is "
    "advected away continuously, so batch-derived release parameters cannot be "
    "transferred directly to a flow-through interpretation. The second stream "
    "advances tracer interpretation within the advection–dispersion framework. "
    "Fontalvo et al. developed a physical interpretation of interwell partitioning "
    "tracer tests in heterogeneous reservoirs [16]. Velasco-Lozano et al. modeled "
    "chemical tracer transport under two-phase flow in advective-dominated porous "
    "media [33]. Mazo et al. formulated a mathematical model of water- and oil-soluble "
    "tracer transfer in multi-stage fractured wells [34]. These studies treat "
    "transport rigorously, but all assume a known source term — a pulse of "
    "specified mass and duration injected at a known location. When the source is "
    "instead the sustained, slowly exhausting release of a polymer matrix, whose rate "
    "is neither known nor constant, these solutions cannot be applied directly. Field "
    "practice for sustained-release tracers consequently falls back on time-of-arrival "
    "(TOA) analyses — peak time, half-peak time, first moment — which discard "
    "most of the information carried by the BTC. For the breakthrough curves "
    "considered in this work, TOA flow-rate estimates deviate by +162% to +685% "
    "(Section 4.3). In short, the materials stream characterizes release without "
    "transport, and the transport stream assumes the release away; no existing "
    "framework couples a sustained, unknown release source term to the transport "
    "solution."
)

INTRO_P4 = (
    "This paper closes that gap with a coupled release–transport model. The model "
    "decomposes a tracer-proppant BTC into two physically motivated components: a "
    "Gaussian pulse produced by the shut-in accumulation slug — tracer that "
    "diffuses from the matrix during the shut-in, accumulates in the pore space of the "
    "pack, and is swept to the sampling point as a coherent slug when flow resumes "
    "— and an erfc tail produced by sustained matrix-diffusion-controlled release "
    "that persists long after the slug has passed. The two components are joined by a "
    "smooth hyperbolic-tangent transition, giving six parameters that are estimated "
    "simultaneously from a single BTC (Section 2). The validation strategy is "
    "predictive, not descriptive. The flow rate Q is a physical quantity with an "
    "independently known value — the pump setting of the displacement experiment "
    "(Section 3.4) — yet it receives no constraint in the objective function; the "
    "model must recover the correct value from the shape of the curve alone. A "
    "structurally wrong model can always improve its residuals by adjusting "
    "parameters, but it cannot spontaneously converge to a correct independent "
    "physical quantity: if Q converges to the pump setting across four independent "
    "global searches, the model structure must capture the transport physics. The "
    "experimental vehicle is ESP-T, an oleophilic epoxy/Fe₃O₄ tracer proppant "
    "synthesized in-house (Section 3). Three results preview the argument of Section "
    "4: the fitted flow rate reproduces the independently set pump rate within 8% "
    "with Q entirely unconstrained; the Peclet number (Pe = 0.934) independently "
    "corroborates the non-Fickian release mechanism established by K-P batch kinetics "
    "(n = 0.45–0.85); and under two-phase flow the oil-phase tracer mass flux "
    "tracks oil production rates with a Pearson correlation of 0.97. The scope "
    "progresses from single-phase core floods to two-phase production allocation to a "
    "field deployment pathway (Sections 5 and 6), establishing a route to per-stage "
    "production allocation from wellhead samples alone, requiring no downhole tools "
    "and only a single shut-in."
)

doc = Document(TGT)


def find(text, paras, starts=False):
    """Return index of paragraph matching TEXT within the given PARAS snapshot."""
    for i, p in enumerate(paras):
        t = p.text.strip()
        if (t.startswith(text) if starts else t == text):
            return i
    raise RuntimeError(f"not found: {text!r}")


def fill_section(heading, anchor_text, body):
    """Remove empty spacers between HEADING and ANCHOR_TEXT, then insert BODY before anchor.

    The paragraph list must be passed explicitly to find(): a module-global snapshot
    becomes stale once paragraphs are removed/inserted (pitfall documented in
    write_section6.py).
    """
    paras = doc.paragraphs  # fresh snapshot
    i_head = find(heading, paras)
    i_anchor = find(anchor_text, paras, starts=True)
    assert i_anchor > i_head, f"anchor before heading for {heading}"
    removed = 0
    for i in range(i_head + 1, i_anchor):
        if paras[i].text.strip() == "":
            paras[i]._p.getparent().remove(paras[i]._p)
            removed += 1
    print(f"[{heading}] removed {removed} spacer paragraph(s)")
    paras = doc.paragraphs
    i_anchor = find(anchor_text, paras, starts=True)
    for text in body:
        p = paras[i_anchor].insert_paragraph_before(text)
        p.style = doc.styles["Normal"]
        for r in p.runs:
            r.font.bold = False
            r.font.italic = False
    print(f"[{heading}] inserted {len(body)} paragraph(s) before {anchor_text!r}")


# Abstract content between the "Abstract" heading and the "Keywords" heading.
fill_section("Abstract", "Keywords", [ABSTRACT])

# Introduction content between the "1. Introduction" heading and the Section 2 heading.
fill_section("1. Introduction", "2. Coupled Release–Transport Model",
             [INTRO_P1, INTRO_P2, INTRO_P3, INTRO_P4])

doc.save(TGT)
print("Saved:", TGT)

counts = {
    "Abstract": len(ABSTRACT.split()),
    "Intro P1": len(INTRO_P1.split()),
    "Intro P2": len(INTRO_P2.split()),
    "Intro P3": len(INTRO_P3.split()),
    "Intro P4": len(INTRO_P4.split()),
}
total = 0
for k, v in counts.items():
    print(f"{k}: {v} words")
    total += v
print(f"Introduction total: {counts['Intro P1'] + counts['Intro P2'] + counts['Intro P3'] + counts['Intro P4']} words")
