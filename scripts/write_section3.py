# -*- coding: utf-8 -*-
"""
Task 3: Write Section 3 — Experimental Methods (compressed)
into ESP-T_v2_manuscript.docx (skeleton produced by Task 1, Section 2 by Task 2).

Writes content into the blank paragraph that follows each Section 3 heading:
  - lead-in paragraph after the Section 3 H1
  - 3.1 ESP-T Synthesis (one paragraph, ~150-170 words)
  - 3.2 Material Prerequisites for Model Validity (lead text + real Word Table 2
    + closing note; the [Fig. 2] marker paragraph is left untouched)
  - 3.3 K-P Batch Release Kinetics
  - 3.4 Core Displacement: BTC Generation
  - 3.5 Parameter Estimation Strategy (engineering language, two-pass strategy)

Conventions (matching scripts/write_section2.py):
  - '_xyz' in content strings renders as a Word subscript run (M_t -> Mt).
  - The [Fig. 2] marker paragraph in 3.2 is left untouched (replaced later).
  - Table captions: bold "Table N." + regular text; "Table Grid" style.
"""
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX = r"C:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP-T_投稿文件\ESP-T_v2_manuscript.docx"

SUB_RE = re.compile(r"_([A-Za-z0-9]+)")


def add_runs(para, text, bold=False, italic=False):
    """Add runs to *para*; '_xyz' segments become subscript runs."""
    pos = 0
    for m in SUB_RE.finditer(text):
        if m.start() > pos:
            r = para.add_run(text[pos:m.start()])
            r.bold = bold
            r.italic = italic
        r = para.add_run(m.group(1))
        r.font.subscript = True
        r.bold = bold
        r.italic = italic
        pos = m.end()
    if pos < len(text):
        r = para.add_run(text[pos:])
        r.bold = bold
        r.italic = italic


def write_into(para, text):
    """Write *text* into an existing (blank) paragraph."""
    add_runs(para, text)
    return para


def insert_before(anchor, text, center=False):
    """Insert a new paragraph immediately before *anchor*; return it."""
    p = anchor.insert_paragraph_before()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(p, text)
    return p


def main():
    doc = Document(DOCX)
    paras = doc.paragraphs

    # Idempotency guard: refuse to run twice on the same document.
    if any("was synthesized in three steps" in p.text for p in paras):
        raise SystemExit("Section 3 already written; refusing to duplicate content. "
                         "Restore the skeleton from git first (git checkout -- <docx>).")

    def find_heading(text):
        for i, p in enumerate(paras):
            if p.text.strip() == text:
                return i
        raise SystemExit(f"heading not found: {text!r}")

    i_3 = find_heading("3. Experimental Methods")
    i_31 = find_heading("3.1 ESP-T Synthesis")
    i_32 = find_heading("3.2 Material Prerequisites for Model Validity")
    i_33 = find_heading("3.3 K-P Batch Release Kinetics")
    i_34 = find_heading("3.4 Core Displacement: BTC Generation")
    i_35 = find_heading("3.5 Parameter Estimation Strategy")

    # --- Lead-in after Section 3 heading ------------------------------------
    write_into(paras[i_3 + 1], (
        "This section describes the synthesis and characterization of the tracer "
        "proppant, the batch release experiments that establish its release mechanism, "
        "the core displacement protocol that generated the breakthrough curves, and the "
        "two-pass strategy used to estimate model parameters. Procedural detail beyond "
        "these essentials is deferred to the Supplementary Material."
    ))

    # --- 3.1 ESP-T Synthesis (one paragraph) ---------------------------------
    write_into(paras[i_31 + 1], (
        "ESP-T was synthesized in three steps (Fig. S1). (i) Stearic acid-modified "
        "nano-Fe₃O₄ (nano-Fe₃O₄@SA) was prepared by co-precipitation: 2.703 g FeCl₃ "
        "and 1.15 g FeCl₂·4H₂O were dissolved in 100 mL deionized water at 80 °C under "
        "N₂ purge, with 2×10⁻⁵ mol MnCl₂·6H₂O added as a dopant; 5.5 mL NH₃·H₂O was "
        "then added and the reaction proceeded at pH 10 for 2 h. The precipitate was "
        "magnetically separated, washed, and sonicated with ethanolic stearic acid for "
        "oleophilic surface modification. (ii) A pre-mixture of 20 mL E51 epoxy resin, "
        "~0.75 g nano-Fe₃O₄@SA (~3.3 wt%), 1 g hollow glass microspheres, and 7 g T31 "
        "curing agent was homogenized. (iii) The pre-mixture was emulsified in a "
        "SiO₂/guar gum aqueous dispersion, cured at 50 °C for 1 h, then rinsed and "
        "dried at 80 °C for 10 h. Pure epoxy microspheres were prepared identically "
        "without nano-Fe₃O₄@SA as a reference. The co-precipitation method also "
        "accommodates other transition metals (Zn, Cu, Co, Ni) and rare earths "
        "(Eu, Dy, Nd) for multi-stage coding. Reagent specifications are listed in "
        "Table S1."
    ))

    # --- 3.2 Material Prerequisites for Model Validity -----------------------
    write_into(paras[i_32 + 1], (
        "The model of Section 2 rests on four material prerequisites: bulk "
        "encapsulation of the tracer within the particle, thermal stability at "
        "downhole temperatures, non-Fickian release kinetics, and oil-phase "
        "selectivity. Each was verified by an independent measurement (Fig. 2) before "
        "any breakthrough curve was interpreted; the results and their implications "
        "for the model are summarized in Table 2."
    ))

    fig2_marker = paras[i_32 + 2]
    if not fig2_marker.text.startswith("[Fig. 2 placeholder"):
        raise SystemExit(f"unexpected paragraph after 3.2 heading: {fig2_marker.text!r}")

    caption = fig2_marker.insert_paragraph_before()
    r_cap = caption.add_run("Table 2.  ")
    r_cap.bold = True
    caption.add_run("Material prerequisites of the coupled release–transport model and "
                    "the evidence supporting each.")

    rows = [
        ("Prerequisite", "Method", "Result", "Implication for Model"),
        ("Tracer in matrix", "SEM-EDS cross-section",
         "Fe distributed throughout particle",
         "Confirms bulk encapsulation → sustained source term valid"),
        ("Thermal stability", "TGA/DTG (air, 10 °C/min)",
         "Decomposition onset 357 °C",
         "No degradation at 80–200 °C downhole → K-P kinetics valid"),
        ("Non-Fickian release", "K-P batch release, M_t/M∞ < 0.6",
         "n = 0.45–0.85",
         "Swelling-diffusion mechanism → erfc tail physically justified"),
        ("Oil-phase selectivity", "WCA + packed-bed filtration",
         "WCA 104.6°; oil/water time ratio 5.53",
         "Tracer signal isolated from water dilution → flux method valid"),
    ]
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    try:
        tbl.style = "Table Grid"
    except Exception:
        pass
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.cell(i, j)
            p = cell.paragraphs[0]
            add_runs(p, val, bold=(i == 0))
    caption._p.addnext(tbl._tbl)  # move table from doc end to after caption

    insert_before(fig2_marker, (
        "Detailed characterization protocols are provided in Supplementary Material "
        "S2, and complete physical property data in Table S4."
    ))

    # --- 3.3 K-P Batch Release Kinetics --------------------------------------
    write_into(paras[i_33 + 1], (
        "ESP-T (5 g, 40–70 mesh) was immersed in 100 mL dodecane in sealed glass vials "
        "held in thermostatic oil baths at 30, 60, 90, and 120 °C. Samples were "
        "withdrawn at 12-h intervals over 14 days, and the tracer metal-ion "
        "concentration was quantified by inductively coupled plasma mass spectrometry "
        "(ICP-MS, PerkinElmer NexION 300X). Release data within M_t/M∞ < 0.6 — the "
        "range in which the Korsmeyer–Peppas power law C/C₀ = K·tⁿ is valid — were "
        "fitted to obtain the kinetic parameters K and n. This validity limit excludes "
        "the late stage, in which the matrix approaches saturation and release no "
        "longer follows the power law. The fitted exponents n = 0.45–0.85 identify the "
        "anomalous (non-Fickian) release mechanism that physically justifies the erfc "
        "tail of the transport model (Section 2.2). The full release dataset is "
        "provided in Table S3."
    ))

    # --- 3.4 Core Displacement: BTC Generation --------------------------------
    write_into(paras[i_34 + 1], (
        "A dynamic displacement apparatus (Fig. S2) was used to generate the "
        "breakthrough curves interpreted in Section 4. A steel core was packed with "
        "ESP-T, sealed with 200-mesh screens, and mounted in a core holder under 5 MPa "
        "confining pressure. For the single-phase experiment, the pack was saturated "
        "with dodecane at 5 mL/min and shut in for 96 h to allow tracer accumulation — "
        "a shut-in duration representative of typical periods between fracturing and "
        "flowback in field operations. Displacement was then resumed at a pump setting "
        "of 0.50 mL/min. Effluent was sampled at 4-min intervals (2 mL per sample, "
        "21 samples) and analyzed by ICP-MS to construct the BTC. The pump setting is "
        "the independent flow-rate reference against which the fitted flow rate Q is "
        "compared in the self-calibration test of Section 4.2. For the two-phase "
        "experiments, dodecane and deionized water were co-injected at oil–water "
        "ratios (OWR) of 4:1, 1:1, and 1:4 and total flow rates of 0.1–0.4 mL/min "
        "under steady-state conditions; samples were collected at 5-min intervals, and "
        "the oil–water ratio and tracer concentration of each sample were recorded "
        "after ICP-MS analysis."
    ))

    # --- 3.5 Parameter Estimation Strategy ------------------------------------
    write_into(paras[i_35 + 1], (
        "All model parameters were estimated from the 21-point single-phase "
        "breakthrough curve by solving the inverse problem in two passes. Pass 1 "
        "(basin location) performed a global search over wide, physically bounded "
        "parameter ranges; Q was bounded to [10, 5000] mL/min — a 500-fold envelope "
        "that admits every flow rate physically achievable with the apparatus. Four "
        "independent runs with different random starting points were executed, and all "
        "four converged to the same region of parameter space, confirming that the "
        "identified basin is robust rather than an artifact of any single "
        "initialization. Pass 2 (local refinement) refined the best Pass-1 solution by "
        "gradient-based minimization, converging within 50–200 iterations and "
        "confirming that Pass 1 had located the correct basin. The identical protocol "
        "— the same bounds, convergence criteria, and number of independent runs — was "
        "applied to all five candidate models compared in Section 4.1, so that model "
        "selection rests on a fair comparison. Throughout, parameters were estimated "
        "from the concentration data alone; Q was subject to no penalty, prior, or "
        "constraint toward the pump setting (Section 2.4). Detailed parameter bounds "
        "are listed in Table S6, and the optimization implementation is described in "
        "Supplementary Material S5.3."
    ))

    doc.save(DOCX)
    print("saved:", DOCX)


def verify():
    """Re-open the document and cross-check Section 3 against the spec."""
    doc = Document(DOCX)
    paras = doc.paragraphs
    idx = [p.text.strip() for p in paras]

    start = idx.index("3. Experimental Methods")
    end = idx.index("4. Results and Discussion")
    all_ok = True

    # Subsection boundaries
    h31 = idx.index("3.1 ESP-T Synthesis")
    h32 = idx.index("3.2 Material Prerequisites for Model Validity")
    h33 = idx.index("3.3 K-P Batch Release Kinetics")
    h34 = idx.index("3.4 Core Displacement: BTC Generation")
    h35 = idx.index("3.5 Parameter Estimation Strategy")

    def words(lo, hi):
        return sum(len(p.text.split()) for p in paras[lo:hi])

    w_lead = words(start + 1, h31)
    w_31 = words(h31 + 1, h32)
    w_32prose = words(h32 + 1, h33)
    w_33 = words(h33 + 1, h34)
    w_34 = words(h34 + 1, h35)
    w_35 = words(h35 + 1, end)
    print(f"  words: lead={w_lead}  3.1={w_31}  3.2(prose)={w_32prose}  "
          f"3.3={w_33}  3.4={w_34}  3.5={w_35}")

    # --- Per-subsection paragraph-count / content checks --------------------
    sec_text = "\n".join(paras[i].text for i in range(start, end))

    # 3.1: exactly one non-empty paragraph; no material-innovation language
    body_31 = [p.text for p in paras[h31 + 1:h32] if p.text.strip()]
    n_paras_31 = len(body_31)
    innovation_terms = ["novel", "innovative", "first time", "state-of-the-art",
                        "superior", "unprecedented", "breakthrough material"]
    found_innov = [t for t in innovation_terms if any(t in b.lower() for b in body_31)]
    checks = {
        "3.1 single paragraph": n_paras_31 == 1,
        "3.1 no innovation language": not found_innov,
        "3.1 three steps (Fig. S1)": "three steps (Fig. S1)" in body_31[0],
        "3.1 Mn doping": "2×10⁻⁵ mol MnCl₂·6H₂O" in body_31[0],
        "3.1 Zn/Cu/Eu/Dy extensibility": "Zn, Cu, Co, Ni" in body_31[0]
        and "Eu, Dy, Nd" in body_31[0],
        "3.1 reagent specs -> Table S1": "Table S1" in body_31[0],
    }
    for name, passed in checks.items():
        print(f"  [{'OK' if passed else 'FAIL'}] {name}")
        all_ok = all_ok and passed

    # 3.2: Table 2 present with 4 data rows; caption; cross-references
    n_tables = len(doc.tables)
    t2 = doc.tables[1] if n_tables > 1 else None
    t2_ok = t2 is not None and len(t2.rows) == 5 and len(t2.columns) == 4
    t2_header = list(t2.rows[0].cells[j].text for j in range(4)) if t2 is not None else []
    checks = {
        "Table 2 exists (5x4)": t2_ok,
        "Table 2 header": t2_header == ["Prerequisite", "Method", "Result",
                                        "Implication for Model"],
        "Table 2 caption": "Table 2.  Material prerequisites of the coupled "
                           "release–transport model" in sec_text,
        "3.2 references Fig. 2": "Fig. 2" in sec_text,
        "3.2 -> SM S2": "Supplementary Material S2" in sec_text,
        "3.2 -> Table S4": "Table S4" in sec_text,
        "3.2 four prerequisites in table": all(
            celltext in [t2.rows[r].cells[0].text for r in range(1, 5)]
            for celltext in ["Tracer in matrix", "Thermal stability",
                             "Non-Fickian release", "Oil-phase selectivity"]),
        "Table 2 not labeled Table 1": "Table 2." in sec_text
        and "Table 1." not in sec_text.replace("Table 1.  Parameter set", ""),
        "Fig. 2 marker intact": "[Fig. 2 placeholder" in sec_text,
    }
    for name, passed in checks.items():
        print(f"  [{'OK' if passed else 'FAIL'}] {name}")
        all_ok = all_ok and passed

    # 3.3
    body_33 = "\n".join(p.text for p in paras[h33 + 1:h34])
    checks = {
        "3.3 Mt/Minf < 0.6": "Mt/M∞ < 0.6" in body_33,
        "3.3 K-P power law": "C/C₀ = K·tⁿ" in body_33,
        "3.3 Korsmeyer-Peppas": "Korsmeyer–Peppas" in body_33,
        "3.3 14 days / 12-h": "14 days" in body_33 and "12-h" in body_33,
        "3.3 ICP-MS instrument": "PerkinElmer NexION 300X" in body_33,
        "3.3 -> Table S3": "Table S3" in body_33,
    }
    for name, passed in checks.items():
        print(f"  [{'OK' if passed else 'FAIL'}] {name}")
        all_ok = all_ok and passed

    # 3.4
    body_34 = "\n".join(p.text for p in paras[h34 + 1:h35])
    checks = {
        "3.4 pump 0.50 mL/min": "0.50 mL/min" in body_34,
        "3.4 5 MPa confining": "5 MPa" in body_34,
        "3.4 96-h shut-in": "96 h" in body_34,
        "3.4 200-mesh screens": "200-mesh" in body_34,
        "3.4 4-min / 21 samples": "4-min" in body_34 and "21 samples" in body_34,
        "3.4 Fig. S2": "Fig. S2" in body_34,
        "3.4 OWR ratios": "4:1, 1:1, and 1:4" in body_34,
        "3.4 total flow 0.1-0.4": "0.1–0.4 mL/min" in body_34,
        "3.4 self-calibration ref": "self-calibration test of Section 4.2" in body_34,
    }
    for name, passed in checks.items():
        print(f"  [{'OK' if passed else 'FAIL'}] {name}")
        all_ok = all_ok and passed

    # 3.5
    body_35 = "\n".join(p.text for p in paras[h35 + 1:end])
    banned = ["population", "crossover", "mutation", "ftol", "gtol", "generations",
              "differential evolution", "random seed"]
    found_banned = [t for t in banned if t in body_35.lower()]
    checks = {
        "3.5 Q bounds [10, 5000]": "[10, 5000] mL/min" in body_35,
        "3.5 500-fold": "500-fold" in body_35,
        "3.5 four independent runs": "Four independent runs" in body_35,
        "3.5 gradient-based": "gradient-based" in body_35,
        "3.5 50-200 iterations": "50–200 iterations" in body_35,
        "3.5 two-pass named": "Pass 1" in body_35 and "Pass 2" in body_35,
        "3.5 five models fair": "all five candidate models" in body_35,
        "3.5 -> Table S6": "Table S6" in body_35,
        "3.5 -> SM S5.3": "S5.3" in body_35,
        "3.5 -> Section 4.1": "Section 4.1" in body_35,
        "3.5 no algorithm terms": not found_banned,
        "3.5 no penalty/prior": "no penalty, prior, or constraint" in body_35,
    }
    for name, passed in checks.items():
        print(f"  [{'OK' if passed else 'FAIL'}] {name}")
        all_ok = all_ok and passed

    # Structural integrity of skeleton outside Section 3
    idx2 = [p.text.strip() for p in paras]
    h4 = idx2.index("4. Results and Discussion")
    heads_ok = "4.1 Model Selection: Single-Process Models Are Insufficient" in idx2[h4:]
    print(f"  [{'OK' if heads_ok else 'FAIL'}] Section 4.1 heading intact after Section 3")
    all_ok = all_ok and heads_ok

    print("VERIFICATION", "PASSED" if all_ok else "FAILED")
    return all_ok


if __name__ == "__main__":
    main()
    ok = verify()
    sys.exit(0 if ok else 1)
