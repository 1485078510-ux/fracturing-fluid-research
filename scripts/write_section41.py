# -*- coding: utf-8 -*-
"""
Task 4: Write Section 4.1 — Model Selection
into ESP-T_v2_manuscript.docx (skeleton by Task 1, Section 2 by Task 2, Section 3 by Task 3).

Writes into the blank paragraph after the 4.1 heading:
  - P1: intro — five candidate models, identical protocol (Section 3.5), Table 3, Fig. 4
  - Table 3 caption + real Word table (5 models, 7 columns; exact values from source
    ESP-T_Final_4-revised.docx Table 4 / v2 source table 2)
  - P2: extra-sum-of-squares F-test  F(3, 14) = 34.70, p < 10^-6
  - P3: single-process models capture ONE feature each; AICc evidence is NECESSARY,
        not SUFFICIENT; "AICc eliminates the single-process alternatives"
  - P4: pointer — Section 4.2 self-calibration is the decisive / positive test

The [Fig. 3] and [Fig. 4] marker paragraphs are left untouched (replaced at
figure-placement). Table captions: bold "Table 3." + regular text; "Table Grid" style.
Model statistics are reproduced verbatim from the source Table 4 (including the
Akaike-weight column).
"""
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX = r"C:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP-T_投稿文件\ESP-T_v2_manuscript.docx"

SUB_RE = re.compile(r"_([A-Za-z0-9]+)")


def add_runs(para, text, bold=False, italic=False):
    """Add runs to *para*; '_xyz' segments become subscript runs."""
    pos = 0
    for m in SUB_RE.finditer(text):
        if m.start() > pos:
            r = para.add_run(text[pos:m.start()])
            r.bold = bold
            r.italic = italic
        r = para.add_run(m.group(1))
        r.font.subscript = True
        r.bold = bold
        r.italic = italic
        pos = m.end()
    if pos < len(text):
        r = para.add_run(text[pos:])
        r.bold = bold
        r.italic = italic


def write_into(para, text):
    """Write *text* into an existing (blank) paragraph."""
    add_runs(para, text)
    return para


def insert_before(anchor, text, center=False):
    """Insert a new paragraph immediately before *anchor*; return it."""
    p = anchor.insert_paragraph_before()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(p, text)
    return p


def main():
    doc = Document(DOCX)
    paras = doc.paragraphs

    # Idempotency guard: refuse to run twice on the same document.
    if any("AICc eliminates the single-process alternatives" in p.text for p in paras):
        raise SystemExit("Section 4.1 already written; refusing to duplicate content. "
                         "Restore from git first (git checkout -- <docx>).")

    idx = [p.text.strip() for p in paras]
    i_41 = idx.index("4.1 Model Selection: Single-Process Models Are Insufficient")
    i_42 = idx.index("4.2 Physical Self-Calibration: The Flow-Rate Test")

    # Skeleton layout inside 4.1: heading, blank, [Fig. 3], [Fig. 4], then 4.2 heading.
    p_blank = paras[i_41 + 1]
    if p_blank.text.strip():
        raise SystemExit(f"unexpected non-blank paragraph after 4.1 heading: {p_blank.text!r}")
    fig3 = paras[i_41 + 2]
    fig4 = paras[i_41 + 3]
    if not fig3.text.startswith("[Fig. 3 placeholder"):
        raise SystemExit(f"unexpected paragraph at 4.1+2: {fig3.text!r}")
    if not fig4.text.startswith("[Fig. 4 placeholder"):
        raise SystemExit(f"unexpected paragraph at 4.1+3: {fig4.text!r}")
    if paras[i_41 + 4].text.strip() != "4.2 Physical Self-Calibration: The Flow-Rate Test":
        raise SystemExit(f"unexpected paragraph at 4.1+4: {paras[i_41 + 4].text!r}")

    # --- P1: introduction (written into the blank paragraph) ------------------
    write_into(p_blank, (
        "Five candidate models were fitted to the 21-point single-phase breakthrough "
        "curve under the identical two-pass protocol of Section 3.5, so that selection "
        "rests on a fair comparison: the dual-component tanh-blended model of Eq. (1), "
        "a single Gaussian (the ADE pulse solution, Source I only), a single erfc (the "
        "sustained-release solution, Source II only), an exponential decay, and the "
        "Korsmeyer–Peppas (K-P) power law. Table 3 reports, for each model, the number "
        "of fitted parameters k, the coefficient of determination R², the root-mean-"
        "square error RMSE, the corrected Akaike information criterion AICc, the "
        "increment ΔAICc relative to the best model, and the Akaike weight. The five "
        "fitted curves are overlaid on the measured BTC, with the ΔAICc ranking, in "
        "Fig. 4."
    ))

    # --- Table 3 caption + table (before the figure markers) ------------------
    # insert_paragraph_before preserves insertion order before the anchor, so insert
    # in final layout order: caption, then P2, P3, P4; the table is then moved to sit
    # directly after the caption. Final layout: P1, caption, Table 3, P2, P3, P4,
    # [Fig. 3], [Fig. 4].
    caption = insert_before(fig3, "")
    r_cap = caption.add_run("Table 3.  ")
    r_cap.bold = True
    caption.add_run("Model-selection statistics for the five candidate models fitted "
                    "to the 21-point single-phase BTC.")

    p2 = insert_before(fig3, (
        "The dual-component model attains R² = 0.9939 with RMSE = 0.0210, an order of "
        "magnitude below every single-process alternative. An extra-sum-of-squares "
        "F-test of the 7-parameter model against the best 4-parameter alternative (the "
        "single Gaussian) gives F(3, 14) = 34.70, p < 10⁻⁶: the reduction in residual "
        "variance attributable to the three additional parameters is far larger than "
        "expected by chance alone."
    ))
    p3 = insert_before(fig3, (
        "The pattern across the five models is diagnostic. Each single-process "
        "alternative captures exactly one feature of the BTC: the single Gaussian "
        "reproduces the peak but misses the tail; the single erfc reproduces the slow "
        "decay but cannot rise to the peak; and the exponential decay and K-P power "
        "law follow only the late tail (Fig. 3). None captures both features "
        "simultaneously, which is why every alternative has ΔAICc ≥ 32.66 — far beyond "
        "the conventional threshold (ΔAICc > 10) at which a candidate is judged to "
        "have essentially no support — and the Akaike weight of the dual-component "
        "model is effectively 1.000, whereas every alternative carries a weight below "
        "10⁻⁴. This evidence is necessary but not sufficient: AICc eliminates the "
        "single-process alternatives, but it cannot by itself establish that the "
        "dual-component structure is physically correct, because an information "
        "criterion rewards flexible curves and cannot distinguish a model that "
        "captures the transport physics from an overfitted one of similar shape."
    ))
    p4 = insert_before(fig3, (
        "The decisive test is therefore the self-calibration of Section 4.2: whether "
        "the fitted flow rate Q converges to the independently known pump setting "
        "(0.50 mL/min, Section 3.4) without any penalty, prior, or constraint toward "
        "that value (Section 2.4). Section 4.2 provides the positive test."
    ))

    rows = [
        ("Model", "k", "R²", "RMSE", "AICc", "ΔAICc", "Weight"),
        ("Dual-component tanh-blended", "7", "0.9939", "0.0210", "−139.70", "0.00", "1.000"),
        ("Single Gaussian", "4", "0.9482", "0.0609", "−107.04", "32.66", "<10⁻⁴"),
        ("Single erfc", "4", "0.7159", "0.1427", "−71.28", "68.42", "<10⁻⁴"),
        ("Exponential decay", "3", "0.7517", "0.1334", "−77.20", "62.51", "<10⁻⁴"),
        ("Korsmeyer–Peppas power law", "3", "−0.0193", "0.2703", "−47.54", "92.16", "<10⁻⁴"),
    ]
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    try:
        tbl.style = "Table Grid"
    except Exception:
        pass
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.cell(i, j)
            add_runs(cell.paragraphs[0], val, bold=(i == 0))
    caption._p.addnext(tbl._tbl)  # move table from doc end to after caption

    doc.save(DOCX)
    print("saved:", DOCX)


def verify():
    """Re-open the document and cross-check Section 4.1 against the spec."""
    doc = Document(DOCX)
    paras = doc.paragraphs
    idx = [p.text.strip() for p in paras]

    start = idx.index("4.1 Model Selection: Single-Process Models Are Insufficient")
    end = idx.index("4.2 Physical Self-Calibration: The Flow-Rate Test")
    all_ok = True

    sec_text = "\n".join(paras[i].text for i in range(start, end))
    body = [p.text for p in paras[start + 1:end] if p.text.strip()]
    words = sum(len(p.split()) for p in body)
    print(f"  words (4.1 prose, excl. caption): {words}")

    # --- Table 3 ---------------------------------------------------------------
    n_tables = len(doc.tables)
    t3 = doc.tables[2] if n_tables > 2 else None
    t3_ok = t3 is not None and len(t3.rows) == 6 and len(t3.columns) == 7
    t3_header = list(t3.rows[0].cells[j].text for j in range(7)) if t3 is not None else []
    exp_rows = [
        ("Dual-component tanh-blended", "7", "0.9939", "0.0210", "−139.70", "0.00", "1.000"),
        ("Single Gaussian", "4", "0.9482", "0.0609", "−107.04", "32.66", "<10⁻⁴"),
        ("Single erfc", "4", "0.7159", "0.1427", "−71.28", "68.42", "<10⁻⁴"),
        ("Exponential decay", "3", "0.7517", "0.1334", "−77.20", "62.51", "<10⁻⁴"),
        ("Korsmeyer–Peppas power law", "3", "−0.0193", "0.2703", "−47.54", "92.16", "<10⁻⁴"),
    ]
    row_match = t3 is not None and all(
        [t3.rows[r + 1].cells[j].text for j in range(7)] == list(exp_rows[r])
        for r in range(5)
    )

    checks = {
        "Table 3 exists (6x7, third table)": t3_ok and n_tables == 3,
        "Table 3 header": t3_header == ["Model", "k", "R²", "RMSE", "AICc", "ΔAICc", "Weight"],
        "Table 3 all 5 model rows exact": row_match,
        "Table 3 caption": "Table 3.  Model-selection statistics" in sec_text,
        "Table 3 not labeled Table 1/2": "Table 3." in sec_text
        and "Table 1." not in sec_text and "Table 2." not in sec_text,
    }
    for name, passed in checks.items():
        print(f"  [{'OK' if passed else 'FAIL'}] {name}")
        all_ok = all_ok and passed

    # --- Statistical content ----------------------------------------------------
    checks = {
        "F-test F(3,14)=34.70 p<1e-6": "F(3, 14) = 34.70, p < 10⁻⁶" in sec_text,
        "R2 = 0.9939 / RMSE 0.0210": "R² = 0.9939" in sec_text and "RMSE = 0.0210" in sec_text,
        "7-parameter vs 4-parameter stated": "7-parameter model" in sec_text
        and "4-parameter alternative" in sec_text,
    }
    for name, passed in checks.items():
        print(f"  [{'OK' if passed else 'FAIL'}] {name}")
        all_ok = all_ok and passed

    # --- Framing (critical) -----------------------------------------------------
    checks = {
        "one-feature-each framing": "captures exactly one feature" in sec_text
        and "cannot rise to the peak" in sec_text,
        "Gaussian fits peak misses tail": "reproduces the peak but misses the tail" in sec_text,
        "erfc captures tail not peak": "reproduces the slow decay but cannot rise to the peak" in sec_text,
        "necessary not sufficient explicit": "necessary but not sufficient" in sec_text,
        "AICc eliminates, not proves": "AICc eliminates the single-process alternatives" in sec_text
        and "cannot by itself establish" in sec_text,
        "no 'AICc proves' language": "proves the model" not in sec_text
        and "AICc proves" not in sec_text,
        "dAICc threshold 10 cited": "ΔAICc > 10" in sec_text,
        "Akaike weight 1.000 vs <1e-4": "1.000" in sec_text and "10⁻⁴" in sec_text,
        "Section 4.2 positive test pointer": "Section 4.2 provides the positive test" in sec_text
        and "decisive test" in sec_text,
        "Q pump setting cross-ref": "0.50 mL/min" in sec_text and "Section 3.4" in sec_text
        and "Section 2.4" in sec_text,
        "identical protocol -> 3.5": "identical two-pass protocol of Section 3.5" in sec_text,
        "Fig. 3 cited": "Fig. 3" in sec_text,
        "Fig. 4 cited": "Fig. 4" in sec_text,
        "Eq. (1) cited": "Eq. (1)" in sec_text,
    }
    for name, passed in checks.items():
        print(f"  [{'OK' if passed else 'FAIL'}] {name}")
        all_ok = all_ok and passed

    # --- Markers and skeleton integrity -----------------------------------------
    markers_ok = "[Fig. 3 placeholder" in sec_text and "[Fig. 4 placeholder" in sec_text
    n_fig_markers_4_1 = sum(1 for p in paras[start + 1:end]
                            if p.text.startswith("[Fig."))
    # Paragraph layout order: P1, caption, table, P2 (F-test), P3 (interpretation),
    # P4 (pointer to 4.2), [Fig. 3], [Fig. 4].
    def para_index(prefix):
        for i in range(start, end):
            if paras[i].text.startswith(prefix):
                return i
        return -1

    i_cap = para_index("Table 3.")
    i_f = para_index("The dual-component model attains")
    i_pattern = para_index("The pattern across the five models")
    i_decisive = para_index("The decisive test is therefore")
    i_fig3 = para_index("[Fig. 3 placeholder")
    i_fig4 = para_index("[Fig. 4 placeholder")
    order_ok = (start < i_cap < i_f < i_pattern < i_decisive < i_fig3 < i_fig4)
    h42 = idx[end:end + 8]  # placeholder-free sanity: heading 4.2 unchanged
    checks = {
        "Fig. 3 + Fig. 4 markers intact": markers_ok and n_fig_markers_4_1 == 2,
        "layout order P1<caption<F-test<interpret<pointer<markers": order_ok,
        "4.2 heading intact after 4.1": idx[end] ==
        "4.2 Physical Self-Calibration: The Flow-Rate Test",
        "word count 300-500": 300 <= words <= 500,
    }
    for name, passed in checks.items():
        print(f"  [{'OK' if passed else 'FAIL'}] {name}")
        all_ok = all_ok and passed

    print("VERIFICATION", "PASSED" if all_ok else "FAILED")
    return all_ok


if __name__ == "__main__":
    main()
    ok = verify()
    sys.exit(0 if ok else 1)
