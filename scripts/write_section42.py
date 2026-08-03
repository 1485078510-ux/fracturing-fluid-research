# -*- coding: utf-8 -*-
"""Task 5: write Section 4.2 (Physical Self-Calibration: The Flow-Rate Test) into the v2 manuscript.

- Replaces the empty spacer paragraph after the 4.2 heading
- Inserts four body paragraphs (decomposition ~200 words + self-calibration ~350 words)
- Strips the descriptive suffix from the Fig. 5 marker (keeps "[Fig. 5 placeholder — see Figure Captions]")
"""
import sys

sys.stdout.reconfigure(encoding="utf-8")
from docx import Document

TGT = "四氧化三铁环氧树脂拟合/ESP-T_投稿文件/ESP-T_v2_manuscript.docx"
HEADING = "4.2 Physical Self-Calibration: The Flow-Rate Test"
MARKER = "[Fig. 5 placeholder — see Figure Captions]"

P1 = (
    "Fitting Eq. (1) to the 21-point single-phase breakthrough curve under the two-pass "
    "protocol of Section 3.5 yields c_b = 0.0459, A = 2334, a = 0.431, α = 107.1 mm, "
    "Q = 0.46 mL/min, t₀ = 25.66 min, and σ = 3.96 min. The fitted curve reproduces the "
    "measured BTC with R² = 0.9939 and RMSE = 0.0210. Fig. 5a displays the decomposition "
    "that makes the two-source structure visible: the shaded Gaussian pulse, which carries "
    "53% of the integrated signal, describes the shut-in accumulation slug (Source I); the "
    "shaded erfc tail, which carries 47%, describes the sustained matrix-diffusion-controlled "
    "release (Source II); and the tanh weight joins the two across a smooth window of roughly "
    "8–10 min centered on the fitted crossover t₀ = 25.66 min. Each component is recognizable "
    "in the data itself: the Gaussian tracks the steep rise and the peak near t = 15 min, "
    "while the erfc tail tracks the slow decay that persists to the end of the measurement "
    "window. The residuals, plotted beneath the fit in Fig. 5a, scatter about zero with no "
    "systematic structure and no feature aligned with the crossover, so the blending "
    "transition introduces no visible fitting artifact at t₀."
)

P2 = (
    "The model has six free parameters — enough degrees of freedom to overfit: of the seven "
    "fitted coefficients in Eq. (1), the transition width σ is effectively determined by its "
    "physically motivated scale, approximately the 4-min sampling interval (Section 2.2), "
    "leaving six values that the data actually determine. However, extra parameters only "
    "improve curve-matching; they do not produce physically correct parameter values. If the "
    "model structure is incorrect, the optimizer will assign Q whatever value helps match the "
    "curve shape, with compensating adjustments in A and α to preserve the fit. Q itself was "
    "unconstrained in the objective function — no penalty, prior, or constraint directed it "
    "toward the pump setting of 0.50 mL/min; it was bounded only by the wide search envelope "
    "[10, 5000] mL/min."
)

P3 = (
    "Yet Q consistently settles at 0.46 ± 0.02 mL/min across four independent global searches, "
    "against the pump setting of 0.50 mL/min — an 8% deviation (Fig. 5b). The search bounds "
    "for Q spanned 10–5000 mL/min (Table S6) — a 500-fold range with no prior or penalty "
    "directing the optimizer toward the pump value. An incorrect model structure could have "
    "achieved comparable R² with Q = 2.0 or Q = 0.05 mL/min simply by compensating with A "
    "and α. Within that envelope nothing points the optimizer toward the pump setting; the "
    "converged value is determined entirely by the timing and shape of the observed BTC."
)

P4 = (
    "The fact that Q emerges from the BTC shape at 0.46 mL/min, entirely unforced and "
    "unconstrained by any prior, constitutes evidence that the model structure captures the "
    "actual transport physics. Q is not a fitting artifact — it is a recovered engineering "
    "parameter. Two derived quantities provide internal consistency checks, though neither "
    "constitutes independent verification: the fitted mean residence time (MRT = 37.4 min) "
    "agrees with the convective travel time (x/v = 38.6 min) within 3%, and the Peclet number "
    "Pe = x/α = 0.934 places the transport in the dispersion-dominated regime discussed "
    "further in Section 4.4."
)

doc = Document(TGT)

paras = doc.paragraphs
idx_head = next(i for i, p in enumerate(paras) if p.text.strip() == HEADING)
idx_mark = next(
    i for i, p in enumerate(paras)
    if p.text.strip().startswith("[Fig. 5 placeholder")
)
print(f"Heading at paragraph {idx_head}, marker at paragraph {idx_mark}")

# 1. Remove the empty spacer paragraph(s) between heading and marker
removed = 0
for i in range(idx_head + 1, idx_mark):
    if paras[i].text.strip() == "":
        paras[i]._p.getparent().remove(paras[i]._p)
        removed += 1
print(f"Removed {removed} empty spacer paragraph(s)")

# 2. Strip the marker suffix: keep only the marker text
mark_p = paras[idx_mark]
first_run = mark_p.runs[0] if mark_p.runs else mark_p.add_run()
first_run.text = MARKER
first_run.font.bold = False
first_run.font.italic = False
for r in mark_p.runs[1:]:
    r._r.getparent().remove(r._r)
print(f"Marker stripped to: {mark_p.text!r}")

# 3. Insert the four body paragraphs before the marker
for text in (P1, P2, P3, P4):
    p = mark_p.insert_paragraph_before(text)
    for r in p.runs:
        r.font.bold = False
        r.font.italic = False
    # ensure paragraph style is Normal
    p.style = doc.styles["Normal"]
print("Inserted 4 body paragraphs")

doc.save(TGT)
print("Saved:", TGT)

# Word counts
for name, t in (("P1", P1), ("P2", P2), ("P3", P3), ("P4", P4)):
    print(f"{name}: {len(t.split())} words")
