# -*- coding: utf-8 -*-
"""Task 6: write Sections 4.3-4.5 (TOA comparison, Independent Corroboration, Robustness)
into the v2 manuscript.

- Removes the empty spacer paragraph(s) after each heading
- Inserts body paragraphs before each figure marker
- Strips the descriptive suffix from the Fig. 6/7/8 markers
  (keeps "[Fig. N placeholder - see Figure Captions]")
"""
import sys

sys.stdout.reconfigure(encoding="utf-8")
from docx import Document

TGT = "四氧化三铁环氧树脂拟合/ESP-T_投稿文件/ESP-T_v2_manuscript.docx"

# ---------------------------------------------------------------- Section 4.3
S43 = {
    "heading": "4.3 Comparison with Time-of-Arrival Methods",
    "marker": "[Fig. 6 placeholder — see Figure Captions]",
    "body": [
        (
            "A natural question is whether the full coupled model of Eq. (1) is "
            "necessary, or whether simpler time-of-arrival (TOA) methods — which "
            "require no optimization — could suffice for estimating the per-interval "
            "flow rate Q. Three TOA approaches were evaluated against the same "
            "21-point single-phase breakthrough curve interpreted in Section 4.2 "
            "(Fig. 6)."
        ),
        (
            "The peak-time method assumes that the convective front arrives at the "
            "sampling point when the concentration reaches its maximum. Applying it "
            "to the measured BTC (t_peak = 15 min) yields Q = 1.31 mL/min, "
            "overestimating the pump setting of 0.50 mL/min by 162%. This gross "
            "overestimate arises because dispersive spreading shifts the apparent "
            "peak earlier than the true convective arrival: the peak marks the "
            "maximum of the slug, not its centroid, and the sustained-release tail "
            "further distorts the concentration envelope. The half-peak method, "
            "which uses the time at which the concentration reaches half its "
            "maximum (t_half ≈ 5 min), performs worse still, returning "
            "Q = 3.93 mL/min (+685%), because the rising limb is similarly affected "
            "by dispersion."
        ),
        (
            "The first-moment method uses the entire BTC rather than a single point: "
            "computing the mean residence time from the first moment of the curve, "
            "MRT = ∫t·C dt / ∫C dt = 37.1 min, and converting it through the same "
            "geometric relation gives Q = 0.53 mL/min (+5.8%), substantially more "
            "accurate than the peak-based methods. However, the first moment "
            "conflates the shut-in slug and the sustained tail into a single number: "
            "it provides no signal decomposition, no mechanistic insight into the "
            "relative contributions of the two release processes, and no estimate of "
            "the Peclet number."
        ),
        (
            "The coupled model of Eq. (1) achieves the best accuracy "
            "(Q = 0.46 mL/min, −8%) while simultaneously decomposing the BTC into "
            "its Gaussian (53%) and erfc (47%) components and recovering "
            "Pe = 0.934 for mechanistic interpretation (Section 4.2). Fig. 6 "
            "summarizes the comparison: the coupled model uniquely combines "
            "flow-rate accuracy with signal decomposition and physical "
            "interpretability."
        ),
    ],
}

# ---------------------------------------------------------------- Section 4.4
S44 = {
    "heading": "4.4 Independent Corroboration: K-P Kinetics and ADE Peclet Number",
    "marker": "[Fig. 7 placeholder — see Figure Captions]",
    "body": [
        (
            "Two completely independent experiments converge on the same physical "
            "picture. In the static batch release study (Section 3.3), ESP-T was "
            "immersed in dodecane in sealed glass vials, the tracer concentration in "
            "the liquid phase was quantified by ICP-MS as a function of time, and "
            "the resulting batch concentration history was fitted to the "
            "Korsmeyer–Peppas power law to extract the release exponent n. The "
            "fitted range n = 0.45–0.85 identifies anomalous (non-Fickian) release, "
            "in which matrix relaxation rather than diffusion alone controls the "
            "rate. In the core displacement experiment, an entirely different "
            "apparatus — a core holder with a packed proppant bed — generated the "
            "breakthrough curve of Section 4.2; fitting the coupled model of "
            "Eq. (1) to that BTC recovered the Peclet number Pe = x/α = 0.934, "
            "placing the transport in the dispersion-dominated (non-piston) regime."
        ),
        (
            "The two experiments are independent in every respect that matters: "
            "different apparatus (a static glass vial versus a dynamic core "
            "holder), different data (a batch concentration history versus a "
            "flowing breakthrough curve), and different fitting targets (the "
            "release parameters K and n versus the transport parameters Q, α, A, "
            "a, t₀, and σ). This is not cross-validation on a common dataset — the "
            "experiments share neither measurements nor fitted parameters, so their "
            "agreement cannot be manufactured by a shared fitting procedure or "
            "inflated by correlated errors."
        ),
        (
            "Despite these differences, both experiments point to the same "
            "physics. In the batch experiment, n < 1 indicates relaxation-limited, "
            "non-Fickian release; in the core experiment, Pe = 0.934 means that "
            "dispersion and convection contribute approximately equally to tracer "
            "spreading, so the front is far from piston-like. Non-Fickian release "
            "and non-piston transport are two manifestations of the same slow, "
            "matrix-dominated exchange process that the coupled model of Section 2 "
            "was constructed to represent. Fig. 7 places the two lines of evidence "
            "side by side: the K-P fits to the batch data (left panel) and the "
            "Peclet annotation on the fitted BTC (right panel)."
        ),
        (
            "Because the experiments are independent, their convergence does more "
            "than restate the Section 4.2 result: it shows that the model "
            "reproduces an independently observed transport regime rather than "
            "merely a single, carefully tuned dataset. The self-calibration test of "
            "Section 4.2 established that the model recovers a known engineering "
            "parameter; the corroboration of this section establishes that the "
            "parameters it recovers describe the same physics seen in an "
            "experiment of a completely different kind."
        ),
    ],
}

# ---------------------------------------------------------------- Section 4.5
S45 = {
    "heading": "4.5 Signal Decomposition and Robustness",
    "marker": "[Fig. 8 placeholder — see Figure Captions]",
    "body": [
        (
            "Separating the fitted curve (Eq. 1) into its two components (Fig. 5a) "
            "shows that the erfc tail accounts for approximately 47% of the tracer "
            "signal integrated over the 105-min measurement window, with the "
            "Gaussian pulse contributing 53%. Nearly half of the tracer detected at "
            "the wellhead therefore originates not from the shut-in accumulation "
            "slug but from sustained, ongoing release during the production period. "
            "This decomposition is a model output, not a direct measurement; its "
            "physical validity rests on independent grounds — the self-calibration "
            "test of Section 4.2, which establishes that the model captures the "
            "actual transport physics, and the independent corroboration of "
            "Section 4.4, which recovers the same two-process picture from a "
            "completely different experiment."
        ),
        (
            "To verify that the decomposition is not an artifact of the chosen "
            "transition width, σ was varied from 1.98 to 11.89 min — a six-fold "
            "range spanning 0.5×–3.0× of the fitted value (3.96 min) — while the "
            "other six parameters were held fixed (Fig. 8). The erfc tail fraction "
            "stays between 46.7% and 47.5% across the entire scan, a spread of "
            "less than one percentage point. If the 47% figure were a parametric "
            "artifact of the blending window, varying σ over a six-fold range "
            "would perturb it substantially; instead, the decomposition is "
            "insensitive to the only parameter that controls how the two "
            "components are separated in time."
        ),
        (
            "The σ insensitivity confirms that the erfc tail is a robust structural "
            "feature of the BTC rather than a tuning outcome. Taken together with "
            "the statistical necessity of the two-component structure (Section "
            "4.1) and the independent corroboration (Section 4.4), it establishes "
            "the decomposition as physically meaningful: the sustained-release "
            "fraction is a property of the tracer system, with a practical "
            "consequence — a single shut-in can support monitoring over an "
            "extended production period."
        ),
    ],
}

SECTIONS = [S43, S44, S45]

doc = Document(TGT)
paras = doc.paragraphs

for sec in SECTIONS:
    idx_head = next(i for i, p in enumerate(paras) if p.text.strip() == sec["heading"])
    fig_no = "".join(ch for ch in sec["marker"] if ch.isdigit())
    idx_mark = next(
        i for i, p in enumerate(paras)
        if p.text.strip().startswith("[Fig. " + fig_no + " placeholder")
    )
    print(f"{sec['heading']}: heading at {idx_head}, marker at {idx_mark}")

    # 1. Remove empty spacer paragraph(s) between heading and marker
    removed = 0
    for i in range(idx_head + 1, idx_mark):
        if paras[i].text.strip() == "":
            paras[i]._p.getparent().remove(paras[i]._p)
            removed += 1
    print(f"  Removed {removed} empty spacer paragraph(s)")

    # 2. Strip the marker suffix: keep only the marker text
    mark_p = paras[idx_mark]
    first_run = mark_p.runs[0] if mark_p.runs else mark_p.add_run()
    first_run.text = sec["marker"]
    first_run.font.bold = False
    first_run.font.italic = False
    for r in mark_p.runs[1:]:
        r._r.getparent().remove(r._r)
    print(f"  Marker stripped to: {mark_p.text!r}")

    # 3. Insert body paragraphs before the marker
    for text in sec["body"]:
        p = mark_p.insert_paragraph_before(text)
        for r in p.runs:
            r.font.bold = False
            r.font.italic = False
        p.style = doc.styles["Normal"]
    print(f"  Inserted {len(sec['body'])} body paragraphs")

doc.save(TGT)
print("Saved:", TGT)

for sec in SECTIONS:
    total = sum(len(t.split()) for t in sec["body"])
    per = [len(t.split()) for t in sec["body"]]
    print(f"{sec['heading']}: {total} words total {per}")
