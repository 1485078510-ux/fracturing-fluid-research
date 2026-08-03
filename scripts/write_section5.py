# -*- coding: utf-8 -*-
"""Task 7: write Section 5 (Two-Phase Production Allocation) into the v2 manuscript.

- 5.1 Physical Context (~90 words): steady multiphase flow, tail-dominated BTC
- 5.2 The Dilution Problem and the Flux Solution (~320 words): C_oil dilution artifact,
  F_O = C_oil x Q_oil mass balance, Fig. 9a/9b
- 5.3 Flux-to-Production-Rate Calibration (~230 words): F_O/F_O,ref vs Q_oil, r = 0.97,
  RMSD = 8.3%, Fig. 9c
- 5.4 Coupling to the BTC Framework (~150 words): K-P -> ADE -> flux -> Q_oil

Conventions identical to Tasks 5/6 (write_section42.py / write_section435.py):
style Normal, explicit non-bold/non-italic runs, Unicode math, spaced em dashes,
literal-underscore subscripts (F_O, C_oil, Q_oil, Q_total, F_O,ref), Fig. 9a/9b/9c citations.
Also strips the descriptive suffix from the [Fig. 9 placeholder] marker (Task-1 deferred note).
"""
import sys

sys.stdout.reconfigure(encoding="utf-8")
from docx import Document

TGT = "四氧化三铁环氧树脂拟合/ESP-T_投稿文件/ESP-T_v2_manuscript.docx"
MARKER = "[Fig. 9 placeholder — see Figure Captions]"

S51 = (
    "Under production conditions, shut-in occurs once during flowback. The majority of the "
    "monitoring period therefore operates under steady multiphase flow, with oil and water "
    "produced concurrently through the proppant pack. In this regime, the BTC is "
    "tail-dominated — the sustained release component (erfc) governs the observed "
    "concentration, while the shut-in accumulation slug (Source I) has already been swept to "
    "the wellhead. Our two-phase core displacement experiments simulate this steady-state "
    "regime directly, holding each OWR and total flow rate fixed while recording the steady "
    "tracer concentration in the produced oil."
)

S52_P1 = (
    "Two-phase displacement experiments were run at three oil–water ratios (OWR = 4:1, 1:1, "
    "1:4) and four total flow rates (0.1–0.4 mL/min) under steady-state conditions "
    "(Section 3.4); the complete dataset is given in Table S5. Fig. 9a plots the oil-phase "
    "tracer concentration C_oil against the total flow rate Q_total for each OWR. C_oil falls "
    "monotonically as Q_total rises — at OWR = 4:1, from 33.51 to 8.33 μg/mL over the "
    "four-fold rate increase. This is a straightforward consequence of dilution: at higher "
    "flow rates, less time is available for tracer to accumulate in each unit volume of "
    "produced fluid. The concentration shows no systematic dependence on OWR, suggesting that "
    "oil and water compete for the proppant surface in rough proportion to their volume "
    "fractions. Concentration alone is therefore a poor proxy for the production rate of a "
    "given stage."
)

S52_P2 = (
    "The way around this is to work with tracer flux rather than concentration. We define the "
    "oil-phase tracer mass flux as F_O = C_oil × Q_oil — the product of the measured "
    "oil-phase concentration and the oil flow rate. At steady state, F_O equals the release "
    "rate from the ESP-T pack: mass leaves the pack at a fixed rate, so the flux measured at "
    "the sampling point cannot depend on how fast fluid sweeps past it. Fig. 9b confirms this "
    "mass balance: within each OWR, F_O is flat across the four total flow rates — at "
    "OWR = 4:1, F_O stays within 2.666–2.779 μg/min while C_oil varies by a factor of four."
)

S52_P3 = (
    "F_O does depend on OWR: the plateau rises from about 0.65 μg/min at OWR = 1:4 to about "
    "1.6 μg/min at 1:1 and about 2.7 μg/min at 4:1. More oil in the mixture means more oil "
    "contacts the proppant surface; the larger the oil-wetted area, the higher the tracer "
    "release rate. The proppant releases tracer at a rate governed by the oil-wetted area, "
    "not by how fast fluid sweeps past it."
)

S53_P1 = (
    "The critical test is whether F_O can be converted into a production rate. Fig. 9c "
    "normalizes the two-phase fluxes by the steady-state flux measured during single-phase "
    "oil displacement, F_O,ref = 3.187 ± 0.15 μg/min (triplicate), and compares the ratio "
    "F_O/F_O,ref with the known oil flow rate Q_oil at each OWR. The normalized flux "
    "collapses the three OWR families onto a single rising trend — approximately 0.84, 0.49, "
    "and 0.21 at the lowest rate, tracking the oil volume fractions 0.8, 0.5, and 0.2 of the "
    "three mixtures. The agreement is close: Pearson r = 0.97, p = 0.006, RMSD = 8.3%, "
    "confirming that per-interval oil production rates can be recovered from F_O measurements "
    "when the total flow rate is known, even with a water phase present."
)

S53_P2 = (
    "The primary source of uncertainty in this calibration is the reproducibility of the "
    "single-phase reference — F_O,ref varies by ±4.7% across triplicate measurements — which "
    "propagates into every normalized flux. Two caveats bound the result. First, the "
    "correlation rests on three OWR conditions (n = 3 independent oil-fraction levels); a "
    "larger OWR matrix would strengthen the calibration. Second, F_O,ref is batch-specific: "
    "it reflects the release behavior of a particular ESP-T pack, so batch-specific F_O,ref "
    "determination is advisable for field deployment. Within these limits, the flux method "
    "converts a concentration measurement corrupted by dilution into a quantitative "
    "production allocation that survives the presence of a water phase."
)

S54 = (
    "Under production conditions the OWR and the total flow rate both evolve with time, so the "
    "steady-state calibration of Section 5.3 must be embedded in a dynamic description; the "
    "BTC framework provides it. K-P provides the temperature-dependent release rate K(T) and "
    "the transport mechanism indicator n(T). The erfc tail amplitude a, fitted from the BTC "
    "(Section 4.2, a = 0.431), quantifies the sustained-release source strength — the rate at "
    "which the ESP-T pack feeds tracer into the produced oil at the wellbore temperature. "
    "Combined with the K-P temperature calibration, the steady-state concentration (erfc tail "
    "plateau) yields the per-stage oil flow rate Q_oil through the flux relation "
    "F_O = C_oil × Q_oil. This closes the loop: static K-P → dynamic ADE → two-phase flux → "
    "per-stage oil production rate, so that a batch measurement on a few grams of ESP-T "
    "becomes a quantitative production allocation for every stage of a fractured well."
)

doc = Document(TGT)
paras = doc.paragraphs


def find(text, starts=False):
    for i, p in enumerate(paras):
        t = p.text.strip()
        if (t.startswith(text) if starts else t == text):
            return i
    raise RuntimeError(f"not found: {text!r}")


def fill_section(heading, anchor_text, body, anchor_is_marker=False):
    """Remove empty spacers between HEADING and ANCHOR_TEXT, then insert BODY before anchor."""
    i_head = find(heading)
    i_anchor = find(anchor_text, starts=True)
    assert i_anchor > i_head, f"anchor before heading for {heading}"
    # 1. Remove empty spacer paragraphs between heading and anchor
    removed = 0
    for i in range(i_head + 1, i_anchor):
        if paras[i].text.strip() == "":
            paras[i]._p.getparent().remove(paras[i]._p)
            removed += 1
    print(f"[{heading}] removed {removed} spacer paragraph(s)")
    # Re-locate the anchor after removals
    i_anchor = find(anchor_text, starts=True)
    # 2. Insert body paragraphs before the anchor
    for text in body:
        p = paras[i_anchor].insert_paragraph_before(text)
        p.style = doc.styles["Normal"]
        for r in p.runs:
            r.font.bold = False
            r.font.italic = False
    print(f"[{heading}] inserted {len(body)} paragraph(s) before {anchor_text!r}")


# Process in document order so re-finding by text stays correct.
fill_section("5.1 Physical Context", "5.2 The Dilution Problem and the Flux Solution", [S51])
fill_section(
    "5.2 The Dilution Problem and the Flux Solution",
    "[Fig. 9 placeholder",
    [S52_P1, S52_P2, S52_P3],
    anchor_is_marker=True,
)
fill_section("5.3 Flux-to-Production-Rate Calibration", "5.4 Coupling to the BTC Framework", [S53_P1, S53_P2])
fill_section("5.4 Coupling to the BTC Framework", "6. Field Deployment Pathway", [S54])

# 3. Strip the descriptive suffix from the Fig. 9 marker (Task-1 deferred note)
paras = doc.paragraphs
idx_mark = next(i for i, p in enumerate(paras) if p.text.strip().startswith("[Fig. 9 placeholder"))
mark_p = paras[idx_mark]
first_run = mark_p.runs[0] if mark_p.runs else mark_p.add_run()
first_run.text = MARKER
first_run.font.bold = False
first_run.font.italic = False
for r in mark_p.runs[1:]:
    r._r.getparent().remove(r._r)
print(f"Marker stripped to: {mark_p.text!r}")

doc.save(TGT)
print("Saved:", TGT)

# Word counts
for name, t in (("S51", S51), ("S52_P1", S52_P1), ("S52_P2", S52_P2), ("S52_P3", S52_P3),
                ("S53_P1", S53_P1), ("S53_P2", S53_P2), ("S54", S54)):
    print(f"{name}: {len(t.split())} words")
print(f"Section 5 total: {sum(len(t.split()) for t in (S51, S52_P1, S52_P2, S52_P3, S53_P1, S53_P2, S54))} words")
