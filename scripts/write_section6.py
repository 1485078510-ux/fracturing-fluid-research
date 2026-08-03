# -*- coding: utf-8 -*-
"""Task 8: write Section 6 (Field Deployment Pathway) into the v2 manuscript.

- Deployment steps (~200 words): 6-step roadmap from lab to field (stage tagging via
  multi-element doping, single shut-in, flowback wellhead sampling, per-element BTC fit,
  flux-method production allocation, trend-based optimization). Fig. 10 marker kept
  directly after this paragraph (skeleton position).
- Limitations (~200 words): five proactive scope boundaries (single-interval scale,
  dodecane model oil, epoxy stability in aggressive environments, n = 3 OWR levels,
  batch-specific F_O,ref) stated as boundaries, not weaknesses.

Conventions identical to Tasks 2-7 (write_section5.py et al.):
style Normal, explicit non-bold/non-italic runs, Unicode math/subscripts
(H₂S, CO₂), spaced em dashes, literal-underscore subscripts
(Q_i, Q_oil_i(t), F_O, F_O,ref), "Eq. (1)" in-text, "120 °C" with space.
Also strips the descriptive suffix from the [Fig. 10 placeholder] marker
(Task-1 deferred note, same as Tasks 5-7 did for Figs. 6-9).
"""
import sys

sys.stdout.reconfigure(encoding="utf-8")
from docx import Document

TGT = "四氧化三铁环氧树脂拟合/ESP-T_投稿文件/ESP-T_v2_manuscript.docx"
MARKER = "[Fig. 10 placeholder — see Figure Captions]"

S6_DEPLOY = (
    "Field deployment builds directly on the laboratory framework of Sections 2–5 and "
    "proceeds in six steps (Fig. 10). (1) Stage tagging: each fracture stage receives ESP-T "
    "particles doped with a distinct metal or rare-earth element — Mn, Zn, Cu, Eu, or Dy — "
    "so that every stage carries a unique ICP-MS fingerprint, allowing per-stage signals to "
    "be separated from a single wellhead sample. (2) Shut-in: a single shut-in after "
    "fracturing suffices; tracer accumulates within each stage, and the resulting slug is "
    "swept toward the wellbore when flowback begins. (3) Flowback sampling: wellhead samples "
    "are collected at intervals guided by Eq. (1) and the wellbore geometry, which fix the "
    "expected pulse arrival time and hence the sampling schedule. (4) Per-stage flow rates: "
    "each element's breakthrough curve is fitted with the dual-component model, recovering "
    "the flow rate Q_i of its stage. (5) Production allocation: during steady production, "
    "periodic wellhead sampling feeds the flux method of Section 5, yielding the per-stage "
    "oil production rate Q_oil_i(t). (6) Optimization: trends in Q_i and Q_oil_i(t) "
    "identify declining stages and guide well-spacing and completion design. Because "
    "sampling and ICP-MS analysis are entirely surface operations, no downhole tooling is "
    "required at any step — the method converts routine flowback and production sampling "
    "into a quantitative map of stage-by-stage productivity."
)

S6_LIMIT = (
    "Several limitations of the present study should be noted; each defines the scope of the "
    "method as currently validated. First, all experiments were conducted at the "
    "single-interval, laboratory scale with uniform proppant packing; multi-interval field "
    "configurations, where inter-stage interference and wellbore mixing act, require "
    "dedicated validation. Second, dodecane served as the model oil; crude oil, with its "
    "higher viscosity and more complex composition, and transient flow conditions during "
    "early flowback remain untested. Third, the chemical stability of the epoxy matrix in "
    "aggressive wellbore environments — H₂S, CO₂, high-salinity brines, and temperatures "
    "above 120 °C — has not been evaluated. Fourth, the flux calibration rests on three OWR "
    "levels (n = 3 independent points); a larger OWR matrix would strengthen the F_O–OWR "
    "relation established in Section 5.3. Fifth, F_O,ref is batch-specific, and "
    "batch-specific F_O,ref determination is recommended before field deployment. None of "
    "these boundaries invalidates the release–transport mechanism demonstrated here; they "
    "delimit the conditions under which per-stage production allocation is presently "
    "quantified, and each maps onto a defined follow-up experiment or field protocol."
)

doc = Document(TGT)
paras = doc.paragraphs


def find(text, paras, starts=False):
    """Return index of paragraph matching TEXT within the given PARAS snapshot."""
    for i, p in enumerate(paras):
        t = p.text.strip()
        if (t.startswith(text) if starts else t == text):
            return i
    raise RuntimeError(f"not found: {text!r}")


def fill_section(heading, anchor_text, body):
    """Remove empty spacers between HEADING and ANCHOR_TEXT, then insert BODY before anchor.

    The paragraph list must be passed explicitly to find(): a module-global snapshot
    becomes stale once paragraphs are removed/inserted, and Python closure scoping
    would otherwise make find() index a different list than the one being modified.
    """
    paras = doc.paragraphs  # fresh snapshot: prior calls may have detached elements
    i_head = find(heading, paras)
    i_anchor = find(anchor_text, paras, starts=True)
    assert i_anchor > i_head, f"anchor before heading for {heading}"
    removed = 0
    for i in range(i_head + 1, i_anchor):
        if paras[i].text.strip() == "":
            paras[i]._p.getparent().remove(paras[i]._p)
            removed += 1
    print(f"[{heading}] removed {removed} spacer paragraph(s)")
    paras = doc.paragraphs
    i_anchor = find(anchor_text, paras, starts=True)
    for text in body:
        p = paras[i_anchor].insert_paragraph_before(text)
        p.style = doc.styles["Normal"]
        for r in p.runs:
            r.font.bold = False
            r.font.italic = False
    print(f"[{heading}] inserted {len(body)} paragraph(s) before {anchor_text!r}")


# 1. Deployment paragraph: between heading and the Fig. 10 marker (skeleton position).
fill_section("6. Field Deployment Pathway", "[Fig. 10 placeholder", [S6_DEPLOY])
# 2. Limitations paragraph: before the next heading (after the marker).
fill_section("6. Field Deployment Pathway", "7. Conclusions", [S6_LIMIT])

# 3. Strip the descriptive suffix from the Fig. 10 marker (Task-1 deferred note).
paras = doc.paragraphs
idx_mark = next(i for i, p in enumerate(paras) if p.text.strip().startswith("[Fig. 10 placeholder"))
mark_p = paras[idx_mark]
first_run = mark_p.runs[0] if mark_p.runs else mark_p.add_run()
first_run.text = MARKER
first_run.font.bold = False
first_run.font.italic = False
for r in mark_p.runs[1:]:
    r._r.getparent().remove(r._r)
print(f"Marker stripped to: {mark_p.text!r}")

doc.save(TGT)
print("Saved:", TGT)

for name, t in (("S6_DEPLOY", S6_DEPLOY), ("S6_LIMIT", S6_LIMIT)):
    print(f"{name}: {len(t.split())} words")
print(f"Section 6 total: {len(S6_DEPLOY.split()) + len(S6_LIMIT.split())} words")
