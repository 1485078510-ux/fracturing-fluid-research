# -*- coding: utf-8 -*-
"""Task 9: write Section 7 (Conclusions) into the v2 manuscript.

- Lead-in sentence + 5 numbered conclusion points, each a standalone finding,
  verbatim from the task-9 brief (design doc §7). No new information introduced;
  engineering register.
- Numbers match the manuscript body exactly: Q_fit = 0.46 mL/min vs
  Q_pump = 0.50 mL/min (8% deviation, bounds 10–5000 mL/min), K-P n = 0.45–0.85,
  Pe = 0.934, OWR 4:1 to 1:4, Pearson r = 0.97, RMSD = 8.3%.

Conventions identical to Tasks 2-8 (write_section6.py et al.):
style Normal, explicit non-bold/non-italic runs, Unicode math/subscripts
(C¹ superscript as in Section 2.2), spaced em dashes, en-dash ranges,
literal-underscore subscripts (Q_fit, Q_pump, F_O, Q_oil), "Eq." style
abbreviations as used in the body.
"""
import sys

sys.stdout.reconfigure(encoding="utf-8")
from docx import Document

TGT = "四氧化三铁环氧树脂拟合/ESP-T_投稿文件/ESP-T_v2_manuscript.docx"

S7_LEAD = (
    "This study developed and experimentally validated a coupled release–transport "
    "framework for interpreting tracer-proppant breakthrough curves in hydraulic "
    "fracturing diagnostics. The principal conclusions are:"
)

S7_POINTS = [
    "1. A coupled release–transport model was developed that decomposes a "
    "tracer-proppant BTC into a shut-in accumulation slug (Gaussian component, ADE "
    "instantaneous-pulse solution) and sustained matrix-diffusion-controlled release "
    "(erfc component, ADE semi-infinite boundary solution), linked by a C¹-continuous "
    "tanh transition.",
    "2. Physical self-calibration: the flow rate Q_fit = 0.46 mL/min converges to the "
    "independently set pump rate Q_pump = 0.50 mL/min (deviation 8%) with Q entirely "
    "unconstrained in the objective function (search bounds 10–5000 mL/min), confirming "
    "the model captures actual transport physics rather than merely fitting the curve.",
    "3. Independent corroboration: K-P batch kinetics (n = 0.45–0.85, non-Fickian "
    "transport) and ADE Peclet number (Pe = 0.934, non-piston displacement) converge on "
    "the same physical picture from completely separate experiments — different "
    "apparatus, different data, different fitting targets.",
    "4. Two-phase extension: the oil-phase tracer mass flux F_O eliminates the dilution "
    "artifact inherent in concentration-based interpretation and tracks the oil "
    "production rate Q_oil across OWR 4:1 to 1:4 (Pearson r = 0.97, RMSD = 8.3%).",
    "5. The framework enables per-stage production allocation from wellhead samples "
    "alone, requiring no downhole tools and only a single shut-in. Multi-stage field "
    "validation with crude oil under transient conditions remains the next step.",
]

doc = Document(TGT)


def find(text, paras, starts=False):
    """Return index of paragraph matching TEXT within the given PARAS snapshot."""
    for i, p in enumerate(paras):
        t = p.text.strip()
        if (t.startswith(text) if starts else t == text):
            return i
    raise RuntimeError(f"not found: {text!r}")


def fill_section(heading, anchor_text, body):
    """Remove empty spacers between HEADING and ANCHOR_TEXT, then insert BODY before anchor.

    The paragraph list must be passed explicitly to find(): a module-global snapshot
    becomes stale once paragraphs are removed/inserted (pitfall documented in
    write_section6.py).
    """
    paras = doc.paragraphs  # fresh snapshot
    i_head = find(heading, paras)
    i_anchor = find(anchor_text, paras, starts=True)
    assert i_anchor > i_head, f"anchor before heading for {heading}"
    removed = 0
    for i in range(i_head + 1, i_anchor):
        if paras[i].text.strip() == "":
            paras[i]._p.getparent().remove(paras[i]._p)
            removed += 1
    print(f"[{heading}] removed {removed} spacer paragraph(s)")
    paras = doc.paragraphs
    i_anchor = find(anchor_text, paras, starts=True)
    for text in body:
        p = paras[i_anchor].insert_paragraph_before(text)
        p.style = doc.styles["Normal"]
        for r in p.runs:
            r.font.bold = False
            r.font.italic = False
    print(f"[{heading}] inserted {len(body)} paragraph(s) before {anchor_text!r}")


# Conclusions content between the "7. Conclusions" heading and "References".
fill_section("7. Conclusions", "References", [S7_LEAD] + S7_POINTS)

doc.save(TGT)
print("Saved:", TGT)

total = len(S7_LEAD.split()) + sum(len(p.split()) for p in S7_POINTS)
print(f"Section 7 total: {total} words (lead-in {len(S7_LEAD.split())}, "
      f"points {total - len(S7_LEAD.split())})")
