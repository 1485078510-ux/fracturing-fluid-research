"""Add new refs to existing 30-ref list. Simple append + cite + renumber."""
import re,os,shutil
shutil.copy2(r'c:\Users\郝\Desktop\claude\ESP-T_Final.docx',
             os.path.join(os.environ['USERPROFILE'],'Desktop','ESP-T_投稿文件','ESP-T_Final.docx'))
from docx import Document; from docx.shared import Pt; from docx.oxml.ns import qn

doc = Document(os.path.join(os.environ['USERPROFILE'],'Desktop','ESP-T_投稿文件','ESP-T_Final.docx'))
ref_start=None
for i,p in enumerate(doc.paragraphs):
    if p.text.strip()=='References': ref_start=i; break

PROTECTED=set()
for i,p in enumerate(doc.paragraphs):
    for run in p.runs:
        if run._element.findall(qn('w:drawing')) or run._element.findall(qn('w:object')): PROTECTED.add(i); break
    for r in p._element:
        tag=r.tag.split('}')[-1] if '}' in r.tag else r.tag
        if tag in ['oMath','oMathPara']: PROTECTED.add(i); break

ref_entries={}; ref_indices=[]
for i in range(ref_start+1,len(doc.paragraphs)):
    t=doc.paragraphs[i].text.strip(); m=re.match(r'\[(\d+)\]\s+(.*)',t)
    if m: ref_entries[int(m.group(1))]=m.group(2); ref_indices.append(i)

existing_texts=set(v.strip().lower()[:80] for v in ref_entries.values())
print(f'Current refs: {len(ref_entries)}')

# User's new refs + authoritative + fitting-related
new_refs = [
    'ROGNER H H. An assessment of world hydrocarbon resources [J]. Annual Review of Energy and the Environment, 1997, 22(1): 217-262.',
    'IEA. World Energy Employment 2025 [R]. Paris: International Energy Agency, 2025.',
    'HE X, ZHANG P, GAO Y, et al. Challenges and countermeasures for beneficial development of unconventional oil and gas resources in China [J]. China Petroleum Exploration, 2025, 30(1): 28-43.',
    'MALYAVKO E, UPADHYE V, HUSEIN N. Research of operational dynamics of a well with two hydraulic fractures with use of marked proppant [C]. SPE-215624, SPE International Hydraulic Fracturing Technology Conference, Muscat, 2023.',
    'KING G E. Thirty years of gas shale fracturing: What have we learned? [C]. SPE-133456, SPE ATCE, Florence, 2010.',
    'MAXWELL S C. Microseismic hydraulic fracture imaging [J]. Interpretation, 2014, 2(3): SJ1-SJ13.',
    'MOLENAAR M M et al. First downhole application of DAS for fracturing diagnostics [J]. SPE Drilling & Completion, 2012, 27(1): 32-38.',
    'LIANG F et al. A comprehensive review on proppant technologies [J]. Petroleum, 2016, 2(1): 26-39.',
    'KORSMEYER R W et al. Mechanisms of solute release from porous hydrophilic polymers [J]. International Journal of Pharmaceutics, 1983, 15(1): 25-35.',
    'CIPOLLA C L, WALLACE J. Stimulated reservoir volume: A misnomer? [C]. SPE-168596, SPE HFTC, The Woodlands, 2014.',
    # Fitting/modeling related for Introduction
    'RITGER P L, PEPPAS N A. A simple equation for description of solute release I. Fickian and non-Fickian release [J]. Journal of Controlled Release, 1987, 5(1): 23-36.',
    'PEPPAS N A, SAHLIN J J. A simple equation for the description of solute release. III. Coupling of diffusion and relaxation [J]. International Journal of Pharmaceutics, 1989, 57(2): 169-172.',
    'VAN GENUCHTEN M T, ALVES W J. Analytical solutions of the one-dimensional convective-dispersive solute transport equation [R]. USDA Technical Bulletin No. 1661, 1982.',
    'SHOOK G M, POPE G A, ASAKAWA K. Determining reservoir properties and flood performance from tracer test analysis [C]. SPE-124614, SPE ATCE, New Orleans, 2009.',
    'LIU J, WANG H, ZHANG T, et al. Study on interpretation method of multistage fracture tracer flowback curve in tight oil reservoirs [J]. ACS Omega, 2024, 9: 10852-10864.',
]

to_add = [(j,rt) for j,rt in enumerate(new_refs) if rt.strip().lower()[:80] not in existing_texts]
nn = max(ref_entries.keys())+1; li = max(ref_indices)
for j,(_,rt) in enumerate(to_add):
    if li+j+1<len(doc.paragraphs) and not doc.paragraphs[li+j+1].text.strip():
        run = doc.paragraphs[li+j+1].add_run(f'[{nn+j}] {rt}'); run.font.name='Times New Roman'; run.font.size=Pt(9)

total = nn+len(to_add)-1
added_nums = [nn+j for j in range(len(to_add))]
print(f'Added {len(to_add)} refs [{nn}-{total}]. Total list: {total}')

# SIMPLE CITATIONS: append new ref numbers to existing brackets
# Para 7 (context): append Rogner, IEA2025, HeX, King, Maxwell, Cipolla to first bracket
para7_new = [nn, nn+1, nn+2, nn+4, nn+5, nn+9]
for run in doc.paragraphs[7].runs:
    m = re.search(r'\[(\d+)', run.text)
    if m and str(para7_new[0]) not in run.text:
        old = f'[{m.group(1)}'
        new_nums = ','.join(str(x) for x in para7_new)
        run.text = run.text.replace(old, f'[{new_nums},{m.group(1)}', 1)
        print(f'[7] Added energy/fracturing refs')
        break

# Para 11 (tracer context): Malyavko, Molenaar, Shook, Liu, Van Genuchten
para11_new = [nn+3, nn+6, nn+13, nn+14, nn+12]
for run in doc.paragraphs[11].runs:
    m = re.search(r'\[(\d+)', run.text)
    if m and str(para11_new[0]) not in run.text:
        old = f'[{m.group(1)}'
        new_nums = ','.join(str(x) for x in para11_new)
        run.text = run.text.replace(old, f'[{new_nums},{m.group(1)}', 1)
        print(f'[11] Added tracer/fitting refs')
        break

# Para 11: Korsmeyer + Ritger + Peppas for K-P
for run in doc.paragraphs[11].runs:
    if 'power law [' in run.text and str(nn+8) not in run.text:
        run.text = run.text.replace('power law [', f'power law [{nn+8},{nn+10},{nn+11},')
        print(f'[11] Added K-P fitting refs')
        break

# Para 13: Liang proppant review
for run in doc.paragraphs[13].runs:
    m = re.search(r'\[(\d+)', run.text)
    if m and str(nn+7) not in run.text:
        old = f'[{m.group(1)}'
        run.text = run.text.replace(old, f'[{nn+7},{m.group(1)}', 1)
        print(f'[13] Added Liang ref')
        break

doc.save(os.path.join(os.environ['USERPROFILE'],'Desktop','ESP-T_投稿文件','ESP-T_Final.docx'))

# ===== RENUMBER =====
doc2 = Document(os.path.join(os.environ['USERPROFILE'],'Desktop','ESP-T_投稿文件','ESP-T_Final.docx'))
ref_entries2={}; ref_indices2=[]
for i in range(ref_start+1,len(doc2.paragraphs)):
    t=doc2.paragraphs[i].text.strip(); m=re.match(r'\[(\d+)\]\s+(.*)',t)
    if m: ref_entries2[int(m.group(1))]=m.group(2); ref_indices2.append(i)

first={}
for i,p in enumerate(doc2.paragraphs):
    if i>=ref_start or i in PROTECTED: continue
    for m in re.finditer(r'(?<!\d)\[(\d+)\](?!\d)',p.text):
        n=int(m.group(1))
        if n in ref_entries2 and n not in first: first[n]=i
    for m in re.finditer(r'(?<!\d)\[(\d+)[–-](\d+)\](?!\d)',p.text):
        a,b=int(m.group(1)),int(m.group(2))
        for n in range(a,b+1):
            if n in ref_entries2 and n not in first: first[n]=i

ordered=sorted(first.keys()); o2n={old:new for new,old in enumerate(ordered,1)}
for i,p in enumerate(doc2.paragraphs):
    if i>=ref_start or i in PROTECTED: continue
    for run in p.runs:
        t=run.text
        def repl(m):
            inner=m.group(1); parts=re.split(r'\s*,\s*',inner); new_p=[]
            for part in parts:
                part=part.strip(); rm=re.match(r'(\d+)[–-](\d+)',part)
                if rm:
                    a,b=int(rm.group(1)),int(rm.group(2))
                    na=o2n.get(a,a); nb=o2n.get(b,b)
                    new_p.append(f'{na}-{nb}' if na!=nb else str(na))
                else: n=int(part); new_p.append(str(o2n.get(n,n)))
            return '['+','.join(new_p)+']'
        nt=re.sub(r'\[(\d+(?:[–-]\d+)?(?:,\s*\d+(?:[–-]\d+)?)*)\]',repl,t)
        if nt!=t: run.text=nt

for idx in ref_indices2:
    p=doc2.paragraphs[idx]
    for r in list(p._element.findall(qn('w:r'))):
        if not r.findall(qn('w:drawing')): p._element.remove(r)
for j,old_num in enumerate(ordered):
    if old_num in ref_entries2 and j<len(ref_indices2):
        run=doc2.paragraphs[ref_indices2[j]].add_run(f'[{j+1}] {ref_entries2[old_num]}')
        run.font.name='Times New Roman'; run.font.size=Pt(9)
for j in range(len(ordered),len(ref_indices2)):
    p=doc2.paragraphs[ref_indices2[j]]
    for r in list(p._element.findall(qn('w:r'))):
        if not r.findall(qn('w:drawing')): p._element.remove(r)
doc2.save(os.path.join(os.environ['USERPROFILE'],'Desktop','ESP-T_投稿文件','ESP-T_Final.docx'))

doc3=Document(os.path.join(os.environ['USERPROFILE'],'Desktop','ESP-T_投稿文件','ESP-T_Final.docx'))
imgs=sum(len(r._element.findall(qn('w:drawing'))) for p in doc3.paragraphs for r in p.runs)
ref_n=sum(1 for i in range(ref_start+1,len(doc3.paragraphs)) if doc3.paragraphs[i].text.strip().startswith('['))
print(f'FINAL: {ref_n} refs, {imgs} images, {len(ordered)} cited')
