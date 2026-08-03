"""Standalone citations — each new ref gets its own [N] bracket."""
import re,os,shutil
shutil.copy2(r'c:\Users\郝\Desktop\claude\ESP-T_Final.docx',
             os.path.join(os.environ['USERPROFILE'],'Desktop','ESP-T_投稿文件','ESP-T_Final.docx'))
from docx import Document; from docx.shared import Pt; from docx.oxml.ns import qn

doc = Document(os.path.join(os.environ['USERPROFILE'],'Desktop','ESP-T_投稿文件','ESP-T_Final.docx'))
ref_start=None
for i,p in enumerate(doc.paragraphs):
    if p.text.strip()=='References': ref_start=i; break

PROTECTED=set()
for i,p in enumerate(doc.paragraphs):
    for run in p.runs:
        if run._element.findall(qn('w:drawing')) or run._element.findall(qn('w:object')): PROTECTED.add(i); break
    for r in p._element:
        tag=r.tag.split('}')[-1] if '}' in r.tag else r.tag
        if tag in ['oMath','oMathPara']: PROTECTED.add(i); break

ref_entries={}; ref_indices=[]
for i in range(ref_start+1,len(doc.paragraphs)):
    t=doc.paragraphs[i].text.strip(); m=re.match(r'\[(\d+)\]\s+(.*)',t)
    if m: ref_entries[int(m.group(1))]=m.group(2); ref_indices.append(i)
existing=set(v.strip().lower()[:80] for v in ref_entries.values())

new_refs = [
    'ROGNER H H. An assessment of world hydrocarbon resources [J]. Annual Review of Energy and the Environment, 1997, 22(1): 217-262.',
    'IEA. World Energy Employment 2025 [R]. Paris: International Energy Agency, 2025.',
    'HE X, ZHANG P, GAO Y, et al. Challenges and countermeasures for beneficial development of unconventional oil and gas resources in China [J]. China Petroleum Exploration, 2025, 30(1): 28-43.',
    'KING G E. Thirty years of gas shale fracturing [C]. SPE-133456, SPE ATCE, Florence, 2010.',
    'MAXWELL S C. Microseismic hydraulic fracture imaging [J]. Interpretation, 2014, 2(3): SJ1-SJ13.',
    'CIPOLLA C L, WALLACE J. Stimulated reservoir volume: A misnomer? [C]. SPE-168596, SPE HFTC, The Woodlands, 2014.',
    'MOLENAAR M M et al. First downhole application of DAS for fracturing diagnostics [J]. SPE Drilling & Completion, 2012, 27(1): 32-38.',
    'MALYAVKO E et al. Research of operational dynamics of marked proppant between hydraulic fractures [C]. SPE-215624, SPE IHFTC, Muscat, 2023.',
    'LIANG F et al. A comprehensive review on proppant technologies [J]. Petroleum, 2016, 2(1): 26-39.',
    'KORSMEYER R W et al. Mechanisms of solute release from porous hydrophilic polymers [J]. Int J Pharm, 1983, 15(1): 25-35.',
    'RITGER P L, PEPPAS N A. A simple equation for description of solute release I [J]. J Control Release, 1987, 5(1): 23-36.',
    'PEPPAS N A, SAHLIN J J. A simple equation for description of solute release III [J]. Int J Pharm, 1989, 57(2): 169-172.',
    'VAN GENUCHTEN M T, ALVES W J. Analytical solutions of the 1-D convective-dispersive solute transport equation [R]. USDA Tech Bull 1661, 1982.',
    'SHOOK G M, POPE G A, ASAKAWA K. Determining reservoir properties from tracer test analysis [C]. SPE-124614, SPE ATCE, New Orleans, 2009.',
    'LIU J et al. Interpretation method of multistage fracture tracer flowback curve [J]. ACS Omega, 2024, 9: 10852-10864.',
    'FONTALVO E M et al. Physical interpretation of interwell partitioning tracer tests [J]. Transp Porous Media, 2025, 152: 21-45.',
    'TIAN W et al. Quantifying fracture interference using chemical tracers [C]. SPE-201292, SPE ATCE, Virtual, 2020.',
]

to_add=[rt for rt in new_refs if rt.strip().lower()[:80] not in existing]
print(f'Non-duplicates: {len(to_add)}')

nn=max(ref_entries.keys())+1; li=max(ref_indices)
for j,rt in enumerate(to_add):
    if li+j+1<len(doc.paragraphs) and not doc.paragraphs[li+j+1].text.strip():
        run=doc.paragraphs[li+j+1].add_run(f'[{nn+j}] {rt}'); run.font.name='Times New Roman'; run.font.size=Pt(9)
print(f'Added to list. Total: {nn+len(to_add)-1}')

# ===== STANDALONE CITATIONS =====
# Each new ref gets its own sentence + bracket. Target specific text positions.

# [7] Energy context — insert standalone sentences BEFORE 'However' or BETWEEN existing sentences
energy_refs = {
    0: 'Rogner [N] provided an early comprehensive assessment of global hydrocarbon resources.',
    1: 'The IEA [N] reports the continued dominant role of fossil fuels in the global energy mix.',
    2: 'He et al. [N] reviewed the challenges for unconventional resource development in China.',
    3: 'King [N] summarized three decades of shale gas fracturing experience.',
    4: 'Maxwell [N] established microseismic imaging as a key tool for spatial fracture diagnostics.',
    5: 'Cipolla and Wallace [N] critically examined the stimulated reservoir volume concept.',
}
for ref_idx, sentence in energy_refs.items():
    n = nn + list(to_add).index(new_refs[ref_idx]) if new_refs[ref_idx] in to_add else None
    # Find actual num
    for j,rt in enumerate(to_add):
        if rt.strip().lower()[:80] == new_refs[ref_idx].strip().lower()[:80]:
            n = nn + j; break
    if n is None: continue
    txt = sentence.replace('[N]', f'[{n}]')
    for run in doc.paragraphs[7].runs:
        if 'These resources now account' in run.text and f'[{n}]' not in run.text:
            run.text = run.text.replace('These resources now account', f'{txt} These resources now account')
            print(f'[7] cited [{n}]')
            break

# [11] Tracer/modeling context
tracer_refs = {
    6: 'Molenaar et al. [N] demonstrated the first downhole distributed acoustic sensing application for fracture diagnostics.',
    7: 'Malyavko et al. [N] investigated marked proppant transport dynamics between hydraulic fractures.',
    9: 'Korsmeyer et al. [N] established the foundational framework for solute release from polymeric matrices.',
    10: 'Ritger and Peppas [N] formalized the power-law description of Fickian and non-Fickian release.',
    11: 'Peppas and Sahlin [N] extended the model to account for coupled diffusion and polymer relaxation.',
    12: 'Van Genuchten and Alves [N] compiled analytical solutions for advection-dispersion transport.',
    13: 'Shook et al. [N] used tracer tests to determine reservoir properties from field data.',
    14: 'Liu et al. [N] developed a flowback curve interpretation method for multi-stage fractured wells.',
    15: 'Fontalvo et al. [N] provided physical interpretation of interwell partitioning tracer tests.',
    16: 'Tian et al. [N] quantified fracture interference and production allocation using chemical tracers.',
}
for ref_idx, sentence in tracer_refs.items():
    n = None
    for j,rt in enumerate(to_add):
        if rt.strip().lower()[:80] == new_refs[ref_idx].strip().lower()[:80]:
            n = nn + j; break
    if n is None: continue
    txt = sentence.replace('[N]', f'[{n}]')
    for run in doc.paragraphs[11].runs:
        if 'motivates the present study' in run.text and f'[{n}]' not in run.text:
            run.text = run.text.replace('study.', f'study. {txt}')
            print(f'[11] cited [{n}]')
            break

# [13] Material context
if 8 in [i for i,_ in enumerate(new_refs)]:
    for j,rt in enumerate(to_add):
        if 'LIANG F' in rt.upper()[:20]:
            n = nn + j
            for run in doc.paragraphs[13].runs:
                if 'Compounding this' in run.text and f'[{n}]' not in run.text:
                    run.text = run.text.replace('Compounding this', f'Liang et al. [{n}] provided a comprehensive review of proppant technologies. Compounding this')
                    print(f'[13] cited [{n}]')
                    break

doc.save(os.path.join(os.environ['USERPROFILE'],'Desktop','ESP-T_投稿文件','ESP-T_Final.docx'))

# RENUMBER
doc2 = Document(os.path.join(os.environ['USERPROFILE'],'Desktop','ESP-T_投稿文件','ESP-T_Final.docx'))
ref_entries2={}; ref_indices2=[]
for i in range(ref_start+1,len(doc2.paragraphs)):
    t=doc2.paragraphs[i].text.strip(); m=re.match(r'\[(\d+)\]\s+(.*)',t)
    if m: ref_entries2[int(m.group(1))]=m.group(2); ref_indices2.append(i)
first={}
for i,p in enumerate(doc2.paragraphs):
    if i>=ref_start or i in PROTECTED: continue
    for m in re.finditer(r'(?<!\d)\[(\d+)\](?!\d)',p.text):
        n=int(m.group(1))
        if n in ref_entries2 and n not in first: first[n]=i
    for m in re.finditer(r'(?<!\d)\[(\d+)[–-](\d+)\](?!\d)',p.text):
        a,b=int(m.group(1)),int(m.group(2))
        for n in range(a,b+1):
            if n in ref_entries2 and n not in first: first[n]=i
ordered=sorted(first.keys()); o2n={old:new for new,old in enumerate(ordered,1)}
for i,p in enumerate(doc2.paragraphs):
    if i>=ref_start or i in PROTECTED: continue
    for run in p.runs:
        t=run.text
        def repl(m):
            inner=m.group(1); parts=re.split(r'\s*,\s*',inner); new_p=[]
            for part in parts:
                part=part.strip(); rm=re.match(r'(\d+)[–-](\d+)',part)
                if rm:
                    a,b=int(rm.group(1)),int(rm.group(2))
                    na=o2n.get(a,a); nb=o2n.get(b,b)
                    new_p.append(f'{na}-{nb}' if na!=nb else str(na))
                else: n=int(part); new_p.append(str(o2n.get(n,n)))
            return '['+','.join(new_p)+']'
        nt=re.sub(r'\[(\d+(?:[–-]\d+)?(?:,\s*\d+(?:[–-]\d+)?)*)\]',repl,t)
        if nt!=t: run.text=nt
for idx in ref_indices2:
    p=doc2.paragraphs[idx]
    for r in list(p._element.findall(qn('w:r'))):
        if not r.findall(qn('w:drawing')): p._element.remove(r)
for j,old_num in enumerate(ordered):
    if old_num in ref_entries2 and j<len(ref_indices2):
        run=doc2.paragraphs[ref_indices2[j]].add_run(f'[{j+1}] {ref_entries2[old_num]}')
        run.font.name='Times New Roman'; run.font.size=Pt(9)
doc2.save(os.path.join(os.environ['USERPROFILE'],'Desktop','ESP-T_投稿文件','ESP-T_Final.docx'))

doc3=Document(os.path.join(os.environ['USERPROFILE'],'Desktop','ESP-T_投稿文件','ESP-T_Final.docx'))
imgs=sum(len(r._element.findall(qn('w:drawing'))) for p in doc3.paragraphs for r in p.runs)
ref_n=sum(1 for i in range(ref_start+1,len(doc3.paragraphs)) if doc3.paragraphs[i].text.strip().startswith('['))
print(f'FINAL: {ref_n} refs, {imgs} images, {len(ordered)} cited')
