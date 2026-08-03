# -*- coding: utf-8 -*-
"""Trim verbose sections: Conclusions (1384→~500) and Experiments (1536→~1100)."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt
from pathlib import Path

SRC = Path(r"c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP-T_投稿文件\ESP-T_Final_submit.docx")
doc = Document(str(SRC))

# ═══════════════════════════════════════════════════════════
# 1. REPLACE CONCLUSIONS (Section 4) with concise version
# ═══════════════════════════════════════════════════════════

concise_conclusions = (
    "The central challenge addressed in this work is the disconnect between tracer release kinetics, "
    "measured in batch experiments and fitted with empirical models, and the quantitative interpretation "
    "of tracer breakthrough signals at the wellhead. We addressed this by developing a two-component "
    "piecewise advection-dispersion transport model that decomposes the breakthrough curve (BTC) into "
    "a Gaussian pulse (shut-in accumulation slug) and an erfc tail (sustained matrix-diffusion-controlled "
    "release), linked by a smooth tanh transition, and by designing an oleophilic epoxy/Fe₃O₄ "
    "tracer proppant (ESP-T) to validate the model."  # ₃ = subscript 3, ₄ = subscript 4
)

concise_modeling = (
    "The dual-component model achieves R² = 0.9939 for single-phase BTCs and is statistically "
    "decisive over four simpler alternatives (ΔAICc = 32.7, F-test p < 10⁻⁶, "
    "Akaike weight > 0.9999). The erfc tail accounts for 47% of the integrated tracer signal, "
    "a result robust to a six-fold variation in the transition width (46.8–47.5%). "
    "The model is validated through three independent lines of evidence: (i) the fitted flow rate "
    "(0.46 mL/min) agrees with the pump setting (0.50 mL/min) within 8%; (ii) the fitted residence "
    "time (37.4 min) matches the calculated travel time (38.6 min) within 3%; and (iii) the Peclet "
    "number (Pe = 0.934) independently confirms the non-Fickian transport mechanism identified via "
    "Korsmeyer-Peppas kinetics (n = 0.45–0.85)."
)  # ² = ², Δ = Delta, ⁻⁶ = ⁻⁶, – = en dash

concise_twophase = (
    "Under steady-state two-phase flow, the oil-phase tracer mass flux (F_O = C_oil × Q_oil) "
    "eliminates the water-dilution artifact inherent in concentration measurements. "
    "F_O quantitatively tracks oil production rates across varying oil-water ratios "
    "(Pearson r = 0.97, RMSD = 8.3%), providing a pathway from wellhead tracer concentration "
    "data to per-interval oil production rates at the laboratory scale."
)  # × = ×

concise_material = (
    "The ESP-T proppant, synthesized by encapsulating stearic acid-modified nano-Fe₃O₄ "
    "(∼3.3 wt%) within an epoxy matrix via emulsion polymerization, exhibits oleophilic "
    "character (WCA 104.6°, oil filtration time reduced by 66%), thermal stability to "
    "357 °C, bulk density below water (0.646 g/cm³), and mechanical integrity meeting "
    "industry benchmarks (crush rate 2.9% at 52 MPa, acid solubility 3.3%)."
)  # ∼ = ∼, ° = °, ³ = ³

concise_outlook = (
    "The present validation is limited to single-interval laboratory experiments with dodecane "
    "as the model oil. Extending the approach to field-scale, multi-interval conditions with "
    "crude oil constitutes the necessary next step. The multi-element coding strategy "
    "(Mn/Zn/Cu/Eu/Dy dopants), combined with the demonstrated robustness of the signal "
    "decomposition and the flux-based production allocation framework, provides a foundation "
    "for per-stage production monitoring in multi-fractured horizontal wells without requiring "
    "downhole tools or repeated shut-ins."
)

# Find Section 4 heading and replace all subsequent content until end
found_s4 = False
conclusions_paras = [concise_conclusions, concise_modeling, concise_twophase,
                     concise_material, concise_outlook]
conclusion_idx = 0

for i, para in enumerate(doc.paragraphs):
    if "4. Conclusions" in para.text.strip():
        found_s4 = True
        # Clear next few paragraphs
        for j in range(i+1, min(i+20, len(doc.paragraphs))):
            t = doc.paragraphs[j].text.strip()
            if t.startswith("References") or t.startswith("["):
                break
            if len(t) > 40:  # substantive paragraph
                if conclusion_idx < len(conclusions_paras):
                    for run in doc.paragraphs[j].runs:
                        run.text = ""
                    if doc.paragraphs[j].runs:
                        doc.paragraphs[j].runs[0].text = conclusions_paras[conclusion_idx]
                    else:
                        r = doc.paragraphs[j].add_run(conclusions_paras[conclusion_idx])
                        r.font.size = Pt(11)
                    conclusion_idx += 1
        break

print(f"Conclusions: replaced {conclusion_idx} paragraphs (was ~1384 words → ~450 words)")

# ═══════════════════════════════════════════════════════════
# 2. TRIM EXPERIMENTS SECTION (remove redundant detail)
# ═══════════════════════════════════════════════════════════
# The synthesis descriptions are detailed. We trim the most verbose parts.

# 2.2 Fabrication — shorten the step-by-step nano-Fe3O4 synthesis
trim_targets = [
    ("Upon reaching the target temperature, 2.703 g",
     "Upon reaching 80 °C, 2.703 g (0.01 mol) FeCl₃ and 1.15 g (0.058 mol) "
     "FeCl₂·4H₂O were added with 2×10⁻⁵ mol MnCl₂·6H₂O "
     "under N₂ purge. After 10 min stirring, 5.5 mL ammonia was rapidly introduced (pH 10, 2 h). "
     "The black precipitate was magnetically separated, washed with anhydrous ethanol to neutrality, "
     "then sonicated with ethanolic stearic acid for oleophilic modification. The product was diluted "
     "to 100 mL for storage. Other transition metals (Zn, Cu, Co, Ni) and rare earths (Eu, Dy, Nd) "
     "can substitute Mn for multi-stage tracer coding [18]."),
]

for old_start, new_text in trim_targets:
    for para in doc.paragraphs:
        if para.text.strip().startswith(old_start):
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = new_text
            print(f"Trimmed: 2.2 Fabrication paragraph")
            break

# 2.4 Characterization — merge verbose lists
for para in doc.paragraphs:
    if "The microstructure, elemental composition and distribution" in para.text:
        for run in para.runs:
            run.text = ""
        trimmed = (
            "Microstructure and elemental distribution were characterized by SEM (ZEISS-Sigma 500) "
            "with EDS elemental mapping; surface morphology by optical microscopy (Leica DM2700P); "
            "thermal stability by TGA/DSC (TA Instruments Q500, air, 10 °C/min, 50–800 °C); "
            "wettability by water contact angle (OCA20, 5 µL droplets, n = 5). Physical and "
            "mechanical properties (bulk/apparent density, sphericity, roundness, acid solubility, "
            "crush rate) were evaluated per SY/T 5107-2016. Oil–water transport selectivity "
            "was assessed via packed-bed filtration time (2.0 g proppant, 20 mL fluid, 200-mesh screen)."
        )
        if para.runs:
            para.runs[0].text = trimmed
        print(f"Trimmed: 2.4 Characterization")
        break

# Remove redundant separate paragraphs in 2.4
to_clear_markers = [
    "Surface morphology and particle dispersion",
    "Water contact angle (WCA) measurements were carried out",
    "The physical and mechanical properties as well as field applicability",
    "The oil–water conductivity of the proppants was indirectly assessed",
]
for para in doc.paragraphs:
    for marker in to_clear_markers:
        if para.text.strip().startswith(marker):
            for run in para.runs:
                run.text = ""
            break
print(f"  Cleaned redundant characterization paragraphs")

# ═══════════════════════════════════════════════════════════
# 3. SAVE
# ═══════════════════════════════════════════════════════════
doc.save(str(SRC))

# Re-count
doc2 = Document(str(SRC))
total_w = sum(len(p.text.split()) for p in doc2.paragraphs if p.text.strip())
table_w = sum(len(cell.text.split()) for t in doc2.tables for r in t.rows for c in r.cells)
print(f"\nTrimmed word count: ~{total_w} (body) + ~{table_w} (tables) = ~{total_w+table_w} total")
print(f"Saved: {SRC}")
