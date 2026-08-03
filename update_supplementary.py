# -*- coding: utf-8 -*-
"""Update supplementary material: mark moved figures, keep remaining content."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import RGBColor
from pathlib import Path

BASE = Path(r"c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP-T_投稿文件")
SRC = BASE / "Supplementary_Material.docx"
DST = BASE / "Supplementary_Material_v2.docx"

doc = Document(str(SRC))

# ═══════════════════════════════════════════════════════════
# Add a header note explaining what was moved to main text
# ═══════════════════════════════════════════════════════════
first_para = doc.paragraphs[0]

moved_note = (
    "[Note: Figures S4 (Breakthrough Curve Fitting Detail), S5 (Model Selection), "
    "and S6 (Sigma Sensitivity Analysis), along with their corresponding tables "
    "(Table S5, Table S6) and the tracer proppant comparison table (Table S1/S3), "
    "have been moved to the main manuscript (Section 3.7, new Figures 3-8, 3-9, 3-10; "
    "Section 1, Table 1). These are the core modeling innovation figures. "
    "The remaining supplementary content below covers detailed characterization "
    "methods (Section S2), supplementary characterization results (Section S3), "
    "and the original supplementary data tables.]"
)

new_para = doc.add_paragraph(moved_note)
first_para._element.addprevious(new_para._element)

# ═══════════════════════════════════════════════════════════
# Mark S4, S5, S6 sections as moved
# ═══════════════════════════════════════════════════════════
sections_to_mark = {
    "S4. Breakthrough Curve Fitting Detail": "[MOVED TO MAIN TEXT — Section 3.7, Figure 3-8]",
    "S5. Breakthrough Curve Model Selection": "[MOVED TO MAIN TEXT — Section 3.7, Figure 3-9, Table 3-4]",
    "S6. Sigma Sensitivity Analysis": "[MOVED TO MAIN TEXT — Section 3.7, Figure 3-10, Table 3-5]",
    "S1. Comparison with Previously Reported Tracer Proppants": "[MOVED TO MAIN TEXT — Section 1, Table 1]",
}

for i, para in enumerate(doc.paragraphs):
    for section_title, marker in sections_to_mark.items():
        if section_title in para.text:
            # Add marker after this paragraph
            marker_para = doc.add_paragraph(marker)
            # Style as comment
            for run in marker_para.runs:
                run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
                run.font.italic = True
            para._element.addnext(marker_para._element)
            print(f"Marked: {section_title}")
            break

# ═══════════════════════════════════════════════════════════
# Save
# ═══════════════════════════════════════════════════════════
from docx.shared import RGBColor
doc.save(str(DST))
print(f"\nSaved supplementary material v2: {DST}")
print("Changes:")
print("  - Added header note explaining moved figures")
print("  - Marked S1, S4, S5, S6 as [MOVED TO MAIN TEXT]")
print("  - Remaining SM content intact (S2 characterization methods, S3 results)")
