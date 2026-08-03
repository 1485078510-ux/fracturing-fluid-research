# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt

doc = Document(r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP-T_投稿文件\ESP-T_v5_intro_optimized.docx')

ref_heading = doc.add_heading('References', level=1)

refs = [
    '[1] Miller C, Waters G, Rylander E. Evaluation of production log data from horizontal wells drilled in organic shales. SPE 144326, SPE North American Unconventional Gas Conference, The Woodlands, TX, 2011. https://doi.org/10.2118/144326-MS',
    '[2] King GE. Thirty years of gas shale fracturing: What have we learned? SPE 133456, SPE Annual Technical Conference and Exhibition, Florence, Italy, 2010. https://doi.org/10.2118/133456-MS',
    '[3] Cipolla CL, Wallace J. Stimulated reservoir volume: A misapplied concept? SPE 168596, SPE Hydraulic Fracturing Technology Conference, The Woodlands, TX, 2014. https://doi.org/10.2118/168596-MS',
    '[4] Hill AD, Zhu D. Production Logging: Theoretical and Interpretive Elements. 2nd ed. Richardson, TX: Society of Petroleum Engineers, 2021.',
    '[5] Molenaar MM, Hill DJ, Webster P, Fidan E, Birch B. First downhole application of distributed acoustic sensing for hydraulic-fracturing monitoring and diagnostics. SPE Drilling & Completion, 2012, 27(1): 32-38. https://doi.org/10.2118/140561-PA',
    '[6] Ugueto GA, Huckabee PT, Molenaar MM, Wyker B, Somanchi K. Perforation cluster efficiency of cemented plug and perf limited entry completions: Insights from fiber optics diagnostics. SPE 179124, SPE Hydraulic Fracturing Technology Conference, The Woodlands, TX, 2016. https://doi.org/10.2118/179124-MS',
    '[7] Jin G, Roy B. Hydraulic-fracture geometry characterization using low-frequency DAS signal. The Leading Edge, 2017, 36(12): 975-980. https://doi.org/10.1190/tle36120975.1',
    '[8] Maxwell SC. Microseismic hydraulic fracture imaging: The path toward optimizing shale gas production. The Leading Edge, 2011, 30(3): 340-346. https://doi.org/10.1190/1.3567266',
    '[9] Watkins JW, Mardock ES. Use of radioactive iodine as a tracer in water-flooding operations. Journal of Petroleum Technology, 1954, 6(9): 117-124. https://doi.org/10.2118/349-G',
    '[10] Sanni M, Al-Abbad M, Kokal S, Dugstad O, Hartvig S, Huseby O. Pushing the envelope of residual oil measurement: A field case study of a new class of inter-well chemical tracers. Journal of Petroleum Science and Engineering, 2018, 163: 538-545. https://doi.org/10.1016/j.petrol.2017.12.076',
    '[11] Patidar AK, Joshi D, Dristant U, Choudhury T. A review of tracer testing techniques in porous media specially attributed to the oil and gas industry. Journal of Petroleum Exploration and Production Technology, 2022, 12(12): 3339-3356. https://doi.org/10.1007/s13202-022-01526-w',
    '[12] Yang H, Guo K, Lin L, et al. Application of micro-substance tracer test in fractured horizontal wells. Journal of Petroleum Exploration and Production Technology, 2024, 14(5): 1235-1246.',
    '[13] Zhao B, Panthi K, Mohanty KK. Tracer eluting proppants for hydraulic fracture characterization. Journal of Petroleum Science and Engineering, 2020, 190: 107048. https://doi.org/10.1016/j.petrol.2020.107048',
    '[14] Zhou Y, Liu H, Gao J, et al. Coated proppants with self-suspension and tracer slow-release functions. Journal of Petroleum Science and Engineering, 2022, 208: 109645. https://doi.org/10.1016/j.petrol.2021.109645',
    '[15] Li N, Cheng Q, Gong Z, et al. Release kinetics of rare earth tracer from polymer-coated proppants for hydraulic fracture analysis. Geoenergy Science and Engineering, 2023, 227: 211782. https://doi.org/10.1016/j.geoen.2023.211782',
    '[16] Gong Z, Li N, Kang W, et al. Novel oleophilic tracer-slow-released proppant for monitoring the oil production contribution. Fuel, 2024, 364: 130945. https://doi.org/10.1016/j.fuel.2024.130945',
    '[17] Ren J, Wang Y, Zhang H, et al. Optimizing the preparation of multi-colored dye-tracer proppants: A potential approach for quantitative localization and volume assessment of proppant flowback in multistage fractured horizontal wells. Geoenergy Science and Engineering, 2024, 241: 213053. https://doi.org/10.1016/j.geoen.2024.213053',
    '[18] Malyavko E, Upadhye V, Husein N. Research of operational dynamics of marked proppant transport between hydraulic fractures. SPE 215624, SPE International Hydraulic Fracturing Technology Conference, Muscat, Oman, 2023. https://doi.org/10.2118/215624-MS',
    '[19] Korsmeyer RW, Gurny R, Doelker E, Buri P, Peppas NA. Mechanisms of solute release from porous hydrophilic polymers. International Journal of Pharmaceutics, 1983, 15(1): 25-35. https://doi.org/10.1016/0378-5173(83)90064-9',
    '[20] Ritger PL, Peppas NA. A simple equation for description of solute release I. Fickian and non-Fickian release from non-swellable devices. Journal of Controlled Release, 1987, 5(1): 23-36. https://doi.org/10.1016/0168-3659(87)90034-4',
    '[21] Peppas NA, Sahlin JJ. A simple equation for the description of solute release III. Coupling of diffusion and relaxation. International Journal of Pharmaceutics, 1989, 57(2): 169-172. https://doi.org/10.1016/0378-5173(89)90306-2',
    '[22] Li H, Liu Z, Li Y, et al. Evaluation of the release mechanism of sustained-release tracers from epoxy resin matrices. ACS Omega, 2021, 6(29): 19269-19280. https://doi.org/10.1021/acsomega.1c02759',
    '[23] Gong Z, Li N, Qin M, et al. Magnetic nano-Fe3O4-based oleophilic tracer for stability studies of nano-tracer in oilfield conditions. Colloids and Surfaces A, 2024, 683: 132998. https://doi.org/10.1016/j.colsurfa.2023.132998',
    '[24] van Genuchten MT, Alves WJ. Analytical solutions of the one-dimensional convective-dispersive solute transport equation. USDA Technical Bulletin No. 1661, 1982.',
    '[25] Shook GM, Pope GA, Asakawa K. Determining reservoir properties and flood performance from tracer test analysis. SPE 124614, SPE Annual Technical Conference and Exhibition, New Orleans, LA, 2009. https://doi.org/10.2118/124614-MS',
    '[26] Fontalvo EM, Oliveira MC, Schoeggl F, et al. Physical interpretation of interwell partitioning tracer tests. Transport in Porous Media, 2025, 152: 21-45. https://doi.org/10.1007/s11242-024-02135-w',
    '[27] Velasco-Lozano M, Balhoff M, Diaz-Paulino L, et al. Modeling of chemical tracers for two-phase flow in advective-dominated porous media. SPE Journal, 2024, 29(7): 3718-3731. https://doi.org/10.2118/219475-PA',
    '[28] Mazo AB, Khamidullin MR, Potashev KA, et al. Mathematical model of water- and oil-soluble tracers transfer in multistage hydraulic fracturing. Fluid Dynamics, 2024, 59(3): 427-443. https://doi.org/10.1134/S0015462823602309',
    '[29] Shook GM, Ansari S, Tiah A. Tracer technology for subsurface flow. Springer, 2021. https://doi.org/10.1007/978-3-030-72264-2',
    '[30] Liang F, Sayed M, Al-Muntasheri GA, Chang FF, Li L. A comprehensive review on proppant technologies. Petroleum, 2016, 2(1): 26-39. https://doi.org/10.1016/j.petlm.2015.11.001',
    '[31] Zoveidavianpoor M, Gharibi A, Bin Jaafar MZ. Experimental characterization of a new high-strength ultra-lightweight composite proppant. Journal of Petroleum Science and Engineering, 2018, 170: 1038-1047. https://doi.org/10.1016/j.petrol.2018.06.064',
    '[32] Wang G, Ma Q, Ren L, et al. A comprehensive review of multifunctional proppants. ACS Omega, 2024, 9(44): 44120-44133. https://doi.org/10.1021/acsomega.4c06714',
    '[33] Wei M, Wang Y, Duan Y, et al. Screening and performance evaluation of epoxy resin long-term sustained-release solid tracer. International Journal of Oil, Gas and Coal Technology, 2024, 36(2): 170-196.',
    '[34] Fan Z, Liu H, Zhang Y, et al. Development of self-generated proppant based on modified low-density and low-viscosity epoxy resin and its evaluation. Petroleum Science, 2022, 19: 2127-2136. https://doi.org/10.1016/j.petsci.2022.05.009',
    '[35] Gong Z, Li N, Kang W, et al. Synthesis and characterization of an oleophilic Fe3O4-doped epoxy resin tracer proppant for hydraulic fracturing. Geoenergy Science and Engineering, 2025 (submitted).',
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.style = doc.styles['Normal']
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'

doc.save(r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP-T_投稿文件\ESP-T_v5_intro_optimized.docx')
print(f'OK: {len(refs)} references added. Total paragraphs: {len(doc.paragraphs)}')
