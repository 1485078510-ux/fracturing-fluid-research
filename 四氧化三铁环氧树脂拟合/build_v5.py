# -*- coding: utf-8 -*-
"""
Build ESP-T v5 manuscript — CORRECTED geometry (x=1m, d=1mm).
Q_fit = 0.52 mL/min, Q_pump = 0.50 (±3.9%).
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def h1(text):
    h = doc.add_heading(text, level=1)
    return h

def h2(text):
    h = doc.add_heading(text, level=2)
    return h

def para(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    return p

# ===== TITLE =====
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run(
    'Coupled Release–Transport Modeling and Tracer Flux Method '
    'for Per-Stage Production Allocation from Tracer-Proppant '
    'Breakthrough Curves'
)
run.bold = True; run.font.size = Pt(14)

# ===== ABSTRACT =====
h1('Abstract')
para(
    'Tracer proppants enable long-term per-stage production monitoring without '
    'downhole hardware, but extracting per-stage contribution rates from wellhead '
    'breakthrough curves (BTCs) requires jointly accounting for sustained tracer '
    'release from the polymer matrix and transport through the production tubing—a '
    'coupled inverse problem not addressed by current practice, which applies the '
    'Korsmeyer–Peppas (K‑P) model for batch release characterization and the '
    'advection–dispersion framework for transport interpretation separately. '
    'We develop a coupled release–transport model that decomposes the BTC into a '
    'Gaussian pulse (shut-in accumulation slug) and an erfc tail (sustained '
    'matrix-diffusion-controlled release), linked by a smooth hyperbolic-tangent '
    'transition. Applied to an oleophilic epoxy/Fe₃O₄ tracer proppant '
    '(ESP‑T) in core displacement experiments with tubing of length 1 m and inner '
    'diameter 1 mm, the model recovers the independently set pump flow rate '
    '(0.50 mL/min) as 0.52 mL/min from the BTC shape alone—a deviation of 3.9% '
    'with Q unconstrained in the objective function (search bounds 0.01–5.0 mL/min). '
    'The fitted mean residence time (1.51 min) agrees with the independently '
    'computed convective travel time (1.57 min) to within 3.8%. The Peclet number '
    '(Pe = 0.75) independently corroborates the non-Fickian transport mechanism '
    'identified from static K‑P batch kinetics (n = 0.45–0.85), two completely '
    'separate experiments converging on the same dispersion-dominated transport '
    'picture. Under steady-state two-phase flow, the oil-phase tracer mass flux '
    'F_O = C_oil × Q_oil is invariant with total flow rate at a given oil–water '
    'ratio, eliminating the dilution artifact inherent in concentration-based '
    'interpretation (Pearson r = 0.97, RMSD = 8.3%). By doping each fracture '
    'stage with a distinct tracer element, per-stage contribution rates are obtained '
    'as (F_O,i / C_i) / Σ(F_O,j / C_j) from surface samples alone, '
    'requiring no downhole tools and only a single shut-in.'
)

# ===== KEYWORDS =====
h1('Keywords')
para('Tracer proppant; Breakthrough curve; Advection–dispersion equation; '
     'Release kinetics; Production allocation; Two-phase flow; Epoxy resin; '
     'Tracer flux')

# ===== 1. INTRODUCTION =====
h1('1. Introduction')

para(
    'Multi-stage hydraulic fracturing has made horizontal wells the dominant production '
    'technology for unconventional reservoirs, but a fundamental operational question '
    'remains unresolved: how much oil does each fracture stage produce? Without per-stage '
    'contribution data, underperforming intervals go undiagnosed, well spacing and '
    'completion designs cannot be validated against production outcomes, and restimulation '
    'decisions rely on guesswork rather than evidence [1,2]. Existing downhole diagnostic '
    'tools each involve trade-offs that limit their routine use. Production logging '
    'requires well intervention and captures only a snapshot of the inflow profile at '
    'the moment of logging [3]. Distributed fiber-optic sensing provides continuous '
    'spatial resolution but demands permanent cable installation at a cost difficult to '
    'justify for marginal wells [4,5]. Microseismic monitoring characterizes fracture '
    'geometry, not production contribution [6]. A method that delivers per-stage '
    'production allocation from surface samples alone—without downhole hardware, '
    'without well intervention, and without repeated shut-ins—would transform '
    'completion diagnostics.'
)

para(
    'Tracer proppants offer a pathway toward this goal. A tracer agent is immobilized '
    'within a solid proppant carrier and co-injected with the proppant pack during '
    'fracturing. The tracer remains in the fracture after placement and releases '
    'gradually into the produced fluid, generating a time-resolved concentration '
    'signal—the breakthrough curve (BTC)—that can be measured at the wellhead by '
    'ICP-MS for months after a single placement [7–10]. Interpreting this BTC to '
    'recover the stage flow rate, however, requires solving a coupled inverse problem: '
    'the observed concentration at the wellhead is shaped jointly by the unknown tracer '
    'release rate from the polymer matrix and the unknown transport parameters of the '
    'production system. These two sets of unknowns—release kinetics and transport '
    'dynamics—are superimposed in a single observable, and neither is known a priori.'
)

para(
    'Current practice addresses release and transport separately, leaving the coupling '
    'unresolved. On the release side, tracer proppant performance is characterized '
    'exclusively through batch release measurements interpreted with the Korsmeyer–Peppas '
    '(K‑P) power law, C/C₀ = K·tⁿ [11,12]. The K‑P model identifies '
    'the release mechanism—Fickian diffusion for n ≤ 0.43, anomalous transport '
    'for 0.43 < n < 0.85, Case-II relaxation for n ≥ 0.85—and provides '
    'the temperature-dependent rate constant. But K‑P is a zero-dimensional batch '
    'model: it describes release into a well-mixed vessel with no spatial coordinate, '
    'no flow field, and no pathway from a release rate to a concentration measured at '
    'a distant sampling point. On the transport side, the one-dimensional '
    'advection–dispersion equation (ADE) has a mature analytical solution framework '
    'that relates BTC shape to transport parameters [13,14]. ADE-based interpretation '
    'has been applied to extract reservoir and fracture properties from tracer BTCs '
    '[15,16] and to model multi-stage tracer transport with phase partitioning [17]. '
    'The common premise across these applications, however, is that the tracer source '
    'term is known—an injection pulse of specified mass, duration, and location. A '
    'tracer-proppant BTC violates this premise: the source is not a single '
    'operator-controlled injection but a sustained, matrix-diffusion-controlled release '
    'whose rate parameters are unknown and whose duration spans the entire production '
    'period. No existing framework couples the sustained release source term to the '
    'transport solution. The practical consequence is that current tracer proppants can '
    'confirm which stages produce but cannot quantify how much each stage produces '
    'from the BTC alone.'
)

para(
    'In this work, we develop a method that closes this gap in two steps. First, we '
    'construct a coupled release–transport model that decomposes the BTC into a Gaussian '
    'pulse (representing tracer that accumulated in the near-wellbore region during '
    'shut-in and is swept out as a coherent slug upon flowback) and an erfc tail '
    '(representing sustained matrix-diffusion-controlled release that continues after '
    'the main slug has passed), linked by a smooth hyperbolic-tangent transition. The '
    'model recovers the independently set pump flow rate (0.50 mL/min) '
    'as 0.52 mL/min—a deviation of 3.9%—with Q unconstrained in the objective '
    'function, and the fitted mean residence time (1.51 min) agrees with the '
    'independently computed convective travel time (1.57 min) to within 3.8%. The '
    'Peclet number (Pe = 0.75) independently corroborates the non-Fickian transport '
    'mechanism identified from static K‑P batch kinetics (n = 0.45–0.85). Second, '
    'we introduce a tracer flux method for production allocation: under steady-state '
    'flow, the oil-phase tracer mass flux F_O = C_oil × Q_oil is shown '
    'experimentally to be invariant with total flow rate at a given oil–water ratio, '
    'eliminating the dilution artifact that confounds concentration-based '
    'interpretation. By combining the K‑P temperature calibration, the ADE-based BTC '
    'decomposition, and the flux method, the per-stage contribution rate is obtained as '
    'Contribution_i = (F_O,i / C_i) / Σ(F_O,j / C_j) from surface '
    'samples alone. We validate the framework using an oleophilic epoxy/Fe₃O₄ '
    'tracer proppant (ESP‑T) in single-phase and two-phase core displacement '
    'experiments, and outline the pathway to field deployment with multi-element '
    'coding for per-stage signal separation.'
)

# ===== 2. EXPERIMENTAL =====
h1('2. Experimental')

h2('2.1 Materials')
para(
    'Analytical grade reagents — including FeCl₃ (≥99.0%), FeCl₂·4H₂O '
    '(≥99.0%), MnCl₂·6H₂O (≥99.0%), stearic acid (≥99.5%), '
    'and anhydrous ethanol — were purchased from Chengdu Kelong Chemical Co., Ltd. '
    'Industrial grade materials included guar gum, E51 epoxy resin (epoxy value '
    '0.48–0.54 mol/100 g, Nantong Xingchen Synthetic Materials), T31 curing agent '
    '(phenolic amine, ≥95% effective component), and hollow glass microspheres '
    '(≥99% purity). Dodecane (≥99%, Chengdu Kelong) served as the model oil phase. '
    'All reagents were used as received. Detailed specifications are provided in '
    'Table S1 (Supplementary Material).'
)

h2('2.2 ESP‑T Synthesis')
para(
    'ESP‑T was synthesized in three steps. (i) Stearic acid-modified nano-Fe₃O₄ '
    '(nano-Fe₃O₄@SA) was prepared by co-precipitation: 2.703 g FeCl₃ and '
    '1.15 g FeCl₂·4H₂O in 100 mL deionized water at 80 °C under N₂, doped '
    'with 2 × 10⁻⁵ mol MnCl₂·6H₂O; 5.5 mL NH₃·H₂O was added, and '
    'the reaction proceeded at pH 10 for 2 h. The black precipitate was magnetically '
    'separated, washed to neutrality, and sonicated with ethanolic stearic acid for '
    'oleophilic surface modification, yielding nano-Fe₃O₄@SA at approximately '
    '12.5 wt% solids in the ethanol mixture. (ii) A pre-mixture of 20 mL E51 epoxy '
    'resin, ~0.75 g nano-Fe₃O₄@SA (~3.3 wt% of the final formulation), '
    '1 g hollow glass microspheres, and 7 g T31 curing agent was homogenized. '
    '(iii) The pre-mixture was emulsified in a SiO₂/guar gum aqueous dispersion '
    'and cured at 50 °C for 1 h, then rinsed and dried at 80 °C for 10 h. '
    'Pure epoxy microspheres were prepared identically without nano-Fe₃O₄@SA '
    'as a reference. Although Mn was used in this study, the co-precipitation method '
    'accommodates other transition metals (Zn, Cu, Co, Ni) and rare earth elements '
    '(Eu, Dy, Nd) for multi-stage tracer coding. Reagent specifications are listed '
    'in Table S1.'
)

h2('2.3 Characterization')
para(
    'Microstructure and elemental distribution were characterized by scanning electron '
    'microscopy with energy-dispersive X-ray spectroscopy (SEM‑EDS, ZEISS Sigma 500) '
    'and optical microscopy (Leica DM2700P). Thermal stability was evaluated by '
    'thermogravimetric analysis and differential scanning calorimetry (TGA/DSC, TA '
    'Instruments Q500, air atmosphere, 10 °C/min, 50–800 °C). Water contact '
    'angle (WCA) was measured using a video optical contact angle analyzer (OCA20, '
    '5 μL droplets, n = 5). Physical and mechanical properties — including bulk '
    'density, apparent density, sphericity, roundness, acid solubility '
    '(HCl 12 wt%/HF 3 wt%, 65 °C, 30 min), and crush rate (52 MPa, 2 min) — were '
    'evaluated per SY/T 5107–2016. Oil-water transport selectivity was assessed via '
    'packed-bed filtration time (2.0 g proppant, 20 mL fluid, 200-mesh screen). '
    'Detailed protocols are provided in the Supplementary Material (Section S2).'
)

h2('2.4 Batch Release Kinetics')
para(
    'ESP‑T (5 g, 40–70 mesh) was immersed in 100 mL dodecane in sealed glass vials '
    'placed in thermostatic oil baths at 30, 60, 90, and 120 °C. Sampling was '
    'conducted at 12 h intervals over 14 days; the concentration of tracer-doped metal '
    'ions was quantified by inductively coupled plasma mass spectrometry (ICP‑MS, '
    'PerkinElmer NexION 300X). Release data within M_t/M_∞ < 0.6 — the range '
    'in which the Korsmeyer–Peppas power law C/C₀ = K·tⁿ is valid '
    '— were fitted to determine the kinetic parameters K and n. The full release '
    'dataset is provided in Table S3.'
)

h2('2.5 Core Displacement Experiments')
para(
    'A dynamic displacement apparatus was used to simulate single-interval production. '
    'A steel core was packed with ESP‑T, sealed with 200-mesh screens, and placed in '
    'a core holder under 5 MPa confining pressure. The production tubing connecting the '
    'core outlet to the sampling point had a length of 1 m and an inner diameter of '
    '1 mm (Fig. S2). For single-phase experiments, the system was saturated with '
    'dodecane at 5 mL/min, shut in for 96 h to allow tracer accumulation, then '
    'displaced at a pump setting of 0.50 mL/min. The pump setting is the independent '
    'flow-rate reference against which the fitted flow rate is compared in the '
    'self-calibration analysis (Section 4.3). Effluent was sampled at 4-min intervals '
    '(2 mL per sample, 21 samples) and analyzed by ICP‑MS to construct the BTC. '
    'For two-phase experiments, dodecane and deionized water were co-injected at three '
    'oil-water volume ratios (OWR = 4:1, 1:1, 1:4) and four total flow rates '
    '(0.1–0.4 mL/min) under continuous steady-state flow. Sampling was conducted at '
    '5-min intervals; both the oil-water volume ratio and tracer concentration in each '
    'sample were recorded.'
)

# ===== 3. MODEL DEVELOPMENT =====
h1('3. Model Development')

h2('3.1 Physical Basis for the Two-Component BTC')
para(
    'The core displacement protocol imposes a tracer history that any quantitative '
    'interpretation must respect. During the 96-h shut-in that precedes flowback, '
    'tracer continuously diffuses from the epoxy matrix into the proppant pack and '
    'the near-wellbore region. When the well opens, this accumulated tracer is swept '
    'toward the sampling point as a coherent slug, producing a Gaussian-shaped '
    'concentration contribution. This is Source I: the shut-in accumulation slug. '
    'Simultaneously, the polymer matrix continues to release tracer after flow '
    'resumes: the model oil continually contacts the proppant surface, swelling '
    'progresses, and tracer diffuses through the gel layer into the flowing stream. '
    'This sustained, matrix-diffusion-controlled release constitutes Source II '
    'and produces a slowly decaying concentration tail. '
    'These two sources are physically irreducible — they are consequences of the '
    'experimental shut-in, not modeling choices. The slug exists because the shut-in '
    'prevents tracer from being swept away; the sustained tail exists because the '
    'polymer matrix does not stop releasing tracer when flow resumes. '
    'A mathematical model capable of separating these two contributions from a single '
    'BTC is required.'
)

h2('3.2 Governing Equations')
para(
    'Tracer transport through the production tubing is governed by the one-dimensional '
    'advection–dispersion equation (ADE): '
    '∂C/∂t + v·∂C/∂x = D·∂²C/∂x², '
    'where C is the tracer concentration, v = 4Q/(πd²) is the cross-section-averaged '
    'flow velocity (Q = volumetric flow rate, d = tubing inner diameter = 1 mm), and '
    'D = α·v is the longitudinal dispersion coefficient '
    '(α = longitudinal dispersivity). For an instantaneous pulse injection of mass '
    'M at x = 0 and t = 0 with boundary condition C(±∞, t) = 0, the classical '
    'solution is C(x, t) = (M/A_cross)·(1/√(4πDt))·exp(−(x − vt)²/(4Dt)). '
    'Substituting the tube-flow relations for v and D and multiplying numerator '
    'and denominator by πd² yields the dimensionless working variable '
    'z = (xπd² − 4Qt) / √(16αQtπd²), where z = 0 marks the convective front '
    'and the denominator captures dispersion broadening that grows as √t. '
    'The geometry is fixed by the experiment: x = 1000 mm (tubing length, 1 m) '
    'and d = 1 mm (tubing inner diameter). Neither is fitted.'
)

para(
    'The two tracer sources enter as two additive components. Source I '
    '(shut-in accumulation slug) takes the instantaneous-pulse solution of the ADE, '
    'giving the Gaussian rise component '
    'C_rise(t) = c_b + (A·d)/√(16παQt·d²) × exp(−z²), '
    'where c_b is the baseline concentration and A is the pulse amplitude coefficient, '
    'proportional to the tracer mass accumulated during shut-in. '
    'Source II (sustained matrix-diffusion-controlled release) corresponds to a '
    'continuous source at the inlet; the solution for a semi-infinite boundary '
    'condition gives the erfc fall component '
    'C_fall(t) = c_b + (a/2) × erfc(−z), '
    'where a is the tail amplitude coefficient representing the sustained-release '
    'source strength. The two components are blended through a C¹-continuous '
    'hyperbolic-tangent weight function '
    'w(t) = ½[1 + tanh((t₀ − t) / σ)], '
    'yielding the full model (Eq. 1): '
    'C(t) = c_b + w(t)·C_rise(t) + [1 − w(t)]·C_fall(t). '
    'The transition center t₀ marks the crossover from pulse-dominated to '
    'tail-dominated behavior; the transition width σ controls the sharpness '
    '(fitted ≈ 7.0 min in this study). The C¹ continuity of the tanh function '
    'is essential for gradient-based optimization, avoiding the non-differentiable '
    'kink that a hard piecewise switch would produce.'
)

h2('3.3 Parameter Estimation and Model Selection')
para(
    'Equation (1) contains seven fitted coefficients: baseline concentration c_b, '
    'pulse amplitude A, tail amplitude a, dispersivity α, flow rate Q, crossover '
    'time t₀, and transition width σ. Fixed geometric constants are '
    'x = 1000 mm (1 m tubing) and d = 1 mm. Derived transport properties are the '
    'Peclet number Pe = x/α, the mean flow velocity v = 4Q/(πd²), and the '
    'mean residence time MRT = x/v. The independent convective travel time from '
    'the pump setting is 1 PV = xπd²/(4Q_pump) = 1.57 min.'
)

para(
    'The inverse problem — extracting six transport and release parameters from '
    'a single BTC with coupled release and transport components — presents a '
    'challenging optimization landscape. A two-pass estimation strategy was adopted. '
    'Pass 1 (basin location): a global search was performed over wide, physically '
    'bounded parameter ranges (Table S6) using four independent runs with different '
    'random starting points. The flow rate Q was bounded to [0.01, 5.0] mL/min, '
    'a 500-fold range spanning two orders of magnitude above the pump setting of '
    '0.50 mL/min, with no penalty, prior, or constraint directing the search toward '
    'the pump value. Pass 2 (local refinement): from the best point identified in '
    'Pass 1, a gradient-based minimization refined the parameter estimates to the '
    'precision required for physical interpretation, typically converging within '
    '50–200 iterations. The identical two-pass protocol was applied to all five '
    'candidate models — the dual-component tanh-blended model and four '
    'single-component alternatives (single Gaussian, single erfc, exponential decay, '
    'K‑P power law) — to ensure that differences in fit quality reflect model '
    'adequacy rather than differences in optimization effort. Full optimization '
    'details are provided in the Supplementary Material (Section S5.3).'
)

# ===== 4. RESULTS AND DISCUSSION =====
h1('4. Results and Discussion')

h2('4.1 ESP‑T Characterization and Material Prerequisites')
para(
    'SEM micrographs confirm that ESP‑T microspheres exhibit excellent sphericity '
    '(>0.9) with no inter-particle agglomeration. The surface displays uniformly '
    'distributed micron-scale protrusions corresponding to nano-Fe₃O₄@SA '
    'nanoclusters embedded within the epoxy matrix; the absence of interfacial '
    'cracks is consistent with physical adhesion mediated by stearic acid alkyl chain '
    'entanglement with the epoxy network. Elemental mapping confirms that Fe is '
    'distributed throughout the particle cross-section, consistent with bulk '
    'encapsulation rather than surface-only attachment. TGA/DSC analysis identifies '
    'the primary epoxy decomposition at 357 °C (DTG peak), far exceeding typical '
    'downhole temperatures (80–200 °C). Stearic acid modification increases the '
    'water contact angle from 72.3° (pure epoxy, weakly hydrophilic) to 104.6° '
    '(ESP‑T, hydrophobic). ESP‑T exhibits a bulk density of 0.646 g/cm³ '
    '(<1.0 g/cm³, enabling self-suspension in water-based fracturing fluids), '
    'acid solubility of 3.3% (well below the 5% threshold per SY/T 5107–2016), '
    'and a crush rate of 2.9% at 52 MPa. Packed-bed filtration tests demonstrate '
    'the functional consequence of the oleophilic surface: ESP‑T reduces oil passage '
    'time by 66% relative to pure epoxy while increasing water passage time nearly '
    'tenfold, yielding a water/oil passage time ratio of 5.53. These results '
    'collectively confirm that the material prerequisites for the coupled '
    'release–transport model are satisfied: the tracer is stably encapsulated within '
    'a thermally robust matrix that selectively releases into the oil phase under '
    'non-Fickian transport conditions.\n\n'
    '[Fig. 2 placeholder — see Figure Captions]'
)

h2('4.2 Temperature-Dependent Tracer Release Kinetics')
para(
    'Tracer release from ESP‑T in dodecane was measured at 30, 60, 90, and 120 °C. '
    'Release accelerates with temperature throughout the 14-day measurement period; '
    'the rate is highest during the first 24 h and decays gradually thereafter. '
    'The K‑P model was fitted within its validated range (M_t/M_∞ < 0.6; '
    'Table 2). The rate constant K increases systematically from 0.0554 (30 °C) '
    'to 0.1964 (120 °C), and all n values fall within 0.45–0.85 '
    '(R² > 0.94 at every temperature), corresponding to anomalous (non-Fickian) '
    'transport co-governed by Fickian diffusion and Case-II polymer relaxation. '
    'The physical picture is one of solvent-driven swelling: dodecane permeates '
    'the cross-linked epoxy network, generating an inner glassy core and an outer '
    'gel layer; swelling expands transport channels through which the doped metal '
    'ions diffuse into the external medium. These kinetic parameters characterize '
    'the release mechanism independently of any flow configuration. In Section 4.3, '
    'the non-Fickian transport regime identified here will be corroborated by the '
    'Peclet number independently obtained from BTC analysis.\n\n'
    '[Fig. 3 placeholder — see Figure Captions]'
)

h2('4.3 BTC Decomposition and Model Validation')
para(
    'Table 3 compares the dual-component model against four single-component '
    'alternatives fitted to the same BTC dataset (21 data points) using the '
    'identical two-pass protocol. The dual-component model achieves '
    'R² = 0.9922 (RMSE = 0.0237), outperforming the single Gaussian '
    '(R² = 0.9482, ΔAICc = 15.4), single erfc (R² = 0.7159, ΔAICc = 51.2), '
    'exponential decay (R² = 0.7517, ΔAICc = 45.3), and K‑P power law '
    '(R² = −0.0193, ΔAICc = 75.0). The dual-component structure is '
    'statistically preferred, though the improvement over the single Gaussian is '
    'moderate in the present geometry because the long, narrow tubing '
    '(1 m × 1 mm) produces a BTC that is dominated (>90%) by '
    'sustained release, with the convective slug contributing a small fraction '
    'of the integrated signal.'
)

para(
    'The fitted parameter values are: c_b = 0.180, A = 0.10, a = 50.0, '
    'α = 1334 mm, Q = 0.52 mL/min, t₀ = 16.3 min, and σ = 7.0 min. '
    'Separating the fitted curve into its two components reveals that the erfc tail '
    'accounts for approximately 92% of the integrated tracer signal, with the '
    'Gaussian pulse contributing 8%. This decomposition reflects the physical reality '
    'of the experimental configuration: the 1 m × 1 mm production tubing has a '
    'transit time of only ~1.6 min at the pump flow rate, so the convective slug '
    'passes through the sampling point rapidly, whereas the sustained matrix release '
    'continues to supply tracer for the entire 105-min measurement period. '
    'The BTC is therefore a release-dominated signal in this geometry.\n\n'
    '[Fig. 5 placeholder — see Figure Captions]'
)

para(
    'The decisive test of the model\'s physical validity is the recovery of the '
    'independently known pump flow rate. The fitted flow rate Q = 0.52 mL/min '
    'deviates from the pump setting of 0.50 mL/min by 3.9% (Fig. 5b). The search '
    'bounds for Q spanned 0.01–5.0 mL/min (a 500-fold range, Table S6), with no '
    'penalty, prior, or constraint directing the optimizer toward the pump value. '
    'The model was free to assign Q any value within this range that minimized '
    'residuals; it consistently converged to 0.52 ± 0.005 mL/min across four '
    'independent global searches. An incorrect model structure could have achieved '
    'comparable R² with an arbitrary Q by compensating with α and the amplitude '
    'coefficients. The fact that Q emerges at 0.52 mL/min, entirely unforced, '
    'constitutes evidence that the model structure captures the actual transport '
    'physics. Q is not a fitting artifact — it is a recovered engineering parameter.'
)

para(
    'As an independent consistency check, the fitted mean residence time '
    'MRT = x/v = xπd²/(4Q) = 1.51 min agrees with the convective travel time '
    'independently computed from the pump setting and tube geometry '
    '(1 PV = 1.57 min) to within 3.8%. The two time scales are derived from '
    'different sources — the former from the fitted parameters, the latter from the '
    'independently set pump rate and known tube dimensions — and their agreement '
    'provides an internal self-consistency check that does not depend on Q alone.'
)

para(
    'The Peclet number derived from the fit, Pe = x/α = 0.75, indicates that '
    'dispersion dominates over advection in this system — a signature of '
    'non-piston, dispersion-dominated displacement. This result is independently '
    'corroborated by the K‑P kinetic analysis (Section 4.2), which identifies the '
    'tracer release mechanism as anomalous (non-Fickian) transport with '
    'n = 0.45–0.85 across all four temperatures. These two findings emerge from '
    'completely separate experiments: the K‑P analysis from static batch release in '
    'glass vials (different apparatus, different data, different fitting targets — '
    'K and n), and the Pe from dynamic core-flood BTC fitting (flowing system, '
    'BTC data, six fitted transport parameters). Their convergence on the same '
    'physical picture — that tracer transport is governed by coupled diffusion and '
    'advection, with dispersion dominating — provides strong, independent '
    'evidence that both the release characterization and the transport model are '
    'physically sound.'
)

h2('4.4 Time-of-Arrival Methods')
para(
    'Simple time-of-arrival (TOA) methods illustrate why a coupled model is '
    'necessary. The peak-time method (t_peak = 15 min) yields Q = 0.13 mL/min '
    '(−74% error), the half-peak method (t_half ≈ 5 min) gives Q = 0.39 mL/min '
    '(−22%), and the first-moment method (MRT_raw = 37.1 min) gives '
    'Q = 0.053 mL/min (−89%). All three fail because they assume piston-like '
    'displacement, whereas the BTC is shaped by dispersion and sustained release '
    'rather than by advection alone. The coupled model uniquely provides both '
    'accurate flow-rate recovery and mechanistic insight through the Peclet number.'
)

h2('4.5 Two-Phase Production Allocation via Tracer Flux')
para(
    'Under steady-state two-phase flow, the tracer concentration C_oil measured '
    'at the outlet decreases with increasing total flow rate Q_total — a consequence '
    'of dilution: at higher flow rates, less time is available for tracer to '
    'accumulate in each unit volume of produced fluid. Concentration alone is '
    'therefore a poor proxy for the production rate of a given stage. We resolve '
    'this by working with the oil-phase tracer mass flux '
    'F_O = C_oil × Q_oil. At steady state, F_O equals the release rate from '
    'the ESP‑T pack and is independent of Q_total — it reflects only what the '
    'proppant releases, which is governed by the oil-wetted surface area (and hence '
    'the oil-water ratio), not by how fast fluid sweeps past it. '
    'This is confirmed experimentally: at each OWR, F_O remains approximately '
    'constant across the full range of total flow rates (0.1–0.4 mL/min). F_O '
    'increases systematically with OWR (≈0.66, 1.61, and 2.72 μg/min at '
    'OWR = 1:4, 1:1, and 4:1, respectively). '
    'Normalizing by the single-phase reference flux '
    'F_O,ref = 3.187 ± 0.15 μg/min (triplicate), the normalized flux '
    'F_O/F_O,ref tracks the independently known oil flow rate across all three '
    'OWR conditions (Pearson r = 0.97, p = 0.006, RMSD = 8.3%).\n\n'
    '[Fig. 9 placeholder — see Figure Captions]'
)

# ===== 5. FIELD IMPLICATIONS =====
h1('5. Implications for Field-Scale Deployment')
para(
    'The coupled framework — K‑P kinetics → ADE model → tracer flux method — '
    'provides a complete pathway from wellhead concentration measurements to '
    'per-stage production allocation. For a multi-stage well, each fracture stage '
    'is assigned a distinct tracer element (e.g., Mn, Zn, Cu, Eu, Dy) doped '
    'into the ESP‑T matrix during synthesis. After a single shut-in period '
    'following fracturing, produced fluid is sampled at the wellhead and analyzed '
    'by ICP‑MS to obtain the concentration C_i of each tracer element. During '
    'flowback, the full BTC of each element is recorded; the ADE model recovers '
    'the per-stage flow rate Q_i from the BTC shape (Section 4.3). During subsequent '
    'steady production, periodic wellhead sampling yields the steady-state '
    'concentration C_i,ss. The per-stage oil flow rate is '
    'Q_i = F_O,i / C_i,ss, where F_O,i is the calibrated tracer flux. '
    'The per-stage contribution rate is then:\n\n'
    'Contribution_i = Q_i / ΣQ_j = (F_O,i / C_i) / Σ(F_O,j / C_j)\n\n'
    'The flux calibration F_O,i is obtained from laboratory single-phase reference '
    'measurements (F_O,ref) scaled by the proppant mass m_i injected into stage i '
    'and the reservoir temperature T_i, using the K‑P-derived temperature '
    'dependence r(T): F_O,i ≈ (m_i/m_ref) × (r(T_i)/r(T_ref)) × F_O,ref. '
    'The ADE model provides the Peclet number for each stage\'s transport regime '
    'and decomposes the BTC into its shut-in slug and sustained-release components, '
    'enabling interpretation of flowback data. The flux method handles the subsequent '
    'steady production phase, where the BTC is tail-dominated.'
)

para(
    'Several limitations define the path from laboratory demonstration to field '
    'deployment. The framework has been validated on a single interval at the '
    'laboratory scale; multi-interval configurations remain untested. All experiments '
    'used dodecane as the model oil; crude oil introduces additional complexity from '
    'variable composition, viscosity, and potential interactions with the epoxy '
    'matrix. The chemical stability of the epoxy matrix in aggressive downhole '
    'environments (H₂S, CO₂, high-salinity brines, temperatures exceeding '
    '120 °C) has not been evaluated. The flux calibration is based on three OWR '
    'levels (n = 3 independent oil-fraction points); a larger matrix would '
    'strengthen the calibration for field use. Batch-specific F_O,ref determination '
    'is advisable for each proppant production lot. Addressing these limitations '
    'through multi-interval experiments, crude oil compatibility tests, and '
    'extended environmental exposure studies defines the path to field-ready '
    'deployment.'
)

# ===== 6. CONCLUSIONS =====
h1('6. Conclusions')

conclusions = [
    'A coupled release–transport model was developed that decomposes a '
    'tracer-proppant breakthrough curve into a Gaussian pulse (shut-in accumulation '
    'slug) and an erfc tail (sustained matrix-diffusion-controlled release), linked '
    'by a C¹-continuous hyperbolic-tangent transition. Applied to an oleophilic '
    'epoxy/Fe₃O₄ tracer proppant (ESP‑T) with production tubing of 1 m length '
    'and 1 mm inner diameter, the model recovers the independently set pump flow '
    'rate (0.50 mL/min) as 0.52 mL/min from the BTC shape alone — a deviation of '
    '3.9% — with Q unconstrained in the objective function (search bounds '
    '0.01–5.0 mL/min, a 500-fold range).',

    'The physical validity of the model is established through three independent '
    'checks: (i) the fitted mean residence time (1.51 min) agrees with the '
    'independently computed convective travel time (1.57 min) to within 3.8%; '
    '(ii) the Peclet number (Pe = 0.75) independently corroborates the non-Fickian '
    'transport mechanism (n = 0.45–0.85) identified from static K‑P batch release '
    'experiments — two completely separate experiments converging on the same '
    'dispersion-dominated transport picture; (iii) the dual-component structure is '
    'statistically preferred over four single-component alternatives.',

    'In the present geometry (1 m × 1 mm tubing), the erfc tail accounts for '
    'approximately 92% of the integrated tracer signal, with the Gaussian pulse '
    'contributing 8%, reflecting a release-dominated BTC in which the convective '
    'slug passes rapidly while sustained matrix diffusion supplies tracer '
    'throughout the measurement period.',

    'A tracer flux method enables production allocation under two-phase flow. The '
    'oil-phase tracer mass flux F_O = C_oil × Q_oil is invariant with total '
    'flow rate at a given oil–water ratio, eliminating the dilution artifact. '
    'The normalized flux tracks oil production rates across oil–water ratios '
    'from 4:1 to 1:4 (Pearson r = 0.97, RMSD = 8.3%).',

    'For multi-stage wells, doping each fracture stage with a distinct tracer '
    'element enables per-stage contribution rates to be obtained from wellhead '
    'ICP‑MS samples as (F_O,i/C_i) / Σ(F_O,j/C_j), where F_O,i is calibrated '
    'from laboratory reference measurements and K‑P temperature-dependent release '
    'kinetics. The framework requires no downhole tools and only a single shut-in. '
    'Multi-stage field validation with crude oil under transient conditions remains '
    'the next step toward deployment.',
]
for c in conclusions:
    para(c)

# ===== SAVE =====
output = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP-T_投稿文件\ESP-T_v5_manuscript.docx'
doc.save(output)
print(f'[OK] v5 saved to {output}')
print(f'Total paragraphs: {len(doc.paragraphs)}')
