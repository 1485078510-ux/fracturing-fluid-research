# -*- coding: utf-8 -*-
"""
Replace Introduction in the revised manuscript with a carefully crafted version.
5 paragraphs, tight funnel, each citation earns its place.
"""
from docx import Document
from docx.shared import Pt

DOCX = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP-T_投稿文件\ESP-T_Final_4-revised_revised.docx'
doc = Document(DOCX)

intro = [

# P1: FIELD PROBLEM — Why per-stage contribution matters, and why existing tools fall short
# Sets the stage for "we need a surface-only method"
(
'Horizontal wells with multi-stage hydraulic fracturing account for the '
'majority of unconventional oil and gas production, yet a fundamental '
'operational question remains unanswered for most wells: how much does each '
'fracture stage produce? Production logs from hundreds of horizontal wells '
'show that 30\u201340% of perforation clusters contribute negligibly [1], '
'while a minority of stages often dominate total output [2]. Without per-stage '
'data, operators cannot distinguish productive from unproductive intervals, '
'validate completion designs against production outcomes, or identify '
'candidates for re-stimulation. Existing diagnostic methods each address '
'part of this need but carry operational burdens that limit routine '
'deployment. Production logging requires well intervention via coiled tubing '
'or tractor in extended-reach horizontals and captures only a snapshot of '
'the inflow profile [3]. Distributed fiber-optic sensing\u2014both '
'distributed temperature sensing (DTS) and distributed acoustic sensing '
'(DAS)\u2014provides continuous spatial resolution without intervention '
'[4,5] and has enabled real-time cluster-efficiency diagnostics [6], but '
'permanent cable installation and surface instrumentation remain difficult '
'to justify for marginal wells [7]. Microseismic monitoring characterizes '
'fracture geometry and stimulated reservoir volume rather than production '
'contribution per stage [8]. A method that delivers per-stage production '
'allocation from surface samples alone\u2014without downhole tools, '
'permanent installation, or repeated intervention\u2014would transform '
'completion diagnostics.'
),

# P2: TRACER PROPPANTS + THE CORE DIFFICULTY
# Introduces the tracer solution, then exposes why it's not solved yet
(
'Chemical tracers offer a pathway toward this method. Since the first '
'documented oilfield tracer test using radioiodine in 1954 [9], tracer '
'technology has progressed from qualitative inter-well connectivity '
'assessment [10] to quantitative residual oil saturation measurement via '
'partitioning inter-well tracer tests [11], and more recently to stage-level '
'inflow profiling in multi-stage fractured wells [12,13]. Tracer proppants '
'\u2014 in which the tracer agent is immobilized within a solid proppant '
'carrier and co-injected with the proppant pack \u2014 offer a decisive '
'practical advantage: the tracer remains in the fracture after placement and '
'releases gradually over months, so a single wellhead sampling program can '
'support long-term monitoring without repeated injection. Several material '
'platforms have been reported, including ceramic carriers with organic dye '
'coatings [14], carbon quantum-dot-encapsulated ceramics [15], rare-earth-doped '
'polymer matrices for multi-element coding [16], oleophilic '
'Fe\u2083O\u2084/polystyrene microspheres for oil-phase selectivity [17], '
'multi-colored dye-tracer proppants for flowback assessment [18], and '
'quantum-dot-encoded polymer microspheres for fluid-phase-resolved inflow '
'profiling [19]. These demonstrations share a common limitation, however: '
'they can confirm which stages produce, but they cannot quantify how much '
'each stage contributes. The reason is fundamental: the tracer concentration '
'measured at the wellhead is shaped jointly by two unknown functions\u2014the '
'release rate from the polymer matrix and the transport dynamics of the '
'production system\u2014superimposed in a single observable. Neither is '
'known a priori, and current practice addresses them separately.'
),

# P3: RELEASE SIDE CEILING — K-P model, what it does and doesn't do
(
'On the release side, tracer-proppant performance is characterized through '
'batch measurements interpreted with the Korsmeyer\u2013Peppas (K\u2011P) '
'power law, C/C\u2080 = K\u00b7t\u207f [20,21]. The K\u2011P model '
'identifies the release mechanism through the diffusional exponent n '
'(Fickian diffusion for n \u2264 0.43, anomalous transport for '
'0.43 < n < 0.85, Case-II relaxation for n \u2265 0.85) and provides '
'the temperature-dependent rate constant K. Extended by Peppas and Sahlin '
'(1989) to explicitly couple diffusive and relaxational contributions [22], '
'the model has been productively applied to characterize release from '
'rare-earth-doped polymer coatings [16], polystyrene-encapsulated '
'Fe\u2083O\u2084 nanoparticles [17], epoxy-matrix sustained-release '
'systems [23], and nano-Fe\u2083O\u2084 tracers in copolymer matrices [24]. '
'However, K\u2011P is fundamentally a zero-dimensional batch description: '
'it characterizes release into a well-mixed, fixed-volume vessel with no '
'spatial coordinate and no flow field. It can determine how fast the tracer '
'leaves the proppant in a beaker; it cannot predict what concentration '
'will be observed at a distant wellhead sampling point, when that '
'concentration will peak, or how it will decay\u2014the features of a '
'breakthrough curve that encode production information.'
),

# P4: TRANSPORT SIDE CEILING + THE GAP
(
'On the transport side, the one-dimensional advection\u2013dispersion '
'equation (ADE), \u2202C/\u2202t + v\u00b7\u2202C/\u2202x = '
'D\u00b7\u2202\u00b2C/\u2202x\u00b2, provides a mature analytical '
'framework for interpreting breakthrough curves. Van Genuchten and Alves '
'(1982) catalogued analytical solutions for a wide range of boundary and '
'initial conditions [25]. Temporal moment analysis yields the mean residence '
'time (hence effective flow velocity) and variance (hence dispersivity) '
'from a single BTC [26]; full-curve nonlinear inversion has been applied '
'to extract reservoir properties from inter-well tracer tests [27], to '
'interpret partitioning tracer tests in layered reservoirs [28], to model '
'chemical-tracer transport under two-phase flow via analytical solutions '
'[29], and to describe multi-stage tracer transfer with phase partitioning '
'[30]. These applications share a critical premise: the tracer source term '
'is known\u2014an injection pulse of specified mass, duration, and '
'location. A tracer-proppant BTC violates this premise fundamentally. The '
'source is not a single operator-controlled injection but a sustained, '
'matrix-diffusion-controlled release whose rate depends on temperature, '
'solvent composition, and the time-varying concentration gradient at the '
'proppant\u2013fluid interface. No existing framework couples this '
'sustained, unknown release source to the ADE transport solution. The two '
'bodies of knowledge\u2014release kinetics and transport interpretation '
'\u2014have remained separate, and neither alone can deliver the per-stage '
'flow rate from the BTC.'
),

# P5: THIS WORK — What we did, how we validated it
(
'In this work, we close this gap. We construct a coupled '
'release\u2013transport model that decomposes the BTC into two physically '
'motivated components: a Gaussian pulse, representing tracer that '
'accumulated in the near-wellbore region during shut-in and is swept out '
'as a coherent slug upon flowback, and an erfc tail, representing sustained '
'matrix-diffusion-controlled release that continues after the main slug '
'has passed. The two components are linked by a C\u00b9-continuous '
'hyperbolic-tangent transition; six parameters are estimated simultaneously '
'from a single BTC by nonlinear least squares, with no parameter constrained '
'in the objective function. The model is validated through a predictive '
'self-calibration test: the effective flow rate Q is left unconstrained; '
'the model must recover the independently set pump flow rate from the BTC '
'shape alone. We apply the framework to an oleophilic epoxy/Fe\u2083O\u2084 '
'tracer proppant (ESP\u2011T). Epoxy resins were chosen as the proppant '
'matrix for their high mechanical strength, thermal stability, and chemical '
'resistance relative to thermoplastic alternatives [31\u201333]; their '
'capacity for sustained, diffusion-controlled tracer release has been '
'demonstrated in both water-soluble [23] and oil-soluble [34] configurations. '
'The model recovers Q = 0.46 mL/min against the independently set pump rate '
'of 0.50 mL/min (8% deviation), and the Peclet number (Pe = 0.934) '
'independently corroborates the non-Fickian transport mechanism identified '
'from static K\u2011P batch kinetics (n = 0.45\u20130.85)\u2014two '
'completely separate experiments converging on the same physical picture. '
'We then introduce a tracer flux method for two-phase production allocation: '
'the oil-phase tracer mass flux F_O = C_oil \u00d7 Q_oil is shown to be '
'invariant with total flow rate at a given oil\u2013water ratio (Pearson '
'r = 0.97, RMSD = 8.3%), eliminating the dilution artifact inherent in '
'concentration-based interpretation. By doping each fracture stage with a '
'distinct tracer element during synthesis\u2014the co-precipitation method '
'accommodates transition metals (Mn, Zn, Cu, Co, Ni) and rare earths '
'(Eu, Dy, Nd)\u2014per-stage contribution rates are obtained as '
'(F_O,i / C_i) / \u03a3(F_O,j / C_j) from periodic ICP\u2011MS wellhead '
'samples, requiring no downhole tools and only a single shut-in.'
),

]

# Find and replace
i0 = None; i1 = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if p.style.name.startswith('Heading') and t == '1. Introduction': i0 = i
    elif i0 is not None and p.style.name.startswith('Heading') and '2.' in t: i1 = i; break

if i0 and i1:
    for idx in reversed(range(i0+1, i1)):
        doc.paragraphs[idx]._element.getparent().remove(doc.paragraphs[idx]._element)
    ins = doc.paragraphs[i0]._element
    for text in reversed(intro):
        np = doc.add_paragraph(text); np.style = doc.styles['Normal']
        for r in np.runs: r.font.size=Pt(12); r.font.name='Times New Roman'
        ins.addnext(np._element)

doc.save(DOCX)
for i, t in enumerate(intro):
    print(f'P{i+1}: {len(t.split())} words')
print(f'Total: {sum(len(t.split()) for t in intro)} words')
print('Done')
