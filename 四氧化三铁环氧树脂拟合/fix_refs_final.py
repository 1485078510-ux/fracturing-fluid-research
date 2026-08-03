# -*- coding: utf-8 -*-
"""Apply corrected references to v5 manuscript."""
from docx import Document
from docx.shared import Pt

DOCX = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP-T_投稿文件\ESP-T_v5_intro_optimized.docx'
doc = Document(DOCX)

corrected_refs = [
    # [1]-[8] — Completion diagnostics (verified, no changes)
    '[1] Miller C, Waters G, Rylander E. Evaluation of production log data from horizontal wells drilled in organic shales. SPE 144326, SPE North American Unconventional Gas Conference, The Woodlands, TX, 2011. https://doi.org/10.2118/144326-MS',
    '[2] King GE. Thirty years of gas shale fracturing: What have we learned? SPE 133456, SPE Annual Technical Conference and Exhibition, Florence, Italy, 2010. https://doi.org/10.2118/133456-MS',
    '[3] Cipolla CL, Wallace J. Stimulated reservoir volume: A misapplied concept? SPE 168596, SPE Hydraulic Fracturing Technology Conference, The Woodlands, TX, 2014. https://doi.org/10.2118/168596-MS',
    '[4] Hill AD, Zhu D. Production Logging: Theoretical and Interpretive Elements. 2nd ed. Richardson, TX: Society of Petroleum Engineers, 2021.',
    '[5] Molenaar MM, Hill DJ, Webster P, Fidan E, Birch B. First downhole application of distributed acoustic sensing for hydraulic-fracturing monitoring and diagnostics. SPE Drilling & Completion, 2012, 27(1): 32-38. https://doi.org/10.2118/140561-PA',
    '[6] Ugueto GA, Huckabee PT, Molenaar MM, Wyker B, Somanchi K. Perforation cluster efficiency of cemented plug and perf limited entry completions: Insights from fiber optics diagnostics. SPE 179124, SPE Hydraulic Fracturing Technology Conference, The Woodlands, TX, 2016. https://doi.org/10.2118/179124-MS',
    '[7] Jin G, Roy B. Hydraulic-fracture geometry characterization using low-frequency DAS signal. The Leading Edge, 2017, 36(12): 975-980. https://doi.org/10.1190/tle36120975.1',
    '[8] Maxwell SC. Microseismic hydraulic fracture imaging: The path toward optimizing shale gas production. The Leading Edge, 2011, 30(3): 340-346. https://doi.org/10.1190/1.3567266',

    # [9]-[12] — Tracer technology (verified, no changes)
    '[9] Watkins JW, Mardock ES. Use of radioactive iodine as a tracer in water-flooding operations. Journal of Petroleum Technology, 1954, 6(9): 117-124. https://doi.org/10.2118/349-G',
    '[10] Sanni M, Al-Abbad M, Kokal S, Dugstad O, Hartvig S, Huseby O. Pushing the envelope of residual oil measurement: A field case study of a new class of inter-well chemical tracers. Journal of Petroleum Science and Engineering, 2018, 163: 538-545. https://doi.org/10.1016/j.petrol.2017.12.076',
    '[11] Patidar AK, Joshi D, Dristant U, Choudhury T. A review of tracer testing techniques in porous media specially attributed to the oil and gas industry. Journal of Petroleum Exploration and Production Technology, 2022, 12(12): 3339-3356. https://doi.org/10.1007/s13202-022-01526-w',
    '[12] Yang H, Guo K, Lin L, et al. Application of micro-substance tracer test in fractured horizontal wells. Journal of Petroleum Exploration and Production Technology, 2024, 14(5): 1235-1246. https://doi.org/10.1007/s13202-024-01765-9',

    # [13]-[18] — Tracer proppant materials
    '[13] Zhao B, Panthi K, Mohanty KK. Tracer eluting proppants for hydraulic fracture characterization. Journal of Petroleum Science and Engineering, 2020, 190: 107048. https://doi.org/10.1016/j.petrol.2020.107048',
    '[14] Zhou Y, Liu H, Gao J, et al. Coated proppants with self-suspension and tracer slow-release functions. Journal of Petroleum Science and Engineering, 2022, 208: 109645. https://doi.org/10.1016/j.petrol.2021.109645',
    '[15] Li N, Cheng Q, Gong Z, Ye X, Peng R, Li Q, Liu X, Li C. Release kinetics of rare earth tracer from polymer-coated proppants for hydraulic fracture analysis. Geoenergy Science and Engineering, 2023, 227: 211782. https://doi.org/10.1016/j.geoen.2023.211782',
    '[16] Gong Z, Li N, Kang W, Qin M, Wu Y, Liu X. Novel oleophilic tracer-slow-released proppant for monitoring the oil production contribution. Fuel, 2024, 364: 130945. https://doi.org/10.1016/j.fuel.2024.130945',
    # FIXED [17]: authors corrected from "Ren J, Wang Y, Zhang H" to actual authors
    '[17] Ren M, Yang B, Yang D, Liu Y, Zhang H, Zhang M, Zhang Y. Optimizing the preparation of multi-colored dye-tracer proppants: A potential approach for quantitative localization and volume assessment of proppant flowback in multistage fractured horizontal wells. Geoenergy Science and Engineering, 2024, 241: 213053. https://doi.org/10.1016/j.geoen.2024.213053',
    '[18] Malyavko E, Upadhye V, Husein N. Research of operational dynamics of a well with two hydraulic fractures with use of marked proppant penetrating into one productive formation. SPE 215624, SPE International Hydraulic Fracturing Technology Conference and Exhibition, Muscat, Oman, 2023. https://doi.org/10.2118/215624-MS',

    # [19]-[21] — K-P kinetics (classic, verified)
    '[19] Korsmeyer RW, Gurny R, Doelker E, Buri P, Peppas NA. Mechanisms of solute release from porous hydrophilic polymers. International Journal of Pharmaceutics, 1983, 15(1): 25-35. https://doi.org/10.1016/0378-5173(83)90064-9',
    '[20] Ritger PL, Peppas NA. A simple equation for description of solute release I. Fickian and non-Fickian release from non-swellable devices in the form of slabs, spheres, cylinders or discs. Journal of Controlled Release, 1987, 5(1): 23-36. https://doi.org/10.1016/0168-3659(87)90034-4',
    '[21] Peppas NA, Sahlin JJ. A simple equation for the description of solute release III. Coupling of diffusion and relaxation. International Journal of Pharmaceutics, 1989, 57(2): 169-172. https://doi.org/10.1016/0378-5173(89)90306-2',

    # FIXED [22]: DOI corrected from 1c02759 to 1c02748
    '[22] Li H, Liu Z, Li Y, Luo H, Cui X, Nie S, Ye K. Evaluation of the release mechanism of sustained-release tracers and its application in horizontal well inflow profile monitoring. ACS Omega, 2021, 6(29): 19269-19280. https://doi.org/10.1021/acsomega.1c02748',
    '[23] Gong Z, Li N, Qin M, Wu Y, Liu X. Magnetic nano-Fe3O4-based oleophilic tracer for stability studies of nano-tracer in oilfield conditions. Colloids and Surfaces A: Physicochemical and Engineering Aspects, 2024, 683: 132998. https://doi.org/10.1016/j.colsurfa.2023.132998',

    # [24]-[28] — ADE transport (verified, minor fixes)
    '[24] van Genuchten MT, Alves WJ. Analytical solutions of the one-dimensional convective-dispersive solute transport equation. USDA Technical Bulletin No. 1661, 1982.',
    '[25] Shook GM, Pope GA, Asakawa K. Determining reservoir properties and flood performance from tracer test analysis. SPE 124614, SPE Annual Technical Conference and Exhibition, New Orleans, LA, 2009. https://doi.org/10.2118/124614-MS',
    # FIXED [26]: updated volume/issue/article and DOI
    '[26] Fontalvo SD, Yutkin MP, Hassanizadeh SM, Radke CJ, Patzek TW. Physical interpretation of interwell partitioning tracer tests for estimation of remaining oil saturation in layered carbonate reservoirs. Transport in Porous Media, 2025, 152(9): 66. https://doi.org/10.1007/s11242-025-02196-y',
    # FIXED [27]: corrected SPE paper number from 219475 to 219730
    '[27] Velasco-Lozano M, Balhoff MT, Diaz-Paulino LE, Lopez-Ramirez S, Galvan-Castro R. Modeling of chemical tracers for two-phase flow in advective-dominated porous media at core scale. SPE Journal, 2024, 29(7): 3718-3731. https://doi.org/10.2118/219730-PA',
    '[28] Mazo AB, Khamidullin MR, Potashev KA, et al. Mathematical model of water- and oil-soluble tracers transfer in multistage hydraulic fracturing. Fluid Dynamics, 2024, 59(3): 427-443. https://doi.org/10.1134/S0015462823602309',
    '[29] Shook GM, Ansari S, Tiah A. Tracer technology for subsurface flow. In: Subsurface Flow: Modeling, Monitoring, and Management. Springer, 2021. https://doi.org/10.1007/978-3-030-72264-2',

    # [30]-[33] — Epoxy proppants
    # FIXED [30]: added missing authors Chang FF, Li L
    '[30] Liang F, Sayed M, Al-Muntasheri GA, Chang FF, Li L. A comprehensive review on proppant technologies. Petroleum, 2016, 2(1): 26-39. https://doi.org/10.1016/j.petlm.2015.11.001',
    # FIXED [31]: DOI corrected from .064 to .030
    '[31] Zoveidavianpoor M, Gharibi A, Bin Jaafar MZ. Experimental characterization of a new high-strength ultra-lightweight composite proppant derived from renewable resources. Journal of Petroleum Science and Engineering, 2018, 170: 1038-1047. https://doi.org/10.1016/j.petrol.2018.06.030',
    '[32] Wang G, Ma Q, Ren L, et al. A comprehensive review of multifunctional proppants. ACS Omega, 2024, 9(44): 44120-44133. https://doi.org/10.1021/acsomega.4c06714',
    '[33] Wei M, Wang Y, Duan Y, et al. Screening and performance evaluation of epoxy resin long-term sustained-release solid tracer. International Journal of Oil, Gas and Coal Technology, 2024, 36(2): 170-196. https://doi.org/10.1504/IJOGCT.2024.139452',

    # [34]-[35] — ESP-T / this work
    '[34] Fan Z, Liu H, Zhang Y, et al. Development of self-generated proppant based on modified low-density and low-viscosity epoxy resin and its evaluation. Petroleum Science, 2022, 19(5): 2127-2136. https://doi.org/10.1016/j.petsci.2022.05.009',
    '[35] Gong Z, Li N, Kang W, et al. Synthesis and characterization of an oleophilic Fe3O4-doped epoxy resin tracer proppant for hydraulic fracturing applications. Geoenergy Science and Engineering, 2025 (submitted).',
]

# Find and replace References section
ref_start = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith('Heading') and 'References' in p.text:
        ref_start = i
        break

if ref_start:
    # Remove old reference paragraphs
    old_indices = []
    for j in range(ref_start + 1, len(doc.paragraphs)):
        if doc.paragraphs[j].style.name.startswith('Heading'):
            break
        old_indices.append(j)
    for idx in sorted(old_indices, reverse=True):
        doc.paragraphs[idx]._element.getparent().remove(doc.paragraphs[idx]._element)

    # Insert corrected references
    from docx.oxml.ns import qn
    insert_after = doc.paragraphs[ref_start]._element
    for text in reversed(corrected_refs):
        new_p = doc.add_paragraph(text)
        new_p.style = doc.styles['Normal']
        for run in new_p.runs:
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
        insert_after.addnext(new_p._element)

    print(f'[OK] Replaced references: {len(old_indices)} old -> {len(corrected_refs)} corrected')
else:
    # No references section — add one
    doc.add_heading('References', level=1)
    for ref in corrected_refs:
        p = doc.add_paragraph(ref)
        p.style = doc.styles['Normal']
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
    print(f'[OK] Added new References section with {len(corrected_refs)} entries')

# Also update Introduction text if it references [17] by old author "Ren J"
# The Introduction text says "Ren et al. (2024)" which is fine since we corrected the reference list

doc.save(DOCX)
print(f'[OK] Saved to {DOCX}')
print(f'Total paragraphs: {len(doc.paragraphs)}')

# Summary of fixes
print('\n=== Fixes applied ===')
print('[17] Authors: Ren J,Wang Y,Zhang H -> Ren M,Yang B,Yang D,Liu Y,Zhang H,Zhang M,Zhang Y')
print('[22] DOI: 1c02759 -> 1c02748')
print('[26] Vol/pages: 152:21-45 -> 152(9):66; DOI updated to 10.1007/s11242-025-02196-y')
print('[27] SPE paper: 219475-PA -> 219730-PA')
print('[30] Authors: added Chang FF, Li L (was missing 2 of 5 authors)')
print('[31] DOI: .064 -> .030')
