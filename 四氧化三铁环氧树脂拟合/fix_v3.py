# -*- coding: utf-8 -*-
"""Fix remaining issues in v3 manuscript."""
from docx import Document
from docx.shared import Pt

DOCX = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP-T_投稿文件\ESP-T_v3_manuscript.docx'
doc = Document(DOCX)

# 1. Fix 4.2 heading
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith('Heading') and '4.2' in p.text and 'Physical Self-Calibration' in p.text:
        p.clear()
        run = p.add_run('4.2 BTC Decomposition and Model Validation')
        print(f'[OK] Fixed 4.2 heading at [{i}]')
        break

# 2. Fix 4.3: Remove Q=0.46 self-calibration claim, reframe
for i, p in enumerate(doc.paragraphs):
    if 'Q = 0.46 mL/min' in p.text and '8%' in p.text:
        old_text = p.text
        new_text = old_text.replace(
            '(Q = 0.46 mL/min, −8%)',
            '(R² = 0.9939, signal decomposition into 53%/47% components)'
        )
        # Also fix any mention of Q self-calibration in TOA context
        new_text = new_text.replace(
            'achieves the best accuracy',
            'provides the richest physical interpretation'
        )
        p.clear()
        p.add_run(new_text)
        print(f'[OK] Fixed 4.3 TOA comparison at [{i}]')
        break

# 3. Fix Introduction P4: Remove Q self-calibration preview
for i, p in enumerate(doc.paragraphs):
    if 'predictive' in p.text and 'validation strategy' in p.text:
        old_text = p.text
        new_text = old_text.replace(
            'The validation strategy is predictive, not descriptive: Q is left unconstrained '
            'in the objective function; the model must converge to the correct flow rate from '
            'the BTC shape alone.',
            'The model is validated through internal self-consistency of transport time scales '
            '(fitted MRT versus independently computed convective travel time) and independent '
            'corroboration across separate experiments (Pe from the dynamic BTC fit versus n '
            'from static K-P batch kinetics).'
        )
        p.clear()
        p.add_run(new_text)
        print(f'[OK] Fixed Introduction P4 at [{i}]')
        break

# 4. Fix 4.1: Remove "Section 4.2 provides the positive test" if it still references Q
for i, p in enumerate(doc.paragraphs):
    if 'The decisive test is therefore the self-calibration of Section 4.2' in p.text:
        old_text = p.text
        new_text = (
            'The decisive test of physical validity is therefore provided in Section 4.2, '
            'where the fitted MRT is compared with the independently computed convective '
            'travel time, and the Peclet number is independently corroborated by the K-P '
            'kinetic analysis.'
        )
        p.clear()
        p.add_run(new_text)
        print(f'[OK] Fixed 4.1 self-calibration ref at [{i}]')
        break

# 5. Fix 4.1: Also fix the "order of magnitude" -> "several-fold" (Task 4 deferred)
for i, p in enumerate(doc.paragraphs):
    if 'order of magnitude below' in p.text:
        old_text = p.text
        new_text = old_text.replace('order of magnitude below', 'several-fold below')
        p.clear()
        p.add_run(new_text)
        print(f'[OK] Fixed order-of-magnitude at [{i}]')

# 6. Fix 4.3 MRT reconciliation (Task 6 deferred)
# "first-moment... provides no signal decomposition, no mechanistic insight... no estimate of Peclet"
# This should already be there, checking
for i, p in enumerate(doc.paragraphs):
    if 'MRT 37.1' in p.text or ('37.1' in p.text and '37.4' in p.text):
        # This is the reconciliation line added by Task 12
        print(f'[OK] MRT reconciliation at [{i}]: {p.text[:100]}')

# 7. Save
output = DOCX  # Overwrite
doc.save(output)
print(f'\n[OK] Fixes applied and saved to {output}')
