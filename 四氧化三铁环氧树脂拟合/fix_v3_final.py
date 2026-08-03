# -*- coding: utf-8 -*-
"""Fix remaining self-calibration references in v3."""
from docx import Document
from docx.shared import Pt

DOCX = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP-T_投稿文件\ESP-T_v3_manuscript.docx'
doc = Document(DOCX)

fixes_done = 0

for i, p in enumerate(doc.paragraphs):
    text = p.text

    # Fix 1: Keywords - replace "Self-calibration" with "Tracer flux"
    if 'Self-calibration;' in text and 'Breakthrough curve' in text:
        new_text = text.replace('Self-calibration;', 'Tracer flux;')
        p.clear(); p.add_run(new_text)
        fixes_done += 1
        print(f'[{i}] Keywords: Self-calibration -> Tracer flux')

    # Fix 2: Section 2.4 heading
    if p.style.name.startswith('Heading') and 'Self-Calibration as Decisive Test' in text:
        p.clear(); run = p.add_run('2.4 Validation Strategy: Internal Consistency and Independent Corroboration')
        fixes_done += 1
        print(f'[{i}] 2.4 heading renamed')

    # Fix 3: Section 2 lead-in mentioning self-calibration
    if 'validation strategy—physical self-calibration' in text.lower() or \
       ('model delivers a self-calibration' in text.lower()):
        new_text = text.replace(
            'validation strategy—physical self-calibration—',
            'validation strategy—internal consistency and independent corroboration—'
        ).replace(
            'model delivers a self-calibration',
            'model provides internal consistency'
        )
        p.clear(); p.add_run(new_text)
        fixes_done += 1
        print(f'[{i}] Section 2 lead-in fixed')

    # Fix 4: Section 3.4 mentioning self-calibration test
    if 'self-calibration test of Section 4.2' in text:
        new_text = text.replace(
            'self-calibration test of Section 4.2',
            'model validation analysis of Section 4.2'
        )
        p.clear(); p.add_run(new_text)
        fixes_done += 1
        print(f'[{i}] 3.4 self-calibration ref fixed')

    # Fix 5: Section 4.4/4.5 referencing Section 4.2 self-calibration
    if 'Section 4.2 self-calibration' in text or 'Section 4.2 result' in text:
        new_text = text.replace('Section 4.2 self-calibration result', 'Section 4.2 validation')
        new_text = new_text.replace('Section 4.2 self-calibration', 'Section 4.2 validation')
        new_text = new_text.replace('Section 4.2 result', 'Section 4.2 result')
        if new_text != text:
            p.clear(); p.add_run(new_text)
            fixes_done += 1
            print(f'[{i}] 4.x self-calibration ref fixed')

    # Fix 6: Any remaining "physical self-calibration" in prose
    if 'physical self-calibration' in text.lower() and not p.style.name.startswith('Heading'):
        new_text = text.replace('physical self-calibration', 'model validation')
        new_text = new_text.replace('Physical self-calibration', 'Model validation')
        if new_text != text:
            p.clear(); p.add_run(new_text)
            fixes_done += 1
            print(f'[{i}] physical self-calibration prose fixed')

    # Fix 7: Figure 5 caption
    if 'Physical self-calibration' in text and 'Fig.' in text and '5' in text:
        new_text = text.replace('Physical self-calibration', 'BTC decomposition and model validation')
        p.clear(); p.add_run(new_text)
        fixes_done += 1
        print(f'[{i}] Fig.5 caption fixed')

# Fix 8: "Q converges" in Introduction - should be about MRT
for i, p in enumerate(doc.paragraphs):
    if 'Q converges to the correct value' in p.text or 'that Q converges' in p.text:
        new_text = p.text.replace(
            'that Q converges to the correct value',
            'that the fitted MRT agrees with the independently computed convective time scale'
        ).replace(
            'Q converges',
            'the fitted MRT agrees with the independently computed convective time scale'
        )
        p.clear(); p.add_run(new_text)
        fixes_done += 1
        print(f'[{i}] Q converges ref fixed')

doc.save(DOCX)
print(f'\n[OK] {fixes_done} fixes applied')
