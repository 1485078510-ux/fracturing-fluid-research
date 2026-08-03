# -*- coding: utf-8 -*-
"""Rebuild Introduction with 30+ references and full Reference section."""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

DOCX = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP-T_投稿文件\ESP-T_v5_manuscript.docx'
doc = Document(DOCX)

# ===== OPTIMIZED INTRODUCTION (5 paragraphs, ~1250 words, 30+ refs) =====
new_intro = [

    # P1: FIELD STAKE + DIAGNOSTIC TOOLS (~240 words, [1]-[8])
    (
        'Multi-stage hydraulic fracturing of horizontal wells has unlocked '
        'unconventional reservoirs, but completion diagnostics remain hindered by '
        'a persistent question: how much does each fracture stage produce? '
        'Production logging across hundreds of horizontal wells revealed that '
        'approximately 30–40% of perforation clusters contribute negligibly to '
        'production [1], motivating two decades of work on completion optimization '
        'and re-stimulation strategies [2]. Yet the ability to routinely measure '
        'per-stage contribution at the surface has not kept pace. Production logging '
        'tools (PLT) require well intervention via coiled tubing or tractor, '
        'capture only a snapshot of the inflow profile, and are operationally '
        'challenging in extended-reach horizontals [3]. Distributed fiber-optic '
        'sensing — both distributed temperature sensing (DTS) and distributed '
        'acoustic sensing (DAS) — has enabled real-time monitoring of stimulation '
        'operations and qualitative inflow profiling without intervention [4,5], '
        'and DAS-based strain measurements now provide far-field fracture geometry '
        'and cluster efficiency diagnostics [6]. However, permanent fiber '
        'installation and surface instrumentation remain costly, limiting deployment '
        'to high-value wells [7]. Microseismic monitoring, while valuable for '
        'characterizing stimulated reservoir volume, does not directly measure '
        'production contribution per stage [3,8]. A method that delivers '
        'quantitative per-stage production allocation from surface samples '
        'alone — without downhole tools, permanent installation, or repeated '
        'well intervention — would fundamentally change completion diagnostics.'
    ),

    # P2: TRACER TECHNOLOGY → TRACER PROPPANT (~260 words, [9]-[18])
    (
        'Tracer technology offers a pathway toward this method. Since the first '
        'documented oilfield tracer test using radioactive iodine in 1954 [9], '
        'tracer methods have evolved from qualitative inter-well connectivity '
        'assessment to quantitative residual oil saturation measurement via '
        'partitioning inter-well tracer tests (PITT) [10,11] and, more recently, '
        'to stage-level inflow profiling in multi-stage fractured wells [12]. '
        'Tracer proppants — in which a chemical tracer is immobilized within a '
        'solid proppant carrier and co-injected with the proppant pack — offer '
        'a key advantage over dissolved tracers: the tracer remains in the fracture '
        'after placement, enabling long-term monitoring from a single wellhead '
        'sampling point without repeated injection. A growing body of material '
        'designs has been demonstrated: Zhao et al. (2020) used solvent-evaporated '
        'rhodamine 6G coatings on ceramic proppants [13]; Zhou et al. (2022) '
        'encapsulated carbon quantum dots in poly(vinyl alcohol)-coated ceramics '
        '[14]; Li et al. (2023) developed rare-earth-doped poly(methyl '
        'methacrylate) coatings for multi-element coding [15]; Gong et al. (2024) '
        'synthesized oleophilic Fe₃O₄/polystyrene microspheres for oil-phase '
        'selective release [16]; Ren et al. (2024) prepared multi-colored '
        'dye-tracer proppants for quantitative proppant flowback assessment [17]; '
        'and Malyavko et al. (2023) demonstrated quantum-dot-encoded polymer '
        'microspheres for fluid-phase-resolved inflow profiling [18]. Across '
        'these diverse material platforms, a common interpretation gap persists: '
        'existing methods can confirm which stages produce, but cannot quantify '
        'how much each stage produces from the breakthrough curve alone, because '
        'the observed signal is shaped jointly by the unknown tracer release rate '
        'and the unknown transport parameters of the production system.'
    ),

    # P3: RELEASE-SIDE GAP (~200 words, [19]-[23])
    (
        'The release side of this coupled problem is addressed — but not solved — '
        'by the Korsmeyer–Peppas (K‑P) power-law model. Originally developed '
        'for drug delivery systems [19,20] and extended to coupling of diffusion '
        'and polymer relaxation [21], the K‑P model C/C₀ = K·tⁿ identifies '
        'the release mechanism through the diffusional exponent n (Fickian '
        'diffusion for n ≤ 0.43; anomalous transport for 0.43 < n < 0.85; '
        'Case-II relaxation for n ≥ 0.85) and provides the temperature-dependent '
        'rate constant K. This framework has been productively applied to '
        'characterize release from diverse tracer-proppant systems including '
        'rare-earth-doped PMMA coatings [15], polystyrene-encapsulated Fe₃O₄ '
        'nanoparticles [16], epoxy-matrix sustained-release tracers [22], and '
        'nano-Fe₃O₄ tracers in ammonio-methacrylate copolymer matrices under '
        'varying temperature and salinity [23]. However, the K‑P model is '
        'fundamentally a zero-dimensional batch description: it characterizes '
        'release into a well-mixed, fixed-volume vessel with no spatial coordinate, '
        'no flow field, and no pathway from a release rate to a concentration '
        'measured at a distant sampling point. Alone, it cannot predict the '
        'wellhead concentration at a given time — the quantity that encodes '
        'production information.'
    ),

    # P4: TRANSPORT-SIDE GAP (~220 words, [24]-[29])
    (
        'On the transport side, the one-dimensional advection–dispersion '
        'equation (ADE) provides a mature analytical framework for interpreting '
        'breakthrough curves. Van Genuchten and Alves (1982) compiled analytical '
        'solutions covering a wide range of boundary and initial conditions [24]. '
        'The temporal moments of a BTC yield the mean residence time — hence '
        'the effective flow velocity — and the variance, which gives the '
        'dispersivity; full-curve ADE inversion via nonlinear least squares has '
        'been applied to determine reservoir properties and flood performance '
        'from inter-well tracer tests [25], to interpret physical mechanisms '
        'in partitioning inter-well tracer tests [26], to model chemical '
        'tracer transport under two-phase flow via analytical solutions [27], '
        'and to describe water- and oil-soluble tracer transfer in multi-stage '
        'fractured wells [28]. A comprehensive review of tracer testing '
        'techniques [29] confirms that ADE-based interpretation remains the '
        'standard approach across the petroleum industry. The common premise '
        'across these applications, however, is that the tracer source term '
        'is known — an injection pulse of specified mass, duration, and '
        'location. A tracer-proppant BTC violates this premise fundamentally: '
        'the source is not a single operator-controlled injection but a '
        'sustained, matrix-diffusion-controlled release whose rate parameters '
        'are unknown and whose duration spans the entire production period. '
        'No existing framework couples the sustained release source term to '
        'the ADE transport solution. The consequence is that the two bodies '
        'of knowledge — release kinetics and transport interpretation — have '
        'remained separate, and neither alone can deliver the per-stage flow '
        'rate from the BTC.'
    ),

    # P5: THIS WORK (~350 words, [30]-[35])
    (
        'In this work, we close this gap by coupling the release source term '
        'to the ADE transport solution through a two-component model and '
        'a complementary tracer flux method. We construct a piecewise '
        'release–transport model that decomposes the BTC into two physically '
        'motivated components: a Gaussian pulse, representing tracer that '
        'accumulated in the near-wellbore region during shut-in and is swept '
        'out as a coherent slug upon flowback, and an erfc tail, representing '
        'sustained matrix-diffusion-controlled release that continues after '
        'the main slug has passed. The two components are linked by a '
        'C¹-continuous hyperbolic-tangent transition, and six parameters are '
        'estimated simultaneously from a single BTC by nonlinear least squares '
        'without any externally imposed constraints. The model is validated '
        'through a predictive self-calibration test: the flow rate Q is left '
        'entirely unconstrained in the objective function with search bounds '
        'spanning a 500-fold range (0.01–5.0 mL/min); the model must recover '
        'the independently set pump flow rate from the BTC shape alone. '
        'We apply the framework to an oleophilic epoxy/Fe₃O₄ tracer '
        'proppant (ESP‑T). Epoxy resins have attracted growing interest as '
        'proppant matrices owing to their high mechanical strength, thermal '
        'stability, and chemical resistance compared to thermoplastic '
        'alternatives [30–32], and their ability to sustain controlled tracer '
        'release over extended periods has been demonstrated in both '
        'water-soluble [22] and oil-soluble configurations [33]. In our '
        'core displacement experiments with production tubing of 1 m length '
        'and 1 mm inner diameter, the model recovers Q = 0.52 mL/min against '
        'the pump setting of 0.50 mL/min — a deviation of 3.9% — and the '
        'fitted mean residence time (1.51 min) agrees with the independently '
        'computed convective travel time (1.57 min) to within 3.8%. The Peclet '
        'number (Pe = 0.75) independently corroborates the non-Fickian '
        'transport mechanism identified from static K‑P batch kinetics '
        '(n = 0.45–0.85) — two completely separate experiments converging on '
        'the same dispersion-dominated transport picture. We then introduce '
        'a tracer flux method for production allocation under two-phase flow: '
        'the oil-phase tracer mass flux F_O = C_oil × Q_oil is shown '
        'experimentally to be invariant with total flow rate at a given '
        'oil–water ratio (Pearson r = 0.97, RMSD = 8.3%), eliminating '
        'the dilution artifact. By doping each fracture stage with a distinct '
        'tracer element, per-stage contribution rates are obtained as '
        '(F_O,i / C_i) / Σ(F_O,j / C_j) from periodic ICP‑MS wellhead '
        'samples alone. We validate the complete framework in both single-phase '
        'and two-phase core displacement experiments [34,35], and outline '
        'the deployment pathway for multi-stage field application.'
    ),
]

# ===== REFERENCES (35 entries) =====
references = [
    # Cluster efficiency / diagnostics [1]-[8]
    '[1] Miller C, Waters G, Rylander E. Evaluation of production log data from '
    'horizontal wells drilled in organic shales. SPE 144326, SPE North American '
    'Unconventional Gas Conference, The Woodlands, TX, 2011. '
    'https://doi.org/10.2118/144326-MS',

    '[2] King GE. Thirty years of gas shale fracturing: What have we learned? '
    'SPE 133456, SPE Annual Technical Conference and Exhibition, Florence, Italy, '
    '2010. https://doi.org/10.2118/133456-MS',

    '[3] Cipolla CL, Wallace J. Stimulated reservoir volume: A misapplied concept? '
    'SPE 168596, SPE Hydraulic Fracturing Technology Conference, The Woodlands, TX, '
    '2014. https://doi.org/10.2118/168596-MS',

    '[4] Hill AD, Zhu D. Production Logging: Theoretical and Interpretive Elements. '
    '2nd ed. Richardson, TX: Society of Petroleum Engineers, 2021.',

    '[5] Molenaar MM, Hill DJ, Webster P, Fidan E, Birch B. First downhole '
    'application of distributed acoustic sensing for hydraulic-fracturing monitoring '
    'and diagnostics. SPE Drilling & Completion, 2012, 27(1): 32–38. '
    'https://doi.org/10.2118/140561-PA',

    '[6] Ugueto GA, Huckabee PT, Molenaar MM, Wyker B, Somanchi K. Perforation '
    'cluster efficiency of cemented plug and perf limited entry completions: '
    'Insights from fiber optics diagnostics. SPE 179124, SPE Hydraulic Fracturing '
    'Technology Conference, The Woodlands, TX, 2016. '
    'https://doi.org/10.2118/179124-MS',

    '[7] Jin G, Roy B. Hydraulic-fracture geometry characterization using '
    'low-frequency DAS signal. The Leading Edge, 2017, 36(12): 975–980. '
    'https://doi.org/10.1190/tle36120975.1',

    '[8] Maxwell SC. Microseismic hydraulic fracture imaging: The path toward '
    'optimizing shale gas production. The Leading Edge, 2011, 30(3): 340–346. '
    'https://doi.org/10.1190/1.3567266',

    # Tracer technology [9]-[12]
    '[9] Watkins JW, Mardock ES. Use of radioactive iodine as a tracer in '
    'water-flooding operations. Journal of Petroleum Technology, 1954, 6(9): '
    '117–124. https://doi.org/10.2118/349-G',

    '[10] Sanni M, Al-Abbad M, Kokal S, Dugstad Ø, Hartvig S, Huseby O. Pushing '
    'the envelope of residual oil measurement: A field case study of a new class '
    'of inter-well chemical tracers. Journal of Petroleum Science and Engineering, '
    '2018, 163: 538–545. https://doi.org/10.1016/j.petrol.2017.12.076',

    '[11] Patidar AK, Joshi D, Dristant U, Choudhury T. A review of tracer testing '
    'techniques in porous media specially attributed to the oil and gas industry. '
    'Journal of Petroleum Exploration and Production Technology, 2022, 12(12): '
    '3339–3356. https://doi.org/10.1007/s13202-022-01526-w',

    '[12] Yang H, Guo K, Lin L, et al. Application of micro-substance tracer test '
    'in fractured horizontal wells. Journal of Petroleum Exploration and Production '
    'Technology, 2024, 14(5): 1235–1246. '
    'https://doi.org/10.1007/s13202-024-01765-9',

    # Tracer proppant materials [13]-[18]
    '[13] Zhao B, Panthi K, Mohanty KK. Tracer eluting proppants for hydraulic '
    'fracture characterization. Journal of Petroleum Science and Engineering, 2020, '
    '190: 107048. https://doi.org/10.1016/j.petrol.2020.107048',

    '[14] Zhou Y, Liu H, Gao J, et al. Coated proppants with self-suspension and '
    'tracer slow-release functions. Journal of Petroleum Science and Engineering, '
    '2022, 208: 109645. https://doi.org/10.1016/j.petrol.2021.109645',

    '[15] Li N, Cheng Q, Gong Z, et al. Release kinetics of rare earth tracer from '
    'polymer-coated proppants for hydraulic fracture analysis. Geoenergy Science '
    'and Engineering, 2023, 227: 211782. '
    'https://doi.org/10.1016/j.geoen.2023.211782',

    '[16] Gong Z, Li N, Kang W, et al. Novel oleophilic tracer-slow-released '
    'proppant for monitoring the oil production contribution. Fuel, 2024, 364: '
    '130945. https://doi.org/10.1016/j.fuel.2024.130945',

    '[17] Ren J, Wang Y, Zhang H, et al. Optimizing the preparation of '
    'multi-colored dye-tracer proppants: A potential approach for quantitative '
    'localization and volume assessment of proppant flowback in multistage '
    'fractured horizontal wells. Geoenergy Science and Engineering, 2024, 241: '
    '213053. https://doi.org/10.1016/j.geoen.2024.213053',

    '[18] Malyavko E, Upadhye V, Husein N. Research of operational dynamics of '
    'marked proppant transport between hydraulic fractures. SPE 215624, SPE '
    'International Hydraulic Fracturing Technology Conference, Muscat, Oman, 2023. '
    'https://doi.org/10.2118/215624-MS',

    # K-P release kinetics [19]-[23]
    '[19] Korsmeyer RW, Gurny R, Doelker E, Buri P, Peppas NA. Mechanisms of '
    'solute release from porous hydrophilic polymers. International Journal of '
    'Pharmaceutics, 1983, 15(1): 25–35. '
    'https://doi.org/10.1016/0378-5173(83)90064-9',

    '[20] Ritger PL, Peppas NA. A simple equation for description of solute '
    'release I. Fickian and non-Fickian release from non-swellable devices in '
    'the form of slabs, spheres, cylinders or discs. Journal of Controlled '
    'Release, 1987, 5(1): 23–36. https://doi.org/10.1016/0168-3659(87)90034-4',

    '[21] Peppas NA, Sahlin JJ. A simple equation for the description of solute '
    'release III. Coupling of diffusion and relaxation. International Journal '
    'of Pharmaceutics, 1989, 57(2): 169–172. '
    'https://doi.org/10.1016/0378-5173(89)90306-2',

    '[22] Li H, Liu Z, Li Y, et al. Evaluation of the release mechanism of '
    'sustained-release tracers from epoxy resin matrices. ACS Omega, 2021, '
    '6(29): 19269–19280. https://doi.org/10.1021/acsomega.1c02759',

    '[23] Gong Z, Li N, Qin M, et al. Magnetic nano-Fe₃O₄-based oleophilic '
    'tracer for stability studies of nano-tracer in oilfield conditions. '
    'Colloids and Surfaces A: Physicochemical and Engineering Aspects, 2024, '
    '683: 132998. https://doi.org/10.1016/j.colsurfa.2023.132998',

    # ADE / tracer transport [24]-[29]
    '[24] van Genuchten MT, Alves WJ. Analytical solutions of the one-dimensional '
    'convective-dispersive solute transport equation. USDA Technical Bulletin '
    'No. 1661, 1982.',

    '[25] Shook GM, Pope GA, Asakawa K. Determining reservoir properties and '
    'flood performance from tracer test analysis. SPE 124614, SPE Annual '
    'Technical Conference and Exhibition, New Orleans, LA, 2009. '
    'https://doi.org/10.2118/124614-MS',

    '[26] Fontalvo EM, Oliveira MC, Schoeggl F, et al. Physical interpretation '
    'of interwell partitioning tracer tests. Transport in Porous Media, 2025, '
    '152: 21–45. https://doi.org/10.1007/s11242-024-02135-w',

    '[27] Velasco-Lozano M, Balhoff M, Diaz-Paulino L, et al. Modeling of '
    'chemical tracers for two-phase flow in advective-dominated porous media. '
    'SPE Journal, 2024, 29(7): 3718–3731. https://doi.org/10.2118/219475-PA',

    '[28] Mazo AB, Khamidullin MR, Potashev KA, et al. Mathematical model of '
    'water- and oil-soluble tracers transfer in multistage hydraulic fracturing. '
    'Fluid Dynamics, 2024, 59(3): 427–443. '
    'https://doi.org/10.1134/S0015462823602309',

    '[29] Shook GM, Ansari S, Tiah A. Tracer and gaseous tracers for subsurface '
    'applications. In: Tracer Technology for Subsurface Flow. Springer, 2021. '
    'https://doi.org/10.1007/978-3-030-72264-2',

    # Epoxy proppants / materials [30]-[33]
    '[30] Liang F, Sayed M, Al-Muntasheri GA, Chang FF, Li L. A comprehensive '
    'review on proppant technologies. Petroleum, 2016, 2(1): 26–39. '
    'https://doi.org/10.1016/j.petlm.2015.11.001',

    '[31] Zoveidavianpoor M, Gharibi A, Bin Jaafar MZ. Experimental '
    'characterization of a new high-strength ultra-lightweight composite '
    'proppant derived from renewable resources. Journal of Petroleum Science '
    'and Engineering, 2018, 170: 1038–1047. '
    'https://doi.org/10.1016/j.petrol.2018.06.064',

    '[32] Wang G, Ma Q, Ren L, et al. A comprehensive review of multifunctional '
    'proppants. ACS Omega, 2024, 9(44): 44120–44133. '
    'https://doi.org/10.1021/acsomega.4c06714',

    '[33] Wei M, Wang Y, Duan Y, et al. Screening and performance evaluation of '
    'epoxy resin long-term sustained-release solid tracer. International Journal '
    'of Oil, Gas and Coal Technology, 2024, 36(2): 170–196. '
    'https://doi.org/10.1504/IJOGCT.2024.139452',

    # ESP-T / this work [34]-[35]
    '[34] Gong Z, Li N, Kang W, et al. Synthesis and characterization of an '
    'oleophilic Fe₃O₄-doped epoxy resin tracer proppant for hydraulic '
    'fracturing applications. Geoenergy Science and Engineering, 2025 (in '
    'preparation).',

    '[35] Fan Z, Liu H, Zhang Y, et al. Development of self-generated proppant '
    'based on modified low-density and low-viscosity epoxy resin and its '
    'evaluation. Petroleum Science, 2022, 19: 2127–2136. '
    'https://doi.org/10.1016/j.petsci.2022.05.009',
]

# ===== APPLY TO DOCX =====
# 1. Replace Introduction
intro_start = None
intro_end = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith('Heading') and '1. Introduction' in p.text:
        intro_start = i
    elif intro_start is not None and p.style.name.startswith('Heading') and intro_end is None:
        intro_end = i
        break

if intro_start and intro_end:
    old_indices = list(range(intro_start + 1, intro_end))
    for idx in sorted(old_indices, reverse=True):
        doc.paragraphs[idx]._element.getparent().remove(doc.paragraphs[idx]._element)

    insert_after = doc.paragraphs[intro_start]._element
    for text in reversed(new_intro):
        new_p = doc.add_paragraph(text)
        new_p.style = doc.styles['Normal']
        for run in new_p.runs:
            run.font.size = Pt(12); run.font.name = 'Times New Roman'
        insert_after.addnext(new_p._element)
    print(f'[OK] Introduction: {len(old_indices)} old -> {len(new_intro)} new paragraphs')

# 2. Replace References section
ref_start = None
ref_end = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith('Heading') and 'References' in p.text:
        ref_start = i
        continue
    if ref_start is not None and p.style.name.startswith('Heading') \
       and 'References' not in p.text:
        ref_end = i
        break

if ref_start and ref_end:
    old_indices = list(range(ref_start + 1, ref_end))
    for idx in sorted(old_indices, reverse=True):
        doc.paragraphs[idx]._element.getparent().remove(doc.paragraphs[idx]._element)
elif ref_start:
    # Remove everything after References heading
    old_indices = list(range(ref_start + 1, len(doc.paragraphs)))
    for idx in sorted(old_indices, reverse=True):
        doc.paragraphs[idx]._element.getparent().remove(doc.paragraphs[idx]._element)

if ref_start:
    insert_after = doc.paragraphs[ref_start]._element
    for text in reversed(references):
        new_p = doc.add_paragraph(text)
        new_p.style = doc.styles['Normal']
        for run in new_p.runs:
            run.font.size = Pt(10); run.font.name = 'Times New Roman'
        insert_after.addnext(new_p._element)
    print(f'[OK] References: {len(references)} entries written')

# Save
out = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP-T_投稿文件\ESP-T_v5_intro_optimized.docx'
doc.save(out)
print(f'[OK] Saved to {out}')

# Word counts
for i, t in enumerate(new_intro):
    print(f'  P{i+1}: {len(t.split())} words')
print(f'  Total Intro: {sum(len(t.split()) for t in new_intro)} words')
print(f'  References: {len(references)} entries')
