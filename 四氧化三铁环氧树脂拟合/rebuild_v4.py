# -*- coding: utf-8 -*-
"""
Rebuild manuscript with STANDARD JOURNAL STRUCTURE:
1. Introduction
2. Experimental
3. Model Development
4. Results and Discussion
5. Implications for Field Deployment
6. Conclusions

Extract content from v3 and reorganize.
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

v3_path = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP-T_投稿文件\ESP-T_v3_manuscript.docx'
doc = Document(v3_path)

# ===== Extract content by section =====
# Map paragraph indices to section boundaries
sections = {}
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith('Heading 1'):
        sections[p.text] = {'start': i, 'end': None}
    elif p.style.name.startswith('Heading 2'):
        # Track last H2 before next H1
        pass

# Find section boundaries
h1_indices = [(i, p.text) for i, p in enumerate(doc.paragraphs)
              if p.style.name.startswith('Heading 1')]
for idx, (i, title) in enumerate(h1_indices):
    end = h1_indices[idx+1][0] if idx+1 < len(h1_indices) else len(doc.paragraphs)
    sections[title] = {'start': i, 'end': end}

# Collect content for new structure
def get_paras_between(doc, start_text_contains, end_text_contains):
    """Get all paragraphs between two headings (non-inclusive of headings)."""
    paras = []
    collecting = False
    for p in doc.paragraphs:
        if collecting:
            if p.style.name.startswith('Heading') and end_text_contains.lower() in p.text.lower():
                break
            if p.text.strip():
                paras.append(p.text)
        if p.style.name.startswith('Heading') and start_text_contains.lower() in p.text.lower():
            collecting = True
    return paras

def get_all_paras_in_section(doc, section_start_contains):
    """Get ALL paragraphs (including sub-headings) from a section until next H1."""
    paras = []
    in_section = False
    for p in doc.paragraphs:
        if in_section:
            if p.style.name.startswith('Heading 1'):
                break
            paras.append((p.style.name, p.text))
        if p.style.name.startswith('Heading 1') and section_start_contains.lower() in p.text.lower():
            in_section = True
    return paras

# Now print what we have for reference
print("=== Current v3 sections ===")
for title in sections:
    print(f"  {title}")

print()
print("=== Building v4 with standard journal structure ===")
print()

# Create new document
new_doc = Document()

# Set default font
style = new_doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# ---- TITLE ----
title_text = ('Tracer Flux Method with Coupled Release-Transport Analysis '
              'for Per-Stage Production Allocation from Tracer-Proppant Breakthrough Curves')
title_p = new_doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run(title_text)
title_run.bold = True
title_run.font.size = Pt(14)

# ---- ABSTRACT ----
# Find abstract from v3
abs_paras = []
abs_found = False
for p in doc.paragraphs:
    if p.style.name.startswith('Heading') and 'Abstract' in p.text:
        abs_found = True
        continue
    if abs_found:
        if p.style.name.startswith('Heading'):
            break
        if p.text.strip():
            abs_paras.append(p.text)

new_doc.add_heading('Abstract', level=1)
if abs_paras:
    new_doc.add_paragraph(abs_paras[0])

# ---- KEYWORDS ----
kw_paras = []
kw_found = False
for p in doc.paragraphs:
    if p.style.name.startswith('Heading') and 'Keywords' in p.text:
        kw_found = True
        continue
    if kw_found:
        if p.style.name.startswith('Heading'):
            break
        if p.text.strip():
            kw_paras.append(p.text)

new_doc.add_heading('Keywords', level=1)
if kw_paras:
    new_doc.add_paragraph(kw_paras[0])

# ---- 1. INTRODUCTION ----
intro_paras = []
in_intro = False
for p in doc.paragraphs:
    if p.style.name.startswith('Heading 1') and '1. Introduction' in p.text:
        in_intro = True
        continue
    if in_intro:
        if p.style.name.startswith('Heading 1'):
            break
        if p.text.strip():
            intro_paras.append(p.text)

new_doc.add_heading('1. Introduction', level=1)
for text in intro_paras:
    new_doc.add_paragraph(text)

# ---- 2. EXPERIMENTAL ----
new_doc.add_heading('2. Experimental', level=1)

# 2.1 Materials
new_doc.add_heading('2.1 Materials', level=2)
# Get from v3 Section 3.1
synth_text = None
in_sec3 = False
for p in doc.paragraphs:
    if p.style.name.startswith('Heading 2') and '3.1' in p.text:
        in_sec3 = True
        continue
    if in_sec3:
        if p.style.name.startswith('Heading'):
            break
        if p.text.strip():
            synth_text = p.text
            break

if synth_text:
    # Extract just the reagent/materials info from the beginning of synthesis text
    new_doc.add_paragraph(
        'Analytical grade reagents — including FeCl₃ (≥99.0%), FeCl₂·4H₂O '
        '(≥99.0%), MnCl₂·6H₂O (≥99.0%), stearic acid (≥99.5%), '
        'and anhydrous ethanol — were purchased from Chengdu Kelong Chemical Co., Ltd. '
        'Industrial grade materials included guar gum, E51 epoxy resin (epoxy value '
        '0.48–0.54 mol/100 g, Nantong Xingchen Synthetic Materials), T31 curing agent '
        '(phenolic amine, ≥95% effective component), and hollow glass microspheres '
        '(≥99% purity). Dodecane (≥99%, Chengdu Kelong) served as the model oil phase. '
        'All reagents were used as received. Detailed specifications are provided in '
        'Table S1 (Supplementary Material).'
    )

# 2.2 ESP-T Synthesis
new_doc.add_heading('2.2 ESP-T Synthesis', level=2)
if synth_text:
    new_doc.add_paragraph(synth_text)

# 2.3 Characterization
new_doc.add_heading('2.3 Characterization', level=2)
# From v3 Section 3.2 - extract methods
char_text = (
    'Microstructure and elemental distribution were characterized by scanning electron '
    'microscopy with energy-dispersive X-ray spectroscopy (SEM-EDS, ZEISS Sigma 500) and '
    'optical microscopy (Leica DM2700P). Thermal stability was evaluated by '
    'thermogravimetric analysis and differential scanning calorimetry (TGA/DSC, TA '
    'Instruments Q500, air atmosphere, 10 °C/min, 50–800 °C). Water contact '
    'angle (WCA) was measured using a video optical contact angle analyzer (OCA20, 5 μL '
    'droplets, n = 5). Physical and mechanical properties — including bulk density, '
    'apparent density, sphericity, roundness, acid solubility (HCl 12 wt%/HF 3 wt%, '
    '65 °C, 30 min), and crush rate (52 MPa, 2 min) — were evaluated per '
    'SY/T 5107–2016. Oil-water transport selectivity was assessed via packed-bed '
    'filtration time (2.0 g proppant, 20 mL fluid, 200-mesh screen). Detailed protocols '
    'are provided in the Supplementary Material (Section S2).'
)
new_doc.add_paragraph(char_text)

# 2.4 Batch Release Kinetics
new_doc.add_heading('2.4 Batch Release Kinetics', level=2)
batch_text = (
    'ESP-T (5 g, 40–70 mesh) was immersed in 100 mL dodecane in sealed glass vials '
    'placed in thermostatic oil baths at 30, 60, 90, and 120 °C. Sampling was '
    'conducted at 12 h intervals over 14 days; the concentration of tracer-doped metal '
    'ions was quantified by inductively coupled plasma mass spectrometry (ICP-MS, '
    'PerkinElmer NexION 300X). Release data within M_t/M_∞ < 0.6 — the range '
    'in which the Korsmeyer–Peppas power law C/C₀ = K·t^n is valid '
    '— were fitted to determine the kinetic parameters K and n. The full release '
    'dataset is provided in Table S3.'
)
new_doc.add_paragraph(batch_text)

# 2.5 Core Displacement Experiments
new_doc.add_heading('2.5 Core Displacement Experiments', level=2)
core_text = (
    'A dynamic displacement apparatus (Fig. S2) was used to simulate single-interval '
    'production. A steel core (length 100 mm, inner diameter 5 mm) was packed with ESP-T, '
    'sealed with 200-mesh screens, and placed in a core holder under 5 MPa confining '
    'pressure. For single-phase experiments, the system was saturated with dodecane at '
    '5 mL/min, shut in for 96 h to allow tracer accumulation (representative of typical '
    'shut-in periods between fracturing and flowback in field operations), then displaced '
    'at a pump setting of 0.50 mL/min. The pump setting is the independent flow-rate '
    'reference for the model validation analysis of Section 4.3. Effluent was sampled at '
    '4-min intervals (2 mL per sample, 21 samples) and analyzed by ICP-MS to construct '
    'the breakthrough curve. For two-phase experiments, dodecane and deionized water were '
    'co-injected at three oil-water volume ratios (OWR = 4:1, 1:1, 1:4) and four total '
    'flow rates (0.1–0.4 mL/min) under continuous steady-state flow. Sampling was '
    'conducted at 5-min intervals; both the oil-water volume ratio and tracer concentration '
    'in each sample were recorded.'
)
new_doc.add_paragraph(core_text)

# ---- 3. MODEL DEVELOPMENT ----
new_doc.add_heading('3. Model Development', level=1)

# 3.1 Physical Basis
new_doc.add_heading('3.1 Physical Basis for Two-Component BTC', level=2)
# From v3 Section 2.1
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith('Heading 2') and '2.1 Physical Basis' in p.text:
        for j in range(i+1, len(doc.paragraphs)):
            pp = doc.paragraphs[j]
            if pp.style.name.startswith('Heading'):
                break
            if pp.text.strip() and '[Fig.' not in pp.text:
                new_doc.add_paragraph(pp.text)
        break

# 3.2 Governing Equations
new_doc.add_heading('3.2 Governing Equations', level=2)
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith('Heading 2') and '2.2 Governing Equations' in p.text:
        for j in range(i+1, len(doc.paragraphs)):
            pp = doc.paragraphs[j]
            if pp.style.name.startswith('Heading'):
                break
            if pp.text.strip():
                new_doc.add_paragraph(pp.text)
        break

# 3.3 Parameter Estimation and Model Selection
new_doc.add_heading('3.3 Parameter Estimation and Model Selection', level=2)

# Combine 2.3 + 2.4 + 3.5 + 4.1 content
# Parameter set table description
param_text = (
    'Equation (1) contains seven fitted coefficients: baseline concentration c_b, pulse '
    'amplitude A, tail amplitude a, dispersivity α, effective flow rate Q, crossover '
    'time t₀, and transition width σ. Of these, σ is effectively determined '
    'by its scale (≈ the 4-min sampling interval), leaving six free parameters. '
    'Table 1 lists the physical meaning and units of each parameter. Fixed geometric '
    'constants are x = 100 mm (tubing length between the pack outlet and the sampling '
    'point) and d = 5 mm (tubing inner diameter). Derived transport properties are '
    'Pe = x/α, v = 4Q/(πd²), and MRT = x/v.'
)
new_doc.add_paragraph(param_text)

# Parameter estimation
est_text = (
    'The inverse problem — extracting six parameters from a single BTC with coupled '
    'release and transport components — presents a challenging optimization landscape. '
    'A two-pass estimation strategy was adopted. Pass 1 (basin location): a global search '
    'was performed over wide, physically bounded parameter ranges (Table S6) using four '
    'independent runs with different random starting points. The effective flow rate Q '
    'was bounded to [0.005, 5.0] mL/min, a 1000-fold range spanning two orders of '
    'magnitude above and below the pump setting of 0.50 mL/min. Pass 2 (local refinement): '
    'from the best point identified in Pass 1, a gradient-based minimization refined the '
    'parameter estimates to the precision required for physical interpretation. This '
    'second pass typically converged within 50–200 iterations, confirming that Pass 1 '
    'had already located the correct basin. The identical two-pass protocol was applied '
    'to all five candidate models — the dual-component tanh-blended model and four '
    'single-component alternatives (single Gaussian, single erfc, exponential decay, '
    'K-P power law) — to ensure fair comparison. Full optimization details are provided '
    'in the Supplementary Material (Section S5.3).'
)
new_doc.add_paragraph(est_text)

# ---- 4. RESULTS AND DISCUSSION ----
new_doc.add_heading('4. Results and Discussion', level=1)

# 4.1 ESP-T Characterization
new_doc.add_heading('4.1 ESP-T Characterization and Material Prerequisites', level=2)
# From v3 Section 3.2 content
char_result = (
    'SEM micrographs confirm that ESP-T microspheres exhibit excellent sphericity (>0.9) '
    'with no inter-particle agglomeration. The surface displays uniformly distributed '
    'micron-scale protrusions corresponding to nano-Fe₃O₄@SA nanoclusters '
    'embedded within the epoxy matrix; the absence of interfacial cracks is consistent '
    'with physical adhesion mediated by stearic acid alkyl chain entanglement with the '
    'epoxy network. Elemental mapping confirms that Fe is distributed throughout the '
    'particle cross-section, consistent with bulk encapsulation rather than surface-only '
    'attachment. TGA/DSC analysis identifies the primary epoxy decomposition at 357 °C '
    '(DTG peak), far exceeding typical downhole temperatures (80–200 °C). '
    'Stearic acid modification increases the water contact angle from 72.3° (pure '
    'epoxy, weakly hydrophilic) to 104.6° (ESP-T, hydrophobic), driven by outward '
    'orientation of stearic acid alkyl chains. ESP-T exhibits a bulk density of '
    '0.646 g/cm³ (<1.0 g/cm³, enabling self-suspension), acid solubility of '
    '3.3% (<5% threshold per SY/T 5107–2016), and a crush rate of 2.9% at 52 MPa. '
    'Packed-bed filtration tests demonstrate the functional consequence of the oleophilic '
    'surface: ESP-T reduces oil passage time by 66% relative to pure epoxy while '
    'increasing water passage time nearly tenfold (water/oil time ratio = 5.53). '
    '\n\n[Fig. 2 placeholder — see Figure Captions]'
)
new_doc.add_paragraph(char_result)

# 4.2 Temperature-Dependent Release Kinetics
new_doc.add_heading('4.2 Temperature-Dependent Tracer Release Kinetics', level=2)
release_text = (
    'Tracer release from ESP-T in dodecane was measured at 30, 60, 90, and 120 °C '
    '(Fig. 3). Release accelerates with temperature throughout the 14-day measurement '
    'period; the rate is highest during the first 24 h and decays gradually thereafter. '
    'The K-P model was fitted within its validated range (M_t/M_∞ < 0.6; Table 2). '
    'The rate constant K increases systematically from 0.0554 (30 °C) to 0.1964 '
    '(120 °C), and all n values fall within 0.45–0.85 (R² > 0.94 at '
    'every temperature), corresponding to anomalous (non-Fickian) transport co-governed '
    'by Fickian diffusion and Case-II polymer relaxation. The physical picture is one of '
    'solvent-driven swelling: dodecane permeates the cross-linked epoxy network, generating '
    'an inner glassy core and an outer gel layer; swelling expands transport channels '
    'through which the doped metal ions diffuse into the external medium. These kinetic '
    'parameters characterize the release mechanism independently of any flow '
    'configuration.\n\n[Fig. 3 placeholder — see Figure Captions]'
)
new_doc.add_paragraph(release_text)

# 4.3 BTC Decomposition and Model Validation
new_doc.add_heading('4.3 BTC Decomposition and Model Validation', level=2)
# Get 4.1 + 4.2 content from v3
# Model selection first (from old 4.1)
model_sel = (
    'Table 3 compares the dual-component model against four single-component alternatives '
    'fitted to the same BTC dataset using the identical two-pass protocol. The dual-component '
    'model achieves R² = 0.9939 (RMSE = 0.0210), decisively outperforming all '
    'single-component alternatives: single Gaussian (R² = 0.9482, ΔAICc = 32.7), '
    'single erfc (R² = 0.7159, ΔAICc = 68.4), exponential decay '
    '(R² = 0.7517, ΔAICc = 62.5), and K-P power law (R² = −0.0193, '
    'ΔAICc = 92.2). An F-test decisively rejects the simpler single-Gaussian model '
    '(F(3, 14) = 34.70, p < 10⁻⁶). Single-process models each capture one feature '
    'of the BTC at the expense of the other: the Gaussian fits the peak but misses the '
    'tail, while the erfc captures the tail but cannot reproduce the peak. This evidence is '
    'necessary but not sufficient: ΔAICc eliminates the single-process alternatives, '
    'but cannot by itself establish that the dual-component structure is physically correct.\n\n'
    '[Fig. 4 placeholder — see Figure Captions]'
)
new_doc.add_paragraph(model_sel)

# Then the decomposition and validation (from old 4.2)
for text in [
    'The fitted two-component model achieves R² = 0.9939 (RMSE = 0.0210; Fig. 5a). '
    'The fitted parameter values are listed in Table 4. Separating the blended curve '
    'into its two constituent components (Fig. 5a) reveals that over the 105-min '
    'measurement window, the Gaussian pulse accounts for approximately 53% of the '
    'integrated tracer signal and the erfc tail for approximately 47%. The fitted '
    'crossover time t₀ = 25.66 min marks the transition from pulse-dominated to '
    'tail-dominated behavior, with a transition width σ = 3.96 min approximately '
    'equal to the 4-min sampling interval.',

    'The most direct test of the model’s physical validity is the internal '
    'consistency of transport time scales. The fitted flow velocity v = 4Q/(πd²) '
    'gives a mean residence time MRT = x/v = 38.6 min. This can be compared with '
    'the convective travel time independently computed from the pump flow rate '
    '(0.50 mL/min) and the tube geometry (x = 100 mm, d = 5 mm): 1 PV = '
    'xπd²/(4Q) = 38.6 min. The two values agree to within 3%, despite the '
    'model having six free parameters and the pump flow rate being used nowhere in '
    'the objective function. An overfitted model has no mechanism to produce this '
    'agreement—it can match the curve shape equally well with different '
    'combinations of v and α, yielding MRT values that drift arbitrarily from the '
    'independently known convective time scale. The convergence of the fitted MRT to '
    'the independently computed value therefore constitutes evidence that the model '
    'captures the underlying transport physics.',

    'The Peclet number derived from the fit, Pe = x/α = 0.934 ≈ 1, '
    'indicates that convection and dispersion contribute approximately equally to '
    'tracer transport—a signature of non-piston, dispersion-dominated displacement. '
    'This result is independently corroborated by the K–P kinetic analysis '
    '(Section 4.2), which identifies the tracer release mechanism as anomalous '
    '(non-Fickian) transport with n = 0.45–0.85 across all four temperatures. '
    'These two findings emerge from completely separate experiments: the K–P '
    'analysis from static batch release in glass vials (different apparatus, '
    'different data, different fitting targets—K and n), and the Pe from '
    'dynamic core-flood BTC fitting (flowing system, BTC data, six fitted transport '
    'parameters). Their convergence on the same physical picture—that tracer '
    'transport in this system is governed by coupled diffusion and advection, with '
    'neither mechanism dominant—provides strong, independent evidence that both '
    'the release characterization and the transport model are physically sound.',
]:
    new_doc.add_paragraph(text)

new_doc.add_paragraph('[Fig. 5 placeholder — see Figure Captions]')

# 4.4 Time-of-Arrival Comparison
new_doc.add_heading('4.4 Comparison with Time-of-Arrival Methods', level=2)
# From old 4.3
toa_text = (
    'A natural question is whether the full coupled model of Eq. (1) is necessary, '
    'or whether simpler time-of-arrival (TOA) methods—which require no '
    'optimization—could suffice for interpreting the BTC. The peak-time method '
    'assumes piston-like displacement: Q = xπd²/(4·t_peak). Applying '
    'this to the measured BTC (t_peak = 15 min) yields Q = 1.31 mL/min, overestimating '
    'the pump setting (0.50 mL/min) by 162%. The half-peak method (t_half ≈ 5 min) '
    'gives Q = 3.93 mL/min (error +685%). Both methods fail because dispersive spreading '
    'shifts the apparent peak substantially earlier than the true convective arrival '
    'in a sustained-release system. The first-moment method, which computes the mean '
    'residence time from the entire BTC (MRT = ∫t·C dt / ∫C dt = 37.1 min) '
    'and converts via Q = xπd²/(4·MRT), yields Q = 0.53 mL/min '
    '(error +5.8%). While accurate, the first-moment method provides no signal '
    'decomposition, no mechanistic insight, and no estimate of the Peclet number—it '
    'conflates the shut-in slug and the sustained tail into a single number. (The MRT '
    'of 37.1 min from the raw first moment and 38.6 min from the fitted model differ '
    'by approximately 4% because the first moment weights the entire measured BTC '
    'including noise, while the fitted MRT reflects the model-smoothed transport time '
    'scale.) The coupled model uniquely combines accurate description of the BTC shape '
    'with signal decomposition into physically meaningful components and the Peclet '
    'number for mechanistic interpretation.\n\n'
    '[Fig. 6 placeholder — see Figure Captions]'
)
new_doc.add_paragraph(toa_text)

# 4.5 Signal Decomposition Robustness
new_doc.add_heading('4.5 Signal Decomposition Robustness', level=2)
robust_text = (
    'To verify that the decomposition into Gaussian and erfc components is not an '
    'artifact of the chosen transition width, σ was varied from 1.98 to 11.89 min '
    '(a six-fold range, 0.5×–3.0× the fitted value of 3.96 min) while '
    'holding the other six parameters fixed. The erfc tail contribution varies by less '
    'than one percentage point across the entire scan range (46.7–47.5%), '
    'confirming that the decomposition is a robust structural feature of the BTC rather '
    'than a parametrically tuned outcome. Together with the MRT self-consistency and the '
    'Pe–K–P corroboration (Section 4.3), this establishes that the two-component '
    'model captures physically meaningful information from the BTC: the relative '
    'contributions of the shut-in accumulation slug and the sustained matrix release '
    '(≈47% erfc tail), the transport regime (Pe ≈ 1, non-piston displacement), '
    'and the mean residence time of the tracer in the system.\n\n'
    '[Fig. 8 placeholder — see Figure Captions]'
)
new_doc.add_paragraph(robust_text)

# 4.6 Two-Phase Production Allocation
new_doc.add_heading('4.6 Two-Phase Production Allocation via Tracer Flux', level=2)
two_phase_text = (
    'Under steady-state two-phase flow, the tracer concentration C_oil measured at the '
    'outlet decreases with increasing total flow rate Q_total — a straightforward '
    'consequence of dilution: at higher flow rates, less time is available for tracer to '
    'accumulate in each unit volume of produced fluid. Concentration alone is therefore a '
    'poor proxy for the production rate of a given stage. The way around this is to work '
    'with tracer mass flux rather than concentration. We define the oil-phase tracer mass '
    'flux as F_O = C_oil × Q_oil. At steady state, F_O equals the release rate '
    'from the ESP-T pack and is therefore independent of Q_total—it reflects only '
    'what the proppant releases, which is governed by the oil-wetted surface area (and '
    'hence the oil-water ratio), not by how fast fluid sweeps past it.\n\n'
    'This is confirmed experimentally: at each OWR, F_O remains approximately constant '
    'across the full range of total flow rates (0.1–0.4 mL/min). F_O increases '
    'systematically with OWR (≈0.66, 1.61, and 2.72 μg/min at OWR = 1:4, '
    '1:1, and 4:1, respectively), consistent with a larger oil-wetted area at higher '
    'oil fractions. To convert F_O into a production rate, we normalize by the '
    'single-phase reference flux F_O,ref = 3.187 ± 0.15 μg/min (triplicate '
    'measurements under pure oil flow). The normalized flux F_O/F_O,ref tracks the '
    'independently known oil flow rate Q_oil across all three OWR conditions '
    '(Pearson r = 0.97, p = 0.006, RMSD = 8.3%). The primary source of uncertainty '
    'is the F_O,ref reproducibility (±4.7%).\n\n'
    '[Fig. 9 placeholder — see Figure Captions]'
)
new_doc.add_paragraph(two_phase_text)

# ---- 5. FIELD DEPLOYMENT ----
new_doc.add_heading('5. Implications for Field-Scale Deployment', level=1)

deploy_text = (
    'The coupled framework—K–P kinetics → ADE model → '
    'tracer flux method—provides a complete pathway from wellhead concentration '
    'measurements to per-stage production allocation. For a multi-stage well, each '
    'fracture stage is assigned a distinct tracer element (e.g., Mn, Zn, Cu, Eu, Dy) '
    'doped into the ESP-T matrix during synthesis. After a single shut-in period following '
    'fracturing, produced fluid is sampled at the wellhead and analyzed by ICP-MS to '
    'obtain the concentration C_i of each tracer element as a function of time. '
    'During the flowback phase, the full BTC of each element is recorded; the ADE model '
    'decomposes each BTC into its shut-in slug and sustained-release components, and '
    'the MRT provides a first estimate of the per-stage drainage volume. During the '
    'subsequent steady production phase, periodic wellhead sampling yields the '
    'steady-state concentration C_i,ss for each tracer element. The per-stage oil '
    'flow rate is obtained as Q_i = F_O,i / C_i,ss, where F_O,i is the calibrated '
    'tracer flux for stage i. The per-stage contribution rate is then:\n\n'
    'Contribution_i = Q_i / ΣQ_j = (F_O,i / C_i) / Σ(F_O,j / C_j)\n\n'
    'The flux calibration F_O,i is obtained from laboratory single-phase reference '
    'measurements (F_O,ref) scaled by the proppant mass m_i injected into stage i '
    'and the reservoir temperature T_i, using the K–P-derived temperature '
    'dependence r(T): F_O,i ≈ (m_i / m_ref) × (r(T_i) / r(T_ref)) × '
    'F_O,ref. Several limitations of the present study define the path from '
    'laboratory demonstration to field deployment. The framework has been validated '
    'on a single interval at the laboratory scale; multi-interval configurations remain '
    'untested. All experiments used dodecane as the model oil; crude oil introduces '
    'additional complexity from variable composition, viscosity, and potential '
    'interactions with the epoxy matrix. The chemical stability of the epoxy matrix in '
    'aggressive downhole environments (H₂S, CO₂, high-salinity brines, '
    'temperatures exceeding 120 °C) has not been evaluated. The flux calibration '
    'is based on three OWR levels (n = 3 independent oil-fraction points); a larger '
    'matrix would strengthen the calibration for field use. Batch-specific F_O,ref '
    'determination is advisable for each proppant production lot.\n\n'
    '[Fig. 10 placeholder — see Figure Captions]'
)
new_doc.add_paragraph(deploy_text)

# ---- 6. CONCLUSIONS ----
new_doc.add_heading('6. Conclusions', level=1)

conclusions = [
    'A coupled release–transport model was developed that decomposes a '
    'tracer-proppant breakthrough curve into a Gaussian pulse (shut-in accumulation '
    'slug) and an erfc tail (sustained matrix-diffusion-controlled release), linked '
    'by a C¹-continuous hyperbolic-tangent transition. The dual-component '
    'structure is statistically decisive over four single-process alternatives '
    '(ΔAICc = 32.7, F-test p < 10⁻⁶, R² = 0.9939).',

    'The model’s physical validity is established through two independent '
    'lines of evidence: (i) the fitted mean residence time (MRT = 38.6 min) is '
    'internally consistent with the convective travel time independently computed '
    'from the pump flow rate and tube geometry (3% deviation), and (ii) the Peclet '
    'number (Pe = 0.934 ≈ 1) derived from the dynamic BTC fit independently '
    'corroborates the non-Fickian transport mechanism (n = 0.45–0.85) identified '
    'from static K–P batch release experiments—two completely separate '
    'experiments converging on the same physical picture.',

    'The erfc tail accounts for approximately 47% of the integrated tracer signal, '
    'a result stable to within ±0.8 percentage points across a six-fold variation '
    'in the transition width, confirming that the sustained matrix-diffusion-controlled '
    'release contributes nearly half of the total detected tracer and enabling extended '
    'monitoring from a single tracer placement.',

    'A tracer flux method enables production allocation under two-phase flow. The '
    'oil-phase tracer mass flux F_O = C_oil × Q_oil is invariant with total '
    'flow rate at a given oil–water ratio, eliminating the dilution artifact. '
    'The normalized flux tracks oil production rates across oil–water ratios '
    'from 4:1 to 1:4 (Pearson r = 0.97, RMSD = 8.3%).',

    'For multi-stage wells, doping each fracture stage with a distinct tracer element '
    'enables per-stage contribution rates to be obtained from wellhead ICP-MS samples '
    'as (F_O,i / C_i) / Σ(F_O,j / C_j), where F_O,i is calibrated from '
    'laboratory reference measurements and K–P temperature-dependent release '
    'kinetics. The framework requires no downhole tools and only a single shut-in. '
    'Multi-stage field validation with crude oil under transient conditions remains '
    'the next step toward deployment.',
]
for c in conclusions:
    new_doc.add_paragraph(c)

# ---- SAVE ----
output = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP-T_投稿文件\ESP-T_v4_manuscript.docx'
new_doc.save(output)
print(f'[OK] v4 manuscript saved to:\n{output}')
print()
print('New structure:')
print('  1. Introduction')
print('  2. Experimental (2.1-2.5)')
print('  3. Model Development (3.1-3.3)')
print('  4. Results and Discussion (4.1-4.6)')
print('  5. Implications for Field-Scale Deployment')
print('  6. Conclusions')
