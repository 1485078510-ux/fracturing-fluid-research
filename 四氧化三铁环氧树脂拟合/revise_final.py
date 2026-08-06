# -*- coding: utf-8 -*-
"""Edit ESP-T_Final_4-revised.docx: replace Introduction, update refs, fix tubing dimensions."""
from docx import Document
from docx.shared import Pt

DOCX = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP-T_投稿文件\ESP-T_Final_4-revised.docx'
doc = Document(DOCX)

# ═══ NEW INTRODUCTION (5 paragraphs, 33 refs) ═══
new_intro = [
    # P1
    'Multi-stage hydraulic fracturing has enabled economic production from '
    'unconventional reservoirs, but completion diagnostics remain hindered by '
    'a persistent question: how much does each fracture stage produce? '
    'Production logging across hundreds of horizontal wells revealed that '
    'approximately 30\u201340% of perforation clusters contribute negligibly '
    'to production [1], motivating two decades of work on completion '
    'optimization and re-stimulation strategies [2,3]. Yet the ability to '
    'routinely measure per-stage contribution at the surface has not kept '
    'pace. Production logging tools require well intervention via coiled '
    'tubing or tractor and capture only a snapshot of the inflow profile [4]. '
    'Distributed fiber-optic sensing\u2014both distributed temperature sensing '
    '(DTS) and distributed acoustic sensing (DAS)\u2014has enabled real-time '
    'monitoring of stimulation operations and qualitative inflow profiling '
    'without intervention [5,6], and DAS-based strain measurements now '
    'provide far-field fracture geometry and cluster efficiency diagnostics '
    '[7]. However, permanent fiber installation and surface instrumentation '
    'remain costly, limiting deployment to high-value wells [8]. Microseismic '
    'monitoring, while valuable for characterizing stimulated reservoir '
    'volume, does not directly measure production contribution per stage '
    '[3,9]. A method that delivers quantitative per-stage production '
    'allocation from surface samples alone\u2014without downhole tools, '
    'permanent installation, or repeated well intervention\u2014would '
    'fundamentally change completion diagnostics.',

    # P2
    'Tracer technology offers a pathway toward this method. Since the first '
    'documented oilfield tracer test using radioactive iodine in 1954 [10], '
    'tracer methods have evolved from qualitative inter-well connectivity '
    'assessment to quantitative residual oil saturation measurement via '
    'partitioning inter-well tracer tests [11,12] and, more recently, to '
    'stage-level inflow profiling in multi-stage fractured wells [13]. Tracer '
    'proppants\u2014in which a chemical tracer is immobilized within a solid '
    'proppant carrier and co-injected with the proppant pack\u2014offer a key '
    'advantage over dissolved tracers: the tracer remains in the fracture '
    'after placement, enabling long-term monitoring from surface samples '
    'without repeated injection. A growing body of material designs has been '
    'demonstrated: Zhao et al. (2020) used solvent-evaporated dye coatings '
    'on ceramic proppants [14]; Zhou et al. (2022) encapsulated carbon '
    'quantum dots in polymer-coated ceramics [15]; Li et al. (2023) developed '
    'rare-earth-doped polymer matrices for multi-element coding [16]; Gong '
    'et al. (2024) synthesized oleophilic Fe\u2083O\u2084/polystyrene '
    'microspheres for oil-phase selective release [17]; Ren et al. (2024) '
    'prepared multi-colored dye-tracer proppants for quantitative flowback '
    'assessment [18]; and Malyavko et al. (2023) demonstrated quantum-dot-encoded '
    'polymer microspheres for fluid-phase-resolved inflow profiling [19]. '
    'Across these diverse material platforms, a common interpretation gap '
    'persists: existing methods can confirm which stages produce, but cannot '
    'quantify how much each stage produces from the breakthrough curve alone, '
    'because the observed signal is shaped jointly by the unknown tracer '
    'release rate and the unknown transport parameters of the production system.',

    # P3
    'The release side of this coupled problem is addressed\u2014but not '
    'solved\u2014by the Korsmeyer\u2013Peppas (K\u2011P) power-law model. '
    'Originally developed for drug delivery systems [20,21] and extended to '
    'coupling of diffusion and polymer relaxation [22], the K\u2011P model '
    'C/C\u2080 = K\u00b7t\u207f identifies the release mechanism through '
    'the diffusional exponent n (Fickian diffusion for n \u2264 0.43; '
    'anomalous transport for 0.43 < n < 0.85; Case-II relaxation for '
    'n \u2265 0.85) and provides the temperature-dependent rate constant '
    'K. This framework has been productively applied to characterize release '
    'from diverse tracer-proppant systems including rare-earth-doped polymer '
    'coatings [16], polystyrene-encapsulated Fe\u2083O\u2084 nanoparticles '
    '[17], epoxy-matrix sustained-release tracers [23], and nano-Fe\u2083O\u2084 '
    'tracers in copolymer matrices under varying temperature and salinity [24]. '
    'However, K\u2011P is fundamentally a zero-dimensional batch description: '
    'it characterizes release into a well-mixed vessel with no spatial '
    'coordinate, no flow field, and no pathway from a release rate to a '
    'concentration measured at a distant sampling point. Alone, it cannot '
    'predict the wellhead concentration at a given time\u2014the quantity '
    'that encodes production information.',

    # P4
    'On the transport side, the one-dimensional advection\u2013dispersion '
    'equation (ADE) provides a mature analytical framework for interpreting '
    'breakthrough curves. Van Genuchten and Alves (1982) compiled analytical '
    'solutions covering a wide range of boundary and initial conditions [25]; '
    'the temporal moments of a BTC yield the mean residence time and '
    'dispersivity. Full-curve ADE inversion has been applied to determine '
    'reservoir properties and flood performance from inter-well tracer tests '
    '[26], to interpret physical mechanisms in partitioning tracer tests [27], '
    'to model chemical tracer transport under two-phase flow via analytical '
    'solutions [28], and to describe water- and oil-soluble tracer transfer '
    'in multi-stage fractured wells [29]. A comprehensive review of tracer '
    'testing techniques [30] confirms that ADE-based interpretation remains '
    'the standard approach across the petroleum industry. The common premise '
    'across these applications, however, is that the tracer source term is '
    'known\u2014an injection pulse of specified mass, duration, and location. '
    'A tracer-proppant BTC violates this premise fundamentally: the source '
    'is not a single operator-controlled injection but a sustained, '
    'matrix-diffusion-controlled release whose rate parameters are unknown '
    'and whose duration spans the entire production period. No existing '
    'framework couples the sustained release source term to the ADE transport '
    'solution. The consequence is that the two bodies of knowledge\u2014release '
    'kinetics and transport interpretation\u2014have remained separate, and '
    'neither alone can deliver the per-stage flow rate from the BTC.',

    # P5
    'In this work, we close this gap by coupling the release source term '
    'to the ADE transport solution through a two-component model and a '
    'complementary tracer flux method. We construct a piecewise '
    'release\u2013transport model that decomposes the BTC into two physically '
    'motivated components: a Gaussian pulse (shut-in accumulation slug) and '
    'an erfc tail (sustained matrix-diffusion-controlled release), linked by '
    'a C\u00b9-continuous hyperbolic-tangent transition. Six parameters are '
    'estimated simultaneously from a single BTC by nonlinear least squares. '
    'The model is validated through a predictive self-calibration test: the '
    'flow rate Q is left entirely unconstrained in the objective function; '
    'the model must recover the independently set pump flow rate from the BTC '
    'shape alone. Applied to an oleophilic epoxy/Fe\u2083O\u2084 tracer '
    'proppant (ESP\u2011T) in core displacement experiments, the model '
    'recovers Q = 0.46 mL/min against the pump setting of 0.50 mL/min\u2014a '
    'deviation of 8%. Epoxy resins have attracted growing interest as '
    'proppant matrices owing to their high mechanical strength, thermal '
    'stability, and chemical resistance compared to thermoplastic alternatives '
    '[31\u201333], and their ability to sustain controlled tracer release '
    'over extended periods has been demonstrated in both water-soluble [23] '
    'and oil-soluble configurations [34]. The Peclet number (Pe = 0.934) '
    'independently corroborates the non-Fickian transport mechanism identified '
    'from static K\u2011P batch kinetics (n = 0.45\u20130.85)\u2014two '
    'completely separate experiments converging on the same '
    'dispersion-dominated transport picture. We then introduce a tracer flux '
    'method for production allocation under two-phase flow: the oil-phase '
    'tracer mass flux F_O = C_oil \u00d7 Q_oil is invariant with total '
    'flow rate at a given oil\u2013water ratio (Pearson r = 0.97, RMSD = '
    '8.3%), eliminating the dilution artifact that confounds '
    'concentration-based interpretation. By doping each fracture stage with '
    'a distinct tracer element, per-stage contribution rates are obtained as '
    '(F_O,i / C_i) / \u03a3(F_O,j / C_j) from periodic ICP\u2011MS wellhead '
    'samples alone. We validate the complete framework in single-phase and '
    'two-phase core displacement experiments, and outline the pathway to '
    'multi-stage field deployment with multi-element coding for per-stage '
    'signal separation.',
]

# ═══ FIND AND REPLACE INTRODUCTION ═══
intro_start = None; intro_end = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if p.style.name.startswith('Heading') and t == '1. Introduction':
        intro_start = i
    elif intro_start is not None and p.style.name.startswith('Heading') and '2.' in t:
        intro_end = i
        break

if intro_start and intro_end:
    old = list(range(intro_start+1, intro_end))
    for idx in reversed(old):
        doc.paragraphs[idx]._element.getparent().remove(doc.paragraphs[idx]._element)
    ins = doc.paragraphs[intro_start]._element
    for t in reversed(new_intro):
        np = doc.add_paragraph(t); np.style = doc.styles['Normal']
        for r in np.runs: r.font.size=Pt(12); r.font.name='Times New Roman'
        ins.addnext(np._element)
    print(f'Intro: {len(old)} old -> {len(new_intro)} new paragraphs')
else:
    print(f'Intro positions: start={intro_start} end={intro_end}')

# ═══ UPDATE KEYWORDS ═══
for i, p in enumerate(doc.paragraphs):
    if p.text.startswith('Keywords:'):
        p.clear(); p.add_run('Keywords: Tracer proppant; Breakthrough curve; '
            'Advection\u2013dispersion equation; Release kinetics; Production '
            'allocation; Two-phase flow; Epoxy resin; Tracer flux')
        print('Keywords updated')
        break

# ═══ UPDATE EXPERIMENTAL SECTION: tubing dimensions ═══
for i, p in enumerate(doc.paragraphs):
    if 'core holder under 5 MPa' in p.text and 'single-phase' in p.text:
        old_text = p.text
        # Insert tubing specification after "5 MPa confining pressure"
        new_text = old_text.replace(
            'under 5 MPa confining pressure. For single-phase experiments,',
            'under 5 MPa confining pressure. The production tubing connecting '
            'the core outlet to the sampling point was 1 m in length with an '
            'inner diameter of 1 mm. For single-phase experiments,')
        p.clear(); p.add_run(new_text)
        print('Tubing dimensions added to Experimental')
        break

# ═══ FIX FORMATTING: font consistency ═══
for p in doc.paragraphs:
    for r in p.runs:
        if r.font.name is None or 'Times' not in str(r.font.name):
            r.font.name = 'Times New Roman'
        if r.font.size is None:
            r.font.size = Pt(12)

# ═══ SAVE ═══
out = DOCX.replace('.docx', '_revised.docx')
doc.save(out)
print(f'Saved: {out}')
print(f'Total paragraphs: {len(doc.paragraphs)}')
