# -*- coding: utf-8 -*-
"""Replace Introduction in v5 with optimized version featuring authoritative references."""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

DOCX = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP-T_投稿文件\ESP-T_v5_manuscript.docx'
doc = Document(DOCX)

new_paragraphs = [
    # ===== P1: FIELD STAKE (~190 words) =====
    (
        'Multi-stage hydraulic fracturing has enabled economic production from '
        'unconventional reservoirs, but a fundamental operational question limits '
        'the value of completion diagnostics: how much oil does each fracture stage '
        'produce? Field evidence consistently shows that 30–50% of perforation '
        'clusters contribute negligibly to production [1,2], yet operators lack a '
        'routine method to identify which stages underperform. Without per-stage '
        'contribution data, well spacing, completion design, and restimulation '
        'decisions rely on inference rather than measurement. Existing downhole '
        'diagnostic methods each address part of this need but carry operational '
        'burdens that preclude routine deployment. Production logging requires '
        'well intervention via coiled tubing or tractor and captures only a '
        'snapshot of the inflow profile at the moment of logging [3]. Distributed '
        'fiber-optic sensing — both distributed temperature sensing (DTS) and '
        'distributed acoustic sensing (DAS) — provides continuous spatial resolution '
        'but demands permanent cable installation and surface instrumentation whose '
        'cost is difficult to justify outside high-value wells [4,5]. Microseismic '
        'monitoring characterizes fracture geometry and stimulated reservoir volume '
        'rather than production contribution per stage [6]. A method that delivers '
        'per-stage production allocation from surface samples alone — without '
        'downhole hardware, without well intervention, and requiring only a '
        'single shut-in — would fundamentally change completion diagnostics.'
    ),

    # ===== P2: BOTTLENECK (~175 words) =====
    (
        'Tracer proppants offer a pathway toward this goal. A chemical tracer — '
        'typically a metal ion, rare-earth element, or fluorescent dye — is '
        'immobilized within a solid proppant carrier and co-injected with the '
        'proppant pack during hydraulic fracturing. The tracer remains in the '
        'fracture after placement and releases gradually into the produced fluid '
        'over months, generating a time-resolved concentration signal — the '
        'breakthrough curve (BTC) — measurable at the wellhead by inductively '
        'coupled plasma mass spectrometry (ICP‑MS). A range of tracer-proppant '
        'designs have been demonstrated: ceramic carriers with organic dye coatings '
        '[7], carbon quantum-dot-encapsulated ceramic proppants [8], rare-earth-doped '
        'polymer matrices for multi-element coding [9], oleophilic Fe₃O₄/polymer '
        'microspheres for oil-phase selectivity [10], and multi-colored dye-tracer '
        'proppants for proppant flowback assessment [11]. Interpreting a '
        'tracer-proppant BTC to recover the stage flow rate, however, requires '
        'solving a coupled inverse problem: the observed concentration at the '
        'wellhead is shaped jointly by the unknown tracer release rate from the '
        'polymer matrix and the unknown transport parameters of the production '
        'system. These two sets of unknowns — release kinetics and transport '
        'dynamics — are superimposed in a single observable, and neither is '
        'known a priori.'
    ),

    # ===== P3: PRIOR WORK CEILING (~260 words) =====
    (
        'Current practice addresses release and transport separately, leaving '
        'the coupling unresolved. On the release side, tracer-proppant performance '
        'is characterized through batch release measurements interpreted with the '
        'Korsmeyer–Peppas (K‑P) power law, C/C₀ = K·tⁿ [12,13]. The K‑P '
        'model identifies the release mechanism — Fickian diffusion for '
        'n ≤ 0.43, anomalous transport for 0.43 < n < 0.85, and Case-II '
        'relaxation for n ≥ 0.85 — and provides the temperature-dependent rate '
        'constant K. This framework has been successfully applied to characterize '
        'release from diverse polymer–tracer systems including poly(methyl '
        'methacrylate)-coated rare-earth tracers [9], polystyrene-encapsulated '
        'Fe₃O₄ nanoparticles [10], and epoxy-matrix water-soluble tracers '
        '[14]. However, the K‑P model is a zero-dimensional batch description: '
        'it characterizes release into a well-mixed vessel with no spatial '
        'coordinate, no flow field, and no pathway from a release rate to a '
        'concentration measured at a distant sampling point.'
    ),

    # ===== P4: TRANSPORT-SIDE GAP (~220 words) =====
    (
        'On the transport side, the one-dimensional advection–dispersion equation '
        '(ADE), ∂C/∂t + v·∂C/∂x = D·∂²C/∂x², provides analytical solutions '
        'that relate the shape of a BTC to the transport parameters — flow velocity '
        'v and dispersion coefficient D — of the system through which the tracer '
        'travels [15]. The temporal moments of a BTC yield the mean residence time '
        '(hence the effective flow velocity) and the variance (hence the '
        'dispersivity); full-curve ADE inversion has been applied to extract '
        'reservoir and fracture properties from inter-well tracer tests [16], to '
        'model tracer transport in multi-stage fractured wells with phase '
        'partitioning via analytical solutions [17], and to interpret partitioning '
        'inter-well tracer tests for residual oil saturation [18]. A common premise '
        'across these applications, however, is that the tracer source term is '
        'known — an injection pulse of specified mass, duration, and location. '
        'A tracer-proppant BTC violates this premise fundamentally: the source is '
        'not a single operator-controlled injection but a sustained, '
        'matrix-diffusion-controlled release whose rate parameters are unknown '
        'and whose duration spans the entire production period. No existing '
        'framework couples the sustained, unknown release source term to the ADE '
        'transport solution. The practical consequence is that current tracer '
        'proppants can confirm which stages produce but cannot quantify how much '
        'each stage produces from the BTC alone.'
    ),

    # ===== P5: THIS WORK (~260 words) =====
    (
        'In this work, we develop a method that closes this gap. We construct a '
        'coupled release–transport model that decomposes the BTC into two '
        'physically motivated components: a Gaussian pulse, representing tracer '
        'that accumulated in the near-wellbore region during shut-in and is swept '
        'out as a coherent slug upon flowback, and an erfc tail, representing '
        'sustained matrix-diffusion-controlled release that continues after the '
        'main slug has passed. The two components are linked by a C¹-continuous '
        'hyperbolic-tangent transition, and six parameters are estimated '
        'simultaneously from a single BTC by nonlinear least squares. The model '
        'is validated through a predictive self-calibration test: the flow rate Q '
        'is left entirely unconstrained in the objective function (search bounds '
        '0.01–5.0 mL/min, a 500-fold range); the model must recover the '
        'independently set pump flow rate from the BTC shape alone. Applied to '
        'an oleophilic epoxy/Fe₃O₄ tracer proppant (ESP‑T) in core '
        'displacement experiments with production tubing of 1 m length and 1 mm '
        'inner diameter, the model recovers Q = 0.52 mL/min against the pump '
        'setting of 0.50 mL/min — a deviation of 3.9%. The fitted mean residence '
        'time (1.51 min) agrees with the independently computed convective travel '
        'time (1.57 min) to within 3.8%. The Peclet number (Pe = 0.75) '
        'independently corroborates the non-Fickian transport mechanism identified '
        'from static K‑P batch kinetics (n = 0.45–0.85) — two completely '
        'separate experiments converging on the same dispersion-dominated transport '
        'picture. We then introduce a tracer flux method for steady-state production '
        'allocation: the oil-phase tracer mass flux F_O = C_oil × Q_oil is '
        'shown to be invariant with total flow rate at a given oil–water ratio '
        '(Pearson r = 0.97, RMSD = 8.3%), eliminating the dilution artifact '
        'that confounds concentration-based interpretation. By doping each fracture '
        'stage with a distinct tracer element, per-stage contribution rates are '
        'obtained as (F_O,i / C_i) / Σ(F_O,j / C_j) from periodic wellhead '
        'ICP‑MS samples alone. We validate the complete framework in both '
        'single-phase and two-phase core displacement experiments, and outline '
        'the pathway to field deployment with multi-element coding for per-stage '
        'signal separation.'
    ),
]

# Find and replace Introduction section
intro_start = None
intro_end = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith('Heading') and '1. Introduction' in p.text:
        intro_start = i
    elif intro_start is not None and p.style.name.startswith('Heading') and intro_end is None:
        intro_end = i
        break

if intro_start and intro_end:
    # Remove old intro paragraphs
    old_indices = list(range(intro_start + 1, intro_end))
    for idx in sorted(old_indices, reverse=True):
        doc.paragraphs[idx]._element.getparent().remove(doc.paragraphs[idx]._element)

    # Insert new paragraphs after Introduction heading
    insert_after = doc.paragraphs[intro_start]._element
    for text in reversed(new_paragraphs):
        new_p = doc.add_paragraph(text)
        new_p.style = doc.styles['Normal']
        for run in new_p.runs:
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
        insert_after.addnext(new_p._element)

    print(f'[OK] Introduction replaced: {len(old_indices)} old -> 5 new paragraphs')

# Save with temp name
out = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP-T_投稿文件\ESP-T_v5_intro_optimized.docx'
doc.save(out)
print(f'[OK] Saved to {out}')

# Print word counts
for i, text in enumerate(new_paragraphs):
    wc = len(text.split())
    print(f'  P{i+1}: {wc} words')
print(f'  Total: {sum(len(t.split()) for t in new_paragraphs)} words')
