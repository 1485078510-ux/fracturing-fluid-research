# -*- coding: utf-8 -*-
"""Replace Introduction in v4 manuscript with nature-writing draft."""
from docx import Document
from docx.shared import Pt

DOCX = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP-T_投稿文件\ESP-T_v4_manuscript.docx'
doc = Document(DOCX)

# New Introduction paragraphs
new_intro = [
    # P1: Field stake
    'Multi-stage hydraulic fracturing has made horizontal wells the dominant production '
    'technology for unconventional reservoirs, but a fundamental operational question '
    'remains unresolved: how much oil does each fracture stage produce? Without per-stage '
    'contribution data, underperforming intervals go undiagnosed, well spacing and '
    'completion designs cannot be validated against production outcomes, and restimulation '
    'decisions rely on guesswork rather than evidence [1,2]. Existing downhole diagnostic '
    'tools each involve trade-offs that limit their routine use. Production logging '
    'requires well intervention and captures only a snapshot of the inflow profile at '
    'the moment of logging [3]. Distributed fiber-optic sensing provides continuous '
    'spatial resolution but demands permanent cable installation at a cost difficult to '
    'justify for marginal wells [4,5]. Microseismic monitoring characterizes fracture '
    'geometry, not production contribution [6]. A method that delivers per-stage '
    'production allocation from surface samples alone—without downhole hardware, '
    'without well intervention, and without repeated shut-ins—would transform '
    'completion diagnostics.',

    # P2: Bottleneck
    'Tracer proppants offer a pathway toward this goal. A tracer agent is immobilized '
    'within a solid proppant carrier and co-injected with the proppant pack during '
    'fracturing. The tracer remains in the fracture after placement and releases '
    'gradually into the produced fluid, generating a time-resolved concentration '
    'signal—the breakthrough curve (BTC)—that can be measured at the wellhead by '
    'ICP-MS for months after a single placement [7–10]. Interpreting this BTC to '
    'recover the stage flow rate, however, requires solving a coupled inverse problem: '
    'the observed concentration at the wellhead is shaped jointly by the unknown tracer '
    'release rate from the polymer matrix and the unknown transport parameters of the '
    'production system. These two sets of unknowns—release kinetics and transport '
    'dynamics—are superimposed in a single observable, and neither is known a priori.',

    # P3: Prior work ceiling
    'Current practice addresses release and transport separately, leaving the coupling '
    'unresolved. On the release side, tracer proppant performance is characterized '
    'exclusively through batch release measurements interpreted with the Korsmeyer–Peppas '
    '(K‑P) power law, C/C₀ = K·tⁿ [11,12]. The K‑P model identifies '
    'the release mechanism—Fickian diffusion for n ≤ 0.43, anomalous transport '
    'for 0.43 < n < 0.85, Case-II relaxation for n ≥ 0.85—and provides '
    'the temperature-dependent rate constant. But K‑P is a zero-dimensional batch '
    'model: it describes release into a well-mixed vessel with no spatial coordinate, '
    'no flow field, and no pathway from a release rate to a concentration measured at '
    'a distant sampling point. On the transport side, the one-dimensional '
    'advection–dispersion equation (ADE) has a mature analytical solution framework '
    'that relates BTC shape to transport parameters [13,14]. ADE-based interpretation '
    'has been applied to extract reservoir and fracture properties from tracer BTCs '
    '[15,16] and to model multi-stage tracer transport with phase partitioning [17]. '
    'The common premise across these applications, however, is that the tracer source '
    'term is known—an injection pulse of specified mass, duration, and location. A '
    'tracer-proppant BTC violates this premise: the source is not a single '
    'operator-controlled injection but a sustained, matrix-diffusion-controlled release '
    'whose rate parameters are unknown and whose duration spans the entire production '
    'period. No existing framework couples the sustained release source term to the '
    'transport solution. The practical consequence is that current tracer proppants can '
    'confirm which stages produce but cannot quantify how much each stage produces '
    'from the BTC alone.',

    # P4: This paper
    'In this work, we develop a method that closes this gap in two steps. First, we '
    'construct a coupled release–transport model that decomposes the BTC into a Gaussian '
    'pulse (representing tracer that accumulated in the near-wellbore region during '
    'shut-in and is swept out as a coherent slug upon flowback) and an erfc tail '
    '(representing sustained matrix-diffusion-controlled release that continues after '
    'the main slug has passed), linked by a smooth hyperbolic-tangent transition. The '
    'dual-component structure is validated through internal self-consistency of the '
    'fitted mean residence time with the independently known convective travel time, '
    'and through independent corroboration between the Peclet number from the dynamic '
    'BTC fit and the transport mechanism from static K‑P batch kinetics—two completely '
    'separate experiments converging on the same physical picture. Second, we introduce '
    'a tracer flux method for production allocation: under steady-state flow, the '
    'oil-phase tracer mass flux F_O = C_oil × Q_oil is shown experimentally to be '
    'invariant with total flow rate at a given oil–water ratio, eliminating the dilution '
    'artifact that confounds concentration-based interpretation. By combining the K‑P '
    'temperature calibration, the ADE-based BTC decomposition, and the flux method, '
    'the per-stage contribution rate is obtained as '
    'Contribution_i = (F_O,i / C_i) / Σ(F_O,j / C_j) from surface samples '
    'alone. We validate the framework using an oleophilic epoxy/Fe₃O₄ tracer '
    'proppant (ESP-T) in single-phase and two-phase core displacement experiments, '
    'and outline the pathway to field deployment with multi-element coding for per-stage '
    'signal separation.',
]

# Find Section 1 heading and replace content
intro_start = None
intro_end = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith('Heading') and '1. Introduction' in p.text:
        intro_start = i
    elif intro_start is not None and p.style.name.startswith('Heading') and intro_end is None:
        intro_end = i
        break

if intro_start and intro_end:
    # Remove old intro paragraphs
    old_indices = list(range(intro_start + 1, intro_end))
    for idx in sorted(old_indices, reverse=True):
        doc.paragraphs[idx]._element.getparent().remove(doc.paragraphs[idx]._element)

    # Insert new paragraphs after the Introduction heading
    from docx.oxml.ns import qn
    insert_after = doc.paragraphs[intro_start]._element
    for text in reversed(new_intro):
        new_p = doc.add_paragraph(text)
        new_p.style = doc.styles['Normal']
        for run in new_p.runs:
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
        insert_after.addnext(new_p._element)

    print(f'[OK] Introduction replaced: {len(old_indices)} old paragraphs -> 4 new paragraphs')

# Save to new file (original may be locked by Word)
import os
base = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP-T_投稿文件'
DOCX_OUT = os.path.join(base, 'ESP-T_v4_manuscript.docx')
try:
    doc.save(DOCX_OUT)
except PermissionError:
    DOCX_OUT = os.path.join(base, 'ESP-T_v4_intro_revised.docx')
    doc.save(DOCX_OUT)
print(f'[OK] Saved to {DOCX_OUT}')
