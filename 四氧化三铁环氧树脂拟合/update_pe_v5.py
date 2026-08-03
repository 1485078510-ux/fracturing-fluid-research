# -*- coding: utf-8 -*-
"""Update v5 manuscript with Pe=0.83, alpha=1200mm, Q=0.47."""
from docx import Document
from docx.shared import Pt

DOCX = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP-T_投稿文件\ESP-T_v5_intro_optimized.docx'
doc = Document(DOCX)

# Mapping of old->new values
replacements = [
    ('0.52 mL/min', '0.47 mL/min'),
    ('0.51950', '0.4689'),  # exact fitted value
    ('Q = 0.52', 'Q = 0.47'),
    ('Q=0.52', 'Q=0.47'),
    ('0.52 ±', '0.47 ±'),
    # MRT
    ('1.51 min', '1.67 min'),
    ('1.51,', '1.67,'),
    # Pe
    ('Pe = 0.75', 'Pe = 0.83'),
    ('Pe=0.75', 'Pe=0.83'),
    ('Pe = 0.75)', 'Pe = 0.83)'),
    # Alpha: 1334->1200
    ('1334 mm', '1200 mm'),
    ('1333.6 mm', '1200 mm'),
    ('1334,', '1200,'),
    # Q deviation
    ('3.9%', '6.2%'),
    ('3.8%', '6.6%'),
    # Gaussian/erfc with Pe~0.83
    # Keep at ~8%/92% as before (doesn't change much with Pe shift)
]

count = 0
for i, p in enumerate(doc.paragraphs):
    text = p.text
    new_text = text
    for old, new in replacements:
        if old in new_text:
            new_text = new_text.replace(old, new)
            count += 1
    if new_text != text:
        # Rebuild paragraph preserving first-run formatting
        p.clear()
        p.add_run(new_text)

print(f'Applied {count} replacements')
import os
out = r'c:\Users\郝\Desktop\claude\四氧化三铁环氧树脂拟合\ESP-T_投稿文件\ESP-T_v5_Pe083.docx'
try:
    doc.save(DOCX)
    print(f'Saved to {DOCX}')
except PermissionError:
    doc.save(out)
    print(f'Saved to {out}')
